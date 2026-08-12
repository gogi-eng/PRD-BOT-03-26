# -*- coding: utf-8 -*-
"""Восстановление тела ДИ «ПРОЕКТ Старший мастер» из текстового дампа + полный проход СНиОТ."""
from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import fix_sniot_document as fix
from docx import Document

DUMP = Path(__file__).resolve().parent / "_work_senior_master_fixed.txt"


def _pick_master_files() -> list[Path]:
    satp = fix.find_satp_dir()
    out: list[Path] = []
    for path in satp.glob("*.docx"):
        low = path.name.lower()
        if "мастер" not in low or "_backup_" in low:
            continue
        out.append(path)
    return out


def main() -> int:
    if not DUMP.is_file():
        print("Дамп не найден:", DUMP)
        return 3

    targets = _pick_master_files()
    if not targets:
        print("Файлы «Старший мастер» не найдены в САТП")
        return 3

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    for target in targets:
        before = count = len(Document(target).paragraphs)
        if count >= 20:
            print(f"SKIP (уже есть тело): {target.name} paras={count}")
            continue
        bak = target.with_name(f"{target.stem}_backup_{stamp}{target.suffix}")
        shutil.copy2(target, bak)
        added = fix.replace_body_from_debug_dump(target, DUMP)
        print(f"RESTORED {target.name}: +{added} абзацев (backup: {bak.name})")

    # полный проход СНиОТ на главном результате
    main_out = fix.resolve_target(
        fix.find_satp_dir() / "ПРОЕКТ Старший мастер_оформлен.docx"
    )
    if not main_out.is_file():
        main_out = next(
            (p for p in targets if "оформлен" in p.name.lower()),
            targets[0],
        )
    rep = fix.apply_sniot_rules_to_file(main_out, always_apply=True)
    doc = Document(main_out)
    issues = fix.validate_sniot_document(doc, path=main_out, profile=fix.detect_profile(doc, main_out))
    print("\n=== VALIDATION ===")
    print("file:", main_out.name)
    print("paragraphs:", len(doc.paragraphs))
    print("issues:", len(issues))
    for item in issues:
        print(" -", item)
    print("apply:", rep.get("ok"), rep.get("actions", [])[-3:])
    return 0 if not issues else 1


if __name__ == "__main__":
    raise SystemExit(main())
