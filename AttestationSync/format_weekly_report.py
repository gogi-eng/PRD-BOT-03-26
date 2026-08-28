# -*- coding: utf-8 -*-
"""
Оформление «Еженедельный итог» по Инструкции по делопроизводству 2025
(Минюст РБ № 65): А4, поля 30/8/20/20 мм, TNR 14, по ширине, отступ 1,25 см,
одинарный интервал, номера страниц со 2-й сверху по центру.

Не ДИ САТП: без нумерации 1.8/1.9 и без правки папки Агент / ОБМЕН / САТП.

Пробелы: не более одного подряд (как MULTI_SPACE_RE в fix_sniot_document.py).
Межстрочный интервал тела: одинарный (1.0), перед/после абзаца 0 pt.
Маркеры списка запрещены: не оставлять • /  / ○ / ■ / дефис-списки / bullet numPr.
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

HERE = Path(__file__).resolve().parent
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))

import fix_sniot_document as sniot_fix

WEEKLY_PROJECT_DIR = Path(r"C:\Users\v.dubovik\Desktop\Еженедельный_итог")
DESKTOP = Path.home() / "Desktop"
DAILY_REPORTS_DIR = DESKTOP / "Ежедневные отчёты"
N_REPORTS = Path(
    r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В\Отчеты"
)
USER_AGENT_DIR = Path(
    r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В\Агент"
)
DAILY_REPORT_NAME_RE = re.compile(
    r"^ежедневный_отч[её]т_\d{4}-\d{2}-\d{2}",
    re.IGNORECASE,
)
PROTECTED_DAILY_MESSAGE = (
    "Это дневной отчёт на правку или уже принятый файл. "
    "После правок пользователя его нельзя повторно оформлять, "
    "проверять спеллером или менять текст. "
    "Недельный итог забирает такие файлы как есть (--from-daily)."
)

WEEKLY_CHAPTER_RE = re.compile(r"^(\d+)\.\s+(\S.*)$")
WEEKLY_SUBSECTION_RE = re.compile(r"^\d+\.\d+\.\s+\S")
DATE_ONLY_RE = re.compile(r"^\d{2}\.\d{2}\.\d{4}$")
PERIOD_LINE_RE = re.compile(
    r"^с\s+\d{2}\.\d{2}\.\d{4}\s+по\s+\d{2}\.\d{2}\.\d{4}$",
    re.IGNORECASE,
)


def looks_like_weekly_itog(path: str | Path, text: str = "") -> bool:
    """Имя/путь/шапка отчёта о работе за неделю — не ДИ и не положение."""
    p = Path(path)
    name = p.name.lower().replace("ё", "е")
    path_l = str(p).lower().replace("ё", "е")
    if re.match(r"^ди[\s_]", name) or "должностн" in name:
        return False
    if "еженедельный_итог" in path_l or "еженедельный итог" in path_l:
        return True
    if name.startswith("отчет_о_работе") or "еженедельный итог" in name:
        return True
    if name.startswith("ежедневный_отчет") or "ежедневные отчет" in path_l:
        return True
    title = (text or "")[:900].lower().replace("ё", "е")
    if "о выполненной работе за период" in title and "сниот" in title:
        return True
    if "о выполненной работе за сутки" in title and "сниот" in title:
        return True
    return False


def is_daily_report_filename(path: Path | str) -> bool:
    name = Path(path).name.lower().replace("ё", "е")
    return bool(DAILY_REPORT_NAME_RE.match(name))


def is_accepted_daily_path(path: Path | str) -> bool:
    """Папка «принятые» — пользователь уже поправил, оформлять нельзя."""
    p = Path(path)
    parts = [x.lower().replace("ё", "е") for x in p.parts]
    return "принятые" in parts and is_daily_report_filename(p)


def is_protected_daily_report(path: Path | str) -> bool:
    """Любой дневной .docx в «Ежедневные отчёты» (на_правку или принятые)."""
    p = Path(path)
    if not is_daily_report_filename(p):
        return False
    parts = [x.lower().replace("ё", "е") for x in p.parts]
    return "ежедневные отчеты" in parts or is_accepted_daily_path(p)


def _norm_path(path: Path | str) -> Path:
    return Path(str(path).strip().strip('"')).expanduser()


def _is_temp_test_path(path: Path) -> bool:
    low = str(path).lower()
    return (
        "\\temp\\" in low
        or "/tmp/" in low
        or "pytest" in low
        or "\\appdata\\local\\temp" in low
    )


def assert_weekly_writable(path: Path | str) -> Path:
    """Писать только Desktop/итог, Отчеты, или временные файлы тестов. Не Агент/ОБМЕН/САТП."""
    p = _norm_path(path)
    try:
        resolved = p.resolve()
    except OSError:
        resolved = p
    low = str(resolved).lower().replace("ё", "е")
    if "обмен" in low or "\\сатп\\" in low or "/сатп/" in low:
        raise PermissionError(
            "Еженедельный итог нельзя сохранять в ОБМЕН или САТП."
        )
    try:
        if resolved == USER_AGENT_DIR.resolve() or USER_AGENT_DIR.resolve() in resolved.parents:
            raise PermissionError(
                "Еженедельный итог не оформляется в папке Агент на N:\\. "
                "Это отдельный документ (рабочий стол / Отчеты)."
            )
    except OSError:
        if "\\агент\\" in low or low.endswith("\\агент"):
            raise PermissionError(
                "Еженедельный итог не оформляется в папке Агент на N:\\."
            )
    if _is_temp_test_path(resolved):
        return resolved
    allowed_parents = [WEEKLY_PROJECT_DIR, DESKTOP]
    if N_REPORTS.exists():
        allowed_parents.append(N_REPORTS)
    for root in allowed_parents:
        try:
            root_r = root.resolve()
        except OSError:
            root_r = root
        if resolved.parent == root_r or root_r in resolved.parents:
            if root_r == DESKTOP.resolve() and resolved.parent == root_r:
                n = resolved.name.lower().replace("ё", "е")
                if not (
                    n.startswith("отчет_о_работе")
                    or n.startswith("ежедневный_отчет")
                    or "еженедельный" in n
                ):
                    raise PermissionError(
                        "На рабочем столе оформляю только файлы отчёта о работе / еженедельного итога."
                    )
            return resolved
    raise PermissionError(
        "Оформление еженедельного итога разрешено только в папке "
        f"{WEEKLY_PROJECT_DIR}, на рабочем столе (Отчёт_о_работе…) "
        f"или в {N_REPORTS}."
    )


def is_weekly_chapter_header(text: str) -> bool:
    """«1. Количественные итоги», не «1.1. …» и не «2.14.07.2026»."""
    t = (text or "").replace("\xa0", " ").strip()
    if not t:
        return False
    match = WEEKLY_CHAPTER_RE.match(t)
    if not match:
        match = re.match(r"^(\d+)\s+([А-ЯЁA-Z].+)$", t)
        if not match:
            return False
        body = match.group(2).strip()
    else:
        body = match.group(2).strip()
    if body[:1].isdigit():
        return False
    if WEEKLY_SUBSECTION_RE.match(t):
        return False
    return 1 <= len(body) <= 90


def canonical_weekly_chapter(text: str) -> str:
    t = re.sub(r"\s+", " ", (text or "").replace("\xa0", " ").strip())
    match = WEEKLY_CHAPTER_RE.match(t)
    if match:
        return f"{match.group(1)}. {match.group(2).upper()}"
    match = re.match(r"^(\d+)\s+(.+)$", t)
    if match and not match.group(2)[:1].isdigit():
        return f"{match.group(1)}. {match.group(2).upper()}"
    return t.upper()


def is_weekly_subsection(text: str) -> bool:
    return bool(WEEKLY_SUBSECTION_RE.match((text or "").strip()))


def is_weekly_letterhead(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    u = t.upper().replace("Ё", "Е")
    if "МИНСККОММУНТЕПЛОСЕТЬ" in u:
        return True
    if u.startswith("СЛУЖБА НАДЕЖНОСТИ"):
        return True
    if u == "ОТЧЕТ" or u.startswith("ОТЧЕТ"):
        return True
    if u.startswith("О ВЫПОЛНЕННОЙ РАБОТЕ"):
        return True
    return bool(PERIOD_LINE_RE.match(t))


def is_weekly_signatory_line(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    if t.startswith("Ведущий инженер"):
        return True
    if "Дубовик" in t and "_" in t:
        return True
    return bool(DATE_ONLY_RE.fullmatch(t))


def normalize_spaces_text(text: str, *, full: bool = True) -> str:
    """Двойные пробелы → один; nbsp/таб → пробел. Тот же sanitize, что у СНиОТ."""
    return sniot_fix.sanitize_paragraph_text(text, full=full)


def _iter_doc_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def strip_weekly_double_spaces(doc: Document) -> int:
    """Не более одного пробела в каждом абзаце (включая таблицы)."""
    changed = 0
    for paragraph in _iter_doc_paragraphs(doc):
        original = paragraph.text
        if not original:
            continue
        is_special = (
            is_weekly_chapter_header(original)
            or is_weekly_letterhead(original)
            or is_weekly_signatory_line(original)
        )
        new_text = normalize_spaces_text(original, full=not is_special)
        if new_text == original:
            continue
        bold = None
        if is_weekly_chapter_header(original) or original.strip().upper().replace("Ё", "Е") == "ОТЧЕТ":
            bold = True
        elif is_weekly_subsection(original):
            bold = True
        else:
            bold = False
        sniot_fix.set_paragraph_text(paragraph, new_text, bold=bold)
        changed += 1
    return changed


def validate_weekly_spaces(doc: Document) -> list[str]:
    """Validate ловит два пробела подряд, nbsp и табы в теле."""
    issues: list[str] = []
    for paragraph in _iter_doc_paragraphs(doc):
        original = paragraph.text or ""
        if not original:
            continue
        collapsed = original.replace("\xa0", " ")
        if sniot_fix.MULTI_SPACE_RE.search(collapsed):
            snippet = original[:50]
            issues.append(f"Двойные пробелы в тексте: {snippet}")
        if "\xa0" in original:
            issues.append(f"Неразрывный пробел: {original[:50]}")
        if "\t" in original and not is_weekly_signatory_line(original):
            issues.append(f"Табуляция в тексте: {original[:50]}")
        expected = normalize_spaces_text(
            original,
            full=not (
                is_weekly_chapter_header(original)
                or is_weekly_letterhead(original)
                or is_weekly_signatory_line(original)
            ),
        )
        if expected != original and not sniot_fix.MULTI_SPACE_RE.search(collapsed):
            snippet = original[:50]
            issues.append(f"Лишние символы не убраны: «{snippet}»")
    return issues


def _apply_letterhead_or_chapter(paragraph, *, chapter: bool) -> None:
    sniot_fix.ensure_paragraph_centered(paragraph)
    sniot_fix.clear_first_line_indent(paragraph)
    paragraph.paragraph_format.left_indent = Cm(0)
    sniot_fix.set_single_line_spacing(paragraph)
    for run in paragraph.runs:
        sniot_fix.apply_run_font(run, bold=chapter or (paragraph.text.strip().upper().replace("Ё", "Е") == "ОТЧЕТ"))


def _apply_body(paragraph) -> None:
    sniot_fix.ensure_paragraph_justified(paragraph)
    sniot_fix.ensure_first_line_indent(paragraph)
    sniot_fix.set_single_line_spacing(paragraph)
    for run in paragraph.runs:
        sniot_fix.apply_run_font(run, bold=is_weekly_subsection(paragraph.text))


def _apply_signatory(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sniot_fix.clear_first_line_indent(paragraph)
    paragraph.paragraph_format.left_indent = Cm(0)
    sniot_fix.set_single_line_spacing(paragraph)
    for run in paragraph.runs:
        sniot_fix.apply_run_font(run, bold=False)


def apply_weekly_chapters(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        if not is_weekly_chapter_header(paragraph.text):
            continue
        canonical = canonical_weekly_chapter(paragraph.text)
        if paragraph.text.strip() != canonical:
            sniot_fix.set_paragraph_text(paragraph, canonical, bold=True)
        _apply_letterhead_or_chapter(paragraph, chapter=True)
        changed += 1
    return changed


def apply_weekly_office_to_document(doc: Document) -> dict:
    """Поля, шрифт, интервал, отступы, главы, пробелы. Без нумерации ДИ 1.8/1.9."""
    actions: list[str] = []
    sniot_fix.apply_page_setup_deloproizvodstvo(doc)
    actions.append("Поля А4 30/8/20/20 мм (Инструкция по делопроизводству 2025, п.18)")
    n_space = strip_weekly_double_spaces(doc)
    actions.append(f"Пробелы нормализованы (абзацев: {n_space})")
    n_mark = sniot_fix.remove_list_markers_in_body(doc)
    actions.append(f"Маркеры списка убраны (абзацев: {n_mark})")
    n_ch = apply_weekly_chapters(doc)
    actions.append(f"Заголовки разделов — прописными по центру ({n_ch})")

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            sniot_fix.set_single_line_spacing(paragraph)
            continue
        if is_weekly_chapter_header(text):
            continue
        if is_weekly_letterhead(text):
            _apply_letterhead_or_chapter(paragraph, chapter=False)
            continue
        if is_weekly_signatory_line(text):
            _apply_signatory(paragraph)
            continue
        if sniot_fix.paragraph_is_inside_table(paragraph):
            sniot_fix.set_single_line_spacing(paragraph)
            sniot_fix.clear_first_line_indent(paragraph)
            for run in paragraph.runs:
                sniot_fix.apply_run_font(run, bold=False)
            continue
        _apply_body(paragraph)

    sniot_fix.apply_table_paragraph_no_indent(doc)
    actions.append("Шрифт Times New Roman 14 pt")
    return {"actions": actions}


def validate_weekly_document(doc: Document) -> list[str]:
    issues: list[str] = []
    issues.extend(sniot_fix.validate_page_margins(doc))
    issues.extend(validate_weekly_spaces(doc))
    issues.extend(sniot_fix.validate_list_markers(doc))
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if is_weekly_chapter_header(text):
            if not sniot_fix.is_paragraph_centered(paragraph):
                issues.append(f"Заголовок раздела не по центру: {text[:50]}")
            if text != canonical_weekly_chapter(text):
                issues.append(f"Заголовок раздела не прописными: {text[:50]}")
            if not sniot_fix.paragraph_has_single_line_spacing(paragraph):
                issues.append(f"Заголовок: не одинарный интервал у «{text[:40]}»")
            if not sniot_fix.paragraph_has_zero_block_spacing(paragraph):
                issues.append(f"Заголовок: интервал перед/после абзаца не 0 у «{text[:40]}»")
            continue
        if is_weekly_letterhead(text) or is_weekly_signatory_line(text):
            continue
        if sniot_fix.paragraph_is_inside_table(paragraph):
            continue
        if not sniot_fix.is_paragraph_justified(paragraph):
            issues.append(f"Абзац не по ширине: {text[:50]}")
        indent = sniot_fix.first_line_indent_cm(paragraph)
        if abs(indent - sniot_fix.FIRST_LINE_INDENT_CM) > sniot_fix.FIRST_LINE_INDENT_TOLERANCE_CM:
            issues.append(
                f"Отступ первой строки не 1,25 см ({indent:.2f} см): {text[:50]}"
            )
        if not sniot_fix.paragraph_has_single_line_spacing(paragraph):
            issues.append(f"Тело: не одинарный интервал у «{text[:40]}»")
        if not sniot_fix.paragraph_has_zero_block_spacing(paragraph):
            issues.append(f"Тело: интервал перед/после абзаца не 0 у «{text[:40]}»")
    return issues


def _write_page_numbers(path: Path) -> None:
    data = path.read_bytes()
    path.write_bytes(sniot_fix.fix_page_numbering(data))


def format_weekly_docx(
    path: Path | str,
    *,
    backup: bool = True,
    allow_daily_initial: bool = False,
) -> dict:
    """
    Оформить готовый .docx итога на диске.
    backup=True — копия *_backup_YYYYMMDD_HHMM.docx рядом (не в Агент).
    Принятые / выданные на правку дневные файлы не трогать
    (кроме первой генерации: allow_daily_initial=True).
    """
    target = assert_weekly_writable(path)
    if not target.is_file():
        raise FileNotFoundError(f"Файл не найден: {target}")
    if is_accepted_daily_path(target):
        raise PermissionError(PROTECTED_DAILY_MESSAGE)
    if is_protected_daily_report(target) and not allow_daily_initial:
        raise PermissionError(PROTECTED_DAILY_MESSAGE)
    actions: list[str] = []
    backup_path = None
    if backup:
        stamp = datetime.now().strftime("%Y%m%d_%H%M")
        backup_path = target.with_name(f"{target.stem}_backup_{stamp}{target.suffix}")
        assert_weekly_writable(backup_path)
        shutil.copy2(target, backup_path)
        actions.append(f"Бэкап: {backup_path.name}")

    doc = Document(str(target))
    applied = apply_weekly_office_to_document(doc)
    actions.extend(applied["actions"])
    doc.save(str(target))
    _write_page_numbers(target)

    doc = Document(str(target))
    issues = validate_weekly_document(doc)
    issues.extend(sniot_fix.validate_page_numbering(target.read_bytes()))
    ok = not issues
    if ok:
        actions.append("Проверка: 0 замечаний (поля, пробелы, отступ, номера страниц)")
    return {
        "ok": ok,
        "path": str(target),
        "backup": str(backup_path) if backup_path else None,
        "actions": actions,
        "issues": issues,
    }


def process_weekly_itog_document(input_path: str) -> dict:
    """Вход для DocAgent: оформить на месте, не копировать в папку Агент."""
    target = Path(input_path)
    if is_protected_daily_report(target) or is_accepted_daily_path(target):
        return {
            "ok": True,
            "input": str(target),
            "output": str(target),
            "type": "ezhenedelnyy_itog",
            "mode": "daily_accepted_skip",
            "actions": [PROTECTED_DAILY_MESSAGE],
            "issues": [],
        }
    result = format_weekly_docx(input_path, backup=True)
    summary = {
        "ok": result["ok"],
        "input": result["path"],
        "output": result["path"],
        "type": "ezhenedelnyy_itog",
        "mode": "weekly_itog_office",
        "actions": list(result["actions"]),
        "issues": list(result["issues"]),
    }
    if result["issues"]:
        summary["actions"].insert(
            0,
            f"⚠ Остались замечания ({len(result['issues'])}) — откройте файл и проверьте",
        )
        for item in result["issues"][:8]:
            summary["actions"].append("⚠ " + item)
    else:
        summary["actions"].append("Еженедельный итог оформлен по Инструкции 2025")
    return summary


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Оформить еженедельный итог по Инструкции по делопроизводству 2025",
    )
    parser.add_argument("--target", type=Path, help="Готовый .docx итога")
    parser.add_argument("--check", action="store_true", help="Только проверка, без записи")
    args = parser.parse_args(argv)
    if not args.target:
        print("Укажите --target путь.docx")
        return 3
    target = Path(args.target)
    if args.check:
        doc = Document(str(target))
        issues = validate_weekly_document(doc)
        if target.is_file():
            issues.extend(sniot_fix.validate_page_numbering(target.read_bytes()))
        if issues:
            print("Замечания:")
            for item in issues:
                print(" -", item)
            return 1
        print("Проверка OK")
        return 0
    result = format_weekly_docx(target, backup=True)
    for act in result["actions"]:
        print(act)
    if result["issues"]:
        print("Замечания:")
        for item in result["issues"]:
            print(" -", item)
        return 1
    print(f"Готово: {result['path']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
