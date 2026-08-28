# -*- coding: utf-8 -*-
"""Тест: remove_duplicate_body_title не стирает тело без маркера главы."""
from __future__ import annotations

import inspect
import sys
import tempfile
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

import fix_sniot_document as fix


def _make_doc_with_body() -> bytes:
    doc = Document()
    doc.add_paragraph("ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("Настоящая должностная инструкция определяет функции.")
    doc.add_paragraph("Старший мастер должен знать:")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_remove_duplicate_does_not_wipe_body_without_chapter_marker():
    raw = _make_doc_with_body()
    before = len(Document(BytesIO(raw)).paragraphs)
    cleaned = fix.remove_duplicate_body_title(raw, first_chapter=None)
    after = len(Document(BytesIO(cleaned)).paragraphs)
    assert before >= 3
    assert after == before


def test_validate_empty_body_reports_issue():
    doc = Document()
    profile = fix.DocumentProfile("di", None, False, False, None)
    issues = fix.validate_body_not_empty(doc, profile)
    assert any("пуст" in i.lower() for i in issues)


def test_center_chapter_headers_sets_jc_in_xml():
    doc = Document()
    paragraph = doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    assert fix.center_chapter_headers(doc) == 1
    assert fix.is_paragraph_centered(paragraph)
    assert fix.paragraph_jc(paragraph) == "center"


def _di_satp_profile() -> fix.DocumentProfile:
    return fix.DocumentProfile(
        kind="di",
        first_chapter="1 ОБЩИЕ ПОЛОЖЕНИЯ",
        has_signatories=True,
        has_di_satp_numbering=True,
        tail_chapter_idx=None,
    )


def _build_correct_numbering_doc() -> Document:
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.8. в своей деятельности руководствуется законодательством.")
    doc.add_paragraph("1.8.1. Первый пункт раздела 1.8.")
    doc.add_paragraph("1.9. Старший мастер должен знать:")
    doc.add_paragraph("1.9.1. ТКП первый;")
    doc.add_paragraph("1.9.2. ТКП второй;")
    doc.add_paragraph("1.9.3. основы делопроизводства.")
    doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    doc.add_paragraph("2.1. выполняет следующие функции:")
    doc.add_paragraph("2.1.1. Функция одна.")
    doc.add_paragraph("2.2. Для выполнения возложенных на него функций старший мастер:")
    doc.add_paragraph("2.2.1. Обязанность одна.")
    doc.add_paragraph("3 ПРАВА")
    doc.add_paragraph("3.1. имеет право:")
    doc.add_paragraph("3.1.1. Право одно.")
    doc.add_paragraph("4 ВЗАИМООТНОШЕНИЯ")
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("5.1. несет ответственность:")
    doc.add_paragraph("5.1.1. Ответственность одна.")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник")
    return doc


def test_correct_numbering_not_modified():
    doc = _build_correct_numbering_doc()
    profile = _di_satp_profile()
    before = [p.text for p in doc.paragraphs]
    assert not fix.validate_numbering_blocks(doc, profile)
    changed = fix.fix_numbering_selective(doc, profile)
    after = [p.text for p in doc.paragraphs]
    assert changed == 0
    assert before == after


def test_wrong_prefix_19_fixed_only_in_block():
    doc = _build_correct_numbering_doc()
    profile = _di_satp_profile()
    doc.paragraphs[4].text = "1.5.1. ТКП с ошибкой."
    doc.paragraphs[5].text = "1.5.2. ТКП второй с ошибкой."
    assert fix.validate_numbering_blocks(doc, profile)
    fix.fix_numbering_selective(doc, profile)
    assert doc.paragraphs[4].text.startswith("1.9.1.")
    assert doc.paragraphs[5].text.startswith("1.9.2.")
    assert doc.paragraphs[2].text.startswith("1.8.1.")


def test_apply_body_paragraph_format_justify_and_indent():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    body = doc.add_paragraph("1.4.1. Текст абзаца тела документа.")
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, False, None)
    fix.center_chapter_headers(doc)
    fix.apply_body_paragraph_format(doc, profile)
    assert fix.is_paragraph_justified(body)
    assert abs(fix.first_line_indent_cm(body) - 1.25) < 0.1
    chapter = doc.paragraphs[0]
    assert fix.is_paragraph_centered(chapter)
    assert fix.first_line_indent_cm(chapter) < 0.1


def test_assert_path_writable_rejects_obmen():
    obmen = (
        r"N:\9 - Служба надёжности и охраны труда (СНiОТ)\!!!ОБМЕН"
        r"\РАССМОТРЕНИЕ  ДИ, РИ, ПОЛОЖЕНИЙ\САТП\test.docx"
    )
    with pytest.raises(PermissionError):
        fix.assert_path_writable(obmen)


def test_assert_path_writable_accepts_agent():
    agent = (
        r"N:\9 - Служба надёжности и охраны труда (СНiОТ)\Дубовик В.В\Агент"
        r"\ПРОЕКТ Старший мастер_оформлен.docx"
    )
    if fix.USER_AGENT_DIR.is_dir():
        resolved = fix.assert_path_writable(agent)
        assert fix.is_path_in_user_agent_dir(resolved)
        assert fix.is_path_in_writable_user_dir(resolved)


def test_assert_path_writable_accepts_projects():
    projects = (
        r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В\Проекты"
        r"\ИНСТРУКЦИЯ по эксплуатации ЦТП.docx"
    )
    assert fix.is_path_in_writable_user_dir(projects) is True
    assert fix.is_path_in_user_agent_dir(projects) is False
    if fix.USER_PROJECT_DIR.is_dir():
        resolved = fix.assert_path_writable(projects)
        assert fix.is_path_in_writable_user_dir(resolved)


def test_projects_sample_not_allowed():
    sample = (
        r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В\Проекты"
        r"\что-то_образец.docx"
    )
    assert fix.is_allowed_sample_path(sample) is False


def test_is_allowed_sample_path_obmen_false():
    obmen = (
        r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\!!!ОБМЕН"
        r"\РАССМОТРЕНИЕ  ДИ, РИ, ПОЛОЖЕНИЙ\ПРОЕКТ_образец.docx"
    )
    assert fix.is_allowed_sample_path(obmen) is False


def test_is_allowed_sample_path_agent_obrazec_true():
    sample = fix.USER_AGENT_DIR / "ПРОЕКТ Старший мастер_образец.docx"
    assert fix.is_allowed_sample_path(sample) is True


def test_is_allowed_sample_path_agent_oformlen_false():
    oformlen = fix.USER_AGENT_DIR / "ПРОЕКТ Старший мастер_оформлен.docx"
    assert fix.is_allowed_sample_path(oformlen) is False


def test_find_etalon_prefers_matching_name(tmp_path, monkeypatch):
    monkeypatch.setattr(fix, "USER_AGENT_DIR", tmp_path)
    other = tmp_path / "Другой документ_образец.docx"
    match = tmp_path / "ПРОЕКТ Старший мастер_образец.docx"
    target = tmp_path / "ПРОЕКТ Старший мастер_оформлен.docx"
    for path in (other, match, target):
        Document().save(path)
    found = fix.find_etalon_path(target)
    assert found is not None
    assert found.name == "ПРОЕКТ Старший мастер_образец.docx"


def test_find_etalon_ignores_explicit_obmen(tmp_path, monkeypatch):
    monkeypatch.setattr(fix, "USER_AGENT_DIR", tmp_path)
    sample = tmp_path / "ПРОЕКТ Старший мастер_образец.docx"
    target = tmp_path / "ПРОЕКТ Старший мастер_оформлен.docx"
    Document().save(sample)
    Document().save(target)
    obmen = Path(
        r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\!!!ОБМЕН"
        r"\РАССМОТРЕНИЕ  ДИ, РИ, ПОЛОЖЕНИЙ\чужой_образец.docx"
    )
    found = fix.find_etalon_path(target, explicit=obmen)
    assert found is not None
    assert found.resolve() == sample.resolve()


def test_find_etalon_skips_unrelated_sample(tmp_path, monkeypatch):
    monkeypatch.setattr(fix, "USER_AGENT_DIR", tmp_path)
    unrelated = tmp_path / "Техническая информация_образец.docx"
    target = tmp_path / "ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx"
    Document().save(unrelated)
    Document().save(target)
    assert fix.find_etalon_path(target) is None


def test_sanitize_full_strips_body_tabs():
    text = "1.6.\tВ своей деятельности инженер ЛСиМ руководствуется ТНПА."
    cleaned = fix.sanitize_paragraph_text(text, full=True)
    assert "\t" not in cleaned
    assert "1.6." in cleaned
    assert "ЛСиМ" in cleaned


def test_no_empty_line_after_soglasovano():
    doc = Document()
    doc.add_paragraph("5.1.1. Последний пункт.")
    doc.add_paragraph("")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    doc.add_paragraph("")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("")
    doc.add_paragraph("Начальник службы")
    profile = fix.DocumentProfile("di", None, True, False, None)
    fix.fix_signatory_block_format(doc, profile)
    soglas_idx = fix.find_soglasovano_index(doc)
    assert doc.paragraphs[soglas_idx + 1].text.strip()
    issues = fix.validate_signatory_block(doc, profile)
    assert not any("после «Согласовано»" in i for i in issues)


def test_russian_phrase_local_legal_acts_replaced():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p = doc.add_paragraph(
        "2.2.17. Выполняет локальные правовые акты предприятия (приказы, указания)."
    )
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, False, None)
    changed = fix.apply_russian_phrase_rules(doc, profile)
    assert changed >= 1
    assert "требования локальных правовых актов" in p.text
    assert "Выполняет локальные правовые акты" not in p.text
    assert p.text.startswith("2.2.17.")


def test_phrase_rules_leave_abbreviations():
    doc = Document()
    p = doc.add_paragraph(
        "1.6. Инженер ЛСиМ руководствуется ТКП, НПА, СНиОТ и САТП."
    )
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, False, None)
    fix.apply_russian_phrase_rules(doc, profile)
    assert "ЛСиМ" in p.text
    assert "ТКП" in p.text
    assert "НПА" in p.text
    assert "СНиОТ" in p.text
    assert "САТП" in p.text
    assert "Осим" not in p.text


def test_process_does_not_change_source_wording_on_di():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p = doc.add_paragraph(
        "2.2.17. Выполняет локальные правовые акты предприятия (приказы, указания)."
    )
    doc.add_paragraph(
        "1.6. В своей деятельности инженер руководствуется нормативными правовыми актами."
    )
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, False, None)
    path = Path("ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx")
    fix.process_sniot_document(doc, profile, source_path=path)
    joined = "\n".join(x.text for x in doc.paragraphs)
    assert "Выполняет локальные правовые акты" in joined
    assert "требования локальных правовых актов" not in joined
    assert "нормативными правовыми актами" in joined
    assert "нормативно-правовыми" not in joined
    src = inspect.getsource(fix.process_sniot_document)
    assert "apply_russian_phrase_rules" not in src


def test_remove_triple_empty_lines_in_body():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Первый пункт.")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("1.4.2. Второй пункт.")
    removed = fix.remove_extra_empty_lines_in_body(doc)
    assert removed >= 2
    texts = [p.text.strip() for p in doc.paragraphs]
    assert "" not in texts[2:5]


def test_remove_empty_after_chapter_header_even_before_signatory():
    doc = Document()
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    removed = fix.remove_empty_lines_after_chapter_headers(doc)
    assert removed >= 2
    assert doc.paragraphs[0].text.strip().startswith("5 ОТВЕТСТВЕННОСТЬ")
    assert doc.paragraphs[1].text.strip() == "Разработал:"


def test_signatory_tail_detected_without_razrabotal_marker():
    doc = Document()
    doc.add_paragraph("5.1.1. Последний пункт.")
    doc.add_paragraph("Начальник службы АТП\tА.И.Торопчин")
    doc.add_paragraph("СОГЛАСОВАНО")
    doc.add_paragraph("Заместитель")
    tail = fix.find_signatory_tail_start(doc)
    assert tail == 1
    profile = fix.DocumentProfile("di", None, True, False, None)
    fix.ensure_razrabotal_marker(doc)
    razrab_idx = fix.find_paragraph_index(doc, "Разработал:")
    assert doc.paragraphs[razrab_idx].text.strip() == "Разработал:"
    assert doc.paragraphs[razrab_idx + 1].text.strip().startswith("Начальник")
    fix.fix_signatory_block_format(doc, profile)
    razrab_idx = fix.find_paragraph_index(doc, "Разработал:")
    assert fix.count_empty_lines_before(doc, razrab_idx) == 1


def test_maybe_restore_skips_when_chapters_present():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Пункт.")
    doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    doc.add_paragraph("2.1.1. Функция.")
    doc.add_paragraph("3 ПРАВА")
    doc.add_paragraph("3.1.1. Право.")
    doc.add_paragraph("4 ВЗАИМООТНОШЕНИЯ")
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("5.1.1. Ответственность.")
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = Path(tmp.name)
    doc.save(path)
    try:
        path = path.rename(
            path.with_name("ПРОЕКТ Старший мастер_test_restore.docx")
        )
    except OSError:
        pass
    n, msg = fix.maybe_restore_senior_master_body(path)
    assert n == 0
    assert "Восстановлено" not in msg
    path.unlink(missing_ok=True)


def test_empty_line_before_razrabotal_preserved_after_body_cleanup():
    doc = Document()
    doc.add_paragraph("5.1.1. Последний пункт.")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник")
    profile = fix.DocumentProfile("di", None, True, False, None)
    fix.ensure_single_empty_line_before(doc, "Разработал:")
    razrab_idx = fix.find_paragraph_index(doc, "Разработал:")
    assert fix.count_empty_lines_before(doc, razrab_idx) == 1
    fix.remove_extra_empty_lines_in_body(doc)
    razrab_idx = fix.find_paragraph_index(doc, "Разработал:")
    assert fix.count_empty_lines_before(doc, razrab_idx) == 1


def test_validate_empty_lines_in_body_reports_double():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Текст.")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("1.4.2. Ещё текст.")
    issues = fix.validate_empty_lines_in_body(doc)
    assert any("пустые строки" in i.lower() for i in issues)


def test_materialize_chapter_one_numbering_visible():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("Настоящая должностная инструкция определяет функции.")
    doc.add_paragraph("Настоящую должностную инструкцию должны знать:")
    doc.add_paragraph("1.4. в своей деятельности руководствуется законодательством.")
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, True, None)
    changed = fix.materialize_chapter_one_numbering(doc, profile)
    assert changed >= 2
    assert doc.paragraphs[1].text.startswith("1.1. Настоящая")
    assert doc.paragraphs[2].text.startswith("1.2. Настоящую")
    assert doc.paragraphs[3].text.startswith("1.4. в своей")
    assert not fix.has_word_list_numbering(doc.paragraphs[1])


def test_materialize_chapter_one_preserves_14x_block():
    doc = _build_correct_numbering_doc()
    profile = _di_satp_profile()
    p14_1 = doc.paragraphs[2].text
    fix.materialize_chapter_one_numbering(doc, profile)
    assert doc.paragraphs[2].text == p14_1
    assert doc.paragraphs[2].text.startswith("1.8.1.")
    assert not fix.validate_numbering_blocks(doc, profile)


def test_chapter_one_hidden_numpr_triggers_materialize():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p = doc.add_paragraph("Настоящая должностная инструкция определяет функции.")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    p_pr.append(num_pr)
    doc.add_paragraph("1.4. в своей деятельности руководствуется законодательством.")
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, True, None)
    assert fix.chapter_one_has_hidden_numpr(doc, profile)
    fix.materialize_chapter_one_numbering(doc, profile)
    assert doc.paragraphs[1].text.startswith("1.1. Настоящая")
    assert not fix.has_word_list_numbering(doc.paragraphs[1])


def test_find_section_header_skips_wrong_intro_14():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph(
        "1.4. Назначение, перемещение и увольнение старшего мастера производится приказом."
    )
    doc.add_paragraph("1.4. Старший мастер в своей деятельности руководствуется:")
    idx = fix.find_section_header_index(
        doc, "1.8", "в своей деятельности руководствуется"
    )
    assert idx == 2
    assert "руководствуется" in doc.paragraphs[idx].text


def test_find_section_header_by_phrase_even_if_numbered_11():
    """Вводный 1.4. Назначение… не подменяет раздел, даже если у раздела номер 1.1."""
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph(
        "1.4. Назначение, перемещение и увольнение старшего мастера производится приказом."
    )
    doc.add_paragraph("1.1. Старший мастер в своей деятельности руководствуется:")
    doc.add_paragraph("1.4.1. Пункт списка раздела 1.4.")
    doc.add_paragraph("1.2. Старший мастер должен знать:")
    doc.add_paragraph("1.5.1. ТКП пункт.")
    idx14 = fix.find_section_header_index(
        doc, "1.8", "в своей деятельности руководствуется"
    )
    idx15 = fix.find_section_header_index(doc, "1.9", "должен знать")
    assert idx14 == 2
    assert idx15 == 4
    assert "руководствуется" in doc.paragraphs[idx14].text
    assert "должен знать" in doc.paragraphs[idx15].text


def test_process_keeps_14_block_when_header_was_11():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Настоящая должностная инструкция определяет функции.")
    doc.add_paragraph(
        "1.4. Назначение, перемещение и увольнение старшего мастера производится приказом."
    )
    doc.add_paragraph("1.1. Старший мастер в своей деятельности руководствуется:")
    doc.add_paragraph("1.4.1. Закон Республики Беларусь.")
    doc.add_paragraph("1.4.2. Правила техники безопасности.")
    doc.add_paragraph("1.2. Старший мастер должен знать:")
    doc.add_paragraph("1.5.1. ТКП 608-2025.")
    doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    doc.add_paragraph("2.1. выполняет следующие функции:")
    doc.add_paragraph("2.1.1. Функция одна.")
    doc.add_paragraph("2.2. Для выполнения возложенных на него функций старший мастер:")
    doc.add_paragraph("2.2.1. Обязанность одна.")
    doc.add_paragraph("3 ПРАВА")
    doc.add_paragraph("3.1. имеет право:")
    doc.add_paragraph("3.1.1. Право одно.")
    doc.add_paragraph("4 ВЗАИМООТНОШЕНИЯ")
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("5.1. несет ответственность:")
    doc.add_paragraph("5.1.1. Ответственность одна.")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник")
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    fix.process_sniot_document(doc, profile)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert any(t.startswith("1.8. Старший мастер") and "руководствуется" in t for t in texts)
    assert any(t.startswith("1.9. Старший мастер") and "должен знать" in t for t in texts)
    assert any(t.startswith("1.8.1. Закон") for t in texts)
    assert any(t.startswith("1.9.1. ТКП") for t in texts)
    assert not any(t.startswith("1.4. Старший мастер") and "руководствуется" in t for t in texts)
    assert not any(t.startswith("1.5. Старший мастер") and "должен знать" in t for t in texts)
    assert not fix.validate_numbering_block_starts(doc, profile)


def test_strip_wrong_intro_section_prefixes():
    """Вводный 1.4. Назначение… остаётся; раздел «руководствуется» — 1.8. по образцу."""
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Настоящая должностная инструкция определяет функции.")
    doc.add_paragraph(
        "Назначение, перемещение и увольнение старшего мастера производится приказом."
    )
    doc.add_paragraph("1.4. Старший мастер в своей деятельности руководствуется:")
    doc.add_paragraph("1.4.1. Пункт списка.")
    doc.add_paragraph("1.5. Старший мастер должен знать:")
    doc.add_paragraph("1.5.1. ТКП пункт.")
    profile = _di_satp_profile()
    changed = fix.materialize_chapter_one_numbering(doc, profile)
    assert changed >= 1
    assert doc.paragraphs[2].text.startswith("1.4. Назначение")
    idx = fix.find_section_header_index(
        doc, "1.8", "в своей деятельности руководствуется"
    )
    assert "руководствуется" in doc.paragraphs[idx].text
    assert doc.paragraphs[idx + 1].text.startswith("1.4.1.")


def test_correct_numbering_not_modified_after_conservative_process():
    doc = _build_correct_numbering_doc()
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    before = [p.text for p in doc.paragraphs if p.text.strip()]
    fix.process_sniot_document(doc, profile)
    after = [p.text for p in doc.paragraphs if p.text.strip()]
    assert before == after
    assert not fix.validate_numbering_blocks(doc, profile)


def test_signatory_tab_stops_from_etalon_layout():
    doc = Document()
    doc.add_paragraph("5.1.1. Последний пункт.")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Начальник службы АТП А.И.Торопчин")
    etalon = Document()
    etalon.add_paragraph("5.1.1. Последний пункт.")
    e_raz = etalon.add_paragraph("Разработал:")
    etalon.add_paragraph("Начальник службы АТП\tА.И.Торопчин")
    e_ppr = e_raz._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), "9072")
    tabs.append(tab)
    e_ppr.append(tabs)
    profile = fix.DocumentProfile("di", None, True, False, None)
    fix.apply_signatory_layout_from_etalon(doc, etalon, profile)
    raz_idx = fix.find_razrabotal_index(doc)
    assert fix.paragraph_has_tab_stops(doc.paragraphs[raz_idx])
    assert "\t" in doc.paragraphs[raz_idx + 1].text


def test_apply_signatory_line_spacing_sets_one_point_five():
    from docx.enum.text import WD_LINE_SPACING

    doc = Document()
    doc.add_paragraph("5.1.1. Последний пункт.")
    doc.add_paragraph("")
    p_raz = doc.add_paragraph("Разработал:")
    p_dev = doc.add_paragraph("Инженер\tИ.И. Иванов")
    doc.add_paragraph("")
    p_sog = doc.add_paragraph("Согласовано:")
    p_boss = doc.add_paragraph("Начальник\tП.П. Петров")
    for paragraph in (p_raz, p_dev, p_sog, p_boss):
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.line_spacing = 1.0
    profile = fix.DocumentProfile("di", None, True, False, None)
    changed = fix.apply_signatory_line_spacing(doc, profile)
    assert changed >= 4
    for paragraph in (p_raz, p_dev, p_sog, p_boss):
        assert fix.paragraph_has_one_point_five_spacing(paragraph)
    assert not fix.validate_signatory_line_spacing(doc, profile)


def test_apply_signatory_line_spacing_does_not_touch_body():
    from docx.enum.text import WD_LINE_SPACING

    doc = Document()
    body = doc.add_paragraph("1.4.1. Текст тела документа.")
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    body.paragraph_format.line_spacing = 1.0
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    profile = fix.DocumentProfile("di", None, True, False, None)
    fix.apply_signatory_line_spacing(doc, profile)
    assert body.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
    assert body.paragraph_format.line_spacing == 1.0


def test_fix_signatory_block_format_enforces_line_spacing():
    from docx.enum.text import WD_LINE_SPACING

    doc = Document()
    doc.add_paragraph("5.1.1. Последний пункт.")
    doc.add_paragraph("Разработал:")
    p_dev = doc.add_paragraph("Инженер")
    doc.add_paragraph("Согласовано:")
    p_boss = doc.add_paragraph("Начальник")
    for paragraph in (p_dev, p_boss):
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    profile = fix.DocumentProfile("di", None, True, False, None)
    fix.fix_signatory_block_format(doc, profile)
    assert not fix.validate_signatory_line_spacing(doc, profile)
    assert not fix.validate_signatory_block(doc, profile)


def test_body_indent_after_hanging_style():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p = doc.add_paragraph("1.4.1. Абзац с отступом списка.")
    p_pr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, False, None)
    fix.apply_body_paragraph_format(doc, profile)
    assert abs(fix.first_line_indent_cm(p) - 1.25) < 0.1


def _count_empty_paragraphs(doc: Document) -> int:
    return sum(1 for p in doc.paragraphs if not p.text.strip())


def test_ensure_chapter_header_spacing_does_not_add_empty_after_header():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Первый пункт.")
    doc.add_paragraph("")
    doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    doc.add_paragraph("2.1.1. Функция.")
    before = _count_empty_paragraphs(doc)
    fix.ensure_chapter_header_spacing(doc)
    after = _count_empty_paragraphs(doc)
    assert after <= before
    ch2_idx = fix.find_paragraph_index(doc, "2 ФУНКЦИИ")
    assert doc.paragraphs[ch2_idx + 1].text.strip().startswith("2.1")


def test_prevent_chapter_header_orphan_moves_page_break_to_content():
    doc = Document()
    hdr = doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    body = doc.add_paragraph("2.1.1. Текст функции.")
    fix.set_page_break_before(hdr, True)
    assert fix.paragraph_has_page_break_before(hdr)
    changed = fix.prevent_chapter_header_orphan(doc)
    assert changed >= 1
    assert not fix.paragraph_has_page_break_before(hdr)
    assert fix.paragraph_has_page_break_before(body)
    assert not fix.validate_chapter_header_orphan(doc)


def test_header_not_separated_from_content_by_page_break_only():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Пункт.")
    hdr = doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    doc.add_paragraph("2.1.1. Следующий текст.")
    fix.set_page_break_before(hdr, True)
    fix.prevent_chapter_header_orphan(doc)
    hdr_idx = fix.find_paragraph_index(doc, "2 ФУНКЦИИ")
    content_idx = fix.find_first_nonempty_paragraph_after(doc, hdr_idx)
    assert content_idx is not None
    assert not fix.paragraph_has_page_break_before(hdr)
    assert fix.paragraph_has_page_break_before(doc.paragraphs[content_idx])


def test_process_sniot_document_reduces_empty_after_headers():
    doc = _build_correct_numbering_doc()
    added = 0
    for idx in fix.find_chapter_header_indices(doc):
        if idx + 1 < len(doc.paragraphs) and doc.paragraphs[idx + 1].text.strip():
            fix.insert_empty_paragraph_before(doc.paragraphs[idx + 1])
            added += 1
    assert added >= 2

    def empties_after_headers(d: Document) -> int:
        total = 0
        for hdr_idx in fix.find_chapter_header_indices(d):
            if hdr_idx + 1 < len(d.paragraphs) and not d.paragraphs[hdr_idx + 1].text.strip():
                total += 1
        return total

    before = empties_after_headers(doc)
    assert before >= 2
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    fix.process_sniot_document(doc, profile)
    after = empties_after_headers(doc)
    assert after < before
    assert not fix.validate_chapter_headers(doc)


def _signatory_tail_doc() -> Document:
    doc = Document()
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("5.1. несет ответственность:")
    doc.add_paragraph("5.1.1. Последний пункт ответственности.")
    doc.add_paragraph("")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер\tИ.И. Иванов")
    doc.add_paragraph("")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник\tП.П. Петров")
    return doc


def test_validate_last_two_pages_rejects_page_break_only_on_razrabotal():
    doc = _signatory_tail_doc()
    profile = fix.DocumentProfile("di", None, True, False, 0)
    razrab_idx = fix.find_razrabotal_index(doc)
    fix.set_page_break_before(doc.paragraphs[razrab_idx], True)
    issues = fix.validate_last_two_pages_layout(doc, profile)
    assert any("только перед «Разработал:»" in i for i in issues)
    assert fix.signatories_appear_orphaned(doc, profile)


def test_fix_last_pages_and_signatories_clears_keep_and_formats_block():
    doc = _signatory_tail_doc()
    profile = fix.DocumentProfile("di", None, True, False, 0)
    razrab_idx = fix.find_razrabotal_index(doc)
    body = doc.paragraphs[2]
    body.paragraph_format.keep_with_next = True
    doc.paragraphs[razrab_idx].paragraph_format.keep_with_next = True
    doc.paragraphs[razrab_idx + 1].paragraph_format.keep_with_next = True
    fix.fix_last_pages_and_signatories(doc, profile)
    last_body = fix.find_last_body_paragraph_before_signatories(doc)
    assert last_body is not None
    assert doc.paragraphs[last_body].paragraph_format.keep_with_next
    assert doc.paragraphs[razrab_idx].paragraph_format.keep_with_next
    assert fix.count_empty_lines_before(doc, razrab_idx) == 1
    assert not fix.validate_signatory_block(doc, profile)
    assert not fix.validate_last_two_pages_layout(doc, profile)


def test_fix_last_pages_page_breaks_never_on_razrabotal():
    doc = _signatory_tail_doc()
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    razrab_idx = fix.find_razrabotal_index(doc)
    fix.set_page_break_before(doc.paragraphs[razrab_idx], True)
    strategy = fix.fix_last_pages_page_breaks(doc, profile, force=True)
    assert strategy in ("before_paragraph", "natural")
    assert not fix.paragraph_has_page_break_before(doc.paragraphs[razrab_idx])
    issues = fix.validate_last_two_pages_layout(doc, profile)
    assert not any("только перед «Разработал:»" in i for i in issues)


def test_signatories_not_orphaned_when_body_text_present():
    doc = _signatory_tail_doc()
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    fix.fix_last_pages_and_signatories(doc, profile)
    last_body = fix.find_last_body_paragraph_before_signatories(doc)
    assert last_body is not None
    assert not fix.is_chapter_header(doc.paragraphs[last_body].text)
    assert not fix.signatories_appear_orphaned(doc, profile)


def test_orphaned_signatories_detected_when_only_chapter_header_before():
    doc = Document()
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    doc.add_paragraph("")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник")
    profile = fix.DocumentProfile("di", None, True, False, 0)
    issues = fix.validate_last_two_pages_layout(doc, profile)
    assert any("заголовок главы" in i.lower() for i in issues)
    assert fix.signatories_appear_orphaned(doc, profile)


def test_strip_leading_dot_before_numbering():
    doc = Document()
    doc.add_paragraph(".1.4.1. Текст пункта.")
    doc.add_paragraph(". 1.5.2. Второй пункт.")
    doc.add_paragraph("1.4.3.. Третий пункт с двойной точкой.")
    doc.add_paragraph(".1.4. в своей деятельности руководствуется:")
    changed = fix.strip_leading_dot_before_numbering(doc)
    assert changed == 4
    assert doc.paragraphs[0].text.startswith("1.4.1. Текст")
    assert doc.paragraphs[1].text.startswith("1.5.2. Второй")
    assert doc.paragraphs[2].text.startswith("1.4.3. Третий")
    assert doc.paragraphs[3].text.startswith("1.4. в своей")
    assert not fix.validate_leading_dot_before_numbering(doc)


def test_strip_leading_dot_skips_chapter_headers():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph(".1.4.1. Пункт после главы.")
    changed = fix.strip_leading_dot_before_numbering(doc)
    assert changed == 1
    assert doc.paragraphs[0].text == "1 ОБЩИЕ ПОЛОЖЕНИЯ"


def _insert_tracked_ins_before_first_run(paragraph, text: str = ".") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "8")
    ins.set(qn("w:author"), "Литвинов")
    ins.set(qn("w:date"), "2026-08-12T13:56:00Z")
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = text
    run_el.append(t_el)
    ins.append(run_el)
    first_r = paragraph._p.find(qn("w:r"))
    if first_r is not None:
        first_r.addprevious(ins)
    else:
        paragraph._p.append(ins)


def _append_tracked_del(paragraph, text: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    deleted = OxmlElement("w:del")
    deleted.set(qn("w:id"), "9")
    deleted.set(qn("w:author"), "Литвинов")
    deleted.set(qn("w:date"), "2026-08-12T13:56:00Z")
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:delText")
    t_el.text = text
    run_el.append(t_el)
    deleted.append(run_el)
    paragraph._p.append(deleted)


def test_visible_text_sees_ins_dot_before_511():
    doc = Document()
    p = doc.add_paragraph("5.1.1. Создание условий труда.")
    _insert_tracked_ins_before_first_run(p, ".")
    assert not p.text.startswith(".")
    assert fix.paragraph_visible_text(p).startswith(".5.1.1.")
    assert fix.validate_leading_dot_before_numbering(doc)
    assert fix.validate_unaccepted_revisions(doc)


def test_accept_tracked_changes_keeps_ins_removes_del():
    doc = Document()
    p = doc.add_paragraph("5.1.1. Создание условий труда.")
    _insert_tracked_ins_before_first_run(p, ".")
    _append_tracked_del(p, " (специалиста по охране труда)")
    changed = fix.accept_tracked_changes(doc)
    assert changed >= 2
    assert fix.count_tracked_change_nodes(doc) == (0, 0)
    assert doc.paragraphs[0].text.startswith(".5.1.1.")
    assert "специалиста" not in doc.paragraphs[0].text


def test_process_accepts_revisions_and_strips_dot_before_511():
    doc = _build_correct_numbering_doc()
    idx_511 = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("5.1.1."):
            idx_511 = i
            _insert_tracked_ins_before_first_run(para, ".")
            break
    assert idx_511 is not None
    _append_tracked_del(doc.paragraphs[idx_511], " (специалиста по охране труда)")
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    fix.process_sniot_document(doc, profile)
    assert fix.count_tracked_change_nodes(doc) == (0, 0)
    found = [p.text.strip() for p in doc.paragraphs if p.text.strip().startswith("5.1.1.")]
    assert found
    assert found[0].startswith("5.1.1.")
    assert not found[0].startswith(".")
    assert "специалиста" not in found[0]
    assert not fix.validate_leading_dot_before_numbering(doc)
    assert not fix.validate_unaccepted_revisions(doc)


def test_deduplicate_list_and_manual_numbering():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    p = doc.add_paragraph("1.5.1. 1.5.1. ТКП 608-2025 текст.")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "2")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    changed = fix.deduplicate_manual_and_list_numbering(doc)
    assert changed >= 2
    assert doc.paragraphs[0].text.startswith("1.5.1. ТКП")
    assert "1.5.1. 1.5.1." not in doc.paragraphs[0].text
    assert not fix.has_word_list_numbering(doc.paragraphs[0])
    assert not fix.validate_duplicate_list_numbering(doc)


def test_remove_orphan_word_list_numbering_keeps_etalon_style():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p = doc.add_paragraph("Старший мастер обязан соблюдать требования безопасности.")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    # python-docx default template: numId=9 — decimal, не bullet
    num_id.set(qn("w:val"), "9")
    num_pr.append(num_id)
    p_pr.append(num_pr)
    assert fix.has_word_list_numbering(p)
    changed = fix.remove_word_list_numbering_in_body(doc)
    assert changed == 0
    assert fix.has_word_list_numbering(p)


def test_remove_word_list_when_manual_number_present():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p = doc.add_paragraph("5.1.1. Создание условий труда.")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    p_pr.append(num_pr)
    changed = fix.remove_word_list_numbering_in_body(doc)
    assert changed == 1
    assert not fix.has_word_list_numbering(p)


def test_ensure_di_satp_chapter_one_numbering():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("Настоящая должностная инструкция определяет функции.")
    doc.add_paragraph("Настоящую должностную инструкцию должны знать:")
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, True, None)
    changed = fix.ensure_di_satp_chapter_one_numbering(doc, profile)
    assert changed >= 2
    assert doc.paragraphs[1].text.startswith("1.1. Настоящая")
    assert doc.paragraphs[2].text.startswith("1.2. Настоящую")


def test_insert_signatory_tab_line():
    text = "Начальник службы АТП А.И.Торопчин"
    assert fix.insert_signatory_tab_line(text) == "Начальник службы АТП\tА.И.Торопчин"


def test_copy_signatory_ppr_copies_tab_stops():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    etalon = Document()
    target = Document()
    e_p = etalon.add_paragraph("Разработал:")
    t_p = target.add_paragraph("Разработал:")
    e_ppr = e_p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), "6803")
    tabs.append(tab)
    e_ppr.append(tabs)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    e_ppr.append(jc)

    assert fix.copy_signatory_ppr_from_etalon(e_p, t_p)
    assert fix.paragraph_has_tab_stops(t_p)
    assert fix.paragraph_jc(t_p) in ("both", "justify")
    assert fix.paragraph_has_one_point_five_spacing(t_p)


def test_apply_signatory_layout_from_etalon_syncs_razrabotal_tabs():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    etalon = Document()
    doc = Document()
    doc.add_paragraph("5.1.7. Последний пункт.")
    e_razrab = etalon.add_paragraph("Разработал:")
    doc.add_paragraph("Разработал:")
    e_sig = etalon.add_paragraph("Начальник службы АТП\tА.И.Торопчин")
    doc.add_paragraph("Начальник службы АТП А.И.Торопчин")

    for p in (e_razrab, e_sig):
        ppr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), "6803")
        tabs.append(tab)
        ppr.append(tabs)

    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", True, True, None)
    changed = fix.apply_signatory_layout_from_etalon(doc, etalon, profile)
    assert changed >= 2
    razrab = doc.paragraphs[fix.find_razrabotal_index(doc)]
    sig = doc.paragraphs[fix.find_razrabotal_index(doc) + 1]
    assert fix.paragraph_has_tab_stops(razrab)
    assert fix.paragraph_has_tab_stops(sig)
    assert "\t" in sig.text
    assert not fix.validate_signatory_tab_stops(doc, profile)


def test_normalize_body_paragraph_styles():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p = doc.add_paragraph("1.4.1. Текст абзаца тела документа.")
    p.style = doc.styles["No Spacing"]
    changed = fix.normalize_body_paragraph_styles(doc)
    assert changed == 1
    assert doc.paragraphs[1].style.name == "Normal"


def test_fix_wrong_section_header_numbers():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Старший мастер в своей деятельности руководствуется:")
    doc.add_paragraph("1.4.1. Пункт тела.")
    doc.add_paragraph("1.2. Старший мастер должен знать:")
    doc.add_paragraph("1.5.1. ТКП текст.")
    changed = fix.fix_missing_section_headers(doc)
    assert changed == 2
    assert doc.paragraphs[1].text.startswith("1.8. Старший мастер")
    assert doc.paragraphs[3].text.startswith("1.9. Старший мастер")


def test_max_one_empty_before_chapter_after_process():
    doc = _build_correct_numbering_doc()
    # лишние пустые перед главой 5
    ch5_idx = fix.find_paragraph_index(doc, "5 ОТВЕТСТВЕННОСТЬ")
    for _ in range(3):
        fix.insert_empty_paragraph_before(doc.paragraphs[ch5_idx])
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    fix.process_sniot_document(doc, profile)
    ch5_idx = fix.find_paragraph_index(doc, "5 ОТВЕТСТВЕННОСТЬ")
    assert fix.count_empty_lines_before(doc, ch5_idx) <= 1
    assert not fix.validate_empty_lines_in_body(doc)


def test_indent_preserved_after_etalon_align():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4. в своей деятельности руководствуется:")
    body = doc.add_paragraph("1.4.1. Текст абзаца тела документа.")
    etalon = Document()
    etalon.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    etalon.add_paragraph("1.4. в своей деятельности руководствуется:")
    et_body = etalon.add_paragraph("1.4.1. Текст абзаца тела документа.")
    et_body.paragraph_format.first_line_indent = None
    p_pr = et_body._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is not None:
        p_pr.remove(ind)
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, True, None)
    fix.copy_paragraph_format_from_etalon(
        et_body,
        body,
        profile=profile,
        target_idx=2,
        doc=doc,
    )
    assert abs(fix.first_line_indent_cm(body) - 1.25) < 0.1


def test_single_paragraph_break_for_signatories():
    doc = _signatory_tail_doc()
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    strategy = fix.fix_last_pages_page_breaks(doc, profile, force=True)
    assert strategy == "natural"
    razrab_idx = fix.find_razrabotal_index(doc)
    assert not fix.paragraph_has_page_break_before(doc.paragraphs[razrab_idx])
    last_body = fix.find_last_body_paragraph_before_signatories(doc)
    assert last_body is not None
    assert not fix.paragraph_has_page_break_before(doc.paragraphs[last_body])


def test_apply_signatory_page_break_never_on_razrabotal():
    doc = _signatory_tail_doc()
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    strategy = fix.apply_signatory_page_break(doc, profile)
    assert strategy == "before_paragraph"
    razrab_idx = fix.find_razrabotal_index(doc)
    assert not fix.paragraph_has_page_break_before(doc.paragraphs[razrab_idx])
    last_body = fix.find_last_body_paragraph_before_signatories(doc)
    assert last_body is not None
    assert fix.paragraph_has_page_break_before(doc.paragraphs[last_body])


def test_every_body_numbered_para_has_indent_after_process():
    doc = _build_correct_numbering_doc()
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    fix.process_sniot_document(doc, profile)
    for idx, paragraph in enumerate(doc.paragraphs):
        if not fix.should_apply_body_paragraph_format(paragraph.text, idx, doc):
            continue
        assert fix.first_line_indent_cm(paragraph) >= 1.2, paragraph.text[:60]
    assert not fix.validate_body_paragraph_format(doc, profile)


def test_every_chapter_header_has_one_empty_line_before():
    doc = _build_correct_numbering_doc()
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    ch2_idx = fix.find_paragraph_index(doc, "2 ФУНКЦИИ")
    if ch2_idx > 0 and fix.is_paragraph_empty(doc.paragraphs[ch2_idx - 1]):
        el = doc.paragraphs[ch2_idx - 1]._element
        el.getparent().remove(el)
    fix.process_sniot_document(doc, profile)
    for idx in fix.find_chapter_header_indices(doc):
        if idx == 0:
            continue
        assert fix.count_empty_lines_before(doc, idx) == 1
    assert not fix.validate_chapter_headers(doc)


def test_chapter_titles_centered_body_items_keep_indent():
    """«Заголовки разделов» = названия глав; «1.9. должен знать:» — пункт тела."""
    assert fix.is_chapter_header("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    assert fix.is_chapter_header("2 ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    assert not fix.is_chapter_header("1.9. Старший мастер должен знать:")
    doc = _build_correct_numbering_doc()
    for paragraph in doc.paragraphs:
        if fix.is_section_header(paragraph.text):
            fix.clear_first_line_indent(paragraph)
            fix.ensure_paragraph_centered(paragraph)
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    fix.process_sniot_document(doc, profile)
    ch_found = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if fix.is_chapter_header(text):
            ch_found += 1
            assert fix.is_paragraph_centered(paragraph), text[:60]
            assert abs(fix.first_line_indent_cm(paragraph)) < 0.1, text[:60]
        if text.startswith("1.9. Старший мастер должен знать"):
            assert not fix.is_paragraph_centered(paragraph), text[:60]
            assert abs(fix.first_line_indent_cm(paragraph) - 1.25) < 0.1, text[:60]
    assert ch_found >= 2
    assert not fix.validate_section_headers(doc)
    assert not fix.validate_chapter_headers(doc)


def test_align_spacing_enforces_chapter_empty_line_not_etalon_zero():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Последний пункт главы 1.")
    doc.add_paragraph("")
    doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    doc.add_paragraph("2.1.1. Функция.")
    etalon = Document()
    etalon.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    etalon.add_paragraph("1.4.1. Последний пункт главы 1.")
    etalon.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    etalon.add_paragraph("2.1.1. Функция.")
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, True, None)
    fix.align_spacing_to_etalon(doc, etalon, profile)
    ch2_idx = fix.find_paragraph_index(doc, "2 ФУНКЦИИ")
    assert fix.count_empty_lines_before(doc, ch2_idx) == 1


def test_fix_last_pages_restores_chapter_spacer_after_cleanup():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Пункт.")
    doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    doc.add_paragraph("2.1.1. Функция.")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    profile = fix.DocumentProfile("di", None, True, False, None)
    fix.ensure_chapter_header_spacing(doc)
    fix.fix_last_pages_and_signatories(doc, profile)
    ch2_idx = fix.find_paragraph_index(doc, "2 ФУНКЦИИ")
    assert fix.count_empty_lines_before(doc, ch2_idx) == 1


def test_apply_page_setup_deloproizvodstvo_a4_margins():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = fix.Cm(2)
    changed = fix.apply_page_setup_deloproizvodstvo(doc)
    assert changed >= 1
    assert abs(fix._length_mm(section.left_margin) - fix.MARGIN_LEFT_MM) < 0.1
    assert abs(fix._length_mm(section.right_margin) - 10) < 0.2
    assert fix.MARGIN_RIGHT_MM == 10
    assert abs(fix._length_mm(section.page_width) - fix.PAGE_WIDTH_MM) < 0.1
    assert not fix.validate_page_margins(doc)
    section.right_margin = fix.Mm(8)
    right_issues = fix.validate_page_margins(doc)
    assert any("правое поле" in item for item in right_issues)


def test_apply_body_single_line_spacing_not_signatories():
    doc = Document()
    body = doc.add_paragraph("1.4.1. Текст пункта тела документа.")
    fix.set_one_point_five_line_spacing(body)
    doc.add_paragraph("")
    razrab = doc.add_paragraph("Разработал:")
    fix.set_one_point_five_line_spacing(razrab)
    sign = doc.add_paragraph("Инженер")
    fix.set_one_point_five_line_spacing(sign)
    profile = fix.DocumentProfile("di", None, True, False, None)
    changed = fix.apply_body_single_line_spacing(doc, profile)
    assert changed >= 1
    assert fix.paragraph_has_single_line_spacing(body)
    assert fix.paragraph_has_zero_block_spacing(body)
    assert fix.paragraph_has_one_point_five_spacing(razrab)
    assert fix.paragraph_has_one_point_five_spacing(sign)
    assert not fix.validate_body_line_spacing(doc, profile)


def test_sanitize_paragraph_text_homoglyphs_and_punctuation():
    raw = "Согласованo:  текст  с  пробелами ."
    cleaned = fix.sanitize_paragraph_text(raw, full=True)
    assert cleaned.startswith("Согласовано:")
    assert "  " not in cleaned
    assert " ." not in cleaned
    assert cleaned.endswith(".")


def test_sanitize_paragraph_text_numbering_and_quotes():
    raw = ".1.4.1.. Текст с \"ошибкой\" и СНiОТ,, продолжение.."
    cleaned = fix.sanitize_paragraph_text(raw, full=True)
    assert cleaned.startswith("1.4.1. Текст")
    assert "СНиОТ" in cleaned
    assert ",," not in cleaned
    assert "«ошибкой»" in cleaned
    assert "продолжение." in cleaned


def test_sanitize_paragraph_text_preserves_ellipsis():
    raw = "1.4.1. Текст с многоточием..."
    cleaned = fix.sanitize_paragraph_text(raw, full=True)
    assert cleaned.endswith("...")


def test_strip_unnecessary_characters_skips_chapter_header():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1.  слово .")
    changed = fix.strip_unnecessary_characters(doc)
    assert changed == 1
    assert doc.paragraphs[0].text == "1 ОБЩИЕ ПОЛОЖЕНИЯ"
    assert "слово." in doc.paragraphs[1].text


def test_strip_unnecessary_characters_signatory_marker():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Пункт.")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    doc.add_paragraph("Согласованo:")
    doc.add_paragraph("Начальник")
    changed = fix.strip_unnecessary_characters(doc)
    assert changed >= 1
    sog_idx = fix.find_paragraph_index(doc, "Согласовано:")
    assert doc.paragraphs[sog_idx].text == "Согласовано:"
    assert not fix.validate_unnecessary_characters(doc)


def test_validate_and_strip_double_spaces_on_title():
    doc = Document()
    doc.add_paragraph("ГОСУДАРСТВЕННОЕ  ПРЕДПРИЯТИЕ")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Текст пункта.")
    issues = fix.validate_unnecessary_characters(doc)
    assert any("Двойные пробелы" in item for item in issues)
    changed = fix.strip_unnecessary_characters(doc)
    assert changed >= 1
    assert "  " not in doc.paragraphs[0].text
    assert not any("Двойные пробелы" in item for item in fix.validate_unnecessary_characters(doc))


def test_strip_double_spaces_in_table_cell():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "МИНСККОММУНТЕПЛОСЕТЬ  2026"
    changed = fix.strip_unnecessary_characters(doc)
    assert changed >= 1
    assert "  " not in table.cell(0, 0).text
    assert not any("Двойные пробелы" in item for item in fix.validate_unnecessary_characters(doc))


def test_light_sanitize_keeps_signatory_tab():
    raw = "Ведущий инженер\tВ.В.Дубовик"
    cleaned = fix.sanitize_paragraph_text(raw, full=False)
    assert "\t" in cleaned
    assert "  " not in cleaned.replace("\t", " ")


def test_apply_body_single_line_spacing_zeros_space_after():
    doc = Document()
    body = doc.add_paragraph("1.4.1. Текст пункта тела документа.")
    fix.set_single_line_spacing(body)
    body.paragraph_format.space_after = Pt(12)
    body.paragraph_format.space_before = Pt(6)
    chapter = doc.add_paragraph("2 ФУНКЦИИ")
    chapter.paragraph_format.space_after = Pt(10)
    profile = fix.DocumentProfile("di", None, True, False, None)
    changed = fix.apply_body_single_line_spacing(doc, profile)
    assert changed >= 1
    assert fix.paragraph_has_zero_block_spacing(body)
    assert fix.paragraph_has_zero_block_spacing(chapter)
    assert fix.paragraph_has_single_line_spacing(body)
    issues = fix.validate_body_line_spacing(doc, profile)
    assert not issues
    source = inspect.getsource(fix.validate_sniot_document)
    assert "validate_signatory_block" in source
    assert "validate_page_layout_flags" in source
    assert "validate_table_header_rows_together" in source
    assert "validate_list_markers" in source


def test_reload_document_from_path_reads_from_disk(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Тестовый пункт.")
    path = tmp_path / "sample.docx"
    doc.save(path)
    reloaded, raw = fix.reload_document_from_path(path)
    assert len(raw) > 100
    assert len(reloaded.paragraphs) == len(doc.paragraphs)
    assert reloaded.paragraphs[1].text.startswith("1.4.1.")


def test_validate_full_document_on_disk_matches_in_memory(tmp_path: Path):
    doc = _build_correct_numbering_doc()
    path = tmp_path / "numbered.docx"
    doc.save(path)
    on_disk = fix.validate_full_document_on_disk(path, path_for_profile=path)
    in_mem, raw = fix.reload_document_from_path(path)
    profile = fix.detect_profile(in_mem, path)
    in_memory = fix.validate_sniot_document(
        in_mem, docx_bytes=raw, profile=profile, path=path
    )
    assert on_disk == in_memory


def test_broken_chapter_one_fails_final_validate_on_disk(tmp_path: Path):
    """Финальная перечитка с диска ловит сбой materialize (1.5.1 в блоке 1.4.x)."""
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Настоящая должностная инструкция определяет функции.")
    doc.add_paragraph(
        "1.4. Назначение, перемещение и увольнение старшего мастера производится приказом."
    )
    doc.add_paragraph("1.4. Старший мастер в своей деятельности руководствуется:")
    doc.add_paragraph("1.5.1. Пункт списка раздела 1.4.")
    doc.add_paragraph("1.5. Старший мастер должен знать:")
    doc.add_paragraph("1.5.8. ТКП пункт.")
    path = tmp_path / "broken.docx"
    doc.save(path)
    profile_path = Path("ПРОЕКТ Старший мастер_оформлен.docx")
    issues = fix.validate_full_document_on_disk(path, path_for_profile=profile_path)
    assert issues
    assert any("1.4.x" in i or "вводн" in i.lower() or "1.5.1" in i for i in issues)


def test_correct_doc_passes_final_numbering_gate(tmp_path: Path):
    doc = _build_correct_numbering_doc()
    path = tmp_path / "ok.docx"
    doc.save(path)
    doc2, _ = fix.reload_document_from_path(path)
    profile = fix.detect_profile(doc2, Path("generic.docx"))
    assert fix.validate_numbering_block_starts(doc2, profile) == []
    assert fix.validate_chapter_one_wrong_intro_prefixes(doc2, profile) == []


def test_strip_unnecessary_characters_invisible_chars():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. Текст\u200b с\u00ad nbsp\u00a0и таб\t.")
    changed = fix.strip_unnecessary_characters(doc)
    assert changed == 1
    text = doc.paragraphs[1].text
    assert "\u200b" not in text
    assert "\u00ad" not in text
    assert "\t" not in text
    assert "\xa0" not in text


def test_interpret_word_signatory_layout_together():
    report = fix.interpret_word_signatory_layout(
        body_page=8, signatory_page=8, total_pages=8
    )
    assert report["orphaned"] is False
    assert "вместе" in report["message"]
    assert report["message"].startswith("Word:")


def test_interpret_word_signatory_layout_orphaned():
    report = fix.interpret_word_signatory_layout(
        body_page=7, signatory_page=8, total_pages=8
    )
    assert report["orphaned"] is True
    assert "оторван" in report["message"]


def test_interpret_word_signatory_layout_missing_markers():
    no_sign = fix.interpret_word_signatory_layout(body_page=7, signatory_page=None)
    assert no_sign["orphaned"] is False
    assert "Разработал" in no_sign["message"]
    no_body = fix.interpret_word_signatory_layout(body_page=None, signatory_page=8)
    assert no_body["orphaned"] is False
    assert "хвост" in no_body["message"]


def test_format_word_layout_action_prefix_for_gui():
    msg = fix.format_word_layout_action(
        {"message": "Word: подписанты вместе с текстом (текст стр. 8, «Разработал:» стр. 8)"}
    )
    assert msg.startswith("СНиОТ: Word:")
    wrapped = fix.format_word_layout_action({"message": "проверка не выполнена"})
    assert wrapped == "СНиОТ: Word: проверка не выполнена"


def test_format_word_layout_action_includes_visible_counts():
    msg = fix.format_word_layout_action(
        {
            "message": "Word: подписанты вместе с текстом",
            "razrabotal_count": 2,
            "double_numbering_count": 1,
        }
    )
    assert "Разработал" in msg
    assert "двойная" in msg


def test_dedupe_consecutive_razrabotal_keeps_one():
    doc = Document()
    doc.add_paragraph("5.1.1. Пункт.")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    doc.add_paragraph("Согласовано:")
    profile = fix.DocumentProfile("di", None, True, False, None)
    assert len(fix.iter_razrabotal_paragraph_elements(doc)) == 2
    fix.deduplicate_razrabotal_markers(doc)
    left = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip().upper().startswith("РАЗРАБОТАЛ")
    ]
    assert left == ["Разработал:"]
    assert not fix.validate_duplicate_razrabotal(doc, profile)


def test_dedupe_razrabotal_in_table_and_body():
    doc = Document()
    doc.add_paragraph("5.1.1. Пункт.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Разработал:"
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер")
    profile = fix.DocumentProfile("di", None, True, False, None)
    assert len(fix.iter_razrabotal_paragraph_elements(doc)) == 2
    assert fix.validate_duplicate_razrabotal(doc, profile)
    removed = fix.deduplicate_razrabotal_markers(doc)
    assert removed == 1
    assert len(fix.iter_razrabotal_paragraph_elements(doc)) == 1
    assert not fix.validate_duplicate_razrabotal(doc, profile)


def test_apply_word_list_fixes_writes_number_and_strips_numpr():
    doc = Document()
    p = doc.add_paragraph("Настоящая должностная инструкция определяет функции.")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    p_pr.append(num_pr)
    changed = fix.apply_word_list_fixes(
        doc,
        [
            {
                "text": "Настоящая должностная инструкция определяет функции.",
                "list_string": "1.1.",
                "already_in_text": False,
            }
        ],
    )
    assert changed >= 1
    assert doc.paragraphs[0].text.startswith("1.1.")
    assert not fix.has_word_list_numbering(doc.paragraphs[0])


def test_ensure_razrabotal_does_not_add_second_when_table_has_one():
    doc = Document()
    doc.add_paragraph("5.1.1. Пункт.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Разработал:"
    added = fix.ensure_razrabotal_marker(doc)
    assert added is False
    assert len(fix.iter_razrabotal_paragraph_elements(doc)) == 1


def _make_table_doc(n_rows: int, n_cols: int = 2) -> Document:
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    table = doc.add_table(rows=n_rows, cols=n_cols)
    for i in range(n_rows):
        for j in range(n_cols):
            table.cell(i, j).text = f"R{i + 1}C{j + 1}"
    return doc


def test_keep_table_header_three_rows_together():
    doc = _make_table_doc(5)
    assert fix.validate_table_header_rows_together(doc)
    changed = fix.keep_table_header_rows_together(doc)
    assert changed >= 1
    rows = fix.table_row_elements(doc.tables[0])
    assert len(rows) == 5
    for i in range(3):
        assert fix.table_row_has_cant_split(rows[i])
        assert fix.table_row_has_tbl_header(rows[i])
    assert fix.table_row_has_keep_next(rows[0])
    assert fix.table_row_has_keep_next(rows[1])
    assert not fix.table_row_has_keep_next(rows[2])
    assert not fix.table_row_has_keep_next(rows[3])
    assert fix.validate_table_header_rows_together(doc) == []


def test_keep_table_two_rows_together():
    doc = _make_table_doc(2)
    fix.keep_table_header_rows_together(doc)
    rows = fix.table_row_elements(doc.tables[0])
    assert len(rows) == 2
    assert fix.table_row_has_cant_split(rows[0])
    assert fix.table_row_has_cant_split(rows[1])
    assert fix.table_row_has_tbl_header(rows[0])
    assert fix.table_row_has_tbl_header(rows[1])
    assert fix.table_row_has_keep_next(rows[0])
    assert not fix.table_row_has_keep_next(rows[1])
    assert fix.validate_table_header_rows_together(doc) == []


def test_process_sniot_keeps_table_header_rows():
    doc = _make_table_doc(5)
    profile = fix.detect_profile(doc, Path("generic.docx"))
    fix.process_sniot_document(doc, profile)
    assert fix.validate_table_header_rows_together(doc) == []
    rows = fix.table_row_elements(doc.tables[0])
    assert fix.table_row_has_keep_next(rows[0])
    assert fix.table_row_has_keep_next(rows[1])
    assert not fix.table_row_has_keep_next(rows[3])


def test_page_break_not_between_chapter_and_table():
    doc = Document()
    hdr = doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    table = doc.add_table(rows=3, cols=1)
    table.cell(0, 0).text = "Шапка"
    table.cell(1, 0).text = "Строка 2"
    table.cell(2, 0).text = "Строка 3"
    fix.set_page_break_before(hdr, True)
    fix.prevent_chapter_header_orphan(doc)
    assert not fix.paragraph_has_page_break_before(hdr)
    assert fix.chapter_header_followed_by_table(hdr)


def test_table_cells_have_no_first_line_indent_after_apply():
    doc = Document()
    doc.add_paragraph("1.8.1. Пункт тела с отступом.")
    table = doc.add_table(rows=1, cols=1)
    cell_p = table.cell(0, 0).paragraphs[0]
    cell_p.text = "Текст в ячейке"
    fix.ensure_first_line_indent(cell_p)
    assert fix.first_line_indent_cm(cell_p) > 1.0
    issues_before = fix.validate_table_paragraph_indents(doc)
    assert issues_before
    profile = fix.DocumentProfile("di", None, False, False, None)
    fix.apply_body_paragraph_format(doc, profile)
    cell_p = table.cell(0, 0).paragraphs[0]
    assert abs(fix.first_line_indent_cm(cell_p)) <= 0.08
    assert fix.validate_table_paragraph_indents(doc) == []
    body_issues = [
        i for i in fix.validate_body_paragraph_format(doc, profile) if "таблице" in i
    ]
    assert body_issues == []


def test_process_sniot_clears_table_indent():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.8.1. Первый пункт тела.")
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Шапка таблицы"
    table.cell(1, 0).text = "Строка данных"
    for cell in (table.cell(0, 0), table.cell(1, 0)):
        fix.ensure_first_line_indent(cell.paragraphs[0])
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    fix.process_sniot_document(doc, profile)
    for paragraph in fix.iter_table_cell_paragraphs(doc):
        if paragraph.text.strip():
            assert abs(fix.first_line_indent_cm(paragraph)) <= 0.08, paragraph.text[:40]
    assert fix.validate_table_paragraph_indents(doc) == []


def test_lsim_di_is_not_satp_numbering():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Настоящая инструкция определяет обязанности инженера ЛСиМ.")
    doc.add_paragraph("В своей деятельности инженер руководствуется ТНПА.")
    doc.add_paragraph("1.8. Инженер ЛСиМ должен знать:")
    doc.add_paragraph("ТКП 050-2007 Котлы паровые.")
    doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    profile = fix.detect_profile(doc, Path("ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx"))
    assert profile.has_di_satp_numbering is False


def test_senior_master_still_satp_numbering():
    doc = _build_correct_numbering_doc()
    profile = fix.detect_profile(doc, Path("ПРОЕКТ Старший мастер_оформлен.docx"))
    assert profile.has_di_satp_numbering is True


def test_should_apply_skips_title_paragraphs():
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("КОММУНАЛЬНЫХ ТЕПЛОВЫХ СЕТЕЙ И КОТЕЛЬНЫХ")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст тела документа для проверки отступа.")
    assert fix.should_apply_body_paragraph_format(doc.paragraphs[0].text, 0, doc) is False
    assert fix.should_apply_body_paragraph_format(doc.paragraphs[1].text, 1, doc) is False
    assert fix.should_apply_body_paragraph_format(doc.paragraphs[3].text, 3, doc) is True


def test_restore_org_header_replaces_tkp_on_title():
    doc = Document()
    doc.add_paragraph("КОММУНАЛЬНОЕ УНИТАРНОЕ ПРОИЗВОДСТВЕННОЕ")
    doc.add_paragraph(
        "ТКП 458-2023 (33240) «Правила технической эксплуатации теплоустановок и тепловых сетей потребителей»."
    )
    doc.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    assert fix.title_has_foreign_npa_insert(doc) is True
    assert fix.restore_org_header_if_npa_inserted(doc) == 1
    assert doc.paragraphs[1].text == fix.ORG_HEADER_TEPLOSETI_LINE
    assert fix.title_has_foreign_npa_insert(doc) is False


def test_strip_visual_highlights_removes_marker_and_shading():
    doc = Document()
    paragraph = doc.add_paragraph("Выделенный абзац тела.")
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "FFFFFF")
    p_pr.append(shd)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("x")
    r_pr = run._r.get_or_add_rPr()
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), "yellow")
    r_pr.append(hl)
    assert fix.strip_visual_highlights(doc) >= 2
    xml = paragraph._p.xml
    assert "w:highlight" not in xml
    assert "w:shd" not in xml


def test_restore_lsim_if_filename_matches():
    doc = Document()
    doc.add_paragraph("1.4. Инженер Осим подчиняется начальнику ЛСиМ.")
    n = fix.restore_lsim_if_osim_poisoned(
        doc, Path("ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx")
    )
    assert n == 1
    assert "Осим" not in doc.paragraphs[0].text
    assert "ЛСиМ" in doc.paragraphs[0].text
    doc2 = Document()
    doc2.add_paragraph("Инженер Осим.")
    assert fix.restore_lsim_if_osim_poisoned(doc2, Path("другой.docx")) == 0


def test_copy_same_file_slash_vs_backslash(tmp_path: Path):
    f = tmp_path / "ДИ_оформлен.docx"
    f.write_bytes(b"PK")
    fwd = str(f).replace("\\", "/")
    assert fix.paths_are_same_file(fwd, f) is True
    assert fix.copy_file_if_different(fwd, f) is False
    other = tmp_path / "other.docx"
    assert fix.copy_file_if_different(f, other) is True


def test_oformlen_not_allowed_sample():
    oformlen = fix.USER_AGENT_DIR / "ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx"
    assert fix.is_allowed_sample_path(oformlen) is False


def test_normalize_fonts_unifies_numbered_runs():
    doc = Document()
    paragraph = doc.add_paragraph()
    num_run = paragraph.add_run("1.4.1. ")
    num_run.bold = True
    num_run.font.name = "Calibri"
    num_run.font.size = Pt(11)
    text_run = paragraph.add_run("Текст пункта тела.")
    text_run.font.name = "Times New Roman"
    text_run.font.size = Pt(14)
    text_run.bold = False
    issues_before = fix.validate_fonts(doc)
    assert issues_before
    fix.normalize_document_fonts(doc)
    assert not fix.validate_fonts(doc)
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        r_pr = run._r.find(qn("w:rPr"))
        assert r_pr is not None
        r_fonts = r_pr.find(qn("w:rFonts"))
        assert r_fonts.get(qn("w:ascii")) == "Times New Roman"
        assert r_fonts.get(qn("w:eastAsia")) == "Times New Roman"
        assert r_pr.find(qn("w:sz")).get(qn("w:val")) == "28"
        bold_el = r_pr.find(qn("w:b"))
        assert bold_el is not None
        assert (bold_el.get(qn("w:val")) or "").lower() in ("false", "0")


def test_is_chapter_header_lsim_wording():
    assert fix.is_chapter_header("4. ВЗАИМОДЕЙСТВИЯ")
    assert fix.is_chapter_header("2 ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    assert fix.is_chapter_header("2 Должностные обязанности")
    assert fix.is_chapter_header("3. ПРАВА")
    assert not fix.is_chapter_header("1.4.1. Текст пункта.")
    assert not fix.is_chapter_header("1.8. Инженер ЛСиМ должен знать:")
    assert fix.canonical_chapter_header("4. ВЗАИМОДЕЙСТВИЯ") == "4 ВЗАИМОДЕЙСТВИЯ"


def test_center_chapter_headers_lsim_interaction():
    doc = Document()
    hdr = doc.add_paragraph("4. ВЗАИМОДЕЙСТВИЯ")
    fix.ensure_first_line_indent(hdr)
    assert fix.center_chapter_headers(doc) == 1
    assert hdr.text.strip() == "4 ВЗАИМОДЕЙСТВИЯ"
    assert fix.is_paragraph_centered(hdr)
    assert fix.first_line_indent_cm(hdr) < 0.1


def test_apply_body_indent_no_spacing_style():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    body = doc.add_paragraph("1.1. Текст абзаца тела документа.")
    try:
        body.style = doc.styles["No Spacing"]
    except KeyError:
        pass
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, False, None)
    fix.center_chapter_headers(doc)
    fix.apply_body_paragraph_format(doc, profile)
    assert fix.is_paragraph_justified(body)
    assert abs(fix.first_line_indent_cm(body) - 1.25) < 0.1


def test_ensure_title_page_separated_next_page():
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    for _ in range(8):
        doc.add_paragraph("")
    minsk = doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст главы.")
    changed = fix.ensure_title_page_separated(doc)
    assert changed >= 1
    empty_run = 0
    max_empty = 0
    body_start = fix.find_body_start_index(doc)
    for i in range(body_start):
        if not doc.paragraphs[i].text.strip():
            empty_run += 1
            max_empty = max(max_empty, empty_run)
        else:
            empty_run = 0
    assert max_empty <= 2
    assert not fix.validate_title_page_separated(doc)
    minsk_idx = fix.find_paragraph_index(doc, "Минск 2026")
    assert fix._paragraph_sectpr(doc.paragraphs[minsk_idx]) is None
    sect = fix._paragraph_sectpr(doc.paragraphs[minsk_idx + 1])
    assert sect is not None
    typ = sect.find(qn("w:type"))
    assert typ is not None and typ.get(qn("w:val")) == "nextPage"


def test_fix_page_numbering_uniform_tnr14_not_bold(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Шапка предприятия")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    path = tmp_path / "hdr.docx"
    doc.save(path)
    fixed = fix.fix_page_numbering(path.read_bytes())
    issues = fix.validate_page_numbering(fixed)
    assert not any("жирн" in i.lower() for i in issues)
    assert not any("не Times" in i for i in issues)
    assert not any("не 14" in i for i in issues)
    with zipfile.ZipFile(BytesIO(fixed)) as zin:
        for name in ("word/header1.xml", "word/header2.xml"):
            xml = zin.read(name).decode("utf-8")
            assert "PAGE" in xml
            assert "Times New Roman" in xml
            assert 'w:val="28"' in xml
            assert 'w:val="false"' in xml
        h3 = zin.read("word/header3.xml").decode("utf-8")
        assert "PAGE" not in h3


def test_materialize_signatory_table_to_paragraphs():
    doc = Document()
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("5.1.1. Пункт ответственности.")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Начальник ЛСиМ"
    table.cell(0, 1).text = "Рафеенко Д.Ф. |"
    table.cell(1, 0).text = "Начальник СНиОТ"
    table.cell(1, 1).text = "Хандуратов М.Н. |"
    table.cell(2, 0).text = "Начальник ЮО"
    table.cell(2, 1).text = "Авраменко Г.А."
    assert fix.has_signatory_block(doc)
    n = fix.materialize_signatory_paragraphs_from_tables(doc)
    assert n >= 3
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "Разработал:" in texts
    assert any(t == "Согласовано:" for t in texts)
    raz = texts.index("Разработал:")
    sog = texts.index("Согласовано:")
    assert raz < sog
    assert "Рафеенко" in texts[raz + 1]
    assert "\t" in [p.text for p in doc.paragraphs if "Рафеенко" in p.text][0]
    assert "Хандуратов" in "".join(texts[sog:])
    assert not any(fix.is_agreement_signatory_table(t) for t in doc.tables)
    profile = fix.detect_profile(doc, Path("ДИ инженер ЛСиМ_оформлен.docx"))
    assert profile.has_signatories
    assert profile.has_di_satp_numbering is False
    fix.fix_signatory_block_format(doc, profile)
    assert fix.find_razrabotal_index(doc) >= 0
    sog_idx = fix.find_soglasovano_index(doc)
    assert doc.paragraphs[sog_idx].text.strip() == "Согласовано:"
    assert not any(r.bold for r in doc.paragraphs[sog_idx].runs if r.text)
    assert not fix.validate_signatory_line_spacing(doc, profile)
    assert not fix.validate_agreement_table_converted(doc, profile)


def _attach_numpr(paragraph, num_id: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(nid)
    p_pr.append(num_pr)


def test_strip_leading_list_marker_chars():
    assert fix.strip_leading_list_marker_text("• Требование охраны труда.") == (
        "Требование охраны труда."
    )
    assert fix.strip_leading_list_marker_text("\uf0b7 Текст пункта.") == "Текст пункта."
    assert fix.strip_leading_list_marker_text("○ Второй пункт.") == "Второй пункт."
    assert fix.strip_leading_list_marker_text("■ Третий пункт.") == "Третий пункт."
    assert fix.strip_leading_list_marker_text("– Четвёртый пункт.") == "Четвёртый пункт."
    assert fix.strip_leading_list_marker_text("- Пятый пункт.") == "Пятый пункт."
    assert fix.strip_leading_list_marker_text("1.4.1. Текст пункта.") == "1.4.1. Текст пункта."
    assert fix.text_has_leading_list_marker("• Текст")
    assert not fix.text_has_leading_list_marker("1.4.1. Текст пункта.")
    assert not fix.text_has_leading_list_marker("Ведущий инженер\tВ.В.Дубовик")


def test_sanitize_full_strips_leading_bullet():
    cleaned = fix.sanitize_paragraph_text("•  Требование  охраны.", full=True)
    assert cleaned == "Требование охраны."
    assert "•" not in cleaned


def test_remove_bullet_numpr_without_manual_number():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p = doc.add_paragraph("Соблюдать требования охраны труда.")
    _attach_numpr(p, "1")  # default template: numId=1 — bullet
    assert fix.has_word_list_numbering(p)
    assert fix.is_bullet_word_list(p, fix.numbering_level_map(doc))
    issues = fix.validate_list_markers(doc)
    assert any("bullet" in item.lower() or "Маркер списка" in item for item in issues)
    changed = fix.remove_word_list_numbering_in_body(doc)
    assert changed >= 1
    assert not fix.has_word_list_numbering(p)
    assert not fix.validate_list_markers(doc)


def test_remove_list_markers_in_body_strips_text_and_bullet():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p_text = doc.add_paragraph("• Требование охраны труда.")
    p_bullet = doc.add_paragraph("Обычный абзац без символа.")
    _attach_numpr(p_bullet, "1")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "○ Пункт в таблице."
    issues_before = fix.validate_list_markers(doc)
    assert issues_before
    changed = fix.remove_list_markers_in_body(doc)
    assert changed >= 2
    assert p_text.text == "Требование охраны труда."
    assert not fix.has_word_list_numbering(p_bullet)
    assert table.cell(0, 0).paragraphs[0].text == "Пункт в таблице."
    assert not fix.validate_list_markers(doc)


def test_list_markers_do_not_touch_numbered_item_or_signatory_tab():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    numbered = doc.add_paragraph("1.4.1. Создание условий труда.")
    _attach_numpr(numbered, "9")  # decimal
    doc.add_paragraph("Разработал:")
    sign = doc.add_paragraph("Ведущий инженер\tВ.В.Дубовик")
    fix.remove_list_markers_in_body(doc)
    assert numbered.text.startswith("1.4.1.")
    assert "\t" in sign.text
    assert not fix.text_has_leading_list_marker(numbered.text)


def test_process_sniot_document_clears_list_markers():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("• Требование охраны труда.")
    p = doc.add_paragraph("Ещё один пункт без номера в тексте.")
    _attach_numpr(p, "1")
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, False, None)
    fix.process_sniot_document(doc, profile)
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "Требование охраны труда." in joined
    assert "•" not in joined
    assert not any(fix.has_word_list_numbering(p) and fix.is_bullet_word_list(p, fix.numbering_level_map(doc)) for p in doc.paragraphs)
    assert not fix.validate_list_markers(doc)


def test_format_word_list_label_matches_liststring():
    assert fix.format_word_list_label("1.%1.", [4]) == "1.4."
    assert fix.format_word_list_label("2.2.%1.", [1]) == "2.2.1."
    assert fix.format_word_list_label("%1.%2.", [2, 3]) == "2.3."
    assert fix.normalize_number_token("1.6.\t") == "1.6."
    assert fix.normalize_number_token("2.2.10.") == "2.2.10."


def _add_decimal_numbering(doc: Document, abstract_id: int, num_id: int, lvl_text: str, start: int) -> None:
    from lxml import etree

    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = doc.part.numbering_part.element
    absn = etree.fromstring(
        f'<w:abstractNum xmlns:w="{ns}" w:abstractNumId="{abstract_id}">'
        f'<w:lvl w:ilvl="0"><w:start w:val="{start}"/>'
        f'<w:numFmt w:val="decimal"/><w:lvlText w:val="{lvl_text}"/>'
        f'<w:lvlJc w:val="left"/></w:lvl></w:abstractNum>'
    )
    first_num = root.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(absn)
    else:
        root.append(absn)
    root.append(
        etree.fromstring(
            f'<w:num xmlns:w="{ns}" w:numId="{num_id}">'
            f'<w:abstractNumId w:val="{abstract_id}"/></w:num>'
        )
    )


def test_source_numbering_preserved_not_renumbered():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Настоящая инструкция определяет обязанности инженера ЛСиМ.")
    doc.add_paragraph("1.2. Инженер относится к категории специалистов.")
    doc.add_paragraph("1.6. В своей деятельности инженер руководствуется ТНПА.")
    doc.add_paragraph("1.8. Инженер ЛСиМ должен знать:")
    doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    snapshot = fix.collect_number_tokens(doc)
    assert "1.6." in snapshot
    assert "1.8." in snapshot
    profile = fix.detect_profile(doc, Path("ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx"))
    profile.source_number_tokens = tuple(snapshot)
    assert profile.has_di_satp_numbering is False
    fix.process_sniot_document(doc, profile)
    texts = [p.text.strip() for p in doc.paragraphs]
    assert any(t.startswith("1.6.") for t in texts)
    assert any(t.startswith("1.8.") for t in texts)
    assert not any(t.startswith("1.3.") for t in texts)
    assert not any(t.startswith("1.7.") and "должен знать" in t for t in texts)
    result = fix.collect_number_tokens(doc)
    for token in snapshot:
        assert token in result
    issues = fix.validate_source_numbering_preserved(doc, profile)
    assert issues == []


def test_materialize_word_list_keeps_source_numbers():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p1 = doc.add_paragraph("Инженер ЛСиМ назначается на должность приказом директора.")
    p2 = doc.add_paragraph("Инженер ЛСиМ подчиняется непосредственно начальнику ЛСиМ.")
    _add_decimal_numbering(doc, abstract_id=80, num_id=80, lvl_text="1.%1.", start=4)
    _attach_numpr(p1, "80")
    _attach_numpr(p2, "80")
    tokens = fix.collect_number_tokens(doc)
    assert "1.4." in tokens
    assert "1.5." in tokens
    n = fix.materialize_word_decimal_numbering(doc)
    assert n >= 2
    assert p1.text.startswith("1.4.")
    assert p2.text.startswith("1.5.")
    assert not fix.has_word_list_numbering(p1)
    assert not fix.has_word_list_numbering(p2)


def test_process_strips_highlight_and_does_not_mark_numbers():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    paragraph = doc.add_paragraph("1.1. Настоящая инструкция определяет обязанности.")
    run = paragraph.runs[0]
    r_pr = run._r.get_or_add_rPr()
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), "yellow")
    r_pr.append(hl)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "FFFF00")
    p_pr.append(shd)
    gray = OxmlElement("w:shd")
    gray.set(qn("w:val"), "clear")
    gray.set(qn("w:fill"), "D9D9D9")
    r_pr.append(gray)
    profile = fix.detect_profile(doc, Path("ДИ инженер ЛСиМ_оформлен.docx"))
    profile.source_number_tokens = ("1.1.",)
    fix.process_sniot_document(doc, profile)
    xml = paragraph._p.xml
    assert "w:highlight" not in xml
    assert fix.validate_visual_highlights(doc) == []
    assert any(p.text.strip().startswith("1.1.") for p in doc.paragraphs)


PROCESS_REQUIRED_CALLS = (
    "accept_tracked_changes",
    "keep_table_header_rows_together",
    "apply_table_paragraph_no_indent",
    "remove_list_markers_in_body",
    "materialize_word_decimal_numbering",
    "materialize_signatory_paragraphs_from_tables",
    "ensure_title_page_separated",
    "normalize_document_fonts",
    "strip_visual_highlights",
    "apply_body_paragraph_format",
    "apply_body_single_line_spacing",
    "center_chapter_headers",
    "apply_page_setup_deloproizvodstvo",
    "apply_signatory_tab_stops",
    "strip_unnecessary_characters",
    "place_title_city_year_at_bottom",
    "ensure_acquaintance_sheet_separate_page",
    "normalize_item_number_spacing",
    "collapse_adjacent_duplicate_words_in_document",
    "apply_signatory_fio_one_line",
    "fix_signatory_date_plaques",
    "mark_acquaintance_heading_if_job_mismatch",
    "apply_word_grammar_via_com",
    "run_word_grammar_check_subprocess",
)


def test_process_sniot_document_calls_required_fixers():
    src = inspect.getsource(fix.process_sniot_document)
    missing = [name for name in PROCESS_REQUIRED_CALLS if name not in src]
    assert missing == [], f"process() не вызывает: {missing}"
    apply_src = inspect.getsource(fix.apply_sniot_rules_to_file)
    assert "fix_page_numbering" in apply_src
    assert "run_word_grammar_check_subprocess" in apply_src
    assert "apply_word_grammar_via_com" in inspect.getsource(fix.run_word_grammar_check_subprocess)
    numbering_src = inspect.getsource(fix._fix_numbering_if_needed)
    assert "has_di_satp_numbering" in numbering_src
    assert "СНиОТ: сборка" in inspect.getsource(fix.autofix)
    assert "СНиОТ: сборка" in inspect.getsource(fix.apply_sniot_rules_to_file)


def test_body_starts_with_first_chapter_skips_org_header():
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    assert fix.body_starts_with_first_chapter(doc, "1 ОБЩИЕ ПОЛОЖЕНИЯ") is True
    profile = fix.DocumentProfile("di", "1 ОБЩИЕ ПОЛОЖЕНИЯ", False, False, None)
    issues = fix.validate_sniot_document(doc, profile=profile)
    assert not any("Дубль титула" in item for item in issues)


def test_validate_signatory_tab_stops_skips_markers():
    doc = Document()
    doc.add_paragraph("5.1.1. Пункт.")
    doc.add_paragraph("Разработал:")
    fio = doc.add_paragraph("Начальник ЛСиМ\tРафеенко Д.Ф.")
    fix.ensure_signatory_tab_stops(fio)
    doc.add_paragraph("Согласовано:")
    profile = fix.DocumentProfile("di", None, True, False, None)
    assert fix.validate_signatory_tab_stops(doc, profile) == []


def test_agreement_table_with_fio_header_is_detected():
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Должность"
    table.cell(0, 1).text = "Ф.И.О."
    table.cell(1, 0).text = "Начальник ЛСиМ"
    table.cell(1, 1).text = "Рафеенко Д.Ф."
    table.cell(2, 0).text = "Начальник СНиОТ"
    table.cell(2, 1).text = "Хандуратов М.Н."
    table.cell(3, 0).text = "Начальник ЮО"
    table.cell(3, 1).text = "Авраменко Г.А."
    assert fix.is_agreement_signatory_table(table) is True
    rows = fix.extract_agreement_table_rows(table)
    assert len(rows) == 3
    assert rows[0][1].startswith("Рафеенко")


def test_agreement_table_three_columns_converted():
    doc = Document()
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("5.1.1. Пункт ответственности.")
    table = doc.add_table(rows=4, cols=3)
    table.cell(0, 0).text = "Должность"
    table.cell(0, 1).text = "Ф.И.О."
    table.cell(0, 2).text = "Подпись"
    table.cell(1, 0).text = "Начальник ЛСиМ"
    table.cell(1, 1).text = "Рафеенко Д.Ф."
    table.cell(1, 2).text = ""
    table.cell(2, 0).text = "Начальник СНиОТ"
    table.cell(2, 1).text = "Хандуратов М.Н."
    table.cell(2, 2).text = ""
    table.cell(3, 0).text = "Начальник ЮО"
    table.cell(3, 1).text = "Авраменко Г.А."
    table.cell(3, 2).text = ""
    n = fix.materialize_signatory_paragraphs_from_tables(doc)
    assert n >= 3
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "Разработал:" in joined
    assert "Согласовано:" in joined
    assert "Рафеенко" in joined
    assert not any(fix.is_agreement_signatory_table(t) for t in doc.tables)


def test_acquaintance_sheet_is_not_agreement_table():
    doc = Document()
    table = doc.add_table(rows=2, cols=5)
    table.cell(0, 0).text = "№ п/п"
    table.cell(0, 1).text = "Ф.И.О."
    table.cell(0, 2).text = "Должность"
    table.cell(0, 3).text = "Роспись"
    table.cell(0, 4).text = "Дата"
    table.cell(1, 0).text = "1."
    table.cell(1, 1).text = "Романовский Владимир Владимирович"
    table.cell(1, 2).text = "инженер"
    assert fix.is_agreement_signatory_table(table) is False
    assert fix.materialize_signatory_paragraphs_from_tables(doc) == 0


def _lsim_like_document() -> Document:
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("КОММУНАЛЬНОЕ УНИТАРНОЕ ПРОИЗВОДСТВЕННОЕ")
    doc.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1. ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Настоящая инструкция определяет должностные обязанности.")
    doc.add_paragraph("1.6. В своей деятельности инженер ЛСиМ руководствуется.")
    doc.add_paragraph("1.8. Инженер ЛСиМ должен знать:")
    doc.add_paragraph("ТКП 458-2023 Правила.")
    doc.add_paragraph("2. ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    doc.add_paragraph("2.1. На инженера ЛСиМ возлагаются следующие функции:")
    doc.add_paragraph("2.1.1. проведение технического диагностирования.")
    doc.add_paragraph("3. ПРАВА")
    doc.add_paragraph("Инженер ЛСиМ имеет право:")
    doc.add_paragraph("4. ВЗАИМОДЕЙСТВИЯ")
    doc.add_paragraph("Взаимодействия с другими подразделениями предприятия.")
    doc.add_paragraph("5. ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("5.1. Инженер ЛСиМ несет ответственность:")
    doc.add_paragraph("5.1.1. За неисполнение обязанностей.")
    doc.add_paragraph("2.2.17. Выполняет локальные правовые акты предприятия.")
    doc.add_paragraph(
        "1.6. В своей деятельности инженер ЛСиМ руководствуется нормативными правовыми актами."
    )
    bullet = doc.add_paragraph("• Лишний маркер списка.")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Начальник ЛСиМ"
    table.cell(0, 1).text = "Рафеенко Д.Ф."
    table.cell(1, 0).text = "Начальник СНиОТ"
    table.cell(1, 1).text = "Хандуратов М.Н."
    table.cell(2, 0).text = "Начальник ЮО"
    table.cell(2, 1).text = "Авраменко Г.А."
    return doc


def test_process_lsim_like_applies_layout_and_keeps_source_numbers():
    doc = _lsim_like_document()
    path = Path("ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx")
    profile = fix.detect_profile(doc, path)
    assert profile.has_di_satp_numbering is False
    tokens_before = fix.collect_number_tokens(doc)
    fix.process_sniot_document(doc, profile, source_path=path)
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "1.1." in joined
    assert "1.6." in joined
    assert "1.8. Инженер ЛСиМ должен знать" in joined or "1.8." in joined
    assert "1.9." not in joined
    assert "Разработал:" in joined
    assert "Согласовано:" in joined
    assert "Рафеенко" in joined
    assert not any(fix.is_agreement_signatory_table(t) for t in doc.tables)
    body = next(p for p in doc.paragraphs if p.text.strip().startswith("1.1."))
    assert abs(fix.first_line_indent_cm(body) - 1.25) < 0.15
    ch4 = next(p for p in doc.paragraphs if "ВЗАИМОДЕЙСТВ" in p.text.upper())
    assert fix.is_paragraph_centered(ch4)
    assert not fix.validate_title_page_separated(doc)
    assert not any("Дубль титула" in i for i in fix.validate_sniot_document(doc, profile=profile))
    assert not fix.validate_signatory_tab_stops(doc, profile)
    assert "•" not in joined
    assert "Выполняет локальные правовые акты" in joined
    assert "требования локальных правовых актов" not in joined
    assert "нормативными правовыми актами" in joined
    assert "нормативно-правовыми" not in joined
    tokens_after = fix.collect_number_tokens(doc)
    for token in ("1.1.", "1.6.", "1.8.", "2.1.1.", "5.1.1."):
        assert token in tokens_before
        assert token in tokens_after
    for run in body.runs:
        if run.text.strip():
            assert run.bold is not True


def test_sanitize_spaces_inside_parentheses():
    raw = "продукции ( выполнение )"
    cleaned = fix.sanitize_paragraph_text(raw, full=True)
    assert cleaned == "продукции (выполнение)"
    assert cleaned != "продукции(выполнение)"
    assert "( " not in cleaned
    assert " )" not in cleaned
    assert "продукции (" in cleaned
    assert "продукции(" not in cleaned
    assert fix.fix_spaces_around_parentheses(raw) == "продукции (выполнение)"
    numbered = "1.4.1. Текст ( выполнение)."
    numbered_clean = fix.sanitize_paragraph_text(numbered, full=True)
    assert numbered_clean.startswith("1.4.1.")
    assert "Текст (выполнение)" in numbered_clean
    assert "Текст(" not in numbered_clean
    closing = fix.sanitize_paragraph_text("текст (слово ) далее", full=True)
    assert " )" not in closing
    assert "текст (слово)" in closing
    assert "текст(" not in closing
    already_ok = "продукции (выполнение)"
    assert fix.sanitize_paragraph_text(already_ok, full=True) == already_ok
    assert fix.sanitize_paragraph_text("1.1. Пункт без скобок.", full=True).startswith("1.1.")


def test_paren_spaces_skip_signatory_tab():
    raw = "Инженер ( ЛСиМ )\tИ.И. Иванов"
    cleaned = fix.sanitize_paragraph_text(raw, full=False)
    assert "\t" in cleaned
    assert "Инженер (ЛСиМ)" in cleaned
    assert "Инженер(" not in cleaned


def test_strip_unnecessary_fixes_paren_spaces_in_body():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("2.2.14. обеспечивает при производстве продукции ( выполнение работ);")
    changed = fix.strip_unnecessary_characters(doc)
    assert changed >= 1
    assert "продукции (выполнение" in doc.paragraphs[1].text
    assert "продукции(" not in doc.paragraphs[1].text
    assert not any("Лишние символы" in item for item in fix.validate_unnecessary_characters(doc))


def test_place_title_city_year_at_bottom_centered():
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    minsk = doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст главы.")
    fix.ensure_title_page_separated(doc)
    changed = fix.place_title_city_year_at_bottom(doc)
    assert changed >= 1
    idx = fix.find_title_city_year_index(doc)
    assert idx is not None
    para = doc.paragraphs[idx]
    assert fix.is_paragraph_centered(para)
    assert fix.first_line_indent_cm(para) < 0.1
    p_pr = para._p.find(qn("w:pPr"))
    assert p_pr is not None and p_pr.find(qn("w:framePr")) is not None
    assert not fix.validate_title_city_year(doc)
    body_start = fix.find_body_start_index(doc)
    assert idx < body_start
    for i in range(idx + 1, body_start):
        assert not doc.paragraphs[i].text.strip()


def test_city_year_variants_recognized():
    assert fix.is_city_year_paragraph("Минск-2026")
    assert fix.is_city_year_paragraph("Минск, 2026")
    assert fix.is_city_year_paragraph("Минск 2026")
    assert not fix.is_city_year_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    assert not fix.is_city_year_paragraph("МИНСККОММУНТЕПЛОСЕТЬ")


def test_acquaintance_sheet_separate_page_not_signatories():
    doc = _signatory_tail_doc()
    doc.add_paragraph("Настоящую инструкцию изучил и обязуюсь ею руководствоваться")
    table = doc.add_table(rows=2, cols=5)
    table.cell(0, 0).text = "№ п/п"
    table.cell(0, 1).text = "Ф.И.О."
    table.cell(0, 2).text = "Должность"
    table.cell(0, 3).text = "Роспись"
    table.cell(0, 4).text = "Дата"
    table.cell(1, 1).text = "Романовский Владимир Владимирович"
    assert fix._is_acquaintance_sheet_table(table)
    assert not fix.is_agreement_signatory_table(table)
    n = fix.materialize_signatory_paragraphs_from_tables(doc)
    assert any(fix._is_acquaintance_sheet_table(t) for t in doc.tables)
    assert n == 0 or "Романовский" not in " ".join(p.text for p in doc.paragraphs if "Разработал" in p.text)
    profile = fix.detect_profile(doc, Path("ДИ инженер ЛСиМ_оформлен.docx"))
    fix.fix_last_pages_and_signatories(doc, profile)
    idx = fix.find_acquaintance_sheet_start(doc)
    assert idx is not None
    assert fix.paragraph_has_page_break_before(doc.paragraphs[idx])
    raz = fix.find_razrabotal_index(doc)
    assert not fix.paragraph_has_page_break_before(doc.paragraphs[raz])
    assert not fix.validate_acquaintance_sheet(doc)


def test_process_keeps_acquaintance_page_break():
    doc = _signatory_tail_doc()
    doc.add_paragraph("Настоящую инструкцию изучил и обязуюсь ею руководствоваться")
    profile = fix.detect_profile(doc, Path("ДИ инженер ЛСиМ_оформлен.docx"))
    fix.process_sniot_document(doc, profile)
    idx = fix.find_acquaintance_sheet_start(doc)
    assert idx is not None
    assert fix.paragraph_has_page_break_before(doc.paragraphs[idx])
    raz = fix.find_razrabotal_index(doc)
    assert not fix.paragraph_has_page_break_before(doc.paragraphs[raz])


def test_spelling_error_skips_abbreviations():
    assert fix.is_abbreviation_token("ЛСиМ")
    assert fix.is_abbreviation_token("СНиОТ")
    assert fix.is_abbreviation_token("ТКП")
    assert fix.is_abbreviation_token("НПА")
    assert fix.is_abbreviation_token("ЛПА")
    assert fix.is_abbreviation_token("ООТиЗ")
    assert not fix.is_abbreviation_token("инженер")
    assert not fix.is_abbreviation_token("выполнение")
    assert fix.spelling_error_is_abbreviation("ЛСиМ")
    assert fix.spelling_error_is_abbreviation(" СНиОТ; ")
    assert not fix.spelling_error_is_abbreviation("выполнени")
    assert fix.pick_spelling_suggestion("ЛСиМ", ["Осим", "Лим"]) is None
    assert fix.pick_spelling_suggestion("выполнени", ["выполнение", "выполнить"]) == "выполнение"
    assert "apply_word_grammar_via_com" in inspect.getsource(fix.process_sniot_document)
    assert "SpellingErrors" in inspect.getsource(fix.apply_word_grammar_via_com)
    assert "CheckGrammar" in inspect.getsource(fix.apply_word_grammar_via_com)
    assert "GrammaticalErrors" in inspect.getsource(fix.apply_word_grammar_via_com)
    assert "GetSpellingSuggestions" in inspect.getsource(fix._fix_word_spelling_errors)


def test_word_grammar_runs_after_xml_save_not_swallowed_by_skip():
    """CLI --skip-word только XML; живой GUI после записи вызывает Word отдельно."""
    apply_src = inspect.getsource(fix.apply_sniot_rules_to_file)
    pos_skip = apply_src.find("if skip_word:")
    pos_gram = apply_src.find("run_word_grammar_check_subprocess")
    assert pos_skip >= 0 and pos_gram > pos_skip
    assert "CheckGrammar" in inspect.getsource(fix.apply_word_grammar_via_com)


def test_interpret_word_signatories_ok_on_penultimate_with_acquaintance():
    report = fix.interpret_word_signatory_layout(
        body_page=7, signatory_page=7, total_pages=8, acquaintance_page=8
    )
    assert report["orphaned"] is False
    assert "вместе" in report["message"]
    assert "ознакомления" in report["message"]


def test_collapse_adjacent_duplicate_words():
    assert (
        fix.collapse_adjacent_duplicate_words(
            "заместителя начальника службы службы механизации"
        )
        == "заместителя начальника службы механизации"
    )
    assert fix.collapse_adjacent_duplicate_words("ТКП ТКП 45") == "ТКП 45"
    assert fix.collapse_adjacent_duplicate_words("и и") == "и"
    assert fix.collapse_adjacent_duplicate_words("1.1. 1.1. Текст") == "1.1. 1.1. Текст"


def test_normalize_number_separator_leading_and_one_space():
    assert fix.normalize_number_separator("  1.1.  Текст") == "1.1. Текст"
    assert fix.normalize_number_separator("\t1.2. Пункт") == "1.2. Пункт"
    assert (
        fix.normalize_number_separator("          2.2.17. Соблюдает")
        == "2.2.17. Соблюдает"
    )
    assert fix.normalize_number_separator("1.8.1 Гражданским") == "1.8.1. Гражданским"
    assert fix.normalize_number_separator("1.8.25.  действующими") == "1.8.25. действующими"


def test_acquaintance_phrase_di_oznakomlen_detected():
    heading = (
        "С должностной инструкцией заместителя начальника службы "
        "механизации и автомобильного транспорта ознакомлен"
    )
    assert fix.is_acquaintance_sheet_text(heading)
    assert not fix.is_acquaintance_sheet_text(
        "2.2.29. Обеспечивает ознакомление подчиненных работников"
    )
    doc = _signatory_tail_doc()
    doc.add_paragraph(heading)
    idx = fix.find_acquaintance_sheet_start(doc)
    assert idx is not None
    n = fix.ensure_acquaintance_sheet_separate_page(doc)
    assert n >= 1
    assert fix.paragraph_has_page_break_before(doc.paragraphs[idx])


def test_signatory_date_plaque_own_line_short_underscores():
    doc = _signatory_tail_doc()
    last = doc.paragraphs[-1]
    last.text = "Начальник ОК\tМ.А.Руфкина «___»__________ 2026г."
    n = fix.fix_signatory_date_plaques(doc)
    assert n >= 1
    texts = [p.text for p in doc.paragraphs]
    assert any("М.А.Руфкина" in t and "2026" not in t for t in texts)
    date_lines = [p for p in doc.paragraphs if fix.is_signatory_date_plaque(p.text)]
    assert date_lines
    plaque = date_lines[-1]
    assert plaque.text.startswith("\t")
    assert "«__»_______ 2026г." in plaque.text.replace("\xa0", " ")
    assert plaque.paragraph_format.keep_together is True
    assert "\n" not in plaque.text


def test_acquaintance_job_mismatch_marks_red_heading():
    doc = Document()
    doc.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ заместителя начальника ОМТС")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph(
        "На должность заместителя начальника ОМТС назначается лицо, имеющее образование."
    )
    for i in range(1, 8):
        doc.add_paragraph(f"1.{i}. Пункт общих положений.")
    doc.add_paragraph("")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер\tИ.И. Иванов")
    doc.add_paragraph("")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник\tП.П. Петров")
    heading = doc.add_paragraph(
        "С должностной инструкцией инженера ЛСиМ ознакомлен"
    )
    path = Path("Зам начальника ОМТС_оформлен.docx")
    profile = fix.detect_profile(doc, path)
    fix.process_sniot_document(doc, profile, source_path=path)
    idx = fix.find_acquaintance_sheet_start(doc)
    assert idx is not None
    assert fix.paragraph_has_page_break_before(doc.paragraphs[idx])
    assert fix.paragraph_has_mismatch_shading(doc.paragraphs[idx])
    assert not fix.jobs_equivalent(
        "заместителя начальника ОМТС", "инженера ЛСиМ"
    )
    issues = fix.validate_acquaintance_job_mismatch(doc, path)
    assert issues == []
    _ = heading


def test_acquaintance_job_match_clears_red():
    doc = Document()
    doc.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ заместителя начальника ОМТС")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph(
        "На должность заместителя начальника ОМТС назначается лицо, имеющее образование."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер\tИ.И. Иванов")
    heading = doc.add_paragraph(
        "С должностной инструкцией заместителя начальника ОМТС ознакомлен"
    )
    fix.set_paragraph_shading(heading, "FF0000")
    path = Path("Зам начальника ОМТС_оформлен.docx")
    n = fix.mark_acquaintance_heading_if_job_mismatch(doc, path)
    assert n >= 1
    assert not fix.paragraph_has_mismatch_shading(heading)
    assert fix.jobs_equivalent(
        "заместителя начальника ОМТС", "заместителя начальника ОМТС"
    )


def test_process_collapses_duplicate_word_on_acquaintance():
    doc = _signatory_tail_doc()
    doc.add_paragraph(
        "С должностной инструкцией заместителя начальника службы службы "
        "механизации и автомобильного транспорта ознакомлен"
    )
    profile = fix.detect_profile(doc, Path("Зам начальника ОМТС_оформлен.docx"))
    fix.process_sniot_document(
        doc, profile, source_path=Path("Зам начальника ОМТС_оформлен.docx")
    )
    idx = fix.find_acquaintance_sheet_start(doc)
    assert idx is not None
    text = doc.paragraphs[idx].text
    assert "службы службы" not in text
    assert "службы механизации" in text


def test_find_body_start_unnumbered_obshchie():
    doc = Document()
    doc.add_paragraph("МИНСКИЙ  ГОРОДСКОЙ  ИСПОЛНИТЕЛЬНЫЙ  КОМИТЕТ")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст пункта.")
    doc.add_paragraph("ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    assert fix.find_body_start_index(doc) == 3
    fix.normalize_first_chapter_heading(doc)
    assert doc.paragraphs[3].text.startswith("1 ОБЩИЕ")
    n = fix.restore_chapter_headers(doc)
    assert n >= 1
    assert "2 ФУНКЦИИ" in doc.paragraphs[5].text


def test_format_title_block_aligns_utverzhdayu_and_names():
    doc = Document()
    org = doc.add_paragraph("МИНСКИЙ  ГОРОДСКОЙ  ИСПОЛНИТЕЛЬНЫЙ  КОМИТЕТ")
    org.paragraph_format.first_line_indent = Pt(20)
    two = doc.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ           \tУТВЕРЖДАЮ                           ")
    two.paragraph_format.first_line_indent = Pt(36)
    doc.add_paragraph("                         \t____________ И.И.Иванов")
    date = doc.add_paragraph("                                                                                        «___» ___________ 2026г.")
    date.alignment = None
    minsk = doc.add_paragraph("Минск 2026")
    minsk.alignment = None
    doc.add_paragraph("ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Настоящая должностная инструкция определяет обязанности.")
    profile = fix.detect_profile(doc, Path("Зам начальника ОМТС_оформлен.docx"))
    fix.process_sniot_document(doc, profile)
    body = fix.find_body_start_index(doc)
    assert body > 0
    assert "  " not in doc.paragraphs[0].text
    assert fix.is_paragraph_centered(doc.paragraphs[0])
    tables = [t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]
    assert tables
    left = tables[0].rows[0].cells[0]
    right = tables[0].rows[0].cells[1]
    left_blob = "\n".join(p.text for p in left.paragraphs)
    right_blob = "\n".join(p.text for p in right.paragraphs)
    assert "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ" in left_blob
    assert "номер инструкции" in left_blob
    assert any(fix._is_title_number_underline_line(p.text or "") for p in left.paragraphs)
    assert "УТВЕРЖДАЮ" in right_blob.upper()
    idx = fix.find_title_city_year_index(doc)
    assert idx is not None and idx < body
    assert fix.is_paragraph_centered(doc.paragraphs[idx])
    assert not fix.validate_title_block(doc)
    assert not fix.validate_title_city_year(doc)


def test_title_etalon_constants_used_by_format_and_process():
    """Константы титула должны реально вызываться из format/process, не лежать мёртвым грузом."""
    fmt = inspect.getsource(fix.format_title_block)
    one = inspect.getsource(fix._format_one_title_paragraph)
    city = inspect.getsource(fix.place_title_city_year_at_bottom)
    frame = inspect.getsource(fix._set_bottom_page_frame)
    mand = inspect.getsource(fix.apply_mandatory_layout_fixes)
    proc = inspect.getsource(fix.process_sniot_document)
    assert "format_title_block" in mand
    assert "place_title_city_year_at_bottom" in mand
    assert "apply_mandatory_layout_fixes(doc, profile)" in proc
    combined = fmt + one + city + frame
    combined += inspect.getsource(fix._ensure_title_right_tab)
    combined += inspect.getsource(fix.ensure_paragraph_right)
    combined += inspect.getsource(fix.ensure_paragraph_left)
    combined += inspect.getsource(fix._ensure_title_left_right_stamp)
    combined += inspect.getsource(fix.apply_title_instruction_number_font)
    combined += inspect.getsource(fix._write_title_right_cell)
    combined += inspect.getsource(fix._strip_paragraph_tabs)
    combined += inspect.getsource(fix.compact_title_stamp_date_plaque)
    for name in (
        "TITLE_ORG_ALIGN",
        "TITLE_STAMP_ALIGN",
        "TITLE_DOC_NAME_ALIGN",
        "TITLE_UTVERZHDAYU_BOLD",
        "TITLE_RIGHT_TAB_TWIPS",
        "TITLE_CITY_YEAR_ALIGN",
        "TITLE_CITY_YEAR_FRAME_ATTRS",
        "TITLE_NUMBER_FONT_PT",
        "TITLE_NUMBER_LABEL",
        "TITLE_STAMP_DATE_TEMPLATE",
        "_strip_paragraph_tabs",
        "canonical_title_city_year",
    ):
        assert name in combined, name
    compact = inspect.getsource(fix.compact_title_empty_paragraphs)
    gap = inspect.getsource(fix.ensure_title_stamp_gap_after_header)
    limit_src = inspect.getsource(fix._title_empty_run_limit)
    assert "_title_empty_run_limit" in compact
    assert "TITLE_EMPTY_BETWEEN_BLOCKS" in limit_src
    assert "TITLE_EMPTY_AFTER_STAMP_BEFORE_CITY" in limit_src
    assert "TITLE_EMPTY_BEFORE_STAMP_MAX" in limit_src
    assert "TITLE_EMPTY_BEFORE_STAMP" in gap
    assert fix.TITLE_EMPTY_BEFORE_STAMP == 6
    assert fix.TITLE_EMPTY_BEFORE_STAMP_MIN == 6
    assert fix.TITLE_EMPTY_BEFORE_STAMP_MAX == 8
    assert fix.TITLE_EMPTY_BETWEEN_BLOCKS == 1
    assert fix.TITLE_EMPTY_AFTER_STAMP_BEFORE_CITY == 1
    assert fix.MAX_TITLE_EMPTY_RUN == 1
    assert fix.TITLE_UTVERZHDAYU_BOLD is False
    assert fix.TITLE_ORG_ALIGN == "center"
    assert fix.TITLE_STAMP_ALIGN == "right"
    assert fix.TITLE_DOC_NAME_ALIGN == "left"
    assert fix.canonical_title_city_year("Минск-2026") == "МИНСК 2026"
    assert fix.canonical_title_city_year("Минск, 2026") == "МИНСК 2026"
    assert fix.is_allowed_sample_path(
        Path(
            r"\\srv-data\doc\9 - Служба надёжности и охраны труда (СНиОТ)"
            r"\!!!ОБМЕН\РАССМОТРЕНИЕ  ДИ, РИ, ПОЛОЖЕНИЙ\ОМТС"
            r"\Зам начальника ОМТС_17.08.26.docx"
        )
    ) is False


def test_title_stamp_gap_six_to_eight_empty_lines():
    """После шапки перед таблицей грифа — 6–8 пустых строк (цель 6)."""
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст пункта.")
    fix.format_title_block(doc)
    fix.place_title_city_year_at_bottom(doc)
    fix.ensure_section_break_after_city_year(doc)
    stamps = [t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]
    assert stamps
    empty_before = fix.count_empty_before_stamp_table(doc)
    assert (
        fix.TITLE_EMPTY_BEFORE_STAMP_MIN
        <= empty_before
        <= fix.TITLE_EMPTY_BEFORE_STAMP_MAX
    )
    assert empty_before == fix.TITLE_EMPTY_BEFORE_STAMP
    prev = stamps[0]._tbl.getprevious()
    while prev is not None and prev.tag == qn("w:p"):
        para = Paragraph(prev, doc)
        if fix.is_paragraph_empty(para):
            prev = prev.getprevious()
            continue
        break
    assert prev is not None
    assert "МИНСККОММУНТЕПЛОСЕТЬ" in (Paragraph(prev, doc).text or "").upper()
    idx = fix.find_title_city_year_index(doc)
    assert idx is not None
    city = doc.paragraphs[idx]
    fr = city._p.find(qn("w:pPr")).find(qn("w:framePr"))
    assert fr is not None
    assert fr.get(qn("w:vAnchor")) == "margin"
    holder = Paragraph(city._p.getnext(), doc)
    chapter = doc.paragraphs[fix.find_body_start_index(doc)]
    assert fix._sectpr_is_next_page(fix._paragraph_sectpr(holder))
    assert holder._p.getnext() is chapter._p
    assert "ОБЩИЕ ПОЛОЖЕНИЯ" in chapter.text.upper()
    assert not fix.validate_title_stamp_gap(doc)
    assert not fix.validate_title_page_separated(doc)

    fat = Document()
    fat.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    for _ in range(12):
        fat.add_paragraph("")
    fat.add_paragraph("УТВЕРЖДАЮ")
    fat.add_paragraph("Минск 2026")
    fat.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    fat.add_paragraph("1.1. Текст.")
    fix.format_title_block(fat)
    assert (
        fix.TITLE_EMPTY_BEFORE_STAMP_MIN
        <= fix.count_empty_before_stamp_table(fat)
        <= fix.TITLE_EMPTY_BEFORE_STAMP_MAX
    )
    assert not fix.validate_title_stamp_gap(fat)


def test_format_title_block_applies_etalon_stamp_and_minsk():
    doc = Document()
    org = doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    org.runs[0].bold = True
    utv = doc.add_paragraph("УТВЕРЖДАЮ")
    utv.runs[0].bold = True
    doc.add_paragraph("Минск-2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст пункта.")
    n = fix.format_title_block(doc)
    assert n >= 1
    assert fix.is_paragraph_centered(doc.paragraphs[0])
    tables = [t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]
    assert tables
    stamp = next(
        p
        for cell in tables[0].rows[0].cells
        for p in cell.paragraphs
        if "УТВЕРЖДАЮ" in (p.text or "").upper()
    )
    assert all((r.bold is False) or (r.bold is None) for r in stamp.runs)
    fix.place_title_city_year_at_bottom(doc)
    idx = fix.find_title_city_year_index(doc)
    assert idx is not None
    assert doc.paragraphs[idx].text.strip() == "МИНСК 2026"
    assert fix.is_paragraph_centered(doc.paragraphs[idx])
    p_pr = doc.paragraphs[idx]._p.find(qn("w:pPr"))
    assert p_pr is not None and p_pr.find(qn("w:framePr")) is not None
    assert not fix.validate_title_block(doc)
    assert not fix.validate_title_city_year(doc)


def test_format_title_stamp_table_left_name_right_utverzhdayu():
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ"
    table.cell(0, 2).paragraphs[0].add_run("УТВЕРЖДАЮ").bold = True
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    n = fix.format_title_block(doc)
    assert n >= 1
    stamps = [t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]
    assert stamps
    table = stamps[0]
    cells = fix._row_unique_cells(table.rows[0])
    assert len(cells) == 2
    left_blob = "\n".join(p.text for p in cells[0].paragraphs)
    right = next(
        p
        for cell in cells
        for p in cell.paragraphs
        if "УТВЕРЖДАЮ" in (p.text or "").upper()
    )
    assert "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ" in left_blob
    assert fix.paragraph_jc(cells[0].paragraphs[0]) == "left"
    assert fix.paragraph_jc(right) == "left"
    assert all((r.bold is False) or (r.bold is None) for r in right.runs)


def test_title_left_underline_number_12pt_right_date_keeplines():
    """Слева линия + «номер инструкции» 12 pt; справа УТВЕРЖДАЮ и дата keepLines; Минск не в footer."""
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Начальник службы")
    doc.add_paragraph("____________ И.И.Иванов")
    doc.add_paragraph("«___» ___________ 2026г.")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст пункта.")
    doc.sections[0].footer.paragraphs[0].text = "МИНСК 2026"
    n = fix.format_title_block(doc)
    assert n >= 1
    fix.place_title_city_year_at_bottom(doc)
    stamps = [t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]
    assert stamps
    left, right = stamps[0].rows[0].cells[0], stamps[0].rows[0].cells[1]
    left_texts = [p.text.strip() for p in left.paragraphs if (p.text or "").strip()]
    assert any(fix._is_title_number_underline_line(t) for t in left_texts)
    num = next(p for p in left.paragraphs if fix._is_title_instruction_number_line(p.text or ""))
    assert fix._paragraph_font_half_points(num) == fix.TITLE_NUMBER_FONT_HALF
    assert num.text.strip() == "номер инструкции"
    assert any("УТВЕРЖДАЮ" in (p.text or "").upper() for p in right.paragraphs)
    date_p = next(p for p in right.paragraphs if fix.is_signatory_date_plaque(p.text or ""))
    assert fix._paragraph_has_keep_lines(date_p)
    assert not (date_p.text or "").startswith("\t")
    assert "«___»_______________" in (date_p.text or "").replace("\xa0", " ")
    sign_p = next(
        p
        for p in right.paragraphs
        if fix.TITLE_IOF_RE.search(p.text or "")
    )
    assert "\t" not in (sign_p.text or "")
    assert not fix._paragraph_has_tab_xml(sign_p)
    assert "Д.А." in (sign_p.text or "") or "Иванов" in (sign_p.text or "")
    assert not any(fix.is_city_year_paragraph(p.text or "") for p in doc.sections[0].footer.paragraphs)
    idx = fix.find_title_city_year_index(doc)
    assert idx is not None
    frame = doc.paragraphs[idx]._p.find(qn("w:pPr"))
    assert frame is not None and frame.find(qn("w:framePr")) is not None
    assert not fix.validate_title_stamp_left_right(doc)
    assert not fix.validate_city_year_not_in_headers_footers(doc)


def test_leading_spaces_before_item_number_stripped():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("          1.9.16. Требования по охране труда;")
    doc.add_paragraph("  2.2.10. Готовит предложения.")
    n = fix.normalize_item_number_spacing(doc)
    assert n >= 2
    assert doc.paragraphs[1].text.startswith("1.9.16. ")
    assert doc.paragraphs[2].text.startswith("2.2.10. ")
    assert not fix.validate_item_number_spacing(doc)


def test_xml_tab_before_item_number_is_stripped():
    """w:tab перед 1.1. не виден в paragraph.text — чистим XML runs."""
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    para = doc.add_paragraph()
    run = para.add_run()
    run._r.append(OxmlElement("w:tab"))
    para.add_run("  1.1.  Настоящая инструкция определяет обязанности.")
    assert "\t" in fix.paragraph_xml_visible_text(para)
    n = fix.normalize_item_number_spacing(doc)
    assert n >= 1
    visible = fix.paragraph_xml_visible_text(doc.paragraphs[1])
    assert visible.startswith("1.1. ")
    assert "\t" not in visible
    assert not fix.validate_item_number_spacing(doc)


def test_apply_sniot_copies_to_agent_before_word():
    src = inspect.getsource(fix.apply_sniot_rules_to_file)
    assert "collect_five_layout_issues" in src
    assert "apply_mandatory_layout_fixes" in src
    assert src.find("copy_to_target_unlocking_word") < src.find("run_word_grammar_check_subprocess")
    assert "пять правок XML" in src


def test_prevent_orphan_removes_empty_between_header_and_text():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("")
    doc.add_paragraph("1.1. Текст главы сразу под заголовком.")
    changed = fix.prevent_chapter_header_orphan(doc)
    assert changed >= 1
    hdr = fix.find_paragraph_index(doc, "1 ОБЩИЕ")
    nxt = doc.paragraphs[hdr + 1].text.strip()
    assert nxt.startswith("1.1.")
    assert not fix.validate_chapter_header_orphan(doc)


def test_mandatory_layout_fixes_are_called_from_process():
    src = inspect.getsource(fix.process_sniot_document)
    assert "apply_mandatory_layout_fixes(doc, profile)" in src
    body = inspect.getsource(fix.apply_mandatory_layout_fixes)
    for name in (
        "strip_unnecessary_characters",
        "remove_list_markers_in_body",
        "normalize_item_number_spacing",
        "apply_body_paragraph_format",
        "apply_body_single_line_spacing",
        "apply_table_paragraph_no_indent",
        "ensure_title_page_separated",
        "format_title_block",
        "place_title_city_year_at_bottom",
        "ensure_section_break_after_city_year",
        "prevent_chapter_header_orphan",
        "apply_page_setup_deloproizvodstvo",
        "apply_signatory_fio_one_line",
        "apply_signatory_line_spacing",
        "fix_signatory_date_plaques",
        "keep_signatory_block_with_text",
        "compact_extra_empty_lines_in_signatory_block",
        "ensure_deloproizvodstvo_in_must_know",
        "ensure_missing_developer_block",
        "fix_duty_by_order_commas",
        "expand_first_lpa_abbreviation",
        "remove_city_year_from_headers_footers",
        "check_document_punctuation_after_edit",
        "ensure_title_stamp_gap_after_header",
        "_clear_breaks_between_stamp_table_and_city",
    ):
        assert name in body


def test_punctuation_of_whole_document_checked_after_edit():
    """После правок пунктуацию проверяют по всему документу, не только по изменённому абзацу."""
    apply_src = inspect.getsource(fix.apply_mandatory_layout_fixes)
    pos_text = apply_src.rfind("fix_duty_by_order_commas")
    pos_punct = apply_src.rfind("check_document_punctuation_after_edit")
    assert pos_text >= 0 and pos_punct > pos_text
    val_src = inspect.getsource(fix.validate_sniot_document)
    assert "validate_document_punctuation" in val_src
    assert "strip_unnecessary_characters" in inspect.getsource(
        fix.check_document_punctuation_after_edit
    )
    assert "normalize_must_know_list_punctuation" in inspect.getsource(
        fix.check_document_punctuation_after_edit
    )
    assert "validate_unnecessary_characters" in inspect.getsource(
        fix.validate_document_punctuation
    )
    assert "validate_must_know_list_punctuation" in inspect.getsource(
        fix.validate_document_punctuation
    )

    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст пункта , и слово .")
    doc.add_paragraph("1.2. Список пунктов,, без правки.")
    assert fix.validate_document_punctuation(doc)
    profile = fix.detect_profile(doc, Path("_test_punct_.docx"))
    fix.apply_mandatory_layout_fixes(doc, profile)
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "слово." in joined
    assert "пункта," in joined
    assert ",," not in joined
    assert not fix.validate_document_punctuation(doc)


def test_omts_filename_is_not_satp_scheme():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("В своей деятельности руководствуется нормативными актами.")
    doc.add_paragraph("Заместитель начальника должен знать требования охраны труда.")
    path = Path("Зам начальника ОМТС_оформлен.docx")
    assert fix.has_di_satp_numbering_structure(doc, path) is False
    profile = fix.detect_profile(doc, path)
    assert profile.has_di_satp_numbering is False
    assert profile.kind == "di"


def test_ppr_tab_stop_is_not_leading_item_tab():
    """Стоп табуляции в pPr не должен считаться табом перед «1 ОБЩИЕ»."""
    doc = Document()
    para = doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p_pr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), "284")
    tabs.append(tab)
    p_pr.append(tabs)
    vis = fix.paragraph_xml_visible_text(para)
    assert vis.startswith("1 ОБЩИЕ")
    assert vis[:1] != "\t"
    assert not fix.validate_item_number_spacing(doc)


def test_mandatory_layout_fixes_omts_like_temp_doc(tmp_path: Path):
    """Пять дефектов уходят на TEMP-копии; файл пользователя на N:\\ не трогаем."""
    doc = Document()
    doc.add_paragraph("МИНСКИЙ  ГОРОДСКОЙ  ИСПОЛНИТЕЛЬНЫЙ  КОМИТЕТ")
    doc.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ           \tУТВЕРЖДАЮ")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("")
    doc.add_paragraph("          1.1.  Настоящая инструкция определяет обязанности.")
    ch2 = doc.add_paragraph("2 ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    ch2.paragraph_format.page_break_before = True
    tab_item = doc.add_paragraph()
    tab_run = tab_item.add_run()
    tab_run._r.append(OxmlElement("w:tab"))
    tab_item.add_run("  2.2.10. Готовит предложения.")
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("5.1.1. Несёт ответственность.")
    doc.add_paragraph("")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер\tВ.В.Дубовик «___»__________ 2026г.")
    doc.add_paragraph("")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник\tИ.И.Иванов")
    path = tmp_path / "Зам начальника ОМТС_оформлен.docx"
    doc.save(path)
    loaded = Document(str(path))
    profile = fix.detect_profile(loaded, path)
    assert profile.has_di_satp_numbering is False
    fix.process_sniot_document(loaded, profile, source_path=path)
    assert "  " not in loaded.paragraphs[0].text.replace("\t", " ")
    body = fix.find_body_start_index(loaded)
    assert body > 0
    assert any("1 ОБЩИЕ" in (p.text or "") for p in loaded.paragraphs)
    item = next(p.text for p in loaded.paragraphs if "1.1." in (p.text or ""))
    assert item.lstrip(" \t") == item
    assert item.startswith("1.1. ")
    item2 = next(p.text for p in loaded.paragraphs if "2.2.10." in (p.text or ""))
    assert item2.startswith("2.2.10. ")
    assert "\t" not in fix.paragraph_xml_visible_text(
        next(p for p in loaded.paragraphs if "2.2.10." in (p.text or ""))
    )
    ch2 = next(p for p in loaded.paragraphs if "ФУНКЦИИ" in (p.text or "").upper())
    assert not fix.paragraph_has_page_break_before(ch2)
    idx = fix.find_title_city_year_index(loaded)
    assert idx is not None and idx < body
    assert fix.is_paragraph_centered(loaded.paragraphs[idx])
    hdr = next(i for i, p in enumerate(loaded.paragraphs) if "ОБЩИЕ ПОЛОЖЕНИЯ" in (p.text or "").upper())
    assert loaded.paragraphs[hdr + 1].text.strip()
    assert not fix.paragraph_has_page_break_before(loaded.paragraphs[hdr])
    plaques = [p for p in loaded.paragraphs if fix.is_signatory_date_plaque(p.text)]
    assert plaques
    assert plaques[-1].text.startswith("\t")
    issues = []
    issues.extend(fix.collect_five_layout_issues(loaded, profile))
    assert issues == []


def test_unc_and_n_drive_are_same_agent_folder():
    n_file = (
        r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В\Агент"
        r"\ПРОЕКТ Старший мастер_оформлен.docx"
    )
    unc_file = (
        r"\\srv-data\doc\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В\Агент"
        r"\ПРОЕКТ Старший мастер_оформлен.docx"
    )
    obmen = (
        r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\!!!ОБМЕН"
        r"\РАССМОТРЕНИЕ  ДИ, РИ, ПОЛОЖЕНИЙ\ПРОЕКТ_образец.docx"
    )
    assert fix.is_path_in_user_agent_dir(n_file) is True
    assert fix.is_path_in_user_agent_dir(unc_file) is True
    assert fix.is_path_in_user_agent_dir(obmen) is False
    assert fix.is_allowed_sample_path(unc_file.replace("оформлен", "образец")) is True
    assert fix.is_allowed_sample_path(obmen) is False


def test_format_title_unbolds_org_header_and_validate_catches():
    doc = Document()
    org = doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    org.runs[0].bold = True
    doc.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ")
    utv = doc.add_paragraph("УТВЕРЖДАЮ")
    utv.runs[0].bold = True
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    before = fix.validate_title_block(doc)
    assert any("жирн" in item.lower() for item in before)
    fix.format_title_block(doc)
    after = fix.validate_title_block(doc)
    assert not any("Шапка предприятия жирная" in item for item in after)
    assert not any("УТВЕРЖДАЮ» не должен быть жирным" in item for item in after)


def test_signatory_spacing_stops_before_acquaintance():
    from docx.enum.text import WD_LINE_SPACING

    doc = Document()
    doc.add_paragraph("5.1.1. Пункт.")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер\tИ.И. Иванов")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник\tП.П. Петров")
    acq = doc.add_paragraph(
        "С должностной инструкцией старшего мастера ознакомлен"
    )
    acq.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    acq.paragraph_format.line_spacing = 1.0
    profile = fix.DocumentProfile("di", None, True, False, None)
    fix.apply_signatory_line_spacing(doc, profile)
    assert fix.paragraph_has_one_point_five_spacing(doc.paragraphs[1])
    assert fix.paragraph_has_single_line_spacing(acq) or acq.paragraph_format.line_spacing == 1.0
    assert not fix.validate_signatory_line_spacing(doc, profile)


def test_validate_catches_wrapped_signatory_fio():
    doc = Document()
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер ________________")
    doc.add_paragraph("И.И. Иванов")
    profile = fix.DocumentProfile("di", None, True, False, None)
    issues = fix.validate_signatory_fio_one_line(doc, profile)
    assert any("перенесен" in item.lower() or "ФИО" in item for item in issues)
    n = fix.apply_signatory_fio_one_line(doc)
    assert n >= 1
    texts = [p.text for p in doc.paragraphs]
    assert any("Инженер" in t and "И.И." in t and "\t" in t for t in texts)
    assert not any(t.strip() == "И.И. Иванов" for t in texts)
    assert fix.validate_signatory_fio_one_line(doc, profile) == []


def test_temp_right_margin_10_fio_same_line_date_tab(tmp_path):
    doc = Document()
    doc.sections[0].right_margin = fix.Mm(8)
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Настоящая инструкция определяет обязанности работника.")
    for i in range(2, 8):
        doc.add_paragraph(f"1.{i}. Пункт общих положений для объёма тела.")
    doc.add_paragraph("")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер ________________")
    doc.add_paragraph("И.И. Иванов")
    doc.add_paragraph("")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник службы\tП.П. Петров «___»__________ 2026г.")
    profile = fix.detect_profile(doc, Path("temp_margins_fio.docx"))
    profile.has_signatories = True
    fix.process_sniot_document(doc, profile, source_path=Path("temp_margins_fio.docx"))
    out = tmp_path / "margins_fio_tab.docx"
    doc.save(str(out))
    loaded = Document(str(out))
    assert abs(fix._length_mm(loaded.sections[0].right_margin) - 10) < 0.6
    assert not any(p.text.strip() == "И.И. Иванов" for p in loaded.paragraphs)
    assert any(
        "Инженер" in p.text and "И.И." in p.text and "\t" in p.text
        for p in loaded.paragraphs
    )
    date_p = next(p for p in loaded.paragraphs if fix.is_signatory_date_plaque(p.text))
    iof_p = next(
        p
        for p in loaded.paragraphs
        if "\t" in (p.text or "") and fix.SIGNATORY_NAME_TAIL.search(p.text or "")
    )
    assert date_p.text.startswith("\t")
    assert fix.signatory_first_tab_pos_twips(date_p) == fix.signatory_first_tab_pos_twips(iof_p)
    assert fix.signatory_first_tab_pos_twips(date_p) is not None
    assert not fix.validate_page_margins(loaded)
    date_issues = fix.validate_signatory_date_plaques(loaded, profile)
    assert date_issues == []


def test_compact_signatory_block_removes_gaps_between_people():
    doc = Document()
    doc.add_paragraph("5.1.1. Последний пункт.")
    doc.add_paragraph("")
    doc.add_paragraph("Разработал:")
    doc.add_paragraph("Инженер\tИ.И.Иванов")
    doc.add_paragraph("\t«__»_______ 2026г.")
    doc.add_paragraph("")
    doc.add_paragraph("Согласовано:")
    doc.add_paragraph("Начальник ОК\tМ.А.Руфкина")
    doc.add_paragraph("\t«__»_______ 2026г.")
    doc.add_paragraph("")
    doc.add_paragraph("Начальник ЮО\tГ.А.Авраменко")
    doc.add_paragraph("\t«__»_______ 2026г.")
    profile = fix.DocumentProfile("di", None, True, False, 0)
    n = fix.compact_extra_empty_lines_in_signatory_block(doc, profile)
    assert n >= 1
    texts = [p.text.strip() for p in doc.paragraphs]
    sog = texts.index("Согласовано:")
    after = texts[sog + 1 :]
    assert "" not in after
    issues = fix.validate_signatory_block(doc, profile)
    assert not any("внутри блока согласующих" in i for i in issues)


def test_keep_signatory_block_with_text_sets_keep_next_chain():
    doc = _signatory_tail_doc()
    profile = fix.DocumentProfile("di", None, True, False, 0)
    n = fix.keep_signatory_block_with_text(doc, profile)
    assert n >= 1
    last_body = fix.find_last_body_paragraph_before_signatories(doc)
    raz = fix.find_razrabotal_index(doc)
    assert doc.paragraphs[last_body].paragraph_format.keep_with_next
    assert doc.paragraphs[raz].paragraph_format.keep_with_next
    last = doc.paragraphs[-1]
    assert not last.paragraph_format.keep_with_next
    assert not fix.validate_page_layout_flags(doc, profile)


def test_title_stamp_geometry_matches_user_edit():
    assert fix.TITLE_STAMP_TABLE_COL_TWIPS == ("4967", "5212")
    assert fix.TITLE_STAMP_TABLE_WIDTH_TWIPS == "10179"
    assert not fix.TITLE_STAMP_DATE_TEMPLATE.startswith("\t")
    assert dict(fix.TITLE_CITY_YEAR_FRAME_ATTRS)["w:vAnchor"] == "margin"


def test_title_right_column_has_no_tabs():
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("____________\tД.А.Миронов")
    doc.add_paragraph("\t«___» ___________ 2026г.")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    fix.format_title_block(doc)
    stamps = [t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]
    right = stamps[0].rows[0].cells[1]
    for para in right.paragraphs:
        assert "\t" not in (para.text or "")
        assert not fix._paragraph_has_tab_xml(para)
    assert not fix.validate_title_stamp_left_right(doc)


def test_minsk_not_in_footer_and_not_on_sectpr():
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    doc.sections[0].footer.paragraphs[0].text = "МИНСК 2026"
    fix.ensure_title_page_separated(doc)
    fix.format_title_block(doc)
    fix.place_title_city_year_at_bottom(doc)
    idx = fix.find_title_city_year_index(doc)
    assert idx is not None
    city = doc.paragraphs[idx]
    assert fix._paragraph_sectpr(city) is None
    fr = city._p.find(qn("w:pPr")).find(qn("w:framePr"))
    assert fr is not None
    assert fr.get(qn("w:vAnchor")) == "margin"
    assert not fix.validate_city_year_not_in_headers_footers(doc)
    assert not any(fix.is_city_year_paragraph(p.text or "") for p in doc.sections[0].footer.paragraphs)
    nxt = city._p.getnext()
    assert nxt is not None and nxt.tag == qn("w:p")
    holder = Paragraph(nxt, doc)
    assert fix._sectpr_is_next_page(fix._paragraph_sectpr(holder))
    body = fix.find_body_start_index(doc)
    assert holder._p.getnext() is doc.paragraphs[body]._p
    assert not fix.validate_section_break_after_city_year(doc)


def test_must_know_adds_office_work_basics():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.6. Ведущий инженер должен знать:")
    doc.add_paragraph("нормативные правовые акты;")
    doc.add_paragraph(
        "1.7. В случае временного отсутствия работника "
        "его обязанности исполняет по распоряжению начальника отдела специалист ОМТС."
    )
    n = fix.ensure_deloproizvodstvo_in_must_know(doc)
    assert n >= 1
    texts = [p.text for p in doc.paragraphs]
    assert "основы делопроизводства." in texts
    assert "нормативные правовые акты;" in texts
    assert texts.index("основы делопроизводства.") < texts.index(
        next(t for t in texts if t.startswith("1.7."))
    )
    assert fix.ensure_deloproizvodstvo_in_must_know(doc) == 0
    skip = Document()
    skip.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    skip.add_paragraph("Заместитель начальника должен знать требования охраны труда.")
    skip.add_paragraph("1.2. Далее.")
    assert fix.ensure_deloproizvodstvo_in_must_know(skip) == 0


def test_must_know_last_item_period_and_mid_semicolon():
    """Пример пользователя: 1.9.15. …;  и  1.9.16. основы делопроизводства."""
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.9. Старший мастер должен знать:")
    doc.add_paragraph("1.9.15. коллективный договор;")
    doc.add_paragraph("1.9.16. основы делопроизводства;")
    doc.add_paragraph("2 ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    n = fix.normalize_must_know_list_punctuation(doc)
    assert n >= 1
    texts = [p.text.strip() for p in doc.paragraphs]
    assert "1.9.15. коллективный договор;" in texts
    assert "1.9.16. основы делопроизводства." in texts
    assert not fix.validate_must_know_list_punctuation(doc)

    missing = Document()
    missing.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    missing.add_paragraph("1.9. Старший мастер должен знать:")
    missing.add_paragraph("1.9.15. коллективный договор.")
    missing.add_paragraph("2 ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ")
    assert fix.ensure_deloproizvodstvo_in_must_know(missing) >= 1
    mtexts = [p.text.strip() for p in missing.paragraphs]
    assert "1.9.15. коллективный договор;" in mtexts
    assert any(t.endswith("основы делопроизводства.") for t in mtexts)
    assert not fix.validate_must_know_list_punctuation(missing)


def test_duty_by_order_gets_commas():
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    para = doc.add_paragraph(
        "1.7. В случае временного отсутствия ведущего экономиста "
        "его обязанности исполняет по распоряжению начальника отдела специалист ОМТС, "
        "который приобретает соответствующие права."
    )
    n = fix.fix_duty_by_order_commas(doc)
    assert n == 1
    assert "исполняет, по распоряжению начальника отдела, специалист" in para.text
    assert fix.fix_duty_by_order_commas(doc) == 0


def test_chapter_one_starts_after_minsk_section_break():
    """После «МИНСК 2026» — nextPage; глава 1 сразу, с верха новой страницы."""
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    n = fix.ensure_section_break_after_city_year(doc)
    assert n >= 1
    idx = fix.find_title_city_year_index(doc)
    city = doc.paragraphs[idx]
    nxt = city._p.getnext()
    holder = Paragraph(nxt, doc)
    assert fix.is_paragraph_empty(holder)
    assert fix._sectpr_is_next_page(fix._paragraph_sectpr(holder))
    chapter = doc.paragraphs[fix.find_body_start_index(doc)]
    assert "ОБЩИЕ ПОЛОЖЕНИЯ" in chapter.text.upper()
    assert holder._p.getnext() is chapter._p
    assert not fix.paragraph_has_page_break_before(chapter)
    assert not fix.is_paragraph_empty(chapter)
    assert not fix.validate_section_break_after_city_year(doc)
    assert not fix.validate_title_page_separated(doc)
    profile = fix.detect_profile(doc, Path("ДИ ведущий экономист.docx"))
    fix.process_sniot_document(doc, profile)
    idx = fix.find_title_city_year_index(doc)
    city = doc.paragraphs[idx]
    holder = Paragraph(city._p.getnext(), doc)
    chapter = doc.paragraphs[fix.find_body_start_index(doc)]
    assert fix._sectpr_is_next_page(fix._paragraph_sectpr(holder))
    assert holder._p.getnext() is chapter._p
    assert not fix.validate_section_break_after_city_year(doc)


def test_title_stamp_columns_not_mixed():
    """Левая колонка — название/линия/номер; правая — УТВЕРЖДАЮ, должность, подпись, дата."""
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ"
    table.cell(0, 1).text = "УТВЕРЖДАЮ"
    table.cell(0, 2).text = "Заместитель начальника отдела материально-технического снабжения"
    doc.add_paragraph("____________ И.И.Иванов")
    doc.add_paragraph("«___» ___________ 2026г.")
    doc.add_paragraph("номер инструкции")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    n = fix.format_title_block(doc)
    assert n >= 1
    stamps = [t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]
    assert stamps
    left, right = stamps[0].rows[0].cells[0], stamps[0].rows[0].cells[1]
    left_blob = "\n".join(p.text for p in left.paragraphs)
    right_blob = "\n".join(p.text for p in right.paragraphs)
    assert "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ" in left_blob
    assert "номер инструкции" in left_blob.casefold()
    assert "УТВЕРЖДАЮ" not in left_blob.upper()
    assert "УТВЕРЖДАЮ" in right_blob.upper()
    assert "Заместитель начальника" in right_blob
    assert "Заместитель начальника" not in left_blob
    assert "Иванов" in right_blob
    assert "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ" not in right_blob
    assert "номер инструкции" not in right_blob.casefold()
    assert not any(fix._paragraph_has_tab_xml(p) for p in right.paragraphs)
    assert not fix.validate_title_stamp_left_right(doc)


def test_title_stamp_built_from_loose_paragraphs_without_table():
    """Нет таблицы на титуле — название и УТВЕРЖДАЮ абзацами собираются в 2 колонки."""
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ")
    doc.add_paragraph("ведущего инженера ОМТС")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Начальник ОМТС")
    doc.add_paragraph("____________ И.И.Иванов")
    doc.add_paragraph("«___» ___________ 2026г.")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    assert not [t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]
    n = fix.format_title_block(doc)
    assert n >= 1
    stamps = [t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]
    assert len(stamps) == 1
    cells = fix._row_unique_cells(stamps[0].rows[0])
    assert len(cells) == 2
    left_blob = "\n".join(p.text for p in cells[0].paragraphs)
    right_blob = "\n".join(p.text for p in cells[1].paragraphs)
    assert "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ" in left_blob
    assert "номер инструкции" in left_blob.casefold()
    assert "УТВЕРЖДАЮ" not in left_blob.upper()
    assert "УТВЕРЖДАЮ" in right_blob.upper()
    assert "Начальник ОМТС" in right_blob
    assert "Начальник ОМТС" not in left_blob
    assert "Иванов" in right_blob
    body_start = fix.find_body_start_index(doc)
    loose_utv = [
        p
        for p in doc.paragraphs[:body_start]
        if not fix.paragraph_is_inside_table(p) and "УТВЕРЖДАЮ" in (p.text or "").upper()
    ]
    assert not loose_utv
    assert not fix.validate_title_stamp_left_right(doc)
    assert fix._ensure_title_left_right_stamp(doc) == 0
    assert len([t for t in fix._iter_title_tables(doc) if fix._is_title_stamp_table(t)]) == 1


def test_no_section_break_between_stamp_table_and_minsk():
    """Между таблицей и «МИНСК 2026» есть пустые строки, нет nextPage; разрыв только после города."""
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Начальник службы")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    profile = fix.detect_profile(doc, Path("ДИ ведущий экономист.docx"))
    fix.process_sniot_document(doc, profile)
    table = fix._title_stamp_table(doc)
    assert table is not None
    idx = fix.find_title_city_year_index(doc)
    assert idx is not None
    city = doc.paragraphs[idx]
    assert table._tbl.getnext() is not city._p
    assert fix.count_empty_between_stamp_and_city(doc) == fix.TITLE_EMPTY_AFTER_STAMP_BEFORE_CITY
    el = table._tbl.getnext()
    while el is not None and el is not city._p:
        if el.tag == qn("w:p"):
            para = Paragraph(el, doc)
            assert not fix._sectpr_is_next_page(fix._paragraph_sectpr(para))
            assert not fix.paragraph_has_page_break_before(para)
        el = el.getnext()
    assert not fix.validate_stamp_table_city_gap(doc)
    nxt = city._p.getnext()
    assert nxt is not None and nxt.tag == qn("w:p")
    holder = Paragraph(nxt, doc)
    assert fix._sectpr_is_next_page(fix._paragraph_sectpr(holder))
    p_pr = city._p.find(qn("w:pPr"))
    fr = p_pr.find(qn("w:framePr")) if p_pr is not None else None
    assert fr is not None and fr.get(qn("w:vAnchor")) == "margin"
    holder_pr = holder._p.find(qn("w:pPr"))
    holder_rpr = holder_pr.find(qn("w:rPr")) if holder_pr is not None else None
    if holder_rpr is not None:
        assert holder_rpr.find(qn("w:vanish")) is None
    chapter = doc.paragraphs[fix.find_body_start_index(doc)]
    assert holder._p.getnext() is chapter._p
    assert "ОБЩИЕ ПОЛОЖЕНИЯ" in chapter.text.upper()
    assert not fix.is_paragraph_empty(chapter)


def test_minsk_break_immediately_after_no_visible_gap():
    """После «МИНСК 2026» сразу разрыв раздела; глава 1 со следующей страницы, без видимых вставок."""
    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("Минск 2026")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст.")
    profile = fix.detect_profile(doc, Path("ДИ ведущий экономист.docx"))
    fix.process_sniot_document(doc, profile)
    idx = fix.find_title_city_year_index(doc)
    city = doc.paragraphs[idx]
    nxt = city._p.getnext()
    assert nxt is not None and nxt.tag == qn("w:p")
    holder = Paragraph(nxt, doc)
    assert fix.is_paragraph_empty(holder)
    assert fix._sectpr_is_next_page(fix._paragraph_sectpr(holder))
    assert not (holder.text or "").strip()
    chapter = doc.paragraphs[fix.find_body_start_index(doc)]
    assert holder._p.getnext() is chapter._p
    assert "ОБЩИЕ ПОЛОЖЕНИЯ" in chapter.text.upper()
    p_pr = city._p.find(qn("w:pPr"))
    fr = p_pr.find(qn("w:framePr")) if p_pr is not None else None
    assert fr is not None and fr.get(qn("w:vAnchor")) == "margin"
    assert fix._paragraph_sectpr(city) is None
    holder_pr = holder._p.find(qn("w:pPr"))
    holder_rpr = holder_pr.find(qn("w:rPr")) if holder_pr is not None else None
    if holder_rpr is not None:
        assert holder_rpr.find(qn("w:vanish")) is None
    assert not fix.validate_section_break_after_city_year(doc)


def test_missing_developer_block_inserts_yellow_kto_marker(tmp_path: Path):
    """Нет «Разработал:» и нет И.О.Фамилия — одна строка с жёлтым КТО???."""
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Текст пункта.")
    doc.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("5.1. несет ответственность.")
    n = fix.ensure_missing_developer_block(doc)
    assert n == 1
    line = next(p for p in doc.paragraphs if fix.MISSING_DEVELOPER_MARKER in (p.text or ""))
    assert "Разработал:" in (line.text or "")
    assert "\n" not in (line.text or "")
    assert "\t" in (line.text or "") or any(True for _ in line._p.iter(qn("w:tab")))
    assert fix._paragraph_has_keep_lines(line)
    tab_pos = fix.signatory_first_tab_pos_twips(line)
    assert tab_pos is not None and int(tab_pos) >= fix.MIN_SIGNATORY_TAB_TWIPS
    yellow_runs = [r for r in line.runs if fix.MISSING_DEVELOPER_MARKER in (r.text or "")]
    assert yellow_runs
    r_pr = yellow_runs[0]._r.find(qn("w:rPr"))
    assert r_pr is not None
    hl = r_pr.find(qn("w:highlight"))
    assert hl is not None and hl.get(qn("w:val")) == "yellow"
    kto_paras = [p for p in doc.paragraphs if fix.MISSING_DEVELOPER_MARKER in (p.text or "")]
    assert len(kto_paras) == 1
    assert fix.ensure_missing_developer_block(doc) == 0
    keep = Document()
    keep.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    keep.add_paragraph("1.1. Текст.")
    keep.add_paragraph("Разработал:")
    keep.add_paragraph("Инженер\tИ.И.Иванов")
    keep.add_paragraph("Согласовано:")
    keep.add_paragraph("Начальник\tП.П.Петров")
    assert fix.ensure_missing_developer_block(keep) == 0
    assert not any(fix.MISSING_DEVELOPER_MARKER in (p.text or "") for p in keep.paragraphs)
    fix.strip_visual_highlights(doc)
    assert not fix.validate_visual_highlights(doc)

    src = Document()
    src.add_paragraph("УНИТАРНОЕ ПРЕДПРИЯТИЕ")
    src.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    src.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ ведущий инженер")
    src.add_paragraph("УТВЕРЖДАЮ")
    src.add_paragraph("МИНСК 2026")
    src.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    src.add_paragraph("1.1. Текст пункта.")
    src.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    src.add_paragraph("5.1. несет ответственность.")
    path = tmp_path / "без_подписантов_оформлен.docx"
    src.save(path)
    loaded = Document(str(path))
    profile = fix.detect_profile(loaded, path)
    fix.process_sniot_document(loaded, profile, source_path=path)
    fio_p = next(p for p in loaded.paragraphs if fix.MISSING_DEVELOPER_MARKER in (p.text or ""))
    assert "Разработал:" in (fio_p.text or "")
    yellow = [r for r in fio_p.runs if fix.MISSING_DEVELOPER_MARKER in (r.text or "")]
    assert yellow
    r_pr = yellow[0]._r.find(qn("w:rPr"))
    assert r_pr is not None
    hl = r_pr.find(qn("w:highlight"))
    assert hl is not None and hl.get(qn("w:val")) == "yellow"
    shd = r_pr.find(qn("w:shd"))
    assert shd is not None and (shd.get(qn("w:fill")) or "").upper() == "FFFF00"
    tab_pos = fix.signatory_first_tab_pos_twips(fio_p)
    assert tab_pos is not None and int(tab_pos) >= fix.MIN_SIGNATORY_TAB_TWIPS
    assert fix._paragraph_has_keep_lines(fio_p)


def test_first_lpa_gets_expansion_others_stay_abbrev():
    """Первое ЛПА без расшифровки — с «далее - ЛПА»; повторные не дублируют."""
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    first = doc.add_paragraph(
        "1.8. В своей деятельности руководствуется ЛПА предприятия."
    )
    second = doc.add_paragraph("1.9. Требования ЛПА обязательны для работников.")
    n = fix.expand_first_lpa_abbreviation(doc)
    assert n == 1
    assert "локальным правовым актом (далее - ЛПА)" in first.text
    assert first.text.count("(далее - ЛПА)") == 1
    assert "требования ЛПА" in second.text.casefold() or "Требования ЛПА" in second.text
    assert "(далее" not in second.text
    assert fix.expand_first_lpa_abbreviation(doc) == 0
    assert first.text.count("ЛПА") == 1
    assert second.text.count("ЛПА") == 1
    assert fix.is_abbreviation_token("ЛПА")
    assert fix.spelling_error_is_abbreviation("ЛПА")
    assert fix.pick_spelling_suggestion("ЛПА", ["лпа", "лапа"]) is None

    gen = Document()
    gen.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    g = gen.add_paragraph("1.4.1. Соблюдает требованиями ЛПА и приказами.")
    assert fix.expand_first_lpa_abbreviation(gen) == 1
    assert "локальных правовых актов (далее - ЛПА)" in g.text

    loc = Document()
    loc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    lp = loc.add_paragraph("1.4.2. Указания содержатся в ЛПА службы.")
    assert fix.expand_first_lpa_abbreviation(loc) == 1
    assert "локальном правовом акте (далее - ЛПА)" in lp.text

    done = Document()
    done.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    done.add_paragraph(
        "1.8. Руководствуется локальным правовым актом (далее - ЛПА) предприятия."
    )
    done.add_paragraph("1.9. Иные ЛПА применяются в части.")
    assert fix.expand_first_lpa_abbreviation(done) == 0
    assert "Иные ЛПА" in done.paragraphs[-1].text
    assert done.paragraphs[-1].text.count("(далее") == 0


def test_process_expands_first_lpa_once(tmp_path: Path):
    src = Document()
    src.add_paragraph("МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    src.add_paragraph("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ")
    src.add_paragraph("УТВЕРЖДАЮ")
    src.add_paragraph("МИНСК 2026")
    src.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    src.add_paragraph("1.8. Инженер руководствуется ЛПА организации.")
    src.add_paragraph("1.9. Нарушение ЛПА влечёт ответственность.")
    src.add_paragraph("5 ОТВЕТСТВЕННОСТЬ")
    src.add_paragraph("5.1. несет ответственность.")
    path = tmp_path / "лпа_оформлен.docx"
    src.save(path)
    loaded = Document(str(path))
    profile = fix.detect_profile(loaded, path)
    fix.process_sniot_document(loaded, profile, source_path=path)
    joined = "\n".join(p.text for p in loaded.paragraphs)
    assert joined.count("(далее - ЛПА)") == 1
    assert "локальным правовым актом (далее - ЛПА)" in joined
    assert "Нарушение ЛПА" in joined or "нарушение ЛПА" in joined.casefold()
    assert fix.is_abbreviation_token("ЛПА")


