# -*- coding: utf-8 -*-
"""Сравнение интервалов _оформлен с *_образец.docx в папке Агент."""
from __future__ import annotations

import shutil
import sys
import tempfile
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

import fix_sniot_document as fix

USER_DIR = fix.USER_AGENT_DIR
ETALON = USER_DIR / "ПРОЕКТ Старший мастер_образец.docx"
TARGET = USER_DIR / "ПРОЕКТ Старший мастер_оформлен.docx"


@pytest.fixture(scope="module")
def etalon_doc() -> Document:
    if not ETALON.is_file():
        pytest.skip(f"Образец не найден: {ETALON}")
    return Document(ETALON)


def test_find_etalon_path_from_oformlen():
    if not TARGET.is_file():
        pytest.skip("Нет целевого файла")
    found = fix.find_etalon_path(TARGET)
    assert found is not None
    assert found.name.endswith("_образец.docx")


def test_align_spacing_increases_empty_lines(etalon_doc: Document):
    if not TARGET.is_file():
        pytest.skip("Нет целевого файла")
    with tempfile.TemporaryDirectory() as td:
        work = Path(td) / TARGET.name
        shutil.copy2(TARGET, work)
        doc = Document(work)
        profile = fix.detect_profile(doc, work)
        before = fix.collect_spacing_metrics(doc)["empty_count"]
        inserted = fix.align_spacing_to_etalon(doc, etalon_doc, profile)
        after = fix.collect_spacing_metrics(doc)["empty_count"]
        assert inserted >= 0
        assert after >= before


def test_process_with_etalon_matches_key_metrics(etalon_doc: Document):
    if not TARGET.is_file():
        pytest.skip("Нет целевого файла")
    with tempfile.TemporaryDirectory() as td:
        work = Path(td) / TARGET.name
        shutil.copy2(TARGET, work)
        doc = Document(work)
        profile = fix.detect_profile(doc, work)
        fix.process_sniot_document(doc, profile, etalon_path=ETALON)
        cmp = fix.compare_spacing_to_etalon(doc, etalon_doc)
        delta = cmp["delta"]
        assert abs(delta["empty_count"]) <= 1
        assert abs(delta["empty_after_chapter"]) <= 1
        assert cmp["target"]["razrab_empty_before"] == cmp["etalon"]["razrab_empty_before"]
        assert cmp["target"]["soglas_empty_before"] == cmp["etalon"]["soglas_empty_before"]


def test_apply_sniot_rules_to_file_with_etalon(monkeypatch: pytest.MonkeyPatch):
    if not TARGET.is_file() or not ETALON.is_file():
        pytest.skip("Нет файлов в папке Агент")
    with tempfile.TemporaryDirectory() as td:
        work = Path(td) / TARGET.name
        shutil.copy2(TARGET, work)
        monkeypatch.setattr(fix, "assert_path_writable", lambda p: Path(p).resolve())
        monkeypatch.setattr(fix, "find_etalon_path", lambda _p: ETALON)
        rep = fix.apply_sniot_rules_to_file(work, always_apply=True)
        assert rep.get("applied") is True
        doc = Document(work)
        etalon = Document(ETALON)
        cmp = fix.compare_spacing_to_etalon(doc, etalon)
        assert abs(cmp["delta"]["empty_count"]) <= 1
        assert abs(cmp["delta"]["empty_after_chapter"]) <= 1
