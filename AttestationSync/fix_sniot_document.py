# -*- coding: utf-8 -*-
"""
Автоисправление документов СНиОТ (ДИ, РИ, Положения, инструкции…) по sniot-di-documents.mdc
и Инструкции по делопроизводству 2025 (Минюст РБ, п.18, гл.7).

Источник правил: sniot-di-documents.mdc + Инструкция по делопроизводству 2025 — **важнее** docx-образца.
Главный API: process_sniot_document(path), validate_sniot_document(doc),
validate_full_document_on_disk(path) — обязательная перечитка с диска перед записью в Агент.
"""
from __future__ import annotations

# Метка сборки — печатается в stdout и обязана попасть в окно «Готово» без подмены константой GUI.
SCRIPT_BUILD = "2026-08-24-agent-projects"

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

try:
    import pythoncom
    import win32com.client
except ImportError:
    pythoncom = None  # type: ignore[assignment]
    win32com = None  # type: ignore[assignment]

RULES = """
Правила оформления документов СНиОТ (sniot-di-documents.mdc)
============================================================

Область: ВСЕ документы СНиОТ — ДИ, РИ, Положения, инструкции по эксплуатации и т.п.
Правила агента и этот скрипт — **важнее** любого docx-образца на N:\\.
Каждый пункт → функция(и) в fix_sniot_document.py (DocAgent — п. DocAgent ниже).

--- ТИТУЛ И СТРУКТУРА ---

 1. Приоритет правил над образцом docx на N:\\
    validate_sniot_document — проверка по mdc, не по образцу.

  2. Титул один раз на 1-й странице (sdt / шапка+УТВЕРЖДАЮ); глава 1 — со 2-й
    remove_duplicate_body_title, ensure_title_page_separated (nextPage + titlePg,
    сжатие пустых строк титула), validate_title_page_separated.
    Геометрия титула (константы TITLE_*): шапка по центру, TNR 14 не жирный.
    Таблица 2 колонки без рамок: слева название + подчёркнутая линия +
    «номер инструкции» 12 pt; справа «УТВЕРЖДАЮ» (не жирный), должность,
    линия подписи + пробел + И.О.Фамилия (без табуляции во 2-м столбце),
    следующей строкой плашка даты без таба, keepLines («___»_______________ YYYY г.).
    «МИНСК 2026» внизу 1-й по центру в ТЕЛЕ (framePr vAnchor=margin), не в колонтитуле;
    сразу после «МИНСК» — пустой абзац с разрывом раздела nextPage;
    глава 1 начинается с самого верха стр. 2 (не пустая строка в начале стр. 2).
    После шапки предприятия (в т.ч. «МИНСККОММУНТЕПЛОСЕТЬ») перед таблицей грифа
    — 6–8 пустых строк (цель 6). Остальные промежутки титула — не больше одной.
    format_title_block, ensure_title_stamp_gap_after_header,
    place_title_city_year_at_bottom, validate_title_block, validate_title_stamp_gap,
    validate_title_city_year.
    Не подменять шапку предприятия / уникальное название ДИ (ЛСиМ, Романовский)
    текстом ТКП или чужим титулом. restore_org_header_if_npa_inserted,
    find_agent_sibling_source. Нумерацию 1.8/1.9 не навязывать, если это не ДИ САТП.

 3. «ОБЩИЕ ПОЛОЖЕНИЯ» → «1 ОБЩИЕ ПОЛОЖЕНИЯ»
    normalize_first_chapter_heading.

 4. Безопасное remove_duplicate_body_title — НИКОГДА не стирать тело без маркера главы
    remove_duplicate_body_title, apply_sniot_rules_to_file (откат при body_after_clean==0).

--- ШРИФТ И АБЗАЦЫ ---

  5. Шрифт Times New Roman 14 pt везде; тело включая номера пунктов (1.1., 2.2.1.) —
    обычное начертание, все runs абзаца с одним rPr (ascii/hAnsi/eastAsia/cs, sz 28, не жирный).
    normalize_document_fonts, apply_run_font, set_paragraph_text, validate_fonts.
    Не оставлять жёлтый маркер / заливку абзацев (w:highlight, w:shd на pPr/rPr).
    Не подсвечивать «исправленные» номера и фрагменты, которых не было в исходнике.
    Исключение: если лист ознакомления — для другой должности, чем инструкция,
    на заголовок листа ставится красная заливка абзаца (w:shd fill=FF0000).
    strip_visual_highlights, mark_acquaintance_heading_if_job_mismatch,
    validate_visual_highlights. Жирные только заголовки глав 1–5.

 6. Тело (не заголовки глав 1–5, не подписанты): выравнивание **по ширине**
    apply_body_paragraph_format, ensure_paragraph_justified,
    validate_body_paragraph_format, is_paragraph_justified.

 7. Отступ первой строки **1,25 см** (709 twips) на теле документа, кроме титула,
    заголовков глав («1 ОБЩИЕ ПОЛОЖЕНИЯ» … «5 ОТВЕТСТВЕННОСТЬ»), подписантов
    **и ячеек таблиц**. В таблицах отступ первой строки **не применяется** (0).
    Строки тела «1.8. …», «1.9. … должен знать:», «2.1. …» — **с** отступом 1,25 см
    (это не названия глав). apply_body_paragraph_format, apply_table_paragraph_no_indent,
    ensure_first_line_indent, first_line_indent_cm,
    validate_body_paragraph_format, validate_section_headers.

 8. Названия глав 1–5 (**по словам пользователя — «заголовки разделов»**): **по центру**,
    жирный, капс. Примеры: «1 ОБЩИЕ ПОЛОЖЕНИЯ», «2 ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ».
    Шаблон «N» или «N.» + КАПС, **без подномера** (не «1.1.», не «1.9. должен знать:»).
    center_chapter_headers, canonical_chapter_header, ensure_paragraph_centered,
    restore_chapter_headers, is_chapter_header, apply_header_paragraph_format,
    validate_chapter_headers.

 9. Перед заголовком главы — ровно одна пустая строка; после — **не добавлять** (только сжатие хвоста)
    ensure_chapter_header_spacing, remove_extra_empty_lines_in_body,
    needs_signatory_layout_compression, validate_chapter_headers.

    9a. Заголовок главы **никогда** не сирота внизу страницы без текста следом
    prevent_chapter_header_orphan, validate_chapter_header_orphan.
    page_break_before — на первый абзац текста (или пустую строку перед ним), **не** на заголовок.

    9b. Таблица: первые 3 строки всегда вместе (верхняя не сирота внизу страницы);
    остальные строки можно переносить. Короче 3 строк — все строки вместе.
    Не ставить разрыв страницы между заголовком главы и таблицей сразу под ним.
    keep_table_header_rows_together, validate_table_header_rows_together.
    В ячейках таблиц **нет** абзацного отступа первой строки (1,25 см).
    apply_table_paragraph_no_indent, validate_table_paragraph_indents.

--- НУМЕРАЦИЯ ПУНКТОВ (ДИ САТП) ---

10. Нумерацию исходника **не удалять и не подменять схемой САТП (1.8/1.9)**.
    Только упорядочить и в финале сверить набор номеров исходника с результатом.
    Если пункты уже пронумерованы в тексте или списком Word (ListString / numPr) —
    сохранить те же номера и последовательность; чинить только ошибки:
    `.5.1.1.`, пропуски после потери автономера, двойной numPr+текст, шрифт номера.
    Запрещено снимать 1.1–1.7, менять 2.x на другую схему, вписывать номера,
    которых не было в исходнике. Если нумерация исходника целая — не трогать
    (как gate validate_numbering_blocks). ДИ САТП старшего мастера — по-прежнему
    1.8/1.9 только для этой ДИ (has_di_satp_numbering).
    materialize_word_decimal_numbering, collect_number_tokens,
    validate_source_numbering_preserved, validate_numbering_blocks,
    analyze_numbering_block, fix_numbering_selective.

11. Заголовки разделов («1.9. Старший мастер должен знать:») — без подномера в списке ниже
    find_section_header_index (regex ^N.N.\\s+[А-Я]), is_section_header, renumber_block.

12. Блокировка сохранения, если нумерация/тело резко пропали
    validate_save_integrity, count_numbered_paragraphs, count_nonempty_body_paragraphs.

--- СТРАНИЦЫ ---

13. Нумерация страниц: со 2-й, верх по центру; все PAGE — TNR 14 не жирные
    (одинаковый rPr в header1/header2/header3); titlePg, пустой first-header, футеры без номеров.
    fix_page_numbering, validate_page_numbering.

14. Переносы: шаг 1 natural (снять keep/page_break); шаг 2 только если подписанты оторваны
    fix_last_pages_and_signatories (mode=natural), fix_last_pages_page_breaks,
    clear_paragraph_page_layout, apply_signatory_page_break, prevent_chapter_header_orphan,
    validate_page_layout_flags, validate_chapter_header_orphan, validate_last_two_pages_layout.
    Запрещён разрыв **только** перед «Разработал:»; запрещён разрыв **только** на заголовке главы.
    Предпоследняя страница **может** содержать подписантов вместе с хвостом текста —
    не ставить принудительный page_break, если блок туда помещается. Лист ознакомления
    на ту же страницу не ставить. ensure_acquaintance_sheet_separate_page.

15. Стратегия переноса (шаг 2, только если подписанты оторваны):
    гл. 5 небольшая → page_break_before на **первый абзац текста** гл. 5 (не на заголовок);
    гл. 5 большая → page_break_before на абзац(ы) в гл. 5 + подписи следом;
    рвётся один длинный абзац → page_break_before на этот абзац.

--- ПОДПИСАНТЫ ---

16. «Согласовано:» с двоеточием, **не жирным**; маркеры в любом ДИ, не только САТП.
    Если в исходнике вместо блока — таблица должностей и фамилий: первая строка →
    «Разработал:» + должность+ФИО (Tab); далее пустая; «Согласовано:»; остальные
    строки. Интервал 1,5; по одной пустой перед каждым маркером; без пустой после.
    Таблицу согласования после переноса **убрать** (не оставлять дубль таблица+блок).
    Лист ознакомления (фраза «Настоящую инструкцию изучил…», таблица ФИО/роспись/дата)
    **не** превращать в подписантов; всегда отдельная страница (page_break_before).
    Подписанты — таблица согласования в конце тела, **до** листа ознакомления.
    fix_soglasovano, materialize_signatory_paragraphs_from_tables,
    ensure_acquaintance_sheet_separate_page, validate_acquaintance_sheet,
    validate_signatory_block, validate_agreement_table_converted.

17. Ровно одна пустая строка **перед** «Разработал:» и **перед** «Согласовано:»
    ensure_single_empty_line_before, validate_signatory_block.

18. **Без** пустой строки **после** «Разработал:» и **после** «Согласовано:»
    remove_empty_paragraphs_after_marker, fix_signatory_block_format,
    validate_signatory_block.

19. Межстрочный интервал **1,5** только на блоке подписантов (от «Разработал:» до последней подписи)
    apply_signatory_line_spacing, set_one_point_five_line_spacing, paragraph_has_one_point_five_spacing,
    fix_signatory_block_format, validate_signatory_line_spacing, validate_signatory_block.
    Тело документа и заголовки глав — **не** 1,5; после align_spacing_to_etalon интервал подписантов восстанавливается.

20. Подписанты всегда связаны с текстом — не сироты на отдельной странице без текста
    fix_last_pages_and_signatories, fix_last_pages_page_breaks, signatories_appear_orphaned,
    validate_last_two_pages_layout, validate_page_layout_flags.

20a. Конец документа — шаг 1 (всегда): естественная вёрстка предпоследней/последней страниц
    fix_last_pages_and_signatories(mode=natural): снять разрывы на теле; оформить подписантов 1,5;
    пустые только перед «Разработал:»/«Согласовано:», без пустых между согласующими;
    keepNext от последнего абзаца текста до конца блока (не рвать и не отрывать от текста).

20b. Конец документа — шаг 2 (только если оторваны): fix_last_pages_page_breaks / --fix-page-breaks
    Перенос абзаца(ов) или гл. 5 целиком; **никогда** разрыв только перед «Разработал:».
    force/--fix-page-breaks **не** означает «всегда ставить разрыв»: если XML не видит
    отрыв, оставить естественную вёрстку (подписанты могут остаться на предпоследней).

--- DOCAGENT / ПУТЬ ---

21. Консервативный режим ДИ/РИ/положений/ИОТ СНиОТ: без text_edits и перестройки
    содержания (нумерацию исходника не сдвигать). Проверка русского языка — для
    всех документов (галочка DocAgent по умолчанию включена); аббревиатуры
    (ЛСиМ, СНиОТ, САТП, ТКП, НПА, ТНПА, ЛПА, ФИО, РТС, ЛЭС, АТП, ОК, ЮО, ООТиЗ
    и 2–6 букв заглавного/смешанного регистра) спеллер не проверяет и не правит.
    is_conservative_di_satp (DocAgent: sniot_document.py, agent_core).

22. Кнопка «Оформить документ» — финальный проход fix_sniot_document с always_apply
    apply_sniot_rules_to_file (--always-apply), agent_core.apply_sniot_rules_to_output.

23. Путь из поля «1. Документ» / handoff, не из чата Cursor
    resolve_from_handoff, resolve_target, DOCAGENT_HANDOFF.

24. Латинская «i» в «СНiОТ» в путях → кириллическая «и»
    normalize_sniot_path_text, resolve_target (alt path).

--- ПОРЯДОК ОБРАБОТКИ ---

25. process_sniot_document:
    title (remove_duplicate_body_title до Document) → numbering → font → chapters (center)
    → body (justify + 1,25 см) → spacing → fix_last_pages_and_signatories (шаг 1)
    → fix_last_pages_page_breaks при --fix-page-breaks или signatories_appear_orphaned (шаг 2)
    → keep_table_header_rows_together (все документы)
    → fix_page_numbering после save.
    → place_title_city_year_at_bottom, ensure_acquaintance_sheet_separate_page.
    → Word COM после save: GUI XML-проход --skip-word (чтобы таймаут не убил запись),
    затем apply_sniot_rules_to_output → apply_word_grammar_check
    (apply_word_grammar_via_com: SpellingErrors + CheckGrammar);
    страницы подписантов inspect_signatory_pages_via_word.

36. После XML-проверки — тихо открыть документ в Word (не показывать окно):
    исправить слова с красной чертой (SpellingErrors + GetSpellingSuggestions), кроме аббревиатур;
    спросить номера страниц хвоста и «Разработал:».
    Если «Разработал:» на странице позже, чем хвост текста — подписанты оторваны: шаг 2 переноса.
    Нет Word / файл занят — предупреждение, XML-проверку не отменяет.
    inspect_signatory_pages_via_word, interpret_word_signatory_layout,
    apply_word_grammar_via_com, apply_sniot_rules_to_file.

--- ПРАВИЛА РУССКОГО ЯЗЫКА (канцелярские фразы) ---

26. text_edits и перенумерация исходника — не включать. apply_russian_phrase_rules
    в process() не вызывать (коды ТКП / сдвиг 1.1→1.4 запрещены).
    Проверка русского (DocAgent russian_check) — для всех документов; аббревиатуры
    не правятся. Канцелярские фразы ТКП — только внутри проверки русского, не text_edits.
    По указанию пользователя (не text_edits): в перечень «… должен знать» добавить
    «основы делопроизводства», если строки нет (заголовок заканчивается на «должен знать:»,
    не фраза «должен знать требования…»). Пунктуация перечисления: промежуточные «;»,
    последний «.» (если вставка стала последним пунктом — точка, не точка с запятой).
    Пример: «1.9.15. коллективный договор;» / «1.9.16. основы делопроизводства.»
    Первое «ЛПА» без расшифровки → «локальным правовым актом (далее - ЛПА)»
    (падеж/число по контексту); дальше «ЛПА» не расшифровывать повторно.
    Спеллер по-прежнему не правит аббревиатуру ЛПА.
    ensure_deloproizvodstvo_in_must_know, normalize_must_know_list_punctuation,
    fix_duty_by_order_commas, expand_first_lpa_abbreviation
    (вызов из apply_mandatory_layout_fixes).
    После любой правки — обязательно проверить пунктуацию ВСЕГО документа
    (титул, тело, таблицы, подписанты) и пунктуацию перечислений «должен знать».
    Без этой проверки документ не «Готово».
    check_document_punctuation_after_edit, validate_document_punctuation.

--- ИНСТРУКЦИЯ ПО ДЕЛОПРОИЗВОДСТВУ 2025 (Минюст РБ) ---

27. Формат А4 (210×297 мм), поля: левое 30 мм, правое 10 мм, верх/низ 20 мм.
    Минюст №65 п.18 — правое ≥8 мм; для документов СНиОТ этого агента правое 10 мм
    (указание пользователя). apply_page_setup_deloproizvodstvo, validate_page_margins.

28. Текст на А4 — одинарный межстрочный интервал (1.0), не 1,15/1,5; интервал перед/после
    абзаца тела = 0 pt (не «зазор» между строками). Заголовки глав — без space_after.
    Блок подписантов — межстрочный 1,5. Общее правило для всех агентов (DocAgent, этот
    скрипт, еженедельный итог). apply_body_single_line_spacing, validate_body_line_spacing,
    set_single_line_spacing, apply_signatory_line_spacing.

29. Выравнивание текста по ширине; абзацный отступ 12,5 мм (гл.7, п.363, 377) — см. п.7–8.
    apply_body_paragraph_format (1,25 см = 12,5 мм).

30. Нумерация листов: со 2-й, верх по центру, TNR 14 pt (гл.7, п.396–398)
    fix_page_numbering, validate_page_numbering, CENTERED_PAGE_HEADER.

31. Точка перед номером пункта (.1.4.1.) — убрать; двойная точка после номера — убрать
    strip_leading_dot_before_numbering, validate_leading_dot_before_numbering.

32. Названия разделов и глав организационных документов — прописными (гл.7, п.395)
    center_chapter_headers, restore_chapter_headers, validate_chapter_headers.

33. Лишние/ошибочные символы в тексте — убрать или нормализовать.
    Обязательно для любого типа (ДИ, РИ, положение, инструкция, еженедельный итог):
    не больше одного пробела подряд; nbsp→обычный пробел; validate ловит остаток.
    Тело — полная очистка; титул — сжать двойные пробелы; подписанты — пробелы сжать,
    Tab должность↔ФИО не ломать. strip_unnecessary_characters, sanitize_paragraph_text,
    validate_unnecessary_characters. Нулевая ширина, soft hyphen, таб→пробел (не у
    подписантов), trim, латинская o/i в «Согласовано»/«СНиОТ», «,,»/«..» (не «…»),
    пробел перед «.», кавычки «»; нумерация — через fix_erroneous_numbering_prefix.
    Пробел сразу после «(» и сразу перед «)» убрать:
    «продукции ( выполнение )» → «продукции (выполнение)».
    Пробел между словом и открывающей «(» **оставить**. Не ломать нумерацию 1.1.
    fix_spaces_around_parentheses.
    Ведущие маркеры списка (•, , ○, ■, дефис-список) — убрать (п. 36).
    Ведущие пробелы/табы перед номером пункта (  1.1. / \\t1.2.) — убрать.
    После номера (1.1. / 1.1.1. / заголовок «1 ОБЩИЕ») — ровно один пробел.
    normalize_item_number_spacing, normalize_number_separator,
    validate_item_number_spacing.
    Соседние одинаковые слова в одном абзаце («службы службы») — схлопнуть.
    collapse_adjacent_duplicate_words, validate_adjacent_duplicate_words.

34. **Обязательно перед записью в папку Агент:** перечитать весь docx с диска (не объект из памяти)
    и validate_sniot_document по **всем** правилам mdc; при любом замечании — **не перезаписывать**
    *_оформлен.docx пользователя (запись отменена, черновик остаётся в _work_)
    reload_document_from_path, validate_full_document_on_disk, apply_sniot_rules_to_file.

    35. Финальная перечитка обязана ловить: нет видимых 1.1.–1.7. в гл. 1; блок 1.8.x не с 1.8.1.;
    блок 1.9.x не с 1.9.1.; дубль «Разработал:»; отступы; оторванные подписанты.
    validate_final_document_gate, validate_numbering_block_starts, validate_chapter_one_intro_numpr.

    36. Никаких маркеров списка Word (все агенты, все типы: ДИ, РИ, положение, инструкция,
    еженедельный итог): не оставлять • /  / ○ / ■ / дефис-списки / w:numPr bullet.
    Только нумерованные пункты (1.1., 2.2.1.) или обычный текст. Не путать с номерами
    страниц, префиксом «1.4.1.» и табуляцией подписантов.
    strip_leading_list_marker_text, remove_list_markers_in_body,
    remove_word_list_numbering_in_body (bullet без ручного номера),
    validate_list_markers.

    37. Таблицы — шапка не сирота: первые 3 строки вместе (cantSplit + tblHeader + keepNext 1→2→3);
    хвост таблицы можно на следующую страницу. Для всех ДИ/РИ/положений/инструкций.
    keep_table_header_rows_together, validate_table_header_rows_together.

    38. Город-год на титуле — канон «МИНСК 2026» (прописные, пробел; год из исходника).
    Низ 1-й страницы, по центру, без отступа 1,25; framePr yAlign=bottom
    (TITLE_CITY_YEAR_*). Не дублировать на стр. 2.
    place_title_city_year_at_bottom, validate_title_city_year.

    39. Лист ознакомления — всегда отдельная страница с самого верха
    (page_break_before на первый абзац). Фразы: «Настоящую инструкцию изучил…»
    и «С должностной инструкцией … ознакомлен». Не путать с «Разработал:».
    Если должность в листе ознакомления не совпадает с инструкцией — красная
    заливка только заголовка листа. Совпадение — красный снять.
    is_acquaintance_sheet_text, ensure_acquaintance_sheet_separate_page,
    mark_acquaintance_heading_if_job_mismatch, validate_acquaintance_sheet,
    validate_acquaintance_job_mismatch, _is_acquaintance_sheet_table.

    40. После XML — тихая проверка Word: красная волнистая черта = орфография (SpellingErrors).
    Для каждого Range взять GetSpellingSuggestions и заменить слово, кроме аббревиатур
    (ЛСиМ, СНиОТ, ТКП, НПА, 2–6 букв Caps/лесенка) — их не трогать, даже если Word подчёркивает.
    Зелёная грамматика: CheckGrammar / GrammaticalErrors, без pymorphy. Не IgnoreAll на весь документ.
    Копия TEMP, затем сохранить в pipeline. Word занят — предупреждение, не падать.
    apply_word_grammar_via_com, run_word_grammar_check_subprocess.

    41. Плашка даты у подписантов («___» ________ 20__ г.) — новая строка,
    выровнять с началом И.О.Ф. (ведущий Tab, позиция таба = таб строки должности),
    без переноса на две строки (укоротить подчёркивание; keepLines на плашке).
    Должность и ФИО — Tab только на одной строке; ФИО не переносить
    (keepLines, короче линия, неразрывные пробелы в И.О.Ф.). Не трогать дату «УТВЕРЖДАЮ».
    apply_signatory_fio_one_line, fix_signatory_date_plaques, apply_signatory_line_spacing,
    validate_signatory_date_plaques, validate_signatory_fio_one_line.

Коды выхода: 0 OK, 1 ошибки валидации, 2 файл занят Word, 3 файл не найден.
"""

_TNR14_NOT_BOLD_RPR = (
    "<w:rPr>"
    '<w:rFonts w:ascii="Times New Roman" w:eastAsia="Times New Roman" '
    'w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
    '<w:b w:val="false"/><w:bCs w:val="false"/>'
    '<w:sz w:val="28"/><w:szCs w:val="28"/>'
    "</w:rPr>"
)

EMPTY_FIRST_HEADER = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:p/>
</w:hdr>"""

CENTERED_PAGE_HEADER = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    "<w:p>"
    f"<w:pPr><w:jc w:val=\"center\"/>{_TNR14_NOT_BOLD_RPR}</w:pPr>"
    f"<w:r>{_TNR14_NOT_BOLD_RPR}<w:fldChar w:fldCharType=\"begin\"/></w:r>"
    f"<w:r>{_TNR14_NOT_BOLD_RPR}<w:instrText xml:space=\"preserve\"> PAGE </w:instrText></w:r>"
    f"<w:r>{_TNR14_NOT_BOLD_RPR}<w:fldChar w:fldCharType=\"separate\"/></w:r>"
    f"<w:r>{_TNR14_NOT_BOLD_RPR}<w:t>2</w:t></w:r>"
    f"<w:r>{_TNR14_NOT_BOLD_RPR}<w:fldChar w:fldCharType=\"end\"/></w:r>"
    "</w:p></w:hdr>"
).encode("utf-8")

EMPTY_FOOTER = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:p/>
</w:ftr>"""

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
HEADER_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
)
HEADER_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
)
NS = {"w": W_NS}
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
FONT_SIZE_EMU = int(FONT_SIZE)
NUM_PREFIX = re.compile(r"^\d+(?:\.\d+)+\.\s*")
LEADING_DOT_BEFORE_NUMBER = re.compile(r"^\s*\.\s*(\d+(?:\.\d+)+\.)\s*")
DOUBLE_DOT_AFTER_NUMBER = re.compile(r"^(\d+(?:\.\d+)+\.)\.(?:\.|\s*)")
SECTION_HEADER = re.compile(r"^\d+\.\d+\.\s+[А-ЯЁA-Z]")
CHAPTER_HEADER = re.compile(r"^\d+\.?\s+[А-ЯЁA-Z][А-ЯЁA-Z\s\-—,«»]+$")
CHAPTER_HEADING_KEYWORDS = (
    "ОБЩИЕ ПОЛОЖЕНИЯ",
    "ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ",
    "ФУНКЦИИ",
    "ПРАВА",
    "ВЗАИМООТНОШЕНИЯ",
    "ВЗАИМОДЕЙСТВИ",
    "ОТВЕТСТВЕННОСТЬ",
)

# Глава 1 по образцу Word (ListString): 1.1–1.7 вводные, затем 1.8 руководствуется, 1.9 должен знать.
# Нельзя ставить повторные «1.4.» / «1.5.» — это откатывает принятую нумерацию.
DI_SATP_CH1_GUIDED = "1.8"
DI_SATP_CH1_KNOW = "1.9"

DI_SATP_SECTION_HEADERS: tuple[tuple[str, str], ...] = (
    (DI_SATP_CH1_GUIDED, "в своей деятельности руководствуется"),
    (DI_SATP_CH1_KNOW, "должен знать"),
    ("2.1", "выполняет следующие функции"),
    ("2.2", "Для выполнения возложенных на него функций"),
    ("3.1", "имеет право"),
    ("5.1", "несет ответственность"),
)

DI_SATP_CHAPTER_ONE_NUMBERED: tuple[tuple[str, str], ...] = (
    ("1.1", "Настоящая должностная инструкция определяет"),
    ("1.2", "Настоящую должностную инструкцию должны знать"),
    ("1.3", "На должность старшего мастера назначается"),
)

DI_SATP_CHAPTER_ONE_NUMPR_MARKERS: tuple[str, ...] = (
    "Назначение, перемещение и увольнение",
    "Старший мастер подчиняется",
    "Старшему мастеру подчиняются",
    "Старший мастер проходит первичную проверку",
)

SIGNATORY_NAME_TAIL = re.compile(
    r"(\s)([А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁ][а-яё\-]+(?:\s+[А-ЯЁ][а-яё\-]+)?)\s*$"
)
SIGNATORY_IOF_ONLY_RE = re.compile(
    r"^[_\s«»\"']*([А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁ][а-яё\-]+(?:\s+[А-ЯЁ][а-яё\-]+)?)\s*$"
)


def configure_stdio_utf8() -> None:
    """Windows: stdout subprocess и GUI читают UTF-8 без «ромбиков»."""
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except Exception:
                pass
FIRST_LINE_INDENT_CM = 1.25
FIRST_LINE_INDENT_TWIPS = 709
FIRST_LINE_INDENT_TOLERANCE_CM = 0.08
SIGNATORY_LINE_SPACING = 1.5
SIGNATORY_LINE_SPACING_TOLERANCE = 0.05
DEFAULT_SIGNATORY_TAB_TWIPS = "6804"
MIN_SIGNATORY_TAB_TWIPS = 6803  # ≥12 см (1 см ≈ 567 twips; 6803 — округление Word)
MISSING_DEVELOPER_MARKER = "КТО???"
MISSING_DEVELOPER_HIGHLIGHT = "yellow"
MISSING_DEVELOPER_FILL = "FFFF00"
# Титул: снято с правки пользователя 18.08.26 (ведущий инженер ОМТС, только чтение).
# Путь ОБМЕН не образец DocAgent и не master_sample_path — константы здесь.
# Шапка по центру; таблица 2 колонки без рамок:
# слева название + линия номера + «номер инструкции» 12 pt;
# справа «УТВЕРЖДАЮ», должность, линия+И.О.Ф., плашка даты keepLines.
# «МИНСК YYYY» внизу 1-й страницы по центру в теле (framePr), не header/footer.
TITLE_ORG_ALIGN = "center"
TITLE_STAMP_ALIGN = "right"
TITLE_STAMP_IN_TABLE_ALIGN = "left"
TITLE_DOC_NAME_ALIGN = "left"
TITLE_CITY_YEAR_ALIGN = "center"
TITLE_UTVERZHDAYU_BOLD = False
TITLE_RIGHT_TAB_TWIPS = "9026"
TITLE_STAMP_RIGHT_INDENT_TWIPS = "-1333"
TITLE_EMPTY_BETWEEN_BLOCKS = 1
MAX_TITLE_EMPTY_RUN = TITLE_EMPTY_BETWEEN_BLOCKS
# После таблицы грифа и перед «МИНСК YYYY» — пустые строки (не разрыв страницы).
TITLE_EMPTY_AFTER_STAMP_BEFORE_CITY = TITLE_EMPTY_BETWEEN_BLOCKS
# После шапки предприятия (в т.ч. «МИНСККОММУНТЕПЛОСЕТЬ») перед таблицей грифа.
TITLE_EMPTY_BEFORE_STAMP = 6
TITLE_EMPTY_BEFORE_STAMP_MIN = 6
TITLE_EMPTY_BEFORE_STAMP_MAX = 8
# Геометрия грифа — с правки пользователя 18.08.2026 (ведущий инженер ОМТС).
# Правая колонка ~5212 twips: табуляцию во 2-м столбце не ставить (ломает ФИО/дату).
TITLE_STAMP_TABLE_COL_TWIPS = ("4967", "5212")
TITLE_STAMP_TABLE_WIDTH_TWIPS = "10179"
TITLE_STAMP_TABLE_IND_TWIPS = "-431"
# В правом столбце грифа табуляция запрещена (переносит ФИО/дату и остаётся в XML).
TITLE_STAMP_DATE_TEMPLATE = "«___»_______________ {year}г."
TITLE_NUMBER_LABEL = "номер инструкции"
TITLE_NUMBER_FONT_PT = 12
TITLE_NUMBER_FONT_HALF = "24"
TITLE_UNDERLINE_LINE = "____________________"
TITLE_IOF_RE = re.compile(
    r"([А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁ][а-яё\-]+(?:\s+[А-ЯЁ][а-яё\-]+)?)"
)
TITLE_CITY_YEAR_TEMPLATE = "МИНСК {year}"
TITLE_CITY_YEAR_FRAME_ATTRS: tuple[tuple[str, str], ...] = (
    ("w:wrap", "none"),
    ("w:vAnchor", "margin"),
    ("w:hAnchor", "margin"),
    ("w:xAlign", "center"),
    ("w:yAlign", "bottom"),
)
# Инструкция по делопроизводству 2025, п.18 (лист А4): левое 30, правое ≥8, верх/низ 20.
# Для документов СНиОТ этого агента правое поле 10 мм (указание пользователя).
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
MARGIN_LEFT_MM = 30
MARGIN_RIGHT_MM = 10
MARGIN_TOP_MM = 20
MARGIN_BOTTOM_MM = 20
MARGIN_TOLERANCE_MM = 1.5
# Первые N строк любой таблицы держать вместе (не сирота-шапка).
TABLE_HEADER_KEEP_ROWS = 3
DELOPROIZVODSTVO_REF = (
    "Инструкция по делопроизводству в государственных органах, иных организациях "
    "(утв. Минюстом РБ 29.08.2025 № 65)"
)
TITLE_DUPLICATE_MARKERS = ("УТВЕРЖД", "ДОЛЖНОСТН", "номер инструкции", "Минсккоммунтеплосеть")
ORG_HEADER_TEPLOSETI_LINE = "КОММУНАЛЬНЫХ ТЕПЛОВЫХ СЕТЕЙ И КОТЕЛЬНЫХ"
NPA_ON_TITLE_RE = re.compile(r"(?i)^\s*(ТКП\b|ТР\s|ГОСТ\b|СНиП\b|Правила\s|«Правила)")
SATP_STRUCTURE_HINTS = (
    "старший мастер",
    "сатп",
)
HIGHLIGHT_FILL_HEX = frozenset(
    {
        "FFFF00",
        "00FF00",
        "00FFFF",
        "FF00FF",
        "FF0000",
        "0000FF",
        "YELLOW",
        "GREEN",
        "CYAN",
        "MAGENTA",
        "RED",
        "BLUE",
        "D9D9D9",
        "C0C0C0",
        "E7E6E6",
        "F2F2F2",
        "A6A6A6",
        "BFBFBF",
        "D0D0D0",
    }
)
ACQUAINTANCE_MISMATCH_FILL = "FF0000"
SIGNATORY_DATE_PLAQUE_RE = re.compile(
    r'[«"„]?\s*_{2,}\s*[»"]?\s*_{2,}\s*20[_0-9]{2,4}\s*г\.?',
    re.IGNORECASE,
)
_ADJACENT_DUP_WORD_RE = re.compile(
    r"(?<![0-9А-Яа-яЁёA-Za-z])"
    r"([А-Яа-яЁёA-Za-z][А-Яа-яЁёA-Za-z0-9\-]*)"
    r"(?:[ \t\xa0]+\1)+"
    r"(?![0-9А-Яа-яЁёA-Za-z])",
    re.IGNORECASE,
)
_ITEM_NUMBER_LINE_RE = re.compile(
    r"^[\s\t]*(\d+(?:\.\d+)+)\.?([ \t]*)(.*)$",
    re.DOTALL,
)
_CHAPTER_NUMBER_LINE_RE = re.compile(
    r"^[\s\t]*(\d+\.?)[ \t]+(.*)$",
    re.DOTALL,
)
_JOB_FROM_APPOINTMENT_RE = re.compile(
    r"на должность\s+(.+?)\s+назначается",
    re.IGNORECASE,
)
_JOB_FROM_ACQ_RE = re.compile(
    r"с\s+должностной\s+инструкцией\s+(.+?)\s+ознакомлен",
    re.IGNORECASE | re.DOTALL,
)
_JOB_FROM_TITLE_RE = re.compile(
    r"должностн(?:ая|ой)\s+инструкци[яи]\s+(.+)",
    re.IGNORECASE,
)

DEFAULT_TARGET_NAME = "ПРОЕКТ Старший мастер_оформлен.docx"
DEFAULT_TARGET_PLUS_NAME = "ПРОЕКТ Старший мастер_оформлен+.docx"
_USER_BASE_N = r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В"
_USER_BASE_UNC = r"\\srv-data\doc\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В"
USER_AGENT_DIR = Path(_USER_BASE_N + r"\Агент")
USER_AGENT_DIR_UNC = Path(_USER_BASE_UNC + r"\Агент")
USER_PROJECT_DIR = Path(_USER_BASE_N + r"\Проекты")
USER_PROJECT_DIR_UNC = Path(_USER_BASE_UNC + r"\Проекты")
WRITABLE_USER_DIRS: tuple[Path, ...] = (
    USER_AGENT_DIR,
    USER_AGENT_DIR_UNC,
    USER_PROJECT_DIR,
    USER_PROJECT_DIR_UNC,
)
_WRITABLE_FOLDER_NAMES = frozenset({"агент", "проекты", "проект"})
# ОБМЕН / САТП больше не источник образца. Пусто, чтобы старый обход не ходил туда.
READONLY_SAMPLE_DIRS: tuple[Path, ...] = ()
SAMPLE_NAME_MARK = "образец"
SAMPLE_EXTENSIONS = {".docx"}
WORK_DIR = Path(r"C:\Users\v.dubovik\AttestationSync")
DOCAGENT_HANDOFF = Path(r"C:\Users\v.dubovik\DocAgent\handoff\request_latest.json")
DOCAGENT_FORMATTERS = Path(r"C:\Users\v.dubovik\DocAgent\formatters")
try:
    _docagent_root = str(DOCAGENT_FORMATTERS.parent)
    if _docagent_root not in sys.path:
        sys.path.insert(0, _docagent_root)
    from formatters.word_com import convert_legacy_word_to_docx as convert_legacy_word_to_docx
except Exception:
    convert_legacy_word_to_docx = None  # type: ignore[assignment]
WORK_FILE = WORK_DIR / "_work_sniot_document.docx"
OUT_FILE = WORK_DIR / "_work_sniot_document_fixed.docx"
SENIOR_MASTER_DUMP = WORK_DIR / "_work_senior_master_fixed.txt"
MIN_NUMBERED_SENIOR_MASTER = 80

CHAPTER_TITLES: tuple[tuple[int, str], ...] = (
    (1, "ОБЩИЕ ПОЛОЖЕНИЯ"),
    (2, "ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ"),
    (3, "ПРАВА"),
    (4, "ВЗАИМООТНОШЕНИЯ"),
    (5, "ОТВЕТСТВЕННОСТЬ"),
)


def paragraph_text_normalized(paragraph: Paragraph) -> str:
    """Текст абзаца без nbsp/нулевой ширины — для проверки «пустой» строки."""
    return (paragraph.text or "").replace("\xa0", " ").replace("\u200b", "").strip()


def is_paragraph_empty(paragraph: Paragraph) -> bool:
    return not paragraph_text_normalized(paragraph)


def find_body_start_index(doc: Document) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph_text_normalized(paragraph)
        if not text:
            continue
        upper = text.upper()
        if (
            is_chapter_header(text)
            or upper.startswith("1 ОБЩИЕ")
            or upper.startswith("ОБЩИЕ ПОЛОЖЕНИЯ")
            or is_unnumbered_chapter_title(text)
        ):
            return i
    return 0


def _shd_fill_value(shd) -> str:
    return (shd.get(qn("w:fill")) or shd.get(qn("w:val")) or "").upper()


def _xml_belongs_to_missing_developer_marker(el) -> bool:
    """Жёлтый маркер «КТО???» — не снимать (исключение из strip_visual_highlights)."""
    node = el
    while node is not None:
        if node.tag in (qn("w:r"), qn("w:p")):
            blob = "".join((t.text or "") for t in node.iter(qn("w:t")))
            if MISSING_DEVELOPER_MARKER in blob:
                return True
            if node.tag == qn("w:p"):
                break
        node = node.getparent()
    return False


def strip_visual_highlights(doc: Document) -> int:
    """
    Снять жёлтый маркер и заливку абзацев, которую Word показывает как «выделение».
    Ячейки таблиц: только цветной маркер, не служебный auto/белый.
    Исключение: жёлтый фон маркера «КТО???» в блоке «Разработал:».
    """
    changed = 0
    root = doc.element
    for hl in list(root.iter(qn("w:highlight"))):
        if _xml_belongs_to_missing_developer_marker(hl):
            continue
        parent = hl.getparent()
        if parent is not None:
            parent.remove(hl)
            changed += 1
    for shd in list(root.iter(qn("w:shd"))):
        parent = shd.getparent()
        if parent is None:
            continue
        if _xml_belongs_to_missing_developer_marker(shd):
            continue
        parent_tag = parent.tag.split("}")[-1]
        fill = _shd_fill_value(shd)
        if parent_tag == "tcPr":
            if fill in HIGHLIGHT_FILL_HEX:
                parent.remove(shd)
                changed += 1
            continue
        parent.remove(shd)
        changed += 1
    return changed


def title_has_foreign_npa_insert(doc: Document) -> bool:
    body_start = find_body_start_index(doc)
    for paragraph in doc.paragraphs[:body_start]:
        if NPA_ON_TITLE_RE.match(paragraph.text.strip()):
            return True
    return False


def restore_org_header_if_npa_inserted(doc: Document) -> int:
    """Вернуть строку шапки предприятия, если на титул попал абзац ТКП/Правил."""
    body_start = find_body_start_index(doc)
    title_blob = " ".join(p.text for p in doc.paragraphs[:body_start]).upper()
    if "МИНСККОММУНТЕПЛОСЕТЬ" not in title_blob and "КОММУНАЛЬНОЕ УНИТАРНОЕ" not in title_blob:
        return 0
    changed = 0
    for paragraph in doc.paragraphs[:body_start]:
        if not NPA_ON_TITLE_RE.match(paragraph.text.strip()):
            continue
        set_paragraph_text(paragraph, ORG_HEADER_TEPLOSETI_LINE)
        ensure_paragraph_centered(paragraph)
        clear_first_line_indent(paragraph)
        changed += 1
    return changed


def restore_lsim_if_osim_poisoned(doc: Document, path: Path | str | None) -> int:
    """Ядовитая замена ЛСиМ→Осим/Сим (спеллер) — откатить, если в имени файла есть ЛСиМ."""
    if path is None or "лсим" not in Path(path).name.lower():
        return 0
    changed = 0
    paragraphs = list(doc.paragraphs)
    paragraphs.extend(iter_table_cell_paragraphs(doc))
    for paragraph in paragraphs:
        original = paragraph.text
        new_text = original.replace("Сим Сим", "ЛСиМ").replace("Осим", "ЛСиМ")
        if new_text == original:
            continue
        set_paragraph_text(paragraph, new_text)
        changed += 1
    return changed


def convert_legacy_word_to_temp_docx(src: Path) -> Path:
    """Старый .doc/.rtf → временный docx в TEMP (не писать _converted в папку Агент)."""
    td = Path(tempfile.mkdtemp(prefix="sniot_src_"))
    dst = td / "source.docx"
    if convert_legacy_word_to_docx is None:
        raise RuntimeError("Нет модуля Word для чтения .doc/.rtf")
    return Path(convert_legacy_word_to_docx(src, dst))


def open_word_document_readonly(path: Path) -> Document:
    """Прочитать .docx напрямую; .doc/.rtf — через Word во временный docx."""
    suffix = path.suffix.lower()
    if suffix in {".doc", ".rtf"}:
        return Document(str(convert_legacy_word_to_temp_docx(path)))
    return Document(str(path))


def read_docx_bytes_any(path: Path) -> bytes:
    """Байты OOXML: .docx как есть, .doc/.rtf после конвертации в TEMP."""
    suffix = path.suffix.lower()
    if suffix in {".doc", ".rtf"}:
        return convert_legacy_word_to_temp_docx(path).read_bytes()
    return path.read_bytes()


def find_agent_sibling_source(target: Path) -> Path | None:
    """Исходник рядом в папке Агент/Проекты: .doc / .rtf / *_converted.docx / имя без _оформлен."""
    if not is_path_in_writable_user_dir(target):
        return None
    name = target.name
    if "_оформлен" not in name.lower():
        return None
    stem = re.sub(r"_оформлен.*$", "", name, flags=re.IGNORECASE)
    folder = target.parent
    for cand in (
        folder / f"{stem}_converted.docx",
        folder / f"{stem}.docx",
    ):
        try:
            if cand.is_file() and cand.resolve() != target.resolve():
                return cand
        except OSError:
            continue
    return None


def title_or_identity_looks_corrupted(doc: Document, path: Path) -> bool:
    if title_has_foreign_npa_insert(doc):
        return True
    name = path.name.lower()
    blob = " ".join(p.text for p in doc.paragraphs[:50])
    if "лсим" in name and "Осим" in blob:
        return True
    return False


def find_signatory_tail_start(doc: Document) -> int | None:
    """
    Начало блока подписантов: «Разработал:», иначе первая строка подписи перед «Согласовано».
    """
    for i, paragraph in enumerate(doc.paragraphs):
        upper = paragraph_text_normalized(paragraph).upper()
        if upper.startswith("РАЗРАБОТАЛ"):
            return i
    try:
        soglas_idx = find_soglasovano_index(doc)
    except ValueError:
        return None
    idx = soglas_idx - 1
    while idx >= 0 and is_paragraph_empty(doc.paragraphs[idx]):
        idx -= 1
    if idx < 0:
        return soglas_idx
    text = paragraph_text_normalized(doc.paragraphs[idx])
    upper = text.upper()
    if "\t" in doc.paragraphs[idx].text or "начальник" in upper or "инженер" in upper:
        return idx
    return soglas_idx


def get_body_spacing_end_index(doc: Document) -> int:
    """Верхняя граница (не включая) для схлопывания пустых строк в теле."""
    tail = find_signatory_tail_start(doc)
    return tail if tail is not None else len(doc.paragraphs)


@dataclass
class DocumentProfile:
    """Профиль документа — какие правила применять."""

    kind: str  # di | ri | polozhenie | generic
    first_chapter: str | None
    has_signatories: bool
    has_di_satp_numbering: bool
    tail_chapter_idx: int | None
    source_number_tokens: tuple[str, ...] = ()


def chapter_header_body(text: str) -> str:
    return re.sub(r"^\d+\.?\s*", "", text.strip()).upper()


def is_unnumbered_chapter_title(text: str) -> bool:
    """«ОБЩИЕ ПОЛОЖЕНИЯ» / «ПРАВА» без номера — заголовок главы, не пункт 1.8."""
    t = re.sub(r"\s+", " ", (text or "").replace("\xa0", " ").strip()).upper().rstrip(".")
    if not t or re.match(r"^\d+", t):
        return False
    return any(t == title for _, title in CHAPTER_TITLES)


def canonical_chapter_header(text: str) -> str:
    """«4. ВЗАИМОДЕЙСТВИЯ» → «4 ВЗАИМОДЕЙСТВИЯ»; чужие названия глав не подменять."""
    t = (text or "").replace("\xa0", " ").strip()
    folded = re.sub(r"\s+", " ", t.upper()).rstrip(".")
    for expected_num, title in CHAPTER_TITLES:
        if folded == title:
            return f"{expected_num} {title}"
    match = re.match(r"^(\d+)\.?\s+(.+)$", t)
    if not match:
        return t
    num = match.group(1)
    body = re.sub(r"\s+", " ", match.group(2).strip().rstrip(".")).upper()
    for expected_num, title in CHAPTER_TITLES:
        if body == title:
            return f"{expected_num} {title}"
    return f"{num} {body}"


def is_chapter_header(text: str) -> bool:
    t = (text or "").replace("\xa0", " ").replace("\u200b", "").strip()
    if not t or re.match(r"^\d+\.\d", t):
        return False
    if is_unnumbered_chapter_title(t):
        return True
    if CHAPTER_HEADER.match(t):
        return True
    match = re.match(r"^(\d+)\.?\s+(.+)$", t)
    if not match:
        return False
    try:
        num = int(match.group(1))
    except ValueError:
        return False
    if num < 1 or num > 9:
        return False
    body = re.sub(r"\s+", " ", match.group(2).strip().rstrip(".")).upper()
    if len(body) > 80:
        return False
    if any(key in body for key in CHAPTER_HEADING_KEYWORDS):
        return True
    return any(body == title for _, title in CHAPTER_TITLES)


def normalize_number_separator(text: str) -> str:
    """
    Номер с начала абзаца: без ведущих пробелов/табов; после номера ровно один пробел.
    Пункты: «1.1. Текст» / «1.8.1. Текст» (точка после номера обязательна).
    Заголовок главы: «1 ОБЩИЕ» — один пробел, без навязывания точки.
    """
    if not text:
        return text
    t = text.replace("\xa0", " ")
    stripped = t.lstrip(" \t")
    if not stripped:
        return ""
    if is_chapter_header(stripped):
        match = _CHAPTER_NUMBER_LINE_RE.match(stripped)
        if not match:
            return stripped
        rest = match.group(2).lstrip(" \t")
        return f"{match.group(1)} {rest}" if rest else match.group(1)
    match = _ITEM_NUMBER_LINE_RE.match(t)
    if not match:
        return t
    num = match.group(1)
    rest = (match.group(3) or "").lstrip(" \t")
    if rest:
        return f"{num}. {rest}"
    return f"{num}."


def is_section_header(text: str) -> bool:
    """Строка тела вида «1.9. Старший мастер…», не название главы и не пункт «1.9.1.»."""
    t = text.strip()
    if not t or is_chapter_header(t):
        return False
    return bool(SECTION_HEADER.match(t))


def get_signatory_start_index(doc: Document) -> int | None:
    try:
        return find_razrabotal_index(doc)
    except ValueError:
        return None


def should_apply_body_paragraph_format(text: str, idx: int, doc: Document) -> bool:
    """Абзацы тела: по ширине + отступ 1,25 см; не названия глав 1–5, не подписанты, не титул."""
    t = text.strip()
    if not t:
        return False
    if is_chapter_header(t):
        return False
    body_start = find_body_start_index(doc)
    if body_start and idx < body_start:
        return False
    razrab_idx = get_signatory_start_index(doc)
    if razrab_idx is not None and idx >= razrab_idx:
        return False
    upper = t.upper()
    if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")):
        return False
    return True


def is_paragraph_justified(paragraph: Paragraph) -> bool:
    jc = paragraph_jc(paragraph)
    if jc in ("both", "justify", "distribute"):
        return True
    try:
        return paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    except Exception:
        return False


def ensure_paragraph_justified(paragraph: Paragraph) -> None:
    """Выравнивание по ширине: python-docx + w:jc val=both."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_pr = paragraph._p.get_or_add_pPr()
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), "both")


def ensure_first_line_indent(paragraph: Paragraph, cm: float = FIRST_LINE_INDENT_CM) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(cm)
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:firstLine"), str(FIRST_LINE_INDENT_TWIPS))
    # hanging/left из стиля списка конфликтуют с отступом первой строки
    for attr in ("w:hanging", "w:left", "w:start"):
        key = qn(attr)
        if ind.get(key) is not None:
            del ind.attrib[key]


def clear_first_line_indent(paragraph: Paragraph) -> None:
    """Снять отступ первой строки (заголовки глав и разделов «1.5. …»)."""
    paragraph.paragraph_format.first_line_indent = None
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        return
    for attr in ("w:firstLine", "w:hanging", "w:left", "w:start"):
        key = qn(attr)
        if ind.get(key) is not None:
            del ind.attrib[key]
    if not ind.attrib:
        p_pr.remove(ind)


def first_line_indent_cm(paragraph: Paragraph) -> float:
    pf = paragraph.paragraph_format
    if pf.first_line_indent is not None:
        return pf.first_line_indent.cm
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is not None:
        ind = p_pr.find(qn("w:ind"))
        if ind is not None:
            raw = ind.get(qn("w:firstLine"))
            if raw:
                try:
                    return int(raw) / 567.0
                except ValueError:
                    pass
    return 0.0


def paragraph_jc(paragraph: Paragraph) -> str | None:
    p_pr = paragraph._p.pPr
    if p_pr is not None and p_pr.jc is not None:
        val = p_pr.jc.val
        if val is None:
            pass
        elif isinstance(val, str):
            return val.lower()
        else:
            name = getattr(val, "name", None)
            if name:
                return str(name).lower()
            return str(val).lower()
    if paragraph.alignment is not None:
        al = paragraph.alignment
        name = getattr(al, "name", None)
        if name:
            return str(name).lower()
    return None


def is_paragraph_centered(paragraph: Paragraph) -> bool:
    jc = paragraph_jc(paragraph)
    if jc == "center":
        return True
    try:
        return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        return False


def ensure_paragraph_centered(paragraph: Paragraph) -> None:
    """Выравнивание по центру: python-docx + w:jc в pPr (устойчиво в Word)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = paragraph._p.get_or_add_pPr()
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), "center")


def is_paragraph_right(paragraph: Paragraph) -> bool:
    return paragraph_jc(paragraph) == "right"


def ensure_paragraph_right(paragraph: Paragraph) -> None:
    """Выравнивание по правому краю (гриф «УТВЕРЖДАЮ»)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), TITLE_STAMP_ALIGN)


def ensure_paragraph_left(paragraph: Paragraph) -> None:
    """Выравнивание по левому краю (наименование ДИ на титуле)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = paragraph._p.get_or_add_pPr()
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), TITLE_DOC_NAME_ALIGN)


def _apply_title_stamp_right_indent(paragraph: Paragraph) -> bool:
    """Правый вынос грифа (эталон: w:ind right=-1333 twips)."""
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    if ind.get(qn("w:right")) == TITLE_STAMP_RIGHT_INDENT_TWIPS:
        return False
    ind.set(qn("w:right"), TITLE_STAMP_RIGHT_INDENT_TWIPS)
    return True


def _ensure_title_right_tab(paragraph: Paragraph) -> None:
    """Правый табулятор для двухколоночного титула «название | УТВЕРЖДАЮ»."""
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:tabs"))
    if old is not None:
        p_pr.remove(old)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), TITLE_RIGHT_TAB_TWIPS)
    tabs.append(tab)
    p_pr.append(tabs)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), TITLE_DOC_NAME_ALIGN)


def find_chapter_header_indices(doc: Document) -> list[int]:
    return [
        i
        for i, p in enumerate(doc.paragraphs)
        if is_chapter_header(paragraph_text_normalized(p) or p.text)
    ]


def find_first_nonempty_paragraph_after(doc: Document, idx: int) -> int | None:
    for j in range(idx + 1, len(doc.paragraphs)):
        if not is_paragraph_empty(doc.paragraphs[j]):
            return j
    return None


def paragraph_has_page_break_before(paragraph: Paragraph) -> bool:
    if paragraph.paragraph_format.page_break_before:
        return True
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is not None and p_pr.find(qn("w:pageBreakBefore")) is not None:
        return True
    return False


def page_break_target_after_chapter_header(doc: Document, hdr_idx: int) -> int | None:
    """
    Куда ставить page_break_before после заголовка главы:
    пустая строка перед текстом (если есть) или первый непустой абзац.
    """
    content_idx = find_first_nonempty_paragraph_after(doc, hdr_idx)
    if content_idx is None:
        return None
    if content_idx > hdr_idx + 1:
        return hdr_idx + 1
    return content_idx


def set_page_break_before(paragraph: Paragraph, enabled: bool) -> None:
    paragraph.paragraph_format.page_break_before = enabled
    p_pr = paragraph._p.get_or_add_pPr()
    el = p_pr.find(qn("w:pageBreakBefore"))
    if enabled:
        if el is None:
            el = OxmlElement("w:pageBreakBefore")
            p_pr.append(el)
    elif el is not None:
        p_pr.remove(el)


EXIT_OK = 0
EXIT_VALIDATION_FAIL = 1
EXIT_FILE_LOCKED = 2
EXIT_NOT_FOUND = 3


def strip_all_number_prefixes(text: str) -> str:
    """Снять все ведущие «1.5.1.» (в т.ч. дубли «1.5.1. 1.5.1.»)."""
    t = text.strip()
    while NUM_PREFIX.match(t):
        t = NUM_PREFIX.sub("", t, count=1).strip()
    return t


def strip_number(text: str) -> str:
    return strip_all_number_prefixes(text)


def paragraph_has_manual_number(text: str) -> bool:
    return bool(NUM_PREFIX.match((text or "").strip()))


def collapse_duplicate_manual_prefix(text: str) -> str:
    """«1.5.1. 1.5.1. Текст» → «1.5.1. Текст»."""
    t = (text or "").strip()
    if not t:
        return t
    prefixes: list[str] = []
    rest = t
    while True:
        match = NUM_PREFIX.match(rest)
        if not match:
            break
        prefixes.append(match.group(0).strip())
        rest = rest[match.end() :].strip()
    if len(prefixes) <= 1:
        return t
    return f"{prefixes[0]} {rest}".strip() if rest else prefixes[0]


def has_erroneous_numbering_prefix(text: str) -> bool:
    """Лишняя точка перед «1.4.1.» или двойная точка «1.4.1..»."""
    t = (text or "").strip()
    if not t:
        return False
    return bool(LEADING_DOT_BEFORE_NUMBER.match(t) or DOUBLE_DOT_AFTER_NUMBER.match(t))


def fix_erroneous_numbering_prefix(text: str) -> str:
    """
    Убрать ошибочную точку перед нумерацией и двойную точку после номера.

    Примеры:
    - «.1.4.1. Текст» / «. 1.4.1. Текст» → «1.4.1. Текст»
    - «1.4.1.. Текст» / «1.4.1..Текст» → «1.4.1. Текст»
    """
    t = (text or "").strip()
    if not t:
        return text or ""
    new_t = t
    match = LEADING_DOT_BEFORE_NUMBER.match(new_t)
    if match:
        prefix = match.group(1)
        rest = new_t[match.end() :].lstrip()
        rest = re.sub(r"^\.+\s*", "", rest)
        new_t = f"{prefix} {rest}".strip() if rest else prefix
    match = DOUBLE_DOT_AFTER_NUMBER.match(new_t)
    if match:
        prefix = match.group(1)
        rest = new_t[match.end() :].lstrip()
        new_t = f"{prefix} {rest}".strip() if rest else prefix
    return new_t


INVISIBLE_AND_CONTROL_RE = re.compile(
    r"[\u200b-\u200f\ufeff\u00ad\x00-\x08\x0b\x0c\x0e-\x1f]"
)
MULTI_SPACE_RE = re.compile(r" {2,}")
SPACE_BEFORE_PUNCT_RE = re.compile(r"([а-яёА-ЯЁ])\s+([.,;:!?])(?![.])")
SPACE_AFTER_OPEN_PAREN_RE = re.compile(r"\(\s+")
SPACE_BEFORE_CLOSE_PAREN_RE = re.compile(r"\s+\)")
DOUBLE_COMMA_RE = re.compile(r",{2,}")
DOUBLE_DOT_NOT_ELLIPSIS_RE = re.compile(r"(?<!\.)\.\.(?!\.)")
CYRILLIC_QUOTED_RE = re.compile(r'"([а-яёА-ЯЁ][^"]*?[а-яёА-ЯЁ0-9])"')
CITY_YEAR_RE = re.compile(
    r"^(?:г\.?\s*)?минск(?:\s*[-–,]\s*|\s+)20\d{2}(?:\s*г\.?)?$",
    re.IGNORECASE,
)
MUST_KNOW_HEADER_RE = re.compile(r"должен знать\s*:?\s*$", re.IGNORECASE)
OFFICE_WORK_BASICS_LINE = "основы делопроизводства;"
OFFICE_WORK_BASICS_NEEDLE = "основы делопроизводства"
_ENUM_TRAILING_PUNCT_RE = re.compile(r"[ \t]*[.;,]+[ \t]*$")
_NUMBERED_LIST_ITEM_RE = re.compile(r"^(\d+(?:\.\d+)+\.)\s+(.*)$")
DUTY_BY_ORDER_RE = re.compile(
    r"(исполняет),?\s+(по распоряжению\s+[^,\n]+?),?\s+(\bспециалист\b)",
    re.IGNORECASE,
)
LPA_BARE_RE = re.compile(r"(?<![А-Яа-яЁёA-Za-z])ЛПА(?![А-Яа-яЁёA-Za-z])")
LPA_EXPANDED_RE = re.compile(
    r"локальн\w*\s+правов\w*\s+акт\w*\s*\(\s*далее\s*[-–—]\s*ЛПА\s*\)",
    re.IGNORECASE,
)
LPA_DALEE_PREFIX_RE = re.compile(r"далее\s*[-–—]\s*$", re.IGNORECASE)
LPA_FULL_PHRASE_TAIL_RE = re.compile(
    r"локальн\w*\s+правов\w*\s+акт\w*\s*$",
    re.IGNORECASE,
)
LPA_EXPANSION_MARKER = "(далее - ЛПА)"
_LPA_FORMS = {
    (False, "nom"): "локальный правовой акт",
    (False, "gen"): "локального правового акта",
    (False, "dat"): "локальному правовому акту",
    (False, "acc"): "локальный правовой акт",
    (False, "ins"): "локальным правовым актом",
    (False, "loc"): "локальном правовом акте",
    (True, "nom"): "локальные правовые акты",
    (True, "gen"): "локальных правовых актов",
    (True, "dat"): "локальным правовым актам",
    (True, "acc"): "локальные правовые акты",
    (True, "ins"): "локальными правовыми актами",
    (True, "loc"): "локальных правовых актах",
}
_ABBREV_TOKEN_RE = re.compile(r"[A-Za-zА-Яа-яЁё]+")
KNOWN_ABBREVS = frozenset(
    {
        "лсим",
        "сниот",
        "сатп",
        "ткп",
        "нпа",
        "тнпа",
        "лпа",
        "фио",
        "ртс",
        "лэс",
        "атп",
        "ок",
        "юо",
        "оотиз",
        "мктс",
        "смат",
        "иот",
        "ри",
        "ди",
        "дсм",
        "окси",
        "прб",
        "гост",
        "iso",
        "iec",
    }
)

HOMOGLYPH_REPLACEMENTS: tuple[tuple[str, str], ...] = (
    ("Согласованo", "Согласовано"),
    ("СОГЛАСОВАНO", "СОГЛАСОВАНО"),
    ("СогласованO", "Согласовано"),
    ("СНiОТ", "СНиОТ"),
    ("СНIОТ", "СНиОТ"),
    ("СNiОТ", "СНиОТ"),
    ("СNiOT", "СНиОТ"),
)

SIGNATORY_MARKER_HOMOGLYPHS: tuple[tuple[str, str], ...] = (
    ("Согласованo", "Согласовано"),
    ("СОГЛАСОВАНO", "СОГЛАСОВАНО"),
    ("СогласованO", "Согласовано"),
    ("Разработалo", "Разработал"),
    ("РАЗРАБОТАЛO", "РАЗРАБОТАЛ"),
)

# Маркеры списка Word / текста. Не номера страниц, не «1.4.1.», не Tab подписантов.
LIST_MARKER_CHARS: frozenset[str] = frozenset(
    {
        "\u2022",  # •
        "\u2023",  # ‣
        "\u2043",  # ⁃
        "\u25e6",  # ◦
        "\u25cb",  # ○
        "\u25a0",  # ■
        "\u25a1",  # □
        "\u25aa",  # ▪
        "\u25cf",  # ●
        "\u25c6",  # ◆
        "\uf0b7",  #  (Wingdings)
        "\uf0a7",  # 
        "\u00b7",  # ·
        "\u2219",  # ∙
        "*",
        "◦",
        "▪",
        "▸",
        "►",
    }
)
LIST_MARKER_DASHES: frozenset[str] = frozenset({"-", "–", "—"})
_MANUAL_NUMBER_AT_START_RE = re.compile(r"^\d+(\.\d+)+\.\s|^\d+\.\d+")


def _remove_invisible_and_controls(text: str) -> str:
    return INVISIBLE_AND_CONTROL_RE.sub("", text)


def _normalize_whitespace_chars(text: str) -> str:
    t = text.replace("\xa0", " ").replace("\t", " ")
    t = _remove_invisible_and_controls(t)
    return MULTI_SPACE_RE.sub(" ", t)


def _apply_homoglyph_replacements(text: str, pairs: tuple[tuple[str, str], ...]) -> str:
    t = text
    for wrong, right in pairs:
        t = t.replace(wrong, right)
    return t


def _normalize_russian_quotes(text: str) -> str:
    t = text.replace("\u201c", "«").replace("\u201d", "»").replace("\u201e", "«")
    return CYRILLIC_QUOTED_RE.sub(r"«\1»", t)


def _fix_punctuation_artifacts(text: str) -> str:
    t = SPACE_BEFORE_PUNCT_RE.sub(r"\1\2", text)
    t = DOUBLE_COMMA_RE.sub(",", t)
    t = DOUBLE_DOT_NOT_ELLIPSIS_RE.sub(".", t)
    t = fix_spaces_around_parentheses(t)
    return t


def fix_spaces_around_parentheses(text: str) -> str:
    """
    Только внутри скобок: нет пробела сразу после «(» и сразу перед «)».
    «продукции ( выполнение )» → «продукции (выполнение)».
    Пробел перед открывающей скобкой после слова не убирать.
    Нумерацию 1.1. не трогает.
    """
    if not text or ("(" not in text and ")" not in text):
        return text
    t = SPACE_AFTER_OPEN_PAREN_RE.sub("(", text)
    t = SPACE_BEFORE_CLOSE_PAREN_RE.sub(")", t)
    return t


def collapse_adjacent_duplicate_words(text: str) -> str:
    """
    Соседние одинаковые слова в одном абзаце: «службы службы» → «службы».
    Без учёта регистра; «ТКП ТКП» тоже. Не трогает «1.1. 1.1.» (цифры)
    и повторы через союз не рядом.
    """
    if not text:
        return text
    prev = None
    t = text
    while prev != t:
        prev = t
        t = _ADJACENT_DUP_WORD_RE.sub(r"\1", t)
    return t


def _letters_only(word: str) -> str:
    return "".join(ch for ch in (word or "") if ch.isalpha())


def is_abbreviation_token(word: str) -> bool:
    """
    Служебные сокращения: ЛСиМ, СНиОТ, САТП, ТКП, НПА…
    2–6 букв, все заглавные или смешанный регистр с ≥2 заглавными.
    Word (красная черта) и python-спеллер такие токены не правят.
    """
    letters = _letters_only(word)
    if not letters:
        return False
    folded = letters.replace("ё", "е").replace("Ё", "Е").lower()
    if folded in KNOWN_ABBREVS:
        return True
    if not 2 <= len(letters) <= 6:
        return False
    n_upper = sum(1 for ch in letters if ch.isupper())
    n_lower = sum(1 for ch in letters if ch.islower())
    if n_lower == 0 and n_upper >= 2:
        return True
    if n_upper >= 2 and n_lower >= 1:
        return True
    return False


def spelling_error_is_abbreviation(text: str) -> bool:
    """Красная черта Word на аббревиатуре — не заменять."""
    token = (text or "").replace("\xa0", " ").strip().strip(".,;:!?()«»\"'`")
    if not token:
        return False
    if is_abbreviation_token(token):
        return True
    parts = _ABBREV_TOKEN_RE.findall(token)
    return bool(parts) and all(is_abbreviation_token(p) for p in parts)


def pick_spelling_suggestion(original: str, suggestions: list[str]) -> str | None:
    """Первое подходящее исправление Word; аббревиатуры и пустые отбросить."""
    src = (original or "").strip()
    if not src or spelling_error_is_abbreviation(src):
        return None
    for raw in suggestions:
        cand = (raw or "").strip()
        if not cand or cand == src:
            continue
        if spelling_error_is_abbreviation(cand):
            continue
        if len(cand) > max(24, len(src) * 3):
            continue
        return cand
    return None


def _light_sanitize_text(text: str) -> str:
    """Минимальная очистка: титул, заголовки глав, блок подписантов (имена не трогаем)."""
    if not text:
        return text
    preserve_tabs = "\t" in text or bool(SIGNATORY_NAME_TAIL.search(text))
    t = text.replace("\xa0", " ")
    if not preserve_tabs:
        t = t.replace("\t", " ")
    t = _remove_invisible_and_controls(t)
    if preserve_tabs:
        t = t.strip(" \xa0")
    else:
        t = t.strip()
    t = MULTI_SPACE_RE.sub(" ", t)
    t = fix_spaces_around_parentheses(t)
    upper = t.upper()
    if upper.startswith("СОГЛАСОВАН") or upper.startswith("РАЗРАБОТАЛ"):
        t = _apply_homoglyph_replacements(t, SIGNATORY_MARKER_HOMOGLYPHS)
    t = collapse_adjacent_duplicate_words(t)
    return t


def text_has_leading_list_marker(text: str) -> bool:
    """Ведущий маркер списка в тексте абзаца (не «1.4.1.» и не пустая строка)."""
    if not text:
        return False
    stripped = text.replace("\xa0", " ").lstrip(" \t")
    if not stripped or _MANUAL_NUMBER_AT_START_RE.match(stripped):
        return False
    return stripped[0] in LIST_MARKER_CHARS or stripped[0] in LIST_MARKER_DASHES


def strip_leading_list_marker_text(text: str) -> str:
    """Убрать • /  / ○ / ■ / дефис-список в начале; нумерацию 1.4.1. не трогать."""
    if not text:
        return text
    original = text
    t = text.replace("\xa0", " ")
    changed = 0
    while changed < 5:
        stripped = t.lstrip(" \t")
        if not stripped or _MANUAL_NUMBER_AT_START_RE.match(stripped):
            break
        if stripped[0] in LIST_MARKER_CHARS or stripped[0] in LIST_MARKER_DASHES:
            t = stripped[1:]
            changed += 1
            continue
        break
    if not changed:
        return original
    return t.lstrip(" \t")


def sanitize_paragraph_text(text: str, *, full: bool = True) -> str:
    """
    Очистка текста абзаца от лишних символов.

    full=False — только invisible/nbsp/tab/trim (титул, подписанты, заголовки глав).
    full=True — полная очистка тела: homoglyphs, пунктуация, кавычки, нумерация,
    ведущие маркеры списка.
    """
    if not text:
        return text
    if not full:
        return _light_sanitize_text(text)
    t = _light_sanitize_text(text)
    t = t.replace("\t", " ")
    t = _apply_homoglyph_replacements(t, HOMOGLYPH_REPLACEMENTS)
    t = strip_leading_list_marker_text(t)
    t = fix_erroneous_numbering_prefix(t)
    t = _fix_punctuation_artifacts(t)
    t = _normalize_russian_quotes(t)
    t = MULTI_SPACE_RE.sub(" ", t).strip()
    t = collapse_adjacent_duplicate_words(t)
    t = normalize_number_separator(t)
    return t


def should_use_full_text_sanitize(text: str, idx: int, doc: Document) -> bool:
    """Полная очистка только для абзацев тела (не титул, не главы 1–5, не подписанты)."""
    t = text.strip()
    if not t:
        return False
    if is_chapter_header(t):
        return False
    if idx < find_body_start_index(doc):
        return False
    if not should_apply_body_paragraph_format(text, idx, doc):
        return False
    return True


def strip_unnecessary_characters(doc: Document) -> int:
    """Убрать лишние символы во всех абзацах (лёгкий или полный режим по типу абзаца).

    Текст берём из XML (w:t + w:tab): python-docx .text не видит табы и часть пробелов в runs.
    """
    changed = 0
    for idx, paragraph in iter_document_paragraphs(doc):
        original = paragraph_xml_visible_text(paragraph)
        if not original:
            continue
        full = should_use_full_text_sanitize(original, idx, doc)
        new_text = sanitize_paragraph_text(original, full=full)
        if new_text != original:
            set_paragraph_text(paragraph, new_text)
            changed += 1
    return changed


def validate_unnecessary_characters(doc: Document) -> list[str]:
    """Проверка: нет лишних invisible, двойных пробелов, homoglyphs в маркерах и тела."""
    issues: list[str] = []
    razrab_idx = get_signatory_start_index(doc)
    body_start = find_body_start_index(doc)
    for idx, paragraph in iter_document_paragraphs(doc):
        original = paragraph_xml_visible_text(paragraph)
        if not original:
            continue
        is_signatory = razrab_idx is not None and idx >= razrab_idx and original.strip()
        on_title = idx >= 0 and (body_start <= 0 or idx < body_start)
        if INVISIBLE_AND_CONTROL_RE.search(original):
            issues.append(f"Служебные/нулевой ширины символы: {original[:50]}")
            continue
        if "\t" in original and not is_signatory and not on_title and idx >= 0:
            issues.append(f"Табуляция вне блока подписантов: {original[:50]}")
            continue
        collapsed = original.replace("\xa0", " ")
        if MULTI_SPACE_RE.search(collapsed):
            issues.append(f"Двойные пробелы в тексте: {original[:50]}")
        if "\xa0" in original:
            fio_nbsp_ok = is_signatory and bool(SIGNATORY_NAME_TAIL.search(original))
            if not fio_nbsp_ok:
                issues.append(f"Неразрывный пробел: {original[:50]}")
        full = should_use_full_text_sanitize(original, idx, doc)
        expected = sanitize_paragraph_text(original, full=full)
        if expected != original and not MULTI_SPACE_RE.search(collapsed) and "\xa0" not in original:
            snippet = original[:55] + ("…" if len(original) > 55 else "")
            issues.append(f"Лишние символы не убраны: «{snippet}»")
        upper = original.strip().upper()
        if upper.startswith("СОГЛАСОВАН") and "Согласованo" in original:
            issues.append("Латинская «o» в «Согласовано»")
    return issues


def check_document_punctuation_after_edit(doc: Document) -> int:
    """После любой правки — пунктуация всего документа (титул, тело, таблицы, подписанты)."""
    changed = strip_unnecessary_characters(doc)
    changed += normalize_must_know_list_punctuation(doc)
    return changed


def validate_document_punctuation(doc: Document) -> list[str]:
    """Проверка пунктуации всего документа после правки. Без неё документ не «Готово»."""
    issues = validate_unnecessary_characters(doc)
    issues.extend(validate_must_know_list_punctuation(doc))
    return issues


def should_skip_numbering_prefix_fix(text: str, idx: int, doc: Document) -> bool:
    """Не трогать заголовки глав и блок подписантов."""
    t = text.strip()
    if not t or is_chapter_header(t):
        return True
    razrab_idx = get_signatory_start_index(doc)
    if razrab_idx is not None and idx >= razrab_idx:
        return True
    upper = t.upper()
    return upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН"))


def strip_leading_dot_before_numbering(doc: Document) -> int:
    """
    Убрать лишнюю точку перед текстовой нумерацией (Word-список + ручной номер).

    Только абзацы с нумерацией вида 1.4.1. / 1.5.2. — не заголовки глав 1–5.
    """
    changed = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        visible = paragraph_visible_text(paragraph).replace("\xa0", " ")
        t = visible.strip()
        if not t or should_skip_numbering_prefix_fix(t, idx, doc):
            continue
        fixed = fix_erroneous_numbering_prefix(t)
        if fixed != t:
            set_paragraph_text(paragraph, fixed)
            changed += 1
    return changed


def has_word_list_numbering(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    return p_pr.find(qn("w:numPr")) is not None


def remove_word_list_numbering(paragraph: Paragraph) -> bool:
    """Убрать w:numPr — иначе Word показывает номер списка + номер в тексте."""
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return False
    p_pr.remove(num_pr)
    return True


def numbering_level_map(doc: Document) -> dict[tuple[int, int], tuple[str, str]]:
    """(numId, ilvl) → (numFmt, lvlText) из numbering.xml."""
    out: dict[tuple[int, int], tuple[str, str]] = {}
    try:
        root = doc.part.numbering_part.element
    except Exception:
        return out
    abs_info: dict[tuple[int, int], tuple[str, str]] = {}
    for absn in root.findall(qn("w:abstractNum")):
        abs_raw = absn.get(qn("w:abstractNumId"))
        if abs_raw is None:
            continue
        try:
            abs_id = int(abs_raw)
        except ValueError:
            continue
        for lvl in absn.findall(qn("w:lvl")):
            ilvl_raw = lvl.get(qn("w:ilvl")) or "0"
            try:
                ilvl = int(ilvl_raw)
            except ValueError:
                ilvl = 0
            fmt_el = lvl.find(qn("w:numFmt"))
            txt_el = lvl.find(qn("w:lvlText"))
            fmt = (fmt_el.get(qn("w:val")) or "") if fmt_el is not None else ""
            lvl_text = (txt_el.get(qn("w:val")) or "") if txt_el is not None else ""
            abs_info[(abs_id, ilvl)] = (fmt, lvl_text)
    num_to_abs: dict[int, int] = {}
    for num in root.findall(qn("w:num")):
        num_raw = num.get(qn("w:numId"))
        abs_el = num.find(qn("w:abstractNumId"))
        if num_raw is None or abs_el is None:
            continue
        try:
            num_to_abs[int(num_raw)] = int(abs_el.get(qn("w:val")) or "0")
        except ValueError:
            continue
    for num_id, abs_id in num_to_abs.items():
        for (a_id, ilvl), info in abs_info.items():
            if a_id == abs_id:
                out[(num_id, ilvl)] = info
    return out


def paragraph_numbering_ids(paragraph: Paragraph) -> tuple[int, int] | None:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        return None
    ilvl_el = num_pr.find(qn("w:ilvl"))
    try:
        num_id = int(num_id_el.get(qn("w:val")) or "0")
        ilvl = int(ilvl_el.get(qn("w:val")) or "0") if ilvl_el is not None else 0
    except ValueError:
        return None
    return num_id, ilvl


def numbering_looks_like_bullet(fmt: str, lvl_text: str) -> bool:
    if (fmt or "").strip().lower() == "bullet":
        return True
    t = (lvl_text or "").strip()
    if not t or "%" in t:
        return False
    if t in {"o", "O", "-", "–", "—", "*"}:
        return True
    return any(ch in LIST_MARKER_CHARS for ch in t)


def is_bullet_word_list(
    paragraph: Paragraph, level_map: dict[tuple[int, int], tuple[str, str]]
) -> bool:
    ids = paragraph_numbering_ids(paragraph)
    if ids is None:
        return False
    fmt, lvl_text = level_map.get(ids, ("", ""))
    return numbering_looks_like_bullet(fmt, lvl_text)


_LVL_PLACEHOLDER_RE = re.compile(r"%(\d+)")
_NUMBER_TOKEN_RE = re.compile(r"^(\d+(?:\.\d+)+)\.?")


def format_word_list_label(lvl_text: str, counters: list[int]) -> str:
    """Подставить счётчики в lvlText Word: «1.%1.» + [4] → «1.4.»."""

    def _repl(match: re.Match[str]) -> str:
        idx = int(match.group(1)) - 1
        if 0 <= idx < len(counters) and counters[idx]:
            return str(counters[idx])
        return match.group(0)

    return _LVL_PLACEHOLDER_RE.sub(_repl, lvl_text or "")


def normalize_number_token(raw: str) -> str:
    text = (raw or "").replace("\xa0", " ").replace("\t", " ").strip()
    match = _NUMBER_TOKEN_RE.match(text)
    if not match:
        return ""
    token = match.group(1)
    return token if token.endswith(".") else f"{token}."


def numbering_level_details(
    doc: Document,
) -> dict[tuple[int, int], tuple[str, str, int]]:
    """(numId, ilvl) → (numFmt, lvlText, start)."""
    out: dict[tuple[int, int], tuple[str, str, int]] = {}
    try:
        root = doc.part.numbering_part.element
    except Exception:
        return out
    abs_info: dict[tuple[int, int], tuple[str, str, int]] = {}
    for absn in root.findall(qn("w:abstractNum")):
        abs_raw = absn.get(qn("w:abstractNumId"))
        if abs_raw is None:
            continue
        try:
            abs_id = int(abs_raw)
        except ValueError:
            continue
        for lvl in absn.findall(qn("w:lvl")):
            ilvl_raw = lvl.get(qn("w:ilvl")) or "0"
            try:
                ilvl = int(ilvl_raw)
            except ValueError:
                ilvl = 0
            fmt_el = lvl.find(qn("w:numFmt"))
            txt_el = lvl.find(qn("w:lvlText"))
            start_el = lvl.find(qn("w:start"))
            fmt = (fmt_el.get(qn("w:val")) or "") if fmt_el is not None else ""
            lvl_text = (txt_el.get(qn("w:val")) or "") if txt_el is not None else ""
            try:
                start = int(start_el.get(qn("w:val")) or "1") if start_el is not None else 1
            except ValueError:
                start = 1
            abs_info[(abs_id, ilvl)] = (fmt, lvl_text, start)
    num_to_abs: dict[int, int] = {}
    overrides: dict[tuple[int, int], int] = {}
    for num in root.findall(qn("w:num")):
        num_raw = num.get(qn("w:numId"))
        abs_el = num.find(qn("w:abstractNumId"))
        if num_raw is None or abs_el is None:
            continue
        try:
            num_id = int(num_raw)
            abs_id = int(abs_el.get(qn("w:val")) or "0")
        except ValueError:
            continue
        num_to_abs[num_id] = abs_id
        for ov in num.findall(qn("w:lvlOverride")):
            try:
                ilvl = int(ov.get(qn("w:ilvl")) or "0")
            except ValueError:
                ilvl = 0
            start_ov = ov.find(qn("w:startOverride"))
            if start_ov is None:
                continue
            try:
                overrides[(num_id, ilvl)] = int(start_ov.get(qn("w:val")) or "1")
            except ValueError:
                continue
    for num_id, abs_id in num_to_abs.items():
        for (a_id, ilvl), info in abs_info.items():
            if a_id != abs_id:
                continue
            fmt, lvl_text, start = info
            start = overrides.get((num_id, ilvl), start)
            out[(num_id, ilvl)] = (fmt, lvl_text, start)
    return out


def iter_word_decimal_labels(doc: Document):
    """Абзацы тела с десятичным списком Word и вычисленная метка (как ListString)."""
    details = numbering_level_details(doc)
    if not details:
        return
    counters: dict[int, list[int]] = {}
    for paragraph in doc.paragraphs:
        ids = paragraph_numbering_ids(paragraph)
        if ids is None:
            continue
        num_id, ilvl = ids
        fmt, lvl_text, start = details.get(ids, ("", "", 1))
        if numbering_looks_like_bullet(fmt, lvl_text):
            continue
        if not lvl_text or "%" not in lvl_text:
            continue
        slot = counters.setdefault(num_id, [])
        while len(slot) <= ilvl:
            slot.append(0)
        if slot[ilvl] == 0:
            slot[ilvl] = start if start > 0 else 1
        else:
            slot[ilvl] += 1
        for deeper in range(ilvl + 1, len(slot)):
            slot[deeper] = 0
        label = format_word_list_label(lvl_text, slot)
        token = normalize_number_token(label)
        if token:
            yield paragraph, token


def collect_number_tokens(doc: Document) -> list[str]:
    """Номера пунктов исходника: текст + десятичный список Word, без потери порядка."""
    tokens: list[str] = []
    seen: set[str] = set()

    def _add(raw: str) -> None:
        token = normalize_number_token(raw)
        if token and token not in seen:
            seen.add(token)
            tokens.append(token)

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").replace("\xa0", " ").strip()
        if not text or is_chapter_header(text):
            continue
        if _is_signatory_for_list_markers(text, -1, doc):
            continue
        match = NUM_PREFIX.match(text.replace("\t", " "))
        if match:
            _add(match.group(0))
    for _paragraph, token in iter_word_decimal_labels(doc):
        _add(token)
    return tokens


def materialize_word_decimal_numbering(doc: Document) -> int:
    """
    Вписать десятичные автономера Word в текст теми же номерами, что в исходнике.
    Не перенумеровывать. Если в тексте уже есть номер — оставить его, снять только numPr.
    """
    changed = 0
    for paragraph, token in iter_word_decimal_labels(doc):
        text = (paragraph.text or "").replace("\xa0", " ")
        stripped = text.strip()
        if not stripped:
            continue
        existing = normalize_number_token(stripped.replace("\t", " "))
        if existing:
            if remove_word_list_numbering(paragraph):
                changed += 1
            continue
        prefix = token if token.endswith(".") else f"{token}."
        new_text = f"{prefix} {stripped}"
        was_bold = bool(paragraph.runs and paragraph.runs[0].bold)
        header = is_chapter_header(stripped) or is_chapter_header(new_text)
        set_paragraph_text(paragraph, new_text, bold=was_bold or header)
        remove_word_list_numbering(paragraph)
        changed += 1
    return changed


def validate_source_numbering_preserved(
    doc: Document,
    profile: DocumentProfile,
    path: Path | None = None,
) -> list[str]:
    """Набор номеров исходника не потерять; чужие номера не вписывать (не ДИ САТП)."""
    if profile.has_di_satp_numbering:
        return []
    expected = list(profile.source_number_tokens or ())
    if not expected and path is not None:
        sibling = find_agent_sibling_source(path)
        if sibling is not None:
            try:
                src_doc = open_word_document_readonly(sibling)
                expected = collect_number_tokens(src_doc)
            except Exception:
                expected = []
    if not expected:
        return []
    actual = set(collect_number_tokens(doc))
    expected_set = set(expected)
    issues: list[str] = []
    lost = [token for token in expected if token not in actual]
    if lost:
        issues.append("Потеряна нумерация исходника: " + ", ".join(lost[:15]))
    extra = [token for token in collect_number_tokens(doc) if token not in expected_set]
    if extra:
        issues.append("Номера, которых не было в исходнике: " + ", ".join(extra[:15]))
    return issues


def _is_signatory_for_list_markers(text: str, idx: int, doc: Document) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    upper = t.upper()
    if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")):
        return True
    if "\t" in text and SIGNATORY_NAME_TAIL.search(text):
        return True
    razrab_idx = get_signatory_start_index(doc)
    if razrab_idx is not None and idx >= 0 and idx >= razrab_idx:
        return True
    return False


def iter_body_xml_paragraphs(doc: Document):
    """Все w:p в теле документа, включая ячейки таблиц (не колонтитулы)."""
    p_to_idx = {id(paragraph._p): idx for idx, paragraph in enumerate(doc.paragraphs)}
    body = doc.element.body
    if body is None:
        return
    seen: set[int] = set()
    for p_el in body.iter(qn("w:p")):
        key = id(p_el)
        if key in seen:
            continue
        seen.add(key)
        yield p_to_idx.get(key, -1), Paragraph(p_el, doc)


def remove_list_markers_in_body(doc: Document) -> int:
    """
    Снять маркеры списка: bullet numPr (даже без ручного номера) и ведущие
    символы • /  / ○ / ■ / дефис в тексте. Десятичную нумерацию Word не трогать.
    """
    changed = 0
    level_map = numbering_level_map(doc)
    for idx, paragraph in iter_body_xml_paragraphs(doc):
        original = paragraph.text or ""
        if _is_signatory_for_list_markers(original, idx, doc):
            continue
        if is_bullet_word_list(paragraph, level_map):
            if remove_word_list_numbering(paragraph):
                changed += 1
        if not original.strip():
            continue
        new_text = strip_leading_list_marker_text(original)
        if new_text != original:
            was_bold = bool(paragraph.runs and paragraph.runs[0].bold)
            header = is_chapter_header(original.strip()) or is_chapter_header(new_text.strip())
            set_paragraph_text(paragraph, new_text, bold=was_bold or header)
            changed += 1
    return changed


def validate_list_markers(doc: Document) -> list[str]:
    """Замечание, если в теле остался ведущий маркер или Word-список bullet."""
    issues: list[str] = []
    level_map = numbering_level_map(doc)
    for idx, paragraph in iter_body_xml_paragraphs(doc):
        original = paragraph.text or ""
        if _is_signatory_for_list_markers(original, idx, doc):
            continue
        snippet = original.strip()[:50] or "(пустой абзац)"
        if is_bullet_word_list(paragraph, level_map):
            issues.append(f"Маркер списка Word (bullet): {snippet}")
        if original.strip() and text_has_leading_list_marker(original):
            issues.append(f"Ведущий маркер списка в тексте: {snippet}")
    return issues


def deduplicate_manual_and_list_numbering(doc: Document) -> int:
    """
    Двойная нумерация: авто-список Word (numPr) + ручной префикс в тексте.
    Оставляем текст; numPr снимаем. Дубли префикса в тексте схлопываем.
    """
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        collapsed = collapse_duplicate_manual_prefix(text)
        collapsed = fix_erroneous_numbering_prefix(collapsed)
        if collapsed != text:
            set_paragraph_text(paragraph, collapsed)
            text = collapsed
            changed += 1
        if not has_word_list_numbering(paragraph):
            continue
        if paragraph_has_manual_number(text) or is_chapter_header(text) or is_section_header(text):
            if remove_word_list_numbering(paragraph):
                changed += 1
    return changed


def remove_word_list_numbering_in_body(doc: Document) -> int:
    """
    Снять w:numPr там, где уже есть ручной номер в тексте, и bullet-списки
    без ручного номера.

    Скрытый список Word + «1.4.1.» / «5.1.1.» в тексте даёт лишнюю точку/номер.
    Десятичный numPr без текста (гл. 1, как в образце) не трогаем.
    Маркеры • /  / bullet numPr снимаем всегда (не нумерованный пункт).
    """
    changed = 0
    try:
        body_start = find_body_start_index(doc)
    except ValueError:
        body_start = 0
    body_end = get_body_spacing_end_index(doc)
    level_map = numbering_level_map(doc)
    for idx in range(body_start, min(body_end, len(doc.paragraphs))):
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        if not has_word_list_numbering(paragraph):
            continue
        if is_bullet_word_list(paragraph, level_map):
            if remove_word_list_numbering(paragraph):
                changed += 1
            continue
        if not text:
            continue
        if paragraph_has_manual_number(text) or is_chapter_header(text) or is_section_header(text):
            if remove_word_list_numbering(paragraph):
                changed += 1
    return changed


def ensure_single_numbering_model(doc: Document) -> int:
    """numPr + ручной префикс в одном абзаце — оставляем текст, numPr снимаем."""
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text or not has_word_list_numbering(paragraph):
            continue
        if paragraph_has_manual_number(text):
            if remove_word_list_numbering(paragraph):
                changed += 1
    return changed


def ensure_di_satp_chapter_one_numbering(doc: Document, profile: DocumentProfile) -> int:
    """
    Глава 1 до «1.8. … руководствуется»: 1.1.–1.7. в тексте, если нет numPr Word.
    """
    if not profile.has_di_satp_numbering:
        return 0
    changed = 0
    for num, fuzzy in DI_SATP_CHAPTER_ONE_NUMBERED:
        try:
            idx = find_paragraph_index(doc, fuzzy, contains=True)
        except ValueError:
            continue
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        if has_word_list_numbering(paragraph):
            continue
        if re.match(rf"^{re.escape(num)}\.\s", text):
            continue
        if paragraph_has_manual_number(text):
            continue
        set_paragraph_text(paragraph, apply_number(strip_number(text), num))
        changed += 1
    return changed


def copy_numpr_from_etalon_paragraph(etalon_paragraph: Paragraph, target_paragraph: Paragraph) -> bool:
    """Скопировать w:numPr с образца (нумерация Word без текста «1.1.»)."""
    e_ppr = etalon_paragraph._p.find(qn("w:pPr"))
    if e_ppr is None:
        return False
    e_num = e_ppr.find(qn("w:numPr"))
    if e_num is None:
        return False
    t_ppr = target_paragraph._p.get_or_add_pPr()
    existing = t_ppr.find(qn("w:numPr"))
    if existing is not None:
        t_ppr.remove(existing)
    t_ppr.append(deepcopy(e_num))
    return True


def chapter_one_intro_markers() -> list[tuple[str, str]]:
    """Пункты гл. 1 до раздела «1.8. … руководствуется»: 1.1.–1.7."""
    markers = list(DI_SATP_CHAPTER_ONE_NUMBERED)
    for offset, fuzzy in enumerate(DI_SATP_CHAPTER_ONE_NUMPR_MARKERS, start=4):
        markers.append((f"1.{offset}", fuzzy))
    return markers


def chapter_one_has_hidden_numpr(doc: Document, profile: DocumentProfile) -> bool:
    """numPr без текста «1.1.» — в Word номера часто не видны; нужен ручной префикс."""
    if not profile.has_di_satp_numbering:
        return False
    try:
        section_14_idx = find_section_header_index(
            doc, DI_SATP_CH1_GUIDED, "в своей деятельности руководствуется"
        )
    except ValueError:
        return False
    for num, fuzzy in chapter_one_intro_markers():
        try:
            idx = find_paragraph_index(doc, fuzzy, contains=True)
        except ValueError:
            continue
        if idx >= section_14_idx:
            continue
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        if has_word_list_numbering(paragraph) and not re.match(
            rf"^{re.escape(num)}\.\s", text
        ):
            return True
    return False


def materialize_chapter_one_numbering(doc: Document, profile: DocumentProfile) -> int:
    """
    Глава 1: видимые 1.1.–1.7. **в тексте**, скрытый numPr снимаем.

    Раздел «1.8. … руководствуется» ищется по фразе, не по первому «1.4.» —
    поэтому вводный «1.4. Назначение…» не ломает блок 1.8.1.
    """
    if not profile.has_di_satp_numbering:
        return 0
    changed = 0
    try:
        section_14_idx = find_section_header_index(
            doc, DI_SATP_CH1_GUIDED, "в своей деятельности руководствуется"
        )
    except ValueError:
        section_14_idx = len(doc.paragraphs)

    for num, fuzzy in chapter_one_intro_markers():
        try:
            idx = find_paragraph_index(doc, fuzzy, contains=True)
        except ValueError:
            continue
        if idx >= section_14_idx:
            continue
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        canonical = apply_number(strip_number(text), num)
        if text != canonical.strip():
            set_paragraph_text(paragraph, canonical)
            changed += 1
        if has_word_list_numbering(paragraph):
            if remove_word_list_numbering(paragraph):
                changed += 1
    return changed


def strip_wrong_intro_section_prefixes(doc: Document, profile: DocumentProfile) -> int:
    """Больше не снимаем 1.4.–1.7. с вводных — номера должны быть видны в Word."""
    return 0


def _is_inside_tag(el, tag_qn: str) -> bool:
    parent = el.getparent()
    while parent is not None:
        if parent.tag == tag_qn:
            return True
        parent = parent.getparent()
    return False


def _paragraph_element_text(p_el) -> str:
    """Видимый текст абзаца: без удалённых правок (w:del) и полей."""
    parts: list[str] = []
    for node in p_el.iter(qn("w:t")):
        if _is_inside_tag(node, qn("w:del")):
            continue
        parts.append(node.text or "")
    return "".join(parts)


def paragraph_xml_visible_text(paragraph: Paragraph) -> str:
    """Текст абзаца из XML: w:t и символ w:tab в run. Стоп табуляции в pPr — не символ."""
    parts: list[str] = []
    for node in paragraph._p.iter():
        if node.tag == qn("w:tab"):
            if _is_inside_tag(node, qn("w:del")):
                continue
            if _is_inside_tag(node, qn("w:pPr")) or _is_inside_tag(node, qn("w:tabs")):
                continue
            parts.append("\t")
        elif node.tag == qn("w:t"):
            if _is_inside_tag(node, qn("w:del")):
                continue
            parts.append(node.text or "")
    return "".join(parts)


def paragraph_visible_text(paragraph: Paragraph) -> str:
    """Текст как при показе правок Word: вставки видны, удаления скрыты."""
    return _paragraph_element_text(paragraph._p)


def count_tracked_change_nodes(doc: Document) -> tuple[int, int]:
    """Число контейнеров рецензирования: (вставки, удаления)."""
    root = doc.element
    n_ins = sum(1 for _ in root.iter(qn("w:ins"))) + sum(1 for _ in root.iter(qn("w:moveTo")))
    n_del = sum(1 for _ in root.iter(qn("w:del"))) + sum(
        1 for _ in root.iter(qn("w:moveFrom"))
    )
    return n_ins, n_del


def _remove_revision_nodes(root, tag: str) -> int:
    nodes = list(root.iter(qn(tag)))
    removed = 0
    for el in reversed(nodes):
        parent = el.getparent()
        if parent is None:
            continue
        parent.remove(el)
        removed += 1
    return removed


def _unwrap_revision_nodes(root, tag: str) -> int:
    nodes = list(root.iter(qn(tag)))
    changed = 0
    for el in reversed(nodes):
        parent = el.getparent()
        if parent is None:
            continue
        idx = parent.index(el)
        for child in list(el):
            el.remove(child)
            parent.insert(idx, child)
            idx += 1
        parent.remove(el)
        changed += 1
    return changed


def accept_tracked_changes(doc: Document) -> int:
    """
    Принять все правки рецензирования Word: вставки оставить текстом, удаления убрать.

    Без этого в Word видно «не приняты предыдущие правки», а точка из w:ins
    перед «5.1.1.» не видна python-docx в paragraph.text.
    """
    root = doc.element
    changed = 0
    changed += _remove_revision_nodes(root, "w:del")
    changed += _remove_revision_nodes(root, "w:moveFrom")
    changed += _unwrap_revision_nodes(root, "w:ins")
    changed += _unwrap_revision_nodes(root, "w:moveTo")
    for tag in (
        "w:rPrChange",
        "w:pPrChange",
        "w:sectPrChange",
        "w:tblPrChange",
        "w:trPrChange",
        "w:tcPrChange",
        "w:numPrChange",
    ):
        changed += _remove_revision_nodes(root, tag)
    return changed


def _wp_inside_table(p_el) -> bool:
    parent = p_el.getparent()
    while parent is not None:
        if parent.tag == qn("w:tbl"):
            return True
        parent = parent.getparent()
    return False


def iter_razrabotal_paragraph_elements(doc: Document) -> list:
    """Все абзацы «Разработал:» в теле, таблицах и content control — как их видит Word."""
    hits: list = []
    body = doc.element.body
    if body is None:
        return hits
    for p_el in body.iter(qn("w:p")):
        text = _paragraph_element_text(p_el).replace("\xa0", " ").strip()
        if text.upper().startswith("РАЗРАБОТАЛ"):
            hits.append(p_el)
    return hits


def deduplicate_razrabotal_markers(doc: Document) -> int:
    """Убрать лишние «Разработал:». Оставить последний в теле (перед «Согласовано»)."""
    hits = iter_razrabotal_paragraph_elements(doc)
    if len(hits) <= 1:
        return 0
    body_hits = [p_el for p_el in hits if not _wp_inside_table(p_el)]
    keep = body_hits[-1] if body_hits else hits[-1]
    removed = 0
    for p_el in hits:
        if p_el is keep:
            continue
        parent = p_el.getparent()
        if parent is None:
            continue
        parent.remove(p_el)
        removed += 1
    if not iter_razrabotal_paragraph_elements(doc):
        ensure_razrabotal_marker(doc)
    return removed


def sync_chapter_one_numpr_from_etalon(
    doc: Document,
    etalon_doc: Document,
    profile: DocumentProfile,
) -> int:
    """
    Глава 1: numPr Word с *_образец.docx для абзацев 1.1.–1.7. (без текста в строке).
    Не трогает разделы «1.8. … руководствуется» / «1.9. … должен знать».
    """
    if not profile.has_di_satp_numbering or etalon_doc is None:
        return 0
    changed = 0
    try:
        section_14_idx = find_section_header_index(
            doc, DI_SATP_CH1_GUIDED, "в своей деятельности руководствуется"
        )
    except ValueError:
        return 0

    markers = list(DI_SATP_CHAPTER_ONE_NUMBERED) + [
        (f"numpr_{i}", fuzzy) for i, fuzzy in enumerate(DI_SATP_CHAPTER_ONE_NUMPR_MARKERS, start=4)
    ]

    for _num, fuzzy in markers:
        if _num.startswith("numpr_"):
            try:
                d_idx = find_paragraph_index(doc, fuzzy, contains=True)
                e_idx = find_paragraph_index(etalon_doc, fuzzy, contains=True)
            except ValueError:
                continue
            if d_idx >= section_14_idx:
                continue
            d_para = doc.paragraphs[d_idx]
            if paragraph_has_manual_number(d_para.text):
                continue
            if copy_numpr_from_etalon_paragraph(etalon_doc.paragraphs[e_idx], d_para):
                changed += 1
            continue

        try:
            d_idx = find_paragraph_index(doc, fuzzy, contains=True)
            e_idx = find_paragraph_index(etalon_doc, fuzzy, contains=True)
        except ValueError:
            continue
        if d_idx >= section_14_idx:
            continue
        d_para = doc.paragraphs[d_idx]
        e_para = etalon_doc.paragraphs[e_idx]
        if has_word_list_numbering(e_para) and not paragraph_has_manual_number(d_para.text):
            if copy_numpr_from_etalon_paragraph(e_para, d_para):
                changed += 1
        elif not has_word_list_numbering(d_para) and not re.match(
            rf"^{re.escape(_num)}\.\s", d_para.text.strip()
        ):
            set_paragraph_text(d_para, apply_number(strip_number(d_para.text), _num))
            changed += 1
    return changed


def insert_signatory_tab_line(text: str) -> str:
    """Должность и ФИО — через табуляцию, как в образце."""
    t = (text or "").replace("\xa0", " ").strip()
    if not t or "\t" in t:
        return text or ""
    match = SIGNATORY_NAME_TAIL.search(t)
    if not match:
        return text or ""
    pos_part = t[: match.start()].strip()
    name_part = t[match.start() :].strip()
    return f"{pos_part}\t{name_part}"


def _clear_paragraph_tab_stops(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return
    old = p_pr.find(qn("w:tabs"))
    if old is not None:
        p_pr.remove(old)


def ensure_signatory_tab_stops(paragraph: Paragraph) -> None:
    """Табулятор должность↔ФИО / дата: не меньше 12 см."""
    min_pos = MIN_SIGNATORY_TAB_TWIPS
    default_pos = DEFAULT_SIGNATORY_TAB_TWIPS
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    existing = tabs.findall(qn("w:tab"))
    if not existing:
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), default_pos)
        tabs.append(tab)
    else:
        for tab in existing:
            try:
                pos = int(tab.get(qn("w:pos")) or "0")
            except ValueError:
                pos = 0
            if pos < min_pos:
                tab.set(qn("w:pos"), default_pos)
    clear_first_line_indent(paragraph)


def _paragraph_has_tab_xml(paragraph: Paragraph) -> bool:
    if "\t" in (paragraph.text or ""):
        return True
    for _ in paragraph._p.iter(qn("w:tab")):
        return True
    p_pr = paragraph._p.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:tabs")) is not None


def _strip_paragraph_tabs(paragraph: Paragraph) -> int:
    """Убрать Tab из абзаца (текст, w:tab и w:tabs). Пробел вместо таба в тексте."""
    changed = 0
    for run in paragraph.runs:
        raw = run.text or ""
        if "\t" in raw:
            run.text = MULTI_SPACE_RE.sub(" ", raw.replace("\t", " "))
            changed += 1
        for tab in list(run._r.findall(qn("w:tab"))):
            run._r.remove(tab)
            changed += 1
    _clear_paragraph_tab_stops(paragraph)
    return changed


def compact_title_stamp_date_plaque(text: str, year: str | None = None) -> str:
    """Плашка даты на титуле: без таба, длиннее линии, одна строка в правой колонке."""
    t = (text or "").replace("\xa0", " ").strip().lstrip("\t")
    year_m = re.search(r"(20\d{2})", t)
    y = year_m.group(1) if year_m else (year or "20__")
    return TITLE_STAMP_DATE_TEMPLATE.format(year=y)


def apply_signatory_tab_stops(doc: Document, profile: DocumentProfile) -> int:
    """w:tabs на строках должность↔ФИО; маркеры «Разработал:» / «Согласовано:» не трогать."""
    if not profile.has_signatories:
        return 0
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return 0
    changed = 0
    for paragraph in doc.paragraphs[razrab_idx:]:
        text = paragraph_text_normalized(paragraph)
        if not text:
            continue
        upper = text.upper()
        if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")):
            if MISSING_DEVELOPER_MARKER not in (paragraph.text or ""):
                continue
        if is_signatory_date_plaque(text):
            continue
        if "\t" not in paragraph.text and not SIGNATORY_NAME_TAIL.search(text):
            continue
        before = signatory_first_tab_pos_twips(paragraph)
        ensure_signatory_tab_stops(paragraph)
        after = signatory_first_tab_pos_twips(paragraph)
        if before != after or not paragraph_has_tab_stops(paragraph):
            changed += 1
    return changed


def _clean_signatory_cell(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()
    return cleaned.strip("/").strip("|").strip()


def _table_col_count(table) -> int:
    try:
        n = len(table.columns)
        if n:
            return n
    except (ValueError, AttributeError, IndexError):
        pass
    if not table.rows:
        return 0
    return len(table.rows[0].cells)


def _unique_row_cells(cells: list[str]) -> list[str]:
    """Схлопнуть дубли объединённых ячеек python-docx."""
    parts: list[str] = []
    for cell in cells:
        text = _clean_signatory_cell(cell)
        if text and (not parts or text != parts[-1]):
            parts.append(text)
    return parts


def _row_job_and_name(cells: list[str]) -> tuple[str, str] | None:
    parts = _unique_row_cells(cells)
    if len(parts) < 2:
        return None
    left, right = parts[0], parts[1]
    if not left or not right or left == right:
        return None
    packed = left.lower().replace(" ", "").replace(".", "")
    if any(hint in packed for hint in ("должность", "фио", "подпись", "роспись", "дата")):
        return None
    if packed.startswith("разработал") or packed.startswith("согласован"):
        return None
    return left, right


def extract_agreement_table_rows(table) -> list[tuple[str, str]]:
    """Должность и ФИО из таблицы согласования; шапку «должность/ФИО» пропускаем."""
    rows: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for row in table.rows:
        pair = _row_job_and_name([cell.text for cell in row.cells])
        if pair is None or pair in seen:
            continue
        seen.add(pair)
        rows.append(pair)
    return rows


def _is_acquaintance_sheet_table(table) -> bool:
    """Лист ознакомления: Ф.И.О. / роспись / дата — не блок подписантов."""
    blob = " ".join(cell.text for row in table.rows for cell in row.cells).upper()
    cols = _table_col_count(table)
    packed_fio = blob.replace(".", "").replace(" ", "")
    if cols < 4:
        return False
    if "ОЗНАКОМЛ" in blob:
        return True
    if "РОСПИС" in blob and ("ФИО" in packed_fio or "Ф.И.О" in blob):
        return True
    return False


def is_acquaintance_sheet_text(text: str) -> bool:
    """Фраза листа ознакомления (не пункт тела «ознакомление подчиненных»)."""
    packed = " ".join((text or "").replace("\xa0", " ").lower().replace("ё", "е").split())
    if not packed:
        return False
    if "лист ознакомления" in packed:
        return True
    if "изучил" in packed and "обязуюсь" in packed:
        return True
    if "настоящ" in packed and "инструкц" in packed and "изучил" in packed:
        return True
    if (
        "настоящ" in packed
        and "инструкц" in packed
        and "ознакомлен" in packed
        and "руководств" in packed
    ):
        return True
    if "с должностной инструкцией" in packed and "ознакомлен" in packed:
        return True
    if packed.startswith("с должностной инструкци") and "ознакомлен" in packed:
        return True
    return False


def find_acquaintance_sheet_start(doc: Document) -> int | None:
    """Первый абзац листа ознакомления — после подписантов, не в главе 1."""
    try:
        start = find_razrabotal_index(doc)
    except ValueError:
        start = max(0, len(doc.paragraphs) - 30)
    for i in range(start, len(doc.paragraphs)):
        if is_acquaintance_sheet_text(doc.paragraphs[i].text):
            return i
    return None


def ensure_acquaintance_sheet_separate_page(doc: Document) -> int:
    """Лист ознакомления всегда с новой страницы; не рвать перед «Разработал:»."""
    idx = find_acquaintance_sheet_start(doc)
    if idx is None:
        return 0
    changed = 0
    para = doc.paragraphs[idx]
    if not paragraph_has_page_break_before(para):
        set_page_break_before(para, True)
        changed += 1
    clear_first_line_indent(para)
    try:
        raz = find_razrabotal_index(doc)
        if paragraph_has_page_break_before(doc.paragraphs[raz]):
            set_page_break_before(doc.paragraphs[raz], False)
            changed += 1
    except ValueError:
        pass
    return changed


def validate_acquaintance_sheet(doc: Document) -> list[str]:
    """Лист ознакомления — отдельная страница; таблица ФИО не в подписантах."""
    idx = find_acquaintance_sheet_start(doc)
    if idx is None:
        return []
    issues: list[str] = []
    if not paragraph_has_page_break_before(doc.paragraphs[idx]):
        issues.append(
            "Лист ознакомления не на отдельной странице (нет разрыва перед блоком)"
        )
    try:
        raz = find_razrabotal_index(doc)
        if raz >= idx:
            issues.append("Лист ознакомления оказался раньше блока подписантов")
    except ValueError:
        pass
    return issues


def normalize_item_number_spacing(doc: Document) -> int:
    """Убрать ведущие пробелы/табы перед 1.1.; после номера — ровно один пробел.

    Чистим XML runs (w:tab и пробелы в w:t), а не только paragraph.text.
    """
    changed = 0
    razrab_idx = get_signatory_start_index(doc)
    for idx, paragraph in iter_document_paragraphs(doc):
        original = paragraph_xml_visible_text(paragraph)
        if not original:
            continue
        if razrab_idx is not None and idx >= 0 and idx >= razrab_idx:
            continue
        new_text = normalize_number_separator(original)
        if new_text == original:
            continue
        set_paragraph_text(paragraph, new_text, bold=is_chapter_header(new_text))
        changed += 1
    return changed


def collapse_adjacent_duplicate_words_in_document(doc: Document) -> int:
    """Схлопнуть соседние дубли слов во всех абзацах, включая лист ознакомления."""
    changed = 0
    for _idx, paragraph in iter_document_paragraphs(doc):
        original = paragraph.text
        if not original:
            continue
        new_text = collapse_adjacent_duplicate_words(original)
        if new_text == original:
            continue
        set_paragraph_text(paragraph, new_text, bold=is_chapter_header(new_text))
        changed += 1
    return changed


def is_signatory_date_plaque(text: str) -> bool:
    t = (text or "").replace("\xa0", " ").strip().lstrip("\t ")
    if not t or "_" not in t or "20" not in t:
        return False
    return bool(SIGNATORY_DATE_PLAQUE_RE.search(t))


def compact_signatory_date_plaque(text: str) -> str:
    """Жёлтый эталон: Tab + «__»_______ 2026г. — одна строка, короткий подчерк."""
    t = (text or "").replace("\xa0", " ").strip().lstrip("\t")
    year_m = re.search(r"(20\d{2})", t)
    year = f"{year_m.group(1)}г." if year_m else "20__г."
    return f"\t«__»_______ {year}"


def split_date_from_signatory_line(text: str) -> tuple[str, str | None]:
    """Если И.О.Ф. и плашка даты в одной строке — разделить."""
    raw = text or ""
    match = SIGNATORY_DATE_PLAQUE_RE.search(raw)
    if not match:
        return raw, None
    before = raw[: match.start()].rstrip(" \t")
    if not before:
        return raw, None
    if not SIGNATORY_NAME_TAIL.search(before):
        return raw, None
    return before, compact_signatory_date_plaque(match.group(0))


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_paragraph_keep_lines(paragraph: Paragraph, enabled: bool) -> None:
    """keepLines: плашка даты / ФИО не переносятся на две строки."""
    paragraph.paragraph_format.keep_together = bool(enabled)
    p_pr = paragraph._p.get_or_add_pPr()
    el = p_pr.find(qn("w:keepLines"))
    if enabled:
        if el is None:
            p_pr.append(OxmlElement("w:keepLines"))
        return
    if el is not None:
        p_pr.remove(el)


def set_paragraph_keep_with_next(paragraph: Paragraph, enabled: bool) -> None:
    """keepNext: абзац не отрывать от следующего (блок подписантов + хвост текста)."""
    paragraph.paragraph_format.keep_with_next = bool(enabled)
    _set_p_keep_next(paragraph._p, enabled)


def _paragraph_has_keep_next(paragraph: Paragraph) -> bool:
    if paragraph.paragraph_format.keep_with_next:
        return True
    p_pr = paragraph._p.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:keepNext")) is not None


def signatory_first_tab_pos_twips(paragraph: Paragraph) -> str | None:
    """Позиция первого (левого) табулятора — начало ФИО / даты."""
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        return None
    positions: list[int] = []
    for tab in tabs.findall(qn("w:tab")):
        pos = tab.get(qn("w:pos"))
        if not pos:
            continue
        try:
            positions.append(int(pos))
        except ValueError:
            continue
    if not positions:
        return None
    return str(min(positions))


def _copy_signatory_tab_stops(src: Paragraph, dst: Paragraph) -> None:
    """Скопировать w:tabs, чтобы дата стояла ровно под началом ФИО."""
    ensure_signatory_tab_stops(src)
    if not paragraph_has_tab_stops(src):
        ensure_signatory_tab_stops(src)
    e_tabs = src._p.get_or_add_pPr().find(qn("w:tabs"))
    if e_tabs is None:
        ensure_signatory_tab_stops(dst)
        return
    d_pr = dst._p.get_or_add_pPr()
    old = d_pr.find(qn("w:tabs"))
    if old is not None:
        d_pr.remove(old)
    d_pr.append(deepcopy(e_tabs))
    clear_first_line_indent(dst)


def _format_date_plaque_paragraph(date_para: Paragraph, iof_para: Paragraph | None) -> None:
    clear_first_line_indent(date_para)
    if iof_para is not None:
        _copy_signatory_tab_stops(iof_para, date_para)
    else:
        ensure_signatory_tab_stops(date_para)
    set_paragraph_keep_lines(date_para, True)
    set_one_point_five_line_spacing(date_para)


def _previous_iof_paragraph(doc: Document, idx: int, raz: int) -> Paragraph | None:
    for j in range(idx - 1, raz - 1, -1):
        text = doc.paragraphs[j].text or ""
        if not text.strip():
            continue
        if is_signatory_date_plaque(text):
            continue
        if SIGNATORY_NAME_TAIL.search(text):
            return doc.paragraphs[j]
        break
    return None


def glue_signatory_fio(name: str) -> str:
    """Неразрывные пробелы внутри И.О.Ф. — не рвать инициалы и фамилию."""
    t = re.sub(r"\s+", " ", (name or "").replace("\xa0", " ").strip())
    t = re.sub(r"([А-ЯЁ]\.)\s+([А-ЯЁ]\.)", "\\1\u00a0\\2", t)
    t = re.sub(r"([А-ЯЁ]\.)\s+([А-ЯЁ][а-яё])", "\\1\u00a0\\2", t)
    return t


def shorten_signatory_job_underline(job: str) -> str:
    """Укоротить длинную линию подписи, чтобы ФИО осталось на той же строке."""
    t = (job or "").rstrip()
    t = re.sub(r"_{8,}", "______", t)
    t = re.sub(r"[ \t]{2,}", " ", t)
    return t.strip()


def _is_iof_only_signatory_line(text: str) -> bool:
    t = (text or "").replace("\xa0", " ").strip().lstrip("\t")
    if not t or is_signatory_date_plaque(t):
        return False
    upper = t.upper()
    if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")):
        return False
    return bool(SIGNATORY_IOF_ONLY_RE.match(t))


def _looks_like_job_without_fio(text: str) -> bool:
    t = (text or "").replace("\xa0", " ").strip()
    if not t or is_signatory_date_plaque(t):
        return False
    upper = t.upper()
    if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")):
        return False
    if SIGNATORY_NAME_TAIL.search(t) or SIGNATORY_IOF_ONLY_RE.match(t.lstrip("\t")):
        return False
    low = t.lower()
    hints = (
        "начальник",
        "инженер",
        "директор",
        "мастер",
        "заместител",
        "главн",
        "ведущий",
        "специалист",
        "юрист",
        "экономист",
    )
    return any(h in low for h in hints) or "_" in t


def paragraph_has_soft_line_break(paragraph: Paragraph) -> bool:
    for el in paragraph._p.iter():
        if el.tag in (qn("w:br"), qn("w:cr")):
            return True
    return False


def _extract_iof_from_line(text: str) -> str | None:
    t = (text or "").replace("\xa0", " ").strip().lstrip("\t")
    match = SIGNATORY_IOF_ONLY_RE.match(t)
    if match:
        return match.group(1)
    tail = SIGNATORY_NAME_TAIL.search(t)
    if tail:
        return tail.group(2)
    return None


def apply_signatory_fio_one_line(doc: Document) -> int:
    """
    Должность и ФИО только на одной строке через Tab.
    ФИО не переносить: keepLines, короче линия, неразрывные пробелы в И.О.Ф.
    """
    try:
        raz = find_razrabotal_index(doc)
    except ValueError:
        return 0
    changed = 0
    end = _signatory_block_end_index(doc)
    i = raz
    while i < end - 1 and i < len(doc.paragraphs) - 1:
        para = doc.paragraphs[i]
        text = para.text or ""
        if is_acquaintance_sheet_text(text):
            break
        nxt = doc.paragraphs[i + 1]
        nxt_text = nxt.text or ""
        if _looks_like_job_without_fio(text) and _is_iof_only_signatory_line(nxt_text):
            iof = _extract_iof_from_line(nxt_text) or nxt_text.strip().lstrip("\t")
            job = shorten_signatory_job_underline(text)
            set_paragraph_text(para, f"{job}\t{glue_signatory_fio(iof)}", bold=False)
            ensure_signatory_tab_stops(para)
            set_paragraph_keep_lines(para, True)
            set_one_point_five_line_spacing(para)
            parent = nxt._p.getparent()
            if parent is not None:
                parent.remove(nxt._p)
            changed += 1
            end = _signatory_block_end_index(doc)
            continue
        i += 1
    i = raz
    end = _signatory_block_end_index(doc)
    while i < end and i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text or ""
        if is_acquaintance_sheet_text(text):
            break
        upper = text.strip().upper()
        if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")) or is_signatory_date_plaque(text):
            i += 1
            continue
        if paragraph_has_soft_line_break(para) and (
            SIGNATORY_NAME_TAIL.search(text) or _is_iof_only_signatory_line(text)
        ):
            visible = paragraph_xml_visible_text(para) or text
            flat = visible.replace("\n", " ").replace("\r", " ")
            flat = re.sub(r"[^\S\t]+", " ", flat).strip(" ")
            if flat != text:
                set_paragraph_text(para, flat, bold=False)
                text = para.text or ""
                changed += 1
        new_text = insert_signatory_tab_line(text)
        if new_text != text:
            set_paragraph_text(para, new_text, bold=False)
            text = new_text
            changed += 1
        if "\t" in text and SIGNATORY_NAME_TAIL.search(text):
            job, name = text.split("\t", 1)
            rebuilt = f"{shorten_signatory_job_underline(job)}\t{glue_signatory_fio(name)}"
            if rebuilt != (para.text or ""):
                set_paragraph_text(para, rebuilt, bold=False)
                changed += 1
            ensure_signatory_tab_stops(para)
            if not _paragraph_has_keep_lines(para):
                set_paragraph_keep_lines(para, True)
                changed += 1
        i += 1
    return changed


def fix_signatory_date_plaques(doc: Document) -> int:
    """Плашка даты: новая строка, Tab как у И.О.Ф., короткий подчерк, keepLines."""
    changed = apply_signatory_fio_one_line(doc)
    try:
        raz = find_razrabotal_index(doc)
    except ValueError:
        return changed
    i = raz
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text or ""
        if is_acquaintance_sheet_text(text):
            break
        before, date_line = split_date_from_signatory_line(text)
        if date_line is not None:
            set_paragraph_text(para, before, bold=False)
            ensure_signatory_tab_stops(para)
            set_paragraph_keep_lines(para, True)
            new_p = insert_paragraph_after(para)
            set_paragraph_text(new_p, date_line, bold=False)
            _format_date_plaque_paragraph(new_p, para)
            changed += 1
            i += 2
            continue
        if is_signatory_date_plaque(text):
            compact = compact_signatory_date_plaque(text)
            if compact != text:
                set_paragraph_text(para, compact, bold=False)
                changed += 1
            iof = _previous_iof_paragraph(doc, i, raz)
            if iof is not None:
                ensure_signatory_tab_stops(iof)
                set_paragraph_keep_lines(iof, True)
            _format_date_plaque_paragraph(para, iof)
        i += 1
    return changed


def _fold_job_inflections(text: str) -> str:
    t = text
    replacements = (
        (r"\bзам(?:естител[ьяюе]|)\.?\b", "заместитель"),
        (r"\bинженера\b", "инженер"),
        (r"\bинженеру\b", "инженер"),
        (r"\bстаршего\b", "старший"),
        (r"\bмастера\b", "мастер"),
        (r"\bначальника\b", "начальник"),
        (r"\bначальнику\b", "начальник"),
        (r"\bслужбы\b", "служба"),
        (r"\bотдела\b", "отдел"),
    )
    for pat, repl in replacements:
        t = re.sub(pat, repl, t)
    return t


def normalize_job_title(text: str) -> str:
    """Нижний регистр, ё=е, дубли слов, падежи должности к именительному."""
    t = (text or "").replace("\xa0", " ").replace("ё", "е").replace("Ё", "е")
    t = t.lower()
    t = collapse_adjacent_duplicate_words(t)
    t = re.sub(r"[_.,;:()«»\"'+]+", " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    t = _fold_job_inflections(t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def extract_acquaintance_job_title(text: str) -> str:
    packed = " ".join((text or "").replace("\xa0", " ").split())
    match = _JOB_FROM_ACQ_RE.search(packed)
    if not match:
        return ""
    return normalize_job_title(match.group(1))


def _job_from_filename(path: Path | str | None) -> str:
    if path is None:
        return ""
    name = Path(path).stem.lower().replace("ё", "е")
    name = name.replace("_оформлен+", "").replace("_оформлен", "")
    name = re.sub(r"\d{1,2}[.\-]\d{1,2}[.\-]\d{2,4}", " ", name)
    name = re.sub(r"\b(ди|проект|оформлен)\b", " ", name)
    name = re.sub(r"\s+", " ", name).strip(" _-")
    return normalize_job_title(name)


def extract_instruction_job_title(doc: Document, path: Path | str | None = None) -> str:
    """Чья инструкция: назначение / титул / имя файла."""
    raz = get_signatory_start_index(doc)
    limit = raz if raz is not None else min(len(doc.paragraphs), 90)
    for paragraph in doc.paragraphs[:limit]:
        raw = paragraph.text or ""
        if not raw.strip():
            continue
        m_app = _JOB_FROM_APPOINTMENT_RE.search(raw)
        if m_app:
            job = normalize_job_title(m_app.group(1))
            if job:
                return job
        if is_chapter_header(raw) or is_acquaintance_sheet_text(raw):
            continue
        m_title = _JOB_FROM_TITLE_RE.search(raw)
        if m_title:
            job = normalize_job_title(m_title.group(1))
            if job and "утвержд" not in job:
                return job
    from_file = _job_from_filename(path)
    if from_file:
        return from_file
    return ""


def jobs_equivalent(instruction_job: str, acquaintance_job: str) -> bool:
    a = normalize_job_title(instruction_job)
    b = normalize_job_title(acquaintance_job)
    if not a or not b:
        return True
    if a == b:
        return True
    if a in b or b in a:
        return True
    tokens_a = {tok for tok in a.split() if len(tok) > 2}
    tokens_b = {tok for tok in b.split() if len(tok) > 2}
    if not tokens_a or not tokens_b:
        return True
    common = tokens_a & tokens_b
    if tokens_a <= tokens_b or tokens_b <= tokens_a:
        return True
    distinctive_a = tokens_a - {"заместитель", "начальник", "служба", "отдел"}
    distinctive_b = tokens_b - {"заместитель", "начальник", "служба", "отдел"}
    if distinctive_a and distinctive_b and distinctive_a.isdisjoint(distinctive_b):
        return False
    return bool(common) and len(common) >= min(2, len(tokens_a), len(tokens_b))


def paragraph_shd_fill(paragraph: Paragraph) -> str:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return ""
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        return ""
    return _shd_fill_value(shd)


def set_paragraph_shading(paragraph: Paragraph, fill: str | None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for shd in list(p_pr.findall(qn("w:shd"))):
        p_pr.remove(shd)
    if not fill:
        return
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def paragraph_has_mismatch_shading(paragraph: Paragraph) -> bool:
    return paragraph_shd_fill(paragraph) in {ACQUAINTANCE_MISMATCH_FILL, "RED", "C00000"}


def mark_acquaintance_heading_if_job_mismatch(
    doc: Document, path: Path | str | None = None
) -> int:
    """
    Если лист ознакомления — для другой должности, чем эта инструкция,
    заголовок листа залить красным. Совпадение — снять красный.
    """
    idx = find_acquaintance_sheet_start(doc)
    if idx is None:
        return 0
    heading = doc.paragraphs[idx]
    instr = extract_instruction_job_title(doc, path)
    acq = extract_acquaintance_job_title(heading.text)
    if not acq:
        if paragraph_has_mismatch_shading(heading):
            set_paragraph_shading(heading, None)
            return 1
        return 0
    mismatch = not jobs_equivalent(instr, acq)
    has_red = paragraph_has_mismatch_shading(heading)
    if mismatch and not has_red:
        set_paragraph_shading(heading, ACQUAINTANCE_MISMATCH_FILL)
        return 1
    if not mismatch and has_red:
        set_paragraph_shading(heading, None)
        return 1
    return 0


def validate_item_number_spacing(doc: Document) -> list[str]:
    """Нет ведущих пробелов перед номером; после номера ровно один пробел."""
    issues: list[str] = []
    razrab_idx = get_signatory_start_index(doc)
    for idx, paragraph in enumerate(doc.paragraphs):
        original = paragraph_xml_visible_text(paragraph) or (paragraph.text or "")
        if not original.strip():
            continue
        if razrab_idx is not None and idx >= razrab_idx:
            continue
        expected = normalize_number_separator(original)
        if expected == original:
            continue
        snippet = original[:55] + ("…" if len(original) > 55 else "")
        if original[:1] in " \t" or original.startswith("\xa0"):
            issues.append(f"Пробел/таб перед номером пункта: «{snippet}»")
        else:
            issues.append(f"После номера должен быть ровно один пробел: «{snippet}»")
        if len(issues) >= 8:
            break
    return issues


def validate_adjacent_duplicate_words(doc: Document) -> list[str]:
    issues: list[str] = []
    for paragraph in doc.paragraphs:
        original = paragraph.text or ""
        if not original.strip():
            continue
        collapsed = collapse_adjacent_duplicate_words(original)
        if collapsed == original:
            continue
        snippet = original[:55] + ("…" if len(original) > 55 else "")
        issues.append(f"Повтор слова подряд: «{snippet}»")
        if len(issues) >= 5:
            break
    return issues


def validate_signatory_fio_one_line(doc: Document, profile: DocumentProfile) -> list[str]:
    """Перенос ФИО на вторую строку / мягкий разрыв — замечание."""
    if not profile.has_signatories:
        return []
    try:
        raz = find_razrabotal_index(doc)
    except ValueError:
        return []
    issues: list[str] = []
    end = _signatory_block_end_index(doc)
    for i in range(raz, end):
        para = doc.paragraphs[i]
        text = para.text or ""
        if is_acquaintance_sheet_text(text):
            break
        if i + 1 < end:
            nxt = doc.paragraphs[i + 1].text or ""
            if _looks_like_job_without_fio(text) and _is_iof_only_signatory_line(nxt):
                issues.append(
                    "ФИО подписанта перенесено на вторую строку — "
                    "должность и ФИО должны быть на одной строке через Tab"
                )
                break
        if paragraph_has_soft_line_break(para) and SIGNATORY_NAME_TAIL.search(text):
            issues.append("Перенос ФИО внутри строки подписанта (разрыв строки)")
            break
        if "\t" in text and SIGNATORY_NAME_TAIL.search(text):
            if not _paragraph_has_keep_lines(para):
                issues.append(
                    "Строка должность+ФИО без keepLines (ФИО может перенестись)"
                )
                break
    return issues


def validate_signatory_date_plaques(doc: Document, profile: DocumentProfile) -> list[str]:
    if not profile.has_signatories:
        return []
    issues: list[str] = []
    issues.extend(validate_signatory_fio_one_line(doc, profile))
    try:
        raz = find_razrabotal_index(doc)
    except ValueError:
        return issues
    for i in range(raz, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text or ""
        if is_acquaintance_sheet_text(text):
            break
        date_match = SIGNATORY_DATE_PLAQUE_RE.search(text)
        if date_match and SIGNATORY_NAME_TAIL.search(text[: date_match.start()]):
            issues.append("Плашка даты должна быть на отдельной строке, выровнена с И.О.Ф.")
            break
        if not is_signatory_date_plaque(text):
            continue
        compact = compact_signatory_date_plaque(text)
        stripped = text.replace("\xa0", " ")
        if not stripped.startswith("\t"):
            issues.append("Плашка даты не выровнена с началом И.О.Ф. (нет табуляции)")
        iof = _previous_iof_paragraph(doc, i, raz)
        if iof is not None:
            date_pos = signatory_first_tab_pos_twips(para)
            iof_pos = signatory_first_tab_pos_twips(iof)
            if date_pos is None or iof_pos is None or date_pos != iof_pos:
                issues.append(
                    "Дата не выровнена с началом ФИО (табуляция не совпадает)"
                )
        if compact.lstrip("\t") != stripped.strip().lstrip("\t") and stripped.count("_") > 12:
            issues.append("Плашка даты слишком длинная — переносится на две строки")
        if len(issues) >= 5:
            break
    return issues


def validate_acquaintance_job_mismatch(
    doc: Document, path: Path | str | None = None
) -> list[str]:
    idx = find_acquaintance_sheet_start(doc)
    if idx is None:
        return []
    heading = doc.paragraphs[idx]
    acq = extract_acquaintance_job_title(heading.text)
    if not acq:
        return []
    instr = extract_instruction_job_title(doc, path)
    mismatch = not jobs_equivalent(instr, acq)
    has_red = paragraph_has_mismatch_shading(heading)
    if mismatch and not has_red:
        return [
            "Лист ознакомления: должность не совпадает с инструкцией — "
            "заголовок должен быть с красной заливкой"
        ]
    if not mismatch and has_red:
        return ["Ложная красная заливка на листе ознакомления (должности совпадают)"]
    return []


def is_agreement_signatory_table(table) -> bool:
    """Таблица согласования (должность | ФИО), не титул и не лист ознакомления."""
    try:
        cols = _table_col_count(table)
        if cols < 2 or cols > 3 or len(table.rows) < 2:
            return False
    except (ValueError, AttributeError, IndexError):
        return False
    blob = " ".join(cell.text for row in table.rows for cell in row.cells).upper()
    if "УТВЕРЖДАЮ" in blob or "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ" in blob:
        return False
    if _is_acquaintance_sheet_table(table):
        return False
    job_hits = 0
    name_hits = 0
    for row in table.rows:
        pair = _row_job_and_name([cell.text for cell in row.cells])
        if pair is None:
            continue
        left, right = pair
        low = left.lower()
        if any(
            hint in low
            for hint in (
                "начальник",
                "инженер",
                "директор",
                "главн",
                "ведущий",
                "заместител",
            )
        ):
            job_hits += 1
        if SIGNATORY_NAME_TAIL.search(right) or re.search(
            r"[А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁа-яё]", right
        ):
            name_hits += 1
    return job_hits >= 2 or (job_hits >= 1 and name_hits >= 2)


def find_agreement_signatory_table(doc: Document):
    """Последняя таблица согласования в теле (не титул, не лист ознакомления)."""
    found = None
    for candidate in doc.tables:
        if is_agreement_signatory_table(candidate):
            found = candidate
    return found


def _insert_paragraph_before_element(doc: Document, element, text: str | None) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    para = Paragraph(new_p, doc)
    if text:
        set_paragraph_text(para, text, bold=False)
        if "\t" in text:
            ensure_signatory_tab_stops(para)
        set_one_point_five_line_spacing(para)
        clear_first_line_indent(para)
    return para


def materialize_signatory_paragraphs_from_tables(doc: Document) -> int:
    """
    Таблица согласования → блок без таблицы:
    «Разработал:», должность+ФИО (Tab), пустая, «Согласовано:», строки согласующих.
    Первая строка таблицы — разработчик, остальные — согласующие.
    """
    has_raz = False
    has_sog = False
    for paragraph in doc.paragraphs:
        if paragraph_is_inside_table(paragraph):
            continue
        upper = paragraph_text_normalized(paragraph).upper()
        if upper.startswith("РАЗРАБОТАЛ"):
            has_raz = True
        if upper.startswith("СОГЛАСОВАН"):
            has_sog = True
    table = find_agreement_signatory_table(doc)
    if table is None:
        return 0
    rows = extract_agreement_table_rows(table)
    if len(rows) < 2:
        return 0
    tbl = table._tbl
    changed = 0
    if has_raz and has_sog:
        parent = tbl.getparent()
        if parent is not None:
            parent.remove(tbl)
            changed += 1
        return changed
    lines: list[str | None] = []
    if not has_raz:
        lines.append("Разработал:")
        lines.append(f"{rows[0][0]}\t{rows[0][1]}")
        agree_rows = rows[1:]
    else:
        agree_rows = rows
    lines.append(None)
    if not has_sog:
        lines.append("Согласовано:")
    for job, name in agree_rows:
        lines.append(f"{job}\t{name}")
    for text in lines:
        _insert_paragraph_before_element(doc, tbl, text)
        changed += 1
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)
        changed += 1
    return changed


def paragraph_has_tab_stops(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        return False
    return len(tabs.findall(qn("w:tab"))) > 0


def copy_signatory_ppr_from_etalon(
    etalon_paragraph: Paragraph,
    target_paragraph: Paragraph,
) -> bool:
    """Скопировать w:tabs, jc, spacing с образца; без отступа 1,25 см."""
    e_ppr = etalon_paragraph._p.find(qn("w:pPr"))
    if e_ppr is None:
        ensure_paragraph_justified(target_paragraph)
        clear_first_line_indent(target_paragraph)
        set_one_point_five_line_spacing(target_paragraph)
        return True

    t_ppr = target_paragraph._p.get_or_add_pPr()
    changed = False

    for tag in ("w:tabs", "w:jc", "w:spacing"):
        existing = t_ppr.find(qn(tag))
        if existing is not None:
            t_ppr.remove(existing)
        e_el = e_ppr.find(qn(tag))
        if e_el is not None:
            t_ppr.append(deepcopy(e_el))
            changed = True

    if t_ppr.find(qn("w:jc")) is None:
        ensure_paragraph_justified(target_paragraph)
        changed = True

    clear_first_line_indent(target_paragraph)
    if not paragraph_has_one_point_five_spacing(target_paragraph):
        set_one_point_five_line_spacing(target_paragraph)
        changed = True

    return changed


def _etalon_signatory_index_map(etalon_doc: Document, e_razrab: int) -> dict[str, int]:
    """Сопоставление строк подписантов образца по ключу (ФИО / маркер)."""
    mapping: dict[str, int] = {}
    for e_i in range(e_razrab, len(etalon_doc.paragraphs)):
        key = paragraph_match_key(etalon_doc.paragraphs[e_i].text)
        if key and key not in mapping:
            mapping[key] = e_i
    return mapping


def apply_signatory_layout_from_etalon(
    doc: Document,
    etalon_doc: Document,
    profile: DocumentProfile,
) -> int:
    """
    Полная синхронизация блока подписантов с *_образец.docx:
    табуляция в тексте, w:tabs в pPr, выравнивание, интервал 1,5.
    """
    if not profile.has_signatories or etalon_doc is None:
        return 0
    try:
        razrab_idx = find_razrabotal_index(doc)
        e_razrab = find_razrabotal_index(etalon_doc)
    except ValueError:
        return 0

    changed = 0
    d_len = len(doc.paragraphs)
    e_len = len(etalon_doc.paragraphs)
    e_by_key = _etalon_signatory_index_map(etalon_doc, e_razrab)

    for d_i in range(razrab_idx, d_len):
        d_para = doc.paragraphs[d_i]
        d_text = d_para.text
        d_key = paragraph_match_key(d_text)
        e_i = e_by_key.get(d_key) if d_key else None
        if e_i is None:
            e_i = e_razrab + (d_i - razrab_idx)
        if e_i >= e_len:
            break
        e_para = etalon_doc.paragraphs[e_i]
        e_text = e_para.text

        if d_text.strip() and e_text.strip():
            e_key = paragraph_match_key(e_text)
            new_text = d_text
            if paragraph_keys_match(d_key, e_key) and "\t" in e_text and "\t" not in d_text:
                new_text = insert_signatory_tab_line(d_text)
            elif "\t" in e_text and "\t" not in d_text:
                new_text = insert_signatory_tab_line(d_text)
            if new_text != d_text:
                set_paragraph_text(d_para, new_text)
                changed += 1
                d_para = doc.paragraphs[d_i]

        if copy_signatory_ppr_from_etalon(e_para, d_para):
            changed += 1

    return changed


BODY_STYLES_NORMALIZE = frozenset(
    {
        "Plain Text",
        "No Spacing",
        "List Paragraph",
        "List",
        "Body Text",
        "Body Text 2",
        "Body Text 3",
        "Compact",
        "Quote",
        "Intense Quote",
        "HTML Preformatted",
    }
)


def normalize_body_paragraph_styles(doc: Document) -> int:
    """Стили Plain Text / No Spacing ломают отображение отступа 1,25 см в Word."""
    changed = 0
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        return 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if not should_apply_body_paragraph_format(paragraph.text, idx, doc):
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name in BODY_STYLES_NORMALIZE:
            paragraph.style = normal
            changed += 1
    return changed


def apply_number(text: str, number: str) -> str:
    return f"{number}. {strip_number(text)}"


def _path_parts_folded(path: Path | str) -> list[str]:
    return [p.replace("ё", "е").casefold() for p in Path(str(path)).parts]


def _looks_like_user_agent_dir(path: Path | str) -> bool:
    """N: и UNC — один каталог Агент. ОБМЕН / САТП не подходят."""
    parts = _path_parts_folded(path)
    if any("обмен" in p for p in parts):
        return False
    for i, part in enumerate(parts):
        if part == "агент" and i >= 1 and "дубовик" in parts[i - 1]:
            return True
    return False


def _looks_like_writable_user_dir(path: Path | str) -> bool:
    """N: и UNC: Агент или Проекты под «Дубовик …». ОБМЕН / САТП — нет."""
    parts = _path_parts_folded(path)
    if any("обмен" in p for p in parts):
        return False
    for i, part in enumerate(parts):
        if part in _WRITABLE_FOLDER_NAMES and i >= 1 and "дубовик" in parts[i - 1]:
            return True
    return False


def _is_under_writable_bases(path: Path | str) -> bool:
    try:
        p = normalize_sniot_path_text(path)
        resolved = p.resolve()
        for base_raw in WRITABLE_USER_DIRS:
            try:
                base = normalize_sniot_path_text(base_raw).resolve()
            except OSError:
                base = Path(os.path.normpath(str(normalize_sniot_path_text(base_raw))))
            try:
                if resolved == base or resolved.is_relative_to(base):
                    return True
            except (OSError, ValueError, AttributeError):
                pass
            a = os.path.normcase(str(resolved))
            b = os.path.normcase(str(base))
            if a == b or a.startswith(b.rstrip("\\") + "\\"):
                return True
        return False
    except (OSError, ValueError, TypeError):
        return False


def is_path_in_user_agent_dir(path: Path | str) -> bool:
    """Путь внутри папки Агент. N: и UNC — один каталог (образцы / умолчания)."""
    try:
        p = normalize_sniot_path_text(path)
        if _looks_like_user_agent_dir(p):
            return True
        resolved = p.resolve()
        for base_raw in (USER_AGENT_DIR, USER_AGENT_DIR_UNC):
            try:
                base = normalize_sniot_path_text(base_raw).resolve()
            except OSError:
                base = Path(os.path.normpath(str(normalize_sniot_path_text(base_raw))))
            try:
                if resolved == base or resolved.is_relative_to(base):
                    return True
            except (OSError, ValueError, AttributeError):
                pass
            a = os.path.normcase(str(resolved))
            b = os.path.normcase(str(base))
            if a == b or a.startswith(b.rstrip("\\") + "\\"):
                return True
        return False
    except (OSError, ValueError, TypeError):
        return False


def is_path_in_writable_user_dir(path: Path | str) -> bool:
    """Путь в папке Агент или Проекты — зона записи."""
    try:
        p = normalize_sniot_path_text(path)
        if _looks_like_writable_user_dir(p):
            return True
        return _is_under_writable_bases(p)
    except (OSError, ValueError, TypeError):
        return False


def writable_dirs_hint() -> str:
    return f"{USER_AGENT_DIR}\nили\n{USER_PROJECT_DIR}"


def is_path_readonly_sample(path: Path | str) -> bool:
    """ОБМЕН больше не эталон. Оставлено для совместимости: всегда False."""
    return False


def filename_has_sample_mark(name: str) -> bool:
    """В имени файла есть подстрока «образец» (регистр не важен)."""
    return SAMPLE_NAME_MARK in (name or "").casefold()


def is_allowed_sample_path(path: Path | str) -> bool:
    """
    Эталон оформления: только папка Агент, в имени «образец», расширение .docx.
    Файлы *_оформлен.docx без слова «образец» — False. ОБМЕН / САТП — False.
    Существование файла не требуется (проверка пути и имени).
    """
    try:
        p = normalize_sniot_path_text(path)
        name = p.name
        if name.startswith("~$"):
            return False
        if p.suffix.lower() not in SAMPLE_EXTENSIONS:
            return False
        if not filename_has_sample_mark(name):
            return False
        stem_l = p.stem.casefold()
        if stem_l.endswith("_оформлен") or stem_l.endswith("_оформлен+"):
            return False
        return is_path_in_user_agent_dir(p)
    except (OSError, ValueError, TypeError):
        return False


def list_agent_sample_paths() -> list[Path]:
    """Все подходящие *.docx со словом «образец» только в папке Агент (без обхода N:\\)."""
    if not USER_AGENT_DIR.is_dir():
        return []
    found: list[Path] = []
    try:
        for item in USER_AGENT_DIR.iterdir():
            if not item.is_file():
                continue
            if is_allowed_sample_path(item):
                found.append(item)
    except OSError:
        return []
    return found


def _sample_rank(target: Path, sample: Path) -> tuple[int, int]:
    """Чем ближе имя образца к целевому документу, тем выше ранг."""
    target_base = re.sub(r"_оформлен\+?$", "", target.stem, flags=re.IGNORECASE)
    target_base = re.sub(re.escape(SAMPLE_NAME_MARK), "", target_base, flags=re.IGNORECASE)
    target_base = target_base.strip(" _").casefold()
    sample_core = re.sub(re.escape(SAMPLE_NAME_MARK), "", sample.stem, flags=re.IGNORECASE)
    sample_core = sample_core.strip(" _").casefold()
    exact = 1 if sample_core == target_base else 0
    t_tokens = set(re.findall(r"[а-яёa-z0-9]+", target_base, flags=re.IGNORECASE))
    s_tokens = set(re.findall(r"[а-яёa-z0-9]+", sample_core, flags=re.IGNORECASE))
    overlap = len(t_tokens & s_tokens)
    return (exact, overlap)


def assert_path_writable(path: Path | str) -> Path:
    """Guard: запись только в папку Агент или Проекты."""
    p = normalize_sniot_path_text(path)
    if not is_path_in_writable_user_dir(p):
        raise PermissionError(
            f"Запись запрещена вне папок Агент/Проекты: {p}\n"
            f"Разрешено:\n{writable_dirs_hint()}\n"
            "Образец — только файл со словом «образец» в папке Агент."
        )
    return p.resolve()


def resolve_from_handoff() -> Path | None:
    """Путь из DocAgent handoff/request_latest.json (Агент или Проекты)."""
    if not DOCAGENT_HANDOFF.is_file():
        return None
    try:
        data = json.loads(DOCAGENT_HANDOFF.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    raw = (data.get("source_path") or "").strip()
    if not raw:
        return None
    path = Path(raw)
    if path.is_file() and is_path_in_writable_user_dir(path):
        return path
    return None


def find_user_agent_default() -> Path | None:
    """Файл по умолчанию в папке Агент пользователя."""
    preferred = USER_AGENT_DIR / DEFAULT_TARGET_NAME
    if preferred.is_file():
        return preferred
    plus = USER_AGENT_DIR / DEFAULT_TARGET_PLUS_NAME
    if plus.is_file():
        return plus
    if not USER_AGENT_DIR.is_dir():
        return None
    matches = sorted(
        (
            f
            for f in USER_AGENT_DIR.glob("*оформлен*.docx")
            if "_backup_" not in f.name.lower()
        ),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    return matches[0] if matches else None


def detect_document_kind(path: Path) -> str:
    name = path.name.lower()
    if "рабоч" in name and "инструкц" in name:
        return "ri"
    if "омтс" in name or ("зам" in name and "начальник" in name):
        return "di"
    if "должностн" in name or name.startswith("ди ") or "ди_" in name:
        return "di"
    if "положен" in name:
        return "polozhenie"
    if "проект" in name and any(w in name for w in ("мастер", "диспетчер", "инженер", "начальник")):
        return "di"
    if "старш" in name and "мастер" in name:
        return "di"
    return "generic"


def find_first_chapter_text(doc: Document) -> str | None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if is_chapter_header(text):
            return text
        upper = text.upper()
        if upper.startswith("ОБЩИЕ ПОЛОЖЕНИЯ"):
            return "1 ОБЩИЕ ПОЛОЖЕНИЯ"
    return None


def count_nonempty_body_paragraphs(doc: Document) -> int:
    return sum(1 for paragraph in doc.paragraphs if paragraph.text.strip())


def count_numbered_paragraphs(doc: Document) -> int:
    """Абзацы с текстовой нумерацией вида 1.4.1."""
    return sum(1 for p in doc.paragraphs if NUM_PREFIX.match((p.text or "").strip()))


def is_senior_master_di_path(path: Path | str) -> bool:
    """ДИ «Старший мастер» САТП — консервативный режим без learned text_edits."""
    name = Path(path).name.lower()
    if "мастер" in name and ("проект" in name or "оформлен" in name):
        return True
    low = str(path).lower()
    return "сатп" in low and "мастер" in name


def is_conservative_di_satp(path: Path | str) -> bool:
    """Любая ДИ/РИ/положение СНиОТ — без text_edits (слова сравнения файлов / коды ТКП)."""
    if is_senior_master_di_path(path):
        return True
    name = Path(path).name.lower()
    if name.startswith("ди ") or "должностн" in name:
        return True
    if "рабоч" in name and "инструкц" in name:
        return True
    if "положен" in name:
        return True
    return False


def normalize_sniot_path_text(path: Path | str) -> Path:
    """Латинская «i» в «СНiОТ» → кириллическая «и» (СНиОТ); «/» → «\\»."""
    normalized = str(path).replace("/", "\\")
    for wrong in ("СНiОТ", "СНIОТ", "СNiОТ", "СNiOT"):
        normalized = normalized.replace(wrong, "СНиОТ")
    return Path(normalized)


def paths_are_same_file(src: Path | str, dst: Path | str) -> bool:
    """True, если пути с «/» и «\\» указывают на один файл."""
    if not src or not dst:
        return False
    a = Path(str(src).replace("/", "\\"))
    b = Path(str(dst).replace("/", "\\"))
    try:
        if a.exists() and b.exists():
            return os.path.samefile(a, b)
    except OSError:
        pass
    try:
        return os.path.normcase(str(a.resolve())) == os.path.normcase(str(b.resolve()))
    except OSError:
        return os.path.normcase(os.path.abspath(str(a))) == os.path.normcase(
            os.path.abspath(str(b))
        )


def copy_file_if_different(src: Path | str, dst: Path | str) -> bool:
    """copy2 без падения, если источник и приёмник — один файл. True — скопировали."""
    if paths_are_same_file(src, dst):
        return False
    try:
        shutil.copy2(src, dst)
    except shutil.SameFileError:
        return False
    return True


def try_close_word_document(path: Path | str) -> dict:
    """Закрыть ТОЛЬКО этот документ в Word. Word.Quit не вызывать."""
    result = {"was_open": False, "closed": False, "message": ""}
    target = str(path)
    if pythoncom is None or win32com is None:
        result["message"] = "Нет win32com — закройте документ в Word вручную"
        return result
    pythoncom.CoInitialize()
    try:
        try:
            word = win32com.client.GetActiveObject("Word.Application")
        except Exception:
            result["message"] = "Word не запущен"
            return result
        matches = []
        try:
            count = int(word.Documents.Count)
        except Exception:
            count = 0
        target_name = Path(target).name.casefold()
        for i in range(1, count + 1):
            try:
                doc = word.Documents(i)
                full = str(doc.FullName or "")
            except Exception:
                continue
            same = False
            try:
                same = paths_are_same_file(full, target)
            except Exception:
                same = False
            if not same:
                a = os.path.normcase(os.path.normpath(full))
                b = os.path.normcase(os.path.normpath(target))
                same = a == b or (
                    Path(full).name.casefold() == target_name
                    and "агент" in a
                    and "агент" in b
                )
            if same:
                matches.append(doc)
        if not matches:
            result["message"] = "Этот файл в Word не открыт"
            return result
        result["was_open"] = True
        for doc in matches:
            try:
                doc.Save()
            except Exception:
                pass
            try:
                doc.Close(False)
            except Exception as exc:
                result["closed"] = False
                result["message"] = f"Word не отдал файл: {exc}"
                return result
        result["closed"] = True
        result["message"] = "Документ закрыт в Word (окно Word оставлено)"
        return result
    except Exception as exc:
        result["message"] = f"Не удалось закрыть в Word: {exc}"
        return result
    finally:
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def copy_to_target_unlocking_word(src: Path, dst: Path) -> None:
    """Запись в Агент; если Word держит файл — закрыть этот документ и повторить."""
    try:
        copy_file_if_different(src, dst)
        return
    except (PermissionError, OSError) as first:
        closed = try_close_word_document(dst)
        if closed.get("closed"):
            copy_file_if_different(src, dst)
            return
        raise PermissionError(
            "СНиОТ: не удалось сохранить — закройте «_оформлен.docx» в Word и повторите"
        ) from first


def find_marker_index(doc: Document, *markers: str, contains: bool = False) -> int:
    """Найти абзац по одному из маркеров (startswith или contains)."""
    last_err: ValueError | None = None
    for marker in markers:
        try:
            return find_paragraph_index(doc, marker, contains=contains)
        except ValueError as exc:
            last_err = exc
    raise ValueError(f"Paragraph not found: {markers!r}") from last_err


def validate_save_integrity(
    *,
    before_nonempty: int,
    before_numbered: int,
    after_doc: Document,
    profile: DocumentProfile,
) -> list[str]:
    """
    Запрет сохранения, если тело или нумерация резко пропали.
    Для ДИ САТП: если была нумерация — нельзя сохранять с нулём пунктов.
    """
    issues: list[str] = []
    after_nonempty = count_nonempty_body_paragraphs(after_doc)
    after_numbered = count_numbered_paragraphs(after_doc)

    if before_nonempty >= 10:
        drop = (before_nonempty - after_nonempty) / before_nonempty
        if drop > 0.10:
            issues.append(
                f"Потеря текста: было {before_nonempty} абзацев, стало {after_nonempty} "
                f"({drop:.0%}) — сохранение отменено"
            )

    if profile.has_di_satp_numbering or before_numbered >= 5:
        if before_numbered >= 5 and after_numbered == 0:
            issues.append(
                f"Нумерация исчезла: было {before_numbered} пунктов, стало 0 — сохранение отменено"
            )
        elif before_numbered >= 10 and after_numbered < before_numbered * 0.5:
            issues.append(
                f"Нумерация сильно уменьшилась: было {before_numbered}, стало {after_numbered}"
            )
    return issues


def reload_document_from_path(docx_path: Path) -> tuple[Document, bytes]:
    """Перечитать весь docx с диска (финальная проверка — не объект Document из памяти)."""
    docx_bytes = docx_path.read_bytes()
    return Document(BytesIO(docx_bytes)), docx_bytes


def validate_full_document_on_disk(
    docx_path: Path,
    *,
    path_for_profile: Path | None = None,
) -> list[str]:
    """
    Обязательная финальная проверка перед записью в папку Агент:
    1) перечитать файл целиком с диска (не Document из памяти);
    2) validate_sniot_document — все правила mdc;
    3) дополнительные gate-проверки структуры гл. 1 и блоков 1.8.x / 1.9.x.

    При любом замечании apply_sniot_rules_to_file **не** перезаписывает *_оформлен.docx.
    """
    doc, docx_bytes = reload_document_from_path(docx_path)
    profile_path = path_for_profile or docx_path
    profile = detect_profile(doc, profile_path)
    issues = validate_sniot_document(
        doc, docx_bytes=docx_bytes, profile=profile, path=profile_path
    )
    extra = validate_final_document_gate(doc, profile, path=profile_path)
    if extra:
        seen = set(issues)
        for item in extra:
            if item not in seen:
                issues.append(item)
                seen.add(item)
    return issues


def categorize_validation_issue(issue: str) -> str:
    """Группа замечания для отчёта финальной проверки."""
    low = issue.lower()
    if "разработал" in low:
        return "подписанты"
    if "подписант" in low or "согласован" in low or "оторван" in low or "разрыв" in low:
        return "подписанты"
    if "глава 1" in low or "блок 1." in low or "1.4.x" in low or "1.5.x" in low or "1.9.x" in low:
        return "глава 1 / нумерация"
    if "номер" in low or "префикс" in low or "numpr" in low or "нумерац" in low:
        return "нумерация"
    if "отступ" in low or "1,25" in low or "1.25" in low:
        return "отступы"
    if "шрифт" in low or "times" in low or "14 pt" in low:
        return "шрифт"
    if "страниц" in low or "колонтитул" in low or "поля" in low:
        return "страницы"
    return "прочее"


def summarize_validation_issues(issues: list[str]) -> dict[str, int]:
    summary: dict[str, int] = {}
    for item in issues:
        key = categorize_validation_issue(item)
        summary[key] = summary.get(key, 0) + 1
    return summary


def format_validation_summary(issues: list[str]) -> str:
    if not issues:
        return "0 замечаний"
    parts = [
        f"{name}: {count}"
        for name, count in sorted(summarize_validation_issues(issues).items())
    ]
    return f"{len(issues)} замечаний ({'; '.join(parts)})"


def parse_debug_dump(dump_path: Path) -> list[tuple[bool, str]]:
    entries: list[tuple[bool, str]] = []
    for line in dump_path.read_text(encoding="utf-8").splitlines():
        match = re.match(r"^\s*\d+\s+b=(\d+)\s*\|\s*(.*)$", line)
        if match:
            entries.append((bool(int(match.group(1))), match.group(2)))
            continue
        match = re.match(r"^\s*\d+\s+font=\{[^}]*\}\s*\|\s*(.*)$", line)
        if match:
            text = match.group(1)
            bold = "b=1" in line or "Heading" in line
            entries.append((bold, text))
    return entries


def replace_body_from_debug_dump(
    docx_path: Path,
    dump_path: Path,
    *,
    from_marker: str = "1 ОБЩИЕ ПОЛОЖЕНИЯ",
) -> int:
    """
    Заменить тело документа из дампа export_debug (_work_*.txt с нумерацией).
    Титул (sdt/таблицы) не трогаем — только абзацы от первой главы до «Разработал:».
    """
    if not dump_path.is_file():
        raise FileNotFoundError(dump_path)

    body_entries: list[tuple[bool, str]] = []
    started = False
    for bold, text in parse_debug_dump(dump_path):
        t = text.strip()
        if not started:
            if from_marker in text or t.startswith("1 ОБЩИЕ"):
                started = True
            else:
                continue
        if t.lower().startswith("разработал"):
            break
        body_entries.append((bold, t))

    if not body_entries:
        raise ValueError(f"В дампе нет тела от «{from_marker}»: {dump_path}")

    doc = Document(docx_path)
    body_start: int | None = None
    for i, paragraph in enumerate(doc.paragraphs):
        t = paragraph.text.strip()
        if "1 ОБЩИЕ ПОЛОЖЕНИЯ" in t or t.startswith("1 ОБЩИЕ"):
            body_start = i
            break
    if body_start is None:
        for i, paragraph in enumerate(doc.paragraphs):
            if _is_body_start(paragraph.text, "1 ОБЩИЕ ПОЛОЖЕНИЯ"):
                body_start = i
                break
    if body_start is None:
        raise ValueError("Не найдено начало тела (глава 1) в docx")

    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        razrab_idx = len(doc.paragraphs)

    for idx in range(razrab_idx - 1, body_start - 1, -1):
        el = doc.paragraphs[idx]._element
        el.getparent().remove(el)

    try:
        razrab_idx = find_razrabotal_index(doc)
        anchor = doc.paragraphs[razrab_idx]
    except ValueError:
        tail = find_signatory_tail_start(doc)
        anchor = doc.paragraphs[tail] if tail is not None else doc.paragraphs[-1]

    for bold, text in reversed(body_entries):
        new_p = insert_empty_paragraph_before(anchor)
        set_paragraph_text(new_p, text, bold=bold)
        anchor = new_p

    doc.save(docx_path)
    return len(body_entries)


def _is_body_start(text: str, first_chapter: str | None) -> bool:
    t = text.strip()
    if not t:
        return False
    if first_chapter and (first_chapter in t or t.startswith(first_chapter[: min(20, len(first_chapter))])):
        return True
    if is_chapter_header(t):
        return True
    upper = t.upper()
    if upper.startswith("ОБЩИЕ ПОЛОЖЕНИЯ"):
        return True
    if upper in (
        "ФУНКЦИИ И ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ",
        "ПРАВА",
        "ВЗАИМООТНОШЕНИЯ",
        "ОТВЕТСТВЕННОСТЬ",
    ):
        return True
    if re.match(r"^\d+\s+[А-ЯЁ]", t):
        return True
    if re.match(r"^\d+\.\d+\.", t):
        return True
    return False


def _is_title_duplicate_paragraph(text: str) -> bool:
    upper = text.upper()
    return any(marker in upper for marker in TITLE_DUPLICATE_MARKERS)


def _is_org_title_paragraph(text: str) -> bool:
    """Шапка предприятия / УТВЕРЖДАЮ / «Минск 20xx» — титул, не дубль и не глава 1."""
    t = (text or "").replace("\xa0", " ").strip()
    if not t or is_chapter_header(t):
        return False
    if _is_title_duplicate_paragraph(t):
        return True
    upper = t.upper()
    hints = (
        "МИНСКИЙ ГОРОДСКОЙ",
        "КОММУНАЛЬНОЕ УНИТАРНОЕ",
        "ПРЕДПРИЯТИЕ ПО",
        "КОММУНАЛЬНЫХ ТЕПЛОВЫХ",
        "МИНСККОММУНТЕПЛОСЕТЬ",
        "УТВЕРЖДАЮ",
        "НОМЕР ИНСТРУКЦИИ",
    )
    if any(hint in upper for hint in hints):
        return True
    if is_city_year_paragraph(t):
        return True
    if re.fullmatch(r"[_…]{3,}", t.replace(" ", "")):
        return True
    return False


def has_signatory_block(doc: Document) -> bool:
    if find_signatory_tail_start(doc) is not None:
        return True
    try:
        find_soglasovano_index(doc)
        return True
    except ValueError:
        pass
    return any(is_agreement_signatory_table(table) for table in doc.tables)


def has_di_satp_numbering_structure(doc: Document, path: Path | str | None = None) -> bool:
    """ДИ старшего мастера САТП — не любая ДИ с «должен знать» / «руководствуется» (не ОМТС)."""
    name = Path(path).name.lower() if path else ""
    if "омтс" in name:
        return False
    blob = " ".join([name] + [p.text for p in doc.paragraphs[:90]]).lower()
    if not any(hint in blob for hint in SATP_STRUCTURE_HINTS):
        return False
    found = 0
    checks = (
        (DI_SATP_CH1_GUIDED, "в своей деятельности руководствуется"),
        (DI_SATP_CH1_KNOW, "должен знать"),
        ("2 ФУНКЦИИ", "ФУНКЦИИ И ДОЛЖНОСТНЫЕ"),
    )
    for exact, fuzzy in checks:
        try:
            if exact.startswith("2 "):
                find_paragraph_index(doc, exact)
            else:
                find_section_header_index(doc, exact, fuzzy)
            found += 1
        except ValueError:
            try:
                find_paragraph_index(doc, fuzzy, contains=True)
                found += 1
            except ValueError:
                pass
    return found >= 2


def find_tail_chapter_index(doc: Document) -> int | None:
    try:
        return find_paragraph_index(doc, "5 ОТВЕТСТВЕННОСТЬ")
    except ValueError:
        pass
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        razrab_idx = len(doc.paragraphs)
    chapters = [i for i in find_chapter_header_indices(doc) if i < razrab_idx]
    return chapters[-1] if chapters else None


def detect_profile(doc: Document, path: Path) -> DocumentProfile:
    return DocumentProfile(
        kind=detect_document_kind(path),
        first_chapter=find_first_chapter_text(doc),
        has_signatories=has_signatory_block(doc),
        has_di_satp_numbering=has_di_satp_numbering_structure(doc, path),
        tail_chapter_idx=find_tail_chapter_index(doc),
    )


def resolve_target(
    explicit: Path | None = None,
    *,
    use_handoff: bool = False,
) -> Path:
    if use_handoff:
        handoff_path = resolve_from_handoff()
        if handoff_path is not None:
            return handoff_path
        raise FileNotFoundError(f"Handoff не найден или файл недоступен: {DOCAGENT_HANDOFF}")

    if explicit is not None:
        path = explicit.expanduser()
        if not path.is_absolute():
            path = USER_AGENT_DIR / path
        if not path.exists():
            alt = normalize_sniot_path_text(path)
            if alt != path and alt.exists():
                path = alt
        if not path.exists():
            raise FileNotFoundError(path)
        if not is_path_in_writable_user_dir(path):
            raise FileNotFoundError(
                f"Путь вне разрешённых папок Агент/Проекты: {path}\n"
                f"Разрешено:\n{writable_dirs_hint()}"
            )
        return path

    handoff_path = resolve_from_handoff()
    if handoff_path is not None:
        return handoff_path

    default_path = find_user_agent_default()
    if default_path is not None:
        return default_path

    raise FileNotFoundError(
        f"Не найден документ: укажите --target, handoff или положите файл в "
        f"{USER_AGENT_DIR} или {USER_PROJECT_DIR}"
    )


def remove_duplicate_body_title(docx_bytes: bytes, first_chapter: str | None = None) -> bytes:
    """
    Удалить только повтор шапки/УТВЕРЖДАЮ на стр. 2 после sdt-титула.
    НИКОГДА не удалять всё тело, если маркер первой главы не найден.
    """
    with zipfile.ZipFile(BytesIO(docx_bytes), "r") as zin:
        xml = zin.read("word/document.xml")
        root = etree.fromstring(xml)
        body = root.find("w:body", NS)
        if body is None:
            return docx_bytes

        to_remove = []
        passed_sdt = False
        for child in list(body):
            tag = child.tag.split("}")[-1]
            if tag == "sectPr":
                break
            if tag == "sdt":
                passed_sdt = True
                continue
            if not passed_sdt:
                continue
            text = "".join(child.xpath(".//w:t/text()", namespaces=NS))
            text_stripped = text.strip()
            if not text_stripped:
                continue
            if _is_body_start(text_stripped, first_chapter):
                break
            if _is_title_duplicate_paragraph(text_stripped):
                to_remove.append(child)
            else:
                break

        if not to_remove:
            return docx_bytes

        for child in to_remove:
            body.remove(child)

        out = BytesIO()
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = xml if item.filename == "word/document.xml" else zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = etree.tostring(
                        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
                    )
                zout.writestr(item, data)
        return out.getvalue()


def _next_relationship_id(rels_root) -> str:
    nums: list[int] = []
    for rel in rels_root:
        rid = rel.get("Id") or ""
        if rid.startswith("rId"):
            try:
                nums.append(int(rid[3:]))
            except ValueError:
                pass
    return f"rId{(max(nums) if nums else 0) + 1}"


def _ensure_content_type_override(ct_root, part_name: str, content_type: str) -> None:
    for ov in ct_root.findall(f"{{{CT_NS}}}Override"):
        if ov.get("PartName") == part_name:
            return
    ov = etree.SubElement(ct_root, f"{{{CT_NS}}}Override")
    ov.set("PartName", part_name)
    ov.set("ContentType", content_type)


def _ensure_document_rel(rels_root, target: str, rel_type: str) -> str:
    for rel in rels_root:
        if rel.get("Target") == target:
            return rel.get("Id") or ""
    rid = _next_relationship_id(rels_root)
    rel = etree.SubElement(rels_root, f"{{{PKG_REL_NS}}}Relationship")
    rel.set("Id", rid)
    rel.set("Type", rel_type)
    rel.set("Target", target)
    return rid


def _ensure_sect_type(sect, value: str) -> None:
    el = sect.find("w:type", NS)
    if el is None:
        el = etree.Element(f"{{{W_NS}}}type")
        sect.insert(0, el)
    el.set(f"{{{W_NS}}}val", value)


def _ensure_title_pg_xml(sect, enabled: bool) -> None:
    el = sect.find("w:titlePg", NS)
    if enabled and el is None:
        sect.insert(0, etree.Element(f"{{{W_NS}}}titlePg"))
    if not enabled and el is not None:
        sect.remove(el)


def _set_section_header_references(sect, id_default: str, id_even: str, id_first: str) -> None:
    for ref in list(sect.findall("w:headerReference", NS)):
        sect.remove(ref)
    for typ, rid in (("even", id_even), ("default", id_default), ("first", id_first)):
        el = etree.Element(f"{{{W_NS}}}headerReference")
        el.set(f"{{{W_NS}}}type", typ)
        el.set(f"{{{R_NS}}}id", rid)
        sect.insert(0, el)


def fix_page_numbering(docx_bytes: bytes) -> bytes:
    """
    Номера со 2-й страницы, верх по центру; header1/header2/header3 —
    одинаковый TNR 14 не жирный; first (header3) пустой; titlePg.
    """
    with zipfile.ZipFile(BytesIO(docx_bytes), "r") as zin:
        files = {item.filename: zin.read(item.filename) for item in zin.infolist()}

    root = etree.fromstring(files["word/document.xml"])
    rels_name = "word/_rels/document.xml.rels"
    ct_name = "[Content_Types].xml"
    rels_root = etree.fromstring(files[rels_name])
    ct_root = etree.fromstring(files[ct_name])

    files["word/header1.xml"] = CENTERED_PAGE_HEADER
    files["word/header2.xml"] = CENTERED_PAGE_HEADER
    files["word/header3.xml"] = EMPTY_FIRST_HEADER
    for part in ("/word/header1.xml", "/word/header2.xml", "/word/header3.xml"):
        _ensure_content_type_override(ct_root, part, HEADER_CONTENT_TYPE)
    id_default = _ensure_document_rel(rels_root, "header1.xml", HEADER_REL_TYPE)
    id_even = _ensure_document_rel(rels_root, "header2.xml", HEADER_REL_TYPE)
    id_first = _ensure_document_rel(rels_root, "header3.xml", HEADER_REL_TYPE)

    sects = root.findall(".//w:sectPr", NS)
    for i, sect in enumerate(sects):
        _set_section_header_references(sect, id_default, id_even, id_first)
        is_first = i == 0
        is_last = i == len(sects) - 1
        _ensure_title_pg_xml(sect, enabled=is_first)
        if is_first and not is_last:
            _ensure_sect_type(sect, "nextPage")
        for pg in sect.findall("w:pgNumType", NS):
            sect.remove(pg)

    files["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )
    files[rels_name] = etree.tostring(
        rels_root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )
    files[ct_name] = etree.tostring(
        ct_root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )
    for name in list(files):
        if name.startswith("word/footer") and name.endswith(".xml"):
            files[name] = EMPTY_FOOTER

    out = BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    return out.getvalue()


def _xml_has_page_field(xml_bytes: bytes) -> bool:
    text = xml_bytes.decode("utf-8", errors="ignore")
    return " PAGE " in text or ">PAGE<" in text or "instrText" in text and "PAGE" in text


def _header_has_bold_page_number(xml_text: str) -> bool:
    for match in re.finditer(r"<w:b(?:Cs)?\b([^>]*)/?>", xml_text):
        attrs = (match.group(1) or "").lower()
        if "false" in attrs or 'val="0"' in attrs:
            continue
        return True
    return False


def _header_page_style_issues(name: str, xml_text: str) -> list[str]:
    issues: list[str] = []
    if "Times New Roman" not in xml_text:
        issues.append(f"{name}: номер страницы не Times New Roman")
    if 'w:val="28"' not in xml_text:
        issues.append(f"{name}: номер страницы не 14 pt")
    if _header_has_bold_page_number(xml_text):
        issues.append(f"{name}: номер страницы жирный — должен быть TNR 14 обычный")
    return issues


def validate_page_numbering(docx_bytes: bytes) -> list[str]:
    issues: list[str] = []
    try:
        with zipfile.ZipFile(BytesIO(docx_bytes), "r") as zin:
            root = etree.fromstring(zin.read("word/document.xml"))
            sects = root.findall(".//w:sectPr", NS)
            if not sects:
                issues.append("Нет sectPr — нельзя проверить титул и номера страниц")
            else:
                if sects[0].find("w:titlePg", NS) is None:
                    issues.append("Нет titlePg — номер страницы может появиться на титуле")
                if len(sects) > 1:
                    typ = sects[0].find("w:type", NS)
                    val = typ.get(f"{{{W_NS}}}val") if typ is not None else ""
                    if val != "nextPage":
                        issues.append(
                            "Нет разрыва раздела nextPage после титула — глава 1 не на стр. 2"
                        )

            header_names = sorted(
                n for n in zin.namelist() if n.startswith("word/header") and n.endswith(".xml")
            )
            footer_names = sorted(
                n for n in zin.namelist() if n.startswith("word/footer") and n.endswith(".xml")
            )
            if "word/header1.xml" not in header_names:
                issues.append("Нет колонтитула для номера страницы (header1)")

            for name in header_names:
                data = zin.read(name)
                xml_text = data.decode("utf-8", errors="ignore")
                is_first = name.endswith("header3.xml")
                if is_first:
                    if _xml_has_page_field(data):
                        issues.append("header3 (титул) не должен содержать номер страницы")
                else:
                    if not _xml_has_page_field(data):
                        issues.append(f"{name}: нет поля PAGE")
                    if 'w:val="center"' not in xml_text and "<w:jc" not in xml_text:
                        issues.append(f"{name}: номер страницы не по центру")
                    issues.extend(_header_page_style_issues(name, xml_text))

            for name in footer_names:
                if _xml_has_page_field(zin.read(name)):
                    issues.append(f"{name}: номер страницы в футере запрещён")
    except (zipfile.BadZipFile, OSError, etree.XMLSyntaxError) as exc:
        issues.append(f"Не удалось проверить нумерацию страниц: {exc}")
    return issues


def _paragraph_sectpr(paragraph: Paragraph):
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    return p_pr.find(qn("w:sectPr"))


def _final_body_sectpr(doc: Document):
    body = doc.element.body
    if body is None:
        return None
    return body.find(qn("w:sectPr"))


def _set_sectpr_next_page_title(sect) -> None:
    typ = sect.find(qn("w:type"))
    if typ is None:
        typ = OxmlElement("w:type")
        sect.insert(0, typ)
    typ.set(qn("w:val"), "nextPage")
    if sect.find(qn("w:titlePg")) is None:
        sect.insert(0, OxmlElement("w:titlePg"))


def _sectpr_is_next_page(sect) -> bool:
    if sect is None:
        return False
    typ = sect.find(qn("w:type"))
    val = typ.get(qn("w:val")) if typ is not None else "nextPage"
    return val == "nextPage"


def _paragraph_index_of(doc: Document, paragraph: Paragraph) -> int | None:
    target = paragraph._p
    for i, para in enumerate(doc.paragraphs):
        if para._p is target:
            return i
    return None


def _title_stamp_table(doc: Document) -> Table | None:
    tables = [t for t in _iter_title_tables(doc) if _is_title_stamp_table(t)]
    return tables[0] if tables else None


def _strip_section_and_page_break(paragraph: Paragraph) -> int:
    """Снять nextPage/разрыв страницы, не трогая keepLines и интервал."""
    changed = 0
    sect = _paragraph_sectpr(paragraph)
    if sect is not None:
        p_pr = paragraph._p.find(qn("w:pPr"))
        if p_pr is not None:
            p_pr.remove(sect)
            changed += 1
    if paragraph_has_page_break_before(paragraph):
        set_page_break_before(paragraph, False)
        changed += 1
    for run in paragraph.runs:
        for br in list(run._element.findall(qn("w:br"))):
            if br.get(qn("w:type")) == "page":
                run._element.remove(br)
                changed += 1
    return changed


def _extract_city_year_out_of_title_tables(doc: Document) -> int:
    """«МИНСК YYYY» не держать в ячейке грифа — иначе разрыв раздела ломает таблицу."""
    if find_title_city_year_index(doc) is not None:
        return 0
    changed = 0
    for table in _iter_title_tables(doc):
        for row in table.rows:
            for cell in _row_unique_cells(row):
                for para in cell.paragraphs:
                    if not is_city_year_paragraph(para.text or ""):
                        continue
                    canon = canonical_title_city_year(para.text or "") or (
                        MULTI_SPACE_RE.sub(" ", (para.text or "").replace("\xa0", " ")).strip()
                    )
                    new_p = OxmlElement("w:p")
                    table._tbl.addnext(new_p)
                    out = Paragraph(new_p, doc)
                    set_paragraph_text(out, canon, bold=False)
                    set_paragraph_text(para, "", bold=False)
                    changed += 1
                    return changed
    return changed


def count_empty_between_stamp_and_city(doc: Document) -> int:
    """Сколько пустых абзацев сразу после таблицы грифа до «МИНСК YYYY»."""
    table = _title_stamp_table(doc)
    idx = find_title_city_year_index(doc)
    if table is None or idx is None:
        return 0
    city_el = doc.paragraphs[idx]._p
    count = 0
    el = table._tbl.getnext()
    while el is not None and el is not city_el:
        if el.tag != qn("w:p"):
            break
        para = Paragraph(el, doc)
        if is_paragraph_empty(para) and _paragraph_sectpr(para) is None:
            count += 1
            el = el.getnext()
            continue
        break
    return count


def _style_title_gap_empty(paragraph: Paragraph) -> None:
    set_single_line_spacing(paragraph)
    clear_first_line_indent(paragraph)


def _clear_breaks_between_stamp_table_and_city(doc: Document) -> int:
    """Между таблицей грифа и «МИНСК YYYY»: без nextPage; пустые строки оставить/поставить.

    Разрыв раздела после таблицы запрещён. Пустые абзацы перед «МИНСК» обязательны
    (TITLE_EMPTY_AFTER_STAMP_BEFORE_CITY).
    """
    table = _title_stamp_table(doc)
    idx = find_title_city_year_index(doc)
    if table is None or idx is None:
        return 0
    changed = 0
    for row in table.rows:
        for cell in _row_unique_cells(row):
            for para in cell.paragraphs:
                changed += _strip_section_and_page_break(para)
    city_el = doc.paragraphs[idx]._p
    city_after_table = False
    el = table._tbl.getnext()
    while el is not None:
        if el is city_el:
            city_after_table = True
            break
        nxt = el.getnext()
        if el.tag == qn("w:tbl"):
            break
        if el.tag == qn("w:p"):
            para = Paragraph(el, doc)
            if (
                not is_paragraph_empty(para)
                and _paragraph_sectpr(para) is None
                and not is_city_year_paragraph(para.text or "")
            ):
                break
            changed += _strip_section_and_page_break(para)
        el = nxt
    if not city_after_table:
        parent = city_el.getparent()
        if parent is not None:
            parent.remove(city_el)
            table._tbl.addnext(city_el)
            changed += 1
    want = TITLE_EMPTY_AFTER_STAMP_BEFORE_CITY
    have = count_empty_between_stamp_and_city(doc)
    tbl = table._tbl
    if have < want:
        for _ in range(want - have):
            new_p = OxmlElement("w:p")
            tbl.addnext(new_p)
            _style_title_gap_empty(Paragraph(new_p, doc))
            changed += 1
    elif have > want:
        extra = have - want
        prev = city_el.getprevious()
        while extra > 0 and prev is not None and prev is not tbl and prev.tag == qn("w:p"):
            para = Paragraph(prev, doc)
            if not is_paragraph_empty(para) or _paragraph_sectpr(para) is not None:
                break
            parent = prev.getparent()
            nxt_prev = prev.getprevious()
            if parent is None:
                break
            parent.remove(prev)
            changed += 1
            extra -= 1
            prev = nxt_prev
    return changed


def count_empty_before_stamp_table(doc: Document) -> int:
    """Сколько пустых абзацев сразу перед таблицей грифа (шапка → 6–8 пустых → гриф)."""
    table = _title_stamp_table(doc)
    if table is None:
        return 0
    count = 0
    prev = table._tbl.getprevious()
    while prev is not None and prev.tag == qn("w:p"):
        para = Paragraph(prev, doc)
        if is_paragraph_empty(para) and _paragraph_sectpr(para) is None:
            count += 1
            prev = prev.getprevious()
            continue
        break
    return count


def _is_empty_between_stamp_and_city(paragraph: Paragraph, doc: Document) -> bool:
    table = _title_stamp_table(doc)
    idx = find_title_city_year_index(doc)
    if table is None or idx is None:
        return False
    city_el = doc.paragraphs[idx]._p
    el = table._tbl.getnext()
    while el is not None and el is not city_el:
        if el is paragraph._p:
            return True
        el = el.getnext()
    return False


def _paragraph_follows_table(paragraph: Paragraph) -> bool:
    prev = paragraph._p.getprevious()
    return prev is not None and prev.tag == qn("w:tbl")


def _is_empty_in_stamp_gap(paragraph: Paragraph, doc: Document) -> bool:
    table = _title_stamp_table(doc)
    if table is None:
        return False
    prev = table._tbl.getprevious()
    while prev is not None and prev.tag == qn("w:p"):
        if prev is paragraph._p:
            return is_paragraph_empty(paragraph) and _paragraph_sectpr(paragraph) is None
        para = Paragraph(prev, doc)
        if not is_paragraph_empty(para) or _paragraph_sectpr(para) is not None:
            return False
        prev = prev.getprevious()
    return False


def _title_empty_run_limit(paragraph: Paragraph, doc: Document) -> int:
    if _is_empty_in_stamp_gap(paragraph, doc):
        return TITLE_EMPTY_BEFORE_STAMP_MAX
    if _is_empty_between_stamp_and_city(paragraph, doc):
        return TITLE_EMPTY_AFTER_STAMP_BEFORE_CITY
    return TITLE_EMPTY_BETWEEN_BLOCKS


def ensure_title_stamp_gap_after_header(doc: Document) -> int:
    """После шапки предприятия — 6–8 пустых строк, затем таблица «название + УТВЕРЖДАЮ»."""
    table = _title_stamp_table(doc)
    if table is None:
        return 0
    tbl = table._tbl
    count = count_empty_before_stamp_table(doc)
    changed = 0
    if count < TITLE_EMPTY_BEFORE_STAMP_MIN:
        need = TITLE_EMPTY_BEFORE_STAMP - count
        for _ in range(need):
            new_p = OxmlElement("w:p")
            tbl.addprevious(new_p)
            para = Paragraph(new_p, doc)
            set_single_line_spacing(para)
            clear_first_line_indent(para)
            changed += 1
    elif count > TITLE_EMPTY_BEFORE_STAMP_MAX:
        extra = count - TITLE_EMPTY_BEFORE_STAMP_MAX
        prev = tbl.getprevious()
        while extra > 0 and prev is not None and prev.tag == qn("w:p"):
            para = Paragraph(prev, doc)
            if not is_paragraph_empty(para) or _paragraph_sectpr(para) is not None:
                break
            parent = prev.getparent()
            nxt_prev = prev.getprevious()
            if parent is None:
                break
            parent.remove(prev)
            changed += 1
            extra -= 1
            prev = nxt_prev
    return changed


def validate_title_stamp_gap(doc: Document) -> list[str]:
    """Перед таблицей грифа — 6–8 пустых строк после шапки предприятия."""
    if _title_stamp_table(doc) is None:
        return []
    count = count_empty_before_stamp_table(doc)
    if TITLE_EMPTY_BEFORE_STAMP_MIN <= count <= TITLE_EMPTY_BEFORE_STAMP_MAX:
        return []
    return [
        "После шапки предприятия перед таблицей грифа («УТВЕРЖДАЮ») "
        f"должно быть {TITLE_EMPTY_BEFORE_STAMP_MIN}–{TITLE_EMPTY_BEFORE_STAMP_MAX} "
        f"пустых строк, сейчас {count}"
    ]


def compact_title_empty_paragraphs(doc: Document) -> int:
    """Пустые на титуле: 6–8 перед грифом; иначе не больше одной между блоками."""
    body_start = find_body_start_index(doc)
    if body_start <= 0:
        return 0
    removed = 0
    idx = 0
    empty_run = 0
    while idx < body_start:
        paragraph = doc.paragraphs[idx]
        if paragraph_is_inside_table(paragraph):
            empty_run = 0
            idx += 1
            continue
        if is_paragraph_empty(paragraph):
            if _paragraph_follows_table(paragraph):
                empty_run = 0
            empty_run += 1
            limit = _title_empty_run_limit(paragraph, doc)
            if empty_run > limit:
                if _paragraph_sectpr(paragraph) is not None:
                    empty_run = 0
                    idx += 1
                    continue
                el = paragraph._element
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
                    removed += 1
                    body_start -= 1
                    continue
            idx += 1
            continue
        empty_run = 0
        idx += 1
    return removed


def _style_invisible_section_break_holder(paragraph: Paragraph) -> None:
    """Абзац-носитель nextPage: высота 1 pt, без vanish.

    vanish прячет абзац, пока выключено «¶». При «показать все знаки» Word
    показывает скрытый текст — строка вдруг появляется, макет прыгает.
    """
    set_single_line_spacing(paragraph)
    clear_first_line_indent(paragraph)
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "20")
    spacing.set(qn("w:lineRule"), "exact")
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)
    vanish = r_pr.find(qn("w:vanish"))
    if vanish is not None:
        r_pr.remove(vanish)
    for tag, val in (("w:sz", "2"), ("w:szCs", "2")):
        el = r_pr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            r_pr.append(el)
        el.set(qn("w:val"), val)


def ensure_title_page_separated(doc: Document) -> int:
    """
    Титул (шапка / УТВЕРЖДАЮ / «Минск 20xx») — отдельная страница:
    nextPage + titlePg перед главой 1, лишние пустые строки титула убрать.
    """
    changed = compact_title_empty_paragraphs(doc)
    body_start = find_body_start_index(doc)
    if body_start <= 0:
        return changed
    anchor_idx: int | None = None
    for idx in range(body_start - 1, -1, -1):
        paragraph = doc.paragraphs[idx]
        if paragraph_is_inside_table(paragraph):
            continue
        if paragraph_text_normalized(paragraph):
            anchor_idx = idx
            break
    if anchor_idx is None:
        anchor_idx = body_start - 1
    anchor = doc.paragraphs[anchor_idx]
    chapter = doc.paragraphs[body_start]
    if paragraph_has_page_break_before(chapter):
        set_page_break_before(chapter, False)
        changed += 1
    if find_title_city_year_index(doc) is not None:
        # nextPage только после «МИНСК YYYY»; между таблицей грифа и городом разрыва нет.
        changed += _extract_city_year_out_of_title_tables(doc)
        changed += _clear_breaks_between_stamp_table_and_city(doc)
        changed += ensure_section_break_after_city_year(doc)
        changed += _clear_breaks_between_stamp_table_and_city(doc)
        return changed
    if is_city_year_paragraph(anchor.text or ""):
        holder = insert_paragraph_after(anchor)
        _style_invisible_section_break_holder(holder)
        old_sect = _paragraph_sectpr(anchor)
        if old_sect is not None:
            p_pr = anchor._p.find(qn("w:pPr"))
            if p_pr is not None:
                p_pr.remove(old_sect)
            holder._p.get_or_add_pPr().append(old_sect)
        anchor = holder
        changed += 1
    sect = _paragraph_sectpr(anchor)
    if sect is None:
        final = _final_body_sectpr(doc)
        if final is None:
            return changed
        p_pr = anchor._p.get_or_add_pPr()
        sect = deepcopy(final)
        p_pr.append(sect)
        changed += 1
    _set_sectpr_next_page_title(sect)
    final = _final_body_sectpr(doc)
    if final is not None and final is not sect:
        title_pg = final.find(qn("w:titlePg"))
        if title_pg is not None:
            final.remove(title_pg)
            changed += 1
    return changed


def validate_title_page_separated(doc: Document) -> list[str]:
    """Глава 1 должна начинаться после nextPage-секции титула."""
    body_start = find_body_start_index(doc)
    if body_start <= 0:
        return []
    issues: list[str] = []
    empty_run = 0
    found_next_page = False
    for idx in range(body_start):
        paragraph = doc.paragraphs[idx]
        if paragraph_is_inside_table(paragraph):
            empty_run = 0
            continue
        if is_paragraph_empty(paragraph):
            if _paragraph_follows_table(paragraph):
                empty_run = 0
            empty_run += 1
            limit = _title_empty_run_limit(paragraph, doc)
            if empty_run > limit:
                issues.append(
                    "На титуле слишком много пустых строк — титул не умещается на одной странице"
                )
        else:
            empty_run = 0
        sect = _paragraph_sectpr(paragraph)
        if sect is None:
            continue
        typ = sect.find(qn("w:type"))
        val = typ.get(qn("w:val")) if typ is not None else "nextPage"
        if val == "nextPage":
            found_next_page = True
    if not found_next_page:
        issues.append("Титул не отделён: нет разрыва раздела nextPage перед главой 1")
    issues.extend(validate_section_break_after_city_year(doc))
    return issues


def is_city_year_paragraph(text: str) -> bool:
    """«Минск-2026», «Минск, 2026», «Минск 2026» — строка города и года на титуле."""
    t = MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ").strip())
    return bool(CITY_YEAR_RE.match(t))


def canonical_title_city_year(text: str) -> str | None:
    """Канон эталона: «МИНСК 2026» (прописные, пробел). Год берётся из исходника."""
    t = MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ").strip())
    if not CITY_YEAR_RE.match(t):
        return None
    year = re.search(r"20\d{2}", t)
    if year is None:
        return None
    return TITLE_CITY_YEAR_TEMPLATE.format(year=year.group(0))


def find_title_city_year_index(doc: Document) -> int | None:
    body_start = find_body_start_index(doc)
    found: int | None = None
    limit = body_start if body_start > 0 else len(doc.paragraphs)
    for i in range(limit):
        para = doc.paragraphs[i]
        if paragraph_is_inside_table(para):
            continue
        if is_city_year_paragraph(para.text):
            found = i
    return found


def _set_bottom_page_frame(paragraph: Paragraph) -> bool:
    """Прижать абзац к низу страницы (framePr), не смешивая с главой 1."""
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:framePr"))
    wanted = {qn(attr): val for attr, val in TITLE_CITY_YEAR_FRAME_ATTRS}
    if existing is not None and all(existing.get(k) == v for k, v in wanted.items()):
        return False
    if existing is not None:
        p_pr.remove(existing)
    fr = OxmlElement("w:framePr")
    for key, val in wanted.items():
        fr.set(key, val)
    p_pr.insert(0, fr)
    return True


def place_title_city_year_at_bottom(doc: Document) -> int:
    """
    «Минск-2026» — последний абзац титула, по центру, без отступа 1,25,
    прижат к низу 1-й страницы. Не дублировать на стр. 2.
    Между таблицей грифа и этой строкой разрыва раздела/страницы нет;
    пустые строки перед «МИНСК» обязательны.
    """
    body_start = find_body_start_index(doc)
    if body_start <= 0:
        return 0
    changed = _extract_city_year_out_of_title_tables(doc)
    for i in range(len(doc.paragraphs) - 1, body_start - 1, -1):
        para = doc.paragraphs[i]
        if not is_city_year_paragraph(para.text):
            continue
        if _paragraph_sectpr(para) is not None:
            continue
        el = para._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            changed += 1
    body_start = find_body_start_index(doc)
    idx = find_title_city_year_index(doc)
    if idx is None or body_start <= 0:
        return changed
    para = doc.paragraphs[idx]
    chapter = doc.paragraphs[body_start]
    nxt = para._p.getnext()
    holder_follows = (
        nxt is not None
        and nxt.tag == qn("w:p")
        and _sectpr_is_next_page(_paragraph_sectpr(Paragraph(nxt, doc)))
    )
    # Не переносить «МИНСК» перед главой, перепрыгивая holder: иначе nextPage
    # остаётся сразу после таблицы грифа.
    if (
        not holder_follows
        and _title_stamp_table(doc) is None
        and para._p is not chapter._p
        and nxt is not chapter._p
    ):
        el = para._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            chapter._p.addprevious(el)
            changed += 1
        idx = find_title_city_year_index(doc)
        if idx is None:
            return changed
        para = doc.paragraphs[idx]
    canon = canonical_title_city_year(para.text or "")
    current = MULTI_SPACE_RE.sub(" ", (para.text or "").replace("\xa0", " ").strip())
    if canon and canon != current:
        set_paragraph_text(para, canon, bold=False)
        changed += 1
    if TITLE_CITY_YEAR_ALIGN == "center":
        ensure_paragraph_centered(para)
    clear_first_line_indent(para)
    set_single_line_spacing(para)
    if _set_bottom_page_frame(para):
        changed += 1
    changed += _clear_breaks_between_stamp_table_and_city(doc)
    changed += ensure_section_break_after_city_year(doc)
    changed += remove_city_year_from_headers_footers(doc)
    return changed


def validate_title_city_year(doc: Document) -> list[str]:
    """Если на титуле есть город-год — он внизу 1-й страницы, по центру, без отступа."""
    issues: list[str] = []
    body_start = find_body_start_index(doc)
    idx = find_title_city_year_index(doc)
    if idx is None:
        return issues
    para = doc.paragraphs[idx]
    if not is_paragraph_centered(para):
        issues.append("«Минск-2026» на титуле не по центру")
    if first_line_indent_cm(para) > FIRST_LINE_INDENT_TOLERANCE_CM:
        issues.append("У «Минск-2026» на титуле не должно быть отступа 1,25 см")
    p_pr = para._p.find(qn("w:pPr"))
    has_frame = p_pr is not None and p_pr.find(qn("w:framePr")) is not None
    if not has_frame:
        issues.append("«Минск-2026» не прижат к низу 1-й страницы")
    if has_frame:
        v_anchor = p_pr.find(qn("w:framePr")).get(qn("w:vAnchor"))
        if v_anchor == "page":
            issues.append("«МИНСК 2026» привязан к краю страницы (попадает в колонтитул)")
    if _paragraph_sectpr(para) is not None:
        issues.append("«МИНСК 2026» совмещён с разрывом раздела — попадает в колонтитул")
    last_title = idx
    for i in range(idx + 1, body_start if body_start > 0 else len(doc.paragraphs)):
        if paragraph_is_inside_table(doc.paragraphs[i]):
            continue
        if paragraph_text_normalized(doc.paragraphs[i]):
            last_title = i
            break
    else:
        last_title = idx
    if last_title != idx and body_start > 0:
        extra = paragraph_text_normalized(doc.paragraphs[last_title])[:40]
        if extra and not is_city_year_paragraph(extra):
            issues.append(
                "«Минск-2026» не последняя строка титула — стоит в середине 1-й страницы"
            )
    if body_start > 0:
        for i in range(body_start, len(doc.paragraphs)):
            if is_city_year_paragraph(doc.paragraphs[i].text):
                issues.append("«Минск-2026» повторён после титула (стр. 2)")
                break
    issues.extend(validate_city_year_not_in_headers_footers(doc))
    issues.extend(validate_stamp_table_city_gap(doc))
    return issues


def validate_stamp_table_city_gap(doc: Document) -> list[str]:
    """Между таблицей грифа и «МИНСК YYYY»: пустые строки есть, разрыва нет."""
    table = _title_stamp_table(doc)
    idx = find_title_city_year_index(doc)
    if table is None or idx is None:
        return []
    issues: list[str] = []
    city_el = doc.paragraphs[idx]._p
    el = table._tbl.getnext()
    while el is not None and el is not city_el:
        if el.tag == qn("w:p"):
            para = Paragraph(el, doc)
            if _sectpr_is_next_page(_paragraph_sectpr(para)) or paragraph_has_page_break_before(
                para
            ):
                issues.append(
                    "Между таблицей грифа и «МИНСК 2026» не должно быть разрыва раздела/страницы"
                )
                break
        el = el.getnext()
    empty_n = count_empty_between_stamp_and_city(doc)
    want = TITLE_EMPTY_AFTER_STAMP_BEFORE_CITY
    if empty_n < 1:
        issues.append(
            "После таблицы грифа перед «МИНСК 2026» должны быть пустые строки"
        )
    elif empty_n > want:
        issues.append(
            "После таблицы грифа перед «МИНСК 2026» слишком много пустых строк "
            f"(нужно {want}, сейчас {empty_n})"
        )
    return issues


def _title_has_utverzhdayu(text: str) -> bool:
    return "УТВЕРЖДАЮ" in (text or "").upper()


def _compact_title_two_column(text: str) -> str:
    """Сжать пробелы; колонки через Tab: «ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ\\tУТВЕРЖДАЮ»."""
    raw = (text or "").replace("\xa0", " ")
    if "\t" in raw:
        left, right = raw.split("\t", 1)
        left = MULTI_SPACE_RE.sub(" ", left).strip()
        right = MULTI_SPACE_RE.sub(" ", right).strip()
        if left and right:
            return f"{left}\t{right}"
        return MULTI_SPACE_RE.sub(" ", raw).strip()
    stripped = raw.strip()
    match = re.match(r"^(.*\S) {3,}(\S.*)$", stripped)
    if match:
        right = MULTI_SPACE_RE.sub(" ", match.group(2)).strip()
        left = MULTI_SPACE_RE.sub(" ", match.group(1)).strip()
        right_u = right.upper()
        if (
            "УТВЕРЖД" in right_u
            or SIGNATORY_NAME_TAIL.search(right)
            or SIGNATORY_DATE_PLAQUE_RE.search(right)
        ):
            return f"{left}\t{right}"
    return MULTI_SPACE_RE.sub(" ", stripped)


def _is_title_instruction_name_line(text: str) -> bool:
    """Название должности на титуле (левая колонка у «УТВЕРЖДАЮ»), не гриф и не ФИО."""
    t = (text or "").replace("\xa0", " ").strip()
    if not t or is_city_year_paragraph(t) or is_chapter_header(t):
        return False
    if _title_has_utverzhdayu(t):
        return False
    if SIGNATORY_DATE_PLAQUE_RE.search(t) or t.startswith("____"):
        return False
    if SIGNATORY_NAME_TAIL.search(t) and ("_" in t or "\t" in t):
        return False
    low = t.casefold()
    keys = (
        "должностн",
        "инструкц",
        "заместител",
        "начальника отдела",
        "материально-техническ",
        "омтс",
    )
    return any(key in low for key in keys)


def _is_title_approve_stamp_line(text: str) -> bool:
    t = (text or "").replace("\xa0", " ").strip()
    if not t or is_city_year_paragraph(t) or is_chapter_header(t):
        return False
    if _title_has_utverzhdayu(t):
        return True
    if SIGNATORY_DATE_PLAQUE_RE.search(t):
        return True
    if re.search(r"[«\"„].{0,12}20[_0-9]{2,4}", t):
        return True
    if SIGNATORY_NAME_TAIL.search(t) and ("_" in t or "\t" in t):
        return True
    low = t.lstrip(" _\t").lower()
    if low.startswith(("начальник", "директор", "главный инженер", "главный", "заместитель", "руководитель")):
        return True
    return False


def _is_title_instruction_number_line(text: str) -> bool:
    t = (text or "").replace("\xa0", " ").strip().casefold()
    return "номер инструкц" in t


def _is_title_number_underline_line(text: str) -> bool:
    t = (text or "").replace("\xa0", " ").strip()
    if not t or "_" not in t:
        return False
    return not t.replace("_", "").replace(" ", "").replace("\t", "")


def _paragraph_has_keep_lines(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:keepLines")) is not None


def _paragraph_font_half_points(paragraph: Paragraph) -> str:
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        _name, _east, sz, _bold = _run_font_fields(run)
        if sz:
            return sz
    return ""


def _row_unique_cells(row) -> list:
    seen: list = []
    ids: set[int] = set()
    for cell in row.cells:
        tid = id(cell._tc)
        if tid in ids:
            continue
        ids.add(tid)
        seen.append(cell)
    return seen


def _unique_keep_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for raw in items:
        key = MULTI_SPACE_RE.sub(" ", (raw or "").replace("\xa0", " ")).strip()
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(key)
    return out


def _is_title_document_name_line(text: str) -> bool:
    """Тип документа для левой колонки грифа (не должность утверждающего)."""
    t = MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ")).strip()
    if not t or is_city_year_paragraph(t) or is_chapter_header(t):
        return False
    if _is_title_instruction_number_line(t):
        return True
    folded = t.casefold()
    if "должностн" in folded and "инструкц" in folded:
        return True
    if "рабочая инструкция" in folded or "рабочей инструкции" in folded:
        return True
    return folded.startswith("положение")


def _is_title_approver_job_line(text: str) -> bool:
    """Должность утверждающего — только правая колонка."""
    t = MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ")).strip()
    if not t or _is_title_document_name_line(t) or _title_has_utverzhdayu(t):
        return False
    if is_signatory_date_plaque(t) or is_city_year_paragraph(t):
        return False
    low = t.lstrip(" _\t").lower()
    return low.startswith(
        (
            "начальник",
            "директор",
            "главный",
            "заместитель",
            "руководитель",
        )
    )


def _stamp_piece_role(text: str, *, default: str) -> str:
    """left = название документа; right = гриф УТВЕРЖДАЮ; skip = служебное.

    Спорные строки (ОМТС, «ведущий инженер») не перекидывать в другой столбец:
    берём default — индекс колонки или положение относительно «УТВЕРЖДАЮ».
    """
    t = MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ")).strip()
    if not t or is_city_year_paragraph(t) or is_chapter_header(t):
        return "skip"
    if _is_title_number_underline_line(t):
        return "skip"
    if _is_title_instruction_number_line(t) or _is_title_document_name_line(t):
        return "left"
    if _title_has_utverzhdayu(t) or is_signatory_date_plaque(t):
        return "right"
    iof_match = TITLE_IOF_RE.search(t)
    if iof_match:
        iof = iof_match.group(1).strip()
        compact = t.replace(" ", "")
        if "_" in t or "\t" in t or compact == iof.replace(" ", ""):
            return "right"
    if _is_title_approver_job_line(t):
        return "right"
    return default


def _append_stamp_piece(bucket_left: list[str], bucket_right: list[str], text: str, default: str) -> None:
    role = _stamp_piece_role(text, default=default)
    piece = MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ")).strip()
    if role == "left" and piece:
        bucket_left.append(piece)
    elif role == "right" and piece:
        bucket_right.append(piece)


def _is_loose_title_stamp_text(text: str) -> bool:
    t = (text or "").replace("\xa0", " ").strip()
    if not t or is_city_year_paragraph(t) or is_chapter_header(t):
        return False
    if _stamp_piece_role(t, default="skip") in ("left", "right"):
        return True
    if _is_title_instruction_number_line(t) or _is_title_number_underline_line(t):
        return True
    if _is_title_instruction_name_line(t) or _is_title_approve_stamp_line(t):
        return True
    compact = _compact_title_two_column(t)
    return "\t" in compact and (
        _title_has_utverzhdayu(compact) or _is_title_instruction_name_line(compact.split("\t", 1)[0])
    )


def _split_title_two_column_parts(text: str) -> tuple[str, str]:
    compact = _compact_title_two_column(text or "")
    if "\t" in compact:
        left, right = compact.split("\t", 1)
        return left.strip(), right.strip()
    return compact.strip(), ""


def _collect_stamp_column_texts(table: Table) -> tuple[list[str], list[str]]:
    left: list[str] = []
    right: list[str] = []
    for row in table.rows:
        cells = _row_unique_cells(row)
        if not cells:
            continue
        n = len(cells)
        for ci, cell in enumerate(cells):
            if n == 1:
                default = "left"
            elif ci == 0:
                default = "left"
            elif ci == n - 1:
                default = "right"
            else:
                default = "skip"
            for para in cell.paragraphs:
                piece_l, piece_r = _split_title_two_column_parts(para.text or "")
                if piece_r:
                    _append_stamp_piece(left, right, piece_l, "left")
                    _append_stamp_piece(left, right, piece_r, "right")
                elif piece_l:
                    _append_stamp_piece(left, right, piece_l, default)
    return _unique_keep_order(left), _unique_keep_order(right)


def _collect_loose_title_stamp_texts(doc: Document) -> tuple[list[str], list[str]]:
    left: list[str] = []
    right: list[str] = []
    body_start = find_body_start_index(doc)
    limit = body_start if body_start > 0 else len(doc.paragraphs)
    seen_utv = False
    for i in range(limit):
        para = doc.paragraphs[i]
        if paragraph_is_inside_table(para):
            continue
        raw = para.text or ""
        if _title_has_utverzhdayu(raw):
            seen_utv = True
        if not _is_loose_title_stamp_text(raw):
            continue
        piece_l, piece_r = _split_title_two_column_parts(raw)
        default = "right" if seen_utv else "left"
        if piece_r:
            _append_stamp_piece(left, right, piece_l, "left")
            _append_stamp_piece(left, right, piece_r, "right")
            continue
        if piece_l:
            _append_stamp_piece(left, right, piece_l, default)
    return _unique_keep_order(left), _unique_keep_order(right)


def _remove_loose_title_stamp_paragraphs(doc: Document) -> int:
    body_start = find_body_start_index(doc)
    limit = body_start if body_start > 0 else len(doc.paragraphs)
    removed = 0
    for i in range(limit - 1, -1, -1):
        para = doc.paragraphs[i]
        if paragraph_is_inside_table(para):
            continue
        if not _is_loose_title_stamp_text(para.text or ""):
            continue
        el = para._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            removed += 1
    return removed


def _title_date_year(doc: Document, texts: list[str]) -> str:
    for raw in texts:
        found = re.search(r"20\d{2}", raw or "")
        if found:
            return found.group(0)
    idx = find_title_city_year_index(doc)
    if idx is not None:
        found = re.search(r"20\d{2}", doc.paragraphs[idx].text or "")
        if found:
            return found.group(0)
    return "2026"


def _clear_cell_paragraphs_keep_one(cell) -> Paragraph:
    tc = cell._tc
    paras = list(tc.findall(qn("w:p")))
    for extra in paras[1:]:
        tc.remove(extra)
    if not paras:
        new_p = OxmlElement("w:p")
        tc.append(new_p)
        return Paragraph(new_p, cell)
    first = Paragraph(paras[0], cell)
    for run in list(first.runs):
        run._element.getparent().remove(run._element)
    return first


def _add_cell_paragraph(cell) -> Paragraph:
    new_p = OxmlElement("w:p")
    cell._tc.append(new_p)
    return Paragraph(new_p, cell)


def _format_title_cell_paragraph(
    para: Paragraph,
    text: str,
    *,
    size_pt: int | None = None,
    underline: bool = False,
    keep_lines: bool = False,
    right_in_cell: bool = False,
) -> None:
    set_paragraph_text(para, text, bold=False, size_pt=size_pt, underline=underline)
    clear_first_line_indent(para)
    set_single_line_spacing(para)
    if right_in_cell:
        ensure_paragraph_right(para)
    elif TITLE_STAMP_IN_TABLE_ALIGN == "left":
        ensure_paragraph_left(para)
    set_paragraph_keep_lines(para, keep_lines)


def _write_title_left_cell(cell, left_texts: list[str]) -> None:
    names = [
        t
        for t in left_texts
        if t
        and _stamp_piece_role(t, default="left") == "left"
        and not _is_title_instruction_number_line(t)
        and not _is_title_number_underline_line(t)
        and not _title_has_utverzhdayu(t)
        and not is_signatory_date_plaque(t)
        and not is_city_year_paragraph(t)
    ]
    first = _clear_cell_paragraphs_keep_one(cell)
    lines = list(names) if names else []
    paras: list[Paragraph] = []
    if lines:
        _format_title_cell_paragraph(first, lines[0], size_pt=14)
        paras.append(first)
        for extra in lines[1:]:
            para = _add_cell_paragraph(cell)
            _format_title_cell_paragraph(para, extra, size_pt=14)
            paras.append(para)
        line_para = _add_cell_paragraph(cell)
    else:
        line_para = first
    _format_title_cell_paragraph(
        line_para, TITLE_UNDERLINE_LINE, size_pt=14, underline=True
    )
    num_para = _add_cell_paragraph(cell)
    _format_title_cell_paragraph(
        num_para, TITLE_NUMBER_LABEL, size_pt=TITLE_NUMBER_FONT_PT
    )


def _write_title_right_cell(cell, right_texts: list[str], year: str) -> None:
    positions: list[str] = []
    iof = ""
    date_raw = ""
    for raw in right_texts:
        t = MULTI_SPACE_RE.sub(" ", (raw or "").replace("\xa0", " ")).strip()
        if not t or t.upper() == "УТВЕРЖДАЮ":
            continue
        if _stamp_piece_role(t, default="right") == "left":
            continue
        if is_signatory_date_plaque(t) or SIGNATORY_DATE_PLAQUE_RE.search(t):
            date_raw = t
            continue
        iof_match = TITLE_IOF_RE.search(t)
        if iof_match and ("_" in t or "\t" in t or t.strip() == iof_match.group(1)):
            iof = iof_match.group(1).strip()
            rest = (t[: iof_match.start()] + t[iof_match.end() :]).strip(" \t_")
            if rest and not is_signatory_date_plaque(rest) and rest.upper() != "УТВЕРЖДАЮ":
                if not _is_title_number_underline_line(rest):
                    positions.append(rest)
            continue
        if _is_title_number_underline_line(t):
            continue
        positions.append(t)
    first = _clear_cell_paragraphs_keep_one(cell)
    _format_title_cell_paragraph(first, "УТВЕРЖДАЮ", size_pt=14)
    _apply_title_utverzhdayu_weight(first)
    for pos in _unique_keep_order(positions):
        para = _add_cell_paragraph(cell)
        _format_title_cell_paragraph(para, pos, size_pt=14)
    sign = _add_cell_paragraph(cell)
    for run in list(sign.runs):
        run._element.getparent().remove(run._element)
    line_run = sign.add_run(TITLE_UNDERLINE_LINE)
    apply_run_font(line_run, bold=False, size_pt=14)
    _set_run_underline(line_run, True)
    if iof:
        name_run = sign.add_run(" " + iof)
        apply_run_font(name_run, bold=False, size_pt=14)
    _strip_paragraph_tabs(sign)
    clear_first_line_indent(sign)
    set_single_line_spacing(sign)
    ensure_paragraph_left(sign)
    date_para = _add_cell_paragraph(cell)
    compact = compact_title_stamp_date_plaque(date_raw, year)
    _format_title_cell_paragraph(date_para, compact, size_pt=14, keep_lines=True)
    _strip_paragraph_tabs(date_para)
    set_paragraph_keep_lines(date_para, True)


def _insert_two_col_table_after(paragraph: Paragraph) -> Table:
    doc = paragraph.part.document
    table = doc.add_table(rows=1, cols=2)
    paragraph._p.addnext(table._tbl)
    return table


def _replace_stamp_table_with_two_col(old_table: Table) -> Table:
    doc = old_table.part.document
    new_table = doc.add_table(rows=1, cols=2)
    old_table._tbl.addnext(new_table._tbl)
    parent = old_table._tbl.getparent()
    if parent is not None:
        parent.remove(old_table._tbl)
    return new_table


def _title_stamp_insert_anchor(doc: Document) -> Paragraph | None:
    body_start = find_body_start_index(doc)
    if body_start <= 0:
        return None
    last_org: Paragraph | None = None
    for i in range(body_start):
        para = doc.paragraphs[i]
        if paragraph_is_inside_table(para):
            continue
        text = para.text or ""
        if is_city_year_paragraph(text) or _is_loose_title_stamp_text(text):
            continue
        if text.strip():
            last_org = para
    return last_org if last_org is not None else doc.paragraphs[0]


def _fill_two_col_stamp_table(table: Table, left_texts: list[str], right_texts: list[str], year: str) -> None:
    _set_title_stamp_table_no_borders(table)
    _apply_title_stamp_table_geometry(table)
    left_cell, right_cell = table.rows[0].cells[0], table.rows[0].cells[1]
    _write_title_left_cell(left_cell, left_texts)
    _write_title_right_cell(right_cell, right_texts, year)


def _ensure_title_left_right_stamp(doc: Document) -> int:
    """
    Титул: таблица 2 колонки без рамок.
    Слева — название, подчёркнутая линия, «номер инструкции» 12 pt.
    Справа — УТВЕРЖДАЮ, должность, линия+И.О.Ф., плашка даты keepLines.
    Если в исходнике название и «УТВЕРЖДАЮ» идут абзацами без таблицы —
    собрать их в такую таблицу. Уже корректную таблицу не пересобирать.
    """
    stamp_tables = [t for t in _iter_title_tables(doc) if _is_title_stamp_table(t)]
    left_texts: list[str] = []
    right_texts: list[str] = []
    for table in stamp_tables:
        col_l, col_r = _collect_stamp_column_texts(table)
        left_texts.extend(col_l)
        right_texts.extend(col_r)
    loose_l, loose_r = _collect_loose_title_stamp_texts(doc)
    left_texts = _unique_keep_order(left_texts + loose_l)
    right_texts = _unique_keep_order(right_texts + loose_r)
    if not left_texts and not right_texts and not stamp_tables:
        return 0
    year = _title_date_year(doc, left_texts + right_texts)
    changed = 0
    if stamp_tables:
        target = stamp_tables[0]
        cells = _row_unique_cells(target.rows[0]) if target.rows else []
        need_replace = len(target.rows) != 1 or len(cells) != 2
        if (
            not need_replace
            and len(stamp_tables) == 1
            and not loose_l
            and not loose_r
            and not validate_title_stamp_left_right(doc)
        ):
            return 0
        if need_replace:
            target = _replace_stamp_table_with_two_col(target)
            changed += 1
        _fill_two_col_stamp_table(target, left_texts, right_texts, year)
        changed += 1
        for extra in stamp_tables[1:]:
            parent = extra._tbl.getparent()
            if parent is not None:
                parent.remove(extra._tbl)
                changed += 1
    else:
        anchor = _title_stamp_insert_anchor(doc)
        if anchor is None:
            return 0
        table = _insert_two_col_table_after(anchor)
        _fill_two_col_stamp_table(table, left_texts, right_texts, year)
        changed += 1
    changed += _remove_loose_title_stamp_paragraphs(doc)
    return changed


def apply_title_instruction_number_font(doc: Document) -> int:
    """«номер инструкции» — TNR 12 pt, слева, не 14."""
    changed = 0
    targets: list[Paragraph] = []
    body_start = find_body_start_index(doc)
    limit = body_start if body_start > 0 else min(len(doc.paragraphs), 45)
    for i in range(limit):
        para = doc.paragraphs[i]
        if _is_title_instruction_number_line(para.text or ""):
            targets.append(para)
    for table in _iter_title_tables(doc):
        for row in table.rows:
            for cell in _row_unique_cells(row):
                for para in cell.paragraphs:
                    if _is_title_instruction_number_line(para.text or ""):
                        targets.append(para)
    for para in targets:
        before = _paragraph_font_half_points(para)
        _format_title_cell_paragraph(para, TITLE_NUMBER_LABEL, size_pt=TITLE_NUMBER_FONT_PT)
        if before != TITLE_NUMBER_FONT_HALF:
            changed += 1
        else:
            changed += 1
    return changed


def _iter_all_headers_footers(doc: Document):
    for section in doc.sections:
        parts = [
            section.header,
            section.footer,
            getattr(section, "even_page_header", None),
            getattr(section, "even_page_footer", None),
        ]
        if section.different_first_page_header_footer:
            parts.extend([section.first_page_header, section.first_page_footer])
        for hf in parts:
            if hf is not None:
                yield hf


def _move_sectpr_off_city_year(doc: Document) -> int:
    """Разрыв раздела nextPage не на «МИНСК YYYY» — иначе Word кладёт строку в колонтитул."""
    idx = find_title_city_year_index(doc)
    if idx is None:
        return 0
    para = doc.paragraphs[idx]
    sect = _paragraph_sectpr(para)
    if sect is None:
        return 0
    holder = insert_paragraph_after(para)
    p_pr = para._p.find(qn("w:pPr"))
    if p_pr is not None:
        p_pr.remove(sect)
    holder._p.get_or_add_pPr().append(sect)
    _style_invisible_section_break_holder(holder)
    return 1


def _first_title_next_page_para(doc: Document, body_start: int, skip_el) -> Paragraph | None:
    for i in range(max(body_start, 0)):
        para = doc.paragraphs[i]
        if skip_el is not None and para._p is skip_el:
            continue
        if _sectpr_is_next_page(_paragraph_sectpr(para)):
            return para
    return None


def ensure_section_break_after_city_year(doc: Document) -> int:
    """
    Сразу после «МИНСК YYYY» — пустой абзац с разрывом раздела nextPage.
    Глава 1 — первый абзац новой страницы (с самого верха стр. 2).
    """
    body_start = find_body_start_index(doc)
    if body_start <= 0:
        return 0
    idx = find_title_city_year_index(doc)
    if idx is None:
        return 0
    changed = _move_sectpr_off_city_year(doc)
    idx = find_title_city_year_index(doc)
    body_start = find_body_start_index(doc)
    if idx is None or body_start <= 0:
        return changed
    city = doc.paragraphs[idx]
    chapter = doc.paragraphs[body_start]
    city_parent = city._p.getparent()
    chapter_parent = chapter._p.getparent()
    if city_parent is None or chapter_parent is None or city_parent is not chapter_parent:
        return changed

    nxt = city._p.getnext()
    holder: Paragraph | None = None
    if nxt is not None and nxt.tag == qn("w:p"):
        cand = Paragraph(nxt, doc)
        if is_paragraph_empty(cand) and _sectpr_is_next_page(_paragraph_sectpr(cand)):
            holder = cand
    if holder is None:
        holder = insert_paragraph_after(city)
        changed += 1
    _style_invisible_section_break_holder(holder)

    holder_sect = _paragraph_sectpr(holder)
    if holder_sect is None:
        donor_para = _first_title_next_page_para(doc, body_start, holder._p)
        if donor_para is not None:
            donor = _paragraph_sectpr(donor_para)
            d_pr = donor_para._p.find(qn("w:pPr"))
            if d_pr is not None and donor is not None:
                d_pr.remove(donor)
                holder._p.get_or_add_pPr().append(donor)
                holder_sect = donor
                changed += 1
        if holder_sect is None:
            final = _final_body_sectpr(doc)
            if final is None:
                return changed
            holder_sect = deepcopy(final)
            holder._p.get_or_add_pPr().append(holder_sect)
            changed += 1
    _set_sectpr_next_page_title(holder_sect)

    body_start = find_body_start_index(doc)
    for i in range(body_start):
        para = doc.paragraphs[i]
        if para._p is holder._p:
            continue
        sect = _paragraph_sectpr(para)
        if not _sectpr_is_next_page(sect):
            continue
        p_pr = para._p.find(qn("w:pPr"))
        if p_pr is not None and sect is not None:
            p_pr.remove(sect)
            changed += 1

    if city._p.getnext() is not holder._p:
        holder_el = holder._p
        parent = holder_el.getparent()
        if parent is not None:
            parent.remove(holder_el)
            city._p.addnext(holder_el)
            holder = Paragraph(holder_el, doc)
            changed += 1

    chapter = doc.paragraphs[find_body_start_index(doc)]
    if holder._p.getnext() is not chapter._p:
        parent = city._p.getparent()
        if parent is not None and chapter._p.getparent() is parent:
            city_el = city._p
            holder_el = holder._p
            parent.remove(city_el)
            chapter._p.addprevious(city_el)
            parent.remove(holder_el)
            chapter._p.addprevious(holder_el)
            city = Paragraph(city_el, doc)
            holder = Paragraph(holder_el, doc)
            changed += 1

    body_start = find_body_start_index(doc)
    holder_idx = _paragraph_index_of(doc, holder)
    if holder_idx is not None:
        while holder_idx + 1 < body_start:
            mid = doc.paragraphs[holder_idx + 1]
            if not is_paragraph_empty(mid) or _paragraph_sectpr(mid) is not None:
                break
            parent = mid._p.getparent()
            if parent is None:
                break
            parent.remove(mid._p)
            changed += 1
            body_start = find_body_start_index(doc)
            holder_idx = _paragraph_index_of(doc, holder)
            if holder_idx is None:
                break

    chapter = doc.paragraphs[find_body_start_index(doc)]
    if paragraph_has_page_break_before(chapter):
        set_page_break_before(chapter, False)
        changed += 1

    final = _final_body_sectpr(doc)
    if final is not None and final is not holder_sect:
        title_pg = final.find(qn("w:titlePg"))
        if title_pg is not None:
            final.remove(title_pg)
            changed += 1
    return changed


def validate_section_break_after_city_year(doc: Document) -> list[str]:
    """После «МИНСК YYYY» обязан быть nextPage; глава 1 — сразу новая страница."""
    idx = find_title_city_year_index(doc)
    body_start = find_body_start_index(doc)
    if idx is None or body_start <= 0:
        return []
    issues: list[str] = []
    city = doc.paragraphs[idx]
    if _paragraph_sectpr(city) is not None:
        issues.append("«МИНСК 2026» совмещён с разрывом раздела — попадает в колонтитул")
        return issues
    nxt = city._p.getnext()
    if nxt is None or nxt.tag != qn("w:p"):
        issues.append(
            "После «МИНСК 2026» нет разрыва раздела — глава 1 должна начинаться с верха стр. 2"
        )
        return issues
    holder = Paragraph(nxt, doc)
    if not _sectpr_is_next_page(_paragraph_sectpr(holder)):
        issues.append(
            "После «МИНСК 2026» нет разрыва раздела nextPage — глава 1 не с верха новой страницы"
        )
        return issues
    chapter = doc.paragraphs[body_start]
    nxt_after_holder = holder._p.getnext()
    if nxt_after_holder is not chapter._p:
        empty_before_ch1 = (
            nxt_after_holder is not None
            and nxt_after_holder.tag == qn("w:p")
            and is_paragraph_empty(Paragraph(nxt_after_holder, doc))
        )
        if empty_before_ch1:
            issues.append(
                "После разрыва раздела за «МИНСК 2026» глава 1 должна начинаться "
                "сверху новой страницы без пустых строк"
            )
        else:
            issues.append("После разрыва раздела за «МИНСК 2026» глава 1 должна идти сразу")
    if paragraph_has_page_break_before(chapter):
        issues.append("У главы 1 лишний разрыв страницы: разрыв уже стоит после «МИНСК 2026»")
    return issues


def remove_city_year_from_headers_footers(doc: Document) -> int:
    """«МИНСК YYYY» только в теле титула, не в колонтитуле."""
    changed = 0
    seen: set[int] = set()
    for hf in _iter_all_headers_footers(doc):
        hid = id(hf._element)
        if hid in seen:
            continue
        seen.add(hid)
        blob = "".join((t.text or "") for t in hf._element.iter(qn("w:t")))
        if not re.search(r"минск\s*[-–, ]?\s*20\d{2}", blob, re.IGNORECASE):
            continue
        for para in list(hf.paragraphs):
            text = para.text or ""
            if is_city_year_paragraph(text) or re.search(
                r"минск\s*[-–, ]?\s*20\d{2}", text, re.IGNORECASE
            ):
                set_paragraph_text(para, "", bold=False)
                changed += 1
        for node in list(hf._element.iter(qn("w:t"))):
            raw = node.text or ""
            if re.search(r"минск\s*[-–, ]?\s*20\d{2}", raw, re.IGNORECASE):
                node.text = ""
                changed += 1
    return changed


def compact_extra_empty_lines_in_signatory_block(
    doc: Document, profile: DocumentProfile
) -> int:
    """
    В блоке подписантов не оставлять пустые строки между согласующими.
    Разрешена одна пустая перед «Разработал:» и одна перед «Согласовано:»
    (если после «Разработал:» есть строка должности, а не сразу «Согласовано:»).
    """
    if not profile.has_signatories:
        return 0
    try:
        raz = find_razrabotal_index(doc)
        sog = find_soglasovano_index(doc)
    except ValueError:
        return 0
    changed = 0
    remove_empty_paragraphs_after_marker(doc, "Разработал:")
    remove_empty_paragraphs_after_marker(doc, "Согласовано:")
    try:
        raz = find_razrabotal_index(doc)
        sog = find_soglasovano_index(doc)
    except ValueError:
        return changed
    i = (find_acquaintance_sheet_start(doc) or len(doc.paragraphs)) - 1
    while i > raz:
        if i >= len(doc.paragraphs):
            i -= 1
            continue
        para = doc.paragraphs[i]
        if not is_paragraph_empty(para):
            i -= 1
            continue
        keep = False
        try:
            sog = find_soglasovano_index(doc)
        except ValueError:
            break
        if i == sog - 1:
            prev = i - 1
            while prev > raz and is_paragraph_empty(doc.paragraphs[prev]):
                prev -= 1
            prev_t = (
                paragraph_text_normalized(doc.paragraphs[prev]) if prev >= 0 else ""
            )
            if not prev_t.upper().startswith("РАЗРАБОТАЛ"):
                keep = True
        if not keep:
            el = para._p
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                changed += 1
                i = min(i, len(doc.paragraphs) - 1)
                continue
        i -= 1
    return changed


def _apply_title_utverzhdayu_weight(paragraph: Paragraph) -> int:
    """«УТВЕРЖДАЮ» по эталону не жирный (TITLE_UTVERZHDAYU_BOLD)."""
    changed = 0
    if not _title_has_utverzhdayu(paragraph.text or ""):
        return 0
    for run in paragraph.runs:
        want = bool(TITLE_UTVERZHDAYU_BOLD)
        if bool(run.bold) != want:
            apply_run_font(run, bold=want)
            changed += 1
    if not paragraph.runs:
        return changed
    return changed


def _format_one_title_paragraph(para: Paragraph, *, in_table: bool) -> int:
    """Применить константы TITLE_* к одному абзацу титула."""
    original = para.text or ""
    if is_paragraph_empty(para):
        return 0
    if _is_title_instruction_number_line(original):
        _format_title_cell_paragraph(
            para, TITLE_NUMBER_LABEL, size_pt=TITLE_NUMBER_FONT_PT
        )
        return 1
    if _is_title_number_underline_line(original):
        _format_title_cell_paragraph(
            para, TITLE_UNDERLINE_LINE, size_pt=14, underline=True
        )
        return 1
    if in_table and is_signatory_date_plaque(original):
        wanted = compact_title_stamp_date_plaque(original)
        if (para.text or "") != wanted:
            set_paragraph_text(para, wanted, bold=False, size_pt=14)
        clear_first_line_indent(para)
        set_single_line_spacing(para)
        ensure_paragraph_left(para)
        _clear_paragraph_tab_stops(para)
        set_paragraph_keep_lines(para, True)
        return 1
    if in_table and TITLE_IOF_RE.search(original) and ("_" in original or "\t" in original):
        if "\t" in original:
            set_paragraph_text(
                para,
                MULTI_SPACE_RE.sub(" ", original.replace("\t", " ")).strip(),
                bold=False,
                size_pt=14,
            )
        clear_first_line_indent(para)
        set_single_line_spacing(para)
        ensure_paragraph_left(para)
        _strip_paragraph_tabs(para)
        return 1
    compact = _compact_title_two_column(original)
    changed = 0
    is_city = is_city_year_paragraph(compact)
    is_stamp = (not is_city) and _is_title_approve_stamp_line(compact)
    is_job = (not is_city) and (not is_stamp) and _is_title_instruction_name_line(compact)
    if compact != original:
        set_paragraph_text(
            para,
            compact,
            bold=TITLE_UTVERZHDAYU_BOLD if _title_has_utverzhdayu(compact) else False,
        )
        changed += 1
    changed += _apply_title_utverzhdayu_weight(para)
    clear_first_line_indent(para)
    set_single_line_spacing(para)
    if is_city:
        return changed
    if is_stamp:
        if in_table:
            if TITLE_STAMP_IN_TABLE_ALIGN == "left":
                ensure_paragraph_left(para)
            else:
                ensure_paragraph_right(para)
        elif "\t" in (para.text or ""):
            _ensure_title_right_tab(para)
        elif TITLE_STAMP_ALIGN == "right":
            ensure_paragraph_right(para)
            _apply_title_stamp_right_indent(para)
        changed += 1
    elif is_job:
        if TITLE_DOC_NAME_ALIGN == "left":
            ensure_paragraph_left(para)
        changed += 1
    elif TITLE_ORG_ALIGN == "center":
        if not is_paragraph_centered(para):
            ensure_paragraph_centered(para)
            changed += 1
        for run in para.runs:
            if (run.text or "").strip() and run.bold:
                apply_run_font(run, bold=False)
                changed += 1
    return changed


def _is_title_stamp_table(table: Table) -> bool:
    blob = " ".join(cell.text for row in table.rows for cell in row.cells).upper()
    if "УТВЕРЖДАЮ" not in blob and "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ" not in blob:
        return False
    if is_agreement_signatory_table(table) or _is_acquaintance_sheet_table(table):
        return False
    return True


def _iter_title_tables(doc: Document) -> list[Table]:
    body_start = find_body_start_index(doc)
    chapter_el = doc.paragraphs[body_start]._p if body_start > 0 else None
    found: list[Table] = []
    for child in list(doc.element.body):
        if chapter_el is not None and child is chapter_el:
            break
        if child.tag == qn("w:tbl"):
            found.append(Table(child, doc))
    return found


def _set_title_stamp_table_no_borders(table: Table) -> int:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)
    return 1


def _apply_title_stamp_table_geometry(table: Table) -> int:
    """Ширины колонок: две колонки без рамок (слева название, справа гриф)."""
    n_want = len(TITLE_STAMP_TABLE_COL_TWIPS)
    unique = _row_unique_cells(table.rows[0]) if table.rows else []
    if len(unique) != n_want:
        return 0
    changed = 0
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    if tbl_w.get(qn("w:w")) != TITLE_STAMP_TABLE_WIDTH_TWIPS:
        tbl_w.set(qn("w:w"), TITLE_STAMP_TABLE_WIDTH_TWIPS)
        tbl_w.set(qn("w:type"), "dxa")
        changed += 1
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    if tbl_ind.get(qn("w:w")) != TITLE_STAMP_TABLE_IND_TWIPS:
        tbl_ind.set(qn("w:w"), TITLE_STAMP_TABLE_IND_TWIPS)
        tbl_ind.set(qn("w:type"), "dxa")
        changed += 1
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        cols = grid.findall(qn("w:gridCol"))
        if len(cols) == n_want:
            for el, width in zip(cols, TITLE_STAMP_TABLE_COL_TWIPS):
                if el.get(qn("w:w")) != width:
                    el.set(qn("w:w"), width)
                    changed += 1
    for row in table.rows:
        cells = _row_unique_cells(row)
        if len(cells) != n_want:
            continue
        for cell, width in zip(cells, TITLE_STAMP_TABLE_COL_TWIPS):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            if tc_w.get(qn("w:w")) != width:
                tc_w.set(qn("w:w"), width)
                tc_w.set(qn("w:type"), "dxa")
                changed += 1
    return changed


def _format_title_stamp_tables(doc: Document) -> int:
    changed = 0
    for table in _iter_title_tables(doc):
        if not _is_title_stamp_table(table):
            continue
        changed += _set_title_stamp_table_no_borders(table)
        changed += _apply_title_stamp_table_geometry(table)
        cells = _row_unique_cells(table.rows[0]) if table.rows else []
        if len(cells) >= 2:
            for para in cells[1].paragraphs:
                changed += _strip_paragraph_tabs(para)
    return changed


def format_title_block(doc: Document) -> int:
    """
    Титул по константам TITLE_*: шапка по центру; таблица 2 колонки без рамок
    (слева название + линия + «номер инструкции» 12 pt; справа УТВЕРЖДАЮ,
    должность, линия+И.О.Ф., дата keepLines). Без отступа 1,25 см; TNR 14
    кроме «номер инструкции». Второй титул не вставлять.
    «МИНСК YYYY» — place_title_city_year_at_bottom, не в колонтитуле.
    """
    body_start = find_body_start_index(doc)
    if body_start <= 0:
        return 0
    changed = _extract_city_year_out_of_title_tables(doc)
    changed += _ensure_title_left_right_stamp(doc)
    body_start = find_body_start_index(doc)
    for i in range(body_start):
        para = doc.paragraphs[i]
        if paragraph_is_inside_table(para):
            continue
        changed += _format_one_title_paragraph(para, in_table=False)
    changed += _format_title_stamp_tables(doc)
    changed += apply_title_instruction_number_font(doc)
    changed += remove_city_year_from_headers_footers(doc)
    changed += ensure_title_stamp_gap_after_header(doc)
    changed += compact_title_empty_paragraphs(doc)
    changed += _clear_breaks_between_stamp_table_and_city(doc)
    return changed


def validate_title_block(doc: Document) -> list[str]:
    """Наименования по центру; УТВЕРЖДАЮ справа в таблице 2 кол.; линия и 12 pt слева."""
    body_start = find_body_start_index(doc)
    if body_start <= 0:
        return []
    issues: list[str] = []
    for i in range(body_start):
        para = doc.paragraphs[i]
        if paragraph_is_inside_table(para) or is_paragraph_empty(para):
            continue
        text = para.text or ""
        collapsed = text.replace("\xa0", " ").replace("\t", " ")
        if MULTI_SPACE_RE.search(collapsed):
            issues.append(f"Двойные пробелы на титуле: {text[:50]}")
        if first_line_indent_cm(para) > FIRST_LINE_INDENT_TOLERANCE_CM:
            issues.append(f"Отступ 1,25 см на титуле: {text[:40]}")
        if is_city_year_paragraph(text):
            continue
        if _title_has_utverzhdayu(text) or _is_title_approve_stamp_line(text):
            if "\t" not in text and not is_paragraph_right(para) and not is_paragraph_centered(para):
                issues.append("«УТВЕРЖДАЮ» / гриф утверждения не выровнены")
        elif _is_title_instruction_name_line(text) or _is_title_instruction_number_line(text):
            continue
        else:
            if not is_paragraph_centered(para):
                snippet = text.strip()[:40]
                issues.append(f"Наименование на титуле не по центру: {snippet}")
            if any(bool(run.bold) for run in para.runs if (run.text or "").strip()):
                issues.append(f"Шапка предприятия жирная: {text.strip()[:40]}")
        if len(issues) >= 8:
            break
    issues.extend(validate_title_stamp_left_right(doc))
    issues.extend(validate_title_stamp_gap(doc))
    issues.extend(validate_city_year_not_in_headers_footers(doc))
    return issues


def validate_title_stamp_left_right(doc: Document) -> list[str]:
    """Левая колонка: название + линия + «номер инструкции» 12 pt; справа УТВЕРЖДАЮ и дата keepLines."""
    tables = [t for t in _iter_title_tables(doc) if _is_title_stamp_table(t)]
    if not tables:
        body_start = find_body_start_index(doc)
        has_stamp = False
        for i in range(body_start if body_start > 0 else 0):
            if _title_has_utverzhdayu(doc.paragraphs[i].text or ""):
                has_stamp = True
                break
        if has_stamp:
            return ["Титул: нет таблицы 2 колонки (слева название, справа УТВЕРЖДАЮ)"]
        return []
    issues: list[str] = []
    table = tables[0]
    if not table.rows:
        return ["Титул: пустая таблица грифа"]
    cells = _row_unique_cells(table.rows[0])
    if len(cells) != 2:
        issues.append("Титул: таблица грифа должна быть из двух колонок без рамок")
        return issues
    left_blob = "\n".join(p.text or "" for p in cells[0].paragraphs)
    right_blob = "\n".join(p.text or "" for p in cells[1].paragraphs)
    if not any(_is_title_number_underline_line(p.text or "") for p in cells[0].paragraphs):
        issues.append("Титул слева: нет подчёркнутой линии для номера инструкции")
    num_paras = [
        p for p in cells[0].paragraphs if _is_title_instruction_number_line(p.text or "")
    ]
    if not num_paras:
        issues.append("Титул слева: нет подписи «номер инструкции»")
    else:
        sz = _paragraph_font_half_points(num_paras[0])
        if sz and sz != TITLE_NUMBER_FONT_HALF:
            issues.append("Титул: «номер инструкции» должен быть 12 pt, не 14")
    if "УТВЕРЖДАЮ" not in right_blob.upper():
        issues.append("Титул справа: нет «УТВЕРЖДАЮ»")
    if not TITLE_UTVERZHDAYU_BOLD:
        for para in cells[1].paragraphs:
            if not _title_has_utverzhdayu(para.text or ""):
                continue
            if any(bool(run.bold) for run in para.runs if (run.text or "").strip()):
                issues.append("Титул справа: «УТВЕРЖДАЮ» не должен быть жирным")
            break
    date_paras = [p for p in cells[1].paragraphs if is_signatory_date_plaque(p.text or "")]
    if date_paras and not _paragraph_has_keep_lines(date_paras[-1]):
        issues.append("Титул справа: плашка даты без keepLines (переносится на две строки)")
    for para in cells[1].paragraphs:
        if _paragraph_has_tab_xml(para):
            issues.append("Титул справа: осталась табуляция во втором столбце")
            break
    if "УТВЕРЖДАЮ" in left_blob.upper():
        issues.append("«УТВЕРЖДАЮ» должен быть в правой колонке, не в левой")
    if "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ" in right_blob.upper():
        issues.append("Название документа попало в правую колонку грифа")
    if TITLE_NUMBER_LABEL in right_blob.casefold():
        issues.append("«номер инструкции» должен быть слева, не справа")
    if any(_is_title_approver_job_line(p.text or "") for p in cells[0].paragraphs):
        issues.append("Должность утверждающего попала в левую колонку грифа")
    if any(_is_title_document_name_line(p.text or "") for p in cells[1].paragraphs):
        issues.append("Название документа попало в правую колонку грифа")
    return issues


def validate_city_year_not_in_headers_footers(doc: Document) -> list[str]:
    """«МИНСК YYYY» нельзя класть в колонтитул — только тело 1-й страницы."""
    for hf in _iter_all_headers_footers(doc):
        blob = "".join((t.text or "") for t in hf._element.iter(qn("w:t")))
        if re.search(r"минск\s*[-–, ]?\s*20\d{2}", blob, re.IGNORECASE):
            return ["«МИНСК 2026» стоит в колонтитуле — нужно в теле 1-й страницы"]
        for para in hf.paragraphs:
            if is_city_year_paragraph(para.text or ""):
                return ["«МИНСК 2026» стоит в колонтитуле — нужно в теле 1-й страницы"]
    return []


def _length_mm(length) -> float:
    if length is None:
        return 0.0
    return float(length.mm)


def apply_page_setup_deloproizvodstvo(doc: Document) -> int:
    """
    Формат А4 и поля: левое 30 мм, правое 10 мм, верх/низ 20 мм.
    Минюст №65 п.18 допускает правое ≥8 мм; для документов СНиОТ этого агента — 10 мм.
    """
    changed = 0
    for section in doc.sections:
        if abs(_length_mm(section.page_width) - PAGE_WIDTH_MM) > MARGIN_TOLERANCE_MM:
            section.page_width = Mm(PAGE_WIDTH_MM)
            changed += 1
        if abs(_length_mm(section.page_height) - PAGE_HEIGHT_MM) > MARGIN_TOLERANCE_MM:
            section.page_height = Mm(PAGE_HEIGHT_MM)
            changed += 1
        if abs(_length_mm(section.left_margin) - MARGIN_LEFT_MM) > MARGIN_TOLERANCE_MM:
            section.left_margin = Mm(MARGIN_LEFT_MM)
            changed += 1
        if abs(_length_mm(section.right_margin) - MARGIN_RIGHT_MM) > MARGIN_TOLERANCE_MM:
            section.right_margin = Mm(MARGIN_RIGHT_MM)
            changed += 1
        if abs(_length_mm(section.top_margin) - MARGIN_TOP_MM) > MARGIN_TOLERANCE_MM:
            section.top_margin = Mm(MARGIN_TOP_MM)
            changed += 1
        if abs(_length_mm(section.bottom_margin) - MARGIN_BOTTOM_MM) > MARGIN_TOLERANCE_MM:
            section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
            changed += 1
    return changed


def validate_page_margins(doc: Document) -> list[str]:
    """Проверка полей и формата А4 (Инструкция по делопроизводству 2025, п.18)."""
    issues: list[str] = []
    for idx, section in enumerate(doc.sections, start=1):
        prefix = f"Секция {idx}"
        checks = (
            ("ширина", section.page_width, PAGE_WIDTH_MM),
            ("высота", section.page_height, PAGE_HEIGHT_MM),
            ("левое поле", section.left_margin, MARGIN_LEFT_MM),
            ("правое поле", section.right_margin, MARGIN_RIGHT_MM),
            ("верхнее поле", section.top_margin, MARGIN_TOP_MM),
            ("нижнее поле", section.bottom_margin, MARGIN_BOTTOM_MM),
        )
        for label, actual, expected_mm in checks:
            if abs(_length_mm(actual) - expected_mm) > MARGIN_TOLERANCE_MM:
                got = _length_mm(actual)
                extra = ""
                if label == "правое поле":
                    extra = " (СНиОТ 10 мм; Минюст №65 п.18 допускает ≥8 мм)"
                issues.append(
                    f"{prefix}: {label} {got:.1f} мм (нужно {expected_mm} мм{extra})"
                )
    return issues


def set_single_line_spacing(paragraph: Paragraph) -> None:
    """Одинарный интервал 1.0 и before/after = 0 — общее правило делопроизводства."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")


def paragraph_has_single_line_spacing(paragraph: Paragraph) -> bool:
    pf = paragraph.paragraph_format
    if pf.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
        return False
    if pf.line_spacing_rule == WD_LINE_SPACING.DOUBLE:
        return False
    if pf.line_spacing_rule == WD_LINE_SPACING.SINGLE:
        ls = pf.line_spacing
        if ls is None:
            return True
        try:
            return float(ls) <= 1.05
        except (TypeError, ValueError):
            return True
    ls = pf.line_spacing
    if ls is None:
        return True
    try:
        return float(ls) <= 1.05
    except (TypeError, ValueError):
        return True


def paragraph_has_zero_block_spacing(paragraph: Paragraph) -> bool:
    """True, если интервал перед/после абзаца явно 0 (не наследование стиля)."""
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        return False
    before = spacing.get(qn("w:before"))
    after = spacing.get(qn("w:after"))
    if before is None or after is None:
        return False
    try:
        return int(before) == 0 and int(after) == 0
    except (TypeError, ValueError):
        return False


def should_apply_compact_single_spacing(text: str, idx: int, doc: Document) -> bool:
    """Тело, титул, заголовки глав — одинарный + before/after 0. Подписанты — нет."""
    razrab_idx = get_signatory_start_index(doc)
    if razrab_idx is not None and idx >= razrab_idx:
        return False
    upper = (text or "").strip().upper()
    if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")):
        return False
    return True


def apply_body_single_line_spacing(doc: Document, profile: DocumentProfile) -> int:
    """Одинарный интервал и before/after=0 на теле/титуле/главах; подписанты не трогаем (там 1,5)."""
    changed = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if not should_apply_compact_single_spacing(paragraph.text, idx, doc):
            continue
        needs = not paragraph_has_single_line_spacing(
            paragraph
        ) or not paragraph_has_zero_block_spacing(paragraph)
        if needs:
            set_single_line_spacing(paragraph)
            changed += 1
    return changed


def validate_body_line_spacing(doc: Document, profile: DocumentProfile) -> list[str]:
    """Тело и заголовки глав — интервал 1.0 и before/after 0. Подписанты — в validate_signatory_block."""
    issues: list[str] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        t = paragraph.text.strip()
        is_chapter = bool(t) and is_chapter_header(t)
        is_body = should_apply_body_paragraph_format(paragraph.text, idx, doc)
        if not is_body and not is_chapter:
            continue
        snippet = t[:40] or f"абзац {idx + 1}"
        if not paragraph_has_single_line_spacing(paragraph):
            issues.append(
                f"Тело: не одинарный интервал у «{snippet}» (делопроизводство, гл.7)"
            )
        if not paragraph_has_zero_block_spacing(paragraph):
            issues.append(
                f"Тело: интервал перед/после абзаца не 0 у «{snippet}»"
            )
    return issues


def clear_paragraph_page_layout(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.page_break_before = False
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is not None:
        for tag in ("w:keepNext", "w:keepLines", "w:pageBreakBefore"):
            el = p_pr.find(qn(tag))
            if el is not None:
                p_pr.remove(el)
    for run in paragraph.runs:
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                run._element.remove(br)


def get_run_properties(run) -> object | None:
    r_pr = run._r.find(qn("w:rPr"))
    return deepcopy(r_pr) if r_pr is not None else None


def apply_run_font(run, *, bold: bool = False, size_pt: int | None = None) -> None:
    """TNR на run: по умолчанию 14 pt (sz 28); «номер инструкции» — 12 pt (sz 24)."""
    pt = int(size_pt) if size_pt else 14
    half = str(pt * 2)
    run.font.name = FONT_NAME
    run.font.size = Pt(pt)
    run.bold = bold
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), FONT_NAME)
    for tag, val in (("w:sz", half), ("w:szCs", half)):
        el = r_pr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            r_pr.append(el)
        el.set(qn("w:val"), val)
    bold_val = "true" if bold else "false"
    for tag in ("w:b", "w:bCs"):
        el = r_pr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            r_pr.append(el)
        el.set(qn("w:val"), bold_val)
    if MISSING_DEVELOPER_MARKER not in (run.text or ""):
        for tag in ("w:highlight", "w:shd"):
            el = r_pr.find(qn(tag))
            if el is not None:
                r_pr.remove(el)


def _set_run_underline(run, enabled: bool = True) -> None:
    r_pr = run._r.get_or_add_rPr()
    el = r_pr.find(qn("w:u"))
    if enabled:
        if el is None:
            el = OxmlElement("w:u")
            r_pr.append(el)
        el.set(qn("w:val"), "single")
        return
    if el is not None:
        r_pr.remove(el)


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool | None = False,
    size_pt: int | None = None,
    underline: bool = False,
) -> None:
    saved_r_pr = get_run_properties(paragraph.runs[0]) if paragraph.runs else None
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    if saved_r_pr is not None:
        old = run._r.find(qn("w:rPr"))
        if old is not None:
            run._r.remove(old)
        run._r.insert(0, saved_r_pr)
    apply_run_font(run, bold=bool(bold), size_pt=size_pt)
    if bold is True:
        run.bold = True
    elif bold is False:
        run.bold = False
    if underline:
        _set_run_underline(run, True)


def insert_empty_paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def normalize_document_fonts(doc: Document) -> int:
    """TNR 14 на все runs; тело и номера пунктов не жирные; главы — жирные.
    «номер инструкции» на титуле — 12 pt, не 14."""
    changed = 0
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        is_num = _is_title_instruction_number_line(paragraph.text)
        is_ch = is_chapter_header(paragraph_text_normalized(paragraph) or paragraph.text)
        size_pt = TITLE_NUMBER_FONT_PT if is_num else None
        for run in paragraph.runs:
            apply_run_font(run, bold=False if is_num else is_ch, size_pt=size_pt)
            changed += 1
    return changed


def find_paragraph_index(doc: Document, startswith: str, *, contains: bool = False) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if (contains and startswith in text) or text.startswith(startswith):
            return i
    raise ValueError(f"Paragraph not found: {startswith!r}")


def find_section_header_index(
    doc: Document,
    section: str,
    *fuzzy_contains: str,
) -> int:
    """
    Заголовок раздела по фразе («руководствуется»), не по первому «1.4.» в тексте.
    Вводный «1.4. Назначение…» и пункт «1.8.1.» не считаются заголовком раздела.
    Номер в тексте может быть неверным (1.4. вместо 1.8.) — всё равно этот абзац.
    """
    for fuzzy in fuzzy_contains:
        needle = (fuzzy or "").strip()
        if not needle:
            continue
        low = needle.lower()
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if low not in text.lower():
                continue
            if re.match(r"^\d+\.\d+\.\d+", text):
                continue
            return i
    pat = re.compile(rf"^{re.escape(section)}\.\s+[А-ЯЁA-Z]")
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if re.match(r"^\d+\.\d+\.\d+", text):
            continue
        if pat.match(text):
            return i
    raise ValueError(f"Section header not found: {section!r}")


def find_razrabotal_index(doc: Document) -> int:
    try:
        return find_paragraph_index(doc, "Разработал:")
    except ValueError:
        pass
    for i, paragraph in enumerate(doc.paragraphs):
        upper = paragraph_text_normalized(paragraph).upper()
        if upper.startswith("РАЗРАБОТАЛ"):
            return i
    tail = find_signatory_tail_start(doc)
    if tail is not None:
        text = paragraph_text_normalized(doc.paragraphs[tail])
        if text.upper().startswith("РАЗРАБОТАЛ"):
            return tail
    raise ValueError("Paragraph not found: 'Разработал:'")


def find_soglasovano_index(doc: Document) -> int:
    try:
        return find_paragraph_index(doc, "Согласовано:")
    except ValueError:
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph_is_inside_table(paragraph):
                continue
            if paragraph.text.strip().upper().startswith("СОГЛАСОВАН"):
                return i
    raise ValueError("Paragraph not found: 'Согласовано:'")


def body_starts_with_first_chapter(doc: Document, first_chapter: str | None) -> bool:
    if not first_chapter:
        return True
    want = canonical_chapter_header(first_chapter)
    want_body = chapter_header_body(first_chapter)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if _is_org_title_paragraph(text):
            continue
        if text.startswith(first_chapter) or first_chapter in text:
            return True
        if is_chapter_header(text):
            if canonical_chapter_header(text) == want:
                return True
            if chapter_header_body(text) == want_body:
                return True
            lead = first_chapter.split()[0].rstrip(".")
            if lead and text.startswith(lead):
                return True
        upper = text.upper()
        if upper.startswith("ОБЩИЕ ПОЛОЖЕНИЯ") and "ОБЩИЕ" in first_chapter.upper():
            return True
        return False
    return False


def normalize_first_chapter_heading(doc: Document) -> None:
    """«ОБЩИЕ ПОЛОЖЕНИЯ» → «1 ОБЩИЕ ПОЛОЖЕНИЯ» (не путать с шапкой титула)."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        upper = text.upper()
        if upper.startswith("ОБЩИЕ ПОЛОЖЕНИЯ") and not re.match(r"^\d+\s+", text):
            set_paragraph_text(paragraph, "1 " + text, bold=True)
            return
        if is_chapter_header(text) or upper.startswith("1 ОБЩИЕ"):
            return


def count_empty_lines_before(doc: Document, idx: int) -> int:
    count = 0
    pos = idx - 1
    while pos >= 0 and is_paragraph_empty(doc.paragraphs[pos]):
        count += 1
        pos -= 1
    return count


def count_empty_lines_between(doc: Document, start_idx: int, end_idx: int) -> int:
    """Пустые абзацы строго между start_idx и end_idx (не включая границы)."""
    if end_idx <= start_idx + 1:
        return 0
    return sum(
        1
        for i in range(start_idx + 1, end_idx)
        if is_paragraph_empty(doc.paragraphs[i])
    )


def find_etalon_path(target: Path, explicit: Path | str | None = None) -> Path | None:
    """
    Образец только из папки Агент и только если в имени есть «образец».
    Явный путь из config/поля принимается лишь при тех же условиях, иначе игнорируется.
    Предпочитается файл, близкий по имени к целевому (старший мастер → …_образец.docx).
    Нет подходящего файла — None (оформление по правилам mdc, без ОБМЕН).
    """
    if explicit:
        try:
            candidate = normalize_sniot_path_text(explicit)
            if is_allowed_sample_path(candidate) and candidate.is_file():
                return candidate.resolve()
        except (OSError, ValueError, TypeError):
            pass
    samples = list_agent_sample_paths()
    target_l = target.name.casefold()
    if "омтс" in target_l:
        samples = [
            s for s in samples if "омтс" in s.name.casefold() or "мастер" not in s.name.casefold()
        ]
    if not samples:
        return None
    ranked = sorted(
        samples,
        key=lambda sample: (_sample_rank(target, sample), sample.name.casefold()),
        reverse=True,
    )
    best = ranked[0]
    exact, overlap = _sample_rank(target, best)
    if exact == 0 and overlap < 2:
        return None
    return best


def paragraph_match_key(text: str) -> str:
    """Нормализованный ключ для сопоставления абзацев с образцом."""
    t = (text or "").replace("\xa0", " ").replace("\u200b", "").strip()
    if not t:
        return ""
    if is_chapter_header(t):
        t = chapter_header_body(t)
    else:
        t = NUM_PREFIX.sub("", t).strip()
    return re.sub(r"\s+", " ", t).lower()


def paragraph_keys_match(a: str, b: str) -> bool:
    if not a or not b:
        return a == b
    if a == b:
        return True
    if len(a) < 12 or len(b) < 12:
        return a == b
    if a[:48] == b[:48]:
        return True
    return a in b or b in a


def _nonempty_paragraph_sequence(doc: Document) -> list[tuple[int, str]]:
    return [
        (i, paragraph_match_key(p.text))
        for i, p in enumerate(doc.paragraphs)
        if not is_paragraph_empty(p)
    ]


def copy_paragraph_format_from_etalon(
    etalon_paragraph: Paragraph,
    target_paragraph: Paragraph,
    *,
    profile: DocumentProfile,
    target_idx: int,
    doc: Document,
) -> None:
    """Скопировать выравнивание/отступ/интервал с образца, где структура совпадает."""
    text = paragraph_text_normalized(target_paragraph)
    if not text:
        return
    if is_chapter_header(text):
        if is_paragraph_centered(etalon_paragraph):
            ensure_paragraph_centered(target_paragraph)
        return
    upper = text.upper()
    if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")) or (
        profile.has_signatories
        and target_idx >= (get_signatory_start_index(doc) or len(doc.paragraphs))
    ):
        copy_signatory_ppr_from_etalon(etalon_paragraph, target_paragraph)
        return
    if should_apply_body_paragraph_format(text, target_idx, doc):
        if is_paragraph_justified(etalon_paragraph):
            ensure_paragraph_justified(target_paragraph)
        # Правило mdc: 1,25 см — образец может иметь numPr без firstLine в pPr
        ensure_first_line_indent(target_paragraph)


def align_spacing_to_etalon(
    doc: Document,
    etalon_doc: Document,
    profile: DocumentProfile,
) -> int:
    """
    Выровнять пустые строки и формат абзацев по образцу в той же папке.
    Вставляет недостающие и удаляет лишние пустые строки между парами абзацев.
    """
    from difflib import SequenceMatcher

    e_seq = _nonempty_paragraph_sequence(etalon_doc)
    d_seq = _nonempty_paragraph_sequence(doc)
    if len(e_seq) < 3 or len(d_seq) < 3:
        return 0

    e_keys = [k for _, k in e_seq]
    d_keys = [k for _, k in d_seq]
    sm = SequenceMatcher(None, e_keys, d_keys, autojunk=False)
    if sm.ratio() < 0.72:
        return 0

    insertions: list[tuple[int, int]] = []
    removals: list[int] = []

    def _empty_indices_between(start_idx: int, end_idx: int) -> list[int]:
        return [
            i
            for i in range(start_idx + 1, end_idx)
            if is_paragraph_empty(doc.paragraphs[i])
        ]

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "delete":
            continue

        def _pair_range(e_start: int, e_end: int, d_start: int) -> None:
            block = min(e_end - e_start, j2 - j1 if tag == "equal" else j2 - d_start)
            for offset in range(block):
                e_idx, e_key = e_seq[e_start + offset]
                d_idx, d_key = d_seq[d_start + offset]
                if not paragraph_keys_match(e_key, d_key):
                    continue
                copy_paragraph_format_from_etalon(
                    etalon_doc.paragraphs[e_idx],
                    doc.paragraphs[d_idx],
                    profile=profile,
                    target_idx=d_idx,
                    doc=doc,
                )

        def _sync_empty_gaps(e_idx: int, e_next_idx: int, d_idx: int, d_next_idx: int) -> None:
            d_next_text = paragraph_text_normalized(doc.paragraphs[d_next_idx])
            if is_chapter_header(d_next_text):
                # Правило mdc важнее образца: ровно одна пустая строка перед главой.
                need = 1 - count_empty_lines_between(doc, d_idx, d_next_idx)
            else:
                d_text = paragraph_text_normalized(doc.paragraphs[d_idx])
                if is_chapter_header(d_text):
                    # После заголовка главы пустую строку не копируем с образца.
                    need = 0 - count_empty_lines_between(doc, d_idx, d_next_idx)
                else:
                    need = count_empty_lines_between(etalon_doc, e_idx, e_next_idx) - count_empty_lines_between(
                        doc, d_idx, d_next_idx
                    )
            if need > 0:
                insertions.append((d_next_idx, need))
            elif need < 0:
                empties = _empty_indices_between(d_idx, d_next_idx)
                for idx in empties[: -need]:
                    removals.append(idx)

        if tag == "equal":
            _pair_range(i1, i2, j1)
            for k in range(i1, i2):
                if k + 1 >= i2:
                    break
                e_idx, _ = e_seq[k]
                e_next_idx, _ = e_seq[k + 1]
                d_idx, _ = d_seq[j1 + (k - i1)]
                d_next_idx, _ = d_seq[j1 + (k - i1) + 1]
                _sync_empty_gaps(e_idx, e_next_idx, d_idx, d_next_idx)
        elif tag == "replace":
            block = min(i2 - i1, j2 - j1)
            _pair_range(i1, i1 + block, j1)
            for offset in range(block - 1):
                e_idx, _ = e_seq[i1 + offset]
                e_next_idx, _ = e_seq[i1 + offset + 1]
                d_idx, _ = d_seq[j1 + offset]
                d_next_idx, _ = d_seq[j1 + offset + 1]
                _sync_empty_gaps(e_idx, e_next_idx, d_idx, d_next_idx)

    for idx in sorted(set(removals), reverse=True):
        el = doc.paragraphs[idx]._element
        el.getparent().remove(el)

    inserted = 0
    for d_next_idx, need in sorted(insertions, key=lambda item: item[0], reverse=True):
        for _ in range(need):
            insert_empty_paragraph_before(doc.paragraphs[d_next_idx])
            inserted += 1
    return inserted + len(removals)


def collect_spacing_metrics(doc: Document) -> dict:
    """Метрики интервалов для сравнения с образцом."""
    empty_idxs = [i for i, p in enumerate(doc.paragraphs) if is_paragraph_empty(p)]
    ch_idxs = find_chapter_header_indices(doc)
    empty_after_ch = sum(
        1
        for ci in ch_idxs
        if ci + 1 < len(doc.paragraphs) and is_paragraph_empty(doc.paragraphs[ci + 1])
    )
    razrab_empty = 0
    soglas_empty = 0
    try:
        razrab_idx = find_paragraph_index(doc, "Разработал:")
        razrab_empty = count_empty_lines_before(doc, razrab_idx)
    except ValueError:
        razrab_idx = None
    try:
        sog_idx = find_soglasovano_index(doc)
        soglas_empty = count_empty_lines_before(doc, sog_idx)
    except ValueError:
        sog_idx = None
    return {
        "total_paras": len(doc.paragraphs),
        "empty_count": len(empty_idxs),
        "empty_after_chapter": empty_after_ch,
        "razrab_empty_before": razrab_empty,
        "soglas_empty_before": soglas_empty,
        "razrab_idx": razrab_idx,
        "soglas_idx": sog_idx,
    }


def compare_spacing_to_etalon(doc: Document, etalon_doc: Document) -> dict:
    """Сравнить метрики целевого документа с образцом."""
    target = collect_spacing_metrics(doc)
    etalon = collect_spacing_metrics(etalon_doc)
    delta = {
        key: target.get(key, 0) - etalon.get(key, 0)
        for key in (
            "total_paras",
            "empty_count",
            "empty_after_chapter",
            "razrab_empty_before",
            "soglas_empty_before",
        )
    }
    return {"target": target, "etalon": etalon, "delta": delta}


def ensure_di_satp_section_headers(doc: Document) -> None:
    """Проставить номера заголовков разделов, если они потеряны (1.8., 1.9., 2.1., 2.2.)."""
    fix_missing_section_headers(doc)


def validate_di_satp_section_header_numbers(
    doc: Document, profile: DocumentProfile
) -> list[str]:
    """Заголовки разделов ДИ САТП: 1.8., 1.9., 2.1. … — не повторные 1.4./1.5."""
    if not profile.has_di_satp_numbering:
        return []
    issues: list[str] = []
    for num, fuzzy in DI_SATP_SECTION_HEADERS:
        try:
            idx = find_section_header_index(doc, num, fuzzy)
            text = doc.paragraphs[idx].text.strip()
            if not re.match(rf"^{re.escape(num)}\.\s", text):
                issues.append(
                    f"Неверный номер раздела (ожидается {num}.): «{text[:60]}»"
                )
        except ValueError:
            issues.append(f"Не найден заголовок раздела {num}. ({fuzzy[:40]}…)")
    return issues


def validate_chapter_one_wrong_intro_prefixes(
    doc: Document, profile: DocumentProfile
) -> list[str]:
    """Вводные абзацы 1.4.–1.7. должны иметь видимый номер в тексте (не скрытый numPr)."""
    if not profile.has_di_satp_numbering:
        return []
    issues: list[str] = []
    try:
        section_14_idx = find_section_header_index(
            doc, DI_SATP_CH1_GUIDED, "в своей деятельности руководствуется"
        )
    except ValueError:
        return []
    for num, fuzzy in chapter_one_intro_markers():
        try:
            idx = find_paragraph_index(doc, fuzzy, contains=True)
        except ValueError:
            continue
        if idx >= section_14_idx:
            continue
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        if not re.match(rf"^{re.escape(num)}\.\s", text):
            issues.append(
                f"Глава 1: нет видимого номера {num}. в тексте «{text[:55]}»"
            )
        elif has_word_list_numbering(paragraph):
            issues.append(
                f"Глава 1: двойная нумерация (текст {num}. + список Word): «{text[:55]}»"
            )
    return issues


def validate_duplicate_razrabotal(doc: Document, profile: DocumentProfile) -> list[str]:
    if not profile.has_signatories:
        return []
    hits = iter_razrabotal_paragraph_elements(doc)
    if len(hits) > 1:
        return [f"Дубль «Разработал:» — Word видит {len(hits)} раза (тело+таблицы)"]
    return []


def validate_di_satp_chapter_one_numbering(
    doc: Document, profile: DocumentProfile
) -> list[str]:
    if not profile.has_di_satp_numbering:
        return []
    issues: list[str] = []
    try:
        section_14_idx = find_section_header_index(
            doc, DI_SATP_CH1_GUIDED, "в своей деятельности руководствуется"
        )
    except ValueError:
        section_14_idx = len(doc.paragraphs)

    checks: list[tuple[str, str]] = list(DI_SATP_CHAPTER_ONE_NUMBERED) + [
        ("—", fuzzy) for fuzzy in DI_SATP_CHAPTER_ONE_NUMPR_MARKERS
    ]
    for num, fuzzy in checks:
        try:
            idx = find_paragraph_index(doc, fuzzy, contains=True)
        except ValueError:
            issues.append(f"Глава 1: не найден абзац ({fuzzy[:40]}…)")
            continue
        if idx >= section_14_idx:
            continue
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        if has_word_list_numbering(paragraph) or paragraph_has_manual_number(text):
            continue
        issues.append(f"Глава 1: нет нумерации у «{text[:55]}»")
    return issues


def validate_signatory_tab_stops(doc: Document, profile: DocumentProfile) -> list[str]:
    """w:tabs только на строках должность↔ФИО; маркеры «Разработал:» / «Согласовано:» — нет."""
    if not profile.has_signatories:
        return []
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return []
    issues: list[str] = []
    for paragraph in doc.paragraphs[razrab_idx:]:
        text = paragraph_text_normalized(paragraph)
        if not text:
            continue
        upper = text.upper()
        if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")):
            continue
        if "\t" not in paragraph.text and not SIGNATORY_NAME_TAIL.search(text):
            continue
        if not paragraph_has_tab_stops(paragraph):
            issues.append(f"Подписант: нет табуляторов Word (w:tabs): «{text[:55]}»")
            continue
        pos = signatory_first_tab_pos_twips(paragraph)
        if pos is None or int(pos) < MIN_SIGNATORY_TAB_TWIPS:
            issues.append("Подписант: отступ до И.О.Фамилия меньше 12 см")
        if len(issues) >= 8:
            break
    return issues


def validate_signatory_tabs(doc: Document, profile: DocumentProfile) -> list[str]:
    """Строки подписи с ФИО — табуляция должность↔ФИО (как в образце)."""
    if not profile.has_signatories:
        return []
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return []
    issues: list[str] = []
    for paragraph in doc.paragraphs[razrab_idx + 1 :]:
        text = paragraph.text.strip()
        if not text:
            continue
        upper = text.upper()
        if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")):
            continue
        if "\t" in text:
            continue
        if SIGNATORY_NAME_TAIL.search(text):
            issues.append(f"Подписант: нет табуляции должность-ФИО: «{text[:55]}»")
    return issues[:6]


def numbering_block_ranges(doc: Document) -> list[tuple[int, int, str]]:
    """Диапазоны для проверки/исправления нумерации (start, end, prefix)."""
    ranges: list[tuple[int, int, str]] = []
    try:
        i_guided = find_section_header_index(
            doc, DI_SATP_CH1_GUIDED, "в своей деятельности руководствуется"
        )
        i_know = find_section_header_index(doc, DI_SATP_CH1_KNOW, "должен знать")
        i2 = find_marker_index(doc, "2 ФУНКЦИИ", "ФУНКЦИИ И ДОЛЖНОСТНЫЕ", contains=True)
        ranges.append((i_guided + 1, i_know, DI_SATP_CH1_GUIDED))
        ranges.append((i_know + 1, i2, DI_SATP_CH1_KNOW))
        start_funcs = find_section_header_index(doc, "2.1", "выполняет следующие функции")
        start_duties = find_section_header_index(
            doc, "2.2", "Для выполнения возложенных на него функций"
        )
        end_ch2 = find_marker_index(doc, "3 ПРАВА", "3 ПРАВА", contains=True)
        ranges.append((start_funcs + 1, start_duties, "2.1"))
        ranges.append((start_duties + 1, end_ch2, "2.2"))
        ranges.append(
            (
                find_section_header_index(doc, "3.1", "имеет право") + 1,
                find_marker_index(doc, "4 ВЗАИМООТНОШЕНИЯ", "4 ВЗАИМООТНОШЕНИЯ", contains=True),
                "3.1",
            )
        )
        ranges.append(
            (
                find_section_header_index(doc, "5.1", "несет ответственность") + 1,
                find_razrabotal_index(doc),
                "5.1",
            )
        )
    except ValueError:
        pass
    return ranges


def block_list_paragraph_indices(doc: Document, start: int, end: int) -> list[int]:
    """Индексы нумеруемых пунктов внутри блока (без заголовков раздела/главы)."""
    indices: list[int] = []
    for idx in range(start, end):
        text = doc.paragraphs[idx].text.strip()
        if not text:
            continue
        if is_section_header(text) or is_chapter_header(text):
            continue
        indices.append(idx)
    return indices


def parse_numbered_item(text: str) -> tuple[str, int] | None:
    """Из «1.5.3. Текст…» → («1.5», 3)."""
    match = re.match(r"^(\d+(?:\.\d+)+)\.\s*", text.strip())
    if not match:
        return None
    parts = match.group(1).split(".")
    return ".".join(parts[:-1]), int(parts[-1])


def analyze_numbering_block(
    doc: Document,
    start: int,
    end: int,
    prefix: str,
) -> tuple[list[str], list[tuple[int, str]]]:
    """
    Проверка одного блока нумерации.
    Возвращает (сообщения об ошибках, список (idx, ожидаемый_номер) для правки).
    """
    issues: list[str] = []
    fixes: list[tuple[int, str]] = []
    indices = block_list_paragraph_indices(doc, start, end)
    expected_counter = 1
    collected: list[int] = []

    for idx in indices:
        text = doc.paragraphs[idx].text.strip()
        expected_num = f"{prefix}.{expected_counter}"
        parsed = parse_numbered_item(text)

        if parsed is None:
            issues.append(f"Без номера в блоке {prefix}.x: {text[:50]}")
            fixes.append((idx, expected_num))
            expected_counter += 1
            continue

        actual_prefix, actual_sub = parsed
        if actual_prefix != prefix:
            issues.append(
                f"Чужой префикс в блоке {prefix}.x: {actual_prefix}.x → {text[:50]}"
            )
            fixes.append((idx, expected_num))
        elif actual_sub != expected_counter:
            issues.append(
                f"Неверный номер в блоке {prefix}.x: ожидался {expected_counter}, "
                f"есть {actual_sub}: {text[:50]}"
            )
            fixes.append((idx, expected_num))
        else:
            collected.append(actual_sub)
        expected_counter += 1

    if collected and collected != list(range(1, len(collected) + 1)):
        summary = f"Пропуски в нумерации {prefix}.x: {collected}"
        if summary not in issues and not any(prefix in i for i in issues):
            issues.append(summary)

    return issues, fixes


def validate_numbering_blocks(doc: Document, profile: DocumentProfile) -> list[str]:
    """Проверка нумерации по блокам — gate перед любой правкой."""
    if not profile.has_di_satp_numbering:
        return []
    issues: list[str] = []
    for start, end, prefix in numbering_block_ranges(doc):
        block_issues, _ = analyze_numbering_block(doc, start, end, prefix)
        issues.extend(block_issues)
    return issues


def validate_numbering_block_starts(doc: Document, profile: DocumentProfile) -> list[str]:
    """
    Блок 1.8.x должен начинаться с 1.8.1., блок 1.9.x — с 1.9.1. (ловит сбой materialize).
    """
    if not profile.has_di_satp_numbering:
        return []
    issues: list[str] = []
    for start, end, prefix in numbering_block_ranges(doc):
        if start >= end:
            issues.append(
                f"Блок {prefix}.x не найден или пуст (разделы перепутаны — проверьте «1.8. … руководствуется»)"
            )
            continue
        indices = block_list_paragraph_indices(doc, start, end)
        if not indices:
            continue
        first_text = doc.paragraphs[indices[0]].text.strip()
        expected_start = f"{prefix}.1."
        if not first_text.startswith(expected_start):
            issues.append(
                f"Блок {prefix}.x начинается не с {prefix}.1.: «{first_text[:60]}»"
            )
    return issues


def validate_chapter_one_intro_numpr(
    doc: Document, profile: DocumentProfile, path: Path | None
) -> list[str]:
    """Скрытый numPr без текста 1.4.–1.7. — в Word номера не видны."""
    if not profile.has_di_satp_numbering:
        return []
    if path is not None and not is_senior_master_di_path(path):
        return []
    return [
        issue
        for issue in validate_chapter_one_wrong_intro_prefixes(doc, profile)
        if "1.4." in issue or "1.5." in issue or "1.6." in issue or "1.7." in issue
        or "видимого номера" in issue
    ]


def validate_final_document_gate(
    doc: Document,
    profile: DocumentProfile,
    *,
    path: Path | None = None,
) -> list[str]:
    """Дополнительные проверки только для финальной перечитки с диска."""
    issues: list[str] = []
    issues.extend(validate_numbering_block_starts(doc, profile))
    issues.extend(validate_chapter_one_intro_numpr(doc, profile, path))
    issues.extend(validate_source_numbering_preserved(doc, profile, path))
    issues.extend(validate_agreement_table_converted(doc, profile))
    issues.extend(validate_visual_highlights(doc))
    if profile.has_signatories:
        try:
            razrab_idx = find_razrabotal_index(doc)
        except ValueError as exc:
            issues.append(str(exc))
        else:
            if count_empty_lines_before(doc, razrab_idx) > 1:
                issues.append(
                    "Подписанты: больше одной пустой строки перед «Разработал:» — блок может оторваться от текста"
                )
            last_body = find_last_body_paragraph_before_signatories(doc)
            if last_body is None:
                issues.append("Подписанты: нет текста перед блоком подписантов")
            elif is_chapter_header(paragraph_text_normalized(doc.paragraphs[last_body])):
                issues.append(
                    "Подписанты: перед «Разработал:» только заголовок главы — блок оторван от текста"
                )
    return issues


def fix_missing_section_headers(doc: Document) -> int:
    """Проставить или исправить номер заголовка раздела (1.4. → 1.8. и т.п.)."""
    changed = 0
    for num, fuzzy in DI_SATP_SECTION_HEADERS:
        try:
            idx = find_section_header_index(doc, num, fuzzy)
        except ValueError:
            try:
                idx = find_paragraph_index(doc, fuzzy, contains=True)
            except ValueError:
                continue
        text = doc.paragraphs[idx].text.strip()
        if not re.match(rf"^{re.escape(num)}\.\s", text):
            set_paragraph_text(doc.paragraphs[idx], apply_number(text, num))
            changed += 1
    return changed


def fix_numbering_selective(doc: Document, profile: DocumentProfile) -> int:
    """
    Выборочная правка нумерации: только абзацы с ошибками в проблемных блоках.
    Если validate_numbering_blocks пуст — не вызывать (process_sniot_document).
    """
    if not profile.has_di_satp_numbering:
        return 0

    changed = fix_missing_section_headers(doc)
    fix_map: dict[int, str] = {}

    for start, end, prefix in numbering_block_ranges(doc):
        _, fixes = analyze_numbering_block(doc, start, end, prefix)
        for idx, expected_num in fixes:
            fix_map[idx] = expected_num

    for idx, expected_num in sorted(fix_map.items()):
        paragraph = doc.paragraphs[idx]
        old_text = paragraph.text.strip()
        new_text = apply_number(old_text, expected_num)
        if old_text != new_text.strip():
            set_paragraph_text(paragraph, new_text)
            changed += 1

    try:
        start_duties = find_section_header_index(
            doc, "2.2", "Для выполнения возложенных на него функций"
        )
        duties_text = doc.paragraphs[start_duties].text.strip()
        canonical = apply_number(duties_text, "2.2")
        if duties_text != canonical.strip() and not re.match(r"^2\.2\.\s", duties_text):
            set_paragraph_text(doc.paragraphs[start_duties], canonical)
            changed += 1
    except ValueError:
        pass

    return changed


def fix_numbering(doc: Document, profile: DocumentProfile) -> int:
    """Обратная совместимость — только выборочная правка по результатам проверки."""
    if validate_numbering_blocks(doc, profile):
        return fix_numbering_selective(doc, profile)
    return 0


def validate_di_satp_numbering_count(
    doc: Document, profile: DocumentProfile, path: Path | None
) -> list[str]:
    if not profile.has_di_satp_numbering or path is None:
        return []
    if not is_senior_master_di_path(path):
        return []
    numbered = count_numbered_paragraphs(doc)
    if numbered < MIN_NUMBERED_SENIOR_MASTER:
        return [f"Мало нумерованных пунктов: {numbered} (ожидается >={MIN_NUMBERED_SENIOR_MASTER})"]
    return []


def validate_leading_dot_before_numbering(doc: Document) -> list[str]:
    """Лишняя точка перед нумерацией или двойная точка после номера."""
    issues: list[str] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph_visible_text(paragraph).replace("\xa0", " ").strip()
        if not text or should_skip_numbering_prefix_fix(text, idx, doc):
            continue
        if has_erroneous_numbering_prefix(text):
            issues.append(f"Лишняя точка перед нумерацией: {text[:60]}")
    return issues


def validate_unaccepted_revisions(doc: Document) -> list[str]:
    n_ins, n_del = count_tracked_change_nodes(doc)
    if n_ins or n_del:
        return [
            f"Не приняты правки рецензирования Word: вставок {n_ins}, удалений {n_del}"
        ]
    return []


def validate_duplicate_list_numbering(doc: Document) -> list[str]:
    """numPr + ручной префикс в тексте → двойная нумерация в Word."""
    issues: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if collapse_duplicate_manual_prefix(text) != text:
            issues.append(f"Дубль номера в тексте: {text[:60]}")
            continue
        if has_word_list_numbering(paragraph) and (
            paragraph_has_manual_number(text) or is_chapter_header(text) or is_section_header(text)
        ):
            issues.append(f"Двойная нумерация (список Word + текст): {text[:60]}")
    return issues


def _run_is_bold(run) -> bool:
    if run.bold is True:
        return True
    r_pr = run._r.find(qn("w:rPr"))
    if r_pr is None:
        return False
    el = r_pr.find(qn("w:b"))
    if el is None:
        return False
    val = (el.get(qn("w:val")) or "true").lower()
    return val not in ("0", "false", "off")


def _run_font_fields(run) -> tuple[str, str, str, bool]:
    """ascii, eastAsia, sz, bold — для проверки единообразия runs абзаца."""
    ascii_name = ""
    east = ""
    sz = ""
    r_pr = run._r.find(qn("w:rPr"))
    if r_pr is not None:
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is not None:
            ascii_name = r_fonts.get(qn("w:ascii")) or r_fonts.get(qn("w:hAnsi")) or ""
            east = r_fonts.get(qn("w:eastAsia")) or ""
        sz_el = r_pr.find(qn("w:sz"))
        if sz_el is not None:
            sz = sz_el.get(qn("w:val")) or ""
    if run.font.name:
        ascii_name = ascii_name or run.font.name
    if run.font.size:
        sz = sz or str(int(run.font.size.pt * 2) if hasattr(run.font.size, "pt") else 28)
    return (ascii_name, east, sz, _run_is_bold(run))


def validate_fonts(doc: Document) -> list[str]:
    issues: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        is_ch = is_chapter_header(text)
        signatures: list[tuple[str, str, str, bool]] = []
        for run in paragraph.runs:
            if not (run.text or "").strip():
                continue
            ascii_name, east, sz, bold = _run_font_fields(run)
            signatures.append((ascii_name, east, sz, bold))
            snippet = text[:40]
            if ascii_name and ascii_name != FONT_NAME:
                issues.append(f"Шрифт не {FONT_NAME}: {ascii_name!r} в «{snippet}»")
            if east and east != FONT_NAME:
                issues.append(f"Шрифт eastAsia не {FONT_NAME}: {east!r} в «{snippet}»")
            if sz and sz != "28":
                if _is_title_instruction_number_line(text) and sz == TITLE_NUMBER_FONT_HALF:
                    continue
                issues.append(f"Размер не 14 pt ({sz}) в «{snippet}»")
            if not is_ch and bold:
                issues.append(f"Жирный шрифт в теле (номер/текст пункта): «{snippet}»")
        if len(set(signatures)) > 1:
            issues.append(f"В абзаце разные шрифты runs: «{text[:40]}»")
        if len(issues) >= 12:
            break
    return issues


def validate_chapter_headers(doc: Document) -> list[str]:
    issues: list[str] = []
    for idx in find_chapter_header_indices(doc):
        text = doc.paragraphs[idx].text.strip()
        if idx > 0:
            empty_before = count_empty_lines_before(doc, idx)
            if empty_before == 0:
                issues.append(f"Нет пустой строки перед главой: {text[:40]}")
            elif empty_before > 1:
                issues.append(f"Больше одной пустой строки перед главой: {text[:40]}")
        jc = paragraph_jc(doc.paragraphs[idx])
        if not is_paragraph_centered(doc.paragraphs[idx]):
            issues.append(f"Заголовок главы не по центру: {text[:40]}")
        indent = first_line_indent_cm(doc.paragraphs[idx])
        if abs(indent) > FIRST_LINE_INDENT_TOLERANCE_CM:
            issues.append(f"У заголовка главы есть отступ 1,25 см: {text[:40]}")
    return issues


def validate_chapter_header_orphan(doc: Document) -> list[str]:
    """Заголовок главы не отрывается от текста; page_break_before не на заголовке."""
    issues: list[str] = []
    for hdr_idx in find_chapter_header_indices(doc):
        hdr = doc.paragraphs[hdr_idx]
        text = hdr.text.strip()[:40]
        if paragraph_has_page_break_before(hdr):
            first_ch = find_chapter_header_indices(doc)
            is_first = first_ch and hdr_idx == first_ch[0]
            has_title_before = is_first and hdr_idx > 0 and any(
                paragraph_text_normalized(doc.paragraphs[i])
                and not paragraph_is_inside_table(doc.paragraphs[i])
                for i in range(hdr_idx)
            )
            if not has_title_before:
                issues.append(f"Разрыв страницы на заголовке главы (сирота): {text}")
        if hdr.paragraph_format.keep_with_next or hdr.paragraph_format.keep_together:
            issues.append(f"keep_with_next/keep_together на заголовке главы: {text}")
        content_idx = find_first_nonempty_paragraph_after(doc, hdr_idx)
        if content_idx is None:
            issues.append(f"Заголовок главы без текста следом: {text}")
            continue
        if content_idx > hdr_idx + 1:
            issues.append(f"Пустая строка между заголовком и текстом главы: {text}")
    return issues


def _signatory_keep_lines_allowed(text: str) -> bool:
    """keepLines на строке должность+ФИО и на плашке даты — разрешены."""
    t = text or ""
    if is_signatory_date_plaque(t):
        return True
    return "\t" in t and bool(SIGNATORY_NAME_TAIL.search(t))


def validate_page_layout_flags(doc: Document, profile: DocumentProfile) -> list[str]:
    if not profile.has_signatories:
        return []
    issues: list[str] = []
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return issues

    if doc.paragraphs[razrab_idx].paragraph_format.page_break_before:
        issues.append("Запрещён разрыв страницы только перед «Разработал:»")

    chain = _signatory_keep_chain_range(doc)
    keep_lo, keep_hi = (chain[0], chain[1] - 1) if chain else (-1, -1)

    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph_is_inside_table(paragraph):
            continue
        if is_acquaintance_sheet_text(paragraph.text or ""):
            continue
        pf = paragraph.paragraph_format
        in_keep_chain = keep_lo <= idx < keep_hi
        if pf.keep_with_next and not in_keep_chain:
            snippet = paragraph.text.strip()[:40] or f"абзац {idx + 1}"
            issues.append(f"keep_with_next/keep_together запрещены: «{snippet}»")
            break
        if pf.keep_together and not _signatory_keep_lines_allowed(paragraph.text or ""):
            snippet = paragraph.text.strip()[:40] or f"абзац {idx + 1}"
            issues.append(f"keep_with_next/keep_together запрещены: «{snippet}»")
            break
    return issues


def validate_visual_highlights(doc: Document) -> list[str]:
    """Жёлтый маркер и цветная заливка текста/абзацев не должны оставаться."""
    issues: list[str] = []
    n_hl = 0
    for hl in doc.element.iter(qn("w:highlight")):
        if _xml_belongs_to_missing_developer_marker(hl):
            continue
        n_hl += 1
    if n_hl:
        issues.append(f"Осталась подсветка текста (highlight): {n_hl}")
    n_shd = 0
    for shd in doc.element.iter(qn("w:shd")):
        parent = shd.getparent()
        if parent is None:
            continue
        if _xml_belongs_to_missing_developer_marker(shd):
            continue
        parent_tag = parent.tag.split("}")[-1]
        fill = _shd_fill_value(shd)
        if parent_tag == "tcPr":
            if fill in HIGHLIGHT_FILL_HEX:
                n_shd += 1
            continue
        if fill in ("", "AUTO", "CLEAR", "FFFFFF"):
            continue
        if parent_tag == "pPr" and fill in {ACQUAINTANCE_MISMATCH_FILL, "RED", "C00000"}:
            p_el = parent.getparent()
            if p_el is not None and p_el.tag == qn("w:p"):
                blob = "".join((t.text or "") for t in p_el.iter(qn("w:t")))
                if is_acquaintance_sheet_text(blob):
                    continue
        n_shd += 1
    if n_shd:
        issues.append(f"Осталась заливка абзацев/текста: {n_shd}")
    return issues


def validate_agreement_table_converted(doc: Document, profile: DocumentProfile) -> list[str]:
    leftover = [table for table in doc.tables if is_agreement_signatory_table(table)]
    if not leftover:
        return []
    if find_signatory_tail_start(doc) is not None:
        return ["Таблица согласования не убрана после переноса в блок подписантов"]
    return ["Таблица согласования не перенесена в блок «Разработал:» / «Согласовано:»"]


def validate_signatory_block(doc: Document, profile: DocumentProfile) -> list[str]:
    if not profile.has_signatories:
        return []
    issues: list[str] = []
    try:
        razrab_idx = find_razrabotal_index(doc)
        soglas_idx = find_soglasovano_index(doc)
    except ValueError as exc:
        issues.append(str(exc))
        return issues

    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper().startswith("СОГЛАСОВАН") and t != "Согласовано:":
            issues.append(f"«Согласовано» без двоеточия или не то написание: {t!r}")
            break

    if razrab_idx + 1 < len(doc.paragraphs) and is_paragraph_empty(doc.paragraphs[razrab_idx + 1]):
        issues.append("Пустая строка после «Разработал:»")
    if soglas_idx + 1 < len(doc.paragraphs) and is_paragraph_empty(doc.paragraphs[soglas_idx + 1]):
        issues.append("Пустая строка после «Согласовано:»")

    if count_empty_lines_before(doc, razrab_idx) != 1:
        issues.append("Должна быть ровно одна пустая строка перед «Разработал:»")
    prev_sog = soglas_idx - 1
    while prev_sog >= 0 and is_paragraph_empty(doc.paragraphs[prev_sog]):
        prev_sog -= 1
    razrab_then_soglas = prev_sog >= 0 and paragraph_text_normalized(
        doc.paragraphs[prev_sog]
    ).upper().startswith("РАЗРАБОТАЛ")
    if razrab_then_soglas:
        if count_empty_lines_before(doc, soglas_idx) != 0:
            issues.append("Пустая строка после «Разработал:»")
    elif count_empty_lines_before(doc, soglas_idx) != 1:
        issues.append("Должна быть ровно одна пустая строка перед «Согласовано:»")

    acq = find_acquaintance_sheet_start(doc)
    block_end = acq if acq is not None else len(doc.paragraphs)
    for i in range(soglas_idx + 1, block_end):
        if is_paragraph_empty(doc.paragraphs[i]):
            issues.append("Пустая строка внутри блока согласующих — подписанты разрываются")
            break

    issues.extend(validate_signatory_line_spacing(doc, profile))

    for p in doc.paragraphs[soglas_idx : soglas_idx + 1]:
        for run in p.runs:
            if run.text and run.bold:
                issues.append("«Согласовано:» не должно быть жирным")
                break

    return issues


def validate_body_not_empty(doc: Document, profile: DocumentProfile) -> list[str]:
    n = count_nonempty_body_paragraphs(doc)
    if n == 0:
        return ["Тело документа пустое (остался только титул в content control)"]
    if profile.kind == "di" and n < 8:
        return [f"Подозрительно мало текста в теле документа ({n} непустых абзацев)"]
    return []


def recover_body_from_text_dump(docx_path: Path, dump_path: Path) -> int:
    """
    Восстановить абзацы тела из текстового дампа (_full__work_*.txt).
    Титул (sdt) в docx не трогаем — добавляем абзацы через python-docx.
    """
    if not dump_path.is_file():
        raise FileNotFoundError(dump_path)
    entries: list[tuple[bool, str]] = []
    for line in dump_path.read_text(encoding="utf-8").splitlines():
        match = re.match(r"^\s*\d+\s+b=(\d+)\s*\|\s*(.*)$", line)
        if not match:
            continue
        entries.append((bool(int(match.group(1))), match.group(2)))

    doc = Document(docx_path)
    added = 0
    for bold, text in entries:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        if bold:
            run.bold = True
        added += 1
    doc.save(docx_path)
    return added


def validate_sniot_document(
    doc: Document,
    *,
    docx_bytes: bytes | None = None,
    profile: DocumentProfile | None = None,
    path: Path | None = None,
) -> list[str]:
    """Проверка всех правил sniot-di-documents.mdc для любого документа СНиОТ."""
    if profile is None:
        profile = detect_profile(doc, path or Path("_unknown_.docx"))
    issues: list[str] = []

    issues.extend(validate_body_not_empty(doc, profile))

    if profile.first_chapter and not body_starts_with_first_chapter(doc, profile.first_chapter):
        head = " ".join(p.text for p in doc.paragraphs[:12])
        if any(k in head for k in TITLE_DUPLICATE_MARKERS):
            issues.append(
                f"Дубль титула: текст не начинается с «{profile.first_chapter[:40]}»"
            )

    issues.extend(validate_numbering_blocks(doc, profile))
    issues.extend(validate_source_numbering_preserved(doc, profile, path))
    issues.extend(validate_di_satp_section_header_numbers(doc, profile))
    issues.extend(validate_chapter_one_wrong_intro_prefixes(doc, profile))
    issues.extend(validate_di_satp_chapter_one_numbering(doc, profile))
    issues.extend(validate_leading_dot_before_numbering(doc))
    issues.extend(validate_unaccepted_revisions(doc))
    issues.extend(validate_document_punctuation(doc))
    issues.extend(validate_list_markers(doc))
    issues.extend(validate_visual_highlights(doc))
    issues.extend(validate_duplicate_list_numbering(doc))
    issues.extend(validate_di_satp_numbering_count(doc, profile, path))
    issues.extend(validate_fonts(doc))
    issues.extend(validate_page_margins(doc))
    issues.extend(validate_title_page_separated(doc))
    issues.extend(validate_title_city_year(doc))
    issues.extend(validate_title_block(doc))
    issues.extend(validate_chapter_headers(doc))
    issues.extend(validate_chapter_header_orphan(doc))
    issues.extend(validate_body_paragraph_format(doc, profile))
    issues.extend(validate_body_line_spacing(doc, profile))
    issues.extend(validate_section_headers(doc))
    issues.extend(validate_empty_lines_in_body(doc))
    issues.extend(validate_signatory_block(doc, profile))
    issues.extend(validate_agreement_table_converted(doc, profile))
    issues.extend(validate_duplicate_razrabotal(doc, profile))
    issues.extend(validate_signatory_tabs(doc, profile))
    issues.extend(validate_signatory_tab_stops(doc, profile))
    issues.extend(validate_page_layout_flags(doc, profile))
    issues.extend(validate_last_two_pages_layout(doc, profile))
    issues.extend(validate_acquaintance_sheet(doc))
    issues.extend(validate_acquaintance_job_mismatch(doc, path))
    issues.extend(validate_item_number_spacing(doc))
    issues.extend(validate_adjacent_duplicate_words(doc))
    issues.extend(validate_signatory_date_plaques(doc, profile))
    issues.extend(validate_table_header_rows_together(doc))

    if docx_bytes is not None:
        issues.extend(validate_page_numbering(docx_bytes))

    return issues


def validate_di_document(doc: Document, *, docx_bytes: bytes | None = None) -> list[str]:
    """Обратная совместимость."""
    return validate_sniot_document(doc, docx_bytes=docx_bytes)


def restore_chapter_headers(doc: Document) -> int:
    """Восстановить «2 ФУНКЦИИ…», «3 ПРАВА» и др., если номера/капс потеряны."""
    fixed = 0
    search_from = 0
    for num, title in CHAPTER_TITLES:
        canonical = f"{num} {title}"
        found_idx: int | None = None
        for i in range(search_from, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            if not text:
                continue
            if re.match(r"^\d+\.\d", text):
                continue
            body = chapter_header_body(text)
            if body == title or (title in body and len(body) <= len(title) + 8):
                found_idx = i
                break
        if found_idx is None:
            continue
        search_from = found_idx + 1
        if doc.paragraphs[found_idx].text.strip() != canonical:
            set_paragraph_text(doc.paragraphs[found_idx], canonical, bold=True)
            fixed += 1
    return fixed


def center_chapter_headers(doc: Document) -> int:
    """Заголовки глав 1–5 — по центру, жирный, капслок (любая ДИ, не только САТП)."""
    centered = 0
    for idx in find_chapter_header_indices(doc):
        paragraph = doc.paragraphs[idx]
        canonical = canonical_chapter_header(paragraph.text)
        if paragraph.text.strip() != canonical:
            set_paragraph_text(paragraph, canonical, bold=True)
        ensure_paragraph_centered(paragraph)
        pf = paragraph.paragraph_format
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        clear_first_line_indent(paragraph)
        for run in paragraph.runs:
            apply_run_font(run, bold=True)
        centered += 1
    return centered


def _load_russian_phrase_rules_module():
    """PHRASE_REPLACEMENTS из DocAgent/formatters/russian_phrase_rules.py."""
    phrase_path = DOCAGENT_FORMATTERS / "russian_phrase_rules.py"
    if phrase_path.is_file():
        import importlib.util

        spec = importlib.util.spec_from_file_location("russian_phrase_rules", phrase_path)
        if spec is not None and spec.loader is not None:
            mod = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)
            return mod
    return None


def _is_must_know_header(text: str) -> bool:
    t = MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ")).strip()
    return bool(MUST_KNOW_HEADER_RE.search(t))


def _must_know_header_child_prefix(text: str) -> str:
    match = re.match(r"^(\d+(?:\.\d+)*)\.\s+", (text or "").strip())
    if not match:
        return ""
    return match.group(1) + "."


def _must_know_block_end(doc: Document, header_idx: int) -> int:
    prefix = _must_know_header_child_prefix(doc.paragraphs[header_idx].text or "")
    raz = None
    try:
        raz = find_razrabotal_index(doc)
    except ValueError:
        pass
    acq = find_acquaintance_sheet_start(doc)
    limit = len(doc.paragraphs)
    if raz is not None:
        limit = min(limit, raz)
    if acq is not None:
        limit = min(limit, acq)
    i = header_idx + 1
    while i < limit:
        para = doc.paragraphs[i]
        t = paragraph_text_normalized(para)
        if not t:
            i += 1
            continue
        if is_chapter_header(t):
            return i
        num = re.match(r"^(\d+(?:\.\d+)*)\.\s+", t)
        if num and prefix:
            token = num.group(1) + "."
            if not token.startswith(prefix) or token == prefix:
                return i
        i += 1
    return limit


def _is_must_know_list_item(text: str) -> bool:
    t = MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ")).strip()
    if not t:
        return False
    if is_chapter_header(t) or _is_must_know_header(t):
        return False
    if t.endswith(":"):
        return False
    return True


def _must_know_list_item_indices(doc: Document, header_idx: int) -> list[int]:
    end = _must_know_block_end(doc, header_idx)
    items: list[int] = []
    for i in range(header_idx + 1, end):
        para = doc.paragraphs[i]
        if paragraph_is_inside_table(para):
            continue
        t = para.text or ""
        if not _is_must_know_list_item(t):
            continue
        items.append(i)
    return items


def _split_must_know_item(text: str) -> tuple[str, str]:
    stripped = MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ")).strip()
    match = _NUMBERED_LIST_ITEM_RE.match(stripped)
    if match:
        prefix, rest = match.group(1), match.group(2)
        return prefix + " ", _ENUM_TRAILING_PUNCT_RE.sub("", rest).rstrip()
    return "", _ENUM_TRAILING_PUNCT_RE.sub("", stripped).rstrip()


def _format_must_know_item(text: str, is_last: bool) -> str:
    prefix, core = _split_must_know_item(text)
    if not core:
        return MULTI_SPACE_RE.sub(" ", (text or "").replace("\xa0", " ")).strip()
    punct = "." if is_last else ";"
    return f"{prefix}{core}{punct}"


def normalize_must_know_list_punctuation(doc: Document) -> int:
    """В перечне «… должен знать»: промежуточные пункты «;», последний «.»."""
    changed = 0
    for idx, para in enumerate(doc.paragraphs):
        if not _is_must_know_header(para.text or ""):
            continue
        items = _must_know_list_item_indices(doc, idx)
        if not items:
            continue
        last = len(items) - 1
        for pos, pidx in enumerate(items):
            item_para = doc.paragraphs[pidx]
            original = item_para.text or ""
            new_text = _format_must_know_item(original, pos == last)
            if new_text != original:
                set_paragraph_text(item_para, new_text, bold=False)
                changed += 1
    return changed


def validate_must_know_list_punctuation(doc: Document) -> list[str]:
    """Промежуточные пункты перечисления «должен знать» — «;», последний — «.»."""
    issues: list[str] = []
    for idx, para in enumerate(doc.paragraphs):
        if not _is_must_know_header(para.text or ""):
            continue
        items = _must_know_list_item_indices(doc, idx)
        if not items:
            continue
        last = len(items) - 1
        for pos, pidx in enumerate(items):
            original = doc.paragraphs[pidx].text or ""
            expected = _format_must_know_item(original, pos == last)
            if original.strip() != expected:
                issues.append(
                    "Пунктуация перечисления «должен знать»: у промежуточных пунктов «;», "
                    "у последнего «.» (пример: «1.9.15. коллективный договор;» / "
                    f"«1.9.16. основы делопроизводства.») — {original.strip()[:70]}"
                )
                break
    return issues


def _must_know_insert_line(doc: Document, header_idx: int, end: int) -> str:
    """В нумерованном блоке 1.9.x — следующий номер; иначе строка без номера."""
    prefix = _must_know_header_child_prefix(doc.paragraphs[header_idx].text or "")
    last_num = 0
    numbered = False
    if prefix:
        child_re = re.compile(rf"^{re.escape(prefix)}(\d+)\.\s+")
        for para in doc.paragraphs[header_idx + 1 : end]:
            match = child_re.match((para.text or "").strip())
            if match:
                numbered = True
                last_num = max(last_num, int(match.group(1)))
    if numbered:
        return f"{prefix}{last_num + 1}. {OFFICE_WORK_BASICS_LINE}"
    return OFFICE_WORK_BASICS_LINE


def ensure_deloproizvodstvo_in_must_know(doc: Document) -> int:
    """В перечень «… должен знать» добавить «основы делопроизводства», если строки нет.
    Если вставка стала последним пунктом — точка; предыдущие пункты — точка с запятой.
    """
    changed = 0
    for idx, para in enumerate(doc.paragraphs):
        if not _is_must_know_header(para.text or ""):
            continue
        end = _must_know_block_end(doc, idx)
        block = doc.paragraphs[idx:end]
        if any(OFFICE_WORK_BASICS_NEEDLE in (p.text or "").casefold() for p in block):
            continue
        insert_at = end - 1
        while insert_at > idx and is_paragraph_empty(doc.paragraphs[insert_at]):
            insert_at -= 1
        anchor = doc.paragraphs[insert_at]
        new_p = insert_paragraph_after(anchor)
        set_paragraph_text(new_p, _must_know_insert_line(doc, idx, end), bold=False)
        pf = new_p.paragraph_format
        pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
        set_single_line_spacing(new_p)
        ensure_paragraph_justified(new_p)
        changed += 1
    changed += normalize_must_know_list_punctuation(doc)
    return changed


def fix_duty_by_order_commas(doc: Document) -> int:
    """«исполняет, по распоряжению …, специалист» — уточнение с обеих сторон запятыми."""
    changed = 0
    raz = None
    try:
        raz = find_razrabotal_index(doc)
    except ValueError:
        pass
    body_start = find_body_start_index(doc)
    if body_start < 0:
        body_start = 0
    end = raz if raz is not None else len(doc.paragraphs)
    for para in doc.paragraphs[body_start:end]:
        if paragraph_is_inside_table(para):
            continue
        original = para.text or ""
        if "по распоряжению" not in original.casefold():
            continue
        if "исполняет" not in original.casefold() or "специалист" not in original.casefold():
            continue

        def _repl(match: re.Match[str]) -> str:
            mid = match.group(2).strip().strip(",")
            return f"{match.group(1)}, {mid}, {match.group(3)}"

        new_text = DUTY_BY_ORDER_RE.sub(_repl, original)
        if new_text != original:
            set_paragraph_text(para, new_text, bold=False)
            changed += 1
    return changed


def _lpa_case_and_number(prefix: str) -> tuple[bool, str]:
    """Падеж и число расшифровки ЛПА по словам слева. По умолчанию — ед.ч. творительный."""
    tokens = re.findall(r"[А-Яа-яЁё]+", (prefix or "").replace("ё", "е").replace("Ё", "Е"))
    low = [t.lower() for t in tokens]
    tail = low[-5:]
    last = tail[-1] if tail else ""
    plural = False
    case = "ins"
    plural_hints = {
        "иных",
        "других",
        "всех",
        "некоторых",
        "действующих",
        "соответствующих",
        "указанных",
        "настоящих",
        "локальных",
        "правовых",
        "требований",
        "требования",
        "требованиям",
        "требованиями",
        "требованиях",
        "норм",
        "нормам",
        "нормами",
        "актов",
        "перечень",
        "перечня",
    }
    if any(w in plural_hints for w in tail):
        plural = True
    if last in {"в", "во", "на", "о", "об", "обо", "при"}:
        case = "loc"
    elif last in {"с", "со"}:
        case = "ins"
    elif last in {"к", "ко", "по", "согласно"}:
        case = "dat"
    elif last in {"от", "из", "для", "до", "без", "у"}:
        case = "gen"
    elif last.startswith("требован") or last.startswith("норм") or last.startswith("положен"):
        case = "gen"
        if not last.endswith(("е", "ие")):
            plural = True
    elif last.startswith("руководств"):
        case = "ins"
    elif last in {
        "выполняет",
        "применяет",
        "издает",
        "издаёт",
        "разрабатывает",
        "утверждает",
        "соблюдает",
        "исполняет",
    }:
        case = "acc"
    elif last.endswith(("ых", "их")):
        case = "gen"
        plural = True
    elif last.endswith(("ыми", "ими")):
        case = "ins"
        plural = True
    elif last.endswith(("ые", "ие")):
        case = "acc"
        plural = True
    elif last.endswith(("ого", "его")):
        case = "gen"
        plural = False
    elif last.endswith(("ому", "ему")):
        case = "dat"
        plural = False
    elif last.endswith(("ом", "ем")) and last not in {"в", "во", "об"}:
        case = "loc"
        plural = False
    return plural, case


def _lpa_expansion_for_prefix(prefix: str) -> str:
    plural, case = _lpa_case_and_number(prefix)
    form = _LPA_FORMS.get((plural, case)) or _LPA_FORMS[(False, "ins")]
    return f"{form} {LPA_EXPANSION_MARKER}"


def _replace_first_bare_lpa(text: str) -> str | None:
    """Первое «ЛПА» без расшифровки → форма с «(далее - ЛПА)». Иначе None."""
    for match in LPA_BARE_RE.finditer(text or ""):
        start, end = match.start(), match.end()
        left = text[:start]
        if LPA_DALEE_PREFIX_RE.search(left):
            continue
        stripped_left = left.rstrip()
        rest = text[end:]
        if stripped_left.endswith("(") and rest.lstrip().startswith(")"):
            before_paren = stripped_left[: stripped_left.rfind("(")]
            if LPA_FULL_PHRASE_TAIL_RE.search(before_paren.rstrip()):
                return text[:start] + "далее - ЛПА" + text[end:]
        return left + _lpa_expansion_for_prefix(left) + text[end:]
    return None


def expand_first_lpa_abbreviation(doc: Document) -> int:
    """
    Первое «ЛПА» без расшифровки заменить на «… правовым актом (далее - ЛПА)»
    (падеж/число по контексту). Повторные «ЛПА» не трогать.
    Спеллер по-прежнему считает «ЛПА» аббревиатурой.
    """
    body_start = find_body_start_index(doc)
    if body_start < 0:
        body_start = 0
    try:
        raz = find_razrabotal_index(doc)
    except ValueError:
        raz = len(doc.paragraphs)
    targets: list[Paragraph] = []
    for i, para in enumerate(doc.paragraphs):
        if i < body_start or i >= raz:
            continue
        raw = para.text or ""
        if not raw.strip():
            continue
        if is_chapter_header(raw) or is_acquaintance_sheet_text(raw):
            continue
        targets.append(para)
    if any(LPA_EXPANDED_RE.search(p.text or "") for p in targets):
        return 0
    for para in targets:
        new_text = _replace_first_bare_lpa(para.text or "")
        if new_text is None:
            continue
        set_paragraph_text(para, new_text, bold=False)
        return 1
    return 0


def apply_russian_phrase_rules(doc: Document, profile: DocumentProfile) -> int:
    """Канцелярские фразы в тексте абзацев тела (не заголовки, не подписанты)."""
    mod = _load_russian_phrase_rules_module()
    if mod is None:
        return 0
    apply_fn = getattr(mod, "apply_phrase_replacements", None)
    if apply_fn is None:
        return 0
    changed = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if not should_apply_body_paragraph_format(paragraph.text, idx, doc):
            continue
        original = paragraph.text
        new_text, details = apply_fn(original)
        if new_text != original and details:
            set_paragraph_text(paragraph, new_text)
            changed += len(details)
    return changed


def validate_russian_phrases_needed(doc: Document, profile: DocumentProfile) -> bool:
    """True, если в теле есть фразы для замены по russian_phrase_rules."""
    mod = _load_russian_phrase_rules_module()
    if mod is None:
        return False
    apply_fn = getattr(mod, "apply_phrase_replacements", None)
    if apply_fn is None:
        return False
    for idx, paragraph in enumerate(doc.paragraphs):
        if not should_apply_body_paragraph_format(paragraph.text, idx, doc):
            continue
        _, details = apply_fn(paragraph.text)
        if details:
            return True
    return False


def iter_table_cell_paragraphs(doc: Document):
    """Абзацы внутри w:tbl (включая вложенные таблицы)."""
    body = doc.element.body
    if body is None:
        return
    for p_el in body.iter(qn("w:p")):
        if _wp_inside_table(p_el):
            yield Paragraph(p_el, doc)


def _set_table_paragraph_no_indent(paragraph: Paragraph) -> bool:
    """Ячейка таблицы: firstLine=0 (перебивает стиль Normal 1,25 см), без hanging/left."""
    changed = False
    indent_cm = first_line_indent_cm(paragraph)
    if abs(indent_cm) > FIRST_LINE_INDENT_TOLERANCE_CM:
        changed = True
    pf = paragraph.paragraph_format
    if pf.left_indent is not None and abs(pf.left_indent.cm) > FIRST_LINE_INDENT_TOLERANCE_CM:
        pf.left_indent = Cm(0)
        changed = True
    pf.first_line_indent = Cm(0)
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
        changed = True
    if ind.get(qn("w:firstLine")) != "0":
        ind.set(qn("w:firstLine"), "0")
        changed = True
    for attr in ("w:hanging", "w:left", "w:start"):
        key = qn(attr)
        if ind.get(key) is not None:
            del ind.attrib[key]
            changed = True
    return changed


def apply_table_paragraph_no_indent(doc: Document) -> int:
    """В таблицах не применять абзацный отступ первой строки."""
    changed = 0
    for paragraph in iter_table_cell_paragraphs(doc):
        if _set_table_paragraph_no_indent(paragraph):
            changed += 1
    return changed


def validate_table_paragraph_indents(doc: Document) -> list[str]:
    """Замечание, если в ячейке таблицы есть отступ первой строки."""
    issues: list[str] = []
    for paragraph in iter_table_cell_paragraphs(doc):
        indent = first_line_indent_cm(paragraph)
        if abs(indent) <= FIRST_LINE_INDENT_TOLERANCE_CM:
            p_pr = paragraph._p.find(qn("w:pPr"))
            raw = None
            if p_pr is not None:
                ind = p_pr.find(qn("w:ind"))
                if ind is not None:
                    raw = ind.get(qn("w:firstLine"))
            if raw in (None, "0"):
                continue
            try:
                if int(raw) <= 20:
                    continue
            except (TypeError, ValueError):
                continue
        text = (paragraph.text or "").strip() or "(пустая ячейка)"
        snippet = text[:50] + ("…" if len(text) > 50 else "")
        issues.append(
            f"В таблице есть отступ первой строки ({indent:.2f} см): «{snippet}»"
        )
        if len(issues) >= 8:
            break
    return issues


def apply_body_paragraph_format(doc: Document, profile: DocumentProfile) -> int:
    """Тело: по ширине + отступ 1,25 см. Главы, подписанты и таблицы — без отступа."""
    changed = normalize_body_paragraph_styles(doc)
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph_is_inside_table(paragraph):
            continue
        text = paragraph_text_normalized(paragraph) or paragraph.text
        if not should_apply_body_paragraph_format(text, idx, doc):
            continue
        ensure_paragraph_justified(paragraph)
        ensure_first_line_indent(paragraph)
        changed += 1
    changed += apply_table_paragraph_no_indent(doc)
    return changed


def apply_header_paragraph_format(doc: Document) -> int:
    """Названия глав 1–5 («1 ОБЩИЕ ПОЛОЖЕНИЯ») — по центру, без отступа 1,25 см."""
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if is_chapter_header(text):
            clear_first_line_indent(paragraph)
            set_single_line_spacing(paragraph)
            if not is_paragraph_centered(paragraph):
                ensure_paragraph_centered(paragraph)
            changed += 1
    return changed


def validate_body_paragraph_format(doc: Document, profile: DocumentProfile) -> list[str]:
    issues: list[str] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph_is_inside_table(paragraph):
            continue
        if not should_apply_body_paragraph_format(paragraph.text, idx, doc):
            continue
        text = paragraph.text.strip()
        snippet = text[:50] + ("…" if len(text) > 50 else "")
        if not is_paragraph_justified(paragraph):
            issues.append(f"Абзац не по ширине: «{snippet}»")
        indent = first_line_indent_cm(paragraph)
        if abs(indent - FIRST_LINE_INDENT_CM) > FIRST_LINE_INDENT_TOLERANCE_CM:
            issues.append(
                f"Отступ первой строки не 1,25 см ({indent:.2f} см): «{snippet}»"
            )
    issues.extend(validate_table_paragraph_indents(doc))
    return issues


def validate_section_headers(doc: Document) -> list[str]:
    """Строки тела «1.9. … должен знать:» — с отступом 1,25 см, как пункты (не по центру)."""
    issues: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not is_section_header(text):
            continue
        indent = first_line_indent_cm(paragraph)
        if abs(indent - FIRST_LINE_INDENT_CM) > FIRST_LINE_INDENT_TOLERANCE_CM:
            issues.append(
                f"У пункта «{text[:50]}» нет отступа 1,25 см ({indent:.2f} см)"
            )
        if is_paragraph_centered(paragraph):
            issues.append(f"Пункт тела по центру (нужно по ширине): «{text[:50]}»")
    return issues


def validate_empty_lines_in_body(doc: Document) -> list[str]:
    """В теле документа не должно быть двойных/тройных пустых строк."""
    issues: list[str] = []
    body_start = find_body_start_index(doc)
    body_end = get_body_spacing_end_index(doc)

    empty_run = 0
    for i in range(body_start, body_end):
        if not is_paragraph_empty(doc.paragraphs[i]):
            if empty_run >= 2:
                issues.append(
                    f"Двойные/тройные пустые строки в теле (абз. {i - empty_run + 1}…)"
                )
            empty_run = 0
        else:
            empty_run += 1
    return issues


def _collapse_empty_run(
    doc: Document,
    run_start: int,
    run_len: int,
    *,
    body_start: int,
    body_end: int,
) -> list[int]:
    """Какие индексы пустых абзацев удалить в серии run_len подряд."""
    if run_len <= 0:
        return []
    to_remove: list[int] = []
    prev_nonempty = run_start - 1
    while prev_nonempty >= body_start and is_paragraph_empty(doc.paragraphs[prev_nonempty]):
        prev_nonempty -= 1
    if prev_nonempty >= body_start and is_chapter_header(
        paragraph_text_normalized(doc.paragraphs[prev_nonempty])
    ):
        return list(range(run_start, run_start + run_len))

    needs_one = False
    next_idx = run_start + run_len
    if next_idx < len(doc.paragraphs):
        nxt = paragraph_text_normalized(doc.paragraphs[next_idx])
        if is_chapter_header(nxt) or nxt in ("Разработал:", "Согласовано:"):
            needs_one = True
        elif nxt.upper().startswith("СОГЛАСОВАН") or nxt.upper().startswith("РАЗРАБОТАЛ"):
            needs_one = True

    if needs_one:
        if run_len > 1:
            to_remove.extend(range(run_start + 1, run_start + run_len))
    else:
        to_remove.extend(range(run_start, run_start + run_len))
    return to_remove


def remove_empty_lines_after_chapter_headers(doc: Document) -> int:
    """Убрать все пустые строки сразу после заголовков глав."""
    removed = 0
    for hdr_idx in sorted(find_chapter_header_indices(doc), reverse=True):
        while hdr_idx + 1 < len(doc.paragraphs) and is_paragraph_empty(doc.paragraphs[hdr_idx + 1]):
            el = doc.paragraphs[hdr_idx + 1]._element
            el.getparent().remove(el)
            removed += 1
    return removed


def remove_extra_empty_lines_in_body(doc: Document) -> int:
    """Убрать лишние пустые строки в теле; оставить ровно одну перед главой/подписантами."""
    body_start = find_body_start_index(doc)
    body_end = get_body_spacing_end_index(doc)

    to_remove: list[int] = []
    i = body_start
    while i < body_end:
        if not is_paragraph_empty(doc.paragraphs[i]):
            i += 1
            continue
        run_start = i
        while i < body_end and is_paragraph_empty(doc.paragraphs[i]):
            i += 1
        run_len = i - run_start
        to_remove.extend(
            _collapse_empty_run(
                doc, run_start, run_len, body_start=body_start, body_end=body_end
            )
        )

    for idx in reversed(sorted(set(to_remove))):
        el = doc.paragraphs[idx]._element
        el.getparent().remove(el)

    body_end = get_body_spacing_end_index(doc)
    extra: list[int] = []
    i = body_start
    while i < body_end:
        if not is_paragraph_empty(doc.paragraphs[i]):
            i += 1
            continue
        run_start = i
        while i < body_end and is_paragraph_empty(doc.paragraphs[i]):
            i += 1
        run_len = i - run_start
        extra.extend(
            _collapse_empty_run(
                doc, run_start, run_len, body_start=body_start, body_end=body_end
            )
        )

    for idx in reversed(sorted(set(extra))):
        el = doc.paragraphs[idx]._element
        el.getparent().remove(el)

    removed_after_headers = remove_empty_lines_after_chapter_headers(doc)
    return len(to_remove) + len(extra) + removed_after_headers


def maybe_restore_senior_master_body(docx_path: Path) -> tuple[int, str]:
    """
    Если в ДИ «Старший мастер» пропала нумерация или главы — восстановить тело из дампа.
    Не трогаем документ, если главы на месте (только нумерация без номеров — чинит fix_numbering_selective).
    """
    if not is_senior_master_di_path(docx_path):
        return 0, ""
    doc = Document(docx_path)
    numbered = count_numbered_paragraphs(doc)
    chapters = len(find_chapter_header_indices(doc))
    if numbered >= MIN_NUMBERED_SENIOR_MASTER and chapters >= 5:
        return 0, ""
    if chapters >= 5:
        return 0, ""
    if not SENIOR_MASTER_DUMP.is_file():
        return 0, f"Дамп не найден: {SENIOR_MASTER_DUMP.name}"
    try:
        added = replace_body_from_debug_dump(docx_path, SENIOR_MASTER_DUMP)
    except Exception as exc:
        return 0, f"Восстановление из дампа пропущено: {exc}"
    return added, f"Восстановлено {added} абзацев из {SENIOR_MASTER_DUMP.name}"


def renumber_block(doc: Document, start_idx: int, end_idx: int, prefix: str) -> int:
    """Полная перенумерация блока — только для ручного/legacy вызова."""
    counter = 1
    changed = 0
    for idx in range(start_idx, end_idx):
        text = doc.paragraphs[idx].text.strip()
        if not text or is_section_header(text) or is_chapter_header(text):
            continue
        new_text = apply_number(text, f"{prefix}.{counter}")
        if text != new_text.strip():
            set_paragraph_text(doc.paragraphs[idx], new_text)
            changed += 1
        counter += 1
    return changed


def fix_soglasovano(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph_is_inside_table(paragraph):
            continue
        if paragraph.text.strip().upper().startswith("СОГЛАСОВАН"):
            set_paragraph_text(paragraph, "Согласовано:", bold=False)
            return
    raise ValueError("Блок «Согласовано» не найден")


def fix_razrabotal(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph_is_inside_table(paragraph):
            continue
        text = paragraph.text.strip()
        if not text.upper().startswith("РАЗРАБОТАЛ"):
            continue
        if MISSING_DEVELOPER_MARKER in (paragraph.text or ""):
            return
        if text != "Разработал:":
            set_paragraph_text(paragraph, "Разработал:", bold=False)
        return
    raise ValueError("Блок «Разработал:» не найден")


def count_nonempty_paragraphs(doc: Document, start: int, end: int) -> int:
    return sum(1 for i in range(start, end) if doc.paragraphs[i].text.strip())


def remove_empty_paragraphs_after_index(doc: Document, idx: int) -> None:
    while idx + 1 < len(doc.paragraphs) and is_paragraph_empty(doc.paragraphs[idx + 1]):
        el = doc.paragraphs[idx + 1]._element
        el.getparent().remove(el)


def ensure_single_empty_line_before(doc: Document, marker: str) -> None:
    idx = find_paragraph_index(doc, marker)
    if idx == 0:
        insert_empty_paragraph_before(doc.paragraphs[idx])
        return
    if marker.strip().upper().startswith("СОГЛАСОВАН"):
        prev = idx - 1
        while prev >= 0 and is_paragraph_empty(doc.paragraphs[prev]):
            prev -= 1
        if prev >= 0 and paragraph_text_normalized(doc.paragraphs[prev]).upper().startswith(
            "РАЗРАБОТАЛ"
        ):
            empty_count = count_empty_lines_before(doc, idx)
            cur = idx
            for _ in range(empty_count):
                el = doc.paragraphs[cur - 1]._element
                el.getparent().remove(el)
                cur -= 1
            return
    empty_count = count_empty_lines_before(doc, idx)
    if empty_count == 0:
        insert_empty_paragraph_before(doc.paragraphs[idx])
    elif empty_count > 1:
        cur = idx
        for _ in range(empty_count - 1):
            prev = doc.paragraphs[cur - 1]
            if _paragraph_sectpr(prev) is not None:
                if (
                    cur >= 2
                    and is_paragraph_empty(doc.paragraphs[cur - 2])
                    and _paragraph_sectpr(doc.paragraphs[cur - 2]) is None
                ):
                    el = doc.paragraphs[cur - 2]._element
                    parent = el.getparent()
                    if parent is not None:
                        parent.remove(el)
                    cur -= 1
                    continue
                break
            el = prev._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
            cur -= 1


def needs_signatory_layout_compression(doc: Document) -> bool:
    try:
        ch4_idx = find_paragraph_index(doc, "4 ВЗАИМООТНОШЕНИЯ")
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return False
    if count_nonempty_paragraphs(doc, ch4_idx, razrab_idx) >= 8:
        return True
    return count_empty_lines_before(doc, razrab_idx) > 1


def ensure_chapter_header_spacing(doc: Document) -> None:
    """Ровно одна пустая строка перед главой; после заголовка пустую не добавляем."""
    markers = [doc.paragraphs[i].text.strip() for i in find_chapter_header_indices(doc)]
    for n, marker in enumerate(markers):
        idx = find_paragraph_index(doc, marker)
        first_after_minsk = (
            n == 0
            and idx > 0
            and find_title_city_year_index(doc) is not None
        )
        if first_after_minsk:
            while idx >= 2:
                prev = doc.paragraphs[idx - 1]
                if not is_paragraph_empty(prev) or _paragraph_sectpr(prev) is not None:
                    break
                prev2 = doc.paragraphs[idx - 2]
                if not (
                    _sectpr_is_next_page(_paragraph_sectpr(prev2))
                    or is_city_year_paragraph(prev2.text or "")
                ):
                    break
                el = prev._element
                parent = el.getparent()
                if parent is None:
                    break
                parent.remove(el)
                idx = find_paragraph_index(doc, marker)
        elif idx != 0:
            ensure_single_empty_line_before(doc, marker)
        idx = find_paragraph_index(doc, marker)
        remove_empty_paragraphs_after_index(doc, idx)


def prevent_chapter_header_orphan(doc: Document) -> int:
    """
    Заголовок главы не остаётся сиротой: page_break_before переносится на текст главы
    (пустую строку перед текстом, если она есть), не на заголовок.
    keep_with_next на заголовках снимается (минимально, без keep_together на длинных абзацах).
    """
    changed = 0
    for hdr_idx in find_chapter_header_indices(doc):
        hdr = doc.paragraphs[hdr_idx]

        if hdr.paragraph_format.keep_with_next or hdr.paragraph_format.keep_together:
            hdr.paragraph_format.keep_with_next = False
            hdr.paragraph_format.keep_together = False
            p_pr = hdr._p.find(qn("w:pPr"))
            if p_pr is not None:
                for tag in ("w:keepNext", "w:keepLines"):
                    el = p_pr.find(qn(tag))
                    if el is not None:
                        p_pr.remove(el)
            changed += 1

        content_idx = find_first_nonempty_paragraph_after(doc, hdr_idx)
        next_is_chapter = content_idx is not None and is_chapter_header(
            paragraph_text_normalized(doc.paragraphs[content_idx])
            or doc.paragraphs[content_idx].text
        )
        if not next_is_chapter:
            while content_idx is not None and content_idx > hdr_idx + 1:
                empty_el = doc.paragraphs[hdr_idx + 1]._element
                parent = empty_el.getparent()
                if parent is None:
                    break
                parent.remove(empty_el)
                changed += 1
                content_idx = find_first_nonempty_paragraph_after(doc, hdr_idx)
        hdr = doc.paragraphs[hdr_idx]
        target_idx = page_break_target_after_chapter_header(doc, hdr_idx)

        if not paragraph_has_page_break_before(hdr):
            continue

        first_indices = find_chapter_header_indices(doc)
        is_first_chapter = first_indices and hdr_idx == first_indices[0]
        has_title_before = is_first_chapter and hdr_idx > 0 and any(
            paragraph_text_normalized(doc.paragraphs[i])
            and not paragraph_is_inside_table(doc.paragraphs[i])
            for i in range(hdr_idx)
        )
        if has_title_before:
            continue

        set_page_break_before(hdr, False)
        # Не ставить разрыв между заголовком главы и таблицей сразу под ним.
        if (
            target_idx is not None
            and not chapter_header_followed_by_table(hdr)
            and not paragraph_sits_between_chapter_and_table(doc.paragraphs[target_idx])
        ):
            set_page_break_before(doc.paragraphs[target_idx], True)
        changed += 1

    return changed


def _xml_plain_text(el) -> str:
    return "".join((node.text or "") for node in el.iter(qn("w:t"))).strip()


def _next_body_block(el):
    """Следующий непустой абзац или таблица на том же уровне тела."""
    nxt = el.getnext()
    while nxt is not None:
        if nxt.tag == qn("w:tbl"):
            return nxt
        if nxt.tag == qn("w:p") and _xml_plain_text(nxt):
            return nxt
        nxt = nxt.getnext()
    return None


def paragraph_is_inside_table(paragraph: Paragraph) -> bool:
    for anc in paragraph._p.iterancestors():
        if anc.tag == qn("w:tbl"):
            return True
    return False


def chapter_header_followed_by_table(paragraph: Paragraph) -> bool:
    """Заголовок главы сразу над таблицей (пустые строки между ними допускаются)."""
    if not is_chapter_header(paragraph_text_normalized(paragraph)):
        return False
    nxt = _next_body_block(paragraph._p)
    return nxt is not None and nxt.tag == qn("w:tbl")


def paragraph_sits_between_chapter_and_table(paragraph: Paragraph) -> bool:
    """Абзац стоит между заголовком главы и таблицей — разрыв здесь оторвёт шапку таблицы."""
    nxt = paragraph._p.getnext()
    while nxt is not None and nxt.tag == qn("w:p") and not _xml_plain_text(nxt):
        nxt = nxt.getnext()
    if nxt is None or nxt.tag != qn("w:tbl"):
        return False
    prev = paragraph._p.getprevious()
    while prev is not None:
        if prev.tag == qn("w:tbl"):
            return False
        if prev.tag == qn("w:p") and _xml_plain_text(prev):
            return is_chapter_header(_xml_plain_text(prev))
        prev = prev.getprevious()
    return False


def _walk_tables(container):
    tables = getattr(container, "tables", None) or []
    for table in tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _walk_tables(cell)


def iter_document_tables(doc: Document):
    """Все таблицы документа, включая вложенные, без повторов из-за объединённых ячеек."""
    seen: set[int] = set()
    for table in _walk_tables(doc):
        key = id(table._tbl)
        if key in seen:
            continue
        seen.add(key)
        yield table


def iter_document_paragraphs(doc: Document):
    """Абзацы тела и ячеек таблиц без повторов (для пробелов и sanitize)."""
    seen: set[int] = set()
    for idx, paragraph in enumerate(doc.paragraphs):
        key = id(paragraph._p)
        if key in seen:
            continue
        seen.add(key)
        yield idx, paragraph
    for table in iter_document_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    key = id(paragraph._p)
                    if key in seen:
                        continue
                    seen.add(key)
                    yield -1, paragraph


def table_row_elements(table) -> list:
    """Строки w:tr этой таблицы, не строки вложенных таблиц."""
    tbl = table._tbl
    rows = []
    for el in tbl.iter(qn("w:tr")):
        parent_tbl = None
        for anc in el.iterancestors():
            if anc.tag == qn("w:tbl"):
                parent_tbl = anc
                break
        if parent_tbl is tbl:
            rows.append(el)
    return rows


def iter_direct_row_paragraphs(tr):
    for tc in tr.findall(qn("w:tc")):
        for child in tc:
            if child.tag == qn("w:p"):
                yield child


def _flag_is_on(el) -> bool:
    if el is None:
        return False
    val = el.get(qn("w:val"))
    if val is None:
        return True
    return val.lower() not in ("0", "false", "off")


def table_row_has_flag(tr, tag: str) -> bool:
    tr_pr = tr.find(qn("w:trPr"))
    if tr_pr is None:
        return False
    return _flag_is_on(tr_pr.find(qn(tag)))


def table_row_has_cant_split(tr) -> bool:
    return table_row_has_flag(tr, "w:cantSplit")


def table_row_has_tbl_header(tr) -> bool:
    return table_row_has_flag(tr, "w:tblHeader")


def _p_has_keep_next(p_el) -> bool:
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        return False
    return _flag_is_on(p_pr.find(qn("w:keepNext")))


def table_row_has_keep_next(tr) -> bool:
    return any(_p_has_keep_next(p) for p in iter_direct_row_paragraphs(tr))


def _set_tr_flag(tr, tag: str, enabled: bool) -> bool:
    tr_pr = tr.find(qn("w:trPr"))
    if tr_pr is None:
        if not enabled:
            return False
        tr_pr = OxmlElement("w:trPr")
        tr.insert(0, tr_pr)
    el = tr_pr.find(qn(tag))
    if enabled:
        if el is None:
            tr_pr.append(OxmlElement(tag))
            return True
        val = el.get(qn("w:val"))
        if val is not None and val.lower() in ("0", "false", "off"):
            del el.attrib[qn("w:val")]
            return True
        return False
    if el is not None:
        tr_pr.remove(el)
        return True
    return False


def _set_p_keep_next(p_el, enabled: bool) -> bool:
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        if not enabled:
            return False
        p_pr = OxmlElement("w:pPr")
        p_el.insert(0, p_pr)
    el = p_pr.find(qn("w:keepNext"))
    if enabled:
        if el is None:
            p_pr.append(OxmlElement("w:keepNext"))
            return True
        val = el.get(qn("w:val"))
        if val is not None and val.lower() in ("0", "false", "off"):
            del el.attrib[qn("w:val")]
            return True
        return False
    if el is not None:
        p_pr.remove(el)
        return True
    return False


def keep_table_header_rows_together(doc: Document) -> int:
    """
    Первые три строки каждой таблицы остаются вместе.
    Если таблица короче — все её строки вместе.
    Хвост (4-я строка и дальше) можно переносить на следующую страницу.
    """
    changed = 0
    for table in iter_document_tables(doc):
        rows = table_row_elements(table)
        if not rows:
            continue
        keep_n = min(TABLE_HEADER_KEEP_ROWS, len(rows))
        for i, tr in enumerate(rows):
            if i < keep_n:
                if _set_tr_flag(tr, "w:cantSplit", True):
                    changed += 1
                if _set_tr_flag(tr, "w:tblHeader", True):
                    changed += 1
                want_keep_next = i < keep_n - 1
                paras = list(iter_direct_row_paragraphs(tr))
                if not paras and want_keep_next:
                    continue
                for p_el in paras:
                    if _set_p_keep_next(p_el, want_keep_next):
                        changed += 1
            else:
                if _set_tr_flag(tr, "w:tblHeader", False):
                    changed += 1
                for p_el in iter_direct_row_paragraphs(tr):
                    if _set_p_keep_next(p_el, False):
                        changed += 1
    return changed


def validate_table_header_rows_together(doc: Document) -> list[str]:
    """Первые 3 строки таблицы связаны; 4-я не склеивает всю таблицу."""
    issues: list[str] = []
    for t_idx, table in enumerate(iter_document_tables(doc), start=1):
        rows = table_row_elements(table)
        if not rows:
            continue
        keep_n = min(TABLE_HEADER_KEEP_ROWS, len(rows))
        for i, tr in enumerate(rows):
            row_no = i + 1
            if i < keep_n:
                if not table_row_has_cant_split(tr):
                    issues.append(
                        f"Таблица {t_idx}: строка {row_no} может оторваться от шапки "
                        "(первые три строки должны оставаться вместе)"
                    )
                    break
                if not table_row_has_tbl_header(tr):
                    issues.append(
                        f"Таблица {t_idx}: строка {row_no} не помечена как шапка при переносе"
                    )
                    break
                if i < keep_n - 1 and not table_row_has_keep_next(tr):
                    issues.append(
                        f"Таблица {t_idx}: строка {row_no} не связана со следующей "
                        "(верхние три строки должны идти вместе)"
                    )
                    break
            elif i == keep_n and table_row_has_keep_next(tr):
                issues.append(
                    f"Таблица {t_idx}: строка {row_no} склеивает всю таблицу — "
                    "хвост должен переноситься на следующую страницу"
                )
                break
        if len(issues) >= 8:
            break
    return issues


def remove_empty_paragraphs_after_marker(doc: Document, marker: str) -> None:
    idx = find_paragraph_index(doc, marker)
    while idx + 1 < len(doc.paragraphs) and is_paragraph_empty(doc.paragraphs[idx + 1]):
        el = doc.paragraphs[idx + 1]._element
        el.getparent().remove(el)


def ensure_razrabotal_marker(doc: Document) -> bool:
    """Поставить «Разработал:», если блок подписи есть, а маркера нет (включая таблицы)."""
    if iter_razrabotal_paragraph_elements(doc):
        return False
    try:
        find_razrabotal_index(doc)
        return False
    except ValueError:
        pass
    tail = find_signatory_tail_start(doc)
    if tail is None:
        return False
    text = paragraph_text_normalized(doc.paragraphs[tail])
    if text.upper().startswith("РАЗРАБОТАЛ"):
        if MISSING_DEVELOPER_MARKER in (doc.paragraphs[tail].text or ""):
            return False
        if text != "Разработал:":
            set_paragraph_text(doc.paragraphs[tail], "Разработал:", bold=False)
            return True
        return False
    marker = insert_empty_paragraph_before(doc.paragraphs[tail])
    set_paragraph_text(marker, "Разработал:", bold=False)
    return True


def _document_has_signatory_iof(doc: Document) -> bool:
    """В хвосте документа (не титул) есть строка И.О.Фамилия."""
    if find_agreement_signatory_table(doc) is not None:
        return True
    acq = find_acquaintance_sheet_start(doc)
    end = acq if acq is not None else len(doc.paragraphs)
    body_start = find_body_start_index(doc)
    if body_start < 0:
        body_start = 0
    start = max(body_start, end - 25)
    for para in doc.paragraphs[start:end]:
        if paragraph_is_inside_table(para):
            continue
        text = para.text or ""
        if SIGNATORY_NAME_TAIL.search(text) or _is_iof_only_signatory_line(text):
            return True
    return False


def _apply_missing_developer_yellow(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    hl = r_pr.find(qn("w:highlight"))
    if hl is None:
        hl = OxmlElement("w:highlight")
        r_pr.append(hl)
    hl.set(qn("w:val"), MISSING_DEVELOPER_HIGHLIGHT)
    shd = r_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        r_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), MISSING_DEVELOPER_FILL)


def _write_missing_developer_one_line(para: Paragraph) -> None:
    """Одна строка: «Разработал:» + Tab 12 см + жёлтый КТО???, без переноса."""
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    label = para.add_run("Разработал:")
    apply_run_font(label, bold=False)
    tab_run = para.add_run()
    tab_run._r.append(OxmlElement("w:tab"))
    apply_run_font(tab_run, bold=False)
    name_run = para.add_run(MISSING_DEVELOPER_MARKER)
    apply_run_font(name_run, bold=False)
    _apply_missing_developer_yellow(name_run)
    ensure_signatory_tab_stops(para)
    set_paragraph_keep_lines(para, True)
    set_one_point_five_line_spacing(para)
    clear_first_line_indent(para)


def restore_missing_developer_highlights(doc: Document) -> int:
    """После шрифтов/снятия подсветки вернуть жёлтый фон маркера КТО???."""
    changed = 0
    for para in doc.paragraphs:
        if MISSING_DEVELOPER_MARKER not in (para.text or ""):
            continue
        for run in para.runs:
            if MISSING_DEVELOPER_MARKER not in (run.text or ""):
                continue
            _apply_missing_developer_yellow(run)
            changed += 1
    return changed


def ensure_missing_developer_block(doc: Document) -> int:
    """
    Если в исходнике нет «Разработал:» и нет строки И.О.Фамилия —
    вставить одну строку «Разработал:» + Tab 12 см + жёлтый «КТО???».
    Существующие корректные подписанты не трогать.
    """
    if iter_razrabotal_paragraph_elements(doc):
        return 0
    try:
        find_razrabotal_index(doc)
        return 0
    except ValueError:
        pass
    if _document_has_signatory_iof(doc):
        return 0
    sog_idx: int | None
    try:
        sog_idx = find_soglasovano_index(doc)
    except ValueError:
        sog_idx = None
    acq_idx = find_acquaintance_sheet_start(doc)
    if sog_idx is not None:
        marker = insert_empty_paragraph_before(doc.paragraphs[sog_idx])
    elif acq_idx is not None:
        marker = insert_empty_paragraph_before(doc.paragraphs[acq_idx])
    else:
        last = doc.paragraphs[-1] if doc.paragraphs else None
        if last is None:
            return 0
        marker = insert_paragraph_after(last)
    _write_missing_developer_one_line(marker)
    return 1


def set_one_point_five_line_spacing(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = SIGNATORY_LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def paragraph_has_one_point_five_spacing(paragraph: Paragraph) -> bool:
    pf = paragraph.paragraph_format
    if pf.line_spacing_rule != WD_LINE_SPACING.ONE_POINT_FIVE:
        return False
    ls = pf.line_spacing
    if ls is None:
        return True
    try:
        return abs(float(ls) - SIGNATORY_LINE_SPACING) <= SIGNATORY_LINE_SPACING_TOLERANCE
    except (TypeError, ValueError):
        return True


def _signatory_block_end_index(doc: Document) -> int:
    """Конец блока подписантов: лист ознакомления не входит (там не интервал 1,5)."""
    acq = find_acquaintance_sheet_start(doc)
    return acq if acq is not None else len(doc.paragraphs)


def apply_signatory_line_spacing(doc: Document, profile: DocumentProfile) -> int:
    """
    Межстрочный интервал 1,5 на блоке подписантов — от «Разработал:» до последней подписи.
    На абзацы тела и лист ознакомления не распространяется.
    Должность+ФИО — одна строка (Tab, keepLines).
    """
    if not profile.has_signatories:
        return 0
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return 0
    changed = apply_signatory_fio_one_line(doc)
    end = _signatory_block_end_index(doc)
    for paragraph in doc.paragraphs[razrab_idx:end]:
        if not paragraph_text_normalized(paragraph):
            continue
        if not paragraph_has_one_point_five_spacing(paragraph):
            set_one_point_five_line_spacing(paragraph)
            changed += 1
        text = paragraph.text or ""
        if MISSING_DEVELOPER_MARKER in text or (
            "\t" in text and SIGNATORY_NAME_TAIL.search(text)
        ):
            if not _paragraph_has_keep_lines(paragraph):
                set_paragraph_keep_lines(paragraph, True)
                changed += 1
            if MISSING_DEVELOPER_MARKER in text:
                ensure_signatory_tab_stops(paragraph)
    return changed


def validate_signatory_line_spacing(doc: Document, profile: DocumentProfile) -> list[str]:
    if not profile.has_signatories:
        return []
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return []
    end = _signatory_block_end_index(doc)
    for paragraph in doc.paragraphs[razrab_idx:end]:
        if not paragraph_text_normalized(paragraph):
            continue
        if not paragraph_has_one_point_five_spacing(paragraph):
            snippet = paragraph.text.strip()[:40] or "подписант"
            return [f"Блок подписантов: интервал не 1,5 у «{snippet}»"]
    return []


def fix_signatory_block_format(
    doc: Document,
    profile: DocumentProfile,
    *,
    use_etalon_spacing: bool = False,
    etalon_doc: Document | None = None,
) -> None:
    if not profile.has_signatories:
        return
    ensure_razrabotal_marker(doc)
    try:
        fix_razrabotal(doc)
        fix_soglasovano(doc)
    except ValueError:
        return
    if etalon_doc is not None:
        apply_signatory_layout_from_etalon(doc, etalon_doc, profile)
    else:
        apply_signatory_line_spacing(doc, profile)
    try:
        razrab_idx = find_paragraph_index(doc, "Разработал:")
    except ValueError:
        razrab_idx = len(doc.paragraphs)
    for paragraph in doc.paragraphs[razrab_idx:]:
        if not paragraph_text_normalized(paragraph):
            continue
        if paragraph_is_inside_table(paragraph):
            continue
        for run in paragraph.runs:
            apply_run_font(run, bold=False)
        if "\t" not in paragraph.text:
            new_text = insert_signatory_tab_line(paragraph.text)
            if new_text != paragraph.text:
                set_paragraph_text(paragraph, new_text, bold=False)
        if "\t" in paragraph.text:
            ensure_signatory_tab_stops(paragraph)
        clear_first_line_indent(paragraph)
    apply_signatory_tab_stops(doc, profile)
    remove_empty_paragraphs_after_marker(doc, "Разработал:")
    remove_empty_paragraphs_after_marker(doc, "Согласовано:")
    ensure_single_empty_line_before(doc, "Разработал:")
    ensure_single_empty_line_before(doc, "Согласовано:")
    remove_extra_empty_lines_in_body(doc)


def find_last_body_paragraph_before_signatories(doc: Document) -> int | None:
    """Последний непустой абзац тела перед блоком подписантов (не «Разработал:»)."""
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return None
    for idx in range(razrab_idx - 1, -1, -1):
        text = paragraph_text_normalized(doc.paragraphs[idx])
        if not text:
            continue
        upper = text.upper()
        if upper.startswith(("РАЗРАБОТАЛ", "СОГЛАСОВАН")):
            continue
        return idx
    return None


def _tail_chapter_indices_before_signatories(doc: Document) -> list[int]:
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        razrab_idx = len(doc.paragraphs)
    return [i for i in find_chapter_header_indices(doc) if i < razrab_idx]


def validate_last_two_pages_layout(doc: Document, profile: DocumentProfile) -> list[str]:
    """
    Проверка предпоследней/последней страницы (эвристики без рендера Word).
    Цель: текст и подписанты вместе; нет разрыва только перед «Разработал:».
    """
    if not profile.has_signatories:
        return []

    issues: list[str] = []
    issues.extend(validate_signatory_block(doc, profile))
    issues.extend(validate_page_layout_flags(doc, profile))

    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError as exc:
        issues.append(str(exc))
        return issues

    if paragraph_has_page_break_before(doc.paragraphs[razrab_idx]):
        issues.append(
            "Запрещён разрыв страницы только перед «Разработал:» — подписанты оторваны от текста"
        )

    last_body_idx = find_last_body_paragraph_before_signatories(doc)
    if last_body_idx is None:
        issues.append("Подписанты без текста перед ними — блок оторван от содержания")
    elif is_chapter_header(paragraph_text_normalized(doc.paragraphs[last_body_idx])):
        issues.append(
            "Перед «Разработал:» только заголовок главы без текста — подписанты оторваны"
        )

    tail_chapters = _tail_chapter_indices_before_signatories(doc)
    if tail_chapters:
        ch5_idx = tail_chapters[-1]
        ch5_text = paragraph_text_normalized(doc.paragraphs[ch5_idx])
        if paragraph_has_page_break_before(doc.paragraphs[ch5_idx]):
            issues.append(f"Разрыв страницы на заголовке главы (сирота): {ch5_text[:40]}")
        content_idx = find_first_nonempty_paragraph_after(doc, ch5_idx)
        if content_idx is None or content_idx >= razrab_idx:
            issues.append(f"Заголовок главы без текста перед подписантами: {ch5_text[:40]}")
        elif content_idx > ch5_idx + 1:
            issues.append(
                f"Пустая строка между заголовком главы и текстом в хвосте: {ch5_text[:40]}"
            )

    return issues


def _signatory_keep_chain_range(doc: Document) -> tuple[int, int] | None:
    """Диапазон [start, end): хвост текста + блок подписантов, без листа ознакомления."""
    try:
        raz = find_razrabotal_index(doc)
    except ValueError:
        return None
    end = _signatory_block_end_index(doc)
    last_body = find_last_body_paragraph_before_signatories(doc)
    if last_body is not None and not is_chapter_header(
        paragraph_text_normalized(doc.paragraphs[last_body])
    ):
        start = last_body
    else:
        start = raz
        if raz > 0 and is_paragraph_empty(doc.paragraphs[raz - 1]):
            start = raz - 1
    if start >= end:
        return None
    return start, end


def keep_signatory_block_with_text(doc: Document, profile: DocumentProfile) -> int:
    """
    Подписанты не разрываются и не отрываются от основного текста:
    keepNext от последнего абзаца тела до предпоследнего абзаца блока.
    keepLines на должность+ФИО и плашке даты. Лист ознакомления не клеить.
    """
    if not profile.has_signatories:
        return 0
    chain = _signatory_keep_chain_range(doc)
    if chain is None:
        return 0
    start, end = chain
    changed = 0
    for i in range(start, end):
        para = doc.paragraphs[i]
        want_next = i < end - 1
        if _paragraph_has_keep_next(para) != want_next:
            set_paragraph_keep_with_next(para, want_next)
            changed += 1
        text = para.text or ""
        if _signatory_keep_lines_allowed(text) and not _paragraph_has_keep_lines(para):
            set_paragraph_keep_lines(para, True)
            changed += 1
    return changed


def signatories_appear_orphaned(doc: Document, profile: DocumentProfile) -> bool:
    """Эвристика: подписанты оторваны от текста (без рендера Word)."""
    if not profile.has_signatories:
        return False
    markers = (
        "оторван",
        "отдельной странице",
        "только заголовок",
        "только перед «разработал",
        "без текста перед",
        "запрещён разрыв страницы только перед «разработал",
    )
    for issue in validate_last_two_pages_layout(doc, profile):
        low = issue.lower()
        if any(m in low for m in markers):
            return True
    return False


def fix_last_pages_and_signatories(
    doc: Document,
    profile: DocumentProfile,
    *,
    mode: str = "natural",
    use_etalon_spacing: bool = False,
    etalon_doc: Document | None = None,
) -> int:
    """
    Шаг 1 — естественная вёрстка хвоста и блок подписантов.
    Снимает принудительные разрывы на теле; оформляет подписантов (1,5);
    keepNext от хвоста текста до конца блока, без пустых между согласующими.
    """
    changed = 0
    for paragraph in doc.paragraphs:
        before = (
            paragraph.paragraph_format.keep_with_next
            or paragraph.paragraph_format.keep_together
            or paragraph.paragraph_format.page_break_before
        )
        clear_paragraph_page_layout(paragraph)
        if before:
            changed += 1

    if mode != "natural":
        mode = "natural"

    fix_signatory_block_format(
        doc, profile, use_etalon_spacing=use_etalon_spacing, etalon_doc=etalon_doc
    )
    changed += prevent_chapter_header_orphan(doc)

    removed = remove_extra_empty_lines_in_body(doc)
    changed += removed
    ensure_chapter_header_spacing(doc)

    changed += apply_signatory_line_spacing(doc, profile)

    try:
        ensure_single_empty_line_before(doc, "Разработал:")
        ensure_single_empty_line_before(doc, "Согласовано:")
    except ValueError:
        pass

    try:
        remove_empty_paragraphs_after_marker(doc, "Разработал:")
    except ValueError:
        pass
    try:
        remove_empty_paragraphs_after_marker(doc, "Согласовано:")
    except ValueError:
        pass
    if etalon_doc is not None:
        changed += apply_signatory_layout_from_etalon(doc, etalon_doc, profile)
    else:
        changed += apply_signatory_line_spacing(doc, profile)
    changed += ensure_acquaintance_sheet_separate_page(doc)
    changed += fix_signatory_date_plaques(doc)
    changed += compact_extra_empty_lines_in_signatory_block(doc, profile)
    changed += keep_signatory_block_with_text(doc, profile)
    return changed


def fix_last_pages_page_breaks(
    doc: Document, profile: DocumentProfile, *, force: bool = False
) -> str:
    """
    Шаг 2 — перенос страниц только если подписанты оторваны от текста.
    force/--fix-page-breaks не ставит разрыв, если блок помещается на предпоследней
    вместе с хвостом текста.     Никогда не ставит разрыв только перед «Разработал:».
    """
    _ = force
    if not profile.has_signatories:
        return "natural"

    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return "natural"

    set_page_break_before(doc.paragraphs[razrab_idx], False)

    if not signatories_appear_orphaned(doc, profile):
        ensure_acquaintance_sheet_separate_page(doc)
        return "natural"

    if profile.tail_chapter_idx is None:
        ensure_acquaintance_sheet_separate_page(doc)
        return "natural"

    strategy = apply_signatory_page_break(doc, profile)
    ensure_acquaintance_sheet_separate_page(doc)
    return strategy


def apply_signatory_page_break(doc: Document, profile: DocumentProfile) -> str:
    """
    Шаг 2: ровно один page_break_before — на один абзац перед подписантами,
    чтобы текст и подписи оказались на одной странице.
    """
    if profile.tail_chapter_idx is None or not profile.has_signatories:
        return "natural"
    try:
        razrab_idx = find_razrabotal_index(doc)
    except ValueError:
        return "natural"

    for idx in range(profile.tail_chapter_idx, razrab_idx):
        set_page_break_before(doc.paragraphs[idx], False)

    last_body = find_last_body_paragraph_before_signatories(doc)
    if last_body is None:
        return "natural"

    target = last_body
    if is_chapter_header(paragraph_text_normalized(doc.paragraphs[target])):
        content_idx = find_first_nonempty_paragraph_after(doc, profile.tail_chapter_idx)
        if content_idx is None or content_idx >= razrab_idx:
            return "natural"
        target = content_idx

    if (
        paragraph_sits_between_chapter_and_table(doc.paragraphs[target])
        or chapter_header_followed_by_table(doc.paragraphs[target])
    ):
        return "natural"

    set_page_break_before(doc.paragraphs[target], True)
    return "before_paragraph"


def _fix_chapter_one_numbering(
    doc: Document,
    etalon_doc: Document | None,
    profile: DocumentProfile,
    *,
    use_etalon: bool,
) -> None:
    """Глава 1: видимые 1.1.–1.7. в тексте; раздел «1.8. … руководствуется» — по фразе."""
    if not profile.has_di_satp_numbering:
        return
    materialize_chapter_one_numbering(doc, profile)
    ensure_di_satp_chapter_one_numbering(doc, profile)


def _fix_numbering_if_needed(doc: Document, profile: DocumentProfile) -> None:
    """Нумерация разделов САТП — только при ошибках validate. Чужую ДИ не перенумеровывать."""
    if not profile.has_di_satp_numbering:
        return
    if validate_di_satp_section_header_numbers(doc, profile):
        fix_missing_section_headers(doc)
    if validate_numbering_blocks(doc, profile):
        fix_numbering_selective(doc, profile)


def _fix_duplicate_numbering_if_needed(doc: Document) -> None:
    if validate_duplicate_list_numbering(doc) or validate_leading_dot_before_numbering(doc):
        strip_leading_dot_before_numbering(doc)
        deduplicate_manual_and_list_numbering(doc)
        remove_word_list_numbering_in_body(doc)
        ensure_single_numbering_model(doc)


def apply_mandatory_layout_fixes(doc: Document, profile: DocumentProfile) -> int:
    """
    Обязательные правки в конце process_sniot_document и GUI «Оформить документ».
    Иначе следующий проход затирает XML или кнопка оставляет старый файл.
    Не навязывает схему САТП 1.8/1.9.
    """
    changed = 0
    changed += restore_chapter_headers(doc)
    normalize_first_chapter_heading(doc)
    changed += center_chapter_headers(doc)
    changed += apply_header_paragraph_format(doc)
    changed += apply_page_setup_deloproizvodstvo(doc)
    changed += strip_unnecessary_characters(doc)
    changed += remove_list_markers_in_body(doc)
    changed += normalize_item_number_spacing(doc)
    changed += apply_body_paragraph_format(doc, profile)
    changed += apply_header_paragraph_format(doc)
    changed += apply_body_single_line_spacing(doc, profile)
    changed += apply_table_paragraph_no_indent(doc)
    changed += ensure_title_page_separated(doc)
    changed += format_title_block(doc)
    changed += place_title_city_year_at_bottom(doc)
    changed += ensure_section_break_after_city_year(doc)
    changed += compact_title_empty_paragraphs(doc)
    changed += ensure_title_stamp_gap_after_header(doc)
    changed += compact_title_empty_paragraphs(doc)
    changed += _clear_breaks_between_stamp_table_and_city(doc)
    changed += remove_city_year_from_headers_footers(doc)
    changed += prevent_chapter_header_orphan(doc)
    changed += normalize_item_number_spacing(doc)
    changed += ensure_deloproizvodstvo_in_must_know(doc)
    changed += fix_duty_by_order_commas(doc)
    changed += expand_first_lpa_abbreviation(doc)
    changed += ensure_missing_developer_block(doc)
    if has_signatory_block(doc):
        profile.has_signatories = True
    if profile.has_signatories:
        changed += apply_signatory_tab_stops(doc, profile)
        changed += apply_signatory_fio_one_line(doc)
        changed += apply_signatory_line_spacing(doc, profile)
        changed += fix_signatory_date_plaques(doc)
        changed += compact_extra_empty_lines_in_signatory_block(doc, profile)
        changed += apply_signatory_line_spacing(doc, profile)
        changed += keep_signatory_block_with_text(doc, profile)
    changed += restore_missing_developer_highlights(doc)
    changed += check_document_punctuation_after_edit(doc)
    return changed


def collect_five_layout_issues(doc: Document, profile: DocumentProfile) -> list[str]:
    """Замечания только по пяти правкам кнопки — ими нельзя блокировать запись, если они уже исправлены."""
    issues: list[str] = []
    for item in validate_unnecessary_characters(doc):
        low = item.lower()
        if "двойные пробелы" in low or "неразрывный пробел" in low:
            issues.append(item)
    issues.extend(validate_item_number_spacing(doc))
    issues.extend(validate_title_block(doc))
    issues.extend(validate_title_city_year(doc))
    issues.extend(validate_chapter_header_orphan(doc))
    issues.extend(validate_signatory_date_plaques(doc, profile))
    issues.extend(validate_list_markers(doc))
    issues.extend(validate_body_line_spacing(doc, profile))
    issues.extend(validate_table_paragraph_indents(doc))
    if profile.has_signatories:
        issues.extend(validate_signatory_line_spacing(doc, profile))
    return issues


def process_sniot_document(
    doc: Document,
    profile: DocumentProfile,
    *,
    apply_page_breaks: bool = False,
    etalon_path: Path | None = None,
    source_path: Path | None = None,
) -> str:
    """
    Порядок правок (правило 25):
    title (до загрузки Document) → numbering → font → chapters (center)
    → body (justify + 1,25 см) → spacing → fix_last_pages_and_signatories (шаг 1)
    → fix_last_pages_page_breaks при необходимости (шаг 2)
    → keep_table_header_rows_together (все документы).

    Уже исправленные абзацы не перезаписываются: каждый блок правится только
    если validate находит ошибки (сохранение принятых исправлений в _оформлен.docx).
    В начале — принять правки рецензирования Word (иначе «не приняты» и точка перед 5.1.1).
    """
    _ = apply_page_breaks
    etalon_doc: Document | None = None
    if etalon_path and etalon_path.is_file():
        etalon_doc = Document(etalon_path)
    use_etalon = etalon_doc is not None

    accept_tracked_changes(doc)
    snapshot: list[str] = list(profile.source_number_tokens or ())
    if not snapshot and source_path is not None:
        sibling = find_agent_sibling_source(source_path)
        if sibling is not None:
            try:
                snapshot = collect_number_tokens(open_word_document_readonly(sibling))
            except Exception:
                snapshot = []
    if not snapshot:
        snapshot = collect_number_tokens(doc)
    profile.source_number_tokens = tuple(snapshot)
    materialize_word_decimal_numbering(doc)
    strip_visual_highlights(doc)
    restore_org_header_if_npa_inserted(doc)
    restore_lsim_if_osim_poisoned(doc, source_path)
    normalize_first_chapter_heading(doc)
    if materialize_signatory_paragraphs_from_tables(doc):
        profile.has_signatories = True
    if has_signatory_block(doc):
        profile.has_signatories = True
    if ensure_missing_developer_block(doc):
        profile.has_signatories = True
    if profile.has_signatories:
        deduplicate_razrabotal_markers(doc)
        ensure_razrabotal_marker(doc)
        deduplicate_razrabotal_markers(doc)
        restore_missing_developer_highlights(doc)
    strip_unnecessary_characters(doc)
    remove_list_markers_in_body(doc)
    restore_chapter_headers(doc)

    _fix_chapter_one_numbering(doc, etalon_doc, profile, use_etalon=use_etalon)
    _fix_numbering_if_needed(doc, profile)
    _fix_duplicate_numbering_if_needed(doc)

    normalize_document_fonts(doc)
    apply_page_setup_deloproizvodstvo(doc)
    ensure_title_page_separated(doc)
    format_title_block(doc)
    place_title_city_year_at_bottom(doc)
    center_chapter_headers(doc)
    strip_unnecessary_characters(doc)
    _fix_duplicate_numbering_if_needed(doc)
    remove_word_list_numbering_in_body(doc)
    remove_list_markers_in_body(doc)
    normalize_body_paragraph_styles(doc)
    apply_body_paragraph_format(doc, profile)
    apply_body_single_line_spacing(doc, profile)
    apply_header_paragraph_format(doc)
    if not use_etalon:
        if validate_empty_lines_in_body(doc):
            remove_extra_empty_lines_in_body(doc)
        ensure_chapter_header_spacing(doc)
    apply_body_paragraph_format(doc, profile)
    if use_etalon and etalon_doc is not None:
        align_spacing_to_etalon(doc, etalon_doc, profile)
        ensure_chapter_header_spacing(doc)
        if validate_empty_lines_in_body(doc):
            remove_extra_empty_lines_in_body(doc)
        apply_body_paragraph_format(doc, profile)
        apply_body_single_line_spacing(doc, profile)
        apply_header_paragraph_format(doc)

    if use_etalon and etalon_doc is not None:
        apply_signatory_layout_from_etalon(doc, etalon_doc, profile)

    fix_last_pages_and_signatories(
        doc, profile, mode="natural", use_etalon_spacing=use_etalon, etalon_doc=etalon_doc
    )
    apply_signatory_line_spacing(doc, profile)
    apply_body_single_line_spacing(doc, profile)
    apply_header_paragraph_format(doc)
    apply_body_paragraph_format(doc, profile)
    ensure_chapter_header_spacing(doc)
    apply_body_paragraph_format(doc, profile)

    strategy = "natural"
    if signatories_appear_orphaned(doc, profile):
        strategy = fix_last_pages_page_breaks(doc, profile, force=False)
        try:
            razrab_idx = find_razrabotal_index(doc)
            set_page_break_before(doc.paragraphs[razrab_idx], False)
        except ValueError:
            pass
        if use_etalon and etalon_doc is not None:
            apply_signatory_layout_from_etalon(doc, etalon_doc, profile)
        else:
            apply_signatory_line_spacing(doc, profile)
        try:
            ensure_single_empty_line_before(doc, "Разработал:")
            ensure_single_empty_line_before(doc, "Согласовано:")
        except ValueError:
            pass
    ensure_acquaintance_sheet_separate_page(doc)
    place_title_city_year_at_bottom(doc)

    _fix_duplicate_numbering_if_needed(doc)
    if profile.has_di_satp_numbering:
        if validate_di_satp_section_header_numbers(doc, profile):
            fix_missing_section_headers(doc)
        if validate_numbering_blocks(doc, profile):
            fix_numbering_selective(doc, profile)
        materialize_chapter_one_numbering(doc, profile)
    strip_unnecessary_characters(doc)
    remove_word_list_numbering_in_body(doc)
    remove_list_markers_in_body(doc)
    if use_etalon and etalon_doc is not None:
        apply_signatory_layout_from_etalon(doc, etalon_doc, profile)
        apply_signatory_line_spacing(doc, profile)
    normalize_body_paragraph_styles(doc)
    apply_header_paragraph_format(doc)
    apply_body_paragraph_format(doc, profile)
    apply_body_single_line_spacing(doc, profile)
    apply_table_paragraph_no_indent(doc)
    keep_table_header_rows_together(doc)
    apply_signatory_tab_stops(doc, profile)
    if materialize_signatory_paragraphs_from_tables(doc):
        profile.has_signatories = True
        apply_signatory_line_spacing(doc, profile)
        apply_signatory_tab_stops(doc, profile)
    strip_visual_highlights(doc)
    normalize_document_fonts(doc)
    ensure_title_page_separated(doc)
    format_title_block(doc)
    place_title_city_year_at_bottom(doc)
    center_chapter_headers(doc)
    apply_header_paragraph_format(doc)
    apply_body_paragraph_format(doc, profile)
    ensure_acquaintance_sheet_separate_page(doc)
    normalize_item_number_spacing(doc)
    collapse_adjacent_duplicate_words_in_document(doc)
    apply_signatory_fio_one_line(doc)
    fix_signatory_date_plaques(doc)
    format_title_block(doc)
    place_title_city_year_at_bottom(doc)
    prevent_chapter_header_orphan(doc)
    ensure_acquaintance_sheet_separate_page(doc)
    strip_visual_highlights(doc)
    mark_acquaintance_heading_if_job_mismatch(doc, source_path)
    apply_mandatory_layout_fixes(doc, profile)
    # Word: красная черта (орфография) и грамматика — после save:
    # apply_word_grammar_via_com / run_word_grammar_check_subprocess

    return strategy


def process_document(
    doc: Document,
    *,
    apply_page_breaks: bool = False,
    path: Path | None = None,
    etalon_path: Path | None = None,
) -> str:
    """API для любого документа; path нужен для detect_profile."""
    profile = detect_profile(doc, path or Path("_unknown_.docx"))
    if etalon_path is None and path is not None:
        etalon_path = find_etalon_path(path)
    return process_sniot_document(
        doc,
        profile,
        apply_page_breaks=apply_page_breaks,
        etalon_path=etalon_path,
        source_path=path,
    )


def export_debug(doc: Document, path: Path) -> None:
    lines = []
    for i, paragraph in enumerate(doc.paragraphs):
        fonts = {run.font.name for run in paragraph.runs if run.text.strip()}
        lines.append(f"{i+1:4d} font={fonts} | {paragraph.text}")
    path.write_text("\n".join(lines), encoding="utf-8")


WD_ACTIVE_END_PAGE_NUMBER = 3
WD_STATISTIC_PAGES = 2
WD_FIND_STOP = 0
WD_WITHIN_TABLE = 12
WORD_LAYOUT_TIMEOUT_SEC = 20
WORD_GRAMMAR_TIMEOUT_SEC = 25
WD_RUSSIAN = 1049
WORD_LIST_STRING_NUM = re.compile(r"^\d+(?:\.\d+)*\.?$")


def interpret_word_signatory_layout(
    *,
    body_page: int | None,
    signatory_page: int | None,
    total_pages: int | None = None,
    acquaintance_page: int | None = None,
) -> dict:
    """
    Сравнить страницу хвоста текста и страницу «Разработал:».
    orphaned=True — подписанты начались на следующей странице без текста.
    Подписанты на предпоследней вместе с текстом — не отрыв; лист ознакомления — отдельно.
    """
    result = {
        "orphaned": False,
        "body_page": body_page,
        "signatory_page": signatory_page,
        "total_pages": total_pages,
        "acquaintance_page": acquaintance_page,
        "message": "",
    }
    if signatory_page is None:
        result["message"] = (
            "Word: не найден абзац «Разработал:» — страницу подписантов проверить не удалось"
        )
        return result
    if body_page is None:
        result["message"] = (
            f"Word: «Разработал:» на стр. {signatory_page}, хвост текста не найден"
        )
        return result
    pages = f"текст стр. {body_page}, «Разработал:» стр. {signatory_page}"
    if total_pages:
        pages += f", всего стр. {total_pages}"
    if acquaintance_page:
        pages += f", лист ознакомления стр. {acquaintance_page}"
    if signatory_page > body_page:
        result["orphaned"] = True
        result["message"] = f"Word: подписанты оторваны от текста ({pages})"
        return result
    result["message"] = f"Word: подписанты вместе с текстом ({pages})"
    return result


def format_word_layout_action(report: dict) -> str:
    msg = (report or {}).get("message") or "Word: проверка страниц не выполнена"
    if not msg.startswith("Word:"):
        msg = f"Word: {msg}"
    extras: list[str] = []
    raz = (report or {}).get("razrabotal_count")
    if raz is not None:
        extras.append(f"«Разработал:» видно {raz} раз")
    double_n = int((report or {}).get("double_numbering_count") or 0)
    if double_n:
        extras.append(f"двойная нумерация {double_n}")
    if extras:
        msg = f"{msg}; {'; '.join(extras)}"
    return f"СНиОТ: {msg}"


def _word_suggestion_texts(suggestions) -> list[str]:
    texts: list[str] = []
    try:
        count = int(suggestions.Count)
    except Exception:
        return texts
    for i in range(1, count + 1):
        try:
            item = suggestions(i)
            name = getattr(item, "Name", None) or str(item)
            if name:
                texts.append(str(name))
        except Exception:
            continue
    return texts


def _mark_word_abbreviations_no_proofing(wdoc) -> int:
    """Аббревиатуры не проверять красной чертой (NoProofing)."""
    marked = 0
    try:
        words = wdoc.Range().Words
        total = int(words.Count)
    except Exception:
        return 0
    for i in range(1, total + 1):
        try:
            wrd = words(i)
            raw = str(wrd.Text or "")
        except Exception:
            continue
        if not spelling_error_is_abbreviation(raw):
            continue
        try:
            wrd.NoProofing = True
            marked += 1
        except Exception:
            continue
    return marked


def _fix_word_spelling_errors(word, wdoc) -> tuple[int, int, int]:
    """
    Красная волнистая черта = орфография. Заменить по GetSpellingSuggestions,
    кроме аббревиатур. Не IgnoreAll на весь документ.
    """
    fixed = 0
    skipped_abbrev = 0
    remaining = 0
    try:
        errors = wdoc.SpellingErrors
        count = int(errors.Count)
    except Exception:
        return 0, 0, 0
    for i in range(count, 0, -1):
        try:
            rng = errors(i)
            original = str(rng.Text or "")
        except Exception:
            continue
        if spelling_error_is_abbreviation(original):
            skipped_abbrev += 1
            try:
                rng.NoProofing = True
            except Exception:
                pass
            continue
        suggestions: list[str] = []
        try:
            suggestions = _word_suggestion_texts(rng.GetSpellingSuggestions())
        except Exception:
            try:
                token = original.strip()
                if token:
                    suggestions = _word_suggestion_texts(
                        word.GetSpellingSuggestions(token)
                    )
            except Exception:
                suggestions = []
        replacement = pick_spelling_suggestion(original, suggestions)
        if not replacement:
            remaining += 1
            continue
        try:
            trailing = ""
            if original.endswith(" ") and not replacement.endswith(" "):
                trailing = " "
            rng.Text = replacement + trailing
            fixed += 1
        except Exception:
            remaining += 1
    try:
        remaining = max(remaining, int(wdoc.SpellingErrors.Count) - skipped_abbrev)
    except Exception:
        pass
    return fixed, skipped_abbrev, remaining


def apply_word_grammar_via_com(docx_path: Path) -> dict:
    """
    Тихо открыть копию в Word: исправить красные подчёркивания (орфография,
    SpellingErrors + GetSpellingSuggestions), прогнать грамматику (CheckGrammar).
    Аббревиатуры не менять. Файл на диске сохраняется.
    """
    report: dict = {
        "ok": False,
        "available": False,
        "applied": False,
        "fixed": 0,
        "skipped_abbrev": 0,
        "remaining": 0,
        "grammar_remaining": 0,
        "message": "",
        "error": None,
    }
    path = Path(docx_path)
    if not path.is_file():
        report["message"] = "Word: файл не найден — орфографию не правили"
        report["error"] = "not_found"
        return report
    if pythoncom is None or win32com is None:
        report["message"] = "Word: pywin32 не установлен — красные подчёркивания не исправлены"
        report["error"] = "no_pywin32"
        return report

    pythoncom.CoInitialize()
    word = None
    wdoc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            word.Options.CheckGrammarAsYouType = True
            word.Options.CheckGrammarWithSpelling = False
            word.Options.CheckSpellingAsYouType = True
        except Exception:
            pass
        wdoc = word.Documents.Open(
            str(path.resolve()),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
        )
        try:
            wdoc.Content.LanguageID = WD_RUSSIAN
        except Exception:
            pass
        _mark_word_abbreviations_no_proofing(wdoc)
        try:
            wdoc.SpellingChecked = False
            wdoc.GrammarChecked = False
        except Exception:
            pass
        fixed, skipped, remaining = _fix_word_spelling_errors(word, wdoc)
        grammar_remaining = 0
        try:
            word.DisplayAlerts = 0
            wdoc.CheckGrammar()
        except Exception:
            pass
        try:
            grammar_remaining = int(wdoc.GrammaticalErrors.Count)
        except Exception:
            grammar_remaining = 0
        wdoc.Save()
        report.update(
            {
                "ok": True,
                "available": True,
                "applied": True,
                "fixed": fixed,
                "skipped_abbrev": skipped,
                "remaining": remaining,
                "grammar_remaining": grammar_remaining,
                "message": (
                    f"Word: орфография — исправлено {fixed} "
                    f"(красная черта), пропущено аббревиатур {skipped}, "
                    f"осталось {remaining}; грамматика замечаний {grammar_remaining}"
                ),
            }
        )
        return report
    except Exception as exc:
        err = str(exc).strip() or type(exc).__name__
        report["error"] = err
        low = err.lower()
        if "busy" in low or "заблокир" in low or "rpc" in low:
            report["message"] = (
                "Word: не удалось проверить орфографию — закройте документ в Word и повторите"
            )
        else:
            report["message"] = f"Word: ошибка орфографии/грамматики — {err[:180]}"
        return report
    finally:
        try:
            if wdoc is not None:
                wdoc.Close(SaveChanges=False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def format_word_proofing_action(report: dict) -> str:
    msg = (report or {}).get("message") or "Word: орфография не проверена"
    if not msg.startswith("Word:"):
        msg = f"Word: {msg}"
    return f"СНиОТ: {msg}"


def _kill_pid_tree(pid: int) -> None:
    """Windows: убить процесс и потомков (скрытый Word COM тоже)."""
    if not pid or pid <= 0:
        return
    try:
        subprocess.run(
            ["taskkill", "/PID", str(pid), "/T", "/F"],
            capture_output=True,
            timeout=8,
            check=False,
        )
    except Exception:
        pass


def _kill_invisible_word_processes() -> None:
    """Скрытый Word без окна — остаток COM после таймаута. Окна пользователя не трогать."""
    try:
        subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-Command",
                "Get-Process WINWORD -ErrorAction SilentlyContinue |"
                " Where-Object { -not $_.MainWindowTitle } | Stop-Process -Force",
            ],
            capture_output=True,
            timeout=8,
            check=False,
        )
    except Exception:
        pass


def _run_timed_subprocess(cmd: list[str], timeout_sec: int):
    """
    Запуск с жёстким таймаутом. По истечении — дерево процессов и скрытый Word.
    None = таймаут (не ждать минутами).
    """
    creationflags = 0
    if sys.platform == "win32":
        creationflags = int(getattr(subprocess, "CREATE_NEW_PROCESS_GROUP", 0)) | int(
            getattr(subprocess, "CREATE_NO_WINDOW", 0)
        )
    proc = subprocess.Popen(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        errors="replace",
        creationflags=creationflags,
    )
    try:
        stdout, stderr = proc.communicate(timeout=timeout_sec)
        return subprocess.CompletedProcess(cmd, proc.returncode or 0, stdout, stderr)
    except subprocess.TimeoutExpired:
        _kill_pid_tree(proc.pid)
        try:
            proc.kill()
        except Exception:
            pass
        try:
            proc.communicate(timeout=5)
        except Exception:
            pass
        _kill_invisible_word_processes()
        return None


def skip_word_com_requested() -> bool:
    return os.environ.get("SNIOT_SKIP_WORD", "").strip().lower() in {"1", "true", "yes"}


def run_word_grammar_check_subprocess(docx_path: Path) -> dict:
    """Орфография (красная черта) + грамматика в отдельном процессе Word.

    Дочерний процесс вызывает apply_word_grammar_via_com (SpellingErrors +
    GetSpellingSuggestions), кроме аббревиатур. Таймаут — не висеть.
    """
    if skip_word_com_requested():
        return {
            "ok": False,
            "available": False,
            "applied": False,
            "message": "Word: орфография пропущена (SNIOT_SKIP_WORD)",
            "error": "skipped",
        }
    script = Path(__file__).resolve()
    cmd = [
        sys.executable,
        str(script),
        "--word-grammar-check",
        str(Path(docx_path).resolve()),
    ]
    proc = _run_timed_subprocess(cmd, WORD_GRAMMAR_TIMEOUT_SEC)
    if proc is None:
        return {
            "ok": False,
            "available": False,
            "applied": False,
            "message": (
                f"Word: проверка орфографии прервана (больше {WORD_GRAMMAR_TIMEOUT_SEC} с) "
                "— закройте документ в Word, если он открыт"
            ),
            "error": "timeout",
        }
    raw = (proc.stdout or "").strip()
    if not raw:
        err = (proc.stderr or "").strip()[:180]
        return {
            "ok": False,
            "available": False,
            "applied": False,
            "message": f"Word: нет ответа орфографии{(' — ' + err) if err else ''}",
            "error": "empty",
        }
    try:
        data = json.loads(raw.splitlines()[-1])
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass
    return {
        "ok": False,
        "available": False,
        "applied": False,
        "message": f"Word: неразобранный ответ орфографии — {raw[:160]}",
        "error": "bad_json",
    }


def _collect_word_visible_extras(wdoc) -> dict:
    """То, что человек видит в Word: автономер списка, дубль «Разработал:»."""
    extras = {
        "razrabotal_count": 0,
        "double_numbering_count": 0,
        "list_fixes": [],
    }
    for para in wdoc.Paragraphs:
        rng = para.Range
        txt = (rng.Text or "").replace("\r", "").replace("\x07", "").strip()
        if not txt:
            continue
        try:
            lst = str(rng.ListFormat.ListString or "").strip()
        except Exception:
            lst = ""
        if txt.upper().startswith("РАЗРАБОТАЛ"):
            extras["razrabotal_count"] += 1
        if not lst or not WORD_LIST_STRING_NUM.match(lst):
            continue
        if paragraph_has_manual_number(txt):
            extras["double_numbering_count"] += 1
            extras["list_fixes"].append(
                {"text": txt[:120], "list_string": lst, "already_in_text": True}
            )
        elif len(extras["list_fixes"]) < 80:
            extras["list_fixes"].append(
                {"text": txt[:120], "list_string": lst, "already_in_text": False}
            )
    return extras


def apply_word_list_fixes(doc: Document, fixes: list[dict]) -> int:
    """
    Перенести видимый номер Word (ListString) в текст и снять numPr —
    иначе на экране «1.1. 1.1. …» или номер есть, а отступ как у списка.
    """
    if not fixes:
        return 0
    changed = 0
    for item in fixes:
        needle = (item.get("text") or "").strip()
        list_s = (item.get("list_string") or "").strip()
        if not needle or len(needle) < 8 or not list_s:
            continue
        if not WORD_LIST_STRING_NUM.match(list_s):
            continue
        num = list_s.rstrip(".")
        key = strip_number(needle)[:40]
        if not key:
            key = needle[:40]
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if key not in text and not text.startswith(needle[:40]) and needle[:40] not in text:
                continue
            already = bool(item.get("already_in_text")) or paragraph_has_manual_number(text)
            if not already:
                set_paragraph_text(paragraph, apply_number(strip_number(text), num))
                changed += 1
            if has_word_list_numbering(paragraph):
                if remove_word_list_numbering(paragraph):
                    changed += 1
            break
    return changed


def apply_word_visible_repairs(
    out_path: Path,
    *,
    profile: DocumentProfile,
    word_report: dict,
) -> int:
    """После снимка Word: номера из ListString в текст, снять numPr, убрать дубль «Разработал:»."""
    doc = Document(out_path)
    changed = apply_word_list_fixes(doc, word_report.get("list_fixes") or [])
    changed += deduplicate_razrabotal_markers(doc)
    if materialize_signatory_paragraphs_from_tables(doc):
        profile.has_signatories = True
        changed += 1
    changed += strip_visual_highlights(doc)
    if not changed:
        return 0
    apply_body_paragraph_format(doc, profile)
    apply_header_paragraph_format(doc)
    if profile.has_signatories:
        apply_signatory_line_spacing(doc, profile)
        apply_signatory_tab_stops(doc, profile)
    strip_visual_highlights(doc)
    doc.save(out_path)
    return changed


def _word_find_page(wdoc, needle: str) -> int | None:
    text = (needle or "").replace("\t", " ").replace("\xa0", " ").strip()
    if not text:
        return None
    text = text[:80]
    rng = wdoc.Content.Duplicate
    finder = rng.Find
    finder.ClearFormatting()
    found = finder.Execute(
        FindText=text,
        MatchCase=False,
        MatchWholeWord=False,
        MatchWildcards=False,
        Forward=True,
        Wrap=WD_FIND_STOP,
    )
    if not found:
        return None
    return int(rng.Information(WD_ACTIVE_END_PAGE_NUMBER))


def inspect_signatory_pages_via_word(docx_path: Path) -> dict:
    """
    Тихо открыть docx в Microsoft Word и вернуть номера страниц хвоста и «Разработал:».
    Не сохраняет файл. Не показывает окно Word.
    """
    report: dict = {
        "ok": False,
        "available": False,
        "orphaned": False,
        "body_page": None,
        "signatory_page": None,
        "total_pages": None,
        "message": "",
        "error": None,
    }
    path = Path(docx_path)
    if not path.is_file():
        report["message"] = f"Word: файл не найден — {path}"
        report["error"] = "not_found"
        return report

    if pythoncom is None or win32com is None:
        report["message"] = "Word: pywin32 не установлен — страницы не проверены"
        report["error"] = "no_pywin32"
        return report

    doc = Document(path)
    try:
        razrab_idx = find_razrabotal_index(doc)
        razrab_needle = (
            paragraph_text_normalized(doc.paragraphs[razrab_idx])[:40] or "Разработал:"
        )
    except ValueError:
        razrab_needle = "Разработал:"
    last_body_idx = find_last_body_paragraph_before_signatories(doc)
    body_needle = (
        paragraph_text_normalized(doc.paragraphs[last_body_idx])[:50]
        if last_body_idx is not None
        else ""
    )
    acq_idx = find_acquaintance_sheet_start(doc)
    acq_needle = (
        paragraph_text_normalized(doc.paragraphs[acq_idx])[:60]
        if acq_idx is not None
        else ""
    )

    pythoncom.CoInitialize()
    word = None
    wdoc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdoc = word.Documents.Open(
            str(path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        wdoc.Repaginate()
        extras = _collect_word_visible_extras(wdoc)
        report.update(extras)
        total = int(wdoc.ComputeStatistics(WD_STATISTIC_PAGES))
        body_page = _word_find_page(wdoc, body_needle) if body_needle else None
        signatory_page = _word_find_page(wdoc, razrab_needle)
        acquaintance_page = _word_find_page(wdoc, acq_needle) if acq_needle else None
        interpreted = interpret_word_signatory_layout(
            body_page=body_page,
            signatory_page=signatory_page,
            total_pages=total,
            acquaintance_page=acquaintance_page,
        )
        report.update(interpreted)
        if extras["razrabotal_count"] > 1:
            report["message"] = (
                f"{report.get('message')}; Word: дубль «Разработал:» "
                f"({extras['razrabotal_count']} раза)"
            )
        if extras["double_numbering_count"]:
            report["message"] = (
                f"{report.get('message')}; Word: двойная нумерация "
                f"({extras['double_numbering_count']})"
            )
        report["ok"] = True
        report["available"] = True
        return report
    except Exception as exc:
        err = str(exc).strip() or type(exc).__name__
        report["error"] = err
        low = err.lower()
        if "busy" in low or "заблокир" in low or "rpc" in low:
            report["message"] = (
                "Word: не удалось открыть документ — закройте его в Word и повторите проверку страниц"
            )
        else:
            report["message"] = f"Word: ошибка проверки страниц — {err[:180]}"
        return report
    finally:
        try:
            if wdoc is not None:
                wdoc.Close(SaveChanges=False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def run_word_layout_check_subprocess(docx_path: Path) -> dict:
    """Проверка страниц в отдельном процессе — таймаут, чтобы GUI не завис."""
    if skip_word_com_requested():
        return {
            "ok": False,
            "available": False,
            "orphaned": False,
            "message": "Word: проверка страниц пропущена (SNIOT_SKIP_WORD)",
            "error": "skipped",
        }
    script = Path(__file__).resolve()
    cmd = [
        sys.executable,
        str(script),
        "--word-layout-check",
        str(Path(docx_path).resolve()),
    ]
    proc = _run_timed_subprocess(cmd, WORD_LAYOUT_TIMEOUT_SEC)
    if proc is None:
        return {
            "ok": False,
            "available": False,
            "orphaned": False,
            "message": (
                f"Word: проверка страниц прервана (больше {WORD_LAYOUT_TIMEOUT_SEC} с) "
                "— закройте документ в Word, если он открыт"
            ),
            "error": "timeout",
        }
    raw = (proc.stdout or "").strip()
    if not raw:
        err = (proc.stderr or "").strip()[:180]
        return {
            "ok": False,
            "available": False,
            "orphaned": False,
            "message": f"Word: нет ответа проверки страниц{(' — ' + err) if err else ''}",
            "error": "empty",
        }
    try:
        data = json.loads(raw.splitlines()[-1])
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass
    return {
        "ok": False,
        "available": False,
        "orphaned": False,
        "message": f"Word: неразобранный ответ проверки — {raw[:160]}",
        "error": "bad_json",
    }


def repair_orphaned_signatories_after_word(
    out_path: Path,
    *,
    profile: DocumentProfile,
    etalon_path: Path | None,
) -> str:
    """Шаг 2 переноса после того, как Word подтвердил отрыв подписантов."""
    doc = Document(out_path)
    etalon_doc = Document(etalon_path) if etalon_path and etalon_path.is_file() else None
    strategy = apply_signatory_page_break(doc, profile)
    try:
        razrab_idx = find_razrabotal_index(doc)
        set_page_break_before(doc.paragraphs[razrab_idx], False)
    except ValueError:
        pass
    if etalon_doc is not None:
        apply_signatory_layout_from_etalon(doc, etalon_doc, profile)
    apply_signatory_line_spacing(doc, profile)
    strip_visual_highlights(doc)
    doc.save(out_path)
    out_path.write_bytes(fix_page_numbering(out_path.read_bytes()))
    return strategy


def print_issues(title: str, issues: list[str]) -> None:
    if not issues:
        return
    print(title)
    for item in issues:
        try:
            print(f"  - {item}")
        except UnicodeEncodeError:
            print(f"  - {item.encode('utf-8', errors='replace').decode('utf-8', errors='replace')}")


def apply_sniot_rules_to_file(
    target: Path,
    *,
    fix_page_breaks: bool = False,
    always_apply: bool = True,
    skip_word: bool = False,
) -> dict:
    """
    Применить все правила СНиОТ к docx и сохранить на месте.

    Для DocAgent «Оформить документ»: финальный проход после type-specific форматтеров.
    always_apply=True — выполнять process_sniot_document даже если validate OK до правки
    (structure_fix мог оставить шрифт/нумерацию/колонтитулы не по правилам).
    """
    result: dict = {
        "ok": False,
        "applied": False,
        "before_issues": [],
        "after_issues": [],
        "strategy": "natural",
        "actions": [f"СНиОТ: сборка {SCRIPT_BUILD}"],
    }
    if not target.is_file():
        result["actions"].append(f"СНиОТ: файл не найден — {target}")
        return result
    try:
        assert_path_writable(target)
    except PermissionError as exc:
        result["actions"].append(str(exc))
        return result
    if target.suffix.lower() != ".docx":
        result["ok"] = True
        result["actions"].append("СНиОТ: пропуск (не .docx)")
        return result

    preview_doc = Document(target)
    profile = detect_profile(preview_doc, target)
    before_nonempty = count_nonempty_body_paragraphs(preview_doc)
    before_numbered = count_numbered_paragraphs(preview_doc)
    source_bytes = target.read_bytes()
    sibling = find_agent_sibling_source(target)
    source_tokens: tuple[str, ...] = ()
    if sibling is not None:
        use_sibling = title_or_identity_looks_corrupted(preview_doc, target)
        sibling_reason = "убрана чужая вставка на титуле / подмена текста"
        if not use_sibling:
            try:
                sibling_doc = open_word_document_readonly(sibling)
                sib_tokens = collect_number_tokens(sibling_doc)
                cur_tokens = collect_number_tokens(preview_doc)
                if len(sib_tokens) > len(cur_tokens):
                    use_sibling = True
                    sibling_reason = "восстановлена нумерация исходника"
                    source_tokens = tuple(sib_tokens)
            except Exception:
                pass
        if use_sibling:
            source_bytes = read_docx_bytes_any(sibling)
            preview_doc = Document(BytesIO(source_bytes))
            profile = detect_profile(preview_doc, target)
            if not source_tokens:
                source_tokens = tuple(collect_number_tokens(preview_doc))
            profile.source_number_tokens = source_tokens
            before_nonempty = count_nonempty_body_paragraphs(preview_doc)
            before_numbered = count_numbered_paragraphs(preview_doc)
            result["actions"].append(
                f"СНиОТ: исходник «{sibling.name}» — {sibling_reason}"
            )
    cleaned_bytes = remove_duplicate_body_title(source_bytes, profile.first_chapter)

    WORK_DIR.mkdir(parents=True, exist_ok=True)
    safe_stem = re.sub(r"[^\w\-]+", "_", target.stem)[:48] or "doc"
    work_path = WORK_DIR / f"_docagent_sniot_{safe_stem}.docx"
    out_path = WORK_DIR / f"_docagent_sniot_fixed_{safe_stem}.docx"
    work_path.write_bytes(cleaned_bytes)

    restored_n, restored_msg = maybe_restore_senior_master_body(work_path)
    if restored_n:
        result["actions"].append(restored_msg)

    doc = Document(work_path)
    body_after_clean = count_nonempty_body_paragraphs(doc)
    if before_nonempty >= 5 and body_after_clean == 0:
        cleaned_bytes = target.read_bytes()
        work_path.write_bytes(cleaned_bytes)
        doc = Document(work_path)
        result["actions"].append(
            "СНиОТ: удаление дубля титула пропущено — иначе было бы стёрто всё тело документа"
        )
    profile = detect_profile(doc, target)
    if source_tokens:
        profile.source_number_tokens = source_tokens
    before_issues = validate_sniot_document(
        doc, docx_bytes=cleaned_bytes, profile=profile, path=target
    )
    result["before_issues"] = before_issues

    if not always_apply and not before_issues:
        result["ok"] = True
        result["actions"].append("СНиОТ: уже соответствует правилам")
        return result

    etalon_path = find_etalon_path(target)
    if etalon_path:
        result["actions"].append(f"СНиОТ: образец — {etalon_path.name}")
    if profile.has_signatories:
        fix_page_breaks = True

    strategy = process_sniot_document(
        doc,
        profile,
        apply_page_breaks=fix_page_breaks,
        etalon_path=etalon_path,
        source_path=target,
    )
    result["strategy"] = strategy
    apply_mandatory_layout_fixes(doc, profile)

    doc.save(out_path)
    fixed_bytes = fix_page_numbering(out_path.read_bytes())
    out_path.write_bytes(fixed_bytes)

    # Пять правок XML уже в out_path. Сначала запись в Агент — потом Word.
    # Иначе таймаут Word (180 с) убивает процесс, и пользователь видит старый файл.
    skip_word = bool(skip_word) or skip_word_com_requested()
    after_issues = validate_full_document_on_disk(out_path, path_for_profile=target)
    doc, _ = reload_document_from_path(out_path)
    profile = detect_profile(doc, target)
    integrity = validate_save_integrity(
        before_nonempty=before_nonempty,
        before_numbered=before_numbered,
        after_doc=doc,
        profile=profile,
    )
    if integrity:
        result["after_issues"] = integrity + after_issues
        result["ok"] = False
        result["applied"] = False
        result["actions"].append(
            "СНиОТ: запись в папку Агент отменена — документ стал хуже (integrity)"
        )
        for item in integrity:
            result["actions"].append(f"СНиОТ ⛔ {item}")
        result["actions"].append(f"СНиОТ: черновик правки — {out_path}")
        return result

    layout_issues = collect_five_layout_issues(doc, profile)
    result["after_issues"] = after_issues
    if layout_issues:
        result["ok"] = False
        result["applied"] = False
        result["actions"].append(
            "СНиОТ: запись в папку Агент отменена — пять обязательных правок не прошли"
        )
        result["actions"].append(
            f"СНиОТ: финальная перечитка с диска — {format_validation_summary(after_issues)}"
        )
        for item in layout_issues[:12]:
            result["actions"].append(f"СНиОТ ! {item}")
        result["actions"].append(f"СНиОТ: черновик правки — {out_path}")
        return result

    try:
        assert_path_writable(target)
        copy_to_target_unlocking_word(out_path, target)
        result["applied"] = True
        result["actions"].append(f"Сохранено: {target}")
        result["actions"].append(
            "СНиОТ: пять правок XML — пробелы, номера, титул, сироты глав, даты подписантов"
        )
        if before_issues:
            result["actions"].append(
                f"СНиОТ (fix_sniot_document): исправлено замечаний — {len(before_issues)}"
            )
        numbering_note = (
            ", нумерация 1.8.x / 1.9.x"
            if profile.has_di_satp_numbering
            else ""
        )
        result["actions"].append(
            "СНиОТ: дубль титула, TNR 14, поля А4, одинарный интервал тела, номера стр., "
            f"главы, тело 1,25 см, подписанты 1,5{numbering_note}"
        )
        result["actions"].append(f"СНиОТ: перенос страниц — {strategy}")
        result["actions"].append(
            f"СНиОТ: абзацев в теле — {count_nonempty_body_paragraphs(doc)}"
        )
        if after_issues:
            result["ok"] = False
            result["actions"].append(
                "СНиОТ: файл записан (пять правок сделаны), остались другие замечания"
            )
            result["actions"].append(
                f"СНиОТ: финальная перечитка с диска — {format_validation_summary(after_issues)}"
            )
            for item in after_issues[:12]:
                result["actions"].append(f"СНиОТ ! {item}")
            if len(after_issues) > 12:
                result["actions"].append(f"СНиОТ ! … и ещё {len(after_issues) - 12}")
        else:
            result["ok"] = True
            result["actions"].append("СНиОТ: Validation OK (0 issues)")
            result["actions"].append("СНиОТ: финальная перечитка с диска — 0 замечаний (все правила)")
    except PermissionError:
        result["ok"] = False
        result["applied"] = False
        result["actions"].append(
            "СНиОТ: не удалось сохранить — закройте «_оформлен.docx» в Word и повторите"
        )
        result["actions"].append(f"СНиОТ: черновик правки — {out_path}")
        return result

    if skip_word:
        result["actions"].append(
            "СНиОТ: Word COM в XML-проходе пропущен (файл уже записан; "
            "орфография и грамматика Word — отдельный вызов после записи)"
        )
        return result

    grammar_report = run_word_grammar_check_subprocess(out_path)
    result["actions"].append(format_word_proofing_action(grammar_report))
    word_usable = grammar_report.get("error") not in {"timeout", "skipped"}
    if not word_usable:
        result["actions"].append(
            "СНиОТ: Word не ответил вовремя — страницы не проверяли, пять правок XML уже сохранены"
        )
        return result
    if grammar_report.get("applied"):
        gdoc = Document(out_path)
        restore_lsim_if_osim_poisoned(gdoc, target)
        gdoc.save(out_path)
        out_path.write_bytes(fix_page_numbering(out_path.read_bytes()))
        gdoc, _ = reload_document_from_path(out_path)
        gprofile = detect_profile(gdoc, target)
        if not collect_five_layout_issues(gdoc, gprofile):
            try:
                copy_file_if_different(out_path, target)
            except PermissionError:
                result["actions"].append(
                    "СНиОТ: орфография Word готова, но файл в Агент открыт — закройте Word"
                )

    if profile.has_signatories or profile.has_di_satp_numbering:
        word_report = run_word_layout_check_subprocess(out_path)
        result["actions"].append(format_word_layout_action(word_report))
        if word_report.get("error") in {"timeout", "skipped"}:
            result["actions"].append(
                "СНиОТ: Word страницы не ответил — пять правок XML уже сохранены"
            )
            return result
        n_vis = apply_word_visible_repairs(
            out_path, profile=profile, word_report=word_report
        )
        if n_vis:
            result["actions"].append(
                f"СНиОТ: выровнено по тому, что видно в Word — {n_vis}"
            )
            out_path.write_bytes(fix_page_numbering(out_path.read_bytes()))
        if word_report.get("orphaned"):
            strategy = repair_orphaned_signatories_after_word(
                out_path, profile=profile, etalon_path=etalon_path
            )
            result["strategy"] = strategy
            result["actions"].append(
                f"СНиОТ: Word подтвердил отрыв — выполнен перенос страниц ({strategy})"
            )
        try:
            copy_file_if_different(out_path, target)
        except PermissionError:
            pass

    return result

def autofix(
    target: Path,
    *,
    check_only: bool = False,
    dry_run: bool = False,
    fix_page_breaks: bool = False,
    always_apply: bool = False,
    skip_word: bool = False,
) -> int:
    """Полный автоматический цикл. Возвращает код выхода."""
    configure_stdio_utf8()
    print(f"=== СНиОТ: {target.name} ({detect_document_kind(target)}) ===", flush=True)
    print(f"СНиОТ: сборка {SCRIPT_BUILD}", flush=True)

    preview_doc = Document(target)
    profile = detect_profile(preview_doc, target)
    cleaned_bytes = remove_duplicate_body_title(target.read_bytes(), profile.first_chapter)
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    WORK_FILE.write_bytes(cleaned_bytes)

    doc = Document(WORK_FILE)
    profile = detect_profile(doc, target)
    before_issues = validate_sniot_document(doc, docx_bytes=cleaned_bytes, profile=profile, path=target)

    if check_only:
        print_issues("Проблемы:", before_issues)
        print("OK" if not before_issues else "Есть замечания")
        return EXIT_OK if not before_issues else EXIT_VALIDATION_FAIL

    if skip_word:
        os.environ["SNIOT_SKIP_WORD"] = "1"

    if always_apply:
        rep = apply_sniot_rules_to_file(
            target,
            fix_page_breaks=fix_page_breaks,
            always_apply=True,
            skip_word=skip_word,
        )
        before_n = len(rep.get("before_issues") or [])
        print(f"СНиОТ: замечаний до правки — {before_n}")
        print(f"Стратегия переноса: {rep.get('strategy', 'natural')}")
        print_issues("Осталось:", rep.get("after_issues") or [])
        if rep.get("applied"):
            print("Validation: OK" if not rep.get("after_issues") else "Есть замечания после правки")
        for act in rep.get("actions") or []:
            try:
                print(act)
            except UnicodeEncodeError:
                print(act.encode("utf-8", errors="replace").decode("utf-8", errors="replace"))
        if not rep.get("ok"):
            return EXIT_VALIDATION_FAIL if rep.get("after_issues") else EXIT_FILE_LOCKED
        return EXIT_OK

    if not before_issues:
        print("Документ уже соответствует правилам. Правка не требуется.")
        return EXIT_OK

    print_issues("Исправляю:", before_issues)
    etalon_path = find_etalon_path(target)
    if etalon_path:
        print(f"Образец: {etalon_path.name}")
    strategy = process_sniot_document(
        doc,
        profile,
        apply_page_breaks=fix_page_breaks,
        etalon_path=etalon_path,
        source_path=target,
    )

    doc.save(OUT_FILE)
    fixed_bytes = fix_page_numbering(OUT_FILE.read_bytes())
    OUT_FILE.write_bytes(fixed_bytes)
    after_issues = validate_full_document_on_disk(OUT_FILE, path_for_profile=target)
    doc, _ = reload_document_from_path(OUT_FILE)
    profile = detect_profile(doc, target)
    export_debug(doc, OUT_FILE.with_suffix(".txt"))

    print(f"Стратегия переноса: {strategy}")
    print_issues("Осталось:", after_issues)

    if after_issues:
        print("ОШИБКА: после правки остались замечания — запись в папку Агент отменена.")
        print(f"Черновик правки: {OUT_FILE}")
        return EXIT_VALIDATION_FAIL

    print("Validation: OK")

    if dry_run:
        print(f"Dry-run: результат в {OUT_FILE}")
        return EXIT_OK

    backup = target.with_name(
        target.stem + f"_backup_{datetime.now().strftime('%Y%m%d_%H%M')}" + target.suffix
    )
    try:
        assert_path_writable(target)
        assert_path_writable(backup)
        copy_file_if_different(target, backup)
        copy_file_if_different(OUT_FILE, target)
        final_issues = validate_full_document_on_disk(target, path_for_profile=target)
        if final_issues:
            print("ОШИБКА: повторная проверка файла в папке Агент не прошла.")
            print_issues("Замечания:", final_issues)
            return EXIT_VALIDATION_FAIL
    except PermissionError:
        print("ОШИБКА: закройте файл в Word и запустите снова.")
        print(f"Готовый результат сохранён локально: {OUT_FILE}")
        return EXIT_FILE_LOCKED

    print(f"Сохранено: {target}")
    print(f"Бэкап: {backup.name}")
    return EXIT_OK


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Автоисправление документов СНиОТ — ДИ, РИ, Положения и др.",
        epilog="Правила: sniot-di-documents.mdc. Образец docx не важнее правил.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--target", type=Path, default=None, help="Необязательно: другой docx")
    parser.add_argument(
        "--handoff",
        action="store_true",
        help="Взять путь из DocAgent/handoff/request_latest.json",
    )
    parser.add_argument("--check", action="store_true", help="Только проверка (validate_sniot_document)")
    parser.add_argument("--dry-run", action="store_true", help="Без записи на N:\\")
    parser.add_argument(
        "--apply",
        action="store_true",
        help="Явная запись на N:\\ (по умолчанию без --dry-run запись выполняется)",
    )
    parser.add_argument(
        "--fix-page-breaks",
        action="store_true",
        help="Шаг 2: перенос страниц, если подписанты оторваны",
    )
    parser.add_argument(
        "--always-apply",
        action="store_true",
        help="Всегда выполнять process_sniot_document (для DocAgent «Оформить документ»)",
    )
    parser.add_argument(
        "--skip-word",
        action="store_true",
        help="Не вызывать Word COM (если Word занят или завис)",
    )
    parser.add_argument(
        "--show-rules",
        action="store_true",
        help="Показать сводку правил и выйти",
    )
    parser.add_argument(
        "--restore-from-dump",
        type=Path,
        default=None,
        help="Восстановить тело из текстового дампа export_debug (_work_*.txt)",
    )
    parser.add_argument(
        "--word-layout-check",
        type=Path,
        default=None,
        help="Тихая проверка страниц в Word (JSON в stdout), без правки файла",
    )
    parser.add_argument(
        "--word-grammar-check",
        type=Path,
        default=None,
        help="Тихая орфография Word (красные подчёркивания) + грамматика; сохраняет файл",
    )
    args = parser.parse_args()

    if args.show_rules:
        print(RULES.strip())
        sys.exit(EXIT_OK)

    if args.word_layout_check:
        configure_stdio_utf8()
        report = inspect_signatory_pages_via_word(args.word_layout_check)
        print(json.dumps(report, ensure_ascii=False))
        sys.exit(EXIT_OK if report.get("ok") else EXIT_VALIDATION_FAIL)

    if args.word_grammar_check:
        configure_stdio_utf8()
        report = apply_word_grammar_via_com(args.word_grammar_check)
        print(json.dumps(report, ensure_ascii=False))
        sys.exit(EXIT_OK if report.get("ok") else EXIT_VALIDATION_FAIL)

    try:
        target = resolve_target(args.target, use_handoff=args.handoff)
    except FileNotFoundError as exc:
        print(f"ОШИБКА: {exc}")
        sys.exit(EXIT_NOT_FOUND)

    if args.restore_from_dump:
        dump = args.restore_from_dump.expanduser()
        if not dump.is_file():
            dump = Path(__file__).resolve().parent / dump.name
        stamp = datetime.now().strftime("%Y%m%d_%H%M")
        backup = target.with_name(f"{target.stem}_backup_{stamp}{target.suffix}")
        try:
            assert_path_writable(target)
            assert_path_writable(backup)
            copy_file_if_different(target, backup)
            n = replace_body_from_debug_dump(target, dump)
            print(f"Восстановлено {n} абзацев из {dump.name}")
            print(f"Бэкап: {backup}")
        except Exception as exc:
            print(f"ОШИБКА восстановления: {exc}")
            sys.exit(EXIT_VALIDATION_FAIL)
        code = autofix(
            target,
            check_only=False,
            dry_run=False,
            fix_page_breaks=args.fix_page_breaks,
            always_apply=True,
        )
        sys.exit(code)

    # ДИ «Старший мастер»: авто-восстановление из дампа, если нумерация пропала
    if is_senior_master_di_path(target) and not args.check:
        restored_n, restored_msg = maybe_restore_senior_master_body(target)
        if restored_n:
            print(restored_msg)

    dry_run = args.dry_run and not args.apply

    code = autofix(
        target,
        check_only=args.check,
        dry_run=dry_run,
        fix_page_breaks=args.fix_page_breaks,
        always_apply=args.always_apply,
        skip_word=args.skip_word,
    )
    sys.exit(code)


if __name__ == "__main__":
    configure_stdio_utf8()
    main()
