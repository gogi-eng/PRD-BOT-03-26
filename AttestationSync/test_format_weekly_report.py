# -*- coding: utf-8 -*-
"""Оформление еженедельного итога: пробелы и Инструкция по делопроизводству 2025."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

HERE = Path(__file__).resolve().parent
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))

import fix_sniot_document as sniot_fix
import format_weekly_report as weekly


def test_normalize_spaces_collapses_multi_space():
    raw = "один  два   три"
    cleaned = weekly.normalize_spaces_text(raw, full=True)
    assert cleaned == "один два три"
    assert "  " not in cleaned


def test_normalize_spaces_nbsp_and_tabs():
    raw = "один\u00a0\u00a0два\t\tтри"
    cleaned = weekly.normalize_spaces_text(raw, full=True)
    assert cleaned == "один два три"
    assert "\u00a0" not in cleaned
    assert "\t" not in cleaned


def test_validate_weekly_spaces_catches_double_space():
    doc = Document()
    doc.add_paragraph("один  два   три")
    issues = weekly.validate_weekly_spaces(doc)
    assert issues
    assert any("Двойные пробелы" in item for item in issues)


def test_strip_weekly_double_spaces_fixes_paragraph():
    doc = Document()
    p = doc.add_paragraph("один  два   три")
    p.runs[0].font.size = Pt(14)
    changed = weekly.strip_weekly_double_spaces(doc)
    assert changed == 1
    assert doc.paragraphs[0].text == "один два три"
    assert weekly.validate_weekly_spaces(doc) == []


def test_does_not_apply_di_satp_18_19_numbering():
    doc = Document()
    doc.add_paragraph("1.4. Назначение — ведущий инженер.")
    weekly.apply_weekly_office_to_document(doc)
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "1.4. Назначение" in body
    assert "1.8." not in body
    assert "1.9." not in body


def test_weekly_chapter_centered_caps():
    doc = Document()
    doc.add_paragraph("1. Количественные итоги")
    doc.add_paragraph("Текст тела  с  двумя пробелами.")
    weekly.apply_weekly_office_to_document(doc)
    header = doc.paragraphs[0].text
    assert header == "1. КОЛИЧЕСТВЕННЫЕ ИТОГИ"
    assert sniot_fix.is_paragraph_centered(doc.paragraphs[0])
    assert doc.paragraphs[1].text == "Текст тела с двумя пробелами."
    assert sniot_fix.is_paragraph_justified(doc.paragraphs[1])


def test_format_weekly_docx_margins_and_spaces(tmp_path: Path):
    path = tmp_path / "Отчёт_о_работе_тест.docx"
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    doc.add_paragraph("ГОСУДАРСТВЕННОЕ ПРЕДПРИЯТИЕ «МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("ОТЧЁТ")
    doc.add_paragraph("о выполненной работе за период")
    doc.add_paragraph("1. Количественные итоги")
    doc.add_paragraph("За неделю  проведено   мероприятий.")
    doc.save(path)

    result = weekly.format_weekly_docx(path, backup=True)
    assert result["ok"], result["issues"]
    assert result["backup"]
    assert Path(result["backup"]).is_file()

    out = Document(str(path))
    section = out.sections[0]
    assert abs(section.left_margin.mm - 30) < 1.0
    assert abs(section.right_margin.mm - 8) < 1.0
    assert abs(section.top_margin.mm - 20) < 1.0
    assert abs(section.bottom_margin.mm - 20) < 1.0
    texts = [p.text for p in out.paragraphs if p.text.strip()]
    assert any("КОЛИЧЕСТВЕННЫЕ ИТОГИ" in t for t in texts)
    assert any("За неделю проведено мероприятий." in t for t in texts)
    assert weekly.validate_weekly_spaces(out) == []
    page_issues = sniot_fix.validate_page_numbering(path.read_bytes())
    assert not any("header1" in i and "нет поля PAGE" in i for i in page_issues)


def test_assert_weekly_writable_blocks_agent_and_obmen(tmp_path: Path):
    # emulate forbidden names via path string check on ОБМЕН
    try:
        weekly.assert_weekly_writable(r"N:\9 - Служба\!!!ОБМЕН\файл.docx")
        assert False, "ОБМЕН должен быть запрещён"
    except PermissionError:
        pass
    try:
        weekly.assert_weekly_writable(str(weekly.USER_AGENT_DIR / "итог.docx"))
        assert False, "папка Агент должна быть запрещена"
    except PermissionError:
        pass
    # temp pytest path is allowed
    ok = weekly.assert_weekly_writable(tmp_path / "Отчёт_о_работе_тест.docx")
    assert ok.name.endswith(".docx")


def test_weekly_zeros_paragraph_spacing():
    doc = Document()
    body = doc.add_paragraph("За неделю проведено  мероприятий.")
    body.paragraph_format.space_after = Pt(12)
    weekly.apply_weekly_office_to_document(doc)
    assert "  " not in doc.paragraphs[0].text
    assert sniot_fix.paragraph_has_single_line_spacing(doc.paragraphs[0])
    assert sniot_fix.paragraph_has_zero_block_spacing(doc.paragraphs[0])
    issues = weekly.validate_weekly_document(doc)
    assert not any("интервал" in item.lower() for item in issues)


def test_weekly_strips_list_markers():
    doc = Document()
    doc.add_paragraph("1. Количественные итоги")
    doc.add_paragraph("• Проведено мероприятий.")
    doc.add_paragraph("- Второй пункт списка.")
    weekly.apply_weekly_office_to_document(doc)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("Проведено мероприятий." in t for t in texts)
    assert any("Второй пункт списка." in t for t in texts)
    assert not any(t.lstrip().startswith(("•", "-", "–", "○", "■")) for t in texts)
    assert sniot_fix.validate_list_markers(doc) == []


def test_weekly_keeps_numbered_subsection():
    doc = Document()
    doc.add_paragraph("1.1. Подпункт отчёта без маркера.")
    weekly.apply_weekly_office_to_document(doc)
    assert doc.paragraphs[0].text.startswith("1.1.")


def test_looks_like_daily_report_not_di():
    path = r"C:\Users\v.dubovik\Desktop\Ежедневные отчёты\на_правку\Ежедневный_отчёт_2026-08-17.docx"
    assert weekly.looks_like_weekly_itog(path)
    assert weekly.is_daily_report_filename(path)
    assert weekly.is_protected_daily_report(path)


def test_process_weekly_skips_accepted_daily(tmp_path: Path):
    folder = tmp_path / "принятые"
    folder.mkdir()
    path = folder / "Ежедневный_отчёт_2026-08-17.docx"
    doc = Document()
    doc.add_paragraph("Текст после правки пользователя  с  пробелами.")
    doc.save(str(path))
    result = weekly.process_weekly_itog_document(str(path))
    assert result["mode"] == "daily_accepted_skip"
    assert result["ok"] is True
    again = Document(str(path))
    assert "  с  пробелами" in again.paragraphs[0].text
