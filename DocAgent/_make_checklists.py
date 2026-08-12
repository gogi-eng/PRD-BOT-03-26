# -*- coding: utf-8 -*-
"""Чек-листы проверки котельной, теплового пункта, Мини-ТЭЦ."""
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


from _checklist_data import (
    BOILER_DOCS, BOILER_EQ, TP_DOCS, TP_EQ, MINI_DOCS, MINI_EQ,
    NORMS_BOILER, NORMS_TP, NORMS_MINI,
    NOTE_BOILER, NOTE_TP, NOTE_MINI,
)



def set_run_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, *, bold=False, center=False, size=11, space_after=3, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=9, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)


def add_table(doc, headers, rows, col_widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, center=True)
        set_cell_shading(table.rows[0].cells[i], "D9E2F3")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            sz = 8 if ci == 4 else 9
            set_cell_text(table.rows[ri + 1].cells[ci], str(val), bold=False, size=sz,
                          center=(ci in (0, 5, 6, 7)))
    # widths
    for row in table.rows:
        for i, w in enumerate(col_widths_cm):
            row.cells[i].width = Cm(w)
    return table


def setup_page(doc):
    for s in doc.sections:
        s.page_width = Cm(29.7)  # A4 landscape for readability
        s.page_height = Cm(21.0)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.0)
        s.top_margin = Cm(1.2)
        s.bottom_margin = Cm(1.2)


HEADERS_DOC = [
    "№",
    "Что проверить (наличие / правильность ведения)",
    "Кто заполняет / проверяет",
    "Периодичность и когда (после каких мероприятий)",
    "Норма",
    "Да",
    "Нет",
    "Не требуется",
    "Примечание",
]

HEADERS_EQ = [
    "№",
    "Что проверить на оборудовании",
    "Кто выполняет",
    "Периодичность / после каких мероприятий",
    "Норма",
    "Да",
    "Нет",
    "Не требуется",
    "Примечание",
]

W_DOC = [0.7, 6.2, 3.0, 4.4, 4.2, 0.9, 0.9, 1.5, 2.6]
W_EQ = [0.7, 6.2, 2.9, 4.5, 4.2, 0.9, 0.9, 1.5, 2.6]


def blank_checks(n=9):
    # return row template ending with empty Yes/No/N/A/Note
    return ["", "", "", ""]


def make_doc_rows(items):
    rows = []
    for i, it in enumerate(items, 1):
        rows.append([
            str(i),
            it[0],
            it[1],
            it[2],
            it[3],
            "☐",
            "☐",
            "☐",
            "",
        ])
    return rows


def header_block(doc, title, subtitle, norms):
    add_para(doc, title, bold=True, center=True, size=14, space_after=2)
    add_para(doc, subtitle, bold=True, center=True, size=12, space_after=4)
    add_para(
        doc,
        f"Дата составления формы: {datetime.now():%d.%m.%Y}   |   "
        "Объект: _______________________________   |   "
        "Адрес: _______________________________   |   "
        "Проверяющий: _________________   Дата проверки: __________",
        bold=False,
        center=False,
        size=10,
        space_after=4,
    )
    add_para(doc, "Нормативная база (выдержки для данного чек-листа):", bold=True, size=10, space_after=2)
    for n in norms:
        add_para(doc, "• " + n, size=9, space_after=1)
    add_para(
        doc,
        "Условные обозначения столбцов: «Да» — соответствует; «Нет» — нарушение/отсутствие; "
        "«Не требуется» — к данному объекту не относится (например, нет газа / нет сосудов); "
        "«Примечание» — пояснения, номера актов, сроки устранения. "
        "В колонке «Норма» указаны пункты ТНПА. На ОПО/ПОО проверяйте удостоверения на право "
        "обслуживания ПОО (ПМЧС № 31). Записи в журналах: полнота, даты, подписи, хронология.",
        size=9,
        space_after=6,
    )


def build_checklist(path: Path, title: str, subtitle: str, norms, docs, eqs, note: str):
    doc = Document()
    setup_page(doc)
    header_block(doc, title, subtitle, norms)

    add_para(doc, "РАЗДЕЛ 1. ДОКУМЕНТАЦИЯ", bold=True, center=True, size=12, space_after=4)
    add_para(
        doc,
        "Проверяется наличие документов и правильность записей: кто должен писать, как часто, "
        "после каких событий (пуск, останов, ремонт, авария, обход, поверка и т.д.).",
        size=9,
        space_after=4,
    )
    add_table(doc, HEADERS_DOC, make_doc_rows(docs), W_DOC)

    doc.add_paragraph()
    add_para(doc, "РАЗДЕЛ 2. ОБОРУДОВАНИЕ", bold=True, center=True, size=12, space_after=4)
    add_para(
        doc,
        "Проверяется техническое состояние и выполнение регламентных проверок на месте.",
        size=9,
        space_after=4,
    )
    add_table(doc, HEADERS_EQ, make_doc_rows(eqs), W_EQ)

    doc.add_paragraph()
    add_para(doc, "Итог проверки", bold=True, size=11, space_after=2)
    add_para(
        doc,
        "Выявлено нарушений: _____   из них критичных: _____   "
        "Предписание / замечания переданы: _________________   "
        "Срок устранения: __________",
        size=10,
        space_after=4,
    )
    add_para(
        doc,
        "Подпись проверяющего: _____________ / _____________     "
        "Подпись представителя объекта: _____________ / _____________",
        size=10,
        space_after=6,
    )
    add_para(doc, "Примечание:", bold=True, size=10, space_after=2)
    add_para(doc, note, size=9, space_after=2)
    add_para(
        doc,
        "Перед применением сверить актуальность редакций ТНПА (etalonline.by / БелГИСС). "
        "Чек-лист составлен по текстам предоставленных файлов нормативки и не заменяет полный текст Правил/ТКП.",
        size=8,
        space_after=2,
    )
    doc.save(str(path))
    print("OK", path, "docs", len(docs), "eq", len(eqs))



def main():
    out_dir = Path(r"C:\\Users\\v.dubovik\\Desktop\\Чек-листы_проверка_объектов")
    out_dir.mkdir(parents=True, exist_ok=True)
    also = Path(r"C:\\Users\\v.dubovik\\DocAgent\\Чек-листы_проверка_объектов")
    also.mkdir(parents=True, exist_ok=True)

    files = [
        (
            "Чек-лист_Котельная.docx",
            "ЧЕК-ЛИСТ ПРОВЕРКИ",
            "КОТЕЛЬНАЯ (газовая / на щепе / электрическая)",
            NORMS_BOILER,
            BOILER_DOCS,
            BOILER_EQ,
            NOTE_BOILER,
        ),
        (
            "Чек-лист_Тепловой_пункт.docx",
            "ЧЕК-ЛИСТ ПРОВЕРКИ",
            "ТЕПЛОВОЙ ПУНКТ (ИТП / ЦТП)",
            NORMS_TP,
            TP_DOCS,
            TP_EQ,
            NOTE_TP,
        ),
        (
            "Чек-лист_Мини-ТЭЦ.docx",
            "ЧЕК-ЛИСТ ПРОВЕРКИ",
            "МИНИ-ТЭЦ (комбинированная выработка тепла и электроэнергии)",
            NORMS_MINI,
            MINI_DOCS,
            MINI_EQ,
            NOTE_MINI,
        ),
    ]

    for name, title, sub, norms, docs, eqs, note in files:
        for folder in (out_dir, also):
            build_checklist(folder / name, title, sub, norms, docs, eqs, note)

    readme = out_dir / "ЧИТАТЬ_МЕНЯ.txt"
    readme.write_text(
        "Чек-листы для проверки объектов теплоэнергетики\n"
        "=============================================\n\n"
        "1) Чек-лист_Котельная.docx\n"
        "2) Чек-лист_Тепловой_пункт.docx\n"
        "3) Чек-лист_Мини-ТЭЦ.docx\n\n"
        "В каждом файле: Документация + Оборудование.\n"
        "Столбцы: Да | Нет | Не требуется | Примечание.\n"
        "В колонке Норма — ссылки на пункты ТНПА.\n"
        "На ОПО/ПОО — проверка удостоверений на право обслуживания ПОО (ПМЧС №31).\n",
        encoding="utf-8",
    )
    print("DONE", out_dir)


if __name__ == "__main__":
    main()
