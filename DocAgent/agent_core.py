# -*- coding: utf-8 -*-
"""Ядро агента: определение вида документа, примеры, запуск форматтеров."""

from __future__ import annotations

import json
import os
import re
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document

from rules import (
    DOCUMENT_TYPES,
    SAMPLE_POLICY,
    detect_type_from_text,
    filename_looks_like_di,
    is_instruction_aligned_sample,
    looks_like_weekly_itog,
)
from path_resolver import (
    USER_AGENT_DIR,
    USER_PROJECT_DIR,
    assert_path_writable,
    canonical_fs_path,
    is_allowed_sample_path,
    is_path_in_user_agent_dir,
    is_path_in_writable_user_dir,
    is_sniot_doc,
    list_agent_sample_paths,
    live_user_agent_dir,
    pick_best_agent_sample,
    writable_dirs_hint,
    _sample_rank,
)
from formatters.common import apply_basic_office_format

ROOT = Path(__file__).resolve().parent
CONFIG_PATH = ROOT / "config.json"
LOG_DIR = ROOT / "logs"


def load_config() -> dict:
    with open(CONFIG_PATH, encoding="utf-8") as f:
        return json.load(f)


def is_conservative_di_satp(input_path: str, doc_type: str) -> bool:
    """См. formatters.sniot_document.is_conservative_di_satp."""
    from formatters.sniot_document import is_conservative_di_satp as _conservative

    return _conservative(input_path, doc_type)


def log(msg: str) -> None:
    LOG_DIR.mkdir(exist_ok=True)
    line = f"{datetime.now():%Y-%m-%d %H:%M:%S}  {msg}\n"
    with open(LOG_DIR / "agent.log", "a", encoding="utf-8") as f:
        f.write(line)


def read_preview_text(path: str, limit: int = 2500) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in (".doc", ".rtf"):
        return _read_doc_via_word(path, limit)
    if ext != ".docx":
        return ""
    doc = Document(path)
    parts = []
    total = 0

    def _append(text: str) -> bool:
        nonlocal total
        t = text.strip()
        if not t:
            return total <= limit
        parts.append(t)
        total += len(t)
        return total <= limit

    for p in doc.paragraphs:
        if not _append(p.text):
            break
    if total <= limit:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if not _append(cell.text):
                        break
                if total > limit:
                    break
            if total > limit:
                break
    return "\n".join(parts)


def _read_doc_via_word(path: str, limit: int = 2500) -> str:
    """Старый .doc / .rtf — через Word. Word пользователя не закрываем."""
    try:
        from formatters.word_com import open_docx_readonly
    except Exception:
        return ""
    try:
        with open_docx_readonly(path) as doc:
            return (doc.Content.Text or "")[:limit]
    except Exception:
        return ""


def ensure_pywin32() -> None:
    """Проверить пакет pywin32 (нужен для .doc и .rtf)."""
    try:
        import win32com.client  # noqa: F401
    except ImportError as e:
        raise RuntimeError(
            "Для файлов .doc и .rtf нужен пакет pywin32.\n\n"
            "Что сделать:\n"
            "1) Запустите install_deps.bat в папке DocAgent\n"
            "   или в командной строке:\n"
            "   python -m pip install pywin32 python-docx\n"
            "2) Убедитесь, что на компьютере установлен Microsoft Word.\n"
            "3) Снова запустите агент (start_agent.bat)."
        ) from e


def convert_to_docx(path: str) -> str:
    """
    Конвертация .doc / .rtf → .docx во временную папку TEMP (через Word).
    В папку Агент _converted не пишем. .docx возвращает как есть.

    Важно: отдельный скрытый Word (DispatchEx), копия без Zone.Identifier,
    OpenAndRepair, на время открытия отключается File Validation —
    иначе Office блокирует старые .doc/.rtf («проблема с этим файлом»).
    """
    from formatters.word_com import (
        convert_legacy_word_to_docx,
        is_legacy_word_file,
    )

    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return path
    if not is_legacy_word_file(path):
        raise RuntimeError(
            f"Неподдерживаемый формат файла (нужен .docx, .doc или .rtf):\n{path}"
        )

    ensure_pywin32()
    abs_src = os.path.abspath(path)
    if not os.path.isfile(abs_src):
        raise RuntimeError(f"Файл не найден:\n{abs_src}")

    out = Path(tempfile.mkdtemp(prefix="docagent_conv_")) / "converted.docx"

    try:
        convert_legacy_word_to_docx(abs_src, out)
    except Exception as e:
        ext_note = ""
        if ext == ".doc":
            ext_note = (
                "Если это очень старый .doc, Word на компьютере иногда не открывает его "
                "автоматически.\n"
                "Откройте файл двойным щелчком в Word → «Файл» → «Сохранить как» → "
                "тип «Документ Word (*.docx)».\n"
                "Файлы .rtf можно оформлять сразу, без этого шага.\n\n"
            )
        raise RuntimeError(
            "Не удалось открыть старый файл Word (.doc / .rtf).\n\n"
            + ext_note
            + "Либо положите готовый .docx в папку «Агент» и нажмите «Оформить документ».\n\n"
            f"Исходник: {abs_src}\n"
            f"Ошибка: {e}"
        ) from e

    if not os.path.isfile(out):
        raise RuntimeError(f"Не удалось сохранить .docx:\n{out}")
    return str(out)


# совместимость со старым именем
def convert_doc_to_docx(path: str) -> str:
    return convert_to_docx(path)


def detect_document(path: str) -> tuple[str, str]:
    text = read_preview_text(path)
    dtype = detect_type_from_text(os.path.basename(path), text)
    sniot_types = {
        "dolzhnostnaya_instrukciya",
        "rabochaya_instrukciya",
        "prikaz",
        "instrukciya_ot",
        "polozhenie",
    }
    if dtype not in sniot_types and looks_like_weekly_itog(
        os.path.basename(path), text, path
    ):
        dtype = "ezhenedelnyy_itog"
    return dtype, text[:800]


def _tokenize_tokens(name: str) -> set[str]:
    """Значимые куски имени файла для подбора похожего образца."""
    low = (name or "").lower()
    low = re.sub(r"[_\-\.]+", " ", low)
    low = re.sub(r"\.(docx|doc|rtf)$", "", low)
    stop = {
        "docx",
        "doc",
        "rtf",
        "converted",
        "оформлен",
        "рабочая",
        "инструкция",
        "должностная",
        "положение",
        "приказ",
        "разряда",
        "разряд",
        "ого",
        "го",
        "по",
        "и",
        "для",
        "сниот",
        "изменено",
        "исправленный",
        "исправлено",
        "оформлен",
        "копия",
    }
    tokens = set()
    for t in re.findall(r"[a-zа-яё0-9]+", low, flags=re.IGNORECASE):
        t = t.lower()
        if len(t) < 3 or t in stop or t.isdigit():
            continue
        tokens.add(t)
    # числа разрядов полезны: 3, 4, 5
    for m in re.findall(r"\b([3-6])\b", low):
        tokens.add(f"разряд{m}")
    # короткие, но важные маркеры темы
    for short in ("цтп", "итп", "кл", "ри", "ди", "иот"):
        if re.search(rf"(?<![a-zа-яё0-9]){short}(?![a-zа-яё0-9])", low):
            tokens.add(short)
    return tokens


def _topic_tags(name: str) -> set[str]:
    """
    Тема документа по имени файла.
    Нельзя брать образец «силовые КЛ», если оформляем «ЦТП».
    """
    low = (name or "").lower().replace("ё", "е")
    tags: set[str] = set()
    if "цтп" in low or "итп" in low or "теплового пункта" in low or "тепловой пункт" in low:
        tags.add("ctp_itp")
    if (
        "кабельн" in low
        or "силовых кл" in low
        or "силовые кл" in low
        or "0,4-10" in low
        or "0.4-10" in low
        or re.search(r"(?<![a-zа-яё0-9])кл(?![a-zа-яё0-9])", low)
    ):
        tags.add("cable_kl")
    if "слесар" in low:
        tags.add("slesar")
    if "диспетчер" in low:
        tags.add("dispatcher")
    if "положен" in low and "эксплуатац" not in low:
        tags.add("org_polozhenie")
    if "охране труда" in low or re.search(r"(?<![a-zа-яё0-9])иот(?![a-zа-яё0-9])", low):
        tags.add("iot")
    if "эксплуатац" in low:
        tags.add("ekspluatacii")
    return tags


def _examples_topic_compatible(source_path: str, example_path: str) -> bool:
    """Эталон подходит по теме (ЦТП≠КЛ, РИ≠ДИ и т.п.)."""
    src = _topic_tags(os.path.basename(source_path))
    ex = _topic_tags(os.path.basename(example_path))
    if not src or not ex:
        return True
    # явный конфликт тем
    conflicts = [
        ("ctp_itp", "cable_kl"),
        ("slesar", "dispatcher"),
        ("iot", "ekspluatacii"),
        ("org_polozhenie", "ekspluatacii"),
    ]
    for a, b in conflicts:
        if (a in src and b in ex) or (b in src and a in ex):
            return False
    # если у обеих есть «узкие» теги — нужно пересечение
    narrow = {"ctp_itp", "cable_kl", "slesar", "dispatcher", "iot", "org_polozhenie"}
    src_n = src & narrow
    ex_n = ex & narrow
    if src_n and ex_n:
        return bool(src_n & ex_n)
    return True


def _coerce_allowed_sample(path: str | None) -> str | None:
    """Config/поле образца: только Агент + «образец» в имени, иначе None."""
    if not path or not str(path).strip():
        return None
    raw = str(path).strip()
    if is_allowed_sample_path(raw) and os.path.isfile(raw):
        return os.path.normpath(raw)
    return None


def choose_best_example(doc_type: str, source_path: str, limit: int = 20) -> dict | None:
    """
    Лучший образец только из папки Агент, в имени есть «образец».
    Предпочитает имя, близкое к целевому документу; иначе любой *образец* в Агент.
    Пути из ОБМЕН / САТП / config вне Агент — игнорируются.
    """
    preferred_path = None
    try:
        from formatters.text_edits import load_patterns

        prefs = (load_patterns() or {}).get("preferred_samples") or {}
        if doc_type == "prikaz":
            preferred_path = prefs.get("prikaz") or prefs.get("master_prikaz")
        elif doc_type == "dolzhnostnaya_instrukciya":
            preferred_path = (
                prefs.get("dolzhnostnaya_instrukciya")
                or prefs.get("master_di")
                or prefs.get("master_instructions")
            )
        if not preferred_path:
            preferred_path = prefs.get(doc_type)
        cfg = load_config()
        masters = cfg.get("master_samples") or {}
        master = (masters.get(doc_type) or "").strip()
        if master:
            preferred_path = master
        elif doc_type == "dolzhnostnaya_instrukciya" and not preferred_path:
            legacy = (cfg.get("master_sample_path") or "").strip()
            if legacy:
                preferred_path = legacy
    except Exception:
        preferred_path = None

    preferred_path = _coerce_allowed_sample(preferred_path)
    if preferred_path:
        src_abs = os.path.normcase(os.path.abspath(source_path))
        pref_abs = os.path.normcase(os.path.abspath(preferred_path))
        src_tok = _tokenize_tokens(os.path.basename(source_path))
        pref_tok = _tokenize_tokens(os.path.basename(preferred_path))
        if not (src_tok & pref_tok):
            log(
                "AUTO example PREFERRED skipped (нет общих слов в имени): "
                f"{os.path.basename(preferred_path)} ← для {os.path.basename(source_path)}"
            )
        elif pref_abs != src_abs and _examples_topic_compatible(source_path, preferred_path):
            log(f"AUTO example PREFERRED: {preferred_path}")
            return {
                "path": preferred_path,
                "name": os.path.basename(preferred_path),
                "folder": os.path.basename(os.path.dirname(preferred_path)),
                "mtime": os.path.getmtime(preferred_path),
                "aligned": True,
                "score": 999,
                "label": f"[ЭТАЛОН] {os.path.basename(preferred_path)}",
                "common_tokens": [],
            }
        if pref_abs != src_abs:
            log(
                "AUTO example PREFERRED skipped (тема не совпадает или не Агент/образец): "
                f"{os.path.basename(preferred_path)} ← для {os.path.basename(source_path)}"
            )

    items = find_examples(doc_type, limit=max(limit, 40), source_path=source_path)
    if not items:
        picked = pick_best_agent_sample(source_path)
        if picked is None:
            log("AUTO example: нет файла «образец» в папке Агент — оформление по правилам mdc")
            return None
        return {
            "path": str(picked),
            "name": picked.name,
            "folder": picked.parent.name,
            "mtime": picked.stat().st_mtime if picked.is_file() else 0,
            "aligned": True,
            "score": 1,
            "label": f"[Агент] {picked.name}",
            "common_tokens": [],
        }

    src_abs = os.path.normcase(os.path.abspath(source_path))
    src_name = os.path.basename(source_path)
    src_tokens = _tokenize_tokens(src_name)
    src_topics = _topic_tags(src_name)

    scored: list[tuple[float, dict]] = []
    for item in items:
        path = item["path"]
        if not is_allowed_sample_path(path):
            continue
        try:
            p_abs = os.path.normcase(os.path.abspath(path))
        except OSError:
            continue
        if p_abs == src_abs:
            continue
        name_low = item["name"].lower()
        if name_low.startswith("~$"):
            continue
        exact, overlap = _sample_rank(Path(source_path), Path(path))
        tok = _tokenize_tokens(item["name"])
        common = src_tokens & tok
        ex_topics = _topic_tags(item["name"])
        if exact == 0 and overlap < 2 and not common and not (src_topics & ex_topics):
            continue
        score = 0.0
        if item.get("aligned"):
            score += 35
        score += 12 * len(common)
        score += 80 * len(src_topics & ex_topics)
        if "цтп" in src_tokens and "цтп" in tok:
            score += 200
        if "эксплуатац" in src_name.lower() and "эксплуатац" in name_low:
            score += 60
        if ("ctp_itp" in src_topics) and ("cable_kl" in ex_topics):
            score -= 500
        if ("cable_kl" in src_topics) and ("ctp_itp" in ex_topics):
            score -= 500
        score += min(item.get("mtime", 0) / 1e12, 5)
        item = dict(item)
        item["score"] = score
        item["common_tokens"] = sorted(common)
        scored.append((score, item))

    if not scored:
        picked = pick_best_agent_sample(source_path)
        if picked is None:
            return None
        return {
            "path": str(picked),
            "name": picked.name,
            "folder": picked.parent.name,
            "mtime": picked.stat().st_mtime if picked.is_file() else 0,
            "aligned": True,
            "score": 1,
            "label": f"[Агент] {picked.name}",
            "common_tokens": [],
        }
    scored.sort(key=lambda x: (-x[0], -x[1].get("mtime", 0)))
    best = scored[0][1]
    log(
        f"AUTO example type={doc_type} score={best.get('score')} "
        f"file={best.get('name')} common={best.get('common_tokens')}"
    )
    return best


def find_examples(doc_type: str, limit: int = 100, source_path: str | None = None) -> list[dict]:
    """
    Все *образец*.docx только из папки Агент (N:\\ или UNC).
    Без обхода ОБМЕН / N:\\. Интернет-шаблоны не используются.
    """
    if SAMPLE_POLICY.get("allow_web_templates"):
        log("WARNING: web templates enabled — отклонение от политики")

    if doc_type == "ezhenedelnyy_itog":
        return []

    results = []
    folder = live_user_agent_dir()
    folder_name = folder.name if folder is not None else USER_AGENT_DIR.name
    for full_path in list_agent_sample_paths():
        name = full_path.name
        try:
            mtime = full_path.stat().st_mtime
        except OSError:
            continue
        aligned = is_instruction_aligned_sample(name)
        results.append(
            {
                "path": str(full_path),
                "name": name,
                "folder": folder_name,
                "mtime": mtime,
                "aligned": aligned,
                "label": f"[Агент / образец] {name}",
            }
        )
    src_name = os.path.basename(source_path) if source_path else ""
    results.sort(
        key=lambda item: (
            not item["aligned"],
            -len(_tokenize_tokens(src_name) & _tokenize_tokens(item["name"])),
            -item["mtime"],
        )
    )
    seen = set()
    uniq = []
    for item in results:
        if item["name"] in seen:
            continue
        seen.add(item["name"])
        uniq.append(item)
        if len(uniq) >= limit:
            break
    return uniq


def _clean_output_stem(stem: str) -> str:
    """Имя результата = имя исходника + «_оформлен», без накопления суффиксов."""
    s = stem or ""
    while True:
        low = s.lower()
        if low.endswith("_converted"):
            s = s[: -len("_converted")]
            continue
        if low.endswith("_оформлен"):
            s = s[: -len("_оформлен")]
            continue
        if low.endswith("_исправлен_агент"):
            s = s[: -len("_исправлен_агент")]
            continue
        break
    return s


def output_path_for(input_path: str) -> str:
    src = canonical_fs_path(input_path)
    stem = _clean_output_stem(src.stem)
    if is_path_in_writable_user_dir(src):
        return str(src.parent / f"{stem}_оформлен.docx")
    return str(USER_AGENT_DIR / f"{stem}_оформлен.docx")


def _can_write_file(path: str) -> bool:
    """Проверка: можно ли создать/перезаписать файл (Агент или Проекты)."""
    try:
        assert_path_writable(path)
    except PermissionError:
        return False
    try:
        parent = os.path.dirname(path) or "."
        if not os.path.isdir(parent):
            return False
        if os.path.exists(path):
            with open(path, "r+b"):
                pass
            return True
        test = os.path.join(parent, f"~docagent_wtest_{os.getpid()}.tmp")
        with open(test, "wb") as f:
            f.write(b"0")
        os.remove(test)
        return True
    except OSError:
        return False


def _is_desktop_path(path: str | Path) -> bool:
    try:
        resolved = str(Path(path).resolve()).lower()
        desktop = str(Path.home() / "Desktop").lower()
        return resolved.startswith(desktop)
    except OSError:
        return False


def _output_candidates_for_input(input_path: str) -> list[str]:
    """Пути для *_оформлен.docx: папка Агент / рядом с исходником; Desktop только если источник там."""
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = _clean_output_stem(Path(input_path).stem)
    candidates: list[str] = []

    if is_path_in_writable_user_dir(input_path):
        same_dir = Path(input_path).parent
        candidates.extend(
            [
                str(same_dir / f"{stem}_оформлен.docx"),
                str(same_dir / f"{stem}_оформлен_{stamp}.docx"),
            ]
        )
        preferred = output_path_for(input_path)
        for path in (preferred, f"{os.path.splitext(preferred)[0]}_{stamp}.docx"):
            if path not in candidates:
                candidates.append(path)
    else:
        preferred = output_path_for(input_path)
        candidates.extend(
            [
                preferred,
                f"{os.path.splitext(preferred)[0]}_{stamp}.docx",
            ]
        )

    if _is_desktop_path(input_path):
        desktop = Path.home() / "Desktop"
        for path in (
            str(desktop / f"{stem}_оформлен.docx"),
            str(desktop / f"{stem}_оформлен_{stamp}.docx"),
        ):
            if path not in candidates:
                candidates.append(path)

    return candidates


def pick_output_path(input_path: str, work_path: str | None = None) -> tuple[str, str | None]:
    """
    Выбрать путь для *_оформлен.docx.
    Имя — от исходного файла (.doc/.rtf/.docx), не от _converted.
    Для документов из папки Агент — сохранение в ту же папку, не на Desktop.
    """
    candidates = _output_candidates_for_input(input_path)
    note = None
    for i, path in enumerate(candidates):
        if _can_write_file(path):
            if i > 0:
                note = (
                    "Рядом с исходником сохранить не удалось (часто файл уже открыт в Word "
                    f"или нет прав на сетевую папку). Сохранено сюда:\n{path}"
                )
                log(f"OUTPUT fallback -> {path}")
            return path, note

    last = candidates[-1]
    note = (
        "Нет доступа к папке документа. Попытка сохранить:\n"
        f"{last}\n\nЗакройте в Word файлы с окончанием «_оформлен.docx» и повторите."
    )
    return last, note


def friendly_save_error(err: BaseException, out_path: str) -> RuntimeError:
    msg = str(err)
    if isinstance(err, PermissionError) or "Permission denied" in msg or "[Errno 13]" in msg:
        return RuntimeError(
            "Нет доступа к файлу результата (Permission denied).\n\n"
            "Что сделать:\n"
            "1) Закройте в Word файл с окончанием «_оформлен.docx» "
            "(и сам исходник, если он открыт только для чтения с блокировкой).\n"
            "2) Нажмите «Оформить» ещё раз — агент сохранит в папку Агент с новым именем, "
            "если исходный файл открыт в Word.\n\n"
            f"Путь: {out_path}"
        )
    return RuntimeError(f"Ошибка оформления: {err}")


def _apply_russian_language_check(summary: dict, cfg: dict, enabled: bool) -> dict:
    """После оформления — внимательно перечитать весь текст на русский язык."""
    if not enabled:
        return summary
    out = summary.get("output")
    if not out or not os.path.isfile(out):
        return summary
    try:
        from formatters.russian_check import check_and_fix_document

        log(f"RUSSIAN CHECK start: {out}")
        rep = check_and_fix_document(
            out,
            use_yandex=bool(cfg.get("russian_check_yandex", True)),
            use_languagetool=bool(cfg.get("russian_check_languagetool", True)),
            auto_fix=True,
            lt_max_paragraphs=int(cfg.get("russian_check_lt_max_paragraphs", 20)),
        )
        summary["russian_check"] = rep
        summary["actions"].append(
            "Проверка русского языка (весь текст, все типы документов): "
            f"локально {rep.get('local_fixes', 0)}, "
            f"орфография {rep.get('speller_fixes', 0)}, "
            f"грамматика/замечания {rep.get('grammar_notes', 0)}; "
            "аббревиатуры (ЛСиМ, СНиОТ, ТКП…) не правятся"
            + (
                f", пропущено {rep.get('skipped_abbreviations', 0)}"
                if rep.get("skipped_abbreviations")
                else ""
            )
        )
        if rep.get("report_path"):
            summary["actions"].append(f"Отчёт по русскому языку: {rep['report_path']}")
        for s in (rep.get("suggestions") or [])[:5]:
            summary["actions"].append(s)
        log(
            f"RUSSIAN CHECK done spell={rep.get('speller_fixes')} "
            f"notes={rep.get('grammar_notes')}"
        )
        # проверка русского может переписать runs — вернуть эталон подписантов
        try:
            from docx import Document
            from formatters.common import save_docx_unprotected
            from formatters.structure_fix import (
                apply_signatory_block_format,
                resolve_instruction_doc_type,
            )

            dtype = resolve_instruction_doc_type(
                doc_type=summary.get("type") or "",
                doc=Document(out),
                source_path=summary.get("input") or out,
            )
            if dtype in ("rabochaya_instrukciya", "dolzhnostnaya_instrukciya"):
                doc_s = Document(out)
                n_sign = apply_signatory_block_format(doc_s, doc_type=dtype)
                if n_sign:
                    save_docx_unprotected(doc_s, out)
                    summary["actions"].append(
                        f"После проверки русского — подписанты восстановлены ({n_sign})"
                    )
        except Exception as e_sign:
            log(f"signatories restore after RU check: {e_sign}")
    except Exception as e:
        log(f"RUSSIAN CHECK fail: {e}")
        summary["actions"].append(
            "Проверка русского языка не завершена (нет сети или сервис недоступен). "
            "Проверьте вручную: https://languagetool.org/ru и https://text.ru/spelling. "
            f"Ошибка: {e}"
        )
    return summary


def process_document(
    input_path: str,
    doc_type: str,
    example_path: str | None = None,
    use_basic_only: bool = False,
    apply_text_edits_flag: bool = True,
    apply_russian_check_flag: bool | None = None,
    structure_rebuild_flag: bool = False,
    auto_pick_example: bool = True,
) -> dict:
    """
    Главный обработчик.
    example_path — образец только из папки Агент, в имени «образец»; иначе агент ищет сам.
    Нет образца / выбран «Стандарт» (auto_pick_example=False) — оформление по правилам
    Инструкции 2025 без эталона docx (без ОБМЕН / сети).
    apply_text_edits_flag — правки/удаления по сравнению ваших файлов.
    apply_russian_check_flag — перечитать весь текст на правила русского языка.
    structure_rebuild_flag — перестроить содержание по структуре образца (ЦЭМ и т.п.).
    """
    input_path = str(canonical_fs_path(input_path))
    weekly_hit = doc_type == "ezhenedelnyy_itog" or looks_like_weekly_itog(
        os.path.basename(input_path), path=input_path
    )
    if weekly_hit:
        from formatters.weekly_report_office import process_weekly_itog_document

        cfg_weekly = load_config()
        if apply_russian_check_flag is None:
            apply_russian_check_flag = bool(cfg_weekly.get("russian_language_check", True))
        weekly_summary = process_weekly_itog_document(input_path)
        return _apply_russian_language_check(
            weekly_summary, cfg_weekly, bool(apply_russian_check_flag)
        )
    if not is_path_in_writable_user_dir(input_path):
        raise RuntimeError(
            "Редактировать можно только файлы в папке Агент или Проекты:\n"
            f"{writable_dirs_hint()}\n\n"
            f"Указан: {input_path}\n\n"
            "Образец — только файл со словом «образец» в папке Агент."
        )
    cfg = load_config()
    if apply_russian_check_flag is None:
        apply_russian_check_flag = bool(cfg.get("russian_language_check", True))

    def finish(summary: dict) -> dict:
        # ВСЕГДА проверять документ перед выдачей результата
        out_path = summary.get("output") or ""
        try:
            from formatters.publish_check import verify_document_before_publish

            if out_path and os.path.isfile(out_path):
                chk = verify_document_before_publish(
                    out_path, doc_type=summary.get("type") or doc_type
                )
                summary["publish_check"] = chk
                if chk.get("details"):
                    summary.setdefault("actions", []).append(
                        "Проверка перед публикацией: "
                        + (
                            "ОК"
                            if chk.get("ok") and not chk.get("issues")
                            else "есть замечания"
                        )
                    )
                for d in (chk.get("details") or [])[:6]:
                    if d not in (summary.get("actions") or []):
                        summary.setdefault("actions", []).append(d)
                for iss in (chk.get("issues") or [])[:5]:
                    summary.setdefault("actions", []).append("⚠ " + iss)
                    log(f"PUBLISH CHECK issue: {iss}")
        except Exception as e:
            log(f"publish_check error: {e}")
            summary.setdefault("actions", []).append(
                f"проверка перед публикацией: ошибка ({e})"
            )
        # ВСЕГДА в самом конце: примечания + подписанты по эталону
        # Для ДИ — пропуск finalize (ломает нумерацию); финал = fix_sniot_document
        try:
            if summary.get("conservative_di_satp"):
                log("Conservative СНиОТ: skip finalize_notes_and_signatories")
                summary.setdefault("actions", []).append(
                    "СНиОТ: без text_edits и finalize — только оформление; "
                    "проверка русского — по галочке (аббревиатуры не правятся)"
                )
            else:
                from formatters.structure_fix import finalize_notes_and_signatories

                if out_path and os.path.isfile(out_path):
                    fin = finalize_notes_and_signatories(
                        out_path, summary.get("type") or doc_type
                    )
                    summary["finalize_tail"] = fin
                    if fin.get("notes_moved"):
                        summary.setdefault("actions", []).append(
                            f"Примечания перенесены перед «Разработал:» ({fin['notes_moved']})"
                        )
                    if fin.get("razrabotal_added"):
                        summary.setdefault("actions", []).append(
                            "Добавлена строка «Разработал:» перед подписантом"
                        )
                    if fin.get("signatories"):
                        summary.setdefault("actions", []).append(
                            f"Подписанты по эталону ({fin.get('doc_type')}): {fin['signatories']}"
                        )
                    if fin.get("indents_fixed") is not None:
                        summary.setdefault("actions", []).append(
                            f"Отступ 1,25 см на текст (правок: {fin.get('indents_fixed', 0)})"
                        )
                    if fin.get("chapters_styled") or fin.get("chapters_centered"):
                        summary.setdefault("actions", []).append(
                            f"Главы по центру: {fin.get('chapters_styled') or fin.get('chapters_centered')}; "
                            f"пустых строк перед главами: {fin.get('chapter_spacers', 0)}"
                        )
                    ch_rep = fin.get("chapters_repaired") or {}
                    if ch_rep.get("numbered") or ch_rep.get("merged"):
                        summary.setdefault("actions", []).append(
                            f"Восстановлены номера глав: {ch_rep.get('numbered', 0)}, "
                            f"склеено строк: {ch_rep.get('merged', 0)}"
                        )
                    toc = fin.get("contents_page") or {}
                    if toc.get("found"):
                        summary.setdefault("actions", []).append(
                            "«Содержание» на отдельной странице; первая глава с верха листа"
                        )
                    if fin.get("table_fonts") is not None:
                        summary.setdefault("actions", []).append(
                            f"Таблицы тела — 12 пт (правок шрифта: {fin.get('table_fonts', 0)})"
                        )
                    if fin.get("numbers_fixed") is not None:
                        summary.setdefault("actions", []).append(
                            f"Нумерация проверена (+{fin.get('numbers_fixed', 0)})"
                        )
        except Exception as e:
            log(f"finalize_tail error: {e}")
            summary.setdefault("actions", []).append(
                f"финал примечаний/подписантов: ошибка ({e})"
            )
        # ВСЕГДА ПОСЛЕДНИМ: полные правила СНиОТ (fix_sniot_document.py)
        try:
            from formatters.sniot_document import (
                apply_sniot_rules_to_output,
                should_apply_sniot_pass,
            )

            src = summary.get("input") or input_path
            dtype = summary.get("type") or doc_type
            if (
                out_path
                and os.path.isfile(out_path)
                and out_path.lower().endswith(".docx")
                and dtype != "prikaz"
                and should_apply_sniot_pass(src, out_path, dtype)
            ):
                log(f"SNIOT pass start: {out_path}")
                sniot = apply_sniot_rules_to_output(out_path)
                summary["sniot_pass"] = sniot
                if sniot.get("build"):
                    summary["sniot_build"] = sniot["build"]
                    acts = summary.setdefault("actions", [])
                    if sniot["build"] in acts:
                        acts.remove(sniot["build"])
                    acts.insert(0, sniot["build"])
                for act in sniot.get("actions") or []:
                    if act not in (summary.get("actions") or []):
                        summary.setdefault("actions", []).append(act)
                after_n = len(sniot.get("after_issues") or [])
                if sniot.get("ok") and sniot.get("applied"):
                    log(f"SNIOT pass OK: validation issues={after_n}")
                elif sniot.get("applied"):
                    log(f"SNIOT pass done with issues ({after_n}): {sniot.get('after_issues')}")
                    summary.setdefault("actions", []).insert(
                        0,
                        f"⚠ СНиОТ: остались замечания ({after_n}) — проверьте отступы и нумерацию",
                    )
                elif not sniot.get("applied"):
                    log(f"SNIOT pass not applied: {sniot.get('actions')}")
                    acts = sniot.get("actions") or []
                    build = next(
                        (a for a in acts if a.startswith("СНиОТ: сборка")),
                        "",
                    )
                    if any("остались замечания" in a for a in acts):
                        head = (
                            "⚠ СНиОТ не записал «_оформлен» — после правки остались замечания. "
                            "Ниже только то, что не удалось исправить."
                        )
                    elif any("закройте" in a.lower() and "word" in a.lower() for a in acts):
                        head = (
                            "⚠ СНиОТ НЕ применён — закройте _оформлен.docx в Word и повторите."
                        )
                    else:
                        head = "⚠ СНиОТ НЕ применён — смотрите список ниже."
                    if build:
                        head = f"{head} {build}"
                    summary.setdefault("actions", []).insert(0, head)
                else:
                    log(f"SNIOT pass failed: {sniot.get('actions')}")
            elif out_path and dtype != "prikaz":
                log(
                    f"SNIOT pass skipped src={src!r} out={out_path!r} "
                    f"type={dtype} apply={should_apply_sniot_pass(src, out_path, dtype)}"
                )
        except Exception as e:
            log(f"sniot_pass error: {e}")
            summary.setdefault("actions", []).append(
                f"финал СНиОТ (fix_sniot_document): ошибка ({e})"
            )
        # После XML СНиОТ: спеллер/грамматика DocAgent по финальному тексту
        # (все типы документов, включая ДИ; галочка по умолчанию включена).
        summary = _apply_russian_language_check(
            summary,
            cfg,
            bool(apply_russian_check_flag),
        )
        return summary

    # если в имени/титуле ясно написано РИ/ДИ — не оформлять как приказ
    try:
        detected, _preview = detect_document(input_path)
    except Exception:
        detected = doc_type
    name_l = os.path.basename(input_path).lower()
    if filename_looks_like_di(name_l):
        if doc_type != "dolzhnostnaya_instrukciya":
            log(
                f"TYPE override {doc_type} -> dolzhnostnaya_instrukciya (имя «ДИ …»)"
            )
        doc_type = "dolzhnostnaya_instrukciya"
        detected = "dolzhnostnaya_instrukciya"
    strong = {
        "rabochaya_instrukciya": ("рабоч" in name_l and "инструкц" in name_l)
        or detected == "rabochaya_instrukciya",
        "dolzhnostnaya_instrukciya": (
            filename_looks_like_di(name_l)
            or ("должностн" in name_l and "инструкц" in name_l)
            or name_l.startswith("ди ")
            or ("проект" in name_l and "мастер" in name_l)
            or detected == "dolzhnostnaya_instrukciya"
        ),
        "polozhenie": (
            (not filename_looks_like_di(name_l))
            and (
                ("положен" in name_l)
                or ("инструкц" in name_l and "эксплуатац" in name_l)
                or (
                    "инструкц" in name_l
                    and "рабоч" not in name_l
                    and "должностн" not in name_l
                    and "охране труда" not in name_l
                    and "иот" not in name_l
                )
                or detected == "polozhenie"
            )
        ),
        "instrukciya_ot": "охране труда" in name_l or "иот" in name_l,
    }
    if detected in strong and strong.get(detected) and detected != doc_type:
        log(f"TYPE override {doc_type} -> {detected} (имя/титул документа)")
        doc_type = detected

    conservative_di = is_conservative_di_satp(input_path, doc_type)
    if conservative_di:
        apply_text_edits_flag = False
        structure_rebuild_flag = False
        log(
            "Conservative СНиОТ: text_edits и structure_fix отключены; "
            f"проверка русского={'да' if apply_russian_check_flag else 'нет'}"
        )

    if example_path:
        coerced = _coerce_allowed_sample(example_path)
        if coerced is None:
            log(
                "example_path отклонён (не папка Агент или нет слова «образец» в имени): "
                f"{example_path}"
            )
            example_path = None
        else:
            example_path = coerced

    if not example_path and auto_pick_example and doc_type not in (
        "unsupported",
        "ezhenedelnyy_itog",
    ):
        best = choose_best_example(doc_type, input_path)
        if best:
            example_path = best["path"]
            log(f"AUTO-selected example: {example_path}")
        else:
            log("AUTO example: нет подходящего «образец» в папке Агент — Инструкция 2025 без эталона")

    log(
        f"PROCESS type={doc_type} file={input_path} example={example_path} "
        f"text_edits={apply_text_edits_flag} russian_check={apply_russian_check_flag} "
        f"structure_rebuild={structure_rebuild_flag}"
    )

    work_path = input_path
    ext = os.path.splitext(work_path)[1].lower()
    if ext in (".doc", ".rtf"):
        work_path = convert_to_docx(work_path)
        log(f"Converted to {work_path}")

    # --- Перестройка содержания по образцу (Положение о ПК ↔ ЦЭМ) ---
    if structure_rebuild_flag or (doc_type == "polozhenie" and not conservative_di):
        try:
            from formatters.polozhenie_pk_rebuild import (
                is_cem_sample,
                is_pk_polozhenie,
                rebuild_pk_by_cem_structure,
            )

            want = structure_rebuild_flag or (
                is_pk_polozhenie(input_path) and is_cem_sample(example_path)
            )
            if want and (is_pk_polozhenie(input_path) or is_cem_sample(example_path)):
                rebuilt = rebuild_pk_by_cem_structure(
                    input_path,
                    sample_path=example_path,
                    python_exe=cfg.get("python_exe"),
                )
                rebuilt.setdefault("input", input_path)
                rebuilt.setdefault("work", work_path)
                rebuilt.setdefault("type", "polozhenie")
                rebuilt.setdefault("example", example_path)
                rebuilt.setdefault("actions", [])
                # довести отступы / пустые / заголовки по правилам делопроизводства
                try:
                    from formatters.format_polozhenie_pk import format_polozhenie_pk

                    fmt = format_polozhenie_pk(rebuilt["output"], backup=False)
                    rebuilt["actions"].append(
                        f"Оформление: заголовков {fmt.get('headings', 0)}, "
                        f"пустых перед главами +{fmt.get('empties_added', 0)}"
                    )
                except Exception as fe:
                    log(f"format_polozhenie_pk error: {fe}")
                    rebuilt["actions"].append(f"оформление отступов: ошибка ({fe})")
                rebuilt["actions"].insert(
                    0,
                    "Режим: перестройка содержания по структуре образца (не только шрифты)",
                )
                return finish(rebuilt)
        except Exception as e:
            log(f"structure_rebuild skipped/failed: {e}")
            if structure_rebuild_flag:
                # явный запрос — не молча откатываться к «оформить»
                from formatters.ai_handoff import open_cursor_with_task, write_cursor_task

                prompt = write_cursor_task(
                    source_path=input_path,
                    sample_path=example_path,
                    doc_type=doc_type,
                    goal=(
                        "Перестроить положение о производственном контроле МКТС "
                        "по структуре образца П.ЦЭМ 10-02-2023. "
                        f"Локальная перестройка не удалась: {e}"
                    ),
                )
                ok, msg = open_cursor_with_task(prompt)
                raise RuntimeError(
                    "Автоматическая перестройка не удалась — задание передано в Cursor.\n\n"
                    + msg
                ) from e

    out, out_note = pick_output_path(input_path, work_path)
    margins = dict(cfg.get("margins_mm", {"left": 30, "right": 15, "top": 20, "bottom": 20}))
    font_name = cfg.get("font_name", "Times New Roman")
    font_size = int(cfg.get("font_size_pt", 14))
    first_indent = float(cfg.get("first_indent_mm", 12.5))

    # Для РИ в ваших примерах правое поле часто 12,5 мм
    if doc_type == "rabochaya_instrukciya":
        margins["right"] = 12.5
    # Документы СНиОТ этого агента: 30/10/20/20 (Минюст №65 п.18 допускает правое ≥8 мм).
    if is_sniot_doc(Path(input_path)):
        margins["left"] = 30
        margins["right"] = 10
        margins["top"] = 20
        margins["bottom"] = 20

    summary = {
        "input": input_path,
        "work": work_path,
        "output": out,
        "type": doc_type,
        "example": example_path,
        "actions": [],
        "conservative_di_satp": conservative_di,
    }
    if out_note:
        summary["actions"].append(out_note)

    meta = DOCUMENT_TYPES.get(doc_type, DOCUMENT_TYPES["unsupported"])

    try:
        if doc_type == "prikaz" and not use_basic_only:
            from formatters import prikaz_builder

            prikaz_builder.Settings.DIRECTOR_NAME = cfg.get("director_name", "А.А.Вирочкин")
            prikaz_builder.Settings.DIRECTOR_TITLE = cfg.get("director_title", "Директор")
            prikaz_builder.format_prikaz(work_path, out)
            summary["actions"].append("Приказ пересобран по образцам папки «Приказы»")
            summary["mode"] = "prikaz_rebuild"
            return finish(summary)

        # ДИ / РИ / Положение / ОТ — оформление + текстовые правки по вашей практике
        text_types = {
            "rabochaya_instrukciya",
            "dolzhnostnaya_instrukciya",
            "polozhenie",
            "instrukciya_ot",
        }
        if apply_text_edits_flag and doc_type in text_types:
            from formatters.text_edits import apply_text_edits

            edit_report = apply_text_edits(
                work_path,
                out,
                also_basic_format=True,
                apply_name_updates=True,
                apply_list_markers=False,
                doc_type=doc_type,
            )
            summary["actions"].extend(meta.get("notes", []))
            if edit_report.get("staff_category"):
                from formatters.normative_docs_policy import category_label_ru

                summary["actions"].append(
                    "Категория персонала: " + category_label_ru(edit_report["staff_category"])
                )
            summary["actions"].append(
                f"Текстовые правки по вашей практике: всего {edit_report.get('total_edits', 0)} "
                f"(замен {edit_report['replacements']}, кодов ТКП {edit_report['tkp_codes']}, "
                f"удалений абзацев {edit_report['deleted_paragraphs']})"
            )
            for d in edit_report.get("details", [])[:8]:
                summary["actions"].append(d)
            if example_path:
                summary["actions"].append(f"Ориентир по образцу: {os.path.basename(example_path)}")
            summary["mode"] = "basic_plus_text_edits"
            summary["edit_report"] = edit_report
            return finish(summary)

        # прочее / без текстовых правок — только базовое оформление
        result = apply_basic_office_format(
            work_path,
            out,
            margins=margins,
            font_name=font_name,
            font_size_pt=font_size,
            first_indent_mm=first_indent,
            right_margin_mm=margins["right"],
        )
        summary["actions"].extend(meta.get("notes", []))
        summary["actions"].append(
            f"Применено базовое оформление: абзацев {result['paragraphs_touched']}, "
            f"фрагментов текста {result['runs_touched']}"
        )
        if example_path:
            summary["actions"].append(f"Ориентир по образцу: {os.path.basename(example_path)}")
        summary["mode"] = result["mode"]
        summary.update(result)
        return finish(summary)
    except (PermissionError, OSError) as e:
        # повтор в папке Агент (с timestamp), не на Desktop
        if isinstance(e, PermissionError) or getattr(e, "errno", None) == 13 or "Permission denied" in str(e):
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            retry_candidates: list[str] = []
            if is_path_in_writable_user_dir(input_path):
                same_dir = Path(input_path).parent
                stem = _clean_output_stem(Path(work_path).stem)
                retry_candidates.extend(
                    [
                        str(same_dir / f"{stem}_оформлен_{stamp}.docx"),
                        str(same_dir / f"{Path(work_path).stem}_оформлен_{stamp}.docx"),
                    ]
                )
            if _is_desktop_path(input_path):
                desk = Path.home() / "Desktop"
                retry_candidates.append(
                    str(desk / f"{Path(work_path).stem}_оформлен_{stamp}.docx")
                )
            alt = None
            for candidate in retry_candidates:
                if _can_write_file(candidate):
                    alt = candidate
                    break
            if alt is None:
                raise friendly_save_error(e, out) from e
            log(f"SAVE retry -> {alt} after {e}")
            try:
                if doc_type == "prikaz" and not use_basic_only:
                    from formatters import prikaz_builder

                    prikaz_builder.format_prikaz(work_path, alt)
                    summary["output"] = alt
                    summary["actions"].append(
                        "Исходная папка была занята (часто файл открыт в Word). "
                        f"Результат сохранён:\n{alt}"
                    )
                    summary["mode"] = "prikaz_rebuild"
                    return finish(summary)
                text_types = {
                    "rabochaya_instrukciya",
                    "dolzhnostnaya_instrukciya",
                    "polozhenie",
                    "instrukciya_ot",
                }
                if apply_text_edits_flag and doc_type in text_types:
                    from formatters.text_edits import apply_text_edits

                    edit_report = apply_text_edits(
                        work_path,
                        alt,
                        also_basic_format=True,
                        apply_name_updates=True,
                        apply_list_markers=False,
                        doc_type=doc_type,
                    )
                    summary["output"] = alt
                    summary["actions"].append(
                        "Исходная папка была занята. Результат сохранён:\n" + alt
                    )
                    summary["mode"] = "basic_plus_text_edits"
                    summary["edit_report"] = edit_report
                    return finish(summary)
                result = apply_basic_office_format(
                    work_path,
                    alt,
                    margins=margins,
                    font_name=font_name,
                    font_size_pt=font_size,
                    first_indent_mm=first_indent,
                    right_margin_mm=margins["right"],
                )
                summary["output"] = alt
                summary["actions"].append(
                    "Исходная папка была занята. Результат на Рабочем столе:\n" + alt
                )
                summary["mode"] = result["mode"]
                summary.update(result)
                return finish(summary)
            except Exception as e2:
                raise friendly_save_error(e2, alt) from e2
        raise friendly_save_error(e, out) from e
    except Exception as e:
        log(f"PROCESS fail: {type(e).__name__}: {e!r}")
        msg = str(e).strip()
        if not msg or msg.lower() == "none":
            msg = f"{type(e).__name__}: сбой при оформлении документа"
        raise RuntimeError(f"Ошибка оформления: {msg}") from e
