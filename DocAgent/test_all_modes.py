# -*- coding: utf-8 -*-
"""Проверка агента во всех режимах (без GUI mainloop)."""

from __future__ import annotations

import os
import sys
import traceback
from pathlib import Path

ROOT = Path(__file__).resolve().parent
os.chdir(ROOT)
sys.path.insert(0, str(ROOT))

results: list[tuple[str, str, str]] = []


def ok(name: str, detail: str = "") -> None:
    results.append(("OK", name, detail))
    print(f"OK | {name} | {detail}")


def fail(name: str, err: object) -> None:
    results.append(("FAIL", name, str(err)))
    print(f"FAIL | {name} | {err}")


def main() -> int:
    from agent_core import detect_document, find_examples, process_document, load_config
    from rules import DOCUMENT_TYPES, detect_type_from_text
    from learner import get_watcher, record_agent_action, watch_roots
    from formatters.text_edits import load_patterns
    from docx import Document

    ok("imports")

    cfg = load_config()
    assert cfg.get("continuous_learning") is True
    assert cfg.get("allow_web_templates") is False
    ok("config", f"types={len(cfg.get('supported_types', []))} learning=on")

    samples = {
        "prikaz": ("Приказ.docx", "ПРИКАЗ\nО создании комиссии\nПРИКАЗЫВАЮ:\n1. Создать комиссию."),
        "rabochaya_instrukciya": ("РИ_слесарь.docx", "РАБОЧАЯ ИНСТРУКЦИЯ\nСлесаря\n1. Общие положения"),
        "dolzhnostnaya_instrukciya": (
            "ДИ_инженер.docx",
            "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ\nИнженера по охране труда\n1. Общие положения\nДолжен знать:",
        ),
        "polozhenie": ("Положение.docx", "ПОЛОЖЕНИЕ\nо службе надёжности\n1. Общие положения"),
        "instrukciya_ot": (
            "ИОТ.docx",
            "ИНСТРУКЦИЯ ПО ОХРАНЕ ТРУДА\nпри работе на высоте\n1. Общие требования охраны труда",
        ),
        "unsupported": ("письмо.docx", "Уважаемые коллеги, направляю информацию."),
    }
    for expected, (fname, text) in samples.items():
        got = detect_type_from_text(fname, text)
        if got == expected:
            ok(f"detect:{expected}")
        else:
            fail(f"detect:{expected}", f"got={got}")

    for t in DOCUMENT_TYPES:
        if t == "unsupported":
            continue
        try:
            ex = find_examples(t)
            ok(f"examples:{t}", f"count={len(ex)}")
        except Exception as e:
            fail(f"examples:{t}", e)

    w = get_watcher()
    roots = watch_roots()
    ok("learner", f"{w.status_text()} roots={len(roots)}")
    record_agent_action("test_all_modes", {"phase": "unit"})
    ok("record_agent_action")

    test_dir = ROOT / "learning" / "test_docs"
    test_dir.mkdir(parents=True, exist_ok=True)

    path = test_dir / "test_prikaz.docx"
    doc = Document()
    doc.add_paragraph("ПРИКАЗ")
    doc.add_paragraph("О тестовой комиссии")
    doc.add_paragraph("ПРИКАЗЫВАЮ:")
    doc.add_paragraph("- Создать комиссию.")
    doc.add_paragraph("Директор          А.А.Вирочкин")
    doc.save(path)
    ok("create_test_docx", str(path))

    det = detect_document(str(path))
    dtype_det = det[0] if isinstance(det, tuple) else det.get("type")
    ok("detect_document", f"type={dtype_det}")

    r1 = process_document(str(path), "prikaz", example_path=None, apply_text_edits_flag=False)
    ok("process:no_text_edits", os.path.basename(r1.get("output", "")))

    r2 = process_document(str(path), "prikaz", example_path=None, apply_text_edits_flag=True)
    ok("process:with_text_edits", os.path.basename(r2.get("output", "")))

    path2 = test_dir / "test_letter.docx"
    d2 = Document()
    d2.add_paragraph("Уважаемые коллеги, направляю информацию.")
    d2.save(path2)
    r3 = process_document(str(path2), "unsupported", example_path=None, apply_text_edits_flag=True)
    ok("process:unsupported_basic", os.path.basename(r3.get("output", "")))

    # режимы других типов — базовая обработка без полного rewrite
    for dtype, title_line in [
        ("rabochaya_instrukciya", "РАБОЧАЯ ИНСТРУКЦИЯ\nСлесаря\n1. Общие положения"),
        ("dolzhnostnaya_instrukciya", "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ\nИнженера\nДолжен знать:\nТрудовой кодекс"),
        ("polozhenie", "ПОЛОЖЕНИЕ\nо службе\n1. Общие положения"),
        ("instrukciya_ot", "ИНСТРУКЦИЯ ПО ОХРАНЕ ТРУДА\nпри работах\n1. Общие требования"),
    ]:
        p = test_dir / f"test_{dtype}.docx"
        d = Document()
        for line in title_line.split("\n"):
            d.add_paragraph(line)
        d.save(p)
        try:
            rr = process_document(str(p), dtype, example_path=None, apply_text_edits_flag=True)
            ok(f"process:{dtype}", os.path.basename(rr.get("output", "")))
        except Exception as e:
            fail(f"process:{dtype}", e)
            traceback.print_exc()

    data = load_patterns()
    ok("learned_patterns", f"sections={list(data.keys())[:8]}")

    try:
        import tkinter as tk
        from agent_gui import AgentWizard  # noqa: F401
        from avatar_widget import FloatingAvatar  # noqa: F401

        root = tk.Tk()
        root.withdraw()
        root.destroy()
        ok("gui_import")
    except Exception as e:
        fail("gui_import", e)

    # remote_guard (если уже есть)
    try:
        from remote_guard import RemoteAccessGuard, scan_once

        snap = scan_once()
        ok("remote_guard_scan", f"remote={snap.get('remote_active')} tools={len(snap.get('remote_tools', []))}")
        g = RemoteAccessGuard()
        ok("remote_guard_init", g.status_text())
    except ImportError:
        ok("remote_guard", "not installed yet")
    except Exception as e:
        fail("remote_guard", e)

    print("---SUMMARY---")
    fails = [r for r in results if r[0] == "FAIL"]
    print(f"Total={len(results)} OK={len(results) - len(fails)} FAIL={len(fails)}")
    for _, n, d in fails:
        print("FAIL", n, d)
    return 1 if fails else 0


if __name__ == "__main__":
    raise SystemExit(main())
