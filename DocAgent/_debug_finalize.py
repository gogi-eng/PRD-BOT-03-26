# -*- coding: utf-8 -*-
import sys
import traceback
from pathlib import Path

sys.path.insert(0, r"C:\Users\v.dubovik\DocAgent")

from docx import Document
from docx.shared import Cm, Pt

from formatters.common import apply_basic_office_format, save_docx_unprotected
from formatters.structure_fix import (
    _is_chapter_heading_para,
    _paragraph_has_page_break_before,
    finalize_notes_and_signatories,
)

doc = Document()
doc.add_paragraph("Государственное предприятие")
t = doc.add_table(1, 2)
t.cell(0, 0).text = "ИНСТРУКЦИЯ"
t.cell(0, 1).text = "УТВЕРЖДАЮ"
doc.add_paragraph("Минск 2026")
doc.add_paragraph("СОДЕРЖАНИЕ")
doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ ............ 3")
doc.add_paragraph("2 ОБЯЗАННОСТИ ................ 5")
doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
p = doc.add_paragraph("1.1. Первый пункт.")
p.paragraph_format.first_line_indent = Cm(0)
doc.add_paragraph("1.2. Второй.")
doc.add_paragraph("2 ОБЯЗАННОСТИ")
doc.add_paragraph("2.1. Третий.")
bt = doc.add_table(2, 2)
bt.cell(0, 0).text = "A"
bt.cell(0, 1).text = "B"
for cell in bt.rows[0].cells:
    for r in cell.paragraphs[0].runs:
        r.font.size = Pt(14)
doc.add_paragraph("Разработал:")
doc.add_paragraph(
    "Первый заместитель директора – главный инженер\tА.А.Вирочкин"
)

out = str(Path.home() / "Desktop" / "DocAgent_тесты" / "debug_finalize.docx")
Path(out).parent.mkdir(exist_ok=True)
save_docx_unprotected(doc, out)

print(
    "chapter detect:",
    _is_chapter_heading_para("1 ОБЩИЕ ПОЛОЖЕНИЯ"),
    _is_chapter_heading_para("2 ОБЯЗАННОСТИ"),
)

apply_basic_office_format(out, out)
try:
    rep = finalize_notes_and_signatories(out, "polozhenie")
    print("FINALIZE OK")
    keys = [
        "indents_fixed",
        "chapter_spacers",
        "chapters_centered",
        "contents_page",
        "table_fonts",
        "double_spaces_fixed",
        "signatories",
        "numbers_fixed",
        "etalon_styles",
        "indents_error",
        "spacing_error",
        "contents_page_error",
        "chapters_error",
        "signatories_error",
        "numbering_error",
    ]
    for k in keys:
        if k in rep:
            print(k, ":", rep.get(k))
except Exception:
    traceback.print_exc()

doc2 = Document(out)
print("---RESULT---")
for i, p in enumerate(doc2.paragraphs):
    ind = p.paragraph_format.first_line_indent
    tw = ind.twips if ind else 0
    al = p.alignment
    br = _paragraph_has_page_break_before(p)
    st = p.style.name if p.style else None
    print(f"{i:02d} tw={tw} al={al} br={br} style={st} | {p.text[:70]!r}")
if len(doc2.tables) > 1:
    r = doc2.tables[1].cell(0, 0).paragraphs[0].runs
    print("table1 font", r[0].font.size.pt if r and r[0].font.size else None)
