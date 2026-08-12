# -*- coding: utf-8 -*-
"""Проверка всех новых алгоритмов DocAgent на одном тестовом РИ."""
from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from formatters.common import save_docx_unprotected
from formatters.normative_docs_policy import (
    is_directives_header_line,
    is_orphan_presidential_directive_line,
)
from formatters.structure_fix import (
    apply_signatory_block_format,
    finalize_notes_and_signatories,
    replace_soft_line_breaks_with_paragraphs,
    replace_tabs_with_single_space,
    resolve_instruction_doc_type,
)
from formatters.text_edits import apply_text_edits
from formatters.publish_check import verify_document_before_publish


def _add_soft_break_para(doc: Document, left: str, right: str) -> None:
    p = doc.add_paragraph()
    p.add_run(left)
    br_run = p.add_run()
    br_run._r.append(OxmlElement("w:br"))
    p.add_run(right)


def _add_tab_body(doc: Document, left: str, right: str) -> None:
    p = doc.add_paragraph()
    p.add_run(left)
    r = p.add_run()
    r._r.append(OxmlElement("w:tab"))
    p.add_run(right)


def build_sample_ri(path: Path) -> None:
    doc = Document()
    # титул ТОЛЬКО в таблице (типичный случай — paragraphs не видят «рабочая»)
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "РАБОЧАЯ ИНСТРУКЦИЯ"
    table.cell(0, 1).text = "УТВЕРЖДАЮ"
    doc.add_paragraph("оператору котельной")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.1. Оператор должен знать:")
    doc.add_paragraph("Директивы Президента Республики Беларусь:")
    doc.add_paragraph(
        "от 27 декабря 2006 г. № 2 «О дебюрократизации государственного аппарата "
        "и повышении качества обеспечения жизнедеятельности населения»;"
    )
    doc.add_paragraph(
        "от 14.06.2007 № 3 «О приоритетных направлениях укрепления экономической "
        "безопасности государства»;"
    )
    doc.add_paragraph("локальные правовые акты предприятия;")
    # разрыв строки в теле
    _add_soft_break_para(doc, "1.2. Первая часть", "вторая часть того же пункта.")
    # таб в теле (должен стать пробелом)
    _add_tab_body(doc, "1.3. Текст с", "табом в середине.")
    doc.add_paragraph("Примечание. Этот абзац должен уехать перед Разработал.")
    doc.add_paragraph("4 ОТВЕТСТВЕННОСТЬ")
    doc.add_paragraph("4.1. за неисполнение обязанностей.")
    # без «Разработал:» — финал должен добавить
    doc.add_paragraph("Начальник РТС-4\tВ.Н.Ромашкевич")
    doc.add_paragraph("СОГЛАСОВАНО")
    doc.add_paragraph("Начальник СНиОТ Д.А.Миронов")
    doc.add_paragraph("Начальник ЮС\tГ.А.Авраменко")
    save_docx_unprotected(doc, str(path))


def assert_true(cond: bool, msg: str) -> None:
    if not cond:
        raise AssertionError(msg)


def main() -> int:
    out_dir = Path.home() / "Desktop" / "DocAgent_тесты"
    out_dir.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%H%M%S")
    src = out_dir / f"тест_РИ_алгоритмы_{stamp}.docx"
    dst = out_dir / f"тест_РИ_алгоритмы_оформлен_{stamp}.docx"
    report_path = out_dir / f"тест_РИ_отчёт_{stamp}.txt"

    build_sample_ri(src)

    # --- unit checks ---
    assert_true(
        is_orphan_presidential_directive_line(
            "от 27 декабря 2006 г. № 2 «О дебюрократизации государственного аппарата»;"
        ),
        "orphan directive detect",
    )
    assert_true(
        is_directives_header_line("Директивы Президента Республики Беларусь:"),
        "directives header detect",
    )

    doc0 = Document(str(src))
    dtype0 = resolve_instruction_doc_type(doc_type="", doc=doc0, source_path=str(src))
    assert_true(dtype0 == "rabochaya_instrukciya", f"resolve type from table, got {dtype0}")

    # soft break unit
    doc_s = Document()
    _add_soft_break_para(doc_s, "AAA", "BBB")
    n = replace_soft_line_breaks_with_paragraphs(doc_s)
    assert_true(n >= 1 and [p.text for p in doc_s.paragraphs] == ["AAA", "BBB"], "soft break")

    # tab body → space, signatory tabs keep
    doc_t = Document()
    _add_tab_body(doc_t, "тело", "таб")
    p = doc_t.add_paragraph("Начальник СНиОТ\tД.А. Миронов")
    n_tab = replace_tabs_with_single_space(doc_t)
    assert_true(" " in doc_t.paragraphs[0].text and "\t" not in doc_t.paragraphs[0].text, "body tab")
    assert_true("\t" in doc_t.paragraphs[1].text, "sign tab kept")

    # --- full apply_text_edits pipeline ---
    rep = apply_text_edits(
        str(src),
        str(dst),
        also_basic_format=True,
        doc_type="rabochaya_instrukciya",
    )
    fin = finalize_notes_and_signatories(str(dst), "rabochaya_instrukciya")
    chk = verify_document_before_publish(str(dst), doc_type="rabochaya_instrukciya")

    doc = Document(str(dst))
    texts = [(p.text or "") for p in doc.paragraphs]
    joined = "\n".join(texts)

    lines = [
        "=== ОТЧЁТ ТЕСТА АЛГОРИТМОВ ===",
        f"source: {src}",
        f"output: {dst}",
        f"dtype resolve: {dtype0}",
        f"edit total: {rep.get('total_edits')}",
        f"structure: {rep.get('structure')}",
        f"finalize: {fin}",
        f"publish_check fixed: {chk.get('fixed')} ok={chk.get('ok')}",
        "",
        "--- детали правок ---",
        *[str(x) for x in (rep.get("details") or [])],
        "",
        "--- publish details ---",
        *[str(x) for x in (chk.get("details") or [])],
        "",
        "--- хвост документа ---",
    ]
    for t in texts[-15:]:
        lines.append(t.replace("\t", "|"))

    # assertions on result
    errors = []
    if "от 27 декабря 2006" in joined:
        errors.append("НЕ удалён обрывок директивы 27.12.2006")
    if "от 14.06.2007" in joined:
        errors.append("НЕ удалён обрывок директивы 14.06.2007")
    if "Директивы Президента Республики Беларусь:" in joined:
        errors.append("НЕ удалён заголовок Директивы Президента")
    if "Первая часть" in joined and "вторая часть" in joined:
        # должны быть разными абзацами
        idx = [i for i, t in enumerate(texts) if "Первая часть" in t or "вторая часть" in t]
        if len(idx) < 2 and any("Первая часть" in t and "вторая часть" in t for t in texts):
            errors.append("разрыв строки НЕ преобразован в два абзаца")
    if any("Текст с\tтабом" in t or "Текст с\t" in t for t in texts):
        errors.append("таб в теле НЕ заменён на пробел")
    if not any(t.strip().startswith("Разработал") for t in texts):
        errors.append("нет строки Разработал:")
    raz = next(
        (i for i, t in enumerate(texts) if t.strip().startswith("Разработал")), None
    )
    note_idxs = [
        i for i, t in enumerate(texts) if t.strip().lower().startswith("примечание")
    ]
    if not note_idxs:
        errors.append("примечание исчезло")
    elif raz is not None and any(i > raz for i in note_idxs):
        errors.append("примечание после Разработал")
    if not any(t.strip() == "СОГЛАСОВАНО" for t in texts):
        errors.append("нет заголовка СОГЛАСОВАНО (заглавными)")
    if not any("В.Н.Ромашкевич" in t for t in texts):
        errors.append("ФИО не в формате «В.Н.Ромашкевич» (без пробела)")
    sog = next((i for i, t in enumerate(texts) if t.strip() == "СОГЛАСОВАНО"), None)
    if sog is not None and sog > 0 and texts[sog - 1].strip() != "":
        errors.append("нет пустой строки перед СОГЛАСОВАНО")
    # таб-стоп ≥12 см у ФИО (эталон)
    from docx.oxml.ns import qn as _qn

    sign_p = next((p for p in doc.paragraphs if "Ромашкевич" in (p.text or "")), None)
    if sign_p is None:
        errors.append("нет строки с Ромашкевич")
    else:
        pPr = sign_p._p.find(_qn("w:pPr"))
        tabs = pPr.find(_qn("w:tabs")) if pPr is not None else None
        if tabs is None:
            errors.append("нет таба 12 см у ФИО")
        else:
            pos = int(tabs.find(_qn("w:tab")).get(_qn("w:pos")))
            if pos / 567 < 11.5:
                errors.append(f"таб ФИО слишком мал: {pos / 567:.2f} см")
    # отступ 1.25 на теле
    from formatters.common import FIRST_INDENT_CM

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t.startswith("1.1.") and not t.startswith("4.1."):
            continue
        ind = p.paragraph_format.first_line_indent
        cm = ind.cm if ind else 0
        if abs(cm - FIRST_INDENT_CM) > 0.08:
            errors.append(f"отступ «{t[:30]}» = {cm}, нужно {FIRST_INDENT_CM}")

    lines.append("")
    if errors:
        lines.append("РЕЗУЛЬТАТ: ОШИБКИ")
        lines.extend("- " + e for e in errors)
    else:
        lines.append("РЕЗУЛЬТАТ: ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ")

    report_path.write_text("\n".join(lines), encoding="utf-8")
    print("\n".join(lines[-40:]))
    if errors:
        print("FAILED", len(errors))
        return 1
    print("OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
