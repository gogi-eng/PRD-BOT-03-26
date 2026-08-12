# -*- coding: utf-8 -*-
"""Полный тест алгоритмов финального оформления инструкции."""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn

from formatters.common import FIRST_INDENT_CM, save_docx_unprotected
from formatters.instruction_catalog import (
    apply_instruction_catalog_to_text,
    canonical_instruction_list_name,
    load_catalog,
)
from formatters.structure_fix import (
    finalize_notes_and_signatories,
    ensure_razrabotal_heading,
    move_notes_before_razrabotal,
)


def main() -> int:
    errors: list[str] = []
    out_dir = Path.home() / "Desktop" / "DocAgent_тесты"
    out_dir.mkdir(exist_ok=True)
    report_lines: list[str] = ["=== ТЕСТ ФИНАЛЬНЫХ АЛГОРИТМОВ ===", ""]

    # --- 1) названия без номера и ИОТ ---
    cat = load_catalog()
    sample = (
        "ИОТ 116 «При транспортировке и хранению баллонов» ОТ При транспортировке "
        "и хранении баллонов со сжатым и сжиженным газом;"
    )
    out, notes = apply_instruction_catalog_to_text(sample, cat)
    report_lines.append(f"IOT in:  {sample}")
    report_lines.append(f"IOT out: {out}")
    if "ИОТ" in out.upper().replace("ИНСТРУКЦИЯ", ""):
        # слово ИОТ как аббревиатура не должно остаться
        if re_has_iot_abbrev(out):
            errors.append(f"осталось сокращение ИОТ: {out}")
    if "№" in out or re_has_num_ref(out):
        errors.append(f"в перечне остался номер: {out}")
    if "Инструкция по охране труда" not in out:
        errors.append(f"нет полного наименования: {out}")
    if "ОТ При" in out:
        errors.append(f"остался хвост ОТ …: {out}")
    info = cat.get("by_number", {}).get("116")
    if info:
        canon = canonical_instruction_list_name(info)
        report_lines.append(f"canon 116: {canon}")
        if "№" in canon:
            errors.append(f"canon содержит номер: {canon}")

    # --- 2) документ без «Разработал:», с примечанием, без отступа ---
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "РАБОЧАЯ ИНСТРУКЦИЯ"
    t.cell(0, 1).text = "УТВЕРЖДАЮ"
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    p11 = doc.add_paragraph("1.1. Оператор выполняет работы.")
    p11.paragraph_format.first_line_indent = Cm(0)
    doc.add_paragraph(
        "ИОТ 116 «При транспортировке и хранению баллонов» ОТ При транспортировке "
        "и хранении баллонов со сжатым и сжиженным газом;"
    )
    doc.add_paragraph("Примечание. Этот абзац должен уехать перед Разработал.")
    doc.add_paragraph("1.2. Ещё один пункт тела.")
    # НЕТ строки Разработал — только ФИО
    doc.add_paragraph("Начальник РТС-4 В.Н.Ромашкевич")
    doc.add_paragraph("СОГЛАСОВАНО")
    doc.add_paragraph("Начальник СНиОТ Д.А.Миронов")

    src = out_dir / "тест_финал_вход.docx"
    dst = out_dir / "тест_финал_выход.docx"
    save_docx_unprotected(doc, str(src))

    # text edits path lite: catalog on paragraphs then finalize
    from formatters.text_edits import apply_text_edits

    rep = apply_text_edits(
        str(src), str(dst), also_basic_format=True, doc_type="rabochaya_instrukciya"
    )
    fin = finalize_notes_and_signatories(str(dst), "rabochaya_instrukciya")
    report_lines.append(f"edit total={rep.get('total_edits')} fin={fin}")

    doc2 = Document(str(dst))
    texts = [(p.text or "") for p in doc2.paragraphs]
    joined = "\n".join(texts)

    # Разработал:
    if not any(t.strip().startswith("Разработал") for t in texts):
        errors.append("нет строки Разработал:")
    else:
        report_lines.append("OK Разработал:")

    # примечание перед Разработал
    raz = next(i for i, t in enumerate(texts) if t.strip().startswith("Разработал"))
    note_idxs = [i for i, t in enumerate(texts) if t.strip().lower().startswith("примечание")]
    if not note_idxs:
        errors.append("примечание исчезло или не найдено")
    elif any(i > raz for i in note_idxs):
        errors.append("примечание после Разработал")
    elif any(
        re_is_body_item(texts[j])
        for i in note_idxs
        for j in range(i + 1, raz)
        if texts[j].strip()
    ):
        errors.append("между примечанием и Разработал есть пункты тела")
    else:
        report_lines.append("OK примечание перед Разработал")

    # отступ 1.25 на пунктах тела
    for p in doc2.paragraphs:
        t = (p.text or "").strip()
        if not t.startswith("1.1.") and not t.startswith("1.2."):
            continue
        ind = p.paragraph_format.first_line_indent
        cm = ind.cm if ind else 0
        if abs(cm - FIRST_INDENT_CM) > 0.08:
            errors.append(f"отступ пункта «{t[:40]}» = {cm}, нужно {FIRST_INDENT_CM}")
        else:
            report_lines.append(f"OK отступ {cm:.2f} см для {t[:40]}")

    # заголовок без 1.25
    for p in doc2.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("1 ОБЩИЕ"):
            ind = p.paragraph_format.first_line_indent
            cm = ind.cm if ind else 0
            if cm > 0.2:
                errors.append(f"у заголовка отступ {cm}")
            break

    # подписанты
    if not any(t.strip() == "СОГЛАСОВАНО" for t in texts):
        errors.append("нет СОГЛАСОВАНО")
    if not any("В.Н. Ромашкевич" in t or "В.Н.Ромашкевич" in t for t in texts):
        errors.append("нет ФИО разработчика")
    # таб 12 см
    for p in doc2.paragraphs:
        if "Ромашкевич" not in (p.text or ""):
            continue
        pPr = p._p.find(qn("w:pPr"))
        tabs = pPr.find(qn("w:tabs")) if pPr is not None else None
        if tabs is None:
            errors.append("нет таба 12 см у ФИО")
        else:
            pos = int(tabs.find(qn("w:tab")).get(qn("w:pos")))
            if pos / 567 < 11.5:
                errors.append(f"таб ФИО слишком мал: {pos/567:.2f} см")
            else:
                report_lines.append(f"OK таб ФИО {pos/567:.2f} см")
        break

    # ИОТ в теле заменён
    if re_has_iot_abbrev(joined):
        errors.append("в документе осталось сокращение ИОТ")
    if "ОТ При транспортировке" in joined:
        errors.append("остался хвост ОТ При…")

    report_lines.append("")
    report_lines.append("--- хвост ---")
    for t in texts[-12:]:
        report_lines.append(t.replace("\t", "|"))

    report_lines.append("")
    if errors:
        report_lines.append("РЕЗУЛЬТАТ: ОШИБКИ")
        report_lines.extend("- " + e for e in errors)
    else:
        report_lines.append("РЕЗУЛЬТАТ: ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ")

    rep_path = out_dir / "тест_финал_отчёт.txt"
    rep_path.write_text("\n".join(report_lines), encoding="utf-8")
    print("\n".join(report_lines[-35:]))
    return 1 if errors else 0


def re_has_iot_abbrev(text: str) -> bool:
    import re

    return bool(re.search(r"(?i)(?<![А-Яа-яЁёA-Za-z])ИОТ(?![А-Яа-яЁёA-Za-z])", text))


def re_has_num_ref(text: str) -> bool:
    import re

    return bool(re.search(r"№\s*\d+", text))


def re_is_body_item(text: str) -> bool:
    import re

    t = (text or "").strip()
    return bool(re.match(r"^\d+\.\d+(\.\d+)*\.\s", t))


if __name__ == "__main__":
    raise SystemExit(main())
