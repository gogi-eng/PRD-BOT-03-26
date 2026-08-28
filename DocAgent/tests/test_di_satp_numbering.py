# -*- coding: utf-8 -*-
"""Тесты сохранения нумерации ДИ САТП «Старший мастер»."""

from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

FIX_SCRIPT = Path(r"C:\Users\v.dubovik\AttestationSync\fix_sniot_document.py")
DUMP = Path(r"C:\Users\v.dubovik\AttestationSync\_work_senior_master_fixed.txt")


def _load_fix():
    spec = importlib.util.spec_from_file_location("fix_sniot_document_test", FIX_SCRIPT)
    assert spec and spec.loader
    mod = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = mod
    spec.loader.exec_module(mod)
    return mod


@pytest.fixture(scope="module")
def fix_mod():
    if not FIX_SCRIPT.is_file():
        pytest.skip(f"нет {FIX_SCRIPT}")
    return _load_fix()


def test_is_senior_master_di_path(fix_mod):
    assert fix_mod.is_senior_master_di_path("ПРОЕКТ Старший мастер_оформлен.docx")
    assert fix_mod.is_senior_master_di_path(r"N:\...\САТП\ПРОЕКТ Старший мастер.docx")
    assert not fix_mod.is_senior_master_di_path("РИ слесарь.docx")


def test_conservative_di_satp_flag():
    from formatters.sniot_document import is_conservative_di_satp

    assert is_conservative_di_satp(
        r"N:\9\!!!ОБМЕН\САТП\ПРОЕКТ Старший мастер_оформлен.docx",
        "dolzhnostnaya_instrukciya",
    )
    assert is_conservative_di_satp(
        r"N:\Агент\ДИ  инженер  ЛСиМ_2026 (Романовский).doc",
        "dolzhnostnaya_instrukciya",
    )
    assert not is_conservative_di_satp("РИ слесарь.docx", "rabochaya_instrukciya")
    assert not is_conservative_di_satp("приказ.docx", "prikaz")


def test_validate_save_integrity_blocks_numbering_loss(fix_mod):
    doc = Document()
    for t in ("1.4.1. пункт", "1.4.2. пункт", "1.5.1. знать", "1.5.2. знать", "1.5.3. знать"):
        doc.add_paragraph(t)
    profile = fix_mod.DocumentProfile(
        kind="di",
        first_chapter="1 ОБЩИЕ ПОЛОЖЕНИЯ",
        has_signatories=False,
        has_di_satp_numbering=True,
        tail_chapter_idx=None,
    )
    after = Document()
    after.add_paragraph("текст без номеров")
    issues = fix_mod.validate_save_integrity(
        before_nonempty=5,
        before_numbered=5,
        after_doc=after,
        profile=profile,
    )
    assert any("Нумерация исчезла" in x for x in issues)


def test_fix_numbering_adds_prefixes_from_dump(tmp_path, fix_mod):
    if not DUMP.is_file():
        pytest.skip("нет эталонного дампа _work_senior_master_fixed.txt")

    # минимальный docx: глава 1 + подписант
    out = tmp_path / "di_test.docx"
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("Разработал:")
    doc.save(out)

    n = fix_mod.replace_body_from_debug_dump(out, DUMP)
    assert n > 50

    doc2 = Document(out)
    profile = fix_mod.detect_profile(doc2, out)
    assert profile.has_di_satp_numbering
    fix_mod.process_sniot_document(doc2, profile)
    doc2.save(out)

    numbered = fix_mod.count_numbered_paragraphs(doc2)
    assert numbered >= 80
    texts = [p.text for p in doc2.paragraphs]
    assert any(t.startswith("1.8.1.") for t in texts)
    assert any(t.startswith("1.9.1.") for t in texts)
    assert any(t.startswith("2.1.1.") for t in texts)
    assert any(t.startswith("5.1.1.") for t in texts)


def test_center_chapter_headers(tmp_path, fix_mod):
    out = tmp_path / "chapters.docx"
    doc = Document()
    doc.add_paragraph("1 общие положения")
    doc.add_paragraph("текст")
    doc.add_paragraph("2 функции и должностные обязанности")
    doc.save(out)

    doc2 = Document(out)
    profile = fix_mod.detect_profile(doc2, out)
    fix_mod.center_chapter_headers(doc2)

    chapters = fix_mod.find_chapter_header_indices(doc2)
    assert len(chapters) >= 2
    for idx in chapters:
        assert fix_mod.is_paragraph_centered(doc2.paragraphs[idx])
        text = doc2.paragraphs[idx].text.strip()
        assert text.startswith("1 ОБЩИЕ") or text.startswith("2 ФУНКЦИИ")


def test_validate_di_satp_requires_80_numbered(tmp_path, fix_mod):
    out = tmp_path / "ПРОЕКТ Старший мастер_оформлен.docx"
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4.1. один")
    doc.save(out)

    doc2 = Document(out)
    profile = fix_mod.detect_profile(doc2, out)
    profile = fix_mod.DocumentProfile(
        kind=profile.kind,
        first_chapter=profile.first_chapter,
        has_signatories=False,
        has_di_satp_numbering=True,
        tail_chapter_idx=None,
    )
    issues = fix_mod.validate_di_satp_numbering_count(doc2, profile, out)
    assert any("Мало нумерованных" in x for x in issues)


def test_maybe_restore_senior_master_from_dump(tmp_path, fix_mod):
    if not DUMP.is_file():
        pytest.skip("нет эталонного дампа")

    out = tmp_path / "ПРОЕКТ Старший мастер_оформлен.docx"
    doc = Document()
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("текст без номеров")
    doc.add_paragraph("Разработал:")
    doc.save(out)

    n, msg = fix_mod.maybe_restore_senior_master_body(out)
    assert n > 50
    assert "Восстановлено" in msg
    doc2 = Document(out)
    assert fix_mod.count_numbered_paragraphs(doc2) >= 80


def test_text_edits_skips_conservative_di(tmp_path):
    from formatters.text_edits import apply_text_edits

    src = tmp_path / "ПРОЕКТ Старший мастер.docx"
    out = tmp_path / "ПРОЕКТ Старший мастер_оформлен.docx"
    doc = Document()
    doc.add_paragraph("1.4.1. тест")
    doc.save(src)

    rep = apply_text_edits(
        str(src),
        str(out),
        doc_type="dolzhnostnaya_instrukciya",
    )
    assert rep.get("skipped_conservative") is True
    assert rep.get("deleted_paragraphs", 0) == 0


def test_text_edits_skips_any_di(tmp_path):
    from formatters.text_edits import apply_text_edits

    src = tmp_path / "ДИ  инженер  ЛСиМ_2026 (Романовский).docx"
    out = tmp_path / "ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx"
    doc = Document()
    doc.add_paragraph("1.6. Выполняет локальные правовые акты предприятия.")
    doc.save(src)

    rep = apply_text_edits(
        str(src),
        str(out),
        doc_type="dolzhnostnaya_instrukciya",
    )
    assert rep.get("skipped_conservative") is True
    saved = Document(str(out))
    assert "Выполняет локальные правовые акты" in saved.paragraphs[0].text
