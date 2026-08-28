# -*- coding: utf-8 -*-
"""Тесты определения СНиОТ-пути и финального прохода fix_sniot_document."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from formatters.sniot_document import is_sniot_document, should_apply_sniot_pass
from path_resolver import is_sniot_doc, normalize_sniot_path_text
from rules import detect_type_from_text


def test_normalize_latin_i_in_sniot_path():
    bad = r"N:\9 - Служба (СНiОТ)\!!!ОБМЕН\САТП\test.docx"
    fixed = normalize_sniot_path_text(bad)
    assert "СНiОТ" not in fixed
    assert "СНиОТ" in fixed


def test_is_sniot_with_latin_i():
    bad = r"N:\9 - Служба надёжности (СНiОТ)\!!!ОБМЕН\САТП\ПРОЕКТ.docx"
    assert is_sniot_document(bad) is True
    assert is_sniot_doc(Path(bad)) is True


def test_should_apply_for_latin_i_path():
    p = r"N:\9 - Служба (СНiОТ)\САТП\doc.docx"
    assert should_apply_sniot_pass(p, p, "unsupported") is True


def test_detect_proekt_starshiy_master_as_di():
    name = "ПРОЕКТ Старший мастер.docx"
    text = "УТВЕРЖДАЮ\nстарший мастер участка"
    assert detect_type_from_text(name, text) == "dolzhnostnaya_instrukciya"


def test_detect_di_engineer_lsim_not_polozhenie():
    from rules import filename_looks_like_di

    name = "ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx"
    text = (
        "ИНСТРУКЦИЯ\nНастоящая инструкция определяет обязанности инженера ЛСиМ.\n"
        "1 ОБЩИЕ ПОЛОЖЕНИЯ"
    )
    assert filename_looks_like_di(name) is True
    assert detect_type_from_text(name, text) == "dolzhnostnaya_instrukciya"
    slash_name = name
    assert detect_type_from_text(slash_name, "положение о службе") == "dolzhnostnaya_instrukciya"


def test_oformlen_is_not_sample():
    from path_resolver import is_allowed_sample_path, USER_AGENT_DIR

    oformlen = USER_AGENT_DIR / "ДИ  инженер  ЛСиМ_2026 (Романовский)_оформлен.docx"
    assert is_allowed_sample_path(oformlen) is False
    mixed = USER_AGENT_DIR / "что-то_образец_оформлен.docx"
    assert is_allowed_sample_path(mixed) is False


def test_same_file_copy_slash_vs_backslash(tmp_path: Path):
    from path_resolver import copy_file_if_different, paths_are_same_file

    src = tmp_path / "doc_оформлен.docx"
    src.write_bytes(b"PK\x03\x04test")
    fwd = str(src).replace("\\", "/")
    back = str(src)
    assert paths_are_same_file(fwd, back) is True
    assert copy_file_if_different(fwd, back) is False
    other = tmp_path / "other.docx"
    assert copy_file_if_different(src, other) is True
    assert other.read_bytes() == src.read_bytes()


def test_ensure_title_does_not_duplicate_org_header():
    from docx import Document
    from formatters.structure_fix import ensure_title_table_like_sample

    doc = Document()
    doc.add_paragraph("МИНСКИЙ ГОРОДСКИЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ")
    doc.add_paragraph("КОММУНАЛЬНОЕ УНИТАРНОЕ ПРОИЗВОДСТВЕННОЕ")
    doc.add_paragraph("«МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("1 ОБЩИЕ ПОЛОЖЕНИЯ")
    before = [p.text for p in doc.paragraphs]
    n_org = sum(1 for t in before if "МИНСККОММУНТЕПЛОСЕТЬ" in t.upper())
    ensure_title_table_like_sample(doc, "dolzhnostnaya_instrukciya")
    after = [p.text for p in doc.paragraphs]
    n_org_after = sum(1 for t in after if "МИНСККОММУНТЕПЛОСЕТЬ" in t.upper())
    assert n_org_after == n_org


def test_detect_weekly_itog_not_di():
    name = "Отчёт_о_работе_Дубовик_ВВ_27-30.07.2026.docx"
    text = (
        "ГОСУДАРСТВЕННОЕ ПРЕДПРИЯТИЕ «МИНСККОММУНТЕПЛОСЕТЬ»\n"
        "Служба надёжности и охраны труда (СНиОТ)\n"
        "ОТЧЁТ\nо выполненной работе за период\n"
    )
    assert detect_type_from_text(name, text) == "ezhenedelnyy_itog"
    assert detect_type_from_text("ДИ инженер ЛСиМ.docx", text) == "dolzhnostnaya_instrukciya"


def test_should_not_apply_sniot_pass_to_weekly_itog():
    p = r"C:\Users\v.dubovik\Desktop\Еженедельный_итог\Отчёт_о_работе.docx"
    assert should_apply_sniot_pass(p, p, "ezhenedelnyy_itog") is False


def test_learner_rejects_title_and_lsim_poison():
    from learner import is_safe_learned_replacement

    assert (
        is_safe_learned_replacement(
            "КОММУНАЛЬНЫХ ТЕПЛОВЫХ СЕТЕЙ И КОТЕЛЬНЫХ",
            "ТКП 458-2023 «Правила технической эксплуатации теплоустановок и тепловых сетей потребителей».",
        )
        is False
    )
    assert is_safe_learned_replacement("ЛСиМ", "Осим") is False


def test_stdout_filter_collects_script_build():
    prefixes = (
        "Validation:",
        "Сохранено:",
        "СНиОТ:",
        "Скрипт сборка",
        "СНиОТ скрипт сборка",
        "СНиОТ: сборка",
        "Word:",
    )
    line = "СНиОТ: сборка 2026-08-24-agent-projects"
    assert line.startswith(prefixes)


LIVE_BUILD = "2026-08-24-agent-projects"


def test_single_live_build_visible_in_launch_chain():
    """bat → start_agent.py → GUI «Готово» → stdout SCRIPT_BUILD — один номер."""
    agent = Path(r"C:\Users\v.dubovik\DocAgent")
    bat = (agent / "start_agent.bat").read_text(encoding="utf-8")
    start_py = (agent / "start_agent.py").read_text(encoding="utf-8")
    gui = (agent / "agent_gui.py").read_text(encoding="utf-8")
    fmt = (agent / "formatters" / "sniot_document.py").read_text(encoding="utf-8")
    core = (agent / "agent_core.py").read_text(encoding="utf-8")
    script = Path(r"C:\Users\v.dubovik\AttestationSync\fix_sniot_document.py").read_text(
        encoding="utf-8"
    )

    assert 'start "" "%PY%" "%~dp0start_agent.py"' in bat
    assert "Python311\\python.exe" in bat
    assert r"DocAgent\\start_agent\.py" in bat
    assert "Stop-Process" in bat
    assert "taskkill /IM python.exe" not in bat.lower()

    assert f'AGENT_BUILD = "{LIVE_BUILD}"' in start_py
    assert f'SCRIPT_BUILD = "{LIVE_BUILD}"' in script
    assert "MARGIN_RIGHT_MM = 10" in script
    assert "apply_signatory_fio_one_line" in script
    assert f'SNIOT_GUI_BUILD = "{LIVE_BUILD}"' in fmt
    assert "pick_sniot_build_line" in fmt
    assert "read_live_script_build" in fmt
    assert "_parse_script_build_line" in fmt
    assert '"build": build' in fmt
    assert "TITLE_ORG_ALIGN = \"center\"" in script
    assert "TITLE_STAMP_ALIGN = \"right\"" in script
    assert "TITLE_UTVERZHDAYU_BOLD = False" in script
    assert "TITLE_DOC_NAME_ALIGN = \"left\"" in script
    assert "TITLE_CITY_YEAR_TEMPLATE" in script
    assert "TITLE_NUMBER_FONT_PT = 12" in script
    assert "TITLE_NUMBER_LABEL = \"номер инструкции\"" in script
    assert "TITLE_STAMP_DATE_TEMPLATE" in script
    assert "keep_signatory_block_with_text" in script
    assert "_ensure_title_left_right_stamp" in script
    assert "remove_city_year_from_headers_footers" in script
    assert "apply_title_instruction_number_font" in script
    assert "_strip_paragraph_tabs" in script
    assert "ensure_deloproizvodstvo_in_must_know" in script
    assert "fix_duty_by_order_commas" in script
    assert "ensure_section_break_after_city_year" in script
    assert "check_document_punctuation_after_edit" in script
    assert "validate_document_punctuation" in script
    assert '"w:vAnchor", "margin"' in script

    assert "from formatters.sniot_document import pick_sniot_build_line" in gui
    assert "pick_sniot_build_line(result)" in gui
    assert 'f"Готово — {build}"' in gui
    assert "SNIOT_GUI_BUILD" not in gui
    assert "apply_sniot_rules_to_output(out_path)" in core
    assert 'summary["sniot_build"] = sniot["build"]' in core

    from formatters.sniot_document import (
        SNIOT_GUI_BUILD,
        pick_sniot_build_line,
        read_live_script_build,
    )

    assert SNIOT_GUI_BUILD == LIVE_BUILD
    assert read_live_script_build() == LIVE_BUILD
    stale = {
        "actions": ["СНиОТ: сборка 2026-08-17-1025-no-90s-timeout"],
        "sniot_pass": {
            "build": f"СНиОТ: сборка {LIVE_BUILD}",
            "actions": [
                "СНиОТ: сборка 2026-08-17-1025-no-90s-timeout",
                f"СНиОТ: сборка {LIVE_BUILD}",
            ],
        },
    }
    assert pick_sniot_build_line(stale) == f"СНиОТ: сборка {LIVE_BUILD}"


def test_gui_sniot_subprocess_skips_word_and_always_applies():
    src = Path(r"C:\Users\v.dubovik\DocAgent\formatters\sniot_document.py").read_text(
        encoding="utf-8"
    )
    assert 'cmd.append("--always-apply")' in src
    assert 'cmd.append("--skip-word")' in src
    from formatters.sniot_document import SNIOT_FIX_TIMEOUT_SEC, SNIOT_GUI_BUILD, SNIOT_XML_TIMEOUT_SEC

    assert SNIOT_XML_TIMEOUT_SEC >= 15 * 60
    assert SNIOT_FIX_TIMEOUT_SEC >= 15 * 60
    assert "SNIOT_FIX_TIMEOUT_SEC = 90" not in src
    assert "timeout=90" not in src
    assert "таймаут 90 с" not in src
    assert SNIOT_GUI_BUILD == LIVE_BUILD


def test_gui_does_not_show_done_when_sniot_not_applied():
    gui = Path(r"C:\Users\v.dubovik\DocAgent\agent_gui.py").read_text(encoding="utf-8")
    assert 'showerror(f"Правила СНиОТ не применены — {build}"' in gui
    assert "if sniot and not sniot.get(\"applied\")" in gui
    assert "pick_sniot_build_line" in gui


def test_publish_check_skips_di_renumber(tmp_path: Path):
    from docx import Document
    from formatters.publish_check import verify_document_before_publish

    doc = Document()
    doc.add_paragraph("1. ОБЩИЕ ПОЛОЖЕНИЯ")
    doc.add_paragraph("1.4. Настоящая инструкция определяет обязанности инженера ЛСиМ.")
    doc.add_paragraph("1.6. В своей деятельности инженер руководствуется.")
    path = tmp_path / "ДИ  инженер  ЛСиМ_оформлен.docx"
    doc.save(path)
    chk = verify_document_before_publish(
        str(path), doc_type="dolzhnostnaya_instrukciya"
    )
    after = Document(str(path))
    texts = [p.text for p in after.paragraphs]
    assert any("1.4." in t for t in texts)
    assert any("1.6." in t for t in texts)
    assert not any("старт ветки" in i for i in (chk.get("issues") or []))
    assert any("исходник сохранён" in d for d in (chk.get("details") or []))


def test_polish_title_keeps_unique_job_title():
    from docx import Document
    from formatters.structure_fix import _polish_existing_title_table

    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ ИНЖЕНЕРУ ЛАБОРАТОРИИ СВАРКИ И МЕТАЛЛА"
    table.cell(0, 1).text = "УТВЕРЖДАЮ"
    _polish_existing_title_table(table, "dolzhnostnaya_instrukciya")
    left = table.cell(0, 0).text
    assert "ИНЖЕНЕРУ" in left or "ЛАБОРАТОРИИ" in left
    assert left.strip().upper() != "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ"


def test_word_com_has_grammar_and_spelling_entry():
    import inspect

    from formatters.sniot_document import apply_sniot_rules_to_output
    from formatters.word_com import apply_word_grammar_check

    assert callable(apply_word_grammar_check)
    out_src = inspect.getsource(apply_sniot_rules_to_output)
    assert "apply_word_grammar_check" in out_src
    assert "run_sniot_document_fix" in out_src
    src = Path(r"C:\Users\v.dubovik\AttestationSync\fix_sniot_document.py").read_text(
        encoding="utf-8"
    )
    assert "SpellingErrors" in src
    assert "GetSpellingSuggestions" in src
    assert "CheckGrammar" in src
    assert "GrammaticalErrors" in src
    assert "fix_spaces_around_parentheses" in src


def test_gui_xml_skips_word_then_calls_grammar_after_save():
    """XML --skip-word, чтобы не убить запись; грамматика Word — после save."""
    import inspect

    from formatters.sniot_document import apply_sniot_rules_to_output, run_sniot_document_fix

    xml_src = inspect.getsource(run_sniot_document_fix)
    assert 'cmd.append("--skip-word")' in xml_src
    live_src = inspect.getsource(apply_sniot_rules_to_output)
    pos_xml = live_src.find("run_sniot_document_fix")
    pos_gram = live_src.find("apply_word_grammar_check")
    assert pos_xml >= 0 and pos_gram > pos_xml


def test_russian_check_runs_after_sniot_on_all_types():
    import inspect

    from agent_core import process_document

    src = inspect.getsource(process_document)
    assert "apply_russian_check_flag = False" not in src
    pos_sniot = src.rfind("apply_sniot_rules_to_output")
    pos_ru = src.rfind("_apply_russian_language_check")
    assert pos_sniot >= 0 and pos_ru > pos_sniot
