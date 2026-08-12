# -*- coding: utf-8 -*-
"""
Обязательная проверка документа ПЕРЕД выдачей результата (публикацией).

Правило Дубовика: агент всегда проверяет оформленный файл
и по возможности сам исправляет замечания.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from .common import save_docx_unprotected, set_run_font
from .structure_fix import (
    apply_signatory_block_format,
    replace_soft_line_breaks_with_paragraphs,
    replace_tabs_with_single_space,
    resolve_instruction_doc_type,
)

MARKER_RE = re.compile(r"^[\s]*[\-\u2013\u2014\u2022\*◦▪▸►]+")
SIGN_FIO_TAB_CM = 12.0


def verify_document_before_publish(
    docx_path: str,
    *,
    doc_type: str = "",
) -> dict:
    """
    Проверка перед публикацией.
    Возвращает {ok, issues, fixed, details}.
    """
    report = {
        "ok": True,
        "issues": [],
        "fixed": 0,
        "details": [],
        "path": docx_path,
    }
    path = Path(docx_path)
    if not path.exists():
        report["ok"] = False
        report["issues"].append(f"файл не найден: {docx_path}")
        return report

    doc = Document(str(path))
    changed = False

    # --- разрыв строки (Shift+Enter) → обычный перенос абзаца ---
    try:
        n_soft = replace_soft_line_breaks_with_paragraphs(doc)
        if n_soft:
            changed = True
            report["fixed"] += n_soft
            report["details"].append(f"разрыв строки → перенос абзаца ({n_soft})")
    except Exception as e:
        report["details"].append(f"замена разрывов строки: {e}")

    # --- табуляция → один пробел (кроме таба 12 см у подписантов) ---
    try:
        n_tab = replace_tabs_with_single_space(doc)
        if n_tab:
            changed = True
            report["fixed"] += n_tab
            report["details"].append(f"табуляция → пробел ({n_tab})")
    except Exception as e:
        report["details"].append(f"замена табуляции: {e}")

    # --- ОБЯЗАТЕЛЬНО: проверка корректности нумерации перед выдачей ---
    try:
        from .structure_fix import (
            normalize_numbering_style,
            verify_and_fix_numbering,
            collapse_duplicate_number_prefix,
        )

        for p in doc.paragraphs:
            t = p.text or ""
            if not t.strip():
                continue
            new_t = normalize_numbering_style(collapse_duplicate_number_prefix(t))
            if new_t != t:
                _rewrite_para(p, new_t)
                changed = True
                report["fixed"] += 1
                report["details"].append(f"нумерация (стиль): {t[:50]} → {new_t[:50]}")
        chk = verify_and_fix_numbering(doc)
        if chk.get("fixed"):
            changed = True
            report["fixed"] += int(chk["fixed"])
            report["details"].append(
                f"нумерация: исправлено пунктов {chk['fixed']} "
                f"(проверено {chk.get('checked', 0)})"
            )
        # повторная проверка — не оставлять известные сбои
        chk2 = verify_and_fix_numbering(doc)
        if chk2.get("fixed"):
            changed = True
            report["fixed"] += int(chk2["fixed"])
            report["details"].append(f"нумерация (2-й проход): +{chk2['fixed']}")
        left_issues = chk2.get("issues") or chk.get("issues") or []
        if left_issues:
            report["ok"] = False
            report["issues"].append(
                "нумерация: остались замечания — " + "; ".join(left_issues[:5])
            )
        report["numbering_check"] = chk2 if chk2 else chk
    except Exception as e:
        report["ok"] = False
        report["issues"].append(f"нумерация: сбой проверки — {e}")

    # --- маркеры запрещены ---
    for p in doc.paragraphs:
        t = p.text or ""
        if MARKER_RE.match(t):
            new_t = re.sub(r"^[\s]*[\-\u2013\u2014\u2022\*◦▪▸►]+\s*", "", t)
            if new_t != t:
                _rewrite_para(p, new_t)
                changed = True
                report["fixed"] += 1
                report["details"].append(f"убран маркер: {t[:70]}")
            else:
                report["ok"] = False
                report["issues"].append(f"маркер: {t[:70]}")

    # --- подпись / И.О.Фамилия на 12 см (приказы и хвосты с «Директор») ---
    if doc_type == "prikaz" or any(
        (p.text or "").lower().startswith("директор") for p in doc.paragraphs[-8:]
    ):
        target = int(Cm(SIGN_FIO_TAB_CM).twips)
        for p in reversed(list(doc.paragraphs)):
            t = (p.text or "").strip()
            if not t:
                continue
            if t.lower().startswith("директор") or (
                "\t" in (p.text or "") and len(t) < 80
            ):
                _set_left_tab(p, target)
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(0)
                changed = True
                report["fixed"] += 1
                report["details"].append(f"подпись: таб ФИО = {SIGN_FIO_TAB_CM} см")
                break

    # --- комиссия (для приказов) ---
    if doc_type == "prikaz" or "комисси" in " ".join(
        p.text for p in doc.paragraphs[:40]
    ).lower():
        blob = "\n".join(p.text for p in doc.paragraphs)
        if "Заместители председателя" in blob and "Члены комиссии" in blob:
            has_zam1 = bool(
                re.search(
                    r"Заместители председателя комиссии:\s*1\.\s",
                    blob,
                    flags=re.I,
                )
            )
            has_mem1 = bool(
                re.search(r"Члены комиссии:\s*1\.\s", blob, flags=re.I)
            )
            if not has_zam1 or not has_mem1:
                report["ok"] = False
                report["issues"].append(
                    "комиссия: после «Заместители…:» / «Члены…:» сразу в той же "
                    "строке «1. должность И.О.Фамилия» "
                    "(эталон: Приказ_комиссия по расслед._ПРОЕКТ)"
                )

    # --- пустая строка перед заголовком новой главы ---
    try:
        from .structure_fix import ensure_empty_line_before_chapters

        ch = ensure_empty_line_before_chapters(doc)
        if ch:
            changed = True
            report["fixed"] += ch
            report["details"].append(f"пустая строка перед главами: +{ch}")
    except Exception as e:
        report["details"].append(f"интервал глав: {e}")

    # --- титул: «номер инструкции» = 11 пт (эталон КЛ) ---
    try:
        from .structure_fix import apply_title_instruction_number_font

        nfont = apply_title_instruction_number_font(doc)
        if nfont:
            changed = True
            report["fixed"] += nfont
            report["details"].append(f"титул: «номер инструкции» → 11 пт ({nfont})")
    except Exception as e:
        report["details"].append(f"шрифт номера инструкции: {e}")

    # --- РИ: обрывки директив «от … № … «О …»» — удалить ---
    try:
        dtype = (doc_type or "").lower()
        blob_head = "\n".join((p.text or "") for p in doc.paragraphs[:40]).lower()
        is_ri = dtype == "rabochaya_instrukciya" or "рабочая инструкция" in blob_head
        if is_ri:
            from .normative_docs_policy import (
                is_directives_header_line,
                is_orphan_presidential_directive_line,
            )

            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if not t:
                    continue
                if is_orphan_presidential_directive_line(t) or is_directives_header_line(
                    t
                ):
                    _rewrite_para(p, "")
                    changed = True
                    report["fixed"] += 1
                    report["details"].append(f"РИ: удалён обрывок директивы: {t[:80]}")
    except Exception as e:
        report["details"].append(f"проверка обрывков директив: {e}")

    # --- примечания только перед «Разработал:» ---
    try:
        from .structure_fix import move_notes_before_razrabotal

        moved = move_notes_before_razrabotal(doc)
        if moved:
            changed = True
            report["fixed"] += moved
            report["details"].append(
                f"примечания перенесены перед «Разработал:» ({moved} абз.)"
            )
        # контроль: нет «Примечание» в теле до «Разработал» вне хвоста
        texts = [(p.text or "").strip() for p in doc.paragraphs]
        raz = next(
            (
                i
                for i, t in enumerate(texts)
                if t.lower().startswith("разработал")
                or t.lower().startswith("разработчик")
            ),
            None,
        )
        if raz is not None:
            for i, t in enumerate(texts):
                if i >= raz:
                    break
                low = t.lower()
                if re.match(r"^(?:\d+(\.\d+)*\.\s*)?(примечан(ие|ия)|прим\.?)\b", low):
                    # допустимо только в хвосте: после i до raz нет пунктов 1.2. / глав
                    after = texts[i + 1 : raz]
                    if any(
                        re.match(r"^\d+\.\d+(\.\d+)*\.\s", x)
                        or re.match(r"^\d+\s+[А-ЯЁA-Z]", x)
                        for x in after
                        if x
                    ):
                        report["ok"] = False
                        report["issues"].append(
                            "примечание стоит не в конце: должно быть сразу перед «Разработал:»"
                        )
                        break
    except Exception as e:
        report["details"].append(f"проверка примечаний: {e}")

    # --- повторно выровнять подписанты (после таб/разрыв/нумерации/примечаний) ---
    try:
        dtype = resolve_instruction_doc_type(doc_type=doc_type, doc=doc)
        # ещё раз примечания перед подписантами
        from .structure_fix import move_notes_before_razrabotal

        moved2 = move_notes_before_razrabotal(doc)
        if moved2:
            changed = True
            report["fixed"] += moved2
        n_sign = apply_signatory_block_format(doc, doc_type=dtype)
        if n_sign:
            changed = True
            report["fixed"] += n_sign
            report["details"].append(
                f"подписанты по эталону ({dtype}): правок {n_sign}"
            )
    except Exception as e:
        report["details"].append(f"повтор подписантов: {e}")

    # финальный контроль маркеров
    left = [
        (p.text or "")[:50]
        for p in doc.paragraphs
        if MARKER_RE.match(p.text or "")
    ]
    if left:
        report["ok"] = False
        report["issues"].append(f"после проверки остались маркеры: {left[:3]}")

    if changed:
        save_docx_unprotected(doc, str(path))

    if report["issues"]:
        report["ok"] = False
    report["details"].append(
        "проверка перед публикацией: "
        + ("ОК" if report["ok"] and not report["issues"] else "есть замечания")
    )
    return report


def _rewrite_para(paragraph, text: str) -> None:
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        set_run_font(paragraph.runs[0])
    else:
        set_run_font(paragraph.add_run(text))


def _set_left_tab(paragraph, pos_twips: int) -> None:
    pPr = paragraph._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:tabs"))):
        pPr.remove(old)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)
