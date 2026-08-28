# -*- coding: utf-8 -*-
"""
Форматирование проекта приказа по образцам папки:
N:\\9 - Служба надёжности и охраны труда (СНиОТ)\\Дубовик В.В\\Приказы

Что делает скрипт:
1) Берёт текст из вашего черновика .docx
2) Собирает НОВЫЙ файл в виде ваших рабочих приказов
3) Сохраняет рядом с исходником с окончанием _оформлен.docx

Запуск:
  python format_prikaz.py
  python format_prikaz.py "C:\\путь\\к\\черновику.docx"
"""

from __future__ import annotations

import os
import re
import sys
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, Twips

from formatters.common import save_docx_unprotected

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass


# ============================================================
# НАСТРОЙКИ ПО ОБРАЗЦАМ ПРИКАЗОВ МКТС
# ============================================================

class Settings:
    # Поля (как в большинстве примеров: 30 / 15 / 20 / 20 мм)
    MARGIN_LEFT = Cm(3.0)
    MARGIN_RIGHT = Cm(1.5)
    MARGIN_TOP = Cm(2.0)
    MARGIN_BOTTOM = Cm(2.0)

    FONT_NAME = "Times New Roman"
    FONT_SIZE = Pt(14)
    FIRST_INDENT = Mm(12.5)
    # Состав комиссии: метки без абзацного отступа; номера 2+ — под «1.» после метки
    # (эталон N: …Приказ_комиссия по расслед._ПРОЕКТ.docx).
    COMMISSION_ZAM_NUM_LEFT = Cm(6.5)  # «Заместители председателя комиссии: »
    COMMISSION_MEM_NUM_LEFT = Cm(3.8)  # «Члены комиссии: »

    # Место под бланк предприятия сверху (пустые строки)
    BLANK_TOP = 7

    # Подпись: И.О.Фамилия с табуляции 12 см (как ДИ/РИ)
    SIGN_FIO_TAB_CM = 12.0

    # Подпись (в свежих примерах июня 2026 — Вирочкин)
    DIRECTOR_TITLE = "Директор"
    DIRECTOR_NAME = "А.А.Вирочкин"

    CITY = "г.Минск"


# ============================================================
# НИЗКОУРОВНЕВОЕ ФОРМАТИРОВАНИЕ (только допустимые теги Word)
# ============================================================

def set_run_font(run, size=None, bold=None, name=None):
    name = name or Settings.FONT_NAME
    size = size or Settings.FONT_SIZE
    run.font.name = name
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


def set_paragraph_style(
    paragraph,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_indent=None,
    left_indent=None,
    space_before=Pt(0),
    space_after=Pt(0),
):
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = space_before
    pf.space_after = space_after
    pf.first_line_indent = first_indent
    pf.left_indent = left_indent
    pf.widow_control = True


def add_text_paragraph(
    doc,
    text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_indent=None,
    left_indent=None,
    bold=False,
    space_before=Pt(0),
    space_after=Pt(0),
):
    p = doc.add_paragraph()
    set_paragraph_style(
        p,
        alignment=alignment,
        first_indent=first_indent,
        left_indent=left_indent,
        space_before=space_before,
        space_after=space_after,
    )
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_empty(doc, count=1, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    for _ in range(count):
        p = doc.add_paragraph()
        set_paragraph_style(p, alignment=alignment, first_indent=Cm(0))
        set_run_font(p.add_run(""))


def set_cell_border_none(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    # заменить старые границы, если были
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(tcBorders)


def add_title_table(doc, title_text):
    """
    Заголовок «О ...» — как в примерах: таблица 1x2 без рамок,
    текст в левой ячейке, шрифт Times New Roman 14.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # ширина как в примерах (~9360 DXA)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Twips(4500)
    right.width = Twips(4860)

    for cell in (left, right):
        set_cell_border_none(cell)
        for p in cell.paragraphs:
            p.clear()

    p = left.paragraphs[0]
    set_paragraph_style(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Cm(0), space_after=Pt(0))
    run = p.add_run(title_text)
    set_run_font(run)

    p2 = right.paragraphs[0]
    set_paragraph_style(p2, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Cm(0))
    return table


def add_signature(doc, title=None, name=None):
    """Подпись: должность слева, И.О.Фамилия с таба 12 см (как эталон ДИ/РИ)."""
    title = title or Settings.DIRECTOR_TITLE
    name = name or Settings.DIRECTOR_NAME
    name = normalize_commission_fio(name) or name
    add_empty(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_style(
        p,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=Cm(0),
        left_indent=Cm(0),
    )
    # очистить старые табы и поставить ровно 12 см
    pPr = p._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:tabs"))):
        pPr.remove(old)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(int(Cm(Settings.SIGN_FIO_TAB_CM).twips)))
    tabs.append(tab)
    pPr.append(tabs)
    run = p.add_run(f"{title}\t{name}")
    set_run_font(run)
    return p


def configure_page(doc):
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Settings.MARGIN_LEFT
        section.right_margin = Settings.MARGIN_RIGHT
        section.top_margin = Settings.MARGIN_TOP
        section.bottom_margin = Settings.MARGIN_BOTTOM
        section.header_distance = Cm(0)
        section.footer_distance = Cm(1.25)

    style = doc.styles["Normal"]
    style.font.name = Settings.FONT_NAME
    style.font.size = Settings.FONT_SIZE
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # кириллица в стиле
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), Settings.FONT_NAME)


# ============================================================
# РАЗБОР ЧЕРНОВИКА
# ============================================================

def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = text.replace("承", "").replace("压", "")
    # после удаления китайских иероглифов: «элементовных» → «элементов»
    text = text.replace("элементовных", "элементов")
    # убрать артефакты вроде «расследования ,»
    text = re.sub(r" +([,;:.])", r"\1", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" +\n", "\n", text)
    return text.strip()


def extract_raw_text(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return clean_text("\n".join(parts))


def parse_draft(raw: str) -> dict:
    """
    Достаёт из «слепленного» черновика:
    - заголовок «О ...»
    - преамбулу
    - пункты приказа
    - подпись (если есть)
    """
    text = clean_text(raw)

    # Убрать шапку-мусор в начале
    text = re.sub(
        r"^ПРИКАЗ\s+ГОСУДАРСТВЕННОЕ ПРЕДПРИЯТИЕ.*?ПРИКАЗ\s+"
        r"[«\"']?__?[»\"']?\s*_{0,20}\s*202\s*г\.?\s*№\s*_{0,10}\s*"
        r"г\.\s*Минск\s*",
        "",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    text = re.sub(r"^ПРИКАЗ\s+", "", text, flags=re.IGNORECASE)

    # Заголовок
    title_match = re.search(
        r"(О создании постоянно действующей комиссии[^В]{0,200}?оборудования)",
        text,
        flags=re.IGNORECASE,
    )
    title = title_match.group(1).strip() if title_match else (
        "О создании постоянно действующей комиссии по техническому "
        "расследованию нарушений в работе (выходов из строя) "
        "теплотехнического оборудования"
    )

    # Преамбула до ПРИКАЗЫВАЮ
    preamble = ""
    m_pre = re.search(
        r"(В целях обеспечения.*?),\s*ПРИКАЗЫВАЮ\s*:",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if m_pre:
        preamble = clean_text(m_pre.group(1)) + ","
    else:
        # запасной вариант
        m2 = re.search(r"(В целях обеспечения.*?ПРИКАЗЫВАЮ)", text, flags=re.I | re.S)
        if m2:
            preamble = clean_text(re.sub(r",?\s*ПРИКАЗЫВАЮ$", "", m2.group(1), flags=re.I)) + ","

    after = text
    m_after = re.search(r"ПРИКАЗЫВАЮ\s*:\s*(.*)", text, flags=re.IGNORECASE | re.DOTALL)
    if m_after:
        after = m_after.group(1).strip()

    # Подпись в конце
    director_name = Settings.DIRECTOR_NAME
    m_sign = re.search(
        r"Директор\s*[_\s]*([А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁа-яё\-]+)\s*$",
        after,
        flags=re.IGNORECASE,
    )
    if m_sign:
        director_name = m_sign.group(1).replace(" ", "")
        # привести к виду А.А.Фамилия
        director_name = re.sub(r"\.\s+", ".", director_name)
        after = after[: m_sign.start()].strip()

    after = re.sub(r"Директор\s*_{3,}.*$", "", after, flags=re.IGNORECASE).strip()

    points = split_points(after)
    return {
        "title": title,
        "preamble": preamble,
        "points": points,
        "director_name": director_name,
    }


def split_points(body: str) -> list:
    """
    Разбивает текст после ПРИКАЗЫВАЮ на пункты.
    Поддерживает уже пронумерованный и «слипшийся» текст.
    """
    body = clean_text(body)

    # Если пункты уже с номерами 1. 2. 3. ...
    numbered = re.split(r"(?=\b\d+\.\s)", body)
    numbered = [clean_text(x) for x in numbered if clean_text(x)]
    if len(numbered) >= 3 and re.match(r"^\d+\.", numbered[0]):
        return expand_subpoints(numbered)

    # Черновик без номеров — режем по известным зачинам пунктов
    starters = [
        r"Создать постоянно действующую комиссию",
        r"Утвердить персональный состав",
        r"Возложить на Комиссию следующие задачи",
        r"Определить, что основанием",
        r"Порядок расследования",
        r"Начальникам районов тепловых сетей",
        r"Контроль за выполнением",
        r"Плановые заседания Комиссии",
        r"Начальникам РТС и начальникам смен",
    ]
    pattern = "(" + "|".join(starters) + ")"
    parts = re.split(pattern, body)
    chunks = []
    i = 0
    while i < len(parts):
        part = parts[i].strip()
        if not part:
            i += 1
            continue
        if re.match(pattern, part) and i + 1 < len(parts):
            chunks.append(clean_text(part + " " + parts[i + 1]))
            i += 2
        else:
            if not any(chunks) or not re.match(pattern, part):
                # хвост без маркера
                if chunks:
                    chunks[-1] = clean_text(chunks[-1] + " " + part)
                else:
                    chunks.append(part)
            i += 1

    result = []
    for idx, chunk in enumerate(chunks, start=1):
        chunk = re.sub(r"^\d+\.\s*", "", chunk)
        result.append(f"{idx}. {chunk}")
    return expand_subpoints(result)


def expand_subpoints(points: list) -> list:
    """Выделяет подпункты 3.1 / 6.1 из текста пункта на отдельные абзацы."""
    out = []
    for point in points:
        # подпункты вида 3.1. ... 3.2. ...
        parts = re.split(r"(?=\b\d+\.\d+\.\s)", point)
        parts = [clean_text(x) for x in parts if clean_text(x)]
        if len(parts) == 1:
            out.append(parts[0])
            continue
        # первый кусок — заголовок пункта
        out.append(parts[0].rstrip(":").rstrip() + ":")
        for sub in parts[1:]:
            out.append(sub)
    return out


def _comp_line(text: str, style: str) -> dict:
    """Абзац состава комиссии: style = intro|flush|zam_cont|mem_cont."""
    return {"text": text, "style": style}


def format_composition_block(point_text: str) -> list:
    """
    Состав комиссии — как в эталоне «Приказ_комиссия по расслед._ПРОЕКТ»:
    • метки «Председатель / Заместители / Члены / Секретарь» — БЕЗ отступа;
    • после двоеточия сразу (в той же строке) должность и И.О.Фамилия;
    • у замов и членов — отдельная нумерация с 1.; первый номер в строке с меткой;
    • строки 2., 3. … выровнены по уровню номера («1.»).
    """
    low = point_text.lower()
    is_comp = (
        "утвердить персональный состав" in low
        or ("создать" in low and "состав" in low and "комисси" in low)
        or ("в следующем составе" in low and "комисси" in low)
    )
    if not is_comp:
        return [point_text]

    lines: list = []
    # вводная до «Председатель…»
    m = re.match(
        r"^(\d+\.\s*.*?(?:состав|составе)\s*:)\s*(.*)$",
        point_text,
        flags=re.S | re.I,
    )
    if not m:
        m2 = re.search(
            r"(Председатель\s+(?:комиссии|Комиссии)\s*:)",
            point_text,
            flags=re.I,
        )
        if not m2:
            return [point_text]
        head = clean_text(point_text[: m2.start()]).rstrip(":") + ":"
        rest = point_text[m2.start() :]
        lines.append(_comp_line(head, "intro"))
    else:
        lines.append(_comp_line(clean_text(m.group(1)), "intro"))
        rest = m.group(2)

    # Председатель: после «:» сразу должность и ФИО, без отступа.
    m_pred = re.search(
        r"Председатель\s+(?:комиссии|Комиссии)\s*[:\-–—]\s*(.+?)"
        r"(?=\s*Заместители\s+председателя|\s*Члены\s+(?:комиссии|Комиссии)\s*:|\s*Секретар|$)",
        rest,
        flags=re.S | re.I,
    )
    if m_pred:
        pred = clean_text(m_pred.group(1)).rstrip(";.")
        pred = re.sub(r"\s+", " ", pred).strip()
        pred = normalize_commission_fio(pred)
        lines.append(_comp_line(f"Председатель комиссии: {pred}.", "flush"))

    # Заместители: «…комиссии: 1. …» в одной строке; 2+ — с отступом на уровне номера.
    m_zam = re.search(
        r"Заместители\s+председателя\s+(?:комиссии|Комиссии)\s*:\s*(.+?)"
        r"(?=\s*Члены\s+(?:комиссии|Комиссии)\s*:|\s*Секретар|$)",
        rest,
        flags=re.S | re.I,
    )
    if m_zam:
        prepared = _split_commission_items(clean_text(m_zam.group(1)))
        cleaned = []
        for z in prepared:
            z = re.sub(r"^\d+\.\s*", "", z).strip()
            cleaned.append(normalize_commission_fio(z))
        if cleaned:
            end0 = "." if len(cleaned) == 1 else ";"
            lines.append(
                _comp_line(
                    f"Заместители председателя комиссии: 1. {cleaned[0]}{end0}",
                    "flush",
                )
            )
            for i, z in enumerate(cleaned[1:], start=2):
                ending = "." if i == len(cleaned) else ";"
                lines.append(_comp_line(f"{i}. {z}{ending}", "zam_cont"))

    # Члены: то же правило.
    m_mem = re.search(
        r"Члены\s+(?:комиссии|Комиссии)\s*:\s*(.+?)"
        r"(?=\s*Секретар|$)",
        rest,
        flags=re.S | re.I,
    )
    if m_mem:
        prepared = _split_commission_items(clean_text(m_mem.group(1)))
        cleaned = []
        for mem in prepared:
            mem = re.sub(r"^\d+\.\s*", "", mem).strip()
            cleaned.append(normalize_commission_fio(mem))
        if cleaned:
            end0 = "." if len(cleaned) == 1 else ";"
            lines.append(
                _comp_line(f"Члены комиссии: 1. {cleaned[0]}{end0}", "flush")
            )
            for i, mem in enumerate(cleaned[1:], start=2):
                ending = "." if i == len(cleaned) else ";"
                lines.append(_comp_line(f"{i}. {mem}{ending}", "mem_cont"))

    m_sec = re.search(
        r"Секретар\w*\s+(?:комиссии|Комиссии)\s*[:\-–—]\s*(.+)$",
        rest,
        flags=re.S | re.I,
    )
    if m_sec:
        sec = clean_text(m_sec.group(1)).rstrip(";.")
        sec = normalize_commission_fio(sec)
        lines.append(_comp_line(f"Секретарь комиссии: {sec}.", "flush"))

    return lines if len(lines) > 1 else [point_text]


def normalize_commission_fio(text: str) -> str:
    """Привести ФИО к виду И.О.Фамилия (без пробелов в инициалах)."""
    t = clean_text(text or "")
    # (Лукашевич Г.Л.) / (Г.Л. Лукашевич)
    m = re.search(
        r"\(([А-ЯЁ][а-яё\-]+)\s+([А-ЯЁ])\.\s*([А-ЯЁ])\.\)",
        t,
    )
    if m:
        return re.sub(
            re.escape(m.group(0)),
            f"{m.group(2)}.{m.group(3)}.{m.group(1)}",
            t,
            count=1,
        )
    m = re.search(
        r"\(([А-ЯЁ])\.\s*([А-ЯЁ])\.\s*([А-ЯЁ][а-яё\-]+)\)",
        t,
    )
    if m:
        return re.sub(
            re.escape(m.group(0)),
            f"{m.group(1)}.{m.group(2)}.{m.group(3)}",
            t,
            count=1,
        )
    # хвост «Лукашевич Г.Л.» → «Г.Л.Лукашевич»
    m = re.search(
        r"([А-ЯЁ][а-яё\-]+)\s+([А-ЯЁ])\.\s*([А-ЯЁ])\.\s*$",
        t,
    )
    if m and not re.search(r"[А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁ]", t):
        return t[: m.start()] + f"{m.group(2)}.{m.group(3)}.{m.group(1)}"
    # «Г. Л. Лукашевич» / «Г.Л. Лукашевич»
    t = re.sub(r"([А-ЯЁ])\.\s*([А-ЯЁ])\.\s*", r"\1.\2.", t)
    return t


def _split_commission_items(text: str) -> list[str]:
    """
    Разбить перечень замов/членов:
    — по «1. 2. 3.»;
    — по «;»;
    — по должностям «Заместитель…» / «Начальник…».
    """
    t = clean_text(text or "")
    if not t:
        return []
    # нумерация: «1. … 2. …»
    if re.search(r"\d+\.\s+\S", t):
        parts = re.split(r"(?=\d+\.\s+)", t)
        out = [clean_text(p).rstrip(";.").strip() for p in parts if clean_text(p).strip()]
        if len(out) > 1:
            return out
    # точки с запятой
    if ";" in t:
        parts = [clean_text(x).rstrip(";.").strip() for x in t.split(";")]
        parts = [p for p in parts if p]
        if len(parts) > 1:
            return parts
    # склейка должностей без разделителей
    parts = re.split(
        r"(?=(?:Заместитель\s+главного|Начальник(?:и)?\s+|Ведущий\s+инженер))",
        t,
        flags=re.I,
    )
    out = [clean_text(p).rstrip(";.").strip() for p in parts if clean_text(p).strip()]
    if len(out) > 1:
        return out
    return [t.rstrip(";.").strip()] if t.strip() else []


def format_tasks_block(point_text: str) -> list:
    if "Возложить на Комиссию следующие задачи" not in point_text:
        return [point_text]
    parts = re.split(r"(?=\b\d+\.\d+\.\s)", point_text)
    parts = [clean_text(x) for x in parts if clean_text(x)]
    if len(parts) <= 1:
        return [point_text]
    head = parts[0].rstrip(":").rstrip() + ":"
    return [head] + parts[1:]


def format_grounds_block(point_text: str) -> list:
    """
    Пункт про основания — список БЕЗ маркеров (эталон комиссии).
    Каждый пункт на том же отступе 1,25 см.
    """
    if "Определить, что основанием" not in point_text:
        return [point_text]
    m = re.match(r"^(\d+\.\s*Определить, что основанием[^:]*:)\s*(.*)$", point_text, flags=re.S)
    if not m:
        return [point_text]
    head = m.group(1)
    rest = clean_text(m.group(2))
    # убрать уже стоящие маркеры
    rest = re.sub(r"(?m)^[\s]*[\-\u2013\u2014\u2022]+\s*", "", rest)
    rest = rest.replace("– ", "").replace("— ", "").replace("- ", "")
    items = [clean_text(x) for x in rest.split(";") if clean_text(x)]
    lines = [head]
    for i, item in enumerate(items):
        item = item.lstrip("–—•- ").rstrip(".")
        ending = "." if i == len(items) - 1 else ";"
        lines.append(f"{item}{ending}")
    return lines


def strip_leading_marker_line(text: str) -> str:
    """Убрать маркер в начале строки списка."""
    t = text or ""
    t = re.sub(
        r"^[\s]*[\-\u2013\u2014\u2022\u00B7\uF0B7\uF0A7\u25CB\u25A0\u25CF\*◦▪▸►■○●]+\s*",
        "",
        t,
    )
    return t



# ============================================================
# СБОРКА ДОКУМЕНТА
# ============================================================

def build_order(parsed: dict, output_path: str):
    doc = Document()
    configure_page(doc)

    # Место под бланк
    add_empty(doc, Settings.BLANK_TOP, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ПРИКАЗ
    add_text_paragraph(
        doc,
        "ПРИКАЗ",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=Cm(0),
        bold=False,
    )
    add_empty(doc, 2)

    # Место издания
    add_text_paragraph(
        doc,
        Settings.CITY,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=Cm(0),
    )
    add_empty(doc, 1)

    # Заголовок «О ...»
    add_title_table(doc, parsed["title"])
    add_empty(doc, 1)

    # Преамбула
    if parsed["preamble"]:
        add_text_paragraph(
            doc,
            parsed["preamble"],
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_indent=Settings.FIRST_INDENT,
        )
        add_empty(doc, 1)

    # ПРИКАЗЫВАЮ
    add_text_paragraph(
        doc,
        "ПРИКАЗЫВАЮ:",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_indent=Settings.FIRST_INDENT,
        bold=False,
    )
    add_empty(doc, 1)

    # Пункты
    for point in parsed["points"]:
        blocks = [point]
        low_p = point.lower()
        if (
            "утвердить персональный состав" in low_p
            or ("создать" in low_p and "состав" in low_p and "комисси" in low_p)
            or ("в следующем составе" in low_p and "комисси" in low_p)
        ):
            blocks = format_composition_block(point)
        elif "возложить на комиссию следующие задачи" in low_p:
            blocks = format_tasks_block(point)
        elif "определить, что основанием" in low_p:
            blocks = format_grounds_block(point)
        elif re.match(r"^\d+\.\s*Начальникам районов", point):
            blocks = format_tasks_block(point)  # 6.1 / 6.2 / 6.3

        for block in blocks:
            # состав комиссии: метки без отступа; 2+ номера — на уровне «1.»
            if isinstance(block, dict):
                text = strip_leading_marker_line(block.get("text", ""))
                style = block.get("style", "body")
            else:
                text = strip_leading_marker_line(block)
                style = "body"

            if style == "flush":
                fi, li = Cm(0), Cm(0)
            elif style == "zam_cont":
                fi, li = Cm(0), Settings.COMMISSION_ZAM_NUM_LEFT
            elif style == "mem_cont":
                fi, li = Cm(0), Settings.COMMISSION_MEM_NUM_LEFT
            else:
                # intro пункта «1. Создать…» и обычные пункты
                fi, li = Settings.FIRST_INDENT, Cm(0)

            add_text_paragraph(
                doc,
                text,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_indent=fi,
                left_indent=li,
            )

    # Подпись
    # В черновике могла быть Рощина, в свежих примерах — Вирочкин.
    # Берём из настроек (актуальный директор), имя из черновика не навязываем.
    add_signature(doc, Settings.DIRECTOR_TITLE, Settings.DIRECTOR_NAME)

    save_docx_unprotected(doc, output_path)
    verify_prikaz_before_publish(output_path)
    return output_path


def verify_prikaz_before_publish(docx_path: str) -> dict:
    """
    Всегда проверять приказ перед выдачей результата:
    — нет маркеров;
    — подпись: таб ФИО на 12 см;
    — состав комиссии: отдельные номера у замов и членов.
    При возможности — автоисправление.
    """
    report = {"ok": True, "issues": [], "fixed": 0, "details": []}
    doc = Document(docx_path)
    target_twips = int(Cm(Settings.SIGN_FIO_TAB_CM).twips)

    for p in doc.paragraphs:
        t = p.text
        if re.match(r"^[\s]*[\-\u2013\u2014\u2022\*◦▪▸►]+", t or ""):
            new_t = strip_leading_marker_line(t)
            if new_t != t:
                for r in p.runs:
                    r.text = ""
                if p.runs:
                    p.runs[0].text = new_t
                    set_run_font(p.runs[0])
                else:
                    set_run_font(p.add_run(new_t))
                report["fixed"] += 1
                report["details"].append(f"убран маркер: {t[:60]}")

    for p in reversed(list(doc.paragraphs)):
        t = p.text.strip()
        if not t:
            continue
        if t.lower().startswith("директор") or "\t" in p.text:
            pPr = p._element.get_or_add_pPr()
            for old in list(pPr.findall(qn("w:tabs"))):
                pPr.remove(old)
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "left")
            tab.set(qn("w:pos"), str(target_twips))
            tabs.append(tab)
            pPr.append(tabs)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            if "\t" in p.text:
                left, right = p.text.split("\t", 1)
                right2 = normalize_commission_fio(right.strip())
                new_t = f"{left.strip()}\t{right2}"
                if new_t != p.text:
                    for r in p.runs:
                        r.text = ""
                    if p.runs:
                        p.runs[0].text = new_t
                        set_run_font(p.runs[0])
                    else:
                        set_run_font(p.add_run(new_t))
            report["fixed"] += 1
            report["details"].append(
                f"подпись: таб ФИО = {Settings.SIGN_FIO_TAB_CM} см"
            )
            break

    blob = "\n".join(p.text for p in doc.paragraphs)
    if "Заместители председателя" in blob and "Члены комиссии" in blob:
        # после двоеточия сразу «1.» в той же строке (не на новой)
        m = re.search(
            r"Заместители председателя комиссии:\s*1\.\s",
            blob,
            flags=re.I,
        )
        m2 = re.search(r"Члены комиссии:\s*1\.\s", blob, flags=re.I)
        if not m or not m2:
            report["ok"] = False
            report["issues"].append(
                "состав комиссии: после «Заместители…:» / «Члены…:» "
                "сразу в той же строке должно быть «1. должность И.О.Фамилия»"
            )
        # метки без абзацного отступа
        for p in doc.paragraphs:
            t = (p.text or "").strip().lower()
            if t.startswith(
                (
                    "председатель комиссии:",
                    "заместители председателя комиссии:",
                    "члены комиссии:",
                    "секретарь комиссии:",
                )
            ):
                fi = p.paragraph_format.first_line_indent
                if fi is not None and float(fi.cm) > 0.05:
                    p.paragraph_format.first_line_indent = Cm(0)
                    report["fixed"] += 1
                    report["details"].append(
                        f"убран отступ у метки комиссии: {p.text[:50]}"
                    )

    still_marks = [
        p.text[:50]
        for p in doc.paragraphs
        if re.match(r"^[\s]*[\-\u2013\u2014\u2022]", p.text or "")
    ]
    if still_marks:
        report["ok"] = False
        report["issues"].append(f"остались маркеры: {still_marks[:3]}")

    save_docx_unprotected(doc, docx_path)
    return report


def format_prikaz(input_path: str, output_path: str | None = None) -> str:
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Файл не найден: {input_path}")

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        if base.endswith("."):
            base = base[:-1]
        output_path = f"{base}_оформлен{ext}"

    print(f"[1/4] Читаю черновик: {input_path}")
    raw = extract_raw_text(input_path)

    print("[2/4] Разбираю структуру приказа...")
    parsed = parse_draft(raw)
    print(f"       Заголовок: {parsed['title'][:80]}...")
    print(f"       Пунктов: {len(parsed['points'])}")

    print("[3/4] Собираю документ по образцам папки «Приказы»...")
    build_order(parsed, output_path)

    print(f"[4/4] Готово: {output_path}")
    return output_path


# ============================================================
# ТОЧКА ВХОДА
# ============================================================

if __name__ == "__main__":
    default_input = r"C:\Users\v.dubovik\Desktop\Приказ_комиссия по расслед..docx"
    input_file = sys.argv[1] if len(sys.argv) > 1 else default_input
    try:
        out = format_prikaz(input_file)
        print("\nОткройте файл в Word:")
        print(out)
    except Exception as e:
        print(f"[ОШИБКА] {e}")
        sys.exit(1)
