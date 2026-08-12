# -*- coding: utf-8 -*-
"""
Оформление Положения о ПК (по правилам Дубовика / DocAgent):
— поля 30/15/20/20;
— Times New Roman 14 пт;
— заголовки разделов по центру, жирные, отступ 0;
— пустая строка перед заголовком (кроме главы с верха после разрыва);
— после «Содержания» — разрыв; «Введение» тела — с верха страницы;
— текст пунктов — абзацный отступ 1,25 см;
— таблицы тела — 12 пт, отступ 0.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt
from docx.text.paragraph import Paragraph

DEFAULT_PATH = Path(
    r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Промышленная безопасность"
    r"\Положение о производственном контроле_по_образцу_ЦЭМ.docx"
)

FONT = "Times New Roman"
SIZE = Pt(14)


def set_run_font(run, *, bold=False, size=SIZE) -> None:
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(a), FONT)


def zero_spacing(p) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def insert_empty_before(paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addprevious(new_p)
    np = Paragraph(new_p, paragraph._parent)
    zero_spacing(np)
    np.paragraph_format.first_line_indent = Cm(0)
    return np


def set_page_break_before(paragraph, value: bool = True) -> None:
    pPr = paragraph._element.get_or_add_pPr()
    tag = qn("w:pageBreakBefore")
    el = pPr.find(tag)
    if value:
        if el is None:
            pPr.append(OxmlElement("w:pageBreakBefore"))
    elif el is not None:
        pPr.remove(el)


def style_heading(p) -> None:
    text = p.text
    zero_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    set_run_font(p.add_run(text), bold=True, size=SIZE)


def style_body(p, *, indent: bool = True) -> None:
    zero_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(12.5) if indent else Cm(0)
    for r in p.runs:
        set_run_font(r, bold=bool(r.bold), size=SIZE)


def _is_numbered_chapter(t: str) -> bool:
    if not t or len(t) > 200:
        return False
    if re.match(r"^\d+\.\d+", t):
        return False
    if re.match(r"^\d+\.\s", t):  # «1. Разработано…»
        return False
    return bool(re.match(r"^\d+\s+\S", t))


def format_polozhenie_pk(docx_path: str | Path, *, backup: bool = False) -> dict:
    path = Path(docx_path)
    report = {"path": str(path), "headings": 0, "empties_added": 0}

    if backup:
        bak = path.with_name(path.stem + "_до_оформления.docx")
        if not bak.exists():
            shutil.copy2(path, bak)
            report["backup"] = str(bak)

    doc = Document(str(path))
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(3)
        s.right_margin = Cm(1.5)

    paras = list(doc.paragraphs)
    toc_idx = next(
        (i for i, p in enumerate(paras) if (p.text or "").strip() == "Содержание"),
        None,
    )
    intro_hits = [
        i for i, p in enumerate(paras) if (p.text or "").strip() == "Введение"
    ]
    body_start = intro_hits[-1] if intro_hits else None

    if toc_idx is not None:
        set_page_break_before(paras[toc_idx], True)
    if body_start is not None:
        set_page_break_before(paras[body_start], True)

    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if t == "УТВЕРЖДЕНО" or t.startswith("УТВЕРЖДАЮ"):
            zero_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(12.3)  # гриф как в Инструкции по делопроизводству
            for r in p.runs:
                set_run_font(r, bold=False)
        elif t.startswith("Приказ от") or t.startswith("Постановление"):
            zero_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(12.3)
            for r in p.runs:
                set_run_font(r, bold=False)
        elif t == "ПОЛОЖЕНИЕ" or (i < 6 and t.startswith("об организации")):
            zero_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                set_run_font(r, bold=True)

    heading_idxs: list[int] = []
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t:
            continue
        if t in ("Предисловие", "Содержание"):
            heading_idxs.append(i)
            continue
        if body_start is not None and i < body_start:
            continue
        if t in ("Введение", "Приложение А", "Приложение Б"):
            heading_idxs.append(i)
            continue
        if _is_numbered_chapter(t):
            heading_idxs.append(i)

    heading_idxs = sorted(set(heading_idxs))
    report["headings"] = len(heading_idxs)

    for i in reversed(heading_idxs):
        if i == 0:
            continue
        if body_start is not None and i == body_start:
            continue
        if toc_idx is not None and i == toc_idx:
            # пустая перед «Содержание» — нужна
            pass
        prev = paras[i - 1]
        if not (prev.text or "").strip():
            continue
        if body_start is not None and i == body_start:
            continue
        insert_empty_before(paras[i])
        report["empties_added"] += 1

    paras = list(doc.paragraphs)
    toc_idx = next(
        (i for i, p in enumerate(paras) if (p.text or "").strip() == "Содержание"),
        None,
    )
    intro_hits = [
        i for i, p in enumerate(paras) if (p.text or "").strip() == "Введение"
    ]
    body_start = intro_hits[-1] if intro_hits else None

    heading_idxs = []
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t:
            continue
        if t in ("Предисловие", "Содержание"):
            heading_idxs.append(i)
            continue
        if body_start is not None and i < body_start:
            continue
        if t in ("Введение", "Приложение А", "Приложение Б") or _is_numbered_chapter(t):
            heading_idxs.append(i)
    heading_set = set(heading_idxs)

    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t:
            zero_spacing(p)
            p.paragraph_format.first_line_indent = Cm(0)
            continue
        if t in ("УТВЕРЖДЕНО",) or t.startswith("УТВЕРЖДАЮ") or t.startswith("Приказ от") or t.startswith("Постановление"):
            continue
        if t == "ПОЛОЖЕНИЕ" or (i < 8 and t.startswith("об организации")):
            continue

        in_toc = (
            toc_idx is not None
            and body_start is not None
            and toc_idx < i < body_start
        )
        if i in heading_set and not in_toc:
            style_heading(p)
            continue
        if t in ("Предисловие", "Содержание"):
            style_heading(p)
            continue
        if in_toc:
            zero_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                set_run_font(r, bold=False)
            continue
        style_body(p, indent=True)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    zero_spacing(p)
                    p.paragraph_format.first_line_indent = Cm(0)
                    for r in p.runs:
                        set_run_font(r, bold=bool(r.bold), size=Pt(12))

    # маркеры запрещены; подстрочные подписи — 11 пт
    try:
        from formatters.structure_fix import (
            apply_form_caption_fonts,
            strip_text_markers_everywhere,
        )

        report["markers_removed"] = strip_text_markers_everywhere(doc)
        report["caption_fonts"] = apply_form_caption_fonts(doc, size_pt=11)
    except Exception as e:
        report["markers_error"] = str(e)

    doc.save(str(path))

    # кликабельный TOC; титул (утверждение+название+Минск); изменения/приложения —
    # отдельные страницы; лист согласования — без разрыва
    try:
        from formatters.toc_and_structure import apply_document_structure_rules

        report["structure"] = apply_document_structure_rules(str(path), update_toc=True)
    except Exception as e:
        report["structure_error"] = str(e)

    return report


if __name__ == "__main__":
    print(format_polozhenie_pk(DEFAULT_PATH, backup=False))
