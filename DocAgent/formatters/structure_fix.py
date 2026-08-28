# -*- coding: utf-8 -*-
"""
Единая форма титула, корректная нумерация, одинаковые подписанты.

Проблема: снятие Word-автонумерации (numPr) без перевода в текст
ломало нумерацию пунктов. Здесь:
1) перед оформлением — перевести автономера в обычный текст (через Word);
2) не удалять десятичную нумерацию, только маркеры-bullets;
3) привести титул к одной форме (шапка + таблица УТВЕРЖДАЮ);
4) выровнять блок подписантов как в эталоне СМАТ.
"""

from __future__ import annotations

import os
import re
import shutil
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips
from docx.text.paragraph import Paragraph

from .common import (
    FIRST_INDENT_CM,
    TABLE_FONT_PT,
    TITLE_TABLE_FONT_PT,
    configure_page,
    is_contents_heading,
    is_heading_like,
    is_paragraph_in_table,
    is_signatory_like,
    is_title_page_like,
    iter_all_paragraphs,
    needs_first_indent,
    save_docx_unprotected,
    set_run_font,
    set_single_line_spacing,
    zero_para_spacing,
)
from .etalon_format_spec import (
    FIRST_INDENT_TWIPS as ETALON_FIRST_INDENT_TWIPS,
    SIGN_FIO_TAB_CM as ETALON_SIGN_FIO_TAB_CM,
    SIGN_TITLE_MAX_CHARS,
    TITLE_NUMBER_LABEL_PT as TITLE_INSTR_NUMBER_PT,
    apply_chapter_heading_format,
    ensure_etalon_styles,
    resolve_etalon_path,
)

# Эталон ширины «должности … ФИО» (примерно как в образце СМАТ)
SIGN_LINE_WIDTH = 56
# Эталон КЛ + требование: И.О.Фамилия с 12 см
SIGN_FIO_TAB_CM = float(ETALON_SIGN_FIO_TAB_CM)
KL_ETALON_PATH = resolve_etalon_path()
# Межстрочный интервал блока подписантов РИ (эталон СЛЕСАРЬ 30.07.2026 + требование 1,5)
SIGN_LINE_SPACING_RI = 1.5
# ЕДИНСТВЕННЫЙ эталон РИ: копия в DocAgent\etalons\ (основной путь).
# Исходник на N: — запасной, если локальной копии нет.
_RI_ETALON_NAME = "Рабочая инструкция СЛЕСАРЮ 5-го разряда СМАТ 30.07.2026.docx"
_RI_ETALON_LOCAL = str(
    Path(__file__).resolve().parents[1] / "etalons" / _RI_ETALON_NAME
)
_RI_ETALON_NETWORK = (
    r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\!!!ОБМЕН\РАССМОТРЕНИЕ  ДИ, РИ, ПОЛОЖЕНИЙ"
    rf"\ТУ (СМАТ)\{_RI_ETALON_NAME}"
)


def _resolve_ri_etalon_path() -> str:
    """Локальная копия в etalons/ — внутренний шаблон стилей. ОБМЕН/сеть не образец."""
    return _RI_ETALON_LOCAL


RI_ETALON_PATH = _resolve_ri_etalon_path()
RI_ETALON_NETWORK_PATH = _RI_ETALON_NETWORK
# то же для подписантов (отдельный Подписанты.docx не использовать)
SIGNATORIES_SAMPLE_RI = RI_ETALON_PATH
# запас: несколько обычных табов, если позиция таба не сработала
RI_SIGN_TAB_TARGET_UNITS = 10
RI_SIGN_CHARS_PER_TAB = 5
# Обязательные пустые строки титула (эталон СЛЕСАРЬ 30.07.2026)
TITLE_EMPTY_BEFORE_TABLE = 2
TITLE_EMPTY_BEFORE_MINSK = 15
# Подпись «номер инструкции» на титуле — 12 пт (не 14)
TITLE_INSTR_NUMBER_LABEL = "номер инструкции"


def _is_soft_line_break(br_el) -> bool:
    """True для «разрыва строки» (Shift+Enter). Разрывы страницы/колонки не трогаем."""
    typ = br_el.get(qn("w:type"))
    if typ in ("page", "column"):
        return False
    return True


def _expand_vertical_tabs_to_br(p_el) -> int:
    """Символ вертикальной табуляции (\\x0b) в тексте → элемент w:br (разрыв строки)."""
    n = 0
    for t_el in list(p_el.iter(qn("w:t"))):
        text = t_el.text or ""
        if "\x0b" not in text and "\v" not in text:
            continue
        text = text.replace("\v", "\x0b")
        parts = text.split("\x0b")
        if len(parts) < 2:
            continue
        t_el.text = parts[0]
        if parts[0].startswith(" ") or parts[0].endswith(" "):
            t_el.set(qn("xml:space"), "preserve")
        insert_after = t_el
        for part in parts[1:]:
            br = OxmlElement("w:br")
            insert_after.addnext(br)
            new_t = OxmlElement("w:t")
            new_t.text = part
            if part.startswith(" ") or part.endswith(" "):
                new_t.set(qn("xml:space"), "preserve")
            br.addnext(new_t)
            insert_after = new_t
            n += 1
    return n


def _find_soft_br(p_el):
    for el in p_el.iter(qn("w:br")):
        if _is_soft_line_break(el):
            return el
    return None


def _split_paragraph_at_soft_br(p_el, br_el):
    """
    Разрезать абзац по разрыву строки: левая часть остаётся, правая — новый абзац.
    Возвращает элемент нового абзаца (правая часть).
    """
    run = br_el.getparent()
    new_p = OxmlElement("w:p")
    pPr = p_el.find(qn("w:pPr"))
    if pPr is not None:
        new_p.insert(0, deepcopy(pPr))
    p_el.addnext(new_p)

    if run is None:
        parent = br_el.getparent()
        if parent is not None:
            parent.remove(br_el)
        return new_p

    # Перенести всё после текущего run в новый абзац
    following = False
    for child in list(p_el):
        if child is run:
            following = True
            continue
        if following and child.tag != qn("w:pPr"):
            new_p.append(child)

    # Разрезать run: содержимое после br → новый run в начале new_p
    after_in_run = []
    seen_br = False
    for child in list(run):
        if child is br_el:
            seen_br = True
            run.remove(br_el)
            continue
        if seen_br:
            after_in_run.append(child)
            run.remove(child)

    if after_in_run:
        new_r = OxmlElement("w:r")
        rPr = run.find(qn("w:rPr"))
        if rPr is not None:
            new_r.insert(0, deepcopy(rPr))
        for c in after_in_run:
            new_r.append(c)
        pPr_new = new_p.find(qn("w:pPr"))
        if pPr_new is not None:
            pPr_new.addnext(new_r)
        else:
            new_p.insert(0, new_r)

    # Пустой run (только rPr) — убрать
    leftover = [c for c in run if c.tag != qn("w:rPr")]
    if not leftover and run.getparent() is p_el:
        p_el.remove(run)

    return new_p


def replace_soft_line_breaks_with_paragraphs(doc) -> int:
    """
    Знак «разрыв строки» (Shift+Enter / w:br / \\x0b) → обычный перенос абзаца (Enter).
    Разрывы страницы и колонки не меняются.
    """
    total = 0
    for p in list(iter_all_paragraphs(doc)):
        p_el = p._p
        total += _expand_vertical_tabs_to_br(p_el)
        current = p_el
        while True:
            br = _find_soft_br(current)
            if br is None:
                break
            current = _split_paragraph_at_soft_br(current, br)
            total += 1
    return total


def _paragraph_keeps_signatory_tab(text: str) -> bool:
    """
    В блоке подписантов таб на 12 см для ФИО/даты — намеренный, не заменяем на пробел.
    (Отдельное правило: И.О.Фамилия с табуляции 12 см.)
    """
    raw = text or ""
    t = raw.strip()
    if not t:
        return False
    if is_signatory_like(t):
        return True
    # дата бланка на позиции ФИО: «___» _____ 2026г.
    if re.match(r"^[«\"'\t\s]*_+", t) and re.search(r"20\d{2}\s*г", t, flags=re.I):
        return True
    # должность + таб + И.О.Фамилия
    if "\t" in raw and re.search(
        r"[А-ЯЁA-Z]\.[А-ЯЁA-Z]\.[А-ЯЁA-Z][А-ЯЁа-яёA-Za-z\-]+",
        raw.replace(" ", ""),
    ):
        return True
    low = t.lower()
    if low.startswith("директор") or low.startswith("начальник"):
        if "\t" in raw or re.search(
            r"[А-ЯЁA-Z]\.[А-ЯЁA-Z]\.[А-ЯЁA-ZА-Яа-яёA-Za-z\-]+",
            t.replace(" ", ""),
        ):
            return True
    return False


def _replace_tabs_in_paragraph_xml(p_el) -> int:
    """Заменить w:tab и символы \\t в runs на один пробел. Не трогает w:tabs в pPr."""
    n = 0
    for tab in list(p_el.iter(qn("w:tab"))):
        parent = tab.getparent()
        if parent is None or parent.tag != qn("w:r"):
            continue  # определение позиции таба в pPr — не символ
        prev = tab.getprevious()
        next_ = tab.getnext()
        need_space = True
        if prev is not None and prev.tag == qn("w:t") and (prev.text or "").endswith(" "):
            need_space = False
        if next_ is not None and next_.tag == qn("w:t") and (next_.text or "").startswith(" "):
            need_space = False
        if need_space:
            t_el = OxmlElement("w:t")
            t_el.text = " "
            t_el.set(qn("xml:space"), "preserve")
            tab.addnext(t_el)
        parent.remove(tab)
        n += 1
    for t_el in p_el.iter(qn("w:t")):
        txt = t_el.text or ""
        if "\t" not in txt:
            continue
        new = re.sub(r"\t+", " ", txt)
        new = re.sub(r" {2,}", " ", new)
        if new != txt:
            t_el.text = new
            if new.startswith(" ") or new.endswith(" "):
                t_el.set(qn("xml:space"), "preserve")
            n += 1
    return n


def replace_tabs_with_single_space(doc) -> int:
    """
    Знак табуляции по документу → один пробел.
    Исключение: блок подписантов (таб 12 см для ФИО/даты).
    """
    total = 0
    for p in iter_all_paragraphs(doc):
        text = p.text or ""
        if "\t" not in text and not any(
            el.tag == qn("w:tab") and (el.getparent() is not None and el.getparent().tag == qn("w:r"))
            for el in p._p.iter(qn("w:tab"))
        ):
            continue
        if _paragraph_keeps_signatory_tab(text):
            continue
        total += _replace_tabs_in_paragraph_xml(p._p)
    return total


def _set_oxml_para_font(p_el, *, size_pt: int = 14, font_name: str = "Times New Roman") -> None:
    """Выставить шрифт всем run в oxml-абзаце (w:p)."""
    half = str(int(size_pt * 2))
    for r in p_el.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), font_name)
        for tag in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn("w:val"), half)


def is_instruction_number_label(text: str) -> bool:
    """Строка титула «номер инструкции» (без лишнего текста)."""
    t = re.sub(r"\s+", " ", (text or "").strip().lower())
    return t == TITLE_INSTR_NUMBER_LABEL or t.startswith(TITLE_INSTR_NUMBER_LABEL)


def apply_title_instruction_number_font(doc: Document) -> int:
    """На титуле подпись «номер инструкции» — всегда Times New Roman 12 пт."""
    fixed = 0
    paras = list(doc.paragraphs[:45])
    for table in doc.tables:
        blob = " ".join(c.text for row in table.rows for c in row.cells).upper()
        if "УТВЕРЖДАЮ" not in blob and "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ" not in blob:
            continue
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    seen: set[int] = set()
    for p in paras:
        pid = id(p._element)
        if pid in seen or not is_instruction_number_label(p.text):
            continue
        seen.add(pid)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        if p.runs:
            for run in p.runs:
                set_run_font(
                    run,
                    font_size=Pt(TITLE_INSTR_NUMBER_PT),
                    bold=False if run.bold is None else run.bold,
                )
                fixed += 1
        else:
            run = p.add_run(TITLE_INSTR_NUMBER_LABEL)
            set_run_font(run, font_size=Pt(TITLE_INSTR_NUMBER_PT), bold=False)
            fixed += 1
        _set_oxml_para_font(p._element, size_pt=TITLE_INSTR_NUMBER_PT)
    return fixed


def _clear_paragraph_content(paragraph) -> None:
    """
    Удалить содержимое абзаца (runs + гиперссылки), сохранив pPr/закладки.
    Важно: runs внутри w:hyperlink НЕ входят в paragraph.runs — их нужно
    убирать отдельно, иначе текст гиперссылки дублируется при перезаписи.
    """
    p_el = paragraph._element
    keep_tags = {
        qn("w:pPr"),
        qn("w:bookmarkStart"),
        qn("w:bookmarkEnd"),
        qn("w:proofErr"),
    }
    for child in list(p_el):
        if child.tag in keep_tags:
            continue
        p_el.remove(child)


def _set_runs(paragraph, text: str, *, bold=None, size=Pt(14)) -> None:
    """Записать текст в абзац; «\\n» → мягкий перенос строки (Shift+Enter)."""
    text = text if text is not None else ""
    _clear_paragraph_content(paragraph)
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            br_run = paragraph.add_run()
            br = OxmlElement("w:br")
            br_run._r.append(br)
            set_run_font(br_run, bold=bold, font_size=size)
        run = paragraph.add_run(part)
        set_run_font(run, bold=bold, font_size=size)


def read_list_pairs_via_word(docx_path: str) -> list[tuple[str, str]]:
    """
    Вернуть список (текст_абзаца_без_автономера, ListString).
    Word не закрываем у пользователя: отдельный экземпляр или копия файла.
    """
    pairs: list[tuple[str, str]] = []
    try:
        from .word_com import open_docx_readonly
    except Exception:
        return pairs
    try:
        with open_docx_readonly(docx_path) as doc:
            for i in range(1, int(doc.Paragraphs.Count) + 1):
                p = doc.Paragraphs(i)
                try:
                    ls = (p.Range.ListFormat.ListString or "").strip()
                except Exception:
                    ls = ""
                if not ls:
                    continue
                raw = p.Range.Text or ""
                text = raw.replace("\r", "").replace("\x07", "").strip()
                if text:
                    pairs.append((text, ls))
    except Exception:
        return pairs
    return pairs


def materialize_list_numbers(docx_path: str) -> int:
    """
    Вписать автономера Word в текст абзацев и снять numPr.
    Сопоставление по тексту (индексы Word ≠ индексы python-docx из‑за таблиц).
    """
    pairs = read_list_pairs_via_word(docx_path)
    if not pairs:
        convert_auto_numbers_to_text_com(docx_path)
        return 0

    doc = Document(docx_path)
    # индекс: нормализованный текст → список абзацев
    by_text: dict[str, list] = {}
    for p in iter_all_paragraphs(doc):
        key = (p.text or "").strip()
        if not key:
            continue
        by_text.setdefault(key, []).append(p)

    changed = 0
    for text, ls in pairs:
        cands = by_text.get(text) or []
        if not cands:
            continue
        p = cands.pop(0)
        body = (p.text or "").strip()
        if re.match(r"^\d", body):
            pPr = p._element.pPr
            if pPr is not None and pPr.numPr is not None:
                pPr.remove(pPr.numPr)
            continue
        # Длинные пункты ДИ тоже нумеруются списком Word (2.2.1. …) —
        # номер из ListString всегда вписываем, не снимаем numPr втихую.
        if re.fullmatch(r"\d+", ls):
            prefix = f"{ls} "
        elif ls.endswith("."):
            prefix = f"{ls} "
        else:
            prefix = f"{ls} "
        new_text = prefix + (p.text or "").lstrip()
        bold = p.runs[0].bold if p.runs else None
        _set_runs(p, new_text, bold=bold)
        pPr = p._element.pPr
        if pPr is not None and pPr.numPr is not None:
            pPr.remove(pPr.numPr)
        by_text.setdefault(new_text.strip(), []).append(p)
        changed += 1
    save_docx_unprotected(doc, docx_path)
    return changed


def convert_auto_numbers_to_text_com(docx_path: str) -> str:
    """Запасной COM ConvertNumbersToText — без закрытия Word пользователя."""
    try:
        from .word_com import open_docx_readwrite
    except Exception:
        return docx_path
    try:
        with open_docx_readwrite(docx_path) as (doc, _opened):
            try:
                doc.Content.ListFormat.ConvertNumbersToText()
            except Exception:
                pass
    except Exception:
        pass
    return docx_path


def convert_auto_numbers_to_text(docx_path: str) -> str:
    """Перевести автоматическую нумерацию Word в обычный текст."""
    if not docx_path.lower().endswith(".docx"):
        return docx_path
    try:
        materialize_list_numbers(docx_path)
    except Exception:
        convert_auto_numbers_to_text_com(docx_path)
    return docx_path


def _numbering_fmt_map(doc: Document) -> dict[tuple[int, int], str]:
    """(numId, ilvl) -> numFmt val."""
    out: dict[tuple[int, int], str] = {}
    try:
        root = doc.part.numbering_part.element
    except Exception:
        return out
    abs_fmt: dict[tuple[int, int], str] = {}
    for absn in root.findall(qn("w:abstractNum")):
        abs_id = int(absn.get(qn("w:abstractNumId")))
        for lvl in absn.findall(qn("w:lvl")):
            ilvl = int(lvl.get(qn("w:ilvl")))
            fmt = lvl.find(qn("w:numFmt"))
            if fmt is not None:
                abs_fmt[(abs_id, ilvl)] = fmt.get(qn("w:val")) or ""
    num_to_abs: dict[int, int] = {}
    for num in root.findall(qn("w:num")):
        num_id = int(num.get(qn("w:numId")))
        abs_el = num.find(qn("w:abstractNumId"))
        if abs_el is not None:
            num_to_abs[num_id] = int(abs_el.get(qn("w:val")))
    for num_id, abs_id in num_to_abs.items():
        for (a, ilvl), fmt in abs_fmt.items():
            if a == abs_id:
                out[(num_id, ilvl)] = fmt
    return out


def clear_bullet_markers_only(paragraph, fmt_map: dict[tuple[int, int], str]) -> bool:
    """Снять только маркеры (bullet), десятичную нумерацию НЕ трогать."""
    pPr = paragraph._element.pPr
    if pPr is None or pPr.numPr is None:
        return False
    try:
        num_id = int(pPr.numPr.numId.val)
        ilvl = int(pPr.numPr.ilvl.val) if pPr.numPr.ilvl is not None else 0
    except Exception:
        return False
    fmt = (fmt_map.get((num_id, ilvl)) or "").lower()
    if fmt == "bullet":
        pPr.remove(pPr.numPr)
        return True
    # если формат неизвестен — НЕ удаляем (лучше оставить, чем потерять номера)
    return False


# Подстрочные подписи бланков: «(должность)», «(подпись)», «(Ф.И.О.)» и т.п. — 11–12 пт
_CAPTION_TOKEN_RE = re.compile(
    r"^\(\s*(?:"
    r"должность|подпись|ф\.?\s*и\.?\s*о\.?|дата|"
    r"наименование\s+подразделения|нужное\s+подчеркнуть|"
    r"место\s+для\s+печати|м\.?\s*п\.?"
    r")\s*\)$",
    re.IGNORECASE,
)
_CAPTION_LINE_RE = re.compile(
    r"^(?:"
    r"\(\s*(?:должность|подпись|ф\.?\s*и\.?\s*о\.?|дата|"
    r"наименование\s+подразделения|нужное\s+подчеркнуть|"
    r"место\s+для\s+печати|м\.?\s*п\.?)\s*\)"
    r"(?:\s*[\t ]+\s*)*"
    r")+$",
    re.IGNORECASE,
)
_TEXT_MARKER_RE = re.compile(
    r"^[\-\u2013\u2014\u2022\u00B7\uF0B7\uF0A7\u25CB\u25A0\u25A1\u25CF\u25E6\*◦▪▸►■○●]+\s*"
)


def is_form_caption_line(text: str) -> bool:
    """Строка-подсказка под бланком: (должность) / (подпись) / (Ф.И.О.) …"""
    t = (text or "").replace("\xa0", " ").strip()
    if not t or len(t) > 120:
        return False
    low = t.lower().strip(" .")
    if low in ("нужное подчеркнуть", "(нужное подчеркнуть)"):
        return True
    # одна или несколько подписей через таб/пробелы
    parts = re.split(r"[\t]+|\s{2,}", t)
    parts = [p.strip() for p in parts if p.strip()]
    if not parts:
        return False
    if all(_CAPTION_TOKEN_RE.match(p) for p in parts):
        return True
    return bool(_CAPTION_LINE_RE.match(t.replace("\t", "  ")))


def apply_form_caption_fonts(doc: Document, size_pt: float = 11) -> int:
    """
    Подстрочные записи бланков — Times New Roman 11–12 пт (по умолчанию 11).
    Применяется ко всему документу, включая таблицы.
    """
    size = Pt(size_pt)
    changed = 0
    for p in list(iter_all_paragraphs(doc)):
        t = (p.text or "").replace("\xa0", " ").strip()
        if not is_form_caption_line(t):
            continue
        zero_para_spacing(p)
        set_single_line_spacing(p)
        p.paragraph_format.first_line_indent = Cm(0)
        for run in p.runs:
            set_run_font(
                run,
                font_name=run.font.name or "Times New Roman",
                font_size=size,
                bold=False,
            )
            changed += 1
        if not p.runs and t:
            # пустые runs — перезаписать
            _set_runs(p, t, bold=False)
            for run in p.runs:
                set_run_font(
                    run,
                    font_name="Times New Roman",
                    font_size=size,
                    bold=False,
                )
                changed += 1
    return changed


def strip_text_markers_everywhere(doc: Document) -> int:
    """
    Убрать маркеры списков по ВСЕМУ тексту:
    — Word bullet (numPr bullet);
    — символы «–», «—», «•», «-», «*» в начале абзаца.
    Десятичную нумерацию (1.1., 2.3.1.) не трогать.
    """
    removed = 0
    try:
        fmt_map = _numbering_fmt_map(doc)
    except Exception:
        fmt_map = {}
    for p in list(iter_all_paragraphs(doc)):
        if clear_bullet_markers_only(p, fmt_map):
            removed += 1
        t = p.text or ""
        if not t.strip():
            continue
        # не трогать подписи бланков и строки только из подчёркиваний
        if is_form_caption_line(t):
            continue
        new_t, n = _strip_leading_marker_text(t)
        if n and new_t != t:
            bold = p.runs[0].bold if p.runs else None
            _set_runs(p, new_t, bold=bold)
            removed += 1
    return removed


def _strip_leading_marker_text(text: str) -> tuple[str, int]:
    """Убрать ведущий маркер; нумерацию вида 1.2. сохранить."""
    if not text:
        return text, 0
    original = text
    # если сразу идёт цифра номера пункта — не трогать
    if re.match(r"^\d+(\.\d+)*\.\s", text.lstrip()):
        return text, 0
    changed = 0
    t = text
    while True:
        m = _TEXT_MARKER_RE.match(t)
        if not m:
            break
        t = t[m.end() :].lstrip(" \t")
        changed += 1
        if changed > 5:
            break
    # маркер вплотную к букве: «–текст»
    m2 = re.match(r"^([\-\u2013\u2014\u2022\u00B7\uF0B7\uF0A7\u25CB\u25A0\u25CF\*◦▪▸►■○●])([^\d\s].*)$", t)
    if m2:
        t = m2.group(2).lstrip()
        changed += 1
    return (t, 1) if changed else (original, 0)


def fix_broken_number_prefixes(text: str) -> str:
    """Исправить нумерацию: главы без точки, пункты с точками после каждой цифры."""
    t = text
    # ошибочное «2.2.4.3» вместо «2.4.3»
    t = re.sub(r"^2\.2\.4\.3\b", "2.4.3", t)
    t = re.sub(r"^(\d+)\.2\.4\.3\b", r"\1.4.3", t)
    # «3. 1. Текст» → «3.1. Текст» (пробел внутри номера)
    t = re.sub(r"^(\d+)\.\s+(\d+)\.\s+", r"\1.\2. ", t)
    t = re.sub(r"^(\d+)\.\s+(\d+)\s+", r"\1.\2. ", t)
    # ДВОЙНАЯ нумерация: «1.2. 1.2. Текст» → один раз
    t = collapse_duplicate_number_prefix(t)
    t = normalize_numbering_style(t)
    # убрать двойные точки «1.1.. »
    t = re.sub(r"^(\d+(?:\.\d+)*)\.\.+", r"\1.", t)
    return t


def _looks_like_chapter_title(rest: str, *, allow_soft: bool = True) -> bool:
    r = (rest or "").strip()
    if not r or len(r) > 100:
        return False
    # «2.1. Текст» / «1. На оператора» — это НЕ заголовок главы
    if r[:1].isdigit():
        return False
    if re.match(r"^\d+\.", r):
        return False
    low = r.lower()
    if any(
        w in low
        for w in (
            "общие положения",
            "общие требования",
            "обязанност",
            "функци",
            "права",
            "ответственност",
            "заключительн",
            "взаимоотношен",
            "профессиональн",
            "требования",
            "порядок работ",
            "порядок выполнен",
            "характеристик",
            "квалификац",
            "описание работ",
            "охрана труда",
            "безопасн",
            "приложение",
            "термины",
            "сокращен",
            "определен",
            "осмотр",
            "фазировк",
            "объем",
            "нормы",
            "измерен",
            "испытан",
            "авари",
            "правила и порядок",
            "ликвидац",
        )
    ):
        return True
    # заголовок главы часто ЗАГЛАВНЫМИ
    letters = [c for c in r if c.isalpha()]
    if letters and sum(1 for c in letters if c.isupper()) / len(letters) > 0.55:
        return True
    if not allow_soft:
        return False
    # мягкий режим только для «1 ОБЩИЕ…» (номер БЕЗ точки после цифры)
    if (
        len(r) < 80
        and not r.endswith((".", ";", ",", ":"))
        and r[0].isupper()
        and r.count(" ") <= 10
    ):
        if not re.search(
            r"\b(обязан|должен|обеспеч|выполня|руководств|назнача)\w*\b", low
        ):
            return True
    return False


def _is_toc_entry_line(text: str) -> bool:
    """Строка оглавления: «1 ОБЩИЕ… …… 3» / с табом и номером страницы."""
    t = (text or "").strip()
    if not t:
        return False
    if is_contents_heading(t):
        return True
    if re.search(r"[\.·…‧]{2,}\s*\d+\s*$", t):
        return True
    if re.search(r"\t+\d+\s*$", t):
        return True
    # после замены таба на пробел: «1 ОБЩИЕ ПОЛОЖЕНИЯ 7»
    if re.match(r"^\d+([.\s]|\.\d)*", t) and re.search(r"\s\d{1,3}\s*$", t):
        body = re.sub(r"^\d+\.?\s*", "", t)
        body = re.sub(r"\s\d{1,3}\s*$", "", body).strip()
        if 3 <= len(body) <= 120 and _looks_like_chapter_title(body, allow_soft=True):
            return True
    return False


def _find_contents_index(paras: list) -> int | None:
    for i, p in enumerate(paras[:60]):
        if is_contents_heading((p.text or "").strip()):
            return i
    return None


def _toc_region_end(paras: list, toc_idx: int | None = None) -> int | None:
    """Индекс первого абзаца ТЕЛА после оглавления (не строки TOC)."""
    if toc_idx is None:
        toc_idx = _find_contents_index(paras)
    if toc_idx is None:
        return None
    for i in range(toc_idx + 1, min(len(paras), toc_idx + 80)):
        t = (paras[i].text or "").strip()
        if not t:
            continue
        if _is_toc_entry_line(t):
            continue
        # первая глава/пункт тела
        if _is_chapter_heading_para(t) or re.match(r"^\d+\.\d+", t):
            return i
        # голый заголовок «ОБЩИЕ ПОЛОЖЕНИЯ» без номера
        if _CHAPTER_KEYWORD_RE.search(t) and not t.endswith((":", ";", ",")) and t[:1].isupper():
            if not re.match(r"^\d", t):
                return i
    return None


def normalize_numbering_style(text: str) -> str:
    """
    Глава: «1 ОБЩИЕ ПОЛОЖЕНИЯ» (без точки после номера).
    Пункт: «1.1. Текст» / «2.3.4. Текст» (точка после каждой цифры уровня).
    """
    t = text.strip()
    if not t:
        return text

    # «1. ОБЩИЕ…» / «1.ОБЩИЕ…» → «1 ОБЩИЕ…»
    m = re.match(r"^(\d+)\.\s*(.+)$", t)
    if m and _looks_like_chapter_title(m.group(2)):
        return f"{m.group(1)} {m.group(2).lstrip()}"

    # «1 ОБЩИЕ…» уже ок
    m = re.match(r"^(\d+)\s+(.+)$", t)
    if m and "." not in m.group(1) and _looks_like_chapter_title(m.group(2)):
        return f"{m.group(1)} {m.group(2).lstrip()}"

    # пункт/подпункт: «1.1 Текст» / «1.1.Текст» / «1.1. Текст»
    m = re.match(r"^(\d+(?:\.\d+)+)\.?\s*(.*)$", t)
    if m and m.group(2) is not None:
        nums = m.group(1)
        rest = m.group(2).lstrip()
        if rest:
            return f"{nums}. {rest}"
        return f"{nums}."

    return text


def collapse_duplicate_number_prefix(text: str) -> str:
    """
    Убрать повтор номера в начале строки.
    «1.2. 1.2. Текст» → «1.2. Текст»
    «2.1.2. 2.1.2. Текст» → «2.1.2. Текст»
    """
    t = text
    for _ in range(5):
        # только полный повтор с пробелом: «1.2. 1.2. »
        n = re.sub(
            r"^((\d+(?:\.\d+)*)\.)\s+\1(\s+)",
            r"\1\3",
            t,
        )
        if n == t:
            break
        t = n
    return t


def safe_phrase_replace(text: str, old: str, new: str) -> str:
    """
    Замена фразы без появления двойной нумерации.
    Если new = «N.N. » + old, а в тексте номер уже есть — не добавлять второй раз.
    """
    if old not in text:
        return text
    m = re.match(r"^((\d+(?:\.\d+)*)\.)\s+(.*)$", new.strip())
    if m and m.group(2) == old.strip():
        # new добавляет номер к old
        stripped = text.lstrip()
        if re.match(r"^(\d+(?:\.\d+)*)\.\s+", stripped):
            # номер уже есть — не подставлять new внутрь строки
            return text
    result = text.replace(old, new)
    return collapse_duplicate_number_prefix(result)


NUM_ITEM_RE = re.compile(r"^(\d+(?:\.\d+)*)(\.?)\s+(.*)$", re.DOTALL)
# «1 ОБЩИЕ ПОЛОЖЕНИЯ» / «2 ФУНКЦИИ…» — раздел
SECTION_RE = re.compile(r"^(\d+)\s+([А-ЯЁA-Z][А-ЯЁA-Zа-яёA-Za-z\s\-«»\"]{2,})$")


def _parse_num_tuple(num_str: str) -> tuple[int, ...]:
    return tuple(int(x) for x in num_str.split("."))


def _format_num(tup: tuple[int, ...], *, section: bool = False) -> str:
    if section:
        return str(tup[0])
    if len(tup) == 1:
        return f"{tup[0]}"
    return ".".join(str(x) for x in tup) + "."


def _is_section_heading(num: str, rest: str) -> bool:
    """Раздел: короткий заголовок ЗАГЛАВНЫМИ / без точки в конце длинного текста."""
    r = rest.strip()
    if not r:
        return False
    if SECTION_RE.match(f"{num} {r}"):
        return True
    # «1. ОБЩИЕ ПОЛОЖЕНИЯ» тоже раздел
    if len(r) < 90 and r.isupper():
        return True
    if len(r) < 90 and re.match(r"^[А-ЯЁA-Z]", r) and not r.endswith(".") and " " in r:
        # «Общие положения» как заголовок раздела при num из одной цифры
        if "." not in num and any(
            w in r.lower()
            for w in (
                "общие",
                "обязанност",
                "функци",
                "права",
                "ответственност",
                "заключительн",
                "взаимоотношен",
            )
        ):
            return True
    return False


def _next_expected(prev: tuple[int, ...] | None, depth: int) -> tuple[int, ...]:
    """Ожидаемый следующий номер заданной глубины после prev."""
    if prev is None:
        return tuple([1] * depth) if depth > 1 else (1,)
    if depth == len(prev):
        return prev[:-1] + (prev[-1] + 1,)
    if depth > len(prev):
        # углубление: 1.2 → 1.2.1
        base = prev + (1,) * (depth - len(prev))
        return base
    # выход наверх: 1.2.3 при depth=2 → 1.3
    return prev[: depth - 1] + (prev[depth - 1] + 1,)


def verify_and_fix_numbering(doc: Document, *, rewrite_numbers: bool = False) -> dict:
    """
    После оформления — проверить КАЖДЫЙ пункт с номером по порядку.
    По умолчанию номера исходника не переписываем (только отчёт): иначе
    1.6./2.2.10. «подтягивались» к 1.3./2.2.1. и пропадали пункты.
    rewrite_numbers=True — старое поведение (заполнить пропуски), только по явной просьбе.
    """
    report = {
        "checked": 0,
        "fixed": 0,
        "ok": 0,
        "issues": [],
        "fixes": [],
    }
    paras = list(doc.paragraphs)
    start = 0
    for i, p in enumerate(paras[:45]):
        t = p.text.strip()
        if re.match(r"^\d+(\s|\.)", t):
            start = i
            break
        if re.fullmatch(r"минск\s+20\d{2}", t.lower()):
            start = i + 1
    end = len(paras)
    for i in range(len(paras) - 1, max(start, len(paras) // 3), -1):
        low = paras[i].text.strip().lower()
        if low.startswith(("разработал", "согласовано", "с должностной", "с рабочей", "примечание")):
            end = i
            break

    last_section = 0
    last_child: dict[tuple[int, ...], int] = {}

    def clear_deeper_than(prefix: tuple[int, ...]) -> None:
        for key in list(last_child.keys()):
            if len(key) >= len(prefix) and key[: len(prefix)] == prefix:
                # сбросить потомков этого узла при новом соседе
                if key != prefix:
                    del last_child[key]

    for i in range(start, end):
        p = paras[i]
        text = p.text.strip()
        if not text or not re.match(r"^\d", text):
            continue
        # оглавление не перенумеровывать (иначе «1 ОБЩИЕ… … 3» ломает главы тела)
        if is_contents_heading(text) or _is_toc_entry_line(text):
            continue
        m = NUM_ITEM_RE.match(text)
        if not m:
            continue
        num_str, rest = m.group(1), m.group(3)
        # хвост оглавления с номером страницы
        if re.search(r"[\.·…‧]{2,}\s*\d+\s*$", rest) or re.search(r"\t+\d+\s*$", rest):
            continue
        report["checked"] += 1
        cur = _parse_num_tuple(num_str)
        section = _is_section_heading(num_str, rest) or (
            len(cur) == 1 and _is_section_heading(num_str, rest)
        )
        # одиночный номер + короткий заголовок = раздел
        if len(cur) == 1 and len(rest.strip()) < 90 and (
            rest.isupper()
            or any(
                w in rest.lower()
                for w in (
                    "общие положения",
                    "обязанност",
                    "функци",
                    "права",
                    "ответственност",
                    "заключительн",
                    "взаимоотношен",
                    "профессиональн",
                )
            )
        ):
            section = True

        if section:
            expected_n = last_section + 1
            expected = (expected_n,)
            if cur != expected:
                msg = f"абз.{i + 1}: раздел «{num_str}» → «{expected_n}»"
                if rewrite_numbers:
                    new_text = f"{expected_n} {rest.lstrip()}"
                    bold = p.runs[0].bold if p.runs else None
                    _set_runs(p, new_text, bold=bold)
                    report["fixed"] += 1
                    report["fixes"].append(msg)
                    last_section = expected_n
                else:
                    report["ok"] += 1
                    last_section = cur[0]
                report["issues"].append(msg)
            else:
                report["ok"] += 1
                last_section = expected_n
            # новый раздел — сбросить счётчики детей других разделов
            keep_section = last_section
            for key in list(last_child.keys()):
                if not key or key[0] != keep_section:
                    del last_child[key]
            last_child[()] = keep_section
            continue

        # обычный пункт N.M.K…
        # Мягкий режим: НЕ перенумеровывать чужие ветки (это ломало 2.2.9→2.2.1).
        # Принимаем номер как есть; правим только явный дубль сразу после такого же номера.
        parent = cur[:-1]
        prev_tail = last_child.get(parent)

        if parent and last_section and parent[0] != last_section:
            last_section = parent[0]

        if prev_tail is None:
            # первая встреча ветки: номер исходника сохраняем (1.6. не сдвигать в 1.1.)
            if cur[-1] > 1 and len(cur) >= 2:
                expected = parent + (1,)
                msg = f"абз.{i + 1}: старт ветки «{num_str}» → «{'.'.join(map(str, expected))}»"
                if rewrite_numbers:
                    pref = ".".join(str(x) for x in expected) + "."
                    new_text = f"{pref} {rest.lstrip()}"
                    bold = p.runs[0].bold if p.runs else None
                    _set_runs(p, new_text, bold=bold)
                    report["fixed"] += 1
                    report["fixes"].append(msg)
                    last_child[parent] = 1
                else:
                    report["ok"] += 1
                    last_child[parent] = cur[-1]
                report["issues"].append(msg)
            else:
                report["ok"] += 1
                last_child[parent] = cur[-1]
        elif cur[-1] == prev_tail + 1:
            report["ok"] += 1
            last_child[parent] = cur[-1]
        elif cur[-1] > prev_tail + 1:
            expected = parent + (prev_tail + 1,)
            msg = f"абз.{i + 1}: пропуск «{num_str}» → «{'.'.join(map(str, expected))}»"
            if rewrite_numbers:
                pref = ".".join(str(x) for x in expected) + "."
                new_text = f"{pref} {rest.lstrip()}"
                bold = p.runs[0].bold if p.runs else None
                _set_runs(p, new_text, bold=bold)
                report["fixed"] += 1
                report["fixes"].append(msg)
                last_child[parent] = expected[-1]
            else:
                report["ok"] += 1
                last_child[parent] = cur[-1]
            report["issues"].append(msg)
        elif cur[-1] == prev_tail:
            expected = parent + (prev_tail + 1,)
            msg = f"абз.{i + 1}: дубль «{num_str}» → «{'.'.join(map(str, expected))}»"
            if rewrite_numbers:
                pref = ".".join(str(x) for x in expected) + "."
                new_text = f"{pref} {rest.lstrip()}"
                bold = p.runs[0].bold if p.runs else None
                _set_runs(p, new_text, bold=bold)
                report["fixed"] += 1
                report["fixes"].append(msg)
                last_child[parent] = expected[-1]
            else:
                report["ok"] += 1
                last_child[parent] = cur[-1]
            report["issues"].append(msg)
        else:
            # откат назад — только замечание, текст не трогаем
            report["ok"] += 1
            report["issues"].append(
                f"абз.{i + 1}: неожиданный номер «{num_str}» после «{'.'.join(map(str, parent + (prev_tail,)))}» (оставлен)"
            )
            last_child[parent] = max(prev_tail, cur[-1])

        full = parent + (last_child[parent],)
        for key in list(last_child.keys()):
            if len(key) > len(full) and key[: len(full)] == full:
                del last_child[key]

    return report


def normalize_document_structure(docx_path: str, doc_type: str = "unsupported") -> dict:
    """Главный проход: нумерация/титул/подписанты + проверка каждого пункта."""
    report = {
        "bullets_cleared": 0,
        "numbers_fixed": 0,
        "signatories": 0,
        "title_fixed": 0,
        "empties_removed": 0,
        "numbering_check": {},
        "output": docx_path,
    }
    doc = Document(docx_path)
    try:
        doc._source_path = docx_path  # type: ignore[attr-defined]
    except Exception:
        pass
    fmt_map = _numbering_fmt_map(doc)

    # 0) «разрыв строки» (Shift+Enter) → обычный перенос абзаца
    try:
        report["soft_breaks_to_paragraphs"] = replace_soft_line_breaks_with_paragraphs(doc)
    except Exception as e:
        report["soft_breaks_to_paragraphs"] = 0
        report["soft_breaks_error"] = str(e)

    # 0b) табуляция → один пробел (кроме таба 12 см у подписантов)
    try:
        report["tabs_to_space"] = replace_tabs_with_single_space(doc)
    except Exception as e:
        report["tabs_to_space"] = 0
        report["tabs_to_space_error"] = str(e)

    # 0c) уточнить тип документа (РИ часто видна только в таблице титула)
    doc_type = resolve_instruction_doc_type(doc_type=doc_type, doc=doc)
    report["resolved_doc_type"] = doc_type

    # 1) только bullets
    for p in iter_all_paragraphs(doc):
        if clear_bullet_markers_only(p, fmt_map):
            report["bullets_cleared"] += 1

    # 2) текстовые префиксы нумерации (локальные сбои) + маркеры в начале строки
    _marker_re = re.compile(r"^[\s]*(?:[\-\u2013\u2014\u2022\u00B7\*◦▪▸►]+)[\s\u00a0]*")
    for p in doc.paragraphs:
        original = p.text
        if not original.strip():
            continue
        fixed = original
        # убрать «–устав» / «• текст» без пробела
        while True:
            new, n = _marker_re.subn("", fixed, count=1)
            if not n:
                break
            fixed = new.lstrip(" \t")
        m_glued = re.match(r"^([\-\u2013\u2014\u2022\u00B7\*◦▪▸►])([^\d\s].*)$", fixed)
        if m_glued:
            fixed = m_glued.group(2).lstrip()
        if fixed != original:
            report["bullets_cleared"] += 1
        fixed = fix_broken_number_prefixes(fixed)
        if fixed != original:
            _set_runs(p, fixed, bold=p.runs[0].bold if p.runs else None)
            report["numbers_fixed"] += 1
        t = fixed.strip()
        if re.match(r"^\d+\s+[А-ЯЁA-Z]{3,}", t) and len(t) < 80:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)

    # 3) титул — ТОЛЬКО если уже есть таблица как в эталоне: подправить поля.
    #    Склеенный титул без таблицы — собрать таблицу и УДАЛИТЬ старые абзацы (не оставлять пустыми).
    try:
        report["title_fixed"] = ensure_title_table_like_sample(doc, doc_type)
    except Exception:
        report["title_fixed"] = 0

    for p in doc.paragraphs[:12]:
        if is_title_page_like(p.text, p.alignment) or (
            p.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(p.text.strip()) < 80
        ):
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)

    # 3b) убрать лишние пустые / затем восстановить обязательные как в эталоне
    report["empties_removed"] = collapse_title_spacer_empties(doc)
    try:
        report["title_spacers"] = ensure_title_sample_spacers(doc)
    except Exception:
        report["title_spacers"] = 0

    # 4) ВСЕ примечания — в конец документа, сразу перед «Разработал:»
    try:
        report["notes_moved"] = move_notes_before_razrabotal(doc)
    except Exception as e:
        report["notes_moved"] = 0
        report["notes_move_error"] = str(e)

    # 5) двойные пробелы — по всему тексту (все документы)
    try:
        report["double_spaces_fixed"] = collapse_double_spaces_in_document(doc)
    except Exception:
        report["double_spaces_fixed"] = 0

    # 5b) пустые внутри глав убрать; содержание → разрыв; главы по центру
    try:
        report["body_empties_removed"] = remove_empty_paragraphs_in_body(doc)
    except Exception:
        report["body_empties_removed"] = 0
    try:
        report["contents_page"] = separate_contents_onto_own_page(doc)
    except Exception as e:
        report["contents_page_error"] = str(e)
    try:
        ch = format_chapter_headings(doc)
        report["chapter_spacers"] = ch.get("spacers", 0)
        report["chapters_centered"] = ch.get("centered", 0)
    except Exception:
        report["chapter_spacers"] = 0

    # 6) абзацный отступ: титул=0, тело=1,25 (кроме заголовков)
    report["indents_fixed"] = apply_body_first_indents(doc)

    # 6a) шрифт таблиц: титул 14, остальные 12
    try:
        report["table_fonts"] = apply_table_fonts(doc)
    except Exception as e:
        report["table_fonts_error"] = str(e)

    # 6b) «номер инструкции» на титуле — 12 пт (не 14)
    try:
        report["title_number_font"] = apply_title_instruction_number_font(doc)
    except Exception:
        report["title_number_font"] = 0

    # 7) примечания перед «Разработал:» (до подписантов)
    try:
        report["notes_moved"] = int(report.get("notes_moved", 0)) + move_notes_before_razrabotal(
            doc
        )
    except Exception:
        pass

    # 7b) подписанты ПОСЛЕ примечаний и отступов
    report["signatories"] = apply_signatory_block_format(doc, doc_type=doc_type)

    # 8) ОБЯЗАТЕЛЬНО: проверка и правка нумерации ПОСЛЕ всех перестановок абзацев
    chk = verify_and_fix_numbering(doc)
    report["numbering_check"] = chk
    report["numbers_fixed"] += int(chk.get("fixed", 0))
    # контрольный второй проход — поймать оставшиеся сбои
    chk2 = verify_and_fix_numbering(doc)
    report["numbers_fixed"] += int(chk2.get("fixed", 0))
    if chk2.get("fixes"):
        report["numbering_check"]["fixes"] = list(chk.get("fixes") or []) + list(
            chk2.get("fixes") or []
        )
    if chk2.get("issues"):
        report["numbering_check"]["issues"] = list(chk2.get("issues") or [])
    report["numbering_check"]["checked"] = chk2.get("checked", chk.get("checked", 0))
    report["numbering_check"]["ok"] = chk2.get("ok", chk.get("ok", 0))
    report["numbering_check"]["fixed"] = int(report.get("numbers_fixed", 0))

    # 8b) после нумерации снова: примечания + подписанты (эталон не ломать)
    try:
        report["notes_moved"] = int(report.get("notes_moved", 0)) + move_notes_before_razrabotal(
            doc
        )
    except Exception:
        pass
    try:
        report["signatories"] = int(report.get("signatories", 0)) + apply_signatory_block_format(
            doc, doc_type=doc_type
        )
    except Exception:
        pass

    save_docx_unprotected(doc, docx_path)

    # 9) если подписи одни на листе — перенести к ним последний абзац текста
    try:
        report["sign_page_fix"] = move_last_body_if_signatories_orphaned(docx_path)
    except Exception as e:
        report["sign_page_fix"] = {"moved": False, "detail": f"error: {e}"}

    # 10) после переноса страницы — снова проверить нумерацию (текст не менялся, но на всякий случай)
    try:
        doc3 = Document(docx_path)
        chk3 = verify_and_fix_numbering(doc3)
        if chk3.get("fixed"):
            save_docx_unprotected(doc3, docx_path)
            report["numbers_fixed"] += int(chk3.get("fixed", 0))
            report["numbering_check"] = chk3
    except Exception:
        pass

    # отчёт в logs
    try:
        from datetime import datetime
        from pathlib import Path

        log_dir = Path(__file__).resolve().parents[1] / "logs"
        log_dir.mkdir(exist_ok=True)
        chk_log = report.get("numbering_check") or {}
        lines = [
            f"Проверка нумерации {datetime.now():%Y-%m-%d %H:%M:%S}",
            f"Файл: {docx_path}",
            f"Проверено пунктов: {chk_log.get('checked', 0)}",
            f"Без ошибок: {chk_log.get('ok', 0)}",
            f"Исправлено: {report.get('numbers_fixed', 0)}",
            "",
            "Исправления:",
            *(chk_log.get("fixes") or ["(нет)"]),
            "",
            "Замечания:",
            *(chk_log.get("issues") or ["(нет)"]),
        ]
        (log_dir / "numbering_check.txt").write_text("\n".join(lines), encoding="utf-8")
    except Exception:
        pass

    return report


# ---------------------------------------------------------------------------
# Титул и подписанты — эталон: РИ СЛЕСАРЮ 5 СМАТ 30.07.2026
# ---------------------------------------------------------------------------

FIO_RE = re.compile(r"[А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁа-яё\-]+")
DATE_RE = re.compile(r"[«\"'].*\d{2,4}\s*г|20\d{2}\s*г|_+\s*20", re.I)
INSTR_RE = re.compile(r"(должностн\w*\s+инструкци\w*|рабоч\w*\s+инструкци\w*)", re.I)


def _delete_paragraph(paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _set_para_format_no_indent(paragraph, *, align=None) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align is not None:
        paragraph.alignment = align


def _write_cell_lines(cell, lines: list[str], *, bold_first: bool = False) -> None:
    """Записать строки в ячейку (каждая строка — абзац), без лишних пустых. Отступ 0."""
    clean = [ln.rstrip() for ln in lines if ln is not None]
    if not clean:
        clean = [""]
    paras = cell.paragraphs
    for p in paras:
        for r in p.runs:
            r.text = ""
        if p.text:
            p.text = ""
    first = paras[0]
    _set_runs(first, clean[0], bold=True if bold_first else None)
    _set_para_format_no_indent(first)
    first.paragraph_format.first_line_indent = Cm(0)
    first.paragraph_format.left_indent = Cm(0)
    for p in list(paras)[1:]:
        _delete_paragraph(p)
    for ln in clean[1:]:
        np = cell.add_paragraph()
        _set_runs(np, ln)
        _set_para_format_no_indent(np)
        np.paragraph_format.first_line_indent = Cm(0)
        np.paragraph_format.left_indent = Cm(0)


def _set_table_no_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")


def _split_two_columns(text: str) -> tuple[str, str] | None:
    """Разделить строку титула на левую и правую колонки (таб / большие пробелы)."""
    raw = text.replace("\xa0", " ").rstrip()
    if not raw.strip():
        return None
    if "\t" in raw:
        parts = [p.strip() for p in re.split(r"\t+", raw)]
        parts = [p for p in parts if p]
        if len(parts) >= 2:
            return parts[0], parts[-1]
        if len(parts) == 1:
            # табы только слева/справа — одна колонка
            only = parts[0]
            if INSTR_RE.search(only) or only.upper().startswith("УТВЕРЖДАЮ"):
                return (only, "") if INSTR_RE.search(only) else ("", only)
            return None
    m = re.match(r"^(.*?)(?: {3,}|\t+)(.*)$", raw)
    if not m:
        return None
    left, right = m.group(1).strip(), m.group(2).strip()
    if not left and not right:
        return None
    if left and right:
        return left, right
    return left, right


def _is_mashed_title_line(text: str) -> bool:
    t = text.strip()
    if not t or (t.startswith("___") and "«" not in t and not DATE_RE.search(t)):
        return False
    low = t.lower()
    if "минск 20" in low or low in ("минск",):
        return False
    if "номер инструкции" in low:
        return False
    # дата утверждения — часть титула
    if DATE_RE.search(t) and ("«" in t or "_" in t or "г." in low):
        return True
    if INSTR_RE.search(t) and "утверждаю" in low:
        return True
    if INSTR_RE.search(t) or t.upper().startswith("УТВЕРЖДАЮ"):
        return True
    pair = _split_two_columns(text)
    if pair:
        left, right = pair
        if left or right:
            if FIO_RE.search(t) or DATE_RE.search(t) or "утверждаю" in low:
                return True
            if left and right and len(t) < 160:
                return True
    if re.match(r"^(старшему|слесарю|водителю|диспетчеру|инженеру|механику)\b", low):
        return True
    if re.match(r"^(службы|и\s+автомобильного|автомобильного)\b", low) and len(t) < 80:
        return True
    return False


def _find_title_table(doc: Document):
    for table in doc.tables:
        try:
            blob = " ".join(c.text for row in table.rows for c in row.cells).lower()
        except Exception:
            continue
        if "утверждаю" in blob or "инструкция" in blob:
            return table
    return None


def _doc_kind_label(doc_type: str, blob: str = "") -> str:
    """
    Подпись вида документа на титуле.
    Не подставлять «ДОЛЖНОСТНАЯ» по умолчанию — для КЛ/положения это ошибка.
    """
    low = (blob or "").lower()
    hint = (doc_type or "").lower().strip()
    if hint == "rabochaya_instrukciya" or "рабочая инструкц" in low or (
        "рабоч" in low and "инструкц" in low and "должностн" not in low
    ):
        return "РАБОЧАЯ ИНСТРУКЦИЯ"
    if hint == "dolzhnostnaya_instrukciya" or "должностн" in low:
        return "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ"
    if hint == "instrukciya_ot" or "охране труда" in low:
        return "ИНСТРУКЦИЯ"
    if hint == "polozhenie" or "положен" in low:
        # инструкция по эксплуатации / положение — на титуле КЛ слово «ИНСТРУКЦИЯ»
        if "инструкц" in low or "эксплуатац" in low:
            return "ИНСТРУКЦИЯ"
        return "ПОЛОЖЕНИЕ"
    if "эксплуатац" in low and "инструкц" in low:
        return "ИНСТРУКЦИЯ"
    if "инструкц" in low and "должностн" not in low and "рабоч" not in low:
        return "ИНСТРУКЦИЯ"
    return "ИНСТРУКЦИЯ"


def _is_kl_tab_title_layout(doc: Document) -> bool:
    """
    Титул эталона КЛ 31.07.2026: «ИНСТРУКЦИЯ … УТВЕРЖДАЮ» табами в абзацах,
    без таблицы УТВЕРЖДАЮ. Такой титул НЕ пересобирать в таблицу ДИ/РИ.
    """
    if _find_title_table(doc) is not None:
        return False
    for p in doc.paragraphs[:20]:
        t = (p.text or "").replace("\xa0", " ")
        low = t.lower()
        if "утверждаю" in low and ("\t" in t or "инструкция" in low):
            return True
        if "\t" in t and "заместитель" in low and "директор" in low:
            return True
    return False


def _collect_mashed_title_paras(doc: Document) -> list:
    """Абзацы склеенного титула (только до начала тела / подписантов)."""
    found: list = []
    started = False
    for p in doc.paragraphs[:40]:
        t = p.text.strip()
        low = t.lower()
        # тело или подписи — титул закончен (не путать с подписантами!)
        if re.match(r"^\d+\s+[А-ЯЁA-Z]", t) or re.match(r"^\d+\.\d+", t):
            break
        if low.startswith(("разработал", "согласовано", "согласовал")):
            break
        if "ознакомлен" in low and "инструкц" in low:
            break
        if "минск 20" in low and not started:
            break
        if t.startswith("___") and started:
            break
        if "номер инструкции" in low and started:
            break
        if _is_mashed_title_line(p.text) or (
            started and t and _split_two_columns(p.text) is not None
        ):
            started = True
            found.append(p)
            continue
        if started and not t:
            found.append(p)
            continue
        if started and t and not _is_mashed_title_line(p.text):
            break
    return found


def _parse_mashed_columns(paras: list) -> tuple[list[str], list[str]]:
    lefts: list[str] = []
    rights: list[str] = []
    for p in paras:
        t = p.text.replace("\xa0", " ").strip()
        if not t:
            continue
        pair = _split_two_columns(p.text)
        if pair:
            l, r = pair
            if l:
                lefts.append(l)
            if r:
                rights.append(r)
        else:
            # одна колонка — эвристика
            if DATE_RE.search(t) or t.upper().startswith("УТВЕРЖДАЮ") or FIO_RE.search(t) or "начальник" in t.lower() and "служб" not in t.lower()[:20]:
                rights.append(t)
            else:
                lefts.append(t)
    return lefts, rights


def _build_title_rows(lefts: list[str], rights: list[str], doc_type: str) -> tuple[list[str], list[str], list[str]]:
    """Собрать 3 строки таблицы как в эталоне."""
    label = _doc_kind_label(doc_type, " ".join(lefts + rights))
    # убрать дубль заголовка из lefts
    pos_lines: list[str] = []
    for ln in lefts:
        if INSTR_RE.search(ln) and len(ln) < 40:
            label = _doc_kind_label(doc_type, ln)
            continue
        pos_lines.append(ln)

    appr: list[str] = []
    date_line = "«_____» _____________ 2026 г."
    for ln in rights:
        low = ln.lower().strip()
        if low.startswith("утверждаю"):
            continue
        if DATE_RE.search(ln) or (ln.strip().startswith("«") and "_" in ln):
            date_line = ln.strip()
            continue
        appr.append(ln)

    # эталон: row1 left = несколько строк должности, row2 left = продолжение (последняя строка)
    if len(pos_lines) >= 2:
        left_mid = "\n".join(pos_lines[:-1])
        left_bot = pos_lines[-1]
    elif len(pos_lines) == 1:
        left_mid = pos_lines[0]
        left_bot = ""
    else:
        left_mid = ""
        left_bot = ""

    right_mid = "\n".join(appr) if appr else ""
    return (
        [label, "УТВЕРЖДАЮ"],
        [left_mid, right_mid],
        [left_bot, date_line],
    )


def _insert_title_table_after(paragraph, rows_data: list[list[str]]) -> object:
    """Вставить таблицу 3×2 сразу после абзаца (ширины колонок как в эталоне)."""
    doc = paragraph.part.document
    table = doc.add_table(rows=3, cols=2)
    _set_table_no_borders(table)
    try:
        table.autofit = True
        # эталон: левая шире (~9 см), правая (~7.6 см) при поле 16.5 см
        for row in table.rows:
            row.cells[0].width = Cm(9.2)
            row.cells[1].width = Cm(7.6)
    except Exception:
        pass
    for ri, row_vals in enumerate(rows_data):
        for ci, val in enumerate(row_vals):
            lines = str(val).split("\n") if val else [""]
            _write_cell_lines(
                table.rows[ri].cells[ci],
                lines,
                bold_first=(ri == 0),
            )
            for p in table.rows[ri].cells[ci].paragraphs:
                _set_para_format_no_indent(p)
    paragraph._p.addnext(table._tbl)
    return table


def _title_table_looks_corrupted(table) -> bool:
    """Колонки перепутаны / склеены (типичный брак после сжатия пробелов)."""
    try:
        left = "\n".join(c.text for c in [table.rows[1].cells[0], table.rows[2].cells[0]]).lower()
        right = "\n".join(c.text for c in [table.rows[1].cells[1], table.rows[2].cells[1]]).lower()
        mid_l = table.rows[1].cells[0].text.lower()
        mid_r = table.rows[1].cells[1].text.lower()
    except Exception:
        return False
    # должность в дательном падеже оказалась справа
    if any(x in mid_r for x in ("старшему", "слесарю", "водителю", "диспетчеру", "инженеру")):
        return True
    # «тендерного» слева, а «старшему» нет — колонки смешались
    if "тендерного" in mid_l and not any(
        x in mid_l for x in ("старшему", "слесарю", "служб")
    ):
        return True
    if "тендерного" in mid_l and "старшему" in mid_r:
        return True
    # в одной ячейке склеены обе колонки без перевода строки между ролями
    if "старшему" in mid_r and "начальник" in mid_r and "\n" not in table.rows[1].cells[1].text:
        if "транспорт" in mid_r or "______" in mid_r:
            return True
    return False


def _polish_existing_title_table(table, doc_type: str) -> int:
    changed = 0
    try:
        c00 = table.rows[0].cells[0].text.strip()
        c01 = table.rows[0].cells[1].text.strip()
    except Exception:
        return 0
    want = _doc_kind_label(doc_type, c00)
    if not c00:
        if "РАБОЧАЯ" in want or "ДОЛЖНОСТНАЯ" in want:
            _write_cell_lines(table.rows[0].cells[0], [want], bold_first=True)
            changed += 1
    elif INSTR_RE.search(c00) and c00.upper().strip() != want:
        rest = INSTR_RE.sub("", c00).strip(" _")
        if not rest:
            if "РАБОЧАЯ" in want or "ДОЛЖНОСТНАЯ" in want:
                _write_cell_lines(table.rows[0].cells[0], [want], bold_first=True)
                changed += 1
        # иначе в ячейке своё название (инженер ЛСиМ и т.п.) — не затирать
    if not c01 or "утверждаю" not in c01.lower():
        _write_cell_lines(table.rows[0].cells[1], ["УТВЕРЖДАЮ"], bold_first=True)
        changed += 1
    # дата в [2,1], если пусто
    try:
        if len(table.rows) >= 3 and not table.rows[2].cells[1].text.strip():
            _write_cell_lines(table.rows[2].cells[1], ["«_____» _____________ 2026 г."])
            changed += 1
    except Exception:
        pass
    _set_table_no_borders(table)
    return changed


def ensure_title_table_like_sample(doc: Document, doc_type: str = "unsupported") -> int:
    """
    Титул как в эталоне СЛЕСАРЬ 30.07.2026:
    — если таблица уже есть: только подправить, склеенные абзацы УДАЛИТЬ (не оставлять пустыми);
    — если таблицы нет, а титул склеен пробелами/табами: собрать таблицу 3×2 и удалить старые строки.
    """
    changed = 0
    table = _find_title_table(doc)
    mashed = _collect_mashed_title_paras(doc)
    mashed_with_text = [p for p in mashed if p.text.strip()]

    if table is not None:
        corrupted = _title_table_looks_corrupted(table)
        if corrupted and mashed_with_text:
            # пересобрать из исходных колонок
            lefts, rights = _parse_mashed_columns(mashed_with_text)
            rows = _build_title_rows(lefts, rights, doc_type)
            for ri, pair in enumerate(rows):
                for ci, val in enumerate(pair):
                    _write_cell_lines(
                        table.rows[ri].cells[ci],
                        str(val).split("\n") if val else [""],
                        bold_first=(ri == 0),
                    )
            changed += 1
        else:
            changed += _polish_existing_title_table(table, doc_type)
            try:
                mid_left = table.rows[1].cells[0].text.strip()
                mid_right = table.rows[1].cells[1].text.strip()
            except Exception:
                mid_left = mid_right = "x"
            if mashed_with_text and (not mid_left or not mid_right or corrupted):
                lefts, rights = _parse_mashed_columns(mashed_with_text)
                rows = _build_title_rows(lefts, rights, doc_type)
                for ri, pair in enumerate(rows):
                    for ci, val in enumerate(pair):
                        _write_cell_lines(
                            table.rows[ri].cells[ci],
                            str(val).split("\n") if val else [""],
                            bold_first=(ri == 0),
                        )
                changed += 1
        for p in mashed:
            _delete_paragraph(p)
            changed += 1
        return changed

    # таблицы нет
    head_blob = " ".join(p.text for p in doc.paragraphs[:12]).lower()
    has_org = (
        "минсккоммунтеплосеть" in head_blob
        or "исполнительный комитет" in head_blob
        or "коммунальное унитарное" in head_blob
    )
    looks_like_title_mash = mashed_with_text and (
        has_org
        or any("утверждаю" in p.text.lower() for p in mashed_with_text)
        or any(INSTR_RE.search(p.text) for p in mashed_with_text)
    )
    if not looks_like_title_mash:
        # шапка организации уже есть абзацами — не вставлять второй титул
        if has_org:
            return 0
        if _should_skip_invented_title(doc, doc_type):
            return 0
        return create_title_page_from_scratch(doc, doc_type, source_path=getattr(doc, "_source_path", "") or "")

    lefts, rights = _parse_mashed_columns(mashed_with_text)
    rows = _build_title_rows(lefts, rights, doc_type)

    # точка вставки: после «МИНСККОММУНТЕПЛОСЕТЬ» или перед первым mashed
    anchor = None
    for p in doc.paragraphs[:15]:
        if "минсккоммунтеплосеть" in p.text.lower():
            anchor = p
            break
    if anchor is None:
        anchor = mashed_with_text[0]
        # вставить ПЕРЕД mashed: берём предыдущий абзац
        # anchor.addprevious via insert after previous
        prev = None
        for p in doc.paragraphs:
            if p._p is mashed_with_text[0]._p:
                break
            prev = p
        if prev is not None:
            anchor = prev

    _insert_title_table_after(anchor, rows)
    changed += 1
    for p in mashed:
        _delete_paragraph(p)
        changed += 1
    return changed


ORG_HEADER_LINES = (
    "МИНСКИЙ ГОРОДСКОЙ ИСПОЛНИТЕЛЬНЫЙ КОМИТЕТ",
    "КОММУНАЛЬНОЕ УНИТАРНОЕ ПРОИЗВОДСТВЕННОЕ",
    "ПРЕДПРИЯТИЕ ПО ЭКСПЛУАТАЦИИ И РЕМОНТУ",
    "КОММУНАЛЬНЫХ ТЕПЛОВЫХ СЕТЕЙ И КОТЕЛЬНЫХ",
    "«МИНСККОММУНТЕПЛОСЕТЬ»",
)


def _guess_instruction_subject(doc: Document, doc_type: str, source_path: str = "") -> str:
    """Краткое наименование из имени файла / начала тела (НЕ подставлять «работнику» вслепую)."""
    # 1) имя файла: «ИНСТРУКЦИЯ по эксплуатации ЦТП.docx»
    name = Path(source_path or "").stem if source_path else ""
    name = re.sub(r"[_\-]+", " ", name)
    name = re.sub(r"\s+", " ", name).strip()
    if name:
        m = re.search(
            r"инструкци[яию]?\s+(по\s+эксплуатации\s+.+)$",
            name,
            re.I,
        )
        if m:
            return m.group(1).strip()[:120]
        m = re.search(
            r"(должностн\w*\s+инструкци\w*|рабоч\w*\s+инструкци\w*|инструкци\w*)\s+(.+)$",
            name,
            re.I,
        )
        if m and m.group(2).strip():
            subj = m.group(2).strip()
            if subj.lower() not in ("docx", "doc") and len(subj) >= 3:
                return subj[:120]
        # файл целиком как тема, если это не общий шаблон
        low_n = name.lower()
        if "эксплуатац" in low_n or "цтп" in low_n or "положен" in low_n:
            cleaned = re.sub(r"^(инструкци[яию]?|положен\w+)\s+", "", name, flags=re.I).strip()
            if cleaned:
                return cleaned[:120]

    blob = " ".join(p.text for p in doc.paragraphs[:30])
    low = blob.lower()
    m = re.search(
        r"(должностн\w*\s+инструкци\w*\s+)(.+?)(?:\n|$)",
        blob,
        re.I,
    )
    if m and len(m.group(2).strip()) < 120:
        return m.group(2).strip().rstrip(".")
    m = re.search(
        r"(рабоч\w*\s+инструкци\w*\s+)(.+?)(?:\n|$)",
        blob,
        re.I,
    )
    if m and len(m.group(2).strip()) < 120:
        return m.group(2).strip().rstrip(".")
    m = re.search(
        r"настоящ\w+\s+инструкци\w+.+?эксплуатаци\w+\s+(.+?)[.]",
        blob,
        re.I,
    )
    if m and 3 < len(m.group(1).strip()) < 80:
        return ("по эксплуатации " + m.group(1).strip())[:120]
    for p in doc.paragraphs[:40]:
        t = p.text.strip()
        if re.match(r"^1\.1\.", t) and len(t) < 200:
            rest = re.sub(r"^1\.1\.\s*", "", t)
            rest = re.split(r"\s+является\b", rest, maxsplit=1, flags=re.I)[0]
            if 5 < len(rest) < 100:
                return rest.strip()
    if "эксплуатац" in low:
        return "по эксплуатации"
    if "рабоч" in low:
        return "работнику"
    return ""



def _should_skip_invented_title(doc: Document, doc_type: str = "") -> bool:
    """
    Не создавать титул «с нуля», если это уже готовая инструкция по эксплуатации /
    положение без шапки МКТС: иначе агент подменяет название на «работнику» / ДИ.
    """
    paras = list(doc.paragraphs[:25])
    head = " ".join((p.text or "") for p in paras).lower()
    has_org = (
        "минсккоммунтеплосеть" in head
        or "исполнительный комитет" in head
        or "коммунальное унитарное" in head
    )
    if has_org:
        return False
    # есть «УТВЕРЖДАЮ» — титул уже задуман, можно собирать
    if "утверждаю" in head:
        return False
    hint = (doc_type or "").lower()
    if hint == "polozhenie" or "инструкция по эксплуатации" in head or "эксплуатац" in head:
        return True
    if _find_contents_index(paras) is not None:
        return True
    return False


def create_title_page_from_scratch(doc: Document, doc_type: str = "unsupported", source_path: str = "") -> int:
    """
    Если титула нет совсем — создать шапку + таблицу 3×2 + линию номера + Минск,
    как в эталоне СЛЕСАРЬ 30.07.2026.
    Вставку делать ПЕРЕД «СОДЕРЖАНИЕ», а не внутрь оглавления.
    """
    # уже есть таблица УТВЕРЖДАЮ / шапка организации — не дублировать
    if _find_title_table(doc) is not None:
        return 0
    head = " ".join(p.text for p in doc.paragraphs[:15]).lower()
    if (
        "минсккоммунтеплосеть" in head
        or "коммунальное унитарное" in head
        or "исполнительный комитет" in head
    ):
        return 0
    if _should_skip_invented_title(doc, doc_type):
        return 0

    # якорь: «СОДЕРЖАНИЕ» (если есть), иначе первая глава/пункт тела
    body_anchor = None
    for p in doc.paragraphs:
        if is_contents_heading((p.text or "").strip()):
            body_anchor = p
            break
    if body_anchor is None:
        for p in doc.paragraphs:
            t = p.text.strip()
            if _is_toc_entry_line(t):
                continue
            if re.match(r"^\d+(\s|\.)", t):
                body_anchor = p
                break
    if body_anchor is None and doc.paragraphs:
        body_anchor = doc.paragraphs[0]

    subject = _guess_instruction_subject(doc, doc_type, source_path=source_path)
    label = _doc_kind_label(doc_type, " ".join([subject, source_path or ""]))
    left_mid = subject or label
    rows = [
        [label, "УТВЕРЖДАЮ"],
        [left_mid, "Первый заместитель директора –\nглавный инженер"],
        ["", "____________  ____________\n«___»________ 2026г."],
    ]

    if body_anchor is None:
        return 0

    created = 0
    insert_point = body_anchor
    for line in ORG_HEADER_LINES:
        hp = insert_point.insert_paragraph_before(line)
        _set_para_format_no_indent(hp)
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.first_line_indent = Cm(0)
        if hp.runs:
            set_run_font(hp.runs[0], bold=True)
        else:
            run = hp.add_run(line)
            set_run_font(run, bold=True)
        created += 1
    # 2 пустых
    e1 = insert_point.insert_paragraph_before("")
    e2 = insert_point.insert_paragraph_before("")
    _set_para_format_no_indent(e1)
    _set_para_format_no_indent(e2)
    created += 2
    # таблица после последнего пустого перед телом
    _insert_title_table_after(e2, rows)
    created += 1
    ensure_title_sample_spacers(doc)
    created += 1
    return created


_CHAPTER_KEYWORD_RE = re.compile(
    r"(общие положения|общие требования|обязанност|функци|права|"
    r"ответственност|заключительн|взаимоотношен|термины|сокращен|"
    r"определен|осмотр|фазировк|объем|нормы измерен|нормы и|"
    r"измерений|испытан|авари|правила и порядок|ликвидац|"
    r"основные меры|безопасн.*обслуживан|порядок работ|"
    r"профессиональн|квалификац|описание работ)",
    re.I,
)

_NOT_CHAPTER_RE = re.compile(
    r"^(утверждаю|согласовано|утверждено|приказ|ознакомлен|"
    r"с инструкцией|сведения о|журнал|минск|коммунальн|"
    r"предприят|исполнительн|производственн|теплосет)",
    re.I,
)


def _is_chapter_heading_para(text: str) -> bool:
    """Заголовок главы: «1 ОБЩИЕ ПОЛОЖЕНИЯ» / «1. ОБЩИЕ…» (не пункт «1.1.»)."""
    t = (text or "").strip()
    if not t or len(t) > 110:
        return False
    # строки оглавления — не главы тела
    if _is_toc_entry_line(t):
        return False
    # пункты/подпункты — не главы
    if re.match(r"^\d+\.\d+", t):
        return False
    low = t.lower()
    if _NOT_CHAPTER_RE.search(low) or low in ("согласовано", "утверждаю"):
        return False
    # «1. ОБЩИЕ…» / «1.ОБЩИЕ…» — строгий детект
    m = re.match(r"^(\d+)\.\s*(.+)$", t)
    if m:
        rest = m.group(2).strip()
        # «3. специфические требования…» — текст пункта, не глава
        if rest[:1].islower():
            return False
        if _looks_like_chapter_title(rest, allow_soft=False):
            # для формы с точкой — только явные ключевые слова или ЗАГЛАВНЫЕ
            letters = [c for c in rest if c.isalpha()]
            caps = (
                letters
                and sum(1 for c in letters if c.isupper()) / len(letters) > 0.55
            )
            if caps or _CHAPTER_KEYWORD_RE.search(rest):
                return True
    # «1 ОБЩИЕ…» — можно мягче
    m = re.match(r"^(\d+)\s+(.+)$", t)
    if m and _looks_like_chapter_title(m.group(2), allow_soft=True):
        rest = m.group(2).strip()
        if rest[:1].islower():
            return False
        return True
    if low.startswith("приложение") and len(t) < 80:
        return True
    # без номера: только явные ключевые слова глав (не шапка титула ЗАГЛАВНЫМИ)
    if not re.match(r"^\d", t) and not t.endswith((":", ";", ",")) and len(t) < 100:
        if _CHAPTER_KEYWORD_RE.search(t) and t[0].isupper():
            return True
    return False


def _next_chapter_point_number(paras: list, start_idx: int, look_ahead: int = 10) -> int | None:
    """Если вскоре идёт пункт «N.1.» — вернуть номер главы N."""
    for k in range(start_idx + 1, min(start_idx + 1 + look_ahead, len(paras))):
        tk = (paras[k].text or "").strip()
        if not tk:
            continue
        m = re.match(r"^(\d+)\.1(?:\.|\s|$)", tk)
        if m:
            return int(m.group(1))
        # встретили другую уже пронумерованную главу — стоп
        if _is_chapter_heading_para(tk) and re.match(r"^\d+", tk):
            return None
        if re.match(r"^\d+\.\d+", tk):
            m2 = re.match(r"^(\d+)\.", tk)
            return int(m2.group(1)) if m2 else None
    return None


def _is_caps_chapter_continuation(text: str) -> bool:
    """Вторая строка заголовка ЗАГЛАВНЫМИ: «КАБЕЛЬНЫХ ЛИНИЙ»."""
    t = (text or "").strip()
    if not t or len(t) >= 80 or re.match(r"^\d", t) or t.endswith((":", ";", ",")):
        return False
    if _NOT_CHAPTER_RE.search(t.lower()):
        return False
    letters = [c for c in t if c.isalpha()]
    if not letters:
        return False
    return sum(1 for c in letters if c.isupper()) / len(letters) > 0.7


def repair_missing_chapter_numbers(doc: Document) -> dict:
    """
    Восстановить номер у заголовков без цифры и склеить вторую строку ЗАГЛАВНЫМИ:
    «4 ОБЪЕМ…» + «КАБЕЛЬНЫХ ЛИНИЙ» → «4 ОБЪЕМ … КАБЕЛЬНЫХ ЛИНИЙ».
    """
    report = {"numbered": 0, "merged": 0}
    paras = list(doc.paragraphs)
    start, end = _body_bounds(paras)
    toc_end = _toc_region_end(paras)
    if toc_end is not None and toc_end > start:
        start = toc_end
    i = start
    while i < end and i < len(paras):
        t = (paras[i].text or "").strip()
        if not t or is_paragraph_in_table(paras[i]):
            i += 1
            continue
        if _is_toc_entry_line(t) or is_contents_heading(t):
            i += 1
            continue
        if re.match(r"^\d+\.\d+", t):
            i += 1
            continue

        is_numbered_ch = bool(
            re.match(r"^\d+(\s|\.)", t) and _is_chapter_heading_para(t)
        )
        # Только короткие заголовки (не абзацы текста с словом «испытание» и т.п.)
        is_bare_ch = bool(
            not re.match(r"^\d", t) and _is_chapter_heading_para(t)
        )
        if is_bare_ch:
            letters = [c for c in t if c.isalpha()]
            caps_ratio = (
                sum(1 for c in letters if c.isupper()) / len(letters) if letters else 0
            )
            # «ОБЩИЕ ПОЛОЖЕНИЯ» — да; длинное предложение с ключевым словом — нет
            if caps_ratio < 0.55 or len(t) > 100:
                is_bare_ch = False
        if not is_numbered_ch and not is_bare_ch:
            i += 1
            continue

        # склеить продолжения ЗАГЛАВНЫМИ (и для уже пронумерованных глав)
        while True:
            j = i + 1
            while j < end and not (paras[j].text or "").replace("\xa0", " ").strip():
                j += 1
            if j >= end:
                break
            t2 = (paras[j].text or "").strip()
            if re.match(r"^\d+\.\d+", t2) or (
                _is_chapter_heading_para(t2) and re.match(r"^\d", t2)
            ):
                break
            if not _is_caps_chapter_continuation(t2):
                break
            # убрать пустые между заголовком и продолжением
            for k in range(j - 1, i, -1):
                if not (paras[k].text or "").replace("\xa0", " ").strip():
                    _delete_paragraph(paras[k])
            paras = list(doc.paragraphs)
            end = min(end, len(paras))
            # индекс продолжения мог сдвинуться
            j = i + 1
            while j < end and not (paras[j].text or "").replace("\xa0", " ").strip():
                j += 1
            if j >= end:
                break
            t2 = (paras[j].text or "").strip()
            if not _is_caps_chapter_continuation(t2):
                break
            t = f"{(paras[i].text or '').strip()} {t2}".strip()
            _set_runs(paras[i], t)
            _delete_paragraph(paras[j])
            report["merged"] += 1
            paras = list(doc.paragraphs)
            end = min(end, len(paras))

        t = (paras[i].text or "").strip()
        if not re.match(r"^\d+(\s|\.)", t) and _CHAPTER_KEYWORD_RE.search(t):
            n = _next_chapter_point_number(paras, i)
            if n is not None:
                new_t = f"{n} {t}"
                _set_runs(paras[i], new_t)
                report["numbered"] += 1
        i += 1
    return report


def _table_blob_lower(table) -> str:
    texts = []
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                texts.append(p.text or "")
    return " ".join(texts).lower()


def _is_title_approve_table(table) -> bool:
    """Титульная таблица «УТВЕРЖДАЮ» — не первая попавшаяся таблица тела."""
    blob = _table_blob_lower(table)
    return "утверждаю" in blob or (
        "согласовано" in blob and ("инструкция" in blob or "утвержд" in blob)
    )


def _body_bounds(paras: list) -> tuple[int, int]:
    """Индексы тела: после титула … до подписантов / ознакомления."""
    start = 0
    for i, p in enumerate(paras[:80]):
        t = p.text.strip()
        if is_contents_heading(t):
            start = i
            break
        if re.match(r"^\d+(\s|\.)", t) and (
            _looks_like_chapter_title(re.sub(r"^\d+\.?\s*", "", t))
            or re.match(r"^\d+\.\d+", t)
        ):
            start = i
            break
        if re.fullmatch(r"минск\s+20\d{2}", t.lower()):
            start = i + 1
    end = len(paras)
    for i in range(len(paras) - 1, max(start, len(paras) // 3), -1):
        low = paras[i].text.strip().lower()
        if low.startswith(("разработал", "согласовано", "с должностной", "с рабочей")):
            end = i
            break
        if "ознакомлен" in low and "инструкц" in low:
            end = i
            break
        if _is_note_start(paras[i].text):
            end = i
            break
    return start, end


def remove_empty_paragraphs_in_body(doc: Document) -> int:
    """
    Убрать пустые строки внутри текста глав (между пунктами).
    НЕ трогать: ровно одну пустую строку перед заголовком новой главы;
    пустую перед «Разработал» / «Согласовано» / «Содержание».
    """
    paras = list(doc.paragraphs)
    if not paras:
        return 0
    start, end = _body_bounds(paras)

    def _empty(p) -> bool:
        return not (p.text or "").replace("\xa0", " ").strip()

    def _keep_before(nxt: str) -> bool:
        nxt_low = (nxt or "").strip().lower()
        if not nxt:
            return False
        if _is_chapter_heading_para(nxt):
            return True
        if is_contents_heading(nxt):
            return True
        if nxt_low.startswith(("разработал", "согласовано")):
            return True
        return False

    removed = 0
    # удалять с конца, чтобы не сбить XML
    for i in range(end - 1, start, -1):
        if not _empty(paras[i]):
            continue
        keep = False
        if i + 1 < len(paras):
            nxt = (paras[i + 1].text or "").replace("\xa0", " ").strip()
            if _keep_before(nxt):
                keep = True
            elif not nxt:
                # цепочка пустых: оставить только ту, что прямо перед якорем
                k = i + 1
                while k < end and _empty(paras[k]):
                    k += 1
                if k < len(paras) and _keep_before(
                    (paras[k].text or "").replace("\xa0", " ").strip()
                ):
                    if k - i > 1:
                        keep = False
                    else:
                        keep = True
        if keep:
            continue
        _delete_paragraph(paras[i])
        removed += 1
    return removed


def apply_table_fonts(doc: Document, body_pt: int = TABLE_FONT_PT) -> int:
    """
    Шрифт в таблицах: титул «УТВЕРЖДАЮ» — 14 пт, все остальные (тело) — 12 пт.
    Абзацный отступ в ячейках = 0.
    """
    changed = 0
    for table in doc.tables:
        is_title = _is_title_approve_table(table)
        size = Pt(TITLE_TABLE_FONT_PT if is_title else body_pt)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Cm(0)
                    zero_para_spacing(p)
                    set_single_line_spacing(p)
                    for run in p.runs:
                        low = (run.text or "").strip().lower()
                        # подпись «номер инструкции» на титуле — 12 пт
                        want = (
                            Pt(TITLE_INSTR_NUMBER_PT)
                            if low.startswith("номер инструкции")
                            else size
                        )
                        if run.font.size != want or run.font.size is None:
                            set_run_font(
                                run,
                                font_name=run.font.name or "Times New Roman",
                                font_size=want,
                                bold=run.bold,
                            )
                            changed += 1
    return changed


def separate_contents_onto_own_page(doc: Document) -> dict:
    """
    После «Содержания» / «Оглавления» — всегда разрыв страницы;
    первая глава текста начинается с самого верха новой страницы
    (без пустых строк перед ней).
    Также разрыв перед самим заголовком содержания (от титула).
    """
    result = {
        "found": False,
        "break_before_toc": False,
        "break_before_body": False,
        "empties_cleared_before_body": 0,
    }
    paras = list(doc.paragraphs)
    toc_idx = None
    for i, p in enumerate(paras):
        if is_contents_heading(p.text):
            toc_idx = i
            break
    if toc_idx is None:
        return result
    result["found"] = True
    _set_page_break_before(paras[toc_idx], True)
    result["break_before_toc"] = True

    # конец содержания: первая глава, за которой вскоре идёт пункт «1.1.» / «2.1.»
    body_start = None
    for i in range(toc_idx + 1, len(paras)):
        t = (paras[i].text or "").strip()
        if not t:
            continue
        if _is_toc_entry_line(t):
            continue
        if _is_chapter_heading_para(t):
            saw_point = False
            for k in range(i + 1, min(i + 12, len(paras))):
                tk = (paras[k].text or "").strip()
                if not tk:
                    continue
                if re.match(r"^\d+\.\d+", tk):
                    saw_point = True
                    break
                if _is_chapter_heading_para(tk) or is_contents_heading(tk):
                    break
                if _is_toc_entry_line(tk):
                    saw_point = False
                    break
            if saw_point:
                body_start = i
                break
        elif re.match(r"^\d+\.\d+", t):
            body_start = i
            break

    if body_start is not None and body_start > toc_idx:
        # убрать пустые строки сразу перед первой главой — глава с верха страницы
        cleared = 0
        for j in range(body_start - 1, toc_idx, -1):
            if (paras[j].text or "").replace("\xa0", " ").strip():
                break
            _delete_paragraph(paras[j])
            cleared += 1
        if cleared:
            result["empties_cleared_before_body"] = cleared
            paras = list(doc.paragraphs)
            # пересчитать индекс первой главы
            body_start = None
            for i in range(toc_idx + 1, len(paras)):
                t = (paras[i].text or "").strip()
                if _is_chapter_heading_para(t) or re.match(r"^\d+\.\d+", t):
                    if not _is_toc_entry_line(t):
                        body_start = i
                        break
        if body_start is not None:
            _set_page_break_before(paras[body_start], True)
            # на всякий случай: у главы с разрывом — без отступа, по центру если глава
            if _is_chapter_heading_para(paras[body_start].text):
                paras[body_start].alignment = WD_ALIGN_PARAGRAPH.CENTER
                paras[body_start].paragraph_format.first_line_indent = Cm(0)
            result["break_before_body"] = True
            result["body_start"] = body_start
    result["toc_idx"] = toc_idx
    return result


def collapse_double_spaces_in_document(doc: Document) -> int:
    """
    Убрать двойные (и более) пробелы по всему тексту, включая таблицы.
    Табы у подписантов сохраняются; пробелы внутри частей схлопываются.
    """
    changed = 0
    for p in iter_all_paragraphs(doc):
        raw = p.text or ""
        if not raw:
            continue
        if "\t" in raw or _paragraph_keeps_signatory_tab(raw):
            parts = raw.split("\t")
            new_parts = [
                re.sub(r"[ \xa0\u202f\u2009\u200a]{2,}", " ", part) for part in parts
            ]
            new = "\t".join(new_parts)
        else:
            new = raw.replace("\xa0", " ").replace("\u202f", " ")
            new = new.replace("\u2009", " ").replace("\u200a", " ")
            new = re.sub(r" {2,}", " ", new)
        if new == raw:
            continue
        bold = p.runs[0].bold if p.runs else None
        _set_runs(p, new, bold=bold)
        if "\x0b" in new:
            _expand_vertical_tabs_to_br(p._element)
        changed += 1
    return changed


def format_chapter_headings(doc: Document) -> dict:
    """
    Для ВСЕХ документов — как Heading 1 эталона КЛ 31.07.2026:
    центр, жирный, капслок, keepNext, отступ 0;
    перед заголовком — пустая строка (кроме главы после содержания).
    """
    report = {"centered": 0, "spacers": 0, "styled": 0}
    report["spacers"] = ensure_empty_line_before_chapters(doc)
    for p in doc.paragraphs:
        if is_paragraph_in_table(p):
            continue
        if not _is_chapter_heading_para(p.text):
            continue
        prev_al = p.alignment
        apply_chapter_heading_format(p, doc=doc)
        report["styled"] += 1
        if prev_al != WD_ALIGN_PARAGRAPH.CENTER:
            report["centered"] += 1
    return report


def apply_spacing_and_clear_body_empties(doc: Document) -> dict:
    """
    Финальная зачистка интервалов и пустых строк:
    — одинарный межстрочный в тексте/титуле/таблицах;
    — space_before = space_after = 0 везде (в т.ч. таблицы);
    — пустые строки в теле убрать (кроме 1 перед главой / перед Разработал);
    — перед каждым заголовком главы — одна пустая строка.
    """
    report = {"spacing": 0, "empties": 0, "chapter_spacers": 0}
    # сначала интервалы
    for p in iter_all_paragraphs(doc):
        if set_single_line_spacing(p):
            report["spacing"] += 1
        if zero_para_spacing(p):
            report["spacing"] += 1
        if is_paragraph_in_table(p):
            p.paragraph_format.first_line_indent = Cm(0)
    # пустые в теле
    try:
        report["empties"] = remove_empty_paragraphs_in_body(doc)
    except Exception:
        pass
    try:
        report["chapter_spacers"] = ensure_empty_line_before_chapters(doc)
    except Exception:
        pass
    # после вставок пустых перед главами — снова 0 before/after
    for p in doc.paragraphs:
        zero_para_spacing(p)
        set_single_line_spacing(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    zero_para_spacing(p)
                    set_single_line_spacing(p)
                    p.paragraph_format.first_line_indent = Cm(0)
    return report


def ensure_empty_line_before_chapters(doc: Document) -> int:
    """
    Перед КАЖДЫМ заголовком главы — всегда одна пустая строка.
    Исключение: глава с разрывом страницы после содержания — с верха листа,
    пустую перед ней не ставим (и лишние пустые перед ней убираем).
    """
    paras = list(doc.paragraphs)
    if not paras:
        return 0
    start, end = _body_bounds(paras)
    added = 0
    for i in range(min(end, len(paras)) - 1, max(start - 1, 0), -1):
        if i < 0 or i >= len(paras):
            continue
        if not _is_chapter_heading_para(paras[i].text):
            continue
        # после содержания / с разрывом страницы — глава с верха, без пустой
        if _paragraph_has_page_break_before(paras[i]):
            continue
        # уже есть пустая сразу перед заголовком
        if i > 0 and not (paras[i - 1].text or "").replace("\xa0", " ").strip():
            continue
        j = i - 1
        while j >= 0 and not (paras[j].text or "").replace("\xa0", " ").strip():
            j -= 1
        if j < 0:
            continue
        new_p = paras[i].insert_paragraph_before("")
        pf = new_p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        added += 1
        paras = list(doc.paragraphs)
    return added


def collapse_title_spacer_empties(doc: Document) -> int:
    """
    Убрать неоправданные пустые строки на титуле.
    Эталон: после шапки — до 2 пустых, затем таблица, затем сразу линия номера.
    Пустые МЕЖДУ таблицей и «____» / «номер» / «Минск» — удалить.
    """
    body = doc.element.body
    children = list(body.iterchildren())
    removed = 0

    # найти индекс title-table
    tbl_idx = None
    for i, child in enumerate(children):
        if child.tag == qn("w:tbl"):
            # проверить содержимое
            texts = [t.text or "" for t in child.iter(qn("w:t"))]
            blob = " ".join(texts).lower()
            if "утверждаю" in blob or "инструкция" in blob:
                tbl_idx = i
                break
    if tbl_idx is None:
        return 0

    def _p_text(el) -> str:
        return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()

    def _is_empty_p(el) -> bool:
        if el.tag != qn("w:p"):
            return False
        t = _p_text(el)
        return not t or t.replace("\t", "").strip() == ""

    # 1) между шапкой и таблицей: оставить не больше 2 пустых подряд в конце
    # найти последний непустой абзац перед таблицей
    pre_empties = []
    for j in range(tbl_idx - 1, -1, -1):
        el = children[j]
        if el.tag != qn("w:p"):
            break
        if _is_empty_p(el):
            pre_empties.append(el)
        else:
            break
    # pre_empties собраны снизу вверх
    for el in pre_empties[2:]:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            removed += 1

    # обновить список после удалений
    children = list(body.iterchildren())
    tbl_idx = None
    for i, child in enumerate(children):
        if child.tag == qn("w:tbl"):
            texts = [t.text or "" for t in child.iter(qn("w:t"))]
            blob = " ".join(texts).lower()
            if "утверждаю" in blob or "инструкция" in blob:
                tbl_idx = i
                break
    if tbl_idx is None:
        return removed

    # 2) после таблицы до маркера (____ / номер / Минск) — удалить пустые
    stop_idx = None
    for j in range(tbl_idx + 1, min(len(children), tbl_idx + 40)):
        el = children[j]
        if el.tag != qn("w:p"):
            continue
        t = _p_text(el)
        low = t.lower()
        if t.startswith("___") or "номер инструкции" in low or low.startswith("минск"):
            stop_idx = j
            break
        if re.match(r"^\d+\s+[А-ЯЁA-Z]", t):
            stop_idx = j
            break
    if stop_idx is not None:
        for j in range(tbl_idx + 1, stop_idx):
            el = children[j]
            if _is_empty_p(el):
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
                    removed += 1
            elif el.tag == qn("w:p"):
                # остатки даты утверждения вне таблицы — удалить
                t = _p_text(el)
                if DATE_RE.search(t) and ("«" in t or "_" in t) and len(t) < 80:
                    parent = el.getparent()
                    if parent is not None:
                        parent.remove(el)
                        removed += 1

    return removed


def _insert_empty_paragraph_after(paragraph) -> object:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    # вернуть обёртку через parent document — проще работать с element
    return new_p


def ensure_title_sample_spacers(doc: Document) -> int:
    """
    Обязательные пустые строки титула как в эталоне СЛЕСАРЬ 30.07.2026:
    — 2 пустые между «МИНСККОММУНТЕПЛОСЕТЬ» и таблицей;
    — после таблицы: линия + «номер инструкции»;
    — ~15 пустых до «Минск 20xx».
    """
    changed = 0
    body = doc.element.body
    children = list(body.iterchildren())

    def _p_text(el) -> str:
        return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()

    def _is_empty_p(el) -> bool:
        if el.tag != qn("w:p"):
            return False
        t = _p_text(el)
        return not t or t.replace("\t", "").strip() == ""

    tbl_idx = None
    tbl_el = None
    for i, child in enumerate(children):
        if child.tag != qn("w:tbl"):
            continue
        blob = " ".join(t.text or "" for t in child.iter(qn("w:t"))).lower()
        if "утверждаю" in blob or "инструкция" in blob:
            tbl_idx = i
            tbl_el = child
            break
    if tbl_idx is None or tbl_el is None:
        return 0

    # --- перед таблицей: ровно TITLE_EMPTY_BEFORE_TABLE пустых ---
    pre = []
    j = tbl_idx - 1
    while j >= 0 and children[j].tag == qn("w:p") and _is_empty_p(children[j]):
        pre.append(children[j])
        j -= 1
    # pre[0] — ближайший к таблице, pre[-1] — у шапки
    while len(pre) > TITLE_EMPTY_BEFORE_TABLE:
        el = pre.pop()  # лишний у шапки
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            changed += 1
    while len(pre) < TITLE_EMPTY_BEFORE_TABLE:
        new_p = OxmlElement("w:p")
        tbl_el.addprevious(new_p)
        pre.insert(0, new_p)
        changed += 1

    children = list(body.iterchildren())
    tbl_idx = None
    for i, child in enumerate(children):
        if child is tbl_el or (child.tag == qn("w:tbl") and "утверждаю" in " ".join(
            t.text or "" for t in child.iter(qn("w:t"))
        ).lower()):
            tbl_idx = i
            tbl_el = child
            break

    # --- после таблицы: underline + номер инструкции ---
    # найти/создать
    after = children[tbl_idx + 1 :] if tbl_idx is not None else []
    # пропустить пустые сразу после таблицы (их быть не должно)
    k = 0
    while k < len(after) and after[k].tag == qn("w:p") and _is_empty_p(after[k]):
        parent = after[k].getparent()
        if parent is not None:
            parent.remove(after[k])
            changed += 1
        k += 1
    children = list(body.iterchildren())
    tbl_idx = next(i for i, c in enumerate(children) if c is tbl_el)

    def _ensure_para_after(ref_el, text: str, size_pt: int = 14) -> object:
        # если следующий уже нужный — ок
        nxt = ref_el.getnext()
        if nxt is not None and nxt.tag == qn("w:p") and _p_text(nxt) == text.strip():
            _set_oxml_para_font(nxt, size_pt=size_pt)
            return nxt
        # если следующий похож (underline / номер)
        if nxt is not None and nxt.tag == qn("w:p"):
            nt = _p_text(nxt)
            if text.startswith("___") and nt.startswith("___"):
                _set_oxml_para_font(nxt, size_pt=size_pt)
                return nxt
            if (
                TITLE_INSTR_NUMBER_LABEL in text.lower()
                and TITLE_INSTR_NUMBER_LABEL in nt.lower()
            ):
                _set_oxml_para_font(nxt, size_pt=size_pt)
                return nxt
        new_p = OxmlElement("w:p")
        # один run с текстом
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), "Times New Roman")
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(size_pt * 2)))
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)
        new_p.append(r)
        ref_el.addnext(new_p)
        return new_p

    line_el = _ensure_para_after(tbl_el, "___________________________", size_pt=14)
    num_el = _ensure_para_after(
        line_el, TITLE_INSTR_NUMBER_LABEL, size_pt=TITLE_INSTR_NUMBER_PT
    )
    changed += 1

    # --- пустые до Минск ---
    children = list(body.iterchildren())
    # найти num_el и minsk
    try:
        num_idx = next(i for i, c in enumerate(children) if c is num_el)
    except StopIteration:
        num_idx = tbl_idx + 2
    minsk_idx = None
    for i in range(num_idx + 1, min(len(children), num_idx + 40)):
        if children[i].tag == qn("w:p") and _p_text(children[i]).lower().startswith("минск"):
            minsk_idx = i
            break
    if minsk_idx is None:
        return changed

    empties = []
    for i in range(num_idx + 1, minsk_idx):
        if children[i].tag == qn("w:p") and _is_empty_p(children[i]):
            empties.append(children[i])
        elif children[i].tag == qn("w:p") and _p_text(children[i]):
            # непустой между номером и минском — не трогаем содержимое, только пустые
            pass

    # посчитать только подряд идущие пустые сразу после номера
    run = []
    for i in range(num_idx + 1, minsk_idx):
        el = children[i]
        if el.tag == qn("w:p") and _is_empty_p(el):
            run.append(el)
        else:
            break

    while len(run) > TITLE_EMPTY_BEFORE_MINSK:
        el = run.pop()
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            changed += 1
    minsk_el = children[minsk_idx]
    while len(run) < TITLE_EMPTY_BEFORE_MINSK:
        new_p = OxmlElement("w:p")
        minsk_el.addprevious(new_p)
        run.append(new_p)
        changed += 1

    return changed


def apply_body_first_indents(doc: Document) -> int:
    """
    Титул (шапка + таблица УТВЕРЖДАЮ + номер + Минск) — отступ 0.
    Текст пунктов тела — 1,25 см (кроме заголовков, содержания, подписантов).
    В ТАБЛИЦАХ отступ не применяется.
    Межстрочный интервал текста и титула — одинарный.
    """
    changed = 0
    # эталон КЛ: Body Text firstLine = 709 twips (1,25 см)
    first_indent = Twips(ETALON_FIRST_INDENT_TWIPS)
    paras = list(doc.paragraphs)
    start, end = _body_bounds(paras)

    # 1) ВСЕ таблицы — отступ 0 + одинарный интервал + before/after = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    if pf.first_line_indent and pf.first_line_indent.cm > 0.05:
                        changed += 1
                    if pf.left_indent and pf.left_indent.cm > 0.05:
                        changed += 1
                    pf.first_line_indent = Cm(0)
                    pf.left_indent = Cm(0)
                    if set_single_line_spacing(p):
                        changed += 1
                    if zero_para_spacing(p):
                        changed += 1

    # 2) абзацы тела/титула (не из таблиц)
    for i, p in enumerate(paras):
        if is_paragraph_in_table(p):
            continue
        text = p.text.strip()
        if set_single_line_spacing(p):
            changed += 1
        if zero_para_spacing(p):
            changed += 1

        # титул / до тела
        if i < start:
            pf = p.paragraph_format
            if (pf.first_line_indent and pf.first_line_indent.cm > 0.05) or (
                pf.left_indent and pf.left_indent.cm > 0.05
            ):
                changed += 1
            pf.first_line_indent = Cm(0)
            if pf.left_indent is None or pf.left_indent.cm < 20:
                pf.left_indent = Cm(0)
            continue

        # зона подписантов и далее
        if i >= end:
            pf = p.paragraph_format
            pf.first_line_indent = Cm(0)
            continue

        if not text:
            continue

        # содержание и строки оглавления — без отступа
        if is_contents_heading(text) or _is_toc_entry_line(text):
            p.paragraph_format.first_line_indent = Cm(0)
            continue
        # заголовки глав — по центру, отступ 0
        if _is_chapter_heading_para(text):
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if is_heading_like(text):
            p.paragraph_format.first_line_indent = Cm(0)
            continue

        # даты / строки с табом подписанта — не ставить 1,25
        if text.startswith("\t") or (
            "«" in text and "___" in text and len(text) < 80
        ):
            p.paragraph_format.first_line_indent = Cm(0)
            continue

        if needs_first_indent(text, p.alignment):
            pf = p.paragraph_format
            if pf.left_indent is not None and 0 < pf.left_indent.mm < 20:
                pf.left_indent = Cm(0)
            cur = pf.first_line_indent.cm if pf.first_line_indent else 0.0
            if abs(cur - FIRST_INDENT_CM) > 0.05:
                changed += 1
            pf.first_line_indent = first_indent
            if p.alignment is None or p.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            if (
                p.paragraph_format.first_line_indent
                and p.paragraph_format.first_line_indent.cm > 0.1
            ):
                p.paragraph_format.first_line_indent = Cm(0)
                changed += 1
    return changed


def _set_signatory_tab_stop(paragraph, pos_cm: float = SIGN_FIO_TAB_CM) -> None:
    """Табуляция для ФИО на заданной позиции (см от левого поля)."""
    pPr = paragraph._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:tabs"))):
        pPr.remove(old)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(int(Cm(pos_cm).twips)))
    tabs.append(tab)
    pPr.append(tabs)


def _clear_paragraph_tab_stops(paragraph) -> None:
    """Убрать заданные позиции табов — в РИ используются обычные табы документа."""
    pPr = paragraph._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:tabs"))):
        pPr.remove(old)


def document_text_blob(doc: Document, *, limit_paras: int = 80) -> str:
    """Текст документа включая ячейки таблиц (титул часто только в таблице)."""
    parts: list[str] = []
    for i, p in enumerate(doc.paragraphs):
        if i >= limit_paras:
            break
        parts.append(p.text or "")
    for table in doc.tables[:8]:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    return "\n".join(parts)


def resolve_instruction_doc_type(
    *,
    doc_type: str = "",
    doc: Document | None = None,
    source_path: str = "",
) -> str:
    """
    Надёжно определить РИ/ДИ.
    Важно: «РАБОЧАЯ ИНСТРУКЦИЯ» часто только в таблице титула — paragraphs[] её не видят.
    """
    hint = (doc_type or "").lower().strip()
    name = Path(source_path).name.lower() if source_path else ""
    blob = ""
    if doc is not None:
        blob = document_text_blob(doc).lower()
    packed = f"{name}\n{blob}"

    # имя «ДИ …» важнее ошибочного hint=polozhenie
    if re.search(r"(^|[\s_])ди([\s_.]|$)", name) or re.match(r"^ди[\s_]", name):
        return "dolzhnostnaya_instrukciya"

    if hint in ("rabochaya_instrukciya", "dolzhnostnaya_instrukciya", "polozhenie", "instrukciya_ot", "prikaz"):
        return hint

    if (
        "рабочая инструкция" in packed
        or ("рабоч" in name and "инструкц" in name)
        or re.search(r"(^|[\s_])ри([\s_.]|$)", name)
    ):
        return "rabochaya_instrukciya"
    if (
        "должностная инструкция" in packed
        or ("должностн" in name and "инструкц" in name)
        or re.search(r"(^|[\s_])ди([\s_.]|$)", name)
    ):
        return "dolzhnostnaya_instrukciya"
    if "положение" in packed:
        return "polozhenie"
    if (
        "инструкция по эксплуатации" in packed
        or "настоящая инструкция" in packed
        or (
            "инструкц" in name
            and "рабоч" not in name
            and "должностн" not in name
            and "охране труда" not in name
            and "иот" not in name
        )
    ):
        return "polozhenie"
    return hint or "unsupported"


def _is_ri_signatory_style(doc_type: str = "", doc: Document | None = None) -> bool:
    """Устарело: заголовок везде «СОГЛАСОВАНО». Оставлено для совместимости вызовов."""
    dtype = resolve_instruction_doc_type(doc_type=doc_type, doc=doc)
    return dtype == "rabochaya_instrukciya"


def _is_di_signatory_style(doc_type: str = "", doc: Document | None = None) -> bool:
    """Устарело: ФИО везде без пробела. Оставлено для совместимости вызовов."""
    dtype = resolve_instruction_doc_type(doc_type=doc_type, doc=doc)
    return dtype == "dolzhnostnaya_instrukciya"


def _format_signatory_fio(name: str, *, ri_style: bool = False, di_style: bool = False) -> str:
    """
    Для ВСЕХ документов: И.О.Фамилия без пробела — пример И.И.Иванов.
    """
    _ = ri_style
    _ = di_style
    compact = re.sub(
        r"([А-ЯЁA-Z])\.\s*([А-ЯЁA-Z])\.\s*",
        r"\1.\2.",
        (name or "").replace(" ", ""),
    )
    m = re.match(r"^([А-ЯЁA-Z]\.[А-ЯЁA-Z]\.)(.+)$", compact)
    if not m:
        return compact
    return f"{m.group(1)}{m.group(2)}"


def _ri_filler_tabs(title: str) -> str:
    """Число обычных табов до ФИО — как в эталоне СЛЕСАРЬ 30.07.2026."""
    length = len((title or "").strip())
    title_units = max(1, (length + RI_SIGN_CHARS_PER_TAB - 1) // RI_SIGN_CHARS_PER_TAB)
    n = max(3, RI_SIGN_TAB_TARGET_UNITS - title_units)
    return "\t" * n


def _apply_signatory_line_format(
    paragraph, *, ri_style: bool = False, line_spacing: float = SIGN_LINE_SPACING_RI
) -> None:
    """
    Для ВСЕХ документов: отступ 0 в блоке подписантов.
    line_spacing: 1.0 — внутри многострочной должности одного подписанта;
                  1.5 — между подписантами / строка с ФИО / заголовки блока.
    """
    _ = ri_style
    _set_para_format_no_indent(paragraph)
    paragraph.paragraph_format.line_spacing = line_spacing


def _insert_paragraph_after_signatory(paragraph, text: str = ""):
    """Вставить абзац сразу после данного (для переноса должности)."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        _set_runs(p, text)
    return p


def _wrap_job_title_for_signature(title: str, max_chars: int = SIGN_TITLE_MAX_CHARS) -> str:
    """
    Для ВСЕХ документов: если должность доходит до середины строки
    (нет места для рукописной подписи перед ФИО на 12 см) —
    разделить на строки по пробелам.
    Возвращает текст с \\n между частями (далее — отдельные абзацы, не Shift+Enter).
    """
    title = re.sub(r"\s+", " ", (title or "").strip())
    if not title or len(title) <= max_chars:
        return title
    words = title.split(" ")
    lines: list[str] = []
    cur = ""
    for w in words:
        trial = f"{cur} {w}".strip() if cur else w
        if len(trial) <= max_chars:
            cur = trial
            continue
        if cur:
            lines.append(cur)
        if len(w) > max_chars:
            while len(w) > max_chars:
                lines.append(w[:max_chars])
                w = w[max_chars:]
            cur = w
        else:
            cur = w
    if cur:
        lines.append(cur)
    return "\n".join(lines) if lines else title


def _write_signatory_title_name(
    paragraph,
    title: str,
    name: str,
    *,
    ri_style: bool = False,
    di_style: bool = False,
) -> None:
    """
    Должность + таб ≥12 см + ФИО; длинную должность — отдельными абзацами.
    Внутри должности одного подписанта — одинарный интервал;
    строка с ФИО (между подписантами) — полуторный.
    """
    title = re.sub(r"\s+", " ", title).strip()
    title = _wrap_job_title_for_signature(title)
    name = _format_signatory_fio(name, ri_style=ri_style, di_style=di_style)
    lines = [ln.strip() for ln in title.split("\n") if ln.strip()] or [title]
    _set_signatory_tab_stop(paragraph, SIGN_FIO_TAB_CM)

    if len(lines) == 1:
        _apply_signatory_line_format(paragraph, ri_style=ri_style, line_spacing=SIGN_LINE_SPACING_RI)
        _set_signatory_tab_stop(paragraph, SIGN_FIO_TAB_CM)
        _set_runs(paragraph, f"{lines[0]}\t{name}")
        return

    # первая часть должности — одинарный
    _apply_signatory_line_format(paragraph, ri_style=ri_style, line_spacing=1.0)
    _set_runs(paragraph, lines[0])
    anchor = paragraph
    for mid in lines[1:-1]:
        np = _insert_paragraph_after_signatory(anchor, mid)
        _apply_signatory_line_format(np, ri_style=ri_style, line_spacing=1.0)
        anchor = np
    # последняя часть + ФИО — полуторный (граница между подписантами)
    last = _insert_paragraph_after_signatory(anchor, "")
    _apply_signatory_line_format(last, ri_style=ri_style, line_spacing=SIGN_LINE_SPACING_RI)
    _set_signatory_tab_stop(last, SIGN_FIO_TAB_CM)
    _set_runs(last, f"{lines[-1]}\t{name}")


def normalize_handwritten_sign_date(text: str) -> str:
    """
    Дата у подписантов должна умещаться в одну строку (под ФИО на 12 см).
    Длинные подчёркивания сжимаем; пробел между годом и «г» убираем: 2026г.
    Короткий вид не обязателен — важно, чтобы строка не переносилась.
    """
    raw = text.replace("\xa0", " ").strip().lstrip("\t ")
    m_year = re.search(r"(20\d{2})", raw)
    year = m_year.group(1) if m_year else "2026"

    # разобрать «день» и «месяц» из бланка, если похоже на дату
    m = re.match(
        r"^[«\"']?\s*(?P<day>_+)\s*[»\"']?\s*(?P<month>_*)\s*20\d{2}\s*г\.?\s*$",
        raw,
    )
    if m:
        day = m.group("day")
        month = m.group("month") or "______"
        # слишком длинные бланки не помещаются до конца строки — укоротить
        if len(day) >= 5:
            day = "___"
        if len(month) > 10:
            month = "________"
        return f"«{day}»{month} {year}г."

    # запасной компактный бланк (помещается в строку)
    return f"«___»________ {year}г."


def _write_signatory_date_aligned(
    paragraph,
    date_text: str,
    *,
    ri_style: bool = False,
) -> None:
    """Дата на том же уровне ≥12 см, что и ФИО; без переноса. Для ВСЕХ документов."""
    d = normalize_handwritten_sign_date(date_text)
    _apply_signatory_line_format(paragraph, ri_style=ri_style)
    _set_signatory_tab_stop(paragraph, SIGN_FIO_TAB_CM)
    _set_runs(paragraph, "\t" + d)


def _set_keep_with_next(paragraph, enabled: bool = True) -> None:
    """Word: «не отрывать от следующего» — подписи не уезжают на пустую страницу."""
    pPr = paragraph._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:keepNext"))):
        pPr.remove(old)
    if enabled:
        pPr.append(OxmlElement("w:keepNext"))


def _set_page_break_before(paragraph, enabled: bool = True) -> None:
    """Разрыв страницы перед абзацем (перенос абзаца на лист с подписями)."""
    pPr = paragraph._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:pageBreakBefore"))):
        pPr.remove(old)
    if enabled:
        pPr.insert(0, OxmlElement("w:pageBreakBefore"))


def _paragraph_has_page_break_before(paragraph) -> bool:
    pPr = paragraph._element.pPr
    if pPr is None:
        return False
    return pPr.find(qn("w:pageBreakBefore")) is not None


def _is_signatory_job_line(text: str) -> bool:
    """Строка похожа на подписанта (должность и/или ФИО)."""
    t = (text or "").strip()
    if not t:
        return False
    low = t.lower()
    if low in ("визы", "виза") or low.startswith("визы"):
        return True
    if low.startswith("разработал") or low.startswith("согласовано") or low.startswith("согласовал"):
        return True
    jobs = (
        "начальник",
        "ведущий",
        "заместитель",
        "инженер",
        "мастер",
        "специалист",
        "директор",
    )
    if any(k in low for k in jobs) and (FIO_RE.search(t) or len(t) < 80):
        return True
    if FIO_RE.fullmatch(t.replace(" ", "")):
        return True
    return False


def _is_note_start(text: str) -> bool:
    """
    Начало / весь абзац «Примечание…».
    Весь абзац, начинающийся со слова «Примечание» / «Примечания» / «Прим.»,
    переносится в конец перед «Разработал:».
    """
    low = (text or "").strip().lower().replace("ё", "е")
    if not low:
        return False
    # убрать ведущую нумерацию пункта: «3.5. Примечание…»
    low2 = re.sub(r"^\d+(\.\d+)*\.\s*", "", low)
    if re.match(r"^примечан(ие|ия)\b", low2):
        return True
    if re.match(r"^прим\.?\b", low2):
        return True
    return False


def _is_note_structure_break(text: str) -> bool:
    """
    Граница, на которой блок примечания заканчивается.
    Важно: простая нумерация «1. …» / «2. …» внутри примечания — НЕ граница
    (иначе в конец уезжает только заголовок «Примечания:»).
    """
    t = (text or "").strip()
    if not t:
        return False
    low = t.lower().replace("ё", "е")
    if low.startswith(("разработал", "согласовано", "согласовал")):
        return True
    if low in ("визы", "виза") or low.startswith("визы"):
        return True
    if "ознакомлен" in low and "инструкц" in low:
        return True
    if _is_note_start(t):
        return True
    # глава: «2 ОБЯЗАННОСТИ»
    if re.match(r"^\d+\s+[А-ЯЁA-ZА-яё]", t) and not re.match(r"^\d+\.\d+", t):
        return True
    # пункт тела с многоуровневым номером: «1.2. …» / «2.3.4. …»
    if re.match(r"^\d+\.\d+(\.\d+)*\.\s", t):
        return True
    return False


def _find_razrabotal_index(paras: list) -> int | None:
    for i, p in enumerate(paras):
        low = (p.text or "").strip().lower().replace("ё", "е")
        if low.startswith("разработал") or low.startswith("разработчик"):
            return i
    return None


def _collect_note_block_indices(paras: list, start_i: int, limit: int) -> list[int]:
    idxs = [start_i]
    i = start_i + 1
    while i < limit:
        t = (paras[i].text or "").strip()
        if not t:
            j = i + 1
            while j < limit and not (paras[j].text or "").strip():
                j += 1
            if j >= limit or _is_note_structure_break(paras[j].text):
                break
            idxs.append(i)
            i += 1
            continue
        if _is_note_structure_break(t):
            break
        idxs.append(i)
        i += 1
    return idxs


def move_notes_before_razrabotal(doc: Document) -> int:
    """
    Обязательно: все примечания перенести в конец документа,
    сразу перед абзацем «Разработал:» (до СОГЛАСОВАНО / виз / ознакомления).

    Возвращает число перенесённых абзацев (0 — уже на месте или примечаний нет).
    """
    paras = list(doc.paragraphs)
    raz_idx = _find_razrabotal_index(paras)
    if raz_idx is None:
        return 0

    blocks: list[list[int]] = []
    i = 0
    while i < raz_idx:
        if _is_note_start(paras[i].text):
            block = _collect_note_block_indices(paras, i, raz_idx)
            if block:
                blocks.append(block)
                i = block[-1] + 1
                continue
        i += 1

    if not blocks:
        return 0

    # Уже все примечания подряд непосредственно перед «Разработал»?
    note_idxs = {idx for b in blocks for idx in b}
    content_before = [
        k for k in range(raz_idx) if (paras[k].text or "").strip()
    ]
    # хвост перед Разработал: только примечания (идут подряд с конца)
    tail_notes: list[int] = []
    for k in reversed(content_before):
        if k in note_idxs:
            tail_notes.append(k)
            continue
        break
    tail_notes.reverse()
    all_note_content = [idx for idx in sorted(note_idxs) if (paras[idx].text or "").strip()]
    if all_note_content and all_note_content == [
        idx for idx in tail_notes if (paras[idx].text or "").strip()
    ]:
        # все примечания уже в хвосте перед Разработал
        return 0

    elements = []
    for block in blocks:
        for idx in block:
            elements.append(paras[idx]._element)

    if not elements:
        return 0

    raz_el = paras[raz_idx]._element
    # Перенос перед «Разработал» с сохранением порядка блоков.
    for el in elements:
        raz_el.addprevious(el)

    return len(elements)


def _find_signatory_block_bounds(paras: list) -> tuple[int | None, int]:
    """Индексы блока подписантов: start … end (перед «ознакомлен»)."""
    start = None
    end = len(paras)
    for i in range(len(paras) - 1, -1, -1):
        low = paras[i].text.strip().lower()
        if "ознакомлен" in low and "инструкц" in low:
            end = i
            continue
        if low.startswith("разработал") or low.startswith("согласовано") or low.startswith("согласовал"):
            start = i
            # вверх: захватить «Разработал:» и строки должность+ФИО над «Согласовано»
            for j in range(i - 1, max(-1, i - 20), -1):
                tj = (paras[j].text or "").strip()
                lj = tj.lower()
                if not tj:
                    start = j
                    continue
                if lj.startswith("разработал") or lj.startswith("разработчик"):
                    start = j
                    continue
                if lj.startswith("визы") or lj in ("виза", "визы"):
                    start = j
                    continue
                if _is_note_start(tj):
                    break
                if re.match(r"^\d+\.\d+(\.\d+)*\.\s", tj):
                    break
                if re.match(r"^\d+\s+[А-ЯЁA-Z]", tj) and not re.match(r"^\d+\.\d+", tj):
                    break
                if _is_signatory_job_line(tj) or FIO_RE.search(tj):
                    start = j
                    continue
                if DATE_RE.search(tj) or (tj.lstrip("\t ").startswith("«") and "_" in tj):
                    start = j
                    continue
                break
            break
        if low in ("визы", "виза") or low.startswith("визы"):
            start = i
            # выше «ВИЗЫ» часто ещё «Начальник РТС…» (утвердил/разработал)
            for j in range(i - 1, max(-1, i - 10), -1):
                tj = paras[j].text.strip()
                if not tj:
                    start = j
                    continue
                if _is_signatory_job_line(tj):
                    start = j
                    continue
                break
            break
    # если нет «Разработал»/«СОГЛАСОВАНО»/«ВИЗЫ» — искать хвост с должностями+ФИО
    if start is None:
        for i in range(len(paras) - 1, max(-1, len(paras) - 25), -1):
            t = paras[i].text.strip()
            low = t.lower()
            if "ознакомлен" in low:
                end = min(end, i)
                continue
            if FIO_RE.search(t) and any(
                k in low
                for k in (
                    "начальник",
                    "ведущий",
                    "заместитель",
                    "инженер",
                    "мастер",
                )
            ):
                start = i
                for j in range(i - 1, max(-1, i - 15), -1):
                    tj = paras[j].text.strip()
                    if not tj:
                        start = j
                        continue
                    if _is_signatory_job_line(tj):
                        start = j
                        continue
                    break
                break
    return start, end


def _find_last_body_before_signs(paras: list, sign_start: int | None) -> int | None:
    """Последний непустой абзац текста перед блоком подписантов (примечания пропускаем)."""
    if sign_start is None or sign_start <= 0:
        return None
    # индексы абзацев, входящих в примечания перед подписями
    note_idxs: set[int] = set()
    i = 0
    while i < sign_start:
        if _is_note_start(paras[i].text):
            block = _collect_note_block_indices(paras, i, sign_start)
            note_idxs.update(block)
            i = block[-1] + 1
            continue
        i += 1

    for j in range(sign_start - 1, -1, -1):
        t = paras[j].text.strip()
        if not t:
            continue
        if j in note_idxs:
            continue
        if t.upper() in ("ОТВЕТСТВЕННОСТЬ",) or re.match(r"^\d+\s+[А-ЯЁA-Z]", t):
            return None
        return j
    return None


def _word_page_of_paragraph_text(wdoc, text: str) -> int | None:
    """Номер страницы абзаца с данным текстом (через Word COM)."""
    needle = re.sub(r"\s+", " ", (text or "").strip())
    if not needle:
        return None
    # Word Find чувствителен к длине — берём устойчивый фрагмент
    probe = needle[:120]
    try:
        count = int(wdoc.Paragraphs.Count)
    except Exception:
        return None
    for i in range(1, count + 1):
        try:
            raw = wdoc.Paragraphs(i).Range.Text or ""
            t = re.sub(r"\s+", " ", raw.replace("\r", "").replace("\x07", "").strip())
            if not t:
                continue
            if t == needle or t.startswith(probe) or probe.startswith(t[:80]):
                return int(wdoc.Paragraphs(i).Range.Information(3))  # wdActiveEndPageNumber
        except Exception:
            continue
    # запасной путь — Find
    try:
        rng = wdoc.Content
        rng.Find.ClearFormatting()
        ok = rng.Find.Execute(FindText=probe[:60], Forward=True, Wrap=0)
        if ok:
            return int(rng.Information(3))
    except Exception:
        pass
    return None


def move_last_body_if_signatories_orphaned(docx_path: str) -> dict:
    """
    Если подписанты оказались на отдельном листе без последнего абзаца текста —
    перенести последний абзац на их лист (разрыв страницы перед ним).

    Как в правке «_оформлен+»: п. 5.1.5 уезжает вместе с «Разработал»/«СОГЛАСОВАНО».
    """
    result = {
        "checked": False,
        "orphaned": False,
        "moved": False,
        "last_body_page": None,
        "sign_page": None,
        "detail": "",
    }
    doc = Document(docx_path)
    paras = list(doc.paragraphs)
    sign_start, _end = _find_signatory_block_bounds(paras)
    last_body = _find_last_body_before_signs(paras, sign_start)
    if sign_start is None or last_body is None:
        result["detail"] = "no signatory block"
        return result

    last_text = paras[last_body].text.strip()
    sign_text = paras[sign_start].text.strip()
    already = _paragraph_has_page_break_before(paras[last_body])

    orphaned = False
    try:
        from .word_com import open_docx_readonly

        with open_docx_readonly(docx_path) as wdoc:
            p_last = _word_page_of_paragraph_text(wdoc, last_text)
            p_sign = _word_page_of_paragraph_text(wdoc, sign_text)
            result["checked"] = True
            result["last_body_page"] = p_last
            result["sign_page"] = p_sign
            if p_last is not None and p_sign is not None:
                orphaned = int(p_sign) > int(p_last)
                result["detail"] = f"pages last={p_last} sign={p_sign}"
            else:
                # страницы не нашли — если разрыва ещё нет, переносим профилактически
                orphaned = not already
                result["detail"] = f"pages unknown last={p_last} sign={p_sign}"
    except Exception as e:
        orphaned = not already
        result["detail"] = f"com unavailable: {e}"

    result["orphaned"] = orphaned
    if not orphaned:
        return result

    # Перенос: разрыв перед последним абзацем тела; убрать разрывы между ним и «Разработал»
    _set_page_break_before(paras[last_body], True)
    for i in range(last_body + 1, sign_start + 1):
        if _paragraph_has_page_break_before(paras[i]):
            _set_page_break_before(paras[i], False)
    keep_last_body_with_signatories(doc)
    save_docx_unprotected(doc, docx_path)
    result["moved"] = True
    if already:
        result["detail"] += "; refreshed page break / cleared breaks before signs"
    else:
        result["detail"] += "; moved last paragraph to signatories page"
    return result


def ensure_razrabotal_heading(doc: Document) -> int:
    """
    Всегда должна быть строка «Разработал:» перед подписантом-разработчиком.
    Если есть блок подписей/ФИО, а заголовка нет — вставить.
    """
    paras = list(doc.paragraphs)
    start, end = _find_signatory_block_bounds(paras)
    if start is None:
        return 0
    for i in range(start, min(end, len(paras))):
        low = (paras[i].text or "").strip().lower()
        if low.startswith("разработал") or low.startswith("разработчик"):
            return 0

    # точка вставки: перед первой должностью/ФИО (разработчик), иначе перед «Согласовано»
    insert_at = None
    for i in range(start, min(end, len(paras))):
        t = (paras[i].text or "").strip()
        low = t.lower()
        if not t:
            continue
        if low.startswith("согласован"):
            if insert_at is None:
                insert_at = i
            break
        if FIO_RE.search(t) or _is_signatory_job_line(t):
            insert_at = i
            break
    if insert_at is None:
        insert_at = start

    # пропустить ведущие пустые
    while insert_at < end and not (paras[insert_at].text or "").strip():
        insert_at += 1
        if insert_at >= len(paras):
            return 0

    new_p = paras[insert_at].insert_paragraph_before("Разработал:")
    _apply_signatory_line_format(new_p, ri_style=_is_ri_signatory_style("", doc))
    return 1


def finalize_notes_and_signatories(docx_path: str, doc_type: str = "") -> dict:
    """
    Финал оформления инструкции (после всех правок):
    1) примечания → перед «Разработал:»;
    2) всегда строка «Разработал:»;
    3) подписанты по эталону;
    4) отступ 1,25 см на текст (кроме титула/заголовков/подписантов);
    5) проверка нумерации пунктов.
    """
    report = {
        "notes_moved": 0,
        "razrabotal_added": 0,
        "signatories": 0,
        "indents_fixed": 0,
        "numbers_fixed": 0,
        "doc_type": doc_type,
    }
    doc = Document(docx_path)
    dtype = resolve_instruction_doc_type(
        doc_type=doc_type, doc=doc, source_path=docx_path
    )
    report["doc_type"] = dtype

    # сначала «Разработал:» (якорь), затем примечания перед ним
    try:
        report["razrabotal_added"] = ensure_razrabotal_heading(doc)
    except Exception as e:
        report["razrabotal_error"] = str(e)

    try:
        report["notes_moved"] = move_notes_before_razrabotal(doc)
    except Exception as e:
        report["notes_error"] = str(e)

    # пустая строка перед «Разработал:» (как в эталоне СЛЕСАРЬ 30.07.2026)
    try:
        report["empty_before_razrabotal"] = _ensure_empty_before_razrabotal(doc)
    except Exception as e:
        report["empty_before_razrabotal_error"] = str(e)

    # ещё раз названия инструкций в перечне (без ИОТ/номеров)
    try:
        from formatters.instruction_catalog import (
            apply_instruction_catalog_to_text,
            detect_unit_hint,
            load_catalog,
        )

        cat = load_catalog(
            unit_hint=detect_unit_hint(docx_path),
            force_rescan=False,
        )
        n_names = 0
        for p in doc.paragraphs:
            t = p.text or ""
            if not t.strip():
                continue
            low = t.strip().lower()
            if not (
                "иот" in low
                or "инструкция по охране труда" in low
                or re.search(r"(?i)\bиот\s*\d+", t)
            ):
                continue
            new_t, notes = apply_instruction_catalog_to_text(t, cat)
            if new_t != t:
                _set_runs(p, new_t)
                n_names += 1
        report["instruction_names_fixed"] = n_names
    except Exception as e:
        report["instruction_names_error"] = str(e)

    try:
        report["signatories"] = apply_signatory_block_format(doc, doc_type=dtype)
    except Exception as e:
        report["signatories_error"] = str(e)

    # после правок подписантов — примечания снова строго перед «Разработал:»
    try:
        report["notes_moved"] = int(report["notes_moved"]) + move_notes_before_razrabotal(
            doc
        )
    except Exception:
        pass

    try:
        report["empty_before_razrabotal"] = int(
            report.get("empty_before_razrabotal") or 0
        ) + _ensure_empty_before_razrabotal(doc)
    except Exception:
        pass

    # Стили эталона КЛ 31.07.2026 (Normal / Heading 1 / Body Text)
    try:
        report["etalon_styles"] = ensure_etalon_styles(doc)
    except Exception as e:
        report["etalon_styles_error"] = str(e)

    # УНИВЕРСАЛЬНО для ВСЕХ документов (п.1–6 правил Дубовика)
    try:
        report["double_spaces_fixed"] = collapse_double_spaces_in_document(doc)
    except Exception as e:
        report["double_spaces_error"] = str(e)

    try:
        report["indents_fixed"] = apply_body_first_indents(doc)
    except Exception as e:
        report["indents_error"] = str(e)

    # пустые строки и интервалы before/after (в т.ч. таблицы = 0)
    try:
        sp = apply_spacing_and_clear_body_empties(doc)
        report["spacing_cleared"] = sp.get("spacing", 0)
        report["body_empties_removed"] = sp.get("empties", 0)
        report["chapter_spacers"] = sp.get("chapter_spacers", 0)
    except Exception as e:
        report["spacing_error"] = str(e)

    # После содержания — разрыв; первая глава с верха страницы
    try:
        report["contents_page"] = separate_contents_onto_own_page(doc)
    except Exception as e:
        report["contents_page_error"] = str(e)

    # восстановить номера глав без цифры («ОБЪЕМ…» + 4.1 → «4 ОБЪЕМ…»)
    try:
        report["chapters_repaired"] = repair_missing_chapter_numbers(doc)
    except Exception as e:
        report["chapters_repaired_error"] = str(e)

    # заголовки глав: центр + пустая перед (кроме главы после содержания)
    try:
        ch = format_chapter_headings(doc)
        report["chapters_centered"] = ch.get("centered", 0)
        report["chapters_styled"] = ch.get("styled", 0)
        report["chapter_spacers"] = int(report.get("chapter_spacers") or 0) + int(
            ch.get("spacers") or 0
        )
    except Exception as e:
        report["chapters_error"] = str(e)

    # шрифт в таблицах: титул 14, тело/приложения 12
    try:
        report["table_fonts"] = apply_table_fonts(doc)
    except Exception as e:
        report["table_fonts_error"] = str(e)

    # отступы 1,25 см ещё раз (кроме заголовков)
    try:
        report["indents_fixed"] = int(report.get("indents_fixed") or 0) + apply_body_first_indents(
            doc
        )
    except Exception as e:
        report["indents_error"] = str(e)

    # подписанты (в т.ч. перенос длинных должностей)
    try:
        report["signatories"] = int(report.get("signatories") or 0) + apply_signatory_block_format(
            doc, doc_type=dtype
        )
    except Exception as e:
        report["signatories_error"] = str(e)

    # после подписантов снова убрать двойные пробелы (если появились)
    try:
        report["double_spaces_fixed"] = int(report.get("double_spaces_fixed") or 0) + (
            collapse_double_spaces_in_document(doc)
        )
    except Exception:
        pass

    # после зачистки пустых — снова пустая перед Разработал
    try:
        report["empty_before_razrabotal"] = int(
            report.get("empty_before_razrabotal") or 0
        ) + _ensure_empty_before_razrabotal(doc)
    except Exception:
        pass

    try:
        name_l = Path(docx_path).name.lower() if docx_path else ""
        skip_renumber = dtype == "dolzhnostnaya_instrukciya" or name_l.startswith(
            "ди "
        ) or name_l.startswith("ди_")
        if skip_renumber:
            report["numbers_fixed"] = 0
            report["numbering_check"] = {"skipped": "source_preserved"}
        else:
            chk = verify_and_fix_numbering(doc)
            report["numbers_fixed"] = int(chk.get("fixed", 0))
            chk2 = verify_and_fix_numbering(doc)
            report["numbers_fixed"] += int(chk2.get("fixed", 0))
            report["numbering_check"] = chk2
    except Exception as e:
        report["numbering_error"] = str(e)

    # подчистить порчу терминов/подписантов после всех проходов
    try:
        from formatters.russian_check import apply_local_russian_fixes

        n_loc = 0
        for p in list(iter_all_paragraphs(doc)):
            t = p.text or ""
            if not t.strip():
                continue
            new_t, _det = apply_local_russian_fixes(t)
            if new_t != t:
                _set_runs(p, new_t)
                n_loc += 1
        report["local_corruption_fixes"] = n_loc
    except Exception as e:
        report["local_corruption_fixes_error"] = str(e)

    # ВСЕГДА: убрать маркеры списков в тексте; подписи под строками — 11–12 пт
    try:
        report["markers_removed"] = strip_text_markers_everywhere(doc)
    except Exception as e:
        report["markers_error"] = str(e)
    try:
        report["caption_fonts"] = apply_form_caption_fonts(doc, size_pt=11)
    except Exception as e:
        report["caption_fonts_error"] = str(e)

    # ВСЕГДА: кликабельное Содержание (TOC), титул, разрывы изменений/приложений
    # (лист согласования — без разрыва страницы)
    try:
        from formatters.toc_and_structure import apply_document_structure_rules

        # сохранить текущий doc перед повторным открытием в structure
        save_docx_unprotected(doc, docx_path)
        struct = apply_document_structure_rules(docx_path, update_toc=True)
        report["structure"] = struct
        # перечитать после structure (уже сохранён)
        return report
    except Exception as e:
        report["structure_error"] = str(e)

    save_docx_unprotected(doc, docx_path)
    return report


def keep_last_body_with_signatories(doc: Document) -> int:
    """
    Последний абзац текста + весь блок подписантов держать вместе (не рвать по страницам).
    Как в правке «_оформлен+»: последний пункт идёт вместе с «Разработал»/«СОГЛАСОВАНО».
    """
    paras = list(doc.paragraphs)
    start, end = _find_signatory_block_bounds(paras)
    if start is None or start <= 0:
        return 0

    last_body = _find_last_body_before_signs(paras, start)
    if last_body is None:
        last_body = start

    changed = 0
    # от последнего пункта тела до предпоследней строки блока подписей — keepNext
    last_keep = end - 1
    while last_keep > start and not paras[last_keep].text.strip():
        last_keep -= 1
    for i in range(last_body, last_keep):
        _set_keep_with_next(paras[i], True)
        changed += 1
    # последняя строка блока — без keepNext (лист ознакомления может быть на следующей)
    if last_keep >= last_body:
        _set_keep_with_next(paras[last_keep], False)
    return changed


def normalize_signatory_line(text: str, *, ri_style: bool = False) -> str | None:
    """
    Привести строку «должность + ФИО» к канону.
    Заголовок согласования для ВСЕХ документов — «СОГЛАСОВАНО» (заглавными).
    None — если не подписант. Пустая строка — удалить (ВИЗЫ).
    """
    _ = ri_style
    t = text.replace("\xa0", " ").strip()
    if not t:
        return None
    up = t.upper().rstrip(":")
    if up in ("ВИЗЫ", "ВИЗА"):
        return ""  # удалить заголовок «ВИЗЫ»
    if up in ("СОГЛАСОВАНО", "СОГЛАСОВАЛ") or up.startswith("СОГЛАСОВАНО"):
        return "СОГЛАСОВАНО"
    if up.startswith("РАЗРАБОТАЛ"):
        return "Разработал:"

    # все ФИО в строке: должность без дублей, имя = последнее ФИО
    fios = list(FIO_RE.finditer(t))
    if fios:
        name_raw = fios[-1].group(0)
        title = t
        for m in reversed(fios):
            title = title[: m.start()] + " " + title[m.end() :]
        title = re.sub(r"[\t]+", " ", title)
        title = re.sub(r"\s+", " ", title).strip(" \t-–—")
        if not title or len(title) > 90 or title[:1].isdigit():
            return None
        name = re.sub(
            r"([А-ЯЁ])\.\s*([А-ЯЁ])\.\s*",
            r"\1.\2.",
            name_raw.replace(" ", ""),
        )
        return f"{title}\t{name}"

    m = re.match(
        r"^(?P<title>.+?)(?:\s{2,}|\t+|\s+)(?P<name>" + FIO_RE.pattern + r")\s*$",
        t,
    )
    if m:
        title = re.sub(r"\s+", " ", m.group("title")).strip()
        if len(title) > 90 or title[:1].isdigit():
            return None
        name = re.sub(r"([А-ЯЁ])\.\s*([А-ЯЁ])\.\s*", r"\1.\2.", m.group("name").replace(" ", ""))
        # убрать дубль ФИО из должности: «… по качеству М.Б.Повод»
        title_nospace = title.replace(" ", "")
        if name.replace(" ", "") in title_nospace:
            title = re.sub(FIO_RE.pattern + r"\s*$", "", title).strip()
            title = re.sub(r"\s+", " ", title)
        if not title:
            return None
        return f"{title}\t{name}"
    return None


def _ensure_empty_before_soglasovano(paras: list, sog_idx: int) -> int:
    """Перед «СОГЛАСОВАНО» — одна пустая строка (для ВСЕХ документов)."""
    if sog_idx <= 0:
        return 0
    prev = paras[sog_idx - 1]
    if not (prev.text or "").strip():
        return 0
    paras[sog_idx].insert_paragraph_before("")
    return 1


def _ensure_empty_before_razrabotal(doc: Document) -> int:
    """Перед «Разработал:» — одна пустая строка (после примечания / последнего пункта)."""
    paras = list(doc.paragraphs)
    raz = _find_razrabotal_index(paras)
    if raz is None or raz <= 0:
        return 0
    prev = paras[raz - 1]
    if not (prev.text or "").strip():
        return 0
    paras[raz].insert_paragraph_before("")
    return 1


def apply_signatory_block_format(doc: Document, doc_type: str = "") -> int:
    """
    Подписанты — правила для ВСЕХ документов (РИ, ДИ, положение, ИОТ, приказ, другое):
    пустая перед «Разработал:»; должность + таб ≥12 см + ФИО;
    длинную должность перенести отдельными абзацами (одинарный интервал внутри
    должности; полуторный между подписантами); дата в одну строку на 12 см;
    пустая перед «СОГЛАСОВАНО» (заглавными); «ВИЗЫ» убрать.

    ФИО всегда без пробела: И.И.Иванов.
    """
    changed = 0
    _ = doc_type
    paras = list(doc.paragraphs)
    if not paras:
        return 0

    start, end = _find_signatory_block_bounds(paras)
    if start is None:
        return 0

    i = start
    while i < end:
        p = paras[i]
        raw = p.text
        t = raw.strip()
        if not t:
            i += 1
            continue

        fixed = normalize_signatory_line(t)
        # «ВИЗЫ» — удалить
        if fixed == "":
            _delete_paragraph(p)
            changed += 1
            i += 1
            continue
        if (
            t.upper().startswith("СОГЛАСОВАНО")
            or t.upper().startswith("РАЗРАБОТАЛ")
            or t.lower().startswith("согласовал")
            or t.lower().startswith("согласовано")
        ):
            new_t = fixed or t
            is_soglas = (
                new_t.upper().startswith("СОГЛАСОВАНО")
                or new_t.lower().startswith("согласовано")
            )
            if is_soglas:
                new_t = "СОГЛАСОВАНО"
                # пустая строка перед «СОГЛАСОВАНО» — для ВСЕХ документов
                inserted = _ensure_empty_before_soglasovano(paras, i)
                if inserted:
                    changed += inserted
                    paras = list(doc.paragraphs)
                    start, end = _find_signatory_block_bounds(paras)
                    if start is None:
                        break
                    for j in range(start, min(end, len(paras))):
                        if (paras[j].text or "").strip().lower().startswith("согласован"):
                            i = j
                            p = paras[j]
                            raw = p.text
                            break
            bold = True if is_soglas else None
            if new_t != raw or (
                bold and not any(r.bold for r in p.runs if r.text)
            ):
                _set_runs(p, new_t, bold=bold)
                changed += 1
            _apply_signatory_line_format(p)
            i += 1
            continue

        if fixed and FIO_RE.search(t):
            title, name = fixed.split("\t", 1)
            _write_signatory_title_name(p, title, name)
            changed += 1
            i += 1
            continue

        # Сначала даты — иначе ««__»______» + следующая должность склеиваются
        if DATE_RE.search(t) or (t.lstrip("\t ").startswith("«") and "_" in t):
            _write_signatory_date_aligned(p, t)
            changed += 1
            i += 1
            continue

        if i + 1 < end:
            t2 = paras[i + 1].text.strip()
            t3 = paras[i + 2].text.strip() if i + 2 < end else ""

            cont_then_name = (
                t2
                and t3
                and not FIO_RE.search(t2)
                and not DATE_RE.search(t2)
                and len(t2) < 60
                and (
                    FIO_RE.fullmatch(t3.replace(" ", "")) is not None
                    or (FIO_RE.search(t3) and len(t3) < 40 and not DATE_RE.search(t3))
                )
            )
            if cont_then_name and not FIO_RE.search(t):
                # должность на 2 абзацах + ФИО отдельно → сохранить разбивку:
                # части должности — одинарный; строка с ФИО — полуторный
                name = re.sub(
                    r"([А-ЯЁ])\.\s*([А-ЯЁ])\.\s*",
                    r"\1.\2.",
                    t3.replace(" ", ""),
                )
                _set_runs(p, re.sub(r"\s+", " ", t).strip())
                _apply_signatory_line_format(p, line_spacing=1.0)
                p2 = paras[i + 1]
                title2 = re.sub(r"\s+", " ", t2).strip()
                _apply_signatory_line_format(p2, line_spacing=SIGN_LINE_SPACING_RI)
                _set_signatory_tab_stop(p2, SIGN_FIO_TAB_CM)
                _set_runs(p2, f"{title2}\t{name}")
                _delete_paragraph(paras[i + 2])
                changed += 3
                i += 3
                continue

            # две строки: первая — только должность, вторая — должность+ФИО или ФИО
            title_then_fio_line = (
                t2
                and not FIO_RE.search(t)
                and not DATE_RE.search(t)
                and FIO_RE.search(t2)
                and not DATE_RE.search(t2)
            )
            if title_then_fio_line:
                _set_runs(p, re.sub(r"\s+", " ", t).strip())
                _apply_signatory_line_format(p, line_spacing=1.0)
                fixed2 = normalize_signatory_line(t2)
                if fixed2 and "\t" in fixed2:
                    title2, name2 = fixed2.split("\t", 1)
                    _apply_signatory_line_format(
                        paras[i + 1], line_spacing=SIGN_LINE_SPACING_RI
                    )
                    _set_signatory_tab_stop(paras[i + 1], SIGN_FIO_TAB_CM)
                    name2 = _format_signatory_fio(name2)
                    title2_clean = re.sub(r"\s+", " ", title2).strip()
                    _set_runs(
                        paras[i + 1],
                        f"{title2_clean}\t{name2}",
                    )
                else:
                    name2 = re.sub(
                        r"([А-ЯЁ])\.\s*([А-ЯЁ])\.\s*",
                        r"\1.\2.",
                        t2.replace(" ", ""),
                    )
                    # ФИО было на второй строке без должности — перенести на первую? нет:
                    # пользователь держит должность на 1-й, продолжение+ФИО на 2-й
                    _apply_signatory_line_format(
                        paras[i + 1], line_spacing=SIGN_LINE_SPACING_RI
                    )
                    _set_signatory_tab_stop(paras[i + 1], SIGN_FIO_TAB_CM)
                    _set_runs(paras[i + 1], f"\t{name2}")
                changed += 2
                i += 2
                continue

            name_only = bool(t2) and (
                FIO_RE.fullmatch(t2.replace(" ", "")) is not None
                or (FIO_RE.search(t2) is not None and len(t2) < 40 and not DATE_RE.search(t2))
            )
            if name_only and not FIO_RE.search(t):
                title = re.sub(r"\s+", " ", t).strip().rstrip("\t")
                name = re.sub(r"([А-ЯЁ])\.\s*([А-ЯЁ])\.\s*", r"\1.\2.", t2.replace(" ", ""))
                _write_signatory_title_name(p, title, name)
                _delete_paragraph(paras[i + 1])
                changed += 2
                i += 2
                continue

        # обломок: строка только с ФИО, если ФИО уже есть в предыдущей — удалить
        if FIO_RE.fullmatch(t.replace(" ", "")) and i > start:
            prev = paras[i - 1].text if paras[i - 1]._element.getparent() is not None else ""
            if t.replace(" ", "") in prev.replace(" ", ""):
                _delete_paragraph(p)
                changed += 1
                i += 1
                continue

        if is_signatory_like(t):
            _apply_signatory_line_format(p)
        i += 1

    changed += keep_last_body_with_signatories(doc)
    return changed
