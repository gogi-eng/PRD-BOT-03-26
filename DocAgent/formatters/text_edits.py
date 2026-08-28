# -*- coding: utf-8 -*-
"""
Текстовые правки по практике Дубовика В.В.:
сравнение документов в папке РАССМОТРЕНИЕ (исходник vs СНиОТ/Дубовик)
+ правила пользователя: без маркеров, отступ 1,25 см, чистка «должен знать».
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from .common import (
    apply_basic_office_format,
    is_signatory_like,
    iter_all_paragraphs,
    save_docx_unprotected,
    set_run_font,
)
from .normative_docs_policy import (
    CATEGORY_WORKER,
    category_label_ru,
    detect_staff_category,
    extract_profession_tokens,
    is_directives_header_line,
    is_external_approved_rules,
    is_npa_delete_phrase,
    is_npa_title_line,
    is_orphan_presidential_directive_line,
    should_remove_normative_line,
)
from .instruction_catalog import (
    apply_instruction_catalog_to_text,
    detect_unit_hint,
    load_catalog,
)
from .structure_fix import (
    convert_auto_numbers_to_text,
    normalize_document_structure,
    _numbering_fmt_map,
    clear_bullet_markers_only,
    collapse_duplicate_number_prefix,
    fix_broken_number_prefixes,
    replace_soft_line_breaks_with_paragraphs,
    replace_tabs_with_single_space,
    resolve_instruction_doc_type,
    safe_phrase_replace,
)

ROOT = Path(__file__).resolve().parents[1]
PATTERNS_PATH = ROOT / "learned_edit_patterns.json"

# Маркеры в начале строки — запрещены (с пробелом и без: «–устав»)
MARKER_RE = re.compile(
    r"^[\s]*(?:[\-\u2013\u2014\u2022\u00B7\uF0B7\uF0A7\u25CB\u25A0\u25A1\u25CF\u25E6\*◦▪▸►■○●]+)[\s\u00a0]*"
)

# Раздел «должен знать»
MUST_KNOW_START = re.compile(r"должен\s+знать\s*:?", re.IGNORECASE)
MUST_KNOW_END = re.compile(
    r"(имеет\s+право|обязанност|возлагаются\s+следующие\s+функции|"
    r"для\s+исполнения|ответственност|заключительн|"
    r"^\d+\.\d*\.?\s+[А-ЯЁ])",
    re.IGNORECASE,
)
# Что убирать из «должен знать» у РАБОЧИХ (не трогать «законодательство» без № закона)
LAW_LINE_RE = re.compile(
    r"(?i)\b(кодекс\w*|закон(?!одател)\w*|декрет\w*|указ\w*|директив\w*)\b",
)

# Ключевые слова реквизитов — без пробелов внутри (Инструкция п. про УТВЕРЖДАЮ/СОГЛАСОВАНО)
REQUISITE_WORDS = (
    "УТВЕРЖДАЮ",
    "УТВЕРЖДЕНО",
    "СОГЛАСОВАНО",
    "ПРИКАЗЫВАЮ",
    "ПОСТАНОВЛЯЕТ",
    "РЕШИЛ",
    "ОБЯЗЫВАЮ",
    "ПРЕДЛАГАЮ",
)


def _looks_like_title_two_column(text: str) -> bool:
    """Строка титула с двумя колонками (пробелы/табы) — нельзя сжимать пробелы."""
    t = text.replace("\xa0", " ")
    if "\t" in t and len(t.strip()) < 200:
        return True
    if re.search(r"\S {3,}\S", t) and len(t.strip()) < 200:
        low = t.lower()
        if any(
            k in low
            for k in (
                "инструкция",
                "утверждаю",
                "начальник",
                "заместитель",
                "диспетчер",
                "слесар",
                "транспорт",
                "служб",
            )
        ):
            return True
        if re.search(r"_{3,}", t) or re.search(r"[А-ЯЁ]\.[А-ЯЁ]\.", t):
            return True
    # строки подписантов «должность … И.О.Фамилия»
    if re.search(r"[А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁа-яё\-]+", t) and len(t) < 120:
        low = t.lower()
        if any(k in low for k in ("начальник", "ведущий", "специалист", "инженер", "заместитель")):
            return True
    return False


def load_patterns() -> dict:
    with open(PATTERNS_PATH, encoding="utf-8") as f:
        return json.load(f)


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Перезаписать абзац целиком (в т.ч. убрать текст внутри гиперссылок)."""
    from formatters.structure_fix import _set_runs

    bold = paragraph.runs[0].bold if paragraph.runs else None
    _set_runs(paragraph, text, bold=bold)


def strip_leading_marker(text: str) -> tuple[str, int]:
    """Убрать маркер в начале строки (-, •, –, —, *), даже без пробела после него."""
    if not text:
        return text, 0
    original = text
    changed = 0
    while True:
        new, n = MARKER_RE.subn("", text, count=1)
        if not n:
            break
        text = new.lstrip(" \t")
        changed += 1
        if changed > 5:
            break
    # один символ-маркер вплотную к букве
    m = re.match(r"^([\-\u2013\u2014\u2022\u00B7\uF0B7\uF0A7\u25CB\u25A0\u25CF\*◦▪▸►■○●])([^\d\s].*)$", text)
    if m:
        text = m.group(2).lstrip()
        changed += 1
    if changed:
        return text, 1
    return original, 0


def strip_normative_from_local_acts(text: str) -> tuple[str, int]:
    """
    Локальный акт НЕ является нормативным правовым актом.

    Основание — Закон РБ «О нормативных правовых актах»
    (файл: Desktop\\Нормативка\\Закон О НОРМАТИВНЫХ ПРАВОВЫХ АКТАХ на 10.2023.RTF):
      ст. 2 — НПА; ст. 4 — ненормативные правовые акты; ст. 5 — локальные акты
      (отдельная категория); ст. 1 — Закон на локальные акты не распространяется;
      ст. 5 п. 2 — локальные акты должны соответствовать НПА.

    Всегда убирать «нормативн…» / «нормативно-» рядом с «локальн…»:
      локальными нормативными правовыми актами → локальными правовыми актами
      локальными нормативно-правовыми актами → локальными правовыми актами
    """
    if not text or "локальн" not in text.lower():
        return text, 0
    original = text
    # «локальными нормативно-правовыми» → «локальными правовыми»
    text = re.sub(
        r"(?i)(локальн\w*\s+)нормативно-?(правов\w*)",
        r"\1\2",
        text,
    )
    # «локальными нормативными правовыми» → «локальными правовыми»
    text = re.sub(
        r"(?i)(локальн\w*\s+)нормативн\w*\s+(правов\w*)",
        r"\1\2",
        text,
    )
    # «локальный нормативный акт» (без «правов») → «локальный акт»
    text = re.sub(
        r"(?i)(локальн\w*\s+)нормативн\w*\s+(акт\w*)",
        r"\1\2",
        text,
    )
    if text != original:
        # подчистить двойные пробелы только в затронутых местах
        text = re.sub(r" {2,}", " ", text)
        return text, 1
    return original, 0


def normalize_spaces(text: str, *, protect_title_columns: bool = False) -> tuple[str, int]:
    """
    Убрать лишние пробелы (Инструкция по делопроизводству 2025 + типографика).
    protect_title_columns=True — не схлопывать длинные пробелы/табы
    (иначе ломается двухколоночный титул «ИНСТРУКЦИЯ … УТВЕРЖДАЮ»).
    """
    if not text:
        return text, 0
    original = text

    if protect_title_columns:
        # колонки титула / строки с табами подписантов — длинные пробелы и табы НЕ трогаем
        text = text.replace("\u00a0", " ").replace("\u202f", " ")
        return text, (0 if text == original else 1)

    # неразрывные / узкие пробелы и табы → один пробел
    text = text.replace("\u00a0", " ").replace("\u202f", " ").replace("\t", " ")
    text = text.replace("\u2009", " ").replace("\u200a", " ")

    # ключевые слова реквизитов с пробелами внутри: «У Т В Е Р Ж Д А Ю» маловероятно,
    # но «УТВЕРЖДА Ю» / лишние пробелы вокруг
    for word in REQUISITE_WORDS:
        # если слово уже есть без пробелов — ок; схлопнуть варианты с пробелами между буквами
        spaced = r"\s*".join(re.escape(ch) for ch in word)
        text = re.sub(spaced, word, text, flags=re.IGNORECASE)

    # пробел перед знаками препинания
    text = re.sub(r"\s+([,.;:!?»\)\]])", r"\1", text)
    # пробел после открывающих
    text = re.sub(r"([«\(\[])\s+", r"\1", text)

    # инициалы: А. А. / А .А . → А.А.
    text = re.sub(r"\b([А-ЯЁA-Z])\s*\.\s*([А-ЯЁA-Z])\s*\.\s*", r"\1.\2.", text)
    # инициалы + фамилия без пробела после второй точки: А.А. Иванов → А.А.Иванов
    text = re.sub(r"\b([А-ЯЁA-Z]\.[А-ЯЁA-Z]\.)\s+([А-ЯЁA-Z][а-яёa-z\-]+)", r"\1\2", text)

    # № с одним пробелом (или без лишних)
    text = re.sub(r"№\s+", "№ ", text)

    # несколько пробелов → один
    text = re.sub(r" {2,}", " ", text)

    # края строки (не трогаем намеренные отступы Word — только текстовые пробелы)
    text = text.strip()

    # вернуть пробел после № если съели всё (редко)
    changed = 1 if text != original.strip() or "  " in original or "\t" in original else 0
    if text != original:
        # более точный счётчик: было ли сжатие
        changed = 1
    return text, changed


def apply_tkp_codes(text: str, codes: dict) -> tuple[str, int]:
    """ТКП 054-2007 → ТКП 054-2007 (02300). если кода ещё нет."""
    n = 0

    def repl(m: re.Match) -> str:
        nonlocal n
        num = m.group(1)
        code = codes.get(num)
        if not code:
            return m.group(0)
        tail = m.string[m.end() : m.end() + 12]
        if re.match(r"\s*\(\d{4,5}\)", tail):
            return m.group(0)
        n += 1
        return f"{m.group(0)} ({code})"

    new = re.sub(r"ТКП\s+(\d{2,3}-\d{4})", repl, text, flags=re.IGNORECASE)
    return new, n


def strip_redundant_unit(text: str, units: list[str]) -> tuple[str, int]:
    n = 0
    for unit in units:
        pat = re.compile(
            rf"(Ведущий инженер|Инженер|Мастер|Начальник)\s+{re.escape(unit)}\s+"
            rf"(проходит|назначается|может|должен|имеет|исполняет|несет|несёт|руководствуется)",
            flags=re.IGNORECASE,
        )

        def repl(m: re.Match) -> str:
            nonlocal n
            n += 1
            return f"{m.group(1)} {m.group(2)}"

        text = pat.sub(repl, text)
    return text, n


def is_law_code_decree_line(text: str) -> bool:
    """Устаревшее имя: строка НПА (закон/кодекс/декрет/указ/директива)."""
    return is_npa_title_line(text)


def apply_text_edits(
    input_path: str,
    output_path: str,
    *,
    also_basic_format: bool = True,
    apply_name_updates: bool = True,
    apply_list_markers: bool = False,  # маркеры ЗАПРЕЩЕНЫ — не добавляем
    doc_type: str = "",
) -> dict:
    patterns = load_patterns()
    report = {
        "replacements": 0,
        "tkp_codes": 0,
        "deleted_paragraphs": 0,
        "deleted_phrases": 0,
        "markers_removed": 0,
        "spaces_normalized": 0,
        "must_know_laws_removed": 0,
        "unit_stripped": 0,
        "name_updates": 0,
        "instruction_refs_fixed": 0,
        "staff_category": "",
        "details": [],
        "output": output_path,
        "skipped_conservative": False,
    }

    # ДИ САТП «Старший мастер»: learned patterns и structure_fix ломают нумерацию
    from formatters.sniot_document import is_conservative_di_satp

    if is_conservative_di_satp(input_path, doc_type):
        from path_resolver import copy_file_if_different

        copy_file_if_different(input_path, output_path)
        if also_basic_format:
            apply_basic_office_format(output_path, output_path)
        report["skipped_conservative"] = True
        report["details"].append(
            "СНиОТ: text_edits пропущены — только оформление (слова исходника не менять); "
            "нумерацию и подписантов восстанавливает fix_sniot_document"
        )
        report["total_edits"] = 0
        return report

    work = output_path
    # 0) автонумерацию Word → обычный текст (иначе номера пропадут)
    from path_resolver import copy_file_if_different

    copy_file_if_different(input_path, work)
    try:
        n_mat = 0
        from .structure_fix import materialize_list_numbers

        n_mat = materialize_list_numbers(work)
        report["details"].append(f"автонумерация Word переведена в текст ({n_mat} абз.)")
    except Exception as e:
        report["details"].append(f"перевод нумерации: {e}")

    # Титул в таблицу ДО сжатия пробелов (иначе колонки «ИНСТРУКЦИЯ | УТВЕРЖДАЮ» склеятся)
    try:
        from .structure_fix import (
            ensure_title_table_like_sample,
            collapse_title_spacer_empties,
            ensure_title_sample_spacers,
        )

        doc_t = Document(work)
        blob0 = " ".join(p.text for p in doc_t.paragraphs[:20]).lower()
        dtype0 = doc_type or (
            "dolzhnostnaya_instrukciya" if "должностн" in blob0 else (
                "rabochaya_instrukciya" if "рабоч" in blob0 else "unsupported"
            )
        )
        tfix = ensure_title_table_like_sample(doc_t, dtype0)
        empt = collapse_title_spacer_empties(doc_t)
        ensure_title_sample_spacers(doc_t)
        save_docx_unprotected(doc_t, work)
        report["details"].append(
            f"титул как в эталоне СЛЕСАРЬ 30.07 (правок={tfix}, пустых убрано={empt})"
        )
    except Exception as e:
        report["details"].append(f"титул (ранний): {e}")

    if also_basic_format:
        apply_basic_office_format(work, work)
    # иначе уже скопировано

    doc = Document(work)
    try:
        n_soft = replace_soft_line_breaks_with_paragraphs(doc)
        if n_soft:
            report["details"].append(
                f"разрыв строки → перенос абзаца ({n_soft})"
            )
    except Exception as e:
        report["details"].append(f"замена разрывов строки: {e}")

    try:
        n_tab = replace_tabs_with_single_space(doc)
        if n_tab:
            report["details"].append(f"табуляция → пробел ({n_tab})")
    except Exception as e:
        report["details"].append(f"замена табуляции: {e}")

    delete_contains = [x.lower() for x in patterns.get("delete_paragraph_if_contains", [])]
    delete_equals = [x.strip().lower() for x in patterns.get("delete_paragraph_if_equals", [])]

    # --- категория персонала (человеческая логика по НПА) ---
    dtype_detect = resolve_instruction_doc_type(
        doc_type=doc_type, doc=doc, source_path=input_path
    )
    blob_head = " ".join(p.text for p in doc.paragraphs[:40])
    # титул РИ/ДИ часто только в таблице — добавить в blob
    for table in doc.tables[:4]:
        for row in table.rows:
            for cell in row.cells:
                blob_head += " " + (cell.text or "")
    report["details"].append(f"тип документа (для правок): {dtype_detect or 'не определён'}")
    staff_cat = detect_staff_category(
        source_path=input_path,
        doc_type=dtype_detect,
        text_blob=blob_head,
    )
    prof_tokens = extract_profession_tokens(input_path, blob_head)
    report["staff_category"] = staff_cat
    report["details"].append(f"категория персонала: {category_label_ru(staff_cat)}")
    strip_npa = staff_cat == CATEGORY_WORKER

    # каталог актуальных инструкций СНиОТ (номера и названия)
    unit_hint = detect_unit_hint(f"{input_path} {blob_head}")
    instr_catalog = load_catalog(unit_hint=unit_hint, force_rescan=True)
    report["details"].append(
        f"каталог инструкций (Перечень PDF + Приказ): "
        f"{instr_catalog.get('unique_numbers', 0)} номеров"
        f" [{instr_catalog.get('source', '?')}]"
        + (f" (подразделение {unit_hint})" if unit_hint else "")
    )

    in_must_know = False

    # Только body-абзацы документа (не ячейки таблиц титула) — для «должен знать»
    body_paras = list(doc.paragraphs)
    fmt_map = _numbering_fmt_map(doc)

    for p in list(iter_all_paragraphs(doc)):
        original = p.text
        text = original
        if clear_bullet_markers_only(p, fmt_map):
            report["markers_removed"] += 1

        if not text.strip():
            continue

        # починить сломанные номера вроде 2.2.4.3
        text2 = fix_broken_number_prefixes(text)
        if text2 != text:
            text = text2
            report["details"].append(f"нумерация: «{original[:40]}» → «{text[:40]}»")

        low = text.strip().lower()

        # зона «должен знать» / перечни документов
        if MUST_KNOW_START.search(text):
            in_must_know = True
        elif in_must_know and MUST_KNOW_END.search(text) and not MUST_KNOW_START.search(text):
            if not should_remove_normative_line(text, staff_cat, profession_tokens=prof_tokens):
                in_must_know = False

        # РИ: обрывки директив «от 27 декабря 2006 г. № 2 «О дебюрократизации…»» — УДАЛЯТЬ
        is_ri = dtype_detect == "rabochaya_instrukciya" or strip_npa
        if is_ri and (
            is_orphan_presidential_directive_line(text)
            or is_directives_header_line(text)
        ):
            _set_paragraph_text(p, "")
            report["deleted_paragraphs"] += 1
            report["must_know_laws_removed"] += 1
            report["details"].append(
                f"РИ: удалён абзац-директива/обрывок: {original[:90]}"
            )
            continue

        # НПА / лишние Правила — только для рабочих профессий
        # (в «должен знать» и ошибочно попавшие в тело, напр. «примеры работ»)
        if (
            strip_npa
            and not MUST_KNOW_START.search(text)
            and should_remove_normative_line(text, staff_cat, profession_tokens=prof_tokens)
            and (
                in_must_know
                or is_external_approved_rules(text)
                or is_npa_title_line(text)
            )
        ):
            _set_paragraph_text(p, "")
            report["must_know_laws_removed"] += 1
            report["deleted_paragraphs"] += 1
            report["details"].append(
                f"для рабочей профессии убран НПА/лишние Правила: {original[:90]}"
            )
            continue

        # заглушки (на титуле «номер инструкции» под чертой — НЕ удалять)
        if low in delete_equals or any(x in low for x in delete_contains):
            skip_title_stub = "номер инструкции" in low and len(original.strip()) < 40
            # НПА-абзацы из learned — только у рабочих
            is_npa_contain = any(
                is_npa_delete_phrase(x) and x in low for x in delete_contains
            )
            if skip_title_stub:
                pass
            elif is_npa_contain and not strip_npa:
                pass  # специалисту законы/кодексы оставляем
            else:
                _set_paragraph_text(p, "")
                report["deleted_paragraphs"] += 1
                report["details"].append(f"удалён абзац: {original[:80]}")
                continue

        # маркеры в тексте — убрать
        text2, n_mark = strip_leading_marker(text)
        if n_mark:
            text = text2
            report["markers_removed"] += 1
            report["details"].append("убран маркер в начале строки")

        # лишние пробелы (титульные колонки и подписанты — не сжимать табы!)
        protect_cols = _looks_like_title_two_column(text) or is_signatory_like(text)
        if "\t" in text and re.search(r"[А-ЯЁ]\.[А-ЯЁ]\.", text):
            protect_cols = True
        text2, n_sp = normalize_spaces(text, protect_title_columns=protect_cols)
        if n_sp or text2 != text:
            if text2 != text:
                report["spaces_normalized"] += 1
                report["details"].append("убраны лишние пробелы")
            text = text2

        for phr in patterns.get("delete_phrases", []):
            if phr not in text:
                continue
            # НПА-фразы не вычищать у специалистов/руководителей
            if is_npa_delete_phrase(phr) and not strip_npa:
                continue
            text = text.replace(phr, "")
            report["deleted_phrases"] += 1
            report["details"].append(f"убрана фраза: {phr[:60]}")

        for item in patterns.get("replace_phrases", []):
            old, new = item["old"], item["new"]
            if old not in text:
                continue
            try:
                from learner import is_safe_learned_replacement

                if not is_safe_learned_replacement(old, new):
                    continue
            except Exception:
                continue
            text2 = safe_phrase_replace(text, old, new)
            if text2 != text:
                text = text2
                report["replacements"] += 1
                report["details"].append(f"замена: «{old}» -> «{new}»")

        # после всех замен — схлопнуть двойные номера, если всплыли
        text2 = collapse_duplicate_number_prefix(text)
        if text2 != text:
            text = text2
            report["details"].append("убрана двойная нумерация пункта")

        for item in patterns.get("replace_phrases_if_missing_prb", []):
            old, new = item["old"], item["new"]
            idx = text.find(old)
            if idx < 0:
                continue
            window = text[max(0, idx - 40) : idx + len(old) + 80]
            if "промышленной безопасности" in window:
                continue
            text = text.replace(old, new)
            report["replacements"] += 1
            report["details"].append(f"замена (добавлен ПрБ): «{old}»")

        text, n_tkp = apply_tkp_codes(text, patterns.get("tkp_codes", {}))
        report["tkp_codes"] += n_tkp
        if n_tkp:
            report["details"].append(f"добавлены коды ТКП: {n_tkp}")

        # НЕ добавляем маркеры (apply_list_markers игнорируем / всегда False)
        _ = apply_list_markers

        text, n_unit = strip_redundant_unit(text, patterns.get("strip_redundant_unit_suffix", []))
        report["unit_stripped"] += n_unit

        if apply_name_updates:
            for old, new in patterns.get("name_updates", {}).items():
                if old in text:
                    text = text.replace(old, new)
                    report["name_updates"] += 1
                    report["details"].append(f"ФИО: {old} -> {new}")

        # локальный ≠ нормативный — всегда убрать «нормативн…» у локальных актов
        text2, n_loc = strip_normative_from_local_acts(text)
        if n_loc:
            text = text2
            report["replacements"] += 1
            report["details"].append(
                "локальный акт: убрано «нормативн…» → «локальные правовые акты»"
            )

        # названия и номера инструкций — только из папки СНиОТ\Инструкции
        if "инструкц" in text.lower() or "иот" in text.lower():
            text2, notes = apply_instruction_catalog_to_text(text, instr_catalog)
            if text2 != text:
                text = text2
                report["instruction_refs_fixed"] += 1
            for n in notes[:5]:
                report["details"].append(f"инструкция: {n}")

        if text != original:
            _set_paragraph_text(p, text)

    # повторный проход: «должен знать» в одной строке со списком через ;
    for p in body_paras:
        if not MUST_KNOW_START.search(p.text):
            continue
        if ":" not in p.text:
            continue
        if not strip_npa:
            continue
        head, _, tail = p.text.partition(":")
        parts = [x.strip() for x in tail.split(";") if x.strip()]
        kept = [
            x
            for x in parts
            if not should_remove_normative_line(x, staff_cat, profession_tokens=prof_tokens)
        ]
        removed = len(parts) - len(kept)
        if removed:
            new_tail = "; ".join(kept)
            if new_tail and not new_tail.endswith("."):
                new_tail += "."
            _set_paragraph_text(p, (head + ": " + new_tail).strip())
            report["must_know_laws_removed"] += removed
            report["details"].append(
                f"для рабочей профессии из «должен знать» убрано НПА/лишних Правил: {removed}"
            )

    save_docx_unprotected(doc, output_path)

    # единый титул, нумерация, подписанты
    try:
        # ВАЖНО: брать уже определённый тип (doc_type / таблица титула),
        # а не только первые 20 абзацев — иначе РИ → unsupported и подписанты ломаются
        dtype = resolve_instruction_doc_type(
            doc_type=dtype_detect or doc_type,
            doc=Document(output_path),
            source_path=input_path,
        )
        st = normalize_document_structure(output_path, dtype)
        report["structure"] = st
        report["details"].append(
            f"структура ({st.get('resolved_doc_type', dtype)}): "
            f"титул={st.get('title_fixed', 0)}, "
            f"нумерация={st.get('numbers_fixed', 0)}, "
            f"подписанты={st.get('signatories', 0)}, "
            f"разрыв→абзац={st.get('soft_breaks_to_paragraphs', 0)}, "
            f"таб→пробел={st.get('tabs_to_space', 0)}"
        )
        chk = st.get("numbering_check") or {}
        if chk:
            report["details"].append(
                f"проверка пунктов: проверено {chk.get('checked', 0)}, "
                f"ок {chk.get('ok', 0)}, исправлено {chk.get('fixed', 0)}"
            )
            for fx in (chk.get("fixes") or [])[:8]:
                report["details"].append(fx)
    except Exception as e:
        report["details"].append(f"структура (ошибка): {e}")

    uniq = []
    seen = set()
    for d in report["details"]:
        if d not in seen:
            seen.add(d)
            uniq.append(d)
    report["details"] = uniq[:40]
    report["total_edits"] = (
        report["replacements"]
        + report["tkp_codes"]
        + report["deleted_paragraphs"]
        + report["deleted_phrases"]
        + report["markers_removed"]
        + report["spaces_normalized"]
        + report["must_know_laws_removed"]
        + report["unit_stripped"]
        + report["name_updates"]
        + report["instruction_refs_fixed"]
    )
    return report
