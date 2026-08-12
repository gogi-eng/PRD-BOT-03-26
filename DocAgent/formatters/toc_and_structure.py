# -*- coding: utf-8 -*-
"""
Структура документа (ВСЕГДА):
— титул: ТОЛЬКО форма утверждения + название + «Минск» внизу по центру;
  предисловие и далее — со следующей страницы;
— кликабельное Содержание (поле TOC Word с \\h);
— заголовки разделов — стиль Heading 1 (для TOC);
— лист согласования = блок подписантов, БЕЗ разрыва страницы (сразу после текста);
— лист регистрации изменений — отдельная страница;
— приложения — с новой страницы.
"""

from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, Twips
from docx.text.paragraph import Paragraph

from formatters.common import (
    is_contents_heading,
    set_run_font,
    set_single_line_spacing,
    zero_para_spacing,
)
from formatters.structure_fix import (
    _set_page_break_before,
    _set_runs,
    apply_form_caption_fonts,
    apply_signatory_block_format,
    ensure_etalon_styles,
    strip_text_markers_everywhere,
)


def _log(msg: str) -> None:
    try:
        from agent_core import log

        log(msg)
    except Exception:
        pass


def _delete_paragraph(paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _insert_paragraph_after(paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        run = p.add_run(text)
        set_run_font(run, font_name="Times New Roman", font_size=Pt(14))
    return p


def _insert_paragraph_before(paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addprevious(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        run = p.add_run(text)
        set_run_font(run, font_name="Times New Roman", font_size=Pt(14))
    return p


def collapse_extra_empties(doc: Document, *, max_run: int = 1) -> int:
    """Убрать серии пустых абзацев (оставить не больше max_run подряд)."""
    removed = 0
    paras = list(doc.paragraphs)
    run = 0
    # с конца, чтобы не сбивать индексы при удалении
    for i in range(len(paras) - 1, -1, -1):
        t = (paras[i].text or "").strip()
        if not t:
            run += 1
            if run > max_run:
                _delete_paragraph(paras[i])
                removed += 1
        else:
            run = 0
    return removed


def _is_body_chapter_heading(text: str) -> bool:
    t = (text or "").strip()
    if not t or len(t) > 200:
        return False
    if t in ("Введение", "Приложение А", "Приложение Б"):
        return True
    if t.startswith("Приложение А") and len(t) < 25:
        return True
    if t.startswith("Приложение Б") and len(t) < 25:
        return True
    if re.match(r"^\d+\.\d+", t) or re.match(r"^\d+\.\s", t):
        return False
    if re.match(r"^\d+\s+\S", t):
        return True
    return False


def find_structure_anchors(doc: Document) -> dict:
    """Индексы ключевых мест документа."""
    paras = list(doc.paragraphs)
    toc_idx = None
    body_intro = None
    soglas_heading = None
    changes_heading = None
    app_a = None
    app_b = None
    intro_hits: list[int] = []

    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if is_contents_heading(t) and toc_idx is None:
            toc_idx = i
        if t == "Введение":
            intro_hits.append(i)
        if t.startswith("11 ") and "согласован" in t.lower():
            soglas_heading = i
        if t.startswith("12 ") and "изменен" in t.lower():
            changes_heading = i
        if t in ("Приложение А",) or (
            t.startswith("Приложение А") and len(t) < 25 and "Форма" not in t
        ):
            if app_a is None or i > (body_intro or 0):
                app_a = i
        if t in ("Приложение Б",) or (
            t.startswith("Приложение Б") and len(t) < 25 and "Форма" not in t
        ):
            if app_b is None or i > (body_intro or 0):
                app_b = i

    if intro_hits:
        # после содержания — последнее «Введение» до приложений
        body_intro = intro_hits[-1]
        if toc_idx is not None:
            after = [i for i in intro_hits if i > toc_idx]
            if after:
                body_intro = after[0] if len(after) == 1 else after[-1]
                # если два: первое — строка TOC, второе — тело
                if len(after) >= 2:
                    body_intro = after[-1]

    return {
        "toc": toc_idx,
        "body_intro": body_intro,
        "soglas": soglas_heading,
        "changes": changes_heading,
        "app_a": app_a,
        "app_b": app_b,
    }


def apply_heading1_to_body_chapters(doc: Document, body_start: int | None) -> int:
    """Стиль Heading 1 + оформление для заголовков тела (нужно для TOC Word)."""
    ensure_etalon_styles(doc)
    n = 0
    paras = list(doc.paragraphs)
    start = body_start if body_start is not None else 0
    for i, p in enumerate(paras):
        if i < start:
            continue
        t = (p.text or "").strip()
        if not _is_body_chapter_heading(t):
            continue
        try:
            p.style = doc.styles["Heading 1"]
        except Exception:
            pass
        zero_para_spacing(p)
        set_single_line_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        # outline level 0
        pPr = p._element.get_or_add_pPr()
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is None:
            ol = OxmlElement("w:outlineLvl")
            pPr.append(ol)
        ol.set(qn("w:val"), "0")
        # шрифт
        for r in p.runs:
            set_run_font(r, font_name="Times New Roman", font_size=Pt(14), bold=True)
        if not p.runs and t:
            _set_runs(p, t, bold=True)
            for r in p.runs:
                set_run_font(r, font_name="Times New Roman", font_size=Pt(14), bold=True)
        n += 1
    return n


def insert_word_toc_field(paragraph) -> None:
    """Вставить поле TOC с гиперссылками (\\h) — встроенное оглавление Word."""
    # очистить абзац
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)

    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    # \o "1-1" — уровни; \h — гиперссылки; \z — скрыть в веб-виде номера не нужны;
    # \u — использовать уровни структуры
    instr.text = ' TOC \\o "1-1" \\h \\z \\u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    # заглушка до обновления полей
    stub = OxmlElement("w:t")
    stub.text = (
        "«Содержание» — щёлкните правой кнопкой → «Обновить поле», "
        "либо откройте документ (поля обновятся автоматически)."
    )

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r2 = OxmlElement("w:r")
    r2.append(stub)
    paragraph._element.append(r2)
    r_end = OxmlElement("w:r")
    r_end.append(fld_end)
    paragraph._element.append(r_end)


def replace_manual_toc_with_field(doc: Document, anchors: dict) -> int:
    """Удалить ручные строки содержания, вставить поле TOC Word."""
    toc_idx = anchors.get("toc")
    body_intro = anchors.get("body_intro")
    if toc_idx is None or body_intro is None or body_intro <= toc_idx + 1:
        return 0

    paras = list(doc.paragraphs)
    # удалить строки между Содержание и телом (не трогая заголовок и Введение тела)
    deleted = 0
    for i in range(body_intro - 1, toc_idx, -1):
        _delete_paragraph(paras[i])
        deleted += 1

    # после удаления — снова найти Содержание и вставить поле сразу после него
    paras = list(doc.paragraphs)
    toc_idx = next(
        (i for i, p in enumerate(paras) if is_contents_heading((p.text or "").strip())),
        None,
    )
    if toc_idx is None:
        return deleted

    toc_para = paras[toc_idx]
    # оформить заголовок «Содержание»
    zero_para_spacing(toc_para)
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_para.paragraph_format.first_line_indent = Cm(0)
    for r in toc_para.runs:
        set_run_font(r, font_name="Times New Roman", font_size=Pt(14), bold=True)

    field_p = _insert_paragraph_after(toc_para, "")
    insert_word_toc_field(field_p)
    return deleted


def apply_structural_page_breaks(doc: Document, anchors: dict) -> list[str]:
    """Разрывы: титул|Содержание|тело|изменения|приложения.

    Лист согласования — БЕЗ разрыва (идёт сразу после текста).
    """
    done: list[str] = []
    paras = list(doc.paragraphs)

    def break_at(idx: int | None, label: str) -> None:
        if idx is None or idx < 0 or idx >= len(paras):
            return
        _set_page_break_before(paras[idx], True)
        done.append(label)

    # перечитать якоря после правок TOC
    anchors = find_structure_anchors(doc)
    paras = list(doc.paragraphs)

    break_at(anchors.get("toc"), "before_toc")
    break_at(anchors.get("body_intro"), "before_body")
    # согласование НЕ отделять — иначе непонятно, что согласуют
    sog = anchors.get("soglas")
    if sog is not None and 0 <= sog < len(paras):
        _set_page_break_before(paras[sog], False)
        done.append("no_break_before_soglasovanie")
    break_at(anchors.get("changes"), "before_changes_sheet")
    # Приложение А — альбомная секция (см. set_appendix_a_landscape);
    # pageBreakBefore на А убираем — разрыв даёт секция nextPage.
    app_a = anchors.get("app_a")
    if app_a is not None and 0 <= app_a < len(paras):
        _set_page_break_before(paras[app_a], False)
        done.append("appendix_a_no_pagebreak_use_section")
    break_at(anchors.get("app_b"), "before_appendix_b")

    # убрать пустые перед главой тела (с верха листа)
    anchors = find_structure_anchors(doc)
    paras = list(doc.paragraphs)
    bi = anchors.get("body_intro")
    if bi is not None:
        while bi > 0 and not (paras[bi - 1].text or "").strip():
            # не удалять если это единственный разделитель после TOC field —
            # но правило: глава с верха → удалить пустые перед ней
            prev = paras[bi - 1]
            # если у предыдущего pageBreakBefore — это другой блок
            _delete_paragraph(prev)
            paras = list(doc.paragraphs)
            anchors = find_structure_anchors(doc)
            bi = anchors.get("body_intro")
            if bi is None:
                break
            paras = list(doc.paragraphs)
        done.append("cleared_empties_before_body")

    return done


# A4 в twips (Word)
_A4_PORTRAIT_W = "11906"
_A4_PORTRAIT_H = "16838"
_A4_LANDSCAPE_W = "16838"
_A4_LANDSCAPE_H = "11906"


def _clone_pg_mar_from_section(doc: Document) -> OxmlElement:
    """Скопировать поля страницы из первой секции документа."""
    src = doc.sections[0]._sectPr.find(qn("w:pgMar"))
    pgMar = OxmlElement("w:pgMar")
    if src is not None:
        for k, v in src.attrib.items():
            pgMar.set(k, v)
    else:
        # 20/20/30/15 мм ≈ 1134/1134/1701/850 twips
        pgMar.set(qn("w:top"), "1134")
        pgMar.set(qn("w:bottom"), "1134")
        pgMar.set(qn("w:left"), "1701")
        pgMar.set(qn("w:right"), "850")
        pgMar.set(qn("w:header"), "720")
        pgMar.set(qn("w:footer"), "720")
        pgMar.set(qn("w:gutter"), "0")
    return pgMar


def _build_sect_pr(*, landscape: bool, next_page: bool, doc: Document) -> OxmlElement:
    """Собрать w:sectPr с ориентацией и полями."""
    sectPr = OxmlElement("w:sectPr")
    if next_page:
        typ = OxmlElement("w:type")
        typ.set(qn("w:val"), "nextPage")
        sectPr.append(typ)
    pgSz = OxmlElement("w:pgSz")
    if landscape:
        pgSz.set(qn("w:w"), _A4_LANDSCAPE_W)
        pgSz.set(qn("w:h"), _A4_LANDSCAPE_H)
        pgSz.set(qn("w:orient"), "landscape")
    else:
        pgSz.set(qn("w:w"), _A4_PORTRAIT_W)
        pgSz.set(qn("w:h"), _A4_PORTRAIT_H)
    sectPr.append(pgSz)
    sectPr.append(_clone_pg_mar_from_section(doc))
    return sectPr


def _set_paragraph_sect_pr(paragraph, sectPr: OxmlElement) -> None:
    """Поставить/заменить sectPr у абзаца (конец секции — этот абзац)."""
    pPr = paragraph._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:sectPr"))):
        pPr.remove(old)
    pPr.append(sectPr)


def _find_appendix_bounds(doc: Document) -> tuple[int | None, int | None, int | None]:
    """Индексы тела: Приложение А, последний абзац А, Приложение Б (не из TOC)."""
    paras = list(doc.paragraphs)
    # тело начинается после последнего «Введение» / после содержания
    body_start = 0
    intro_hits = [i for i, p in enumerate(paras) if (p.text or "").strip() == "Введение"]
    toc_idx = next(
        (i for i, p in enumerate(paras) if is_contents_heading((p.text or "").strip())),
        None,
    )
    if intro_hits:
        after_toc = [i for i in intro_hits if toc_idx is None or i > toc_idx]
        body_start = after_toc[-1] if after_toc else intro_hits[-1]

    app_a = None
    app_b = None
    for i, p in enumerate(paras):
        if i < body_start:
            continue
        t = (p.text or "").strip()
        if t == "Приложение А" or (
            t.startswith("Приложение А") and len(t) < 25 and "Форма" not in t
        ):
            if app_a is None:
                app_a = i
        if t == "Приложение Б" or (
            t.startswith("Приложение Б") and len(t) < 25 and "Форма" not in t
        ):
            app_b = i
    if app_a is None:
        return None, None, None
    if app_b is not None and app_b > app_a:
        return app_a, app_b - 1, app_b
    return app_a, len(paras) - 1, None


def set_appendix_a_landscape(doc: Document) -> dict:
    """
    Приложение А — альбомная ориентация (отдельная секция nextPage).
    До А и с Приложения Б — книжная.
    """
    report: dict = {"ok": False}
    app_a, last_a, app_b = _find_appendix_bounds(doc)
    report["app_a"] = app_a
    report["last_a"] = last_a
    report["app_b"] = app_b
    if app_a is None or last_a is None:
        report["error"] = "appendix_a_not_found"
        return report
    if app_a == 0:
        report["error"] = "appendix_a_at_start"
        return report

    paras = list(doc.paragraphs)
    # конец предыдущей (книжной) секции — абзац перед «Приложение А»
    prev = paras[app_a - 1]
    _set_paragraph_sect_pr(
        prev, _build_sect_pr(landscape=False, next_page=True, doc=doc)
    )
    # убрать pageBreakBefore у заголовка А — разрыв даёт секция
    _set_page_break_before(paras[app_a], False)

    # конец альбомной секции А
    _set_paragraph_sect_pr(
        paras[last_a], _build_sect_pr(landscape=True, next_page=True, doc=doc)
    )

    # тело документа (последняя секция) — книжная (Приложение Б и далее)
    body_sect = doc.sections[-1]._sectPr
    pgSz = body_sect.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        body_sect.insert(0, pgSz)
    pgSz.set(qn("w:w"), _A4_PORTRAIT_W)
    pgSz.set(qn("w:h"), _A4_PORTRAIT_H)
    if qn("w:orient") in pgSz.attrib:
        del pgSz.attrib[qn("w:orient")]

    if app_b is not None:
        _set_page_break_before(paras[app_b], False)  # nextPage секции А уже переносит

    report["ok"] = True
    report["sections_after"] = len(doc.sections)
    return report


def format_soglasovanie_as_signatories(doc: Document) -> int:
    """Лист согласования — как блок подписантов (без отдельной страницы)."""
    # СОГЛАСОВАНО: → СОГЛАСОВАНО
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.lower().startswith("согласовано"):
            _set_runs(p, "СОГЛАСОВАНО", bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
    return apply_signatory_block_format(doc, doc_type="polozhenie")


TITLE_SPACERS_BEFORE_NAME = 17  # пустые ПОСЛЕ грифа → название ровно посередине листа
TITLE_SPACERS_AFTER_NAME = 14  # пустые ПОСЛЕ названия → «Минск» внизу
# Гриф утверждения как в Инструкции по делопроизводству (2025 / .docx):
# левый отступ 6970 twips ≈ 12,3 см — все строки грифа с одной вертикали.
APPROVE_STAMP_LEFT = Twips(6970)

DELOPROIZVODSTVO_TITLE_SAMPLE = (
    r"C:\Users\v.dubovik\Desktop\Нормативка\Инструкция по делопроизводству.docx"
)


def _is_approve_stamp_line(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    u = t.upper()
    if u == "УТВЕРЖДЕНО" or u.startswith("УТВЕРЖДАЮ"):
        return True
    if t.startswith("Приказ") or t.startswith("Постановление"):
        return True
    if t.startswith("(в ред") or t.startswith("(в редакции"):
        return True
    return False


def _is_doc_title_line(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    u = t.upper()
    if u in ("ПОЛОЖЕНИЕ", "ИНСТРУКЦИЯ", "ПРАВИЛА", "ПОРЯДОК"):
        return True
    if u.startswith("ПОЛОЖЕНИЕ ") or t.startswith("об организации") or t.startswith("о порядке"):
        return True
    if t.startswith("по делопроизводству") or t.startswith("ПО ДЕЛОПРОИЗВОДСТВУ"):
        return True
    return False


def _style_approve_stamp_line(paragraph) -> None:
    """Гриф справа: левый отступ 6970 twips, выравнивание влево (как в образце)."""
    t = (paragraph.text or "").strip()
    if (paragraph.text or "") != t:
        _set_runs(paragraph, t, bold=False)
    zero_para_spacing(paragraph)
    set_single_line_spacing(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = APPROVE_STAMP_LEFT
    pf.right_indent = Cm(0)
    for r in paragraph.runs:
        set_run_font(r, font_name="Times New Roman", font_size=Pt(14), bold=False)


def _insert_title_spacers_after(anchor, count: int, *, center: bool = True):
    """Вставить пустые абзацы-распорки после anchor; вернуть последний."""
    for _ in range(count):
        np = _insert_paragraph_after(anchor, "")
        zero_para_spacing(np)
        set_single_line_spacing(np)
        np.paragraph_format.first_line_indent = Cm(0)
        np.paragraph_format.left_indent = Cm(0)
        if center:
            np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor = np
    return anchor


def _clear_empties_between(paras: list, start_idx: int, end_idx: int) -> int:
    """Удалить пустые абзацы строго между start_idx и end_idx (не включая границы)."""
    removed = 0
    for i in range(end_idx - 1, start_idx, -1):
        if not (paras[i].text or "").strip():
            _delete_paragraph(paras[i])
            removed += 1
    return removed


def format_title_page_only(doc: Document) -> dict:
    """
    Титульный лист: гриф справа; название ровно посередине листа
    (центр по горизонтали и по вертикали); «Минск» внизу по центру.
    Эталон распорок: правка Дубовика (~17 пустых до названия, ~14 после).
    """
    report: dict = {"ok": False, "sample": DELOPROIZVODSTVO_TITLE_SAMPLE}
    paras = list(doc.paragraphs)
    if not paras:
        return report

    approve_idx = next(
        (
            i
            for i, p in enumerate(paras)
            if (p.text or "").strip().upper() in ("УТВЕРЖДЕНО", "УТВЕРЖДАЮ")
            or (p.text or "").strip().upper().startswith("УТВЕРЖДАЮ")
        ),
        None,
    )
    after_title_idx = next(
        (i for i, p in enumerate(paras) if (p.text or "").strip() == "Предисловие"),
        None,
    )
    if after_title_idx is None:
        after_title_idx = next(
            (
                i
                for i, p in enumerate(paras)
                if is_contents_heading((p.text or "").strip())
            ),
            None,
        )
    report["approve"] = approve_idx
    report["after_title"] = after_title_idx
    if approve_idx is None:
        return report

    stamp_end = approve_idx
    for i in range(approve_idx, len(paras)):
        t = (paras[i].text or "").strip()
        if i > approve_idx and (_is_doc_title_line(t) or t.startswith("Минск")):
            break
        if i > approve_idx and after_title_idx is not None and i >= after_title_idx:
            break
        if t and _is_approve_stamp_line(t):
            stamp_end = i
        elif t and not _is_approve_stamp_line(t) and not _is_doc_title_line(t):
            if i == approve_idx + 1:
                stamp_end = i
            else:
                break
        elif not t and i > stamp_end:
            break

    report["stamp_end"] = stamp_end
    for i in range(approve_idx, stamp_end + 1):
        if (paras[i].text or "").strip():
            _style_approve_stamp_line(paras[i])

    paras = list(doc.paragraphs)
    stamp_end = next(
        (
            i
            for i in range(len(paras) - 1, -1, -1)
            if _is_approve_stamp_line((paras[i].text or "").strip())
            and i < (after_title_idx or len(paras))
        ),
        stamp_end,
    )
    title_start = next(
        (
            i
            for i, p in enumerate(paras)
            if _is_doc_title_line((p.text or "").strip())
            and i > stamp_end
            and (after_title_idx is None or i < after_title_idx)
        ),
        None,
    )
    if title_start is None:
        report["error"] = "no_title_name"
        return report

    report["empties_before_cleared"] = _clear_empties_between(paras, stamp_end, title_start)
    paras = list(doc.paragraphs)
    stamp_end = next(
        (
            i
            for i in range(len(paras) - 1, -1, -1)
            if _is_approve_stamp_line((paras[i].text or "").strip())
        ),
        stamp_end,
    )
    title_start = next(
        (
            i
            for i, p in enumerate(paras)
            if _is_doc_title_line((p.text or "").strip()) and i > stamp_end
        ),
        None,
    )
    if title_start is None:
        return report
    _insert_title_spacers_after(paras[stamp_end], TITLE_SPACERS_BEFORE_NAME, center=True)
    report["spacers_before_name"] = TITLE_SPACERS_BEFORE_NAME

    paras = list(doc.paragraphs)
    minsk_idx = next(
        (i for i, p in enumerate(paras) if (p.text or "").strip() == "Минск"),
        None,
    )
    after_now = next(
        (i for i, p in enumerate(paras) if (p.text or "").strip() == "Предисловие"),
        None,
    )
    stop = minsk_idx if minsk_idx is not None else after_now
    title_idxs = [
        i
        for i, p in enumerate(paras)
        if _is_doc_title_line((p.text or "").strip())
        and i > stamp_end
        and (stop is None or i < stop)
    ]
    for i in title_idxs:
        zero_para_spacing(paras[i])
        set_single_line_spacing(paras[i])
        paras[i].alignment = WD_ALIGN_PARAGRAPH.CENTER
        paras[i].paragraph_format.first_line_indent = Cm(0)
        paras[i].paragraph_format.left_indent = Cm(0)
        for r in paras[i].runs:
            set_run_font(r, font_name="Times New Roman", font_size=Pt(14), bold=True)
    title_end = title_idxs[-1] if title_idxs else title_start

    minsk_idx = next(
        (i for i, p in enumerate(paras) if (p.text or "").strip() == "Минск"),
        None,
    )
    if minsk_idx is not None:
        report["empties_after_cleared"] = _clear_empties_between(paras, title_end, minsk_idx)
        paras = list(doc.paragraphs)
        minsk_idx = next(
            (i for i, p in enumerate(paras) if (p.text or "").strip() == "Минск"),
            None,
        )
        title_end = max(
            (
                i
                for i in range(minsk_idx or 0)
                if _is_doc_title_line((paras[i].text or "").strip())
            ),
            default=title_end,
        )
        if minsk_idx is not None:
            _insert_title_spacers_after(
                paras[title_end], TITLE_SPACERS_AFTER_NAME, center=False
            )
            report["spacers_after_name"] = TITLE_SPACERS_AFTER_NAME
        paras = list(doc.paragraphs)
        minsk_idx = next(
            (i for i, p in enumerate(paras) if (p.text or "").strip() == "Минск"),
            None,
        )
        if minsk_idx is not None:
            mp = paras[minsk_idx]
            if (mp.text or "").strip() != "Минск":
                _set_runs(mp, "Минск", bold=False)
            zero_para_spacing(mp)
            set_single_line_spacing(mp)
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mp.paragraph_format.first_line_indent = Cm(0)
            mp.paragraph_format.left_indent = Cm(0)
            for r in mp.runs:
                set_run_font(r, font_name="Times New Roman", font_size=Pt(14), bold=False)

    paras = list(doc.paragraphs)
    after_title_idx = next(
        (i for i, p in enumerate(paras) if (p.text or "").strip() == "Предисловие"),
        None,
    )
    if after_title_idx is None:
        after_title_idx = next(
            (
                i
                for i, p in enumerate(paras)
                if is_contents_heading((p.text or "").strip())
            ),
            None,
        )
    if after_title_idx is not None:
        minsk_now = next(
            (i for i, p in enumerate(paras) if (p.text or "").strip() == "Минск"),
            None,
        )
        if minsk_now is not None:
            for i in range(after_title_idx - 1, minsk_now, -1):
                if not (paras[i].text or "").strip():
                    _delete_paragraph(paras[i])
            paras = list(doc.paragraphs)
            after_title_idx = next(
                (
                    i
                    for i, p in enumerate(paras)
                    if (p.text or "").strip() == "Предисловие"
                ),
                None,
            )
            if after_title_idx is None:
                after_title_idx = next(
                    (
                        i
                        for i, p in enumerate(paras)
                        if is_contents_heading((p.text or "").strip())
                    ),
                    None,
                )
        if after_title_idx is not None:
            _set_page_break_before(paras[after_title_idx], True)
            report["break_after_title"] = True

    report["ok"] = True
    return report


def update_toc_via_word(docx_path: str) -> bool:
    """Обновить поля TOC через Word (кликабельные ссылки появятся сразу)."""
    try:
        import win32com.client  # type: ignore
    except Exception as e:
        _log(f"TOC update: no pywin32 ({e})")
        return False

    word = None
    path = str(Path(docx_path).resolve())
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(path, ReadOnly=False, AddToRecentFiles=False)
        try:
            # 12 = wdFieldTOC; обновить все поля
            doc.Fields.Update()
            # оглавления отдельно
            for toc in doc.TablesOfContents:
                toc.Update()
        except Exception as e:
            _log(f"TOC Update warn: {e}")
        doc.Save()
        doc.Close(False)
        return True
    except Exception as e:
        _log(f"TOC Word COM failed: {e}")
        return False
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def apply_document_structure_rules(docx_path: str, *, update_toc: bool = True) -> dict:
    """
    Полный проход структуры документа.
    """
    report: dict = {"path": docx_path}
    doc = Document(docx_path)

    # пустые вне титула (титульные распорки добавим отдельно)
    report["empties_removed"] = collapse_extra_empties(doc, max_run=1)
    anchors = find_structure_anchors(doc)
    report["anchors"] = {k: v for k, v in anchors.items()}

    report["heading1"] = apply_heading1_to_body_chapters(doc, anchors.get("body_intro"))
    report["toc_lines_removed"] = replace_manual_toc_with_field(doc, anchors)
    report["page_breaks"] = apply_structural_page_breaks(doc, anchors)
    report["signatories"] = format_soglasovanie_as_signatories(doc)
    report["markers"] = strip_text_markers_everywhere(doc)
    report["captions"] = apply_form_caption_fonts(doc, size_pt=11)

    # заголовок «Содержание» / «Предисловие» по центру
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t in ("Содержание", "Предисловие", "Оглавление"):
            zero_para_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                set_run_font(r, font_name="Times New Roman", font_size=Pt(14), bold=True)

    # титул В КОНЦЕ: только утверждение + название + Минск; предисловие — с новой стр.
    report["title_page"] = format_title_page_only(doc)

    # Приложение А — альбомная ориентация (отдельная секция)
    try:
        report["appendix_a_landscape"] = set_appendix_a_landscape(doc)
    except Exception as e:
        report["appendix_a_landscape_error"] = str(e)

    doc.save(docx_path)
    report["toc_updated"] = False
    if update_toc:
        report["toc_updated"] = update_toc_via_word(docx_path)
    return report


if __name__ == "__main__":
    target = (
        r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Промышленная безопасность"
        r"\Положение о производственном контроле_по_образцу_ЦЭМ.docx"
    )
    print(apply_document_structure_rules(target))
