# -*- coding: utf-8 -*-
"""Общие операции форматирования Word (только допустимые теги OOXML)."""

from __future__ import annotations

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, Twips

# Эталон КЛ 31.07.2026: Body Text firstLine=709 twips ≈ 1,25 см; таблицы 12 пт
from .etalon_format_spec import (
    FIRST_INDENT_CM,
    FIRST_INDENT_TWIPS,
    MARGINS_DEFAULT_MM,
    PAGE_HEIGHT_MM,
    PAGE_WIDTH_MM,
    TABLE_FONT_PT,
    TITLE_NUMBER_LABEL_PT,
    ensure_etalon_styles,
)

# Титульная таблица («УТВЕРЖДАЮ») — как основной текст
TITLE_TABLE_FONT_PT = 14


def set_run_font(run, font_name="Times New Roman", font_size=Pt(14), bold=None):
    run.font.name = font_name
    run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


def configure_page(doc, left=30, right=15, top=20, bottom=20, font_name="Times New Roman", font_size_pt=14):
    """Поля и Normal — по эталону КЛ 31.07.2026 (секция титула 30/15/20/20)."""
    for section in doc.sections:
        section.page_width = Mm(PAGE_WIDTH_MM)
        section.page_height = Mm(PAGE_HEIGHT_MM)
        section.left_margin = Mm(left)
        section.right_margin = Mm(right)
        section.top_margin = Mm(top)
        section.bottom_margin = Mm(bottom)

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    # эталон: одинарный (line=240), before/after = 0
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)
    try:
        ensure_etalon_styles(doc)
    except Exception:
        pass


def remove_word_protection(doc) -> int:
    """
    Убрать ограничение редактирования Word из документа.
    Пароль агент никогда не ставит — тег часто копируется из эталона без пароля,
    но Word всё равно может спрашивать «пароль защиты».
    """
    removed = 0
    try:
        settings = doc.settings.element
    except Exception:
        return 0
    for tag in ("documentProtection", "writeProtection", "readOnlyRecommended"):
        el = settings.find(qn(f"w:{tag}"))
        if el is not None:
            settings.remove(el)
            removed += 1
    return removed


def save_docx_unprotected(doc, output_path: str) -> None:
    """Сохранить документ без защиты от редактирования."""
    remove_word_protection(doc)
    doc.save(output_path)


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def is_paragraph_in_table(paragraph) -> bool:
    """Абзац внутри ячейки таблицы Word (титул, приложение и т.п.)."""
    el = paragraph._element
    while el is not None:
        if el.tag == qn("w:tc"):
            return True
        el = el.getparent()
    return False


def set_single_line_spacing(paragraph) -> bool:
    """Одинарный межстрочный интервал. True — если меняли."""
    pf = paragraph.paragraph_format
    changed = False
    if pf.line_spacing_rule != WD_LINE_SPACING.SINGLE:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        changed = True
    # сбросить множитель/точные значения, если остались
    if pf.line_spacing not in (None, 1.0, 1):
        try:
            # при SINGLE line_spacing может быть None или 1
            if isinstance(pf.line_spacing, (int, float)) and abs(float(pf.line_spacing) - 1.0) > 0.05:
                pf.line_spacing = 1.0
                changed = True
        except Exception:
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            changed = True
    return changed


def zero_para_spacing(paragraph) -> bool:
    """
    Интервал перед/после абзаца = 0.
    Обязательно для текста, титула и ячеек таблиц.
    Всегда пишем явное 0 в XML (None = «как в стиле», это не ноль).
    """
    pf = paragraph.paragraph_format
    before = pf.space_before
    after = pf.space_after
    need_before = before is None or (hasattr(before, "pt") and abs(float(before.pt)) > 0.01)
    need_after = after is None or (hasattr(after, "pt") and abs(float(after.pt)) > 0.01)
    # даже если API уже показывает 0 — убедиться, что в XML есть before/after
    if not need_before and not need_after:
        try:
            pPr = paragraph._element.pPr
            sp = pPr.find(qn("w:spacing")) if pPr is not None else None
            if sp is None or sp.get(qn("w:before")) is None or sp.get(qn("w:after")) is None:
                need_before = True
                need_after = True
        except Exception:
            pass
    changed = False
    if need_before:
        pf.space_before = Pt(0)
        changed = True
    if need_after:
        pf.space_after = Pt(0)
        changed = True
    return changed


def clear_list_numbering(paragraph) -> bool:
    """
    Устаревшее имя: раньше снимало ВСЮ автонумерацию и ломало пункты 1.1, 1.2…
    Теперь снимает только маркеры-bullets. Десятичную нумерацию не трогает.
    """
    return clear_bullet_only(paragraph)


def clear_bullet_only(paragraph) -> bool:
    """Снять только bullet; decimal/numbering оставить."""
    pPr = paragraph._element.pPr
    if pPr is None:
        return False
    numPr = pPr.numPr
    if numPr is None:
        return False
    # без карты нумерации документа безопаснее НЕ удалять
    # (десятичные списки сохраняем). Для bullets вызывайте clear_bullet_markers_only.
    return False


def is_contents_heading(text: str) -> bool:
    """Заголовок «Содержание» / «Оглавление»."""
    t = re.sub(r"\s+", " ", (text or "").strip().lower().replace("ё", "е"))
    return t in ("содержание", "оглавление")


def is_heading_like(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if is_contents_heading(t):
        return True
    if t.isupper() and len(t) < 120:
        return True
    if t.startswith(
        ("ГЛАВА", "РАЗДЕЛ", "УТВЕРЖДАЮ", "СОГЛАСОВАНО", "УТВЕРЖДЕНО", "ПРИКАЗ", "ПРИЛОЖЕНИЕ")
    ):
        return True
    low = t.lower()
    # разделы вида «1 ОБЩИЕ ПОЛОЖЕНИЯ» / «2 ФУНКЦИИ…» (не пункты 1.10. …)
    if len(t) < 100 and re.match(r"^\d+\.?\s+[А-ЯЁA-ZА-яё]", t) and not re.match(
        r"^\d+\.\d+", t
    ):
        if any(
            w in low
            for w in (
                "общие положения",
                "общие требования",
                "обязанност",
                "права",
                "ответственност",
                "заключительн",
                "функции",
                "взаимоотношен",
                "требования",
                "порядок",
                "характеристик",
                "квалификац",
                "описание",
                "охрана труда",
                "безопасность",
            )
        ):
            return True
        if t[1:].strip().isupper() or t.isupper():
            return True
        # короткий заголовок главы без точки в конце предложения
        rest = re.sub(r"^\d+\.?\s*", "", t)
        if len(t) < 90 and not t.endswith((".", ";", ",")) and rest[:1].isupper():
            letters = [c for c in rest if c.isalpha()]
            if letters and sum(1 for c in letters if c.isupper()) / len(letters) >= 0.5:
                return True
    # короткий заголовок «1. Общие положения» без точки в конце длинного текста
    if re.match(r"^\d+\.\s+[А-ЯЁA-Z]", t) and len(t) < 90 and not t.endswith("."):
        if not re.search(r"\.\s+\S+\s+\S+\s+\S+", t):
            return True
    return False


def is_title_page_like(text: str, alignment=None) -> bool:
    """Титул / шапка: название вида, организация, год — без абзацного отступа."""
    t = text.strip()
    if not t:
        return False
    low = t.lower()
    # пункты и перечни в ТЕЛЕ инструкции — не титул (нужен отступ 1,25)
    if re.match(r"^\d+\.\d+", t):
        return False
    if t.endswith(";") or t.endswith(","):
        return False
    if low.startswith("иот ") or low.startswith("иот\t") or low.startswith("иот№"):
        return False
    if low.startswith("инструкция по охране труда «") or low.startswith(
        "инструкция по охране труда \""
    ):
        return False
    if low.startswith("примечание"):
        return False
    title_keys = (
        "государственное предприятие",
        "должностная инструкция",
        "рабочая инструкция",
        "утверждаю",
        "утверждено",
    )
    # короткое «Положение о …» только как заголовок документа, не пункт перечня
    if low.startswith("положение о") and len(t) < 80 and not t.endswith("."):
        return True
    # «Минсккоммунтеплосеть» / шапка — короткие строки
    if "минсккоммунтеплосеть" in low and len(t) < 100:
        return True
    if any(k in low for k in title_keys) and len(t) < 120:
        # не считать титулом длинные предложения тела
        if len(t) > 90 and ("руководствуется" in low or "должен" in low):
            return False
        return True
    if alignment == WD_ALIGN_PARAGRAPH.CENTER and len(t) < 100:
        if re.fullmatch(r"\d{4}\s*г\.?", t, flags=re.I):
            return True
        if "г." in low and len(t) < 40:
            return True
    return False


def is_signatory_like(text: str) -> bool:
    """Подписанты, разработчик, согласования — без абзацного отступа."""
    t = text.strip()
    if not t:
        return False
    # длинные пункты тела — НЕ подписанты
    if len(t) > 120:
        return False
    # пункты списка в теле («начальник службы;») — НЕ подписанты
    if t.endswith(";") or t.endswith(","):
        return False
    low = t.lower()
    keys = (
        "разработал",
        "разработчик",
        "согласовал",
        "согласовано",
        "утверждаю",
        "утверждено",
        "визирован",
        "ознакомлен",
        "лист ознакомления",
        "начальник сниот",
        "подпись",
        "расшифровка",
    )
    if any(k in low for k in keys):
        return True
    # короткие строки должностей/ФИО в блоке подписей (без ; в конце)
    if len(t) < 80 and any(
        k in low
        for k in (
            "директор",
            "первый заместитель",
            "главный инженер",
            "начальник ",
            "ведущий ",
        )
    ):
        if not re.match(r"^\d+(\.\d+)*\.", t):
            return True
    if re.search(r"_{3,}", t) and re.search(r"[А-ЯЁ]\.[А-ЯЁ]\.", t):
        return True
    if re.match(r"^[А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁа-яё\-]+$", t):
        return True
    return False


def is_commission_label_line(text: str) -> bool:
    """Метки состава комиссии — без абзацного отступа 1,25 см."""
    t = (text or "").strip().lower()
    return t.startswith(
        (
            "председатель комиссии:",
            "заместители председателя комиссии:",
            "члены комиссии:",
            "секретарь комиссии:",
        )
    )


def needs_first_indent(text: str, alignment=None) -> bool:
    """Нужен ли абзацный отступ 1,25 см."""
    if is_heading_like(text):
        return False
    if is_title_page_like(text, alignment):
        return False
    if is_signatory_like(text):
        return False
    if is_commission_label_line(text):
        return False
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return False
    return True


def apply_basic_office_format(
    input_path: str,
    output_path: str,
    margins=None,
    font_name="Times New Roman",
    font_size_pt=14,
    first_indent_mm=12.5,  # 1,25 см
    right_margin_mm=15,
) -> dict:
    """
    Базовое оформление:
    поля, шрифт, одинарный интервал (текст и титул),
    отступ начала строки 1,25 см (кроме заголовков, титула, подписантов, таблиц).
    В таблицах абзацный отступ НЕ применяется.
    Маркеры списков Word снимаются.
    """
    margins = margins or {"left": 30, "right": right_margin_mm, "top": 20, "bottom": 20}
    # пользователь: всегда 1,25 см
    first_indent = Twips(FIRST_INDENT_TWIPS)

    doc = Document(input_path)
    configure_page(
        doc,
        left=margins["left"],
        right=margins["right"],
        top=margins["top"],
        bottom=margins["bottom"],
        font_name=font_name,
        font_size_pt=font_size_pt,
    )

    changed_runs = 0
    changed_paras = 0
    cleared_markers = 0
    from .structure_fix import _numbering_fmt_map, clear_bullet_markers_only

    fmt_map = _numbering_fmt_map(doc)
    # титул «УТВЕРЖДАЮ» — 14 пт; таблицы тела (даже если первая) — 12 пт
    title_table_ids = set()
    for table in doc.tables:
        blob = " ".join(
            (p.text or "")
            for row in table.rows
            for cell in row.cells
            for p in cell.paragraphs
        ).lower()
        if "утверждаю" in blob or (
            "согласовано" in blob and ("инструкция" in blob or "утвержд" in blob)
        ):
            title_table_ids.add(id(table._tbl))

    for p in iter_all_paragraphs(doc):
        text = p.text.strip()
        in_table = is_paragraph_in_table(p)
        low = text.lower()
        if low == "номер инструкции" or low.startswith("номер инструкции"):
            run_size = Pt(TITLE_NUMBER_LABEL_PT)
        elif in_table:
            in_title_tbl = False
            el = p._element
            while el is not None:
                if el.tag == qn("w:tbl"):
                    in_title_tbl = id(el) in title_table_ids
                    break
                el = el.getparent()
            run_size = Pt(TITLE_TABLE_FONT_PT if in_title_tbl else TABLE_FONT_PT)
        else:
            run_size = Pt(font_size_pt)
        for run in p.runs:
            set_run_font(run, font_name=font_name, font_size=run_size, bold=run.bold)
            changed_runs += 1
        if clear_bullet_markers_only(p, fmt_map):
            cleared_markers += 1
        pf = p.paragraph_format
        # одинарный интервал — и в тексте, и на титуле, и в таблицах
        set_single_line_spacing(p)
        # перед/после абзаца всегда 0 (особенно в таблицах)
        zero_para_spacing(p)
        if not text:
            if in_table:
                pf.first_line_indent = Cm(0)
            continue

        # сброс «висячего» левого отступа маркеров
        if pf.left_indent is not None and pf.left_indent.mm < 20:
            # не трогаем грифы с большим отступом (>= 20 мм)
            if pf.left_indent.mm < 15:
                pf.left_indent = Cm(0)

        # В ТАБЛИЦАХ отступ не применяется; before/after уже 0
        if in_table:
            pf.first_line_indent = Cm(0)
            changed_paras += 1
            continue

        if needs_first_indent(text, p.alignment):
            if p.alignment is None or p.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            left = pf.left_indent
            if left is None or left.mm < 20:
                pf.first_line_indent = first_indent
        else:
            pf.first_line_indent = Cm(0)
            if is_heading_like(text) or is_title_page_like(text, p.alignment):
                if p.alignment is None and is_title_page_like(text, p.alignment):
                    pass
                elif is_heading_like(text) and p.alignment is None:
                    pass
        changed_paras += 1

    save_docx_unprotected(doc, output_path)
    return {
        "output": output_path,
        "paragraphs_touched": changed_paras,
        "runs_touched": changed_runs,
        "cleared_list_markers": cleared_markers,
        "mode": "basic_office",
    }
