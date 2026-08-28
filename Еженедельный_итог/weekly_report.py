# -*- coding: utf-8 -*-
"""
Еженедельный итог работы Дубовика В.В. (СНиОТ).

Что делает скрипт:
1) Определяет период недели (пн–вс, либо пн–сегодня, если неделя ещё идёт).
2) Собирает файлы, созданные/изменённые в этот период (Документы, Загрузки,
   рабочий стол, папка «Фоты проверок», сетевая папка N:\\…\\Дубовик В.В).
3) По именам фото IMG_ГГГГММДД_… группирует проверки по датам.
4) Считает количество: мероприятий, документов рассмотренных, документов разработанных
   (включая задачи СЭД с скриншотов: Ознакомиться/Согласовать → рассмотрено,
   Исполнить → мероприятия).
5) Пишет отчёт Word на рабочий стол и в N:\\…\\Отчеты (если доступен).

Ежедневно пн–пт в 16:30: отчёт за сегодня. Сб и вс — тишина.
Догон 08:30 пн–пт: предыдущий рабочий день (в понедельник — пятница).
снова открыть Word, если файл есть, но не правили (не сохраняли);
ничего не делать, если уже правили (папка «принятые» или сохранение >60 сек).
Задачи планировщика скрипт ставит сам при любом запуске.
Недельный --from-daily склеивает правленные дни КАК ЕСТЬ (без --format).

Запуск: двойной щелчок по «Сделать еженедельный итог.bat»
или: python weekly_report.py
      python weekly_report.py --daily
      python weekly_report.py --from-daily
      python weekly_report.py --from 2026-07-27 --to 2026-07-30
      python weekly_report.py --events 5 --reviewed 12 --developed 4
"""

from __future__ import annotations

import argparse
import datetime as dt
import os
import re
import shutil
import subprocess
import sys
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

try:
    from zoneinfo import ZoneInfo
except ImportError:
    ZoneInfo = None  # type: ignore[misc, assignment]

_ATTESTATION = Path(r"C:\Users\v.dubovik\AttestationSync")
if str(_ATTESTATION) not in sys.path:
    sys.path.insert(0, str(_ATTESTATION))
from format_weekly_report import (
    PROTECTED_DAILY_MESSAGE,
    format_weekly_docx,
    is_accepted_daily_path,
    is_protected_daily_report,
    is_weekly_letterhead,
    is_weekly_signatory_line,
    validate_weekly_document,
)
import fix_sniot_document as sniot_office

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:
    print("Нужна библиотека python-docx. Установите: pip install python-docx")
    sys.exit(1)

DESKTOP = Path.home() / "Desktop"
DOWNLOADS = Path.home() / "Downloads"
DOCUMENTS = Path.home() / "Documents"
PHOTOS = DESKTOP / "Фоты проверок"
N_DUBOVIK = Path(
    r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В"
)
N_REPORTS = N_DUBOVIK / "Отчеты"
DAILY_ROOT = DESKTOP / "Ежедневные отчёты"
DAILY_EDIT_DIR = DAILY_ROOT / "на_правку"
DAILY_ACCEPTED_DIR = DAILY_ROOT / "принятые"
DO_NOT_REFORMAT_MARK = "НЕ ОФОРМЛЯТЬ ПОВТОРНО"
EDITED_MIN_SECONDS = 60
SCHEDULER_TASK_DAILY = "СНиОТ_ежедневный_отчёт"
SCHEDULER_TASK_CATCH = "СНиОТ_ежедневный_отчёт_догон"
PYTHON_EXE = Path(r"C:\Users\v.dubovik\AppData\Local\Programs\Python\Python311\python.exe")

IMG_DATE_RE = re.compile(r"(?:IMG_|PHOTO_)?(\d{4})(\d{2})(\d{2})[_-]?", re.I)
FOLDER_DATE_RE = re.compile(
    r"(\d{1,2})[.\-_](\d{1,2})[.\-_](\d{2,4})|(\d{1,2})[.\-_](\d{1,2})[.\-_]?(\d{2})",
    re.I,
)

DOC_EXTS = {".docx", ".doc", ".pdf", ".xlsx", ".xls", ".rtf", ".txt"}
IMG_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".heic"}
SKIP_DIR_PARTS = {
    "node_modules",
    ".git",
    "__pycache__",
    "venv",
    ".venv",
    "AppData",
    "ViberDownloads",
}


def moscow_now() -> dt.datetime:
    """Сейчас по Москве (на этом ПК обычно уже Москва; иначе Europe/Moscow)."""
    if ZoneInfo is not None:
        try:
            return dt.datetime.now(ZoneInfo("Europe/Moscow")).replace(tzinfo=None)
        except Exception:
            pass
    return dt.datetime.now()


def is_weekend(day: dt.date) -> bool:
    """Суббота и воскресенье — выходные, дневных отчётов нет."""
    return day.weekday() >= 5


def previous_workday(day: dt.date) -> dt.date:
    """Предыдущий рабочий день (пн–пт). С понедельника — пятница."""
    d = day - dt.timedelta(days=1)
    while is_weekend(d):
        d -= dt.timedelta(days=1)
    return d


def week_bounds(today: dt.date | None = None) -> tuple[dt.datetime, dt.datetime]:
    """Понедельник 07:30 — пятница 23:59:59 (или сегодня, если неделя ещё идёт). Сб/вс в период не входят."""
    today = today or moscow_now().date()
    monday = today - dt.timedelta(days=today.weekday())
    start = dt.datetime.combine(monday, dt.time(7, 30))
    friday = monday + dt.timedelta(days=4)
    if is_weekend(today):
        end_day = friday
    else:
        end_day = min(today, friday)
    end = dt.datetime.combine(end_day, dt.time(23, 59, 59))
    return start, end


def day_bounds(
    day: dt.date | None = None,
    now: dt.datetime | None = None,
) -> tuple[dt.datetime, dt.datetime]:
    """Текущие сутки по Москве: с 00:00 до сейчас (если это сегодня) или до 23:59:59."""
    now = now or moscow_now()
    day = day or now.date()
    start = dt.datetime.combine(day, dt.time(0, 0, 0))
    if now.date() == day:
        end = now.replace(microsecond=0)
    else:
        end = dt.datetime.combine(day, dt.time(23, 59, 59))
    if end < start:
        end = start
    return start, end


def daily_filename(day: dt.date) -> str:
    return f"Ежедневный_отчёт_{day.isoformat()}.docx"


def dates_in_period(start: dt.datetime, end: dt.datetime) -> list[dt.date]:
    days: list[dt.date] = []
    cur = start.date()
    last = end.date()
    while cur <= last:
        days.append(cur)
        cur += dt.timedelta(days=1)
    return days


def workdays_in_period(start: dt.datetime, end: dt.datetime) -> list[dt.date]:
    """Только пн–пт."""
    return [d for d in dates_in_period(start, end) if not is_weekend(d)]


def is_internal_report_path(path: Path) -> bool:
    """Свои отчёты и служебные файлы в статистику работы не входят."""
    name = path.name
    if name.startswith("~$"):
        return True
    s = str(path)
    s_low = s.lower().replace("ё", "е")
    n_low = name.lower().replace("ё", "е")
    if "еженедельный_итог" in s_low:
        return True
    if "ежедневные отчеты" in s_low:
        return True
    if n_low.startswith("отчет_о_работе_дубовик"):
        return True
    if n_low.startswith("ежедневный_отчет_"):
        return True
    return False


def ensure_daily_folders(
    edit_dir: Path | None = None,
    accepted_dir: Path | None = None,
) -> tuple[Path, Path]:
    edit_dir = edit_dir or DAILY_EDIT_DIR
    accepted_dir = accepted_dir or DAILY_ACCEPTED_DIR
    edit_dir.mkdir(parents=True, exist_ok=True)
    accepted_dir.mkdir(parents=True, exist_ok=True)
    readme = (edit_dir.parent) / "_ПРОЧТИТЕ.txt"
    if not readme.is_file():
        readme.write_text(
            "ЕЖЕДНЕВНЫЕ ОТЧЁТЫ\n"
            "==================\n\n"
            "Каждый день в 16:30 (если компьютер включён) сюда кладётся Word "
            "за сегодня: папка «на_правку».\n\n"
            "Что сделать:\n"
            "1) Откройте файл Ежедневный_отчёт_ГОД-МЕСЯЦ-ДЕНЬ.docx\n"
            "2) Поправьте текст в Word\n"
            "3) Сохраните (тот же файл) — этого достаточно.\n"
            "   Можно дополнительно скопировать готовый файл в папку «принятые».\n\n"
            "Недельный итог заберёт эти файлы КАК ЕСТЬ: без повторного оформления, "
            "спеллера и замены слов.\n\n"
            "Суббота и воскресенье — выходные, отчётов нет.\n"
            "Если пятницу в 16:30 не правили — в понедельник в 08:30 откроется отчёт за пятницу.\n"
            "Задачи 16:30 и 08:30 планировщик ставит сам; в выходной скрипт сразу выходит.\n",
            encoding="utf-8",
        )
    return edit_dir, accepted_dir


def lookup_daily_in_dir(folder: Path, day: dt.date) -> Path | None:
    if not folder.is_dir():
        return None
    exact = folder / daily_filename(day)
    if exact.is_file():
        return exact
    needle = day.isoformat()
    matches = [
        p
        for p in folder.glob("*.docx")
        if needle in p.name and not p.name.startswith("~$")
    ]
    if not matches:
        return None
    if len(matches) == 1:
        return matches[0]
    return max(matches, key=lambda p: p.stat().st_mtime)


def find_daily_for_day(
    day: dt.date,
    *,
    edit_dir: Path | None = None,
    accepted_dir: Path | None = None,
) -> tuple[Path | None, str]:
    """
    Приоритет: папка «принятые», затем «на_правку» (если сохранили поверх).
    source: accepted | edit | missing
    """
    edit_dir = edit_dir or DAILY_EDIT_DIR
    accepted_dir = accepted_dir or DAILY_ACCEPTED_DIR
    accepted = lookup_daily_in_dir(accepted_dir, day)
    if accepted is not None:
        return accepted, "accepted"
    edited = lookup_daily_in_dir(edit_dir, day)
    if edited is not None:
        return edited, "edit"
    return None, "missing"


def generated_stamp_path(docx: Path) -> Path:
    return docx.with_name(docx.stem + ".generated.txt")


def write_generated_stamp(docx: Path, when: dt.datetime | None = None) -> None:
    when = when or moscow_now()
    generated_stamp_path(docx).write_text(
        when.isoformat(timespec="seconds"), encoding="utf-8"
    )


def read_generated_time(docx: Path) -> dt.datetime | None:
    stamp = generated_stamp_path(docx)
    if not stamp.is_file():
        return None
    try:
        return dt.datetime.fromisoformat(stamp.read_text(encoding="utf-8").strip())
    except ValueError:
        return None


def was_daily_edited_by_user(
    path: Path,
    *,
    min_seconds: int = EDITED_MIN_SECONDS,
) -> bool:
    """
    Правил = сохранил в Word после создания, либо положил в «принятые».
    Открыл Word и закрыл без Save — не правка (mtime почти как при создании).
    """
    if path is None or not path.is_file():
        return False
    parts = [x.lower().replace("ё", "е") for x in path.parts]
    if "принятые" in parts:
        return True
    _ctime, mtime = file_times(path)
    generated = read_generated_time(path)
    base = generated or _ctime
    return (mtime - base).total_seconds() > min_seconds


def catch_up_target_day(now: dt.datetime | None = None) -> dt.date:
    """Рабочий день для догона 08:30: вчера, а в понедельник — пятница."""
    now = now or moscow_now()
    return previous_workday(now.date())


def plan_daily_run(
    day: dt.date,
    *,
    edit_dir: Path | None = None,
    accepted_dir: Path | None = None,
    catch_up: bool = False,
) -> tuple[str, Path | None, str]:
    """
    skip — уже правили (принятые или сохранение спустя >60 сек);
    reopen — файл есть, но не правили (только для догона 08:30);
    create — файла нет, или в 16:30 есть неотредактированный (пересобрать за сегодня).
    """
    path, source = find_daily_for_day(
        day, edit_dir=edit_dir, accepted_dir=accepted_dir
    )
    if source == "accepted" and path is not None:
        return "skip", path, "accepted"
    if path is not None and was_daily_edited_by_user(path):
        return "skip", path, "edited"
    if path is None:
        return "create", None, "missing"
    if catch_up:
        return "reopen", path, "unedited"
    return "create", path, "unedited"


def ensure_windows_daily_tasks() -> list[str]:
    """Поставить/обновить обе задачи планировщика. Вызывается при любом запуске."""
    notes: list[str] = []
    if os.environ.get("WEEKLY_SKIP_SCHEDULER") == "1":
        return ["Планировщик: пропуск (WEEKLY_SKIP_SCHEDULER=1)"]
    if sys.platform != "win32":
        return ["Планировщик: не Windows — задачи не ставлю"]
    py = str(PYTHON_EXE if PYTHON_EXE.is_file() else sys.executable)
    script = str(Path(__file__).resolve())
    jobs = (
        (SCHEDULER_TASK_DAILY, "16:30", f'"{py}" "{script}" --daily'),
        (SCHEDULER_TASK_CATCH, "08:30", f'"{py}" "{script}" --daily --catch-up'),
    )
    for name, at, tr in jobs:
        cmd = [
            "schtasks",
            "/Create",
            "/TN",
            name,
            "/SC",
            "DAILY",
            "/ST",
            at,
            "/F",
            "/RL",
            "LIMITED",
            "/IT",
            "/TR",
            tr,
        ]
        try:
            proc = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                check=False,
                encoding="utf-8",
                errors="replace",
            )
            if proc.returncode == 0:
                notes.append(f"Планировщик: {name} — каждый день в {at}")
            else:
                err = (proc.stderr or proc.stdout or "").strip()[:240]
                notes.append(f"Планировщик: не удалось {name}: {err}")
        except OSError as exc:
            notes.append(f"Планировщик: ошибка {name}: {exc}")
    return notes


def parse_args():
    p = argparse.ArgumentParser(description="Еженедельный отчёт СНиОТ")
    p.add_argument("--from", dest="date_from", help="Начало ГГГГ-ММ-ДД")
    p.add_argument("--to", dest="date_to", help="Конец ГГГГ-ММ-ДД")
    p.add_argument(
        "--out",
        dest="out",
        help="Куда сохранить .docx (по умолчанию — рабочий стол)",
    )
    p.add_argument(
        "--events",
        type=int,
        default=None,
        help="Число мероприятий вручную (иначе считает скрипт)",
    )
    p.add_argument(
        "--reviewed",
        type=int,
        default=None,
        help="Число рассмотренных документов вручную",
    )
    p.add_argument(
        "--developed",
        type=int,
        default=None,
        help="Число разработанных документов вручную",
    )
    p.add_argument(
        "--format",
        dest="format_path",
        help="Только оформить уже готовый .docx итога (с бэкапом)",
    )
    p.add_argument(
        "--check",
        dest="check_path",
        help="Только проверить оформление готового .docx итога",
    )
    p.add_argument(
        "--daily",
        action="store_true",
        help="Ежедневный отчёт за сутки (на правку в Word)",
    )
    p.add_argument(
        "--from-daily",
        action="store_true",
        dest="from_daily",
        help="Недельный итог из уже правленных дневных .docx (без повторного оформления)",
    )
    p.add_argument(
        "--date",
        dest="for_date",
        help="Дата ГГГГ-ММ-ДД для --daily (по умолчанию сегодня по Москве)",
    )
    p.add_argument(
        "--fill-missing",
        action="store_true",
        dest="fill_missing",
        help="Для --from-daily: если дневного файла нет — вставить сырой черновик за этот день",
    )
    p.add_argument(
        "--catch-up",
        "--missed-yesterday",
        dest="catch_up",
        action="store_true",
        help="Догон 08:30: вчерашний день. Нет файла — создать; не правили — открыть Word; правили — ничего",
    )
    p.add_argument(
        "--no-open",
        action="store_true",
        dest="no_open",
        help="Не открывать Word после сохранения",
    )
    return p.parse_args()


def in_range(ts: dt.datetime, start: dt.datetime, end: dt.datetime) -> bool:
    return start <= ts <= end


def file_times(path: Path) -> tuple[dt.datetime, dt.datetime]:
    """
    (время_создания/появления, время_последнего_сохранения).
    На Windows стараемся взять birth time; иначе st_ctime.
    Учёт работы — только по mtime (файл реально сохраняли).
    """
    st = path.stat()
    m = dt.datetime.fromtimestamp(st.st_mtime)
    birth = getattr(st, "st_birthtime", None)
    if birth:
        c = dt.datetime.fromtimestamp(birth)
    else:
        c = dt.datetime.fromtimestamp(st.st_ctime)
    return c, m


def should_skip(path: Path) -> bool:
    parts = {x.lower() for x in path.parts}
    return bool(parts & {x.lower() for x in SKIP_DIR_PARTS})


def is_likely_unedited_download(path: Path, ctime: dt.datetime, mtime: dt.datetime) -> bool:
    """Скачали в «Загрузки» и сразу не правили — в отчёт не берём."""
    if "downloads" not in {p.lower() for p in path.parts}:
        return False
    return abs((mtime - ctime).total_seconds()) < 180


def format_work_span(item: dict, start: dt.datetime, end: dt.datetime) -> str:
    """
    Текст про время правки/сохранения.
    Точного «сколько минут смотрели» ОС не хранит — только создание и сохранение.
    """
    c: dt.datetime = item["ctime"]
    m: dt.datetime = item["mtime"]
    saved = m.strftime("%d.%m.%Y %H:%M")
    # оба события в периоде и сохранение позже создания — оценка окна работы
    if in_range(c, start, end) and m > c:
        delta = m - c
        hours = delta.total_seconds() / 3600
        if 0.05 <= hours <= 10:  # от ~3 мин до 10 ч — похоже на одну сессию
            return (
                f"правка примерно с {c.strftime('%d.%m.%Y %H:%M')} "
                f"по {m.strftime('%H:%M')} "
                f"(~{int(delta.total_seconds() // 60)} мин), сохранение {saved}"
            )
        if c.date() == m.date():
            return (
                f"в работе {c.strftime('%d.%m.%Y')} "
                f"с {c.strftime('%H:%M')} до {m.strftime('%H:%M')}, "
                f"сохранение {saved}"
            )
    return f"сохранён {saved}"


def estimate_session_minutes(item: dict, start: dt.datetime, end: dt.datetime) -> int:
    """Оценка минут по разнице создание→сохранение в периоде (0, если нельзя оценить)."""
    c, m = item["ctime"], item["mtime"]
    if not in_range(m, start, end):
        return 0
    if not in_range(c, start, end) or m <= c:
        return 0
    mins = int((m - c).total_seconds() // 60)
    if mins < 3 or mins > 10 * 60:
        return 0
    return mins


def scan_docs(roots: list[Path], start: dt.datetime, end: dt.datetime) -> list[dict]:
    """Только файлы, которые СОХРАНЯЛИ в периоде (mtime), на ПК или N:."""
    found = []
    for root in roots:
        if not root.exists():
            continue
        try:
            iterator = root.rglob("*")
        except OSError:
            continue
        for path in iterator:
            try:
                if not path.is_file():
                    continue
                if should_skip(path):
                    continue
                if path.suffix.lower() not in DOC_EXTS:
                    continue
                if is_internal_report_path(path):
                    continue
                c, m = file_times(path)
                # главное правило: документ правили и сохранили в периоде
                if not in_range(m, start, end):
                    continue
                if is_likely_unedited_download(path, c, m):
                    continue
                found.append(
                    {
                        "path": path,
                        "name": path.name,
                        "ctime": c,
                        "mtime": m,
                        "touched": m,  # для отчёта — момент сохранения
                        "saved_in_period": True,
                        "created_in_period": in_range(c, start, end),
                        "rel": str(path),
                    }
                )
            except OSError:
                continue
    found.sort(key=lambda x: x["touched"])
    return found


def date_from_img_name(name: str) -> dt.date | None:
    m = IMG_DATE_RE.search(name)
    if not m:
        return None
    try:
        return dt.date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    except ValueError:
        return None


def scan_photos(photo_root: Path, start: dt.datetime, end: dt.datetime) -> dict[dt.date, list[Path]]:
    by_day: dict[dt.date, list[Path]] = defaultdict(list)
    if not photo_root.exists():
        return by_day
    start_d, end_d = start.date(), end.date()
    for path in photo_root.rglob("*"):
        try:
            if not path.is_file() or path.suffix.lower() not in IMG_EXTS:
                continue
            d = date_from_img_name(path.name)
            if d is None:
                # папка вида «14_07_26 …»
                for part in path.parts:
                    fm = FOLDER_DATE_RE.search(part)
                    if not fm:
                        continue
                    g = fm.groups()
                    if g[0] and g[1] and g[2]:
                        day, month, year = int(g[0]), int(g[1]), int(g[2])
                    elif g[3] and g[4] and g[5]:
                        day, month, year = int(g[3]), int(g[4]), int(g[5])
                    else:
                        continue
                    if year < 100:
                        year += 2000
                    try:
                        d = dt.date(year, month, day)
                    except ValueError:
                        d = None
                    break
            if d is None or d < start_d or d > end_d:
                continue
            by_day[d].append(path)
        except OSError:
            continue
    for d in by_day:
        by_day[d].sort(key=lambda p: p.name)
    return dict(sorted(by_day.items()))


def find_sed_screens(downloads: Path, start: dt.datetime, end: dt.datetime) -> list[Path]:
    """Скриншоты СЭД в Загрузках за период (по имени и/или дате файла)."""
    out: list[Path] = []
    if not downloads.exists():
        return out
    name_hints = ("экран", "screen", "снимок", "сед", "задач", "ознакомит", "согласов")
    for path in downloads.iterdir():
        try:
            if not path.is_file():
                continue
            if path.suffix.lower() not in {".png", ".jpg", ".jpeg", ".webp", ".bmp"}:
                continue
            c, m = file_times(path)
            if not (in_range(c, start, end) or in_range(m, start, end)):
                continue
            low = path.name.lower()
            if any(h in low for h in name_hints):
                out.append(path)
        except OSError:
            continue
    out.sort(key=lambda p: p.stat().st_mtime)
    return out


def ocr_image_windows(path: Path) -> str:
    """Распознать текст скриншота через встроенный Windows OCR."""
    script = Path(__file__).resolve().parent / "win_ocr.ps1"
    if not script.is_file():
        return ""
    out_txt = Path(__file__).resolve().parent / f"_ocr_tmp_{os.getpid()}.txt"
    try:
        import subprocess

        proc = subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(script),
                "-ImagePath",
                str(path),
                "-OutFile",
                str(out_txt),
            ],
            capture_output=True,
            timeout=90,
            check=False,
        )
        if out_txt.is_file() and out_txt.stat().st_size > 0:
            text = out_txt.read_text(encoding="utf-8").strip()
            return text
        err = (proc.stderr or b"").decode("utf-8", errors="replace").strip()
        out = (proc.stdout or b"").decode("utf-8", errors="replace").strip()
        if err or out:
            print(f"OCR предупреждение ({path.name}): {(err or out)[:240]}")
        return ""
    except Exception as e:
        print(f"OCR не удался для {path.name}: {e}")
        return ""
    finally:
        if out_txt.exists():
            try:
                out_txt.unlink()
            except OSError:
                pass


_SED_ACTION_RE = re.compile(
    r"(?i)\b(ознакомиться|согласовать|исполнить)\b\s*[«\"'«\"]?\s*(.+?)(?="
    r"\b(?:ознакомиться|согласовать|исполнить|отсутствие|приказ|внутренние|"
    r"файлы|подписанный|акт\b|\(\s*(?:понедельник|вторник|среда|четверг|пятница|"
    r"суббота|воскресенье)\s*\))"
    r"|$)",
)


def _normalize_sed_title(title: str) -> str:
    t = (title or "").strip(" \t\r\n«»\"'„“.…-_")
    t = re.sub(r"\s+", " ", t)
    # обрезать хвосты OCR: «__. (среда)», даты-дни недели
    t = re.sub(
        r"\s*[_.,]{0,4}\s*\(\s*(?:понедельник|вторник|среда|четверг|пятница|суббота|воскресенье)\s*\)\s*$",
        "",
        t,
        flags=re.I,
    )
    t = re.sub(r"[_…]{2,}\s*$", "", t)
    # частые ошибки OCR
    repl = (
        ("инструщия", "инструкция"),
        ("инструщии", "инструкции"),
        ("нупедого", "нулевого"),
        ("нупевого", "нулевого"),
        ("еженедепьное", "еженедельное"),
        ("допжностная", "должностная"),
        ("територий", "территорий"),
        ("мд ", "№ "),
        ("ng3", "№3"),
        ("ftc-", "РТС-"),
    )
    for a, b in repl:
        t = re.sub(a, b, t, flags=re.I)
    return t[:160].strip(" .,_-")


def parse_sed_tasks_from_text(text: str) -> list[dict]:
    """
    Из текста OCR вытащить задачи СЭД:
    Ознакомиться / Согласовать / Исполнить + название.
    """
    if not text or not text.strip():
        return []
    # склеить переносы, даты-заголовки оставить как разделители
    blob = re.sub(r"[\r\n]+", " ", text)
    blob = re.sub(r"\s+", " ", blob).strip()

    tasks: list[dict] = []
    for m in _SED_ACTION_RE.finditer(blob):
        action = m.group(1).capitalize()
        # normalize action spelling
        al = action.lower()
        if al.startswith("ознаком"):
            action = "Ознакомиться"
            kind = "reviewed"
        elif al.startswith("соглас"):
            action = "Согласовать"
            kind = "reviewed"
        else:
            action = "Исполнить"
            kind = "execute"
        title = _normalize_sed_title(m.group(2))
        if len(title) < 3:
            continue
        # отбросить мусор вроде одной даты
        if re.fullmatch(r"\d{1,2}\.\d{1,2}\.\d{2,4}", title):
            continue
        tasks.append(
            {
                "action": action,
                "title": title,
                "kind": kind,
                "raw": m.group(0)[:200],
            }
        )

    # дедуп: одинаковое действие + похожее начало названия
    seen: set[str] = set()
    uniq: list[dict] = []
    for t in tasks:
        title_key = re.sub(r"[^a-zа-яё0-9]+", "", t["title"].lower())[:36]
        key = t["action"].lower() + "|" + title_key
        if not title_key or key in seen:
            continue
        # если уже есть более длинный похожий — пропуск короткого
        skip = False
        for prev in list(seen):
            if prev.startswith(t["action"].lower() + "|") and (
                title_key in prev.split("|", 1)[-1] or prev.split("|", 1)[-1] in title_key
            ):
                skip = True
                break
        if skip:
            continue
        seen.add(key)
        uniq.append(t)
    return uniq


def extract_sed_tasks(screens: list[Path]) -> dict:
    """OCR всех скриншотов СЭД → задачи и счётчики."""
    all_tasks: list[dict] = []
    per_file: list[tuple[Path, list[dict], str]] = []
    for path in screens:
        text = ocr_image_windows(path)
        # если имя похоже на экран, но OCR пустой — всё равно учтём файл
        tasks = parse_sed_tasks_from_text(text)
        # эвристика: если в имени нет «экран/сед», а задач нет — пропуск
        low = path.name.lower()
        looks_sed = any(
            h in low for h in ("экран", "screen", "снимок", "сед", "задач")
        )
        if not tasks and not looks_sed and text:
            # текст есть, но без глаголов задач — не СЭД
            if not re.search(r"(?i)ознакомиться|согласовать|исполнить", text):
                continue
        per_file.append((path, tasks, text[:500]))
        all_tasks.extend(tasks)

    # общий дедуп между скриншотами
    seen: set[str] = set()
    uniq: list[dict] = []
    for t in all_tasks:
        key = (t["action"] + "|" + t["title"][:50]).lower()
        key = re.sub(r"[^a-zа-яё0-9|]+", "", key)
        if key in seen:
            continue
        seen.add(key)
        uniq.append(t)

    n_acq = sum(1 for t in uniq if t["action"] == "Ознакомиться")
    n_agr = sum(1 for t in uniq if t["action"] == "Согласовать")
    n_exe = sum(1 for t in uniq if t["action"] == "Исполнить")
    return {
        "tasks": uniq,
        "per_file": per_file,
        "n_acquaint": n_acq,
        "n_agree": n_agr,
        "n_execute": n_exe,
        "n_reviewed": n_acq + n_agr,
        "n_total": len(uniq),
    }


def set_run_font(run, bold=False, size=14):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(a), "Times New Roman")


def add_p(doc, text, bold=False, center=False, first=True, size=14):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.first_line_indent = Cm(1.25) if first and not center else Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(p.add_run(text), bold=bold, size=size)


def add_h(doc, text):
    add_p(doc, text, bold=True, first=False)


def add_bul(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.left_indent = Cm(0.75)
    pf.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(p.add_run("– " + text), size=14)


def group_docs_by_day(docs: list[dict]) -> dict[dt.date, list[dict]]:
    g: dict[dt.date, list[dict]] = defaultdict(list)
    for item in docs:
        g[item["touched"].date()].append(item)
    return dict(sorted(g.items()))


def normalize_edition_key(name: str) -> str:
    """Свести разные редакции одного документа к одному ключу."""
    s = name.lower()
    s = re.sub(r"\.(docx?|pdf|xlsx?|rtf|txt)$", "", s)
    for pat in (
        r"_оформлен\+?",
        r"_formatted",
        r"_проект",
        r"_эталон",
        r"_backup",
        r"_тест\w*",
        r"_с_проверками",
        r"_с_количествами",
        r"\s*\(\d+\)",
        r"_копия",
        r"— копия",
        r"- копия",
        r"_out\d*",
        r"_in\b",
        r"_final",
        r"\.{2,}",
    ):
        s = re.sub(pat, "", s, flags=re.I)
    s = re.sub(r"[\s_\-.–—]+", " ", s).strip()
    s = re.sub(r"\s*\d{1,2}[./]\d{1,2}[./]\d{2,4}.*$", "", s)
    return s[:100]


def dedupe_editions(docs: list[dict]) -> list[dict]:
    """Оставить по одной записи на логический документ (последняя редакция)."""
    best: dict[str, dict] = {}
    for item in docs:
        key = normalize_edition_key(item["name"])
        if not key or key.startswith("~$"):
            continue
        prev = best.get(key)
        if prev is None or item["touched"] >= prev["touched"]:
            best[key] = item
    return sorted(best.values(), key=lambda x: x["touched"])


def interesting_name(name: str) -> bool:
    low = name.lower()
    keys = (
        "акт",
        "приказ",
        "инструкц",
        "отчет",
        "отчёт",
        "комисс",
        "опо",
        "jsa",
        "риск",
        "протокол",
        "сед",
        "ноль",
        "травмат",
        "персонал",
        "провер",
        "ди ",
        "ри ",
        "докладн",
        "распоряж",
        "памятк",
        "карт",
    )
    return any(k in low for k in keys)


# Документы, которые обычно «разрабатываем / готовим сами»
DEVELOPED_KEYS = (
    "приказ",
    "акт",
    "докладн",
    "jsa",
    "карт",
    "инструкц",
    "проект",
    "оформлен",
    "черновик",
    "положен",
    "распоряж",
    "памятк",
    "отчет",
    "отчёт",
    "справк",
    "информац по смен",
    "персональн",
)

# Документы, которые чаще «рассматриваем / согласовываем / изучаем»
REVIEWED_KEYS = (
    "согласов",
    "подписанн",
    "ознакомит",
    "протокол",
    "re_",
    "закон",
    "ткп",
    "правил",
    "норм",
    "комментар",
    "zero",
    "нулев",
    "сед",
)

# По именам — признаки мероприятий (проверки, совещания, комиссии)
EVENT_DOC_KEYS = (
    "проверк",
    "акт",
    "совещани",
    "комисс",
    "обследован",
    "освидетельств",
    "ноль",
    "травмат",
    "нот",
)


def classify_document(item: dict, start: dt.datetime, end: dt.datetime) -> str:
    """
    Вернуть: developed | reviewed | other.

    Считаем ТОЛЬКО документы, сохранённые в периоде (mtime уже отфильтрован в scan_docs).
    developed — создан и сохранён на этой неделе (подготовка нового документа);
    reviewed  — файл был раньше, на этой неделе правили/сохраняли
                (рассмотрение, изучение, сверка с нормами, правка по замечаниям).
    """
    name = item["name"]
    low = name.lower()
    created_in = bool(item.get("created_in_period")) or in_range(item["ctime"], start, end)
    saved_in = in_range(item["mtime"], start, end)

    if not saved_in:
        return "other"

    # явные признаки рассмотрения / согласования / изучения норм
    if any(k in low for k in REVIEWED_KEYS) and not created_in:
        return "reviewed"
    if any(k in low for k in REVIEWED_KEYS) and any(
        k in low for k in ("подписанн", "протокол", "закон", "ткп", "re_", "норм")
    ):
        return "reviewed"

    if created_in and any(k in low for k in DEVELOPED_KEYS):
        return "developed"
    if created_in and interesting_name(name):
        return "developed"

    # существовал раньше — на этой неделе сохранили после правки/сверки
    if (not created_in) and interesting_name(name):
        return "reviewed"

    if created_in:
        return "developed"
    return "reviewed"


def count_stats(
    docs: list[dict],
    photos: dict[dt.date, list[Path]],
    start: dt.datetime,
    end: dt.datetime,
    sed_info: dict | None = None,
) -> dict:
    """
    Мероприятия = уникальные выезды/проверки по фото (день+папка)
                 + отдельные «событийные» документы (акты проверок и т.п.),
                   если по этому дню ещё не было фото-мероприятия
                 + задачи СЭД «Исполнить».
    Рассмотренные = файлы + задачи СЭД «Ознакомиться»/«Согласовать».
    """
    developed = []
    reviewed = []
    event_docs = []

    for item in docs:
        kind = classify_document(item, start, end)
        if kind == "developed":
            developed.append(item)
        elif kind == "reviewed":
            reviewed.append(item)
        low = item["name"].lower()
        if any(k in low for k in EVENT_DOC_KEYS) and interesting_name(item["name"]):
            event_docs.append(item)

    # Уникальные мероприятия по фото: (дата, папка объекта)
    photo_events = set()
    for day, files in photos.items():
        for f in files:
            photo_events.add((day, f.parent.name))

    # Документы-мероприятия, не перекрытые фото того же дня
    photo_days = set(photos.keys())
    extra_event_docs = []
    seen_names = set()
    for item in event_docs:
        key = item["name"].lower()
        if key in seen_names:
            continue
        seen_names.add(key)
        day = item["touched"].date()
        if day in photo_days and any(k in key for k in ("акт", "проверк", "нот", "травмат")):
            continue
        extra_event_docs.append(item)

    sed_info = sed_info or {}
    sed_reviewed = int(sed_info.get("n_reviewed") or 0)
    sed_execute = int(sed_info.get("n_execute") or 0)
    sed_total = int(sed_info.get("n_total") or 0)

    events_count = len(photo_events) + len(extra_event_docs) + sed_execute
    reviewed_n = len(reviewed) + sed_reviewed

    work_files = developed + reviewed
    session_mins = sum(estimate_session_minutes(x, start, end) for x in work_files)
    saves = sorted(x["mtime"] for x in work_files) if work_files else []

    return {
        "events": events_count,
        "events_photo": len(photo_events),
        "events_docs": len(extra_event_docs),
        "events_sed_execute": sed_execute,
        "photo_event_list": sorted(photo_events, key=lambda x: (x[0], x[1])),
        "event_docs": extra_event_docs,
        "developed": developed,
        "reviewed": reviewed,
        "developed_n": len(developed),
        "reviewed_n": reviewed_n,
        "reviewed_files_n": len(reviewed),
        "sed_reviewed_n": sed_reviewed,
        "sed_execute_n": sed_execute,
        "sed_total_n": sed_total,
        "sed_tasks": list(sed_info.get("tasks") or []),
        "work_session_minutes": session_mins,
        "work_first_save": saves[0] if saves else None,
        "work_last_save": saves[-1] if saves else None,
        "work_files_n": len(work_files),
    }


def collect_work(
    start: dt.datetime,
    end: dt.datetime,
    *,
    events: int | None = None,
    reviewed: int | None = None,
    developed: int | None = None,
) -> tuple[list[dict], dict[dt.date, list[Path]], list[Path], dict]:
    """Сбор файлов, фото и СЭД за период (тот же, что у недельного итога)."""
    roots = [DESKTOP, DOWNLOADS, N_DUBOVIK]
    docs = scan_docs(roots, start, end)
    docs = [d for d in docs if not is_internal_report_path(d["path"])]
    docs = dedupe_editions(docs)
    photos = scan_photos(PHOTOS, start, end)
    sed = find_sed_screens(DOWNLOADS, start, end)
    print(f"Скриншотов СЭД: {len(sed)}")
    sed_info = extract_sed_tasks(sed)
    print(
        f"Задач СЭД распознано: {sed_info['n_total']} "
        f"(ознакомиться/согласовать={sed_info['n_reviewed']}, "
        f"исполнить={sed_info['n_execute']})"
    )
    stats = count_stats(docs, photos, start, end, sed_info=sed_info)
    stats["manual"] = False
    if events is not None:
        stats["events"] = events
        stats["manual"] = True
    if reviewed is not None:
        stats["reviewed_n"] = reviewed
        stats["manual"] = True
    if developed is not None:
        stats["developed_n"] = developed
        stats["manual"] = True
    return docs, photos, sed, stats


def build_report(
    start: dt.datetime,
    end: dt.datetime,
    docs: list[dict],
    photos: dict[dt.date, list[Path]],
    sed: list[Path],
    out_path: Path,
    stats: dict,
    *,
    kind: str = "weekly",
    format_office: bool = True,
) -> Path:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(3)
        s.right_margin = Cm(0.8)

    d1 = start.strftime("%d.%m.%Y")
    d2 = end.strftime("%d.%m.%Y")
    today_s = moscow_now().strftime("%d.%m.%Y")
    is_daily = kind == "daily"

    add_p(doc, "ГОСУДАРСТВЕННОЕ ПРЕДПРИЯТИЕ «МИНСККОММУНТЕПЛОСЕТЬ»", True, True, False, 14)
    add_p(doc, "Служба надёжности и охраны труда (СНиОТ)", False, True, False, 14)
    doc.add_paragraph()
    add_p(doc, "ОТЧЁТ", True, True, False, 16)
    if is_daily:
        add_p(doc, "о выполненной работе за сутки", False, True, False)
        add_p(doc, d1, True, True, False)
    else:
        add_p(doc, "о выполненной работе за период", False, True, False)
        add_p(doc, f"с {d1} по {d2}", True, True, False)
    doc.add_paragraph()
    add_p(
        doc,
        "Исполнитель: ведущий инженер по промышленной безопасности СНиОТ Дубовик В.В.",
        first=False,
    )
    if is_daily:
        add_p(
            doc,
            "Это ежедневный отчёт НА ПРАВКУ. Поправьте текст в Word и сохраните "
            "этот же файл (или положите копию в папку «принятые»). "
            "Недельный итог заберёт файл как есть, без повторного оформления, "
            f"спеллера и замены слов. Метка: {DO_NOT_REFORMAT_MARK}.",
            first=True,
        )
    add_p(
        doc,
        "Документ сформирован автоматически скриптом «Еженедельный итог». "
        "В количественные итоги по файлам входят только документы, "
        "которые были отредактированы и СОХРАНЕНЫ на компьютере или на диске N: "
        "в указанный период (по времени последнего сохранения). "
        "Время правки — оценка по созданию/сохранению файла (ОС не хранит "
        "точную длительность «изучения» без сохранения).",
        first=True,
    )

    add_h(doc, "1. Количественные итоги")
    n_docs = len(docs)
    n_photo_days = len(photos)
    n_photos = sum(len(v) for v in photos.values())
    ev = stats["events"]
    rev = stats["reviewed_n"]
    dev = stats["developed_n"]

    period_phrase = f"За сутки {d1}" if is_daily else f"За период с {d1} по {d2}"
    add_p(
        doc,
        f"{period_phrase} проведено мероприятий: {ev}; "
        f"документов рассмотрено: {rev}; документов разработано: {dev}.",
        first=True,
    )
    add_bul(doc, f"Мероприятий (всего): {ev}")
    add_bul(
        doc,
        f"  из них проверок объектов по фотофиксации: {stats['events_photo']}; "
        f"иных мероприятий по документам: {stats['events_docs']}; "
        f"задач СЭД «Исполнить»: {stats.get('events_sed_execute', 0)}",
    )
    add_bul(
        doc,
        "Документов рассмотрено (правка / изучение / сверка с нормами / "
        f"ознакомление / согласование): {rev}",
    )
    add_bul(
        doc,
        f"  из них сохранено на ПК/N:: {stats.get('reviewed_files_n', rev)}; "
        f"задач СЭД «Ознакомиться»/«Согласовать»: {stats.get('sed_reviewed_n', 0)}",
    )
    add_bul(doc, f"Документов разработано (созданы и сохранены в периоде): {dev}")
    if stats.get("work_first_save") and stats.get("work_last_save"):
        add_bul(
            doc,
            "Работа с файлами (по времени сохранений): с "
            f"{stats['work_first_save'].strftime('%d.%m.%Y %H:%M')} "
            f"по {stats['work_last_save'].strftime('%d.%m.%Y %H:%M')} "
            f"(файлов с сохранением: {stats.get('work_files_n', 0)}).",
        )
    sess = int(stats.get("work_session_minutes") or 0)
    if sess > 0:
        add_bul(
            doc,
            f"Оценка суммарного времени правок по файлам "
            f"(где есть создание и сохранение в периоде, сессия до 10 ч): "
            f"примерно {sess // 60} ч {sess % 60} мин. "
            f"Это не хронометраж «изучения без сохранения».",
        )
    add_bul(
        doc,
        f"Справочно: файлов с сохранением за период — {n_docs}; "
        f"дней с фото — {n_photo_days}; снимков — {n_photos}; "
        f"скриншотов СЭД — {len(sed)}; задач СЭД распознано — {stats.get('sed_total_n', 0)}.",
    )
    if stats.get("manual"):
        add_p(
            doc,
            "Часть показателей задана вручную параметрами запуска (--events / --reviewed / --developed).",
            first=True,
        )
    else:
        add_p(
            doc,
            "Цифры посчитаны автоматически по файлам и фото. "
            "При необходимости поправьте вручную в этом абзаце перед сдачей.",
            first=True,
        )

    add_h(doc, "1.1. Мероприятия (детализация)")
    if stats["photo_event_list"] or stats["event_docs"]:
        for day, folder in stats["photo_event_list"]:
            add_bul(
                doc,
                f"{day.strftime('%d.%m.%Y')} — проверка объекта "
                f"(папка «{folder}»). Адрес: _______________",
            )
        for item in stats["event_docs"]:
            add_bul(
                doc,
                f"{item['touched'].strftime('%d.%m.%Y')} — {item['name']}",
            )
    else:
        add_p(doc, "Автоматически выделенных мероприятий не найдено.")

    add_h(doc, "1.2. Разработанные документы")
    if stats["developed"]:
        for item in stats["developed"][:40]:
            span = format_work_span(item, start, end)
            add_bul(
                doc,
                f"{item['name']} — {span}",
            )
        if len(stats["developed"]) > 40:
            add_bul(doc, f"… и ещё: {len(stats['developed']) - 40}")
    else:
        add_p(doc, "Разработанных документов за период не выявлено.")

    add_h(doc, "1.3. Рассмотренные / правленные документы")
    add_p(
        doc,
        "Файлы, которые уже существовали ранее и были сохранены в периоде "
        "(правка замечаний, изучение, сверка с нормативными требованиями и т.п.).",
        first=True,
    )
    if stats["reviewed"]:
        for item in stats["reviewed"][:40]:
            span = format_work_span(item, start, end)
            add_bul(
                doc,
                f"{item['name']} — {span}",
            )
        if len(stats["reviewed"]) > 40:
            add_bul(doc, f"… и ещё: {len(stats['reviewed']) - 40}")
    else:
        add_p(doc, "Сохранённых «рассмотренных» документов за период не выявлено.")

    add_h(doc, "2. Проверки объектов (по дате в имени фото)")
    if not photos:
        add_p(
            doc,
            "В папке «Фоты проверок» за этот период не найдено фото "
            "с датой в имени файла (IMG_ГГГГММДД_…) или в имени папки.",
        )
    else:
        add_p(
            doc,
            "Адреса с табличек на снимках скрипт сам не читает. "
            "Ниже — даты и папки/файлы; при сдаче отчёта допишите адрес объекта.",
            first=True,
        )
        for day, files in photos.items():
            add_h(doc, f"2.{day.strftime('%d.%m.%Y')}")
            folders = sorted({f.parent.name for f in files})
            for folder in folders:
                cnt = sum(1 for f in files if f.parent.name == folder)
                add_bul(doc, f"Папка «{folder}» — снимков: {cnt}. Адрес объекта: _______________")
            examples = ", ".join(f.name for f in files[:5])
            if len(files) > 5:
                examples += f" … (+ещё {len(files) - 5})"
            add_bul(doc, f"Примеры файлов: {examples}")

    add_h(doc, "3. Документы по дням")
    by_day = group_docs_by_day(docs)
    if not by_day:
        add_p(doc, "Документов за период не найдено.")
    else:
        for day, items in by_day.items():
            add_h(doc, f"3.{day.strftime('%d.%m.%Y')}")
            # сначала «важные», потом остальные (короткий список)
            important = [x for x in items if interesting_name(x["name"])]
            other = [x for x in items if x not in important]
            show = important[:25] + other[:10]
            for x in show:
                where = "N:" if str(x["path"]).upper().startswith("N:") else (
                    "Загрузки" if "Downloads" in str(x["path"]) else (
                        "Стол" if "Desktop" in str(x["path"]) else "ПК"
                    )
                )
                add_bul(
                    doc,
                    f"[{where}] {x['name']} — {format_work_span(x, start, end)}",
                )
            rest = len(items) - len(show)
            if rest > 0:
                add_bul(doc, f"… и ещё файлов: {rest}")

    add_h(doc, "4. СЭД")
    sed_tasks = stats.get("sed_tasks") or []
    if sed_tasks:
        n_acq = sum(1 for t in sed_tasks if t["action"] == "Ознакомиться")
        n_agr = sum(1 for t in sed_tasks if t["action"] == "Согласовать")
        n_exe = sum(1 for t in sed_tasks if t["action"] == "Исполнить")
        add_p(
            doc,
            f"Распознано задач СЭД: {len(sed_tasks)} "
            f"(Ознакомиться — {n_acq}, Согласовать — {n_agr}, Исполнить — {n_exe}). "
            f"Ознакомиться и Согласовать входят в «документов рассмотрено»; "
            f"Исполнить — в «мероприятия».",
            first=True,
        )
        for t in sed_tasks[:60]:
            add_bul(doc, f"{t['action']} — {t['title']}")
        if len(sed_tasks) > 60:
            add_bul(doc, f"… и ещё задач: {len(sed_tasks) - 60}")
        if sed:
            add_p(doc, "Исходные скриншоты:", first=False)
            for pth in sed:
                add_bul(doc, pth.name)
    elif sed:
        add_p(
            doc,
            "Найдены скриншоты, но задачи не распознаны (проверьте качество снимка). "
            "Перенесите формулировки вручную:",
        )
        for pth in sed:
            add_bul(doc, pth.name)
    else:
        add_p(
            doc,
            "Скриншот СЭД за период не найден. Сделайте снимок страницы задач "
            "(Ознакомиться / Согласовать / Исполнить) и положите PNG в папку «Загрузки» "
            "— при следующем запуске задачи попадут в итоги автоматически.",
        )

    add_h(doc, "5. Что дописать вручную перед сдачей")
    add_bul(doc, "Адреса объектов по фото (с табличек на снимках).")
    add_bul(
        doc,
        "Если скриншот СЭД был неполный — допишите задачи, которых не хватает.",
    )
    add_bul(doc, "Устные поручения руководства / совещания.")
    add_bul(
        doc,
        "При необходимости исправить числа: мероприятий / рассмотрено / разработано.",
    )

    add_h(doc, "6. Примечание")
    add_p(
        doc,
        f"Скрипт: Desktop\\Еженедельный_итог\\weekly_report.py"
        f"{' --daily' if is_daily else ''}. "
        f"Дата формирования: {today_s}.",
        first=True,
    )
    if is_daily:
        add_p(
            doc,
            f"После правки сохраните файл. Метка: {DO_NOT_REFORMAT_MARK}.",
            first=True,
        )
    doc.add_paragraph()
    add_p(doc, "Ведущий инженер по промышленной безопасности СНиОТ", first=False)
    add_p(doc, "_________________ / В.В. Дубовик /", first=False)
    add_p(doc, today_s, first=False)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    if format_office:
        formatted = format_weekly_docx(
            out_path,
            backup=False,
            allow_daily_initial=is_daily,
        )
        if formatted.get("issues"):
            print("Оформление: есть замечания:")
            for item in formatted["issues"][:12]:
                print(" -", item)
        else:
            print("Оформление: Инструкция по делопроизводству 2025, без двойных пробелов")
    else:
        print("Оформление дневного/принятого файла не повторялось (как есть).")
    return out_path


def _insert_before_sectpr(dest_body, element) -> None:
    sect = dest_body.find(qn("w:sectPr"))
    if sect is not None:
        sect.addprevious(element)
    else:
        dest_body.append(element)


def copy_docx_body_as_is(
    src_path: Path,
    dest_doc: Document,
    *,
    skip_letterhead: bool = True,
) -> int:
    """Скопировать абзацы и таблицы без оформления, спеллера и замен текста."""
    src = Document(str(src_path))
    dest_body = dest_doc.element.body
    skipping = skip_letterhead
    copied = 0
    for child in list(src.element.body):
        if child.tag == qn("w:sectPr"):
            continue
        if skipping and child.tag == qn("w:p"):
            texts = [node.text or "" for node in child.iter(qn("w:t"))]
            text = "".join(texts)
            if (not text.strip()) or is_weekly_letterhead(text) or is_weekly_signatory_line(text):
                continue
            skipping = False
        _insert_before_sectpr(dest_body, deepcopy(child))
        copied += 1
    return copied


def assemble_weekly_from_daily(
    start: dt.datetime,
    end: dt.datetime,
    out_path: Path,
    *,
    edit_dir: Path | None = None,
    accepted_dir: Path | None = None,
    fill_missing: bool = False,
) -> tuple[Path, list[str]]:
    """
    Склеить недельный итог из дневных Word.
    Принятые / сохранённые дневные файлы НЕ прогонять через --format.
    """
    edit_dir, accepted_dir = ensure_daily_folders(edit_dir, accepted_dir)
    notes: list[str] = []
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(0.8)
    sniot_office.apply_page_setup_deloproizvodstvo(doc)

    d1 = start.strftime("%d.%m.%Y")
    d2 = end.strftime("%d.%m.%Y")
    today_s = moscow_now().strftime("%d.%m.%Y")

    add_p(doc, "ГОСУДАРСТВЕННОЕ ПРЕДПРИЯТИЕ «МИНСККОММУНТЕПЛОСЕТЬ»", True, True, False, 14)
    add_p(doc, "Служба надёжности и охраны труда (СНиОТ)", False, True, False, 14)
    doc.add_paragraph()
    add_p(doc, "ОТЧЁТ", True, True, False, 16)
    add_p(doc, "о выполненной работе за период", False, True, False)
    add_p(doc, f"с {d1} по {d2}", True, True, False)
    doc.add_paragraph()
    add_p(
        doc,
        "Исполнитель: ведущий инженер по промышленной безопасности СНиОТ Дубовик В.В.",
        first=False,
    )
    add_p(
        doc,
        "Составлен из дневных отчётов (папки «принятые» и «на_правку»). "
        "Тексты дней включены как есть: без повторного оформления, "
        f"спеллера и замены слов. {DO_NOT_REFORMAT_MARK}.",
        first=True,
    )

    for day in workdays_in_period(start, end):
        add_h(doc, f"День {day.strftime('%d.%m.%Y')}")
        path, source = find_daily_for_day(
            day, edit_dir=edit_dir, accepted_dir=accepted_dir
        )
        if path is None and fill_missing:
            day_start, day_end = day_bounds(
                day, now=dt.datetime.combine(day, dt.time(23, 59, 59))
            )
            print(f"Нет дневного файла за {day.isoformat()} — собираю сырой черновик")
            docs, photos, sed, stats = collect_work(day_start, day_end)
            draft = edit_dir / daily_filename(day)
            build_report(
                day_start,
                day_end,
                docs,
                photos,
                sed,
                draft,
                stats,
                kind="daily",
                format_office=True,
            )
            path, source = draft, "draft"
        if path is None:
            add_p(
                doc,
                f"{day.strftime('%d.%m.%Y')} — дневной отчёт не найден "
                "(компьютер был выключен в 16:30 или файл не сохраняли). "
                "День пропущен.",
                first=True,
            )
            notes.append(f"{day.isoformat()}: нет файла — пропуск")
            continue
        label = {
            "accepted": "принятый (папка «принятые»)",
            "edit": "сохранённый в «на_правку»",
            "draft": "сырой черновик (дня не было — собран сейчас)",
        }.get(source, source)
        add_p(
            doc,
            f"Источник: {path.name} — {label}. Без повторного оформления.",
            first=True,
        )
        copied = copy_docx_body_as_is(path, doc, skip_letterhead=True)
        notes.append(f"{day.isoformat()}: {source} ({copied} фрагм.) {path.name}")

    add_h(doc, "Примечание")
    add_p(
        doc,
        f"Скрипт: Desktop\\Еженедельный_итог\\weekly_report.py --from-daily. "
        f"Дата сборки: {today_s}.",
        first=True,
    )
    doc.add_paragraph()
    add_p(doc, "Ведущий инженер по промышленной безопасности СНиОТ", first=False)
    add_p(doc, "_________________ / В.В. Дубовик /", first=False)
    add_p(doc, today_s, first=False)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    out_path.write_bytes(sniot_office.fix_page_numbering(out_path.read_bytes()))
    return out_path, notes


def _copy_to_n_reports(path: Path) -> None:
    if not N_REPORTS.exists():
        return
    try:
        dest = N_REPORTS / path.name
        shutil.copy2(path, dest)
        print(f"Копия: {dest}")
    except OSError as e:
        print(f"На N: не скопировано ({e})")


def _open_docx(path: Path) -> None:
    try:
        os.startfile(path)  # noqa: S606
    except OSError:
        pass


def _write_daily_run_log(path: Path, day: dt.date) -> None:
    log = Path(__file__).resolve().parent / "_last_daily_run.txt"
    try:
        log.write_text(
            f"{moscow_now().strftime('%Y-%m-%d %H:%M:%S')} МСК\n"
            f"день: {day.isoformat()}\n"
            f"файл: {path}\n",
            encoding="utf-8",
        )
    except OSError:
        pass


def main():
    args = parse_args()
    for line in ensure_windows_daily_tasks():
        print(line)
    if args.check_path:
        path = Path(args.check_path)
        if is_protected_daily_report(path) or is_accepted_daily_path(path):
            print(PROTECTED_DAILY_MESSAGE)
            print("Проверку оформления для дневного отчёта не выполняю.")
            return 0
        doc = Document(str(path))
        issues = validate_weekly_document(doc)
        issues.extend(sniot_office.validate_page_numbering(path.read_bytes()))
        if issues:
            print("Замечания:")
            for item in issues:
                print(" -", item)
            return 1
        print(f"Проверка OK: {path}")
        return 0
    if args.format_path:
        fp = Path(args.format_path)
        if is_protected_daily_report(fp) or is_accepted_daily_path(fp):
            print(PROTECTED_DAILY_MESSAGE)
            return 2
        result = format_weekly_docx(args.format_path, backup=True)
        for act in result["actions"]:
            print(act)
        if result["issues"]:
            print("Замечания:")
            for item in result["issues"]:
                print(" -", item)
            return 1
        print(f"Готово: {result['path']}")
        if not args.no_open:
            _open_docx(Path(result["path"]))
        return 0
    if (args.daily or args.catch_up) and args.from_daily:
        print("Нельзя одновременно --daily/--catch-up и --from-daily.")
        return 2

    if args.daily or args.catch_up:
        today = moscow_now().date()
        if not args.for_date and is_weekend(today):
            print("Выходной (суббота/воскресенье) — дневных отчётов нет.")
            return 0
        edit_dir, accepted_dir = ensure_daily_folders()
        if args.catch_up:
            if args.for_date:
                day = dt.datetime.strptime(args.for_date, "%Y-%m-%d").date()
            else:
                day = catch_up_target_day()
            print(f"Догон 08:30: проверяю рабочий день {day.isoformat()} (МСК)")
        elif args.for_date:
            day = dt.datetime.strptime(args.for_date, "%Y-%m-%d").date()
        else:
            day = today
        if is_weekend(day):
            print("За субботу и воскресенье дневные отчёты не делаем.")
            return 0
        action, existing, source = plan_daily_run(
            day,
            edit_dir=edit_dir,
            accepted_dir=accepted_dir,
            catch_up=bool(args.catch_up),
        )
        if action == "skip":
            print(
                f"Отчёт за {day.isoformat()} уже правили ({source}): {existing}. "
                "Ничего не делаю."
            )
            _write_daily_run_log(existing or Path("skip"), day)
            return 0
        if action == "reopen":
            print(
                f"Отчёт за {day.isoformat()} не правили (не сохраняли). "
                f"Снова открываю Word: {existing}"
            )
            _write_daily_run_log(existing or Path("reopen"), day)
            if not args.no_open and existing is not None:
                _open_docx(existing)
            return 0
        start, end = day_bounds(day)
        print(f"Ежедневный отчёт за {day.isoformat()} (МСК): {start} — {end}")
        docs, photos, sed, stats = collect_work(
            start,
            end,
            events=args.events,
            reviewed=args.reviewed,
            developed=args.developed,
        )
        print(
            f"Итоги: мероприятий={stats['events']}, "
            f"рассмотрено={stats['reviewed_n']}, "
            f"разработано={stats['developed_n']}"
        )
        out = Path(args.out) if args.out else edit_dir / daily_filename(day)
        path = build_report(
            start, end, docs, photos, sed, out, stats, kind="daily", format_office=True
        )
        print(f"Готово на правку: {path}")
        print("Поправьте Word и СОХРАНИТЕ. Иначе завтра в 08:30 файл откроется снова.")
        write_generated_stamp(path)
        _write_daily_run_log(path, day)
        _copy_to_n_reports(path)
        if not args.no_open:
            _open_docx(path)
        return 0

    if args.date_from and args.date_to:
        start = dt.datetime.strptime(args.date_from, "%Y-%m-%d").replace(
            hour=7, minute=30
        )
        end = dt.datetime.strptime(args.date_to, "%Y-%m-%d").replace(
            hour=23, minute=59, second=59
        )
    else:
        start, end = week_bounds()

    if args.from_daily:
        print(f"Неделя из дневных: {start} — {end}")
        name = (
            f"Отчёт_о_работе_Дубовик_ВВ_"
            f"{start.strftime('%d.%m.%Y')}-{end.strftime('%d.%m.%Y')}.docx"
        )
        out = Path(args.out) if args.out else DESKTOP / name
        path, notes = assemble_weekly_from_daily(
            start, end, out, fill_missing=args.fill_missing
        )
        for line in notes:
            print(line)
        print(f"Готово (дневные как есть, без повторного оформления): {path}")
        _copy_to_n_reports(path)
        if not args.no_open:
            _open_docx(path)
        return 0

    print(f"Период: {start} — {end}")
    docs, photos, sed, stats = collect_work(
        start,
        end,
        events=args.events,
        reviewed=args.reviewed,
        developed=args.developed,
    )
    print(
        f"Итоги: мероприятий={stats['events']}, "
        f"рассмотрено={stats['reviewed_n']}, "
        f"разработано={stats['developed_n']}"
    )

    name = (
        f"Отчёт_о_работе_Дубовик_ВВ_"
        f"{start.strftime('%d.%m.%Y')}-{end.strftime('%d.%m.%Y')}.docx"
    )
    out = Path(args.out) if args.out else DESKTOP / name
    path = build_report(start, end, docs, photos, sed, out, stats)
    print(f"Готово: {path}")
    _copy_to_n_reports(path)
    if not args.no_open:
        _open_docx(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
