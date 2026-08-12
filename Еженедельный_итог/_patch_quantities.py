# -*- coding: utf-8 -*-
"""Вставить количественные итоги в базовый отчёт (редакции = 1 документ)."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = Path(
    r"C:\Users\v.dubovik\Desktop\Еженедельный_итог\Отчёт_о_работе_Дубовик_ВВ_27-30.07.2026.docx"
)
OUT = Path(
    r"C:\Users\v.dubovik\Desktop\Еженедельный_итог\Отчёт_о_работе_Дубовик_ВВ_27-30.07.2026.docx"
)
# если исходник открыт в Word — пишем рядом
OUT_ALT = Path(
    r"C:\Users\v.dubovik\Desktop\Еженедельный_итог\Отчёт_о_работе_Дубовик_ВВ_27-30.07.2026_с_количествами.docx"
)

# По содержанию отчёта + уникальные документы (разные редакции = 1)
EVENTS = 3
REVIEWED = 14
DEVELOPED = 7

EVENTS_DETAIL = [
    "Проверка рабочих мест СЭХ в рамках Недели нулевого травматизма (Акт № 3).",
    "Проверка рабочих мест АС в рамках Недели нулевого травматизма (Подписанный Акт_АС).",
    "Посещение Службы наладки и испытаний / проверка выполнения мероприятий "
    "(поручение еженедельного совещания № 17 п. 1.1).",
]

REVIEWED_DETAIL = [
    "Приказ / материалы «О проведении недели «нулевого травматизма»» (№ 559) — согласование.",
    "Рабочие инструкции слесаря по ремонту оборудования — согласование.",
    "Протокол № 3 от 16.07.2026 — ознакомление.",
    "Материалы «О результатах комплексной проверки РТС-5» (№ 475) — на контроле / рассмотрение.",
    "Должностная инструкция начальника сектора — согласование.",
    "Акт_ZERO_СЭХ_24.07.26 — рассмотрение / консультации.",
    "Подписанный Акт_АС — рассмотрение.",
    "Рабочая инструкция машинисту крана автомобильного — согласование.",
    "Рабочая инструкция электрогазосварщику — согласование.",
    "Рабочая инструкция уборщику территорий — согласование.",
    "Рабочая инструкция трактористу 6-го разряда — согласование.",
    "Рабочая инструкция слесарю по ремонту автомобилей — согласование.",
    "Рабочая инструкция слесарю по ремонту — согласование.",
    "Приказ о создании комиссии для проверки знаний по теплоустановкам и тепловым сетям — ознакомление (СЭД).",
]

DEVELOPED_DETAIL = [
    "Акт № 3 по результатам проведения Недели нулевого травматизма в СЭХ.",
    "Текущая информация по состоянию ОПО и ПОО (актуализация).",
    "Докладная записка «Информация по смен. персоналу 24.07.2026+».",
    "Проект / оформление приказа о комиссии по техническому расследованию нарушений "
    "теплотехнического оборудования (несколько редакций файла учтены как 1 документ).",
    "Приказ о создании комиссии для проверки знаний по теплоустановкам и тепловым сетям (доработка).",
    "JSA_01 «Слесарь по ремонту оборудования котельных и пылеприготовительных цехов».",
    "Карта рисков по профессии слесаря по ремонту оборудования котельных и пылеприготовительных цехов.",
]


def set_run_font(run, bold=False, size=14):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(a), "Times New Roman")


def style_para(p, text, bold=False, first=True, bullet=False):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3 if bullet else 6)
    pf.line_spacing = 1.15
    if bullet:
        pf.left_indent = Cm(0.75)
        pf.first_line_indent = Cm(0)
    else:
        pf.left_indent = Cm(0)
        pf.first_line_indent = Cm(1.25) if first else Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.text = ""
    t = ("– " + text) if bullet else text
    run = p.add_run(t) if not p.runs else p.runs[0]
    if not p.runs:
        run = p.add_run(t)
    else:
        # clear and set first run
        p.runs[0].text = t
        for r in p.runs[1:]:
            r.text = ""
        run = p.runs[0]
    set_run_font(run, bold=bold)


def add_styled(doc, text, bold=False, first=True, bullet=False):
    p = doc.add_paragraph()
    style_para(p, text, bold=bold, first=first, bullet=bullet)
    return p


def remove_old_quantity_block(doc: Document) -> None:
    """Удалить ранее вставленный блок 1.1, если есть."""
    idxs = []
    in_block = False
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith("1.1. Количественные"):
            in_block = True
        if in_block:
            if t.startswith("2. ") and "1.1" not in t:
                break
            idxs.append(i)
        if t.startswith("Количественно: мероприятий"):
            idxs.append(i)
    # delete from end
    for i in reversed(sorted(set(idxs))):
        p = doc.paragraphs[i]
        el = p._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def main():
    src = BASE
    bak = BASE.with_name(BASE.stem + "_без_количеств.docx")
    if bak.exists():
        # работаем от чистой копии без количеств
        src = bak
    if not src.exists():
        raise SystemExit(f"Нет файла: {src}")

    if not bak.exists() and BASE.exists():
        try:
            shutil.copy2(BASE, bak)
            print("backup:", bak)
            src = bak
        except OSError:
            pass

    doc = Document(str(src))
    remove_old_quantity_block(doc)

    # точка вставки: абзац перед «2. …»
    insert_after = None
    saw_resume = False
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith("1. Краткое резюме"):
            saw_resume = True
        if saw_resume and t.startswith("2. "):
            insert_after = doc.paragraphs[i - 1]._p
            break
    if insert_after is None:
        raise SystemExit("Не найден раздел 2 для вставки")

    # создаём абзацы в конце и переносим
    created = []
    created.append(add_styled(doc, "1.1. Количественные показатели", bold=True, first=False))
    created.append(
        add_styled(
            doc,
            f"За период с 27.07.2026 по 30.07.2026 проведено мероприятий: {EVENTS}; "
            f"документов рассмотрено: {REVIEWED}; документов разработано: {DEVELOPED}. "
            "Разные редакции одного документа (черновик / проект / оформлен / formatted) "
            "учтены как один документ.",
            first=True,
        )
    )
    created.append(add_styled(doc, "Мероприятия:", bold=True, first=False))
    for x in EVENTS_DETAIL:
        created.append(add_styled(doc, x, bullet=True))
    created.append(
        add_styled(
            doc,
            "Документы рассмотренные (согласование / ознакомление / изучение):",
            bold=True,
            first=False,
        )
    )
    for x in REVIEWED_DETAIL:
        created.append(add_styled(doc, x, bullet=True))
    created.append(
        add_styled(
            doc,
            "Документы разработанные (подготовлены / доработаны):",
            bold=True,
            first=False,
        )
    )
    for x in DEVELOPED_DETAIL:
        created.append(add_styled(doc, x, bullet=True))

    anchor = insert_after
    for p in created:
        el = p._p
        anchor.addnext(el)
        anchor = el

    # строка в раздел «Итоги»
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("1) Закрыт"):
            summary = add_styled(
                doc,
                f"Количественно: мероприятий — {EVENTS}; "
                f"рассмотрено документов — {REVIEWED}; "
                f"разработано документов — {DEVELOPED}.",
                first=True,
            )
            p._p.addprevious(summary._p)
            break

    saved = None
    for target in (OUT, OUT_ALT):
        try:
            doc.save(str(target))
            saved = target
            print("saved", target)
            break
        except PermissionError:
            print("locked:", target)

    if saved is None:
        raise SystemExit("Не удалось сохранить — закройте Word и запустите снова")

    copies = [
        Path(r"C:\Users\v.dubovik\Desktop\Отчёт_о_работе_Дубовик_ВВ_27-30.07.2026_с_количествами.docx"),
        Path(
            r"N:\9 - Служба надёжности и охраны труда (СНиОТ)\Дубовик В.В\Отчеты"
            r"\Отчёт_о_работе_Дубовик_ВВ_27-30.07.2026_с_количествами.docx"
        ),
    ]
    for c in copies:
        try:
            c.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(saved, c)
            print("copied", c)
        except OSError as e:
            print("skip", c, e)

    # проверка
    d2 = Document(str(saved))
    for i, p in enumerate(d2.paragraphs):
        t = (p.text or "").strip()
        if not t:
            continue
        if (
            t.startswith("1.1")
            or "проведено мероприятий" in t.lower()
            or t.startswith("Количественно")
            or t.startswith("Мероприятия:")
            or t.startswith("Документы рассмотренные")
            or t.startswith("Документы разработанные")
        ):
            print(f"{i:03d}| {t[:130]}")

    print(f"EVENTS={EVENTS} REVIEWED={REVIEWED} DEVELOPED={DEVELOPED}")


if __name__ == "__main__":
    main()
