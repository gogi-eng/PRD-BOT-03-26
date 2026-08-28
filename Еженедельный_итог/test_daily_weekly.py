# -*- coding: utf-8 -*-
"""Дневной отчёт и склейка недели из принятых файлов — без сети."""
from __future__ import annotations

import datetime as dt
import os
import sys
from pathlib import Path

from docx import Document

HERE = Path(r"C:\Users\v.dubovik\Desktop\Еженедельный_итог")
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))
ATTEST = Path(r"C:\Users\v.dubovik\AttestationSync")
if str(ATTEST) not in sys.path:
    sys.path.insert(0, str(ATTEST))

import format_weekly_report as weekly_fmt
import weekly_report as wr


def _write_daily_docx(path: Path, body: str) -> None:
    doc = Document()
    doc.add_paragraph("ГОСУДАРСТВЕННОЕ ПРЕДПРИЯТИЕ «МИНСККОММУНТЕПЛОСЕТЬ»")
    doc.add_paragraph("Служба надёжности и охраны труда (СНиОТ)")
    doc.add_paragraph("ОТЧЁТ")
    doc.add_paragraph("о выполненной работе за сутки")
    doc.add_paragraph(body)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def test_day_bounds_today_stops_at_now():
    now = dt.datetime(2026, 8, 17, 16, 30, 0)
    start, end = wr.day_bounds(dt.date(2026, 8, 17), now=now)
    assert start == dt.datetime(2026, 8, 17, 0, 0, 0)
    assert end == now


def test_day_bounds_past_day_full():
    now = dt.datetime(2026, 8, 17, 16, 30, 0)
    start, end = wr.day_bounds(dt.date(2026, 8, 16), now=now)
    assert start == dt.datetime(2026, 8, 16, 0, 0, 0)
    assert end == dt.datetime(2026, 8, 16, 23, 59, 59)


def test_week_bounds_monday_0730():
    start, end = wr.week_bounds(dt.date(2026, 8, 19))  # Wednesday
    assert start.weekday() == 0
    assert start.hour == 7 and start.minute == 30
    assert end.date() == dt.date(2026, 8, 19)


def test_weekend_is_saturday_sunday():
    assert wr.is_weekend(dt.date(2026, 8, 22)) is True  # Saturday
    assert wr.is_weekend(dt.date(2026, 8, 23)) is True  # Sunday
    assert wr.is_weekend(dt.date(2026, 8, 17)) is False  # Monday
    assert wr.is_weekend(dt.date(2026, 8, 21)) is False  # Friday


def test_monday_catch_up_is_friday():
    monday_morning = dt.datetime(2026, 8, 17, 8, 30, 0)
    assert wr.catch_up_target_day(monday_morning) == dt.date(2026, 8, 14)


def test_tuesday_catch_up_is_monday():
    now = dt.datetime(2026, 8, 18, 8, 30, 0)
    assert wr.catch_up_target_day(now) == dt.date(2026, 8, 17)


def test_week_bounds_on_sunday_is_mon_fri():
    start, end = wr.week_bounds(dt.date(2026, 8, 23))  # Sunday
    assert start.date() == dt.date(2026, 8, 17)
    assert end.date() == dt.date(2026, 8, 21)  # Friday
    days = wr.workdays_in_period(start, end)
    assert days == [
        dt.date(2026, 8, 17),
        dt.date(2026, 8, 18),
        dt.date(2026, 8, 19),
        dt.date(2026, 8, 20),
        dt.date(2026, 8, 21),
    ]
    assert all(d.weekday() < 5 for d in days)


def test_daily_filename():
    assert wr.daily_filename(dt.date(2026, 8, 17)) == "Ежедневный_отчёт_2026-08-17.docx"


def test_find_daily_prefers_accepted(tmp_path: Path):
    edit = tmp_path / "на_правку"
    accepted = tmp_path / "принятые"
    edit.mkdir()
    accepted.mkdir()
    day = dt.date(2026, 8, 17)
    name = wr.daily_filename(day)
    (edit / name).write_text("edit", encoding="utf-8")
    (accepted / name).write_text("ok", encoding="utf-8")
    path, source = wr.find_daily_for_day(day, edit_dir=edit, accepted_dir=accepted)
    assert source == "accepted"
    assert path == accepted / name


def test_find_daily_falls_back_to_edit(tmp_path: Path):
    edit = tmp_path / "на_правку"
    accepted = tmp_path / "принятые"
    edit.mkdir()
    accepted.mkdir()
    day = dt.date(2026, 8, 18)
    name = wr.daily_filename(day)
    (edit / name).write_bytes(b"x")
    path, source = wr.find_daily_for_day(day, edit_dir=edit, accepted_dir=accepted)
    assert source == "edit"
    assert path == edit / name


def test_find_daily_missing(tmp_path: Path):
    path, source = wr.find_daily_for_day(
        dt.date(2026, 8, 19),
        edit_dir=tmp_path / "на_правку",
        accepted_dir=tmp_path / "принятые",
    )
    assert path is None
    assert source == "missing"


def test_assemble_keeps_user_text_and_skips_missing(tmp_path: Path):
    edit = tmp_path / "на_правку"
    accepted = tmp_path / "принятые"
    monday = dt.date(2026, 8, 17)
    _write_daily_docx(
        accepted / wr.daily_filename(monday),
        "Пользователь  оставил  двойные пробелы и слово ПРАВКА17.",
    )
    _write_daily_docx(
        edit / wr.daily_filename(dt.date(2026, 8, 18)),
        "Черновик вторника ТЕКСТ18 без повторного оформления.",
    )
    out = tmp_path / "неделя.docx"
    start = dt.datetime(2026, 8, 17, 7, 30)
    end = dt.datetime(2026, 8, 19, 23, 59, 59)
    path, notes = wr.assemble_weekly_from_daily(
        start,
        end,
        out,
        edit_dir=edit,
        accepted_dir=accepted,
        fill_missing=False,
    )
    assert path.is_file()
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Пользователь  оставил  двойные пробелы" in text
    assert "ПРАВКА17" in text
    assert "ТЕКСТ18" in text
    assert "дневной отчёт не найден" in text
    assert any("2026-08-17: accepted" in n for n in notes)
    assert any("2026-08-18: edit" in n for n in notes)
    assert any("2026-08-19: нет файла" in n for n in notes)
    assert "НЕ ОФОРМЛЯТЬ ПОВТОРНО" in text


def test_copy_body_does_not_collapse_spaces(tmp_path: Path):
    src = tmp_path / "Ежедневный_отчёт_2026-08-17.docx"
    _write_daily_docx(src, "слово  слово   слово")
    dest = Document()
    wr.copy_docx_body_as_is(src, dest, skip_letterhead=True)
    joined = "\n".join(p.text for p in dest.paragraphs)
    assert "слово  слово   слово" in joined


def test_internal_report_paths_skipped():
    assert wr.is_internal_report_path(
        Path(r"C:\Users\v.dubovik\Desktop\Ежедневные отчёты\на_правку\Ежедневный_отчёт_2026-08-17.docx")
    )
    assert wr.is_internal_report_path(
        Path(r"C:\Users\v.dubovik\Desktop\Отчёт_о_работе_Дубовик_ВВ_01.01.2026-07.01.2026.docx")
    )
    assert not wr.is_internal_report_path(Path(r"C:\Users\v.dubovik\Desktop\Акт проверки.docx"))


def test_format_refuses_accepted_daily(tmp_path: Path):
    folder = tmp_path / "принятые"
    folder.mkdir()
    path = folder / "Ежедневный_отчёт_2026-08-17.docx"
    _write_daily_docx(path, "текст пользователя")
    try:
        weekly_fmt.format_weekly_docx(path, backup=False)
        assert False, "принятый дневной нельзя оформлять"
    except PermissionError as exc:
        assert "нельзя повторно оформлять" in str(exc)


def test_format_refuses_npravku_without_initial_flag(tmp_path: Path):
    folder = tmp_path / "Ежедневные отчёты" / "на_правку"
    folder.mkdir(parents=True)
    path = folder / "Ежедневный_отчёт_2026-08-17.docx"
    _write_daily_docx(path, "черновик")
    try:
        weekly_fmt.format_weekly_docx(path, backup=False)
        assert False, "дневной на правке нельзя оформлять повторно"
    except PermissionError:
        pass


def test_docagent_skips_protected_daily(tmp_path: Path):
    folder = tmp_path / "Ежедневные отчёты" / "принятые"
    folder.mkdir(parents=True)
    path = folder / "Ежедневный_отчёт_2026-08-17.docx"
    _write_daily_docx(path, "готово")
    result = weekly_fmt.process_weekly_itog_document(str(path))
    assert result["ok"] is True
    assert result["mode"] == "daily_accepted_skip"
    assert "нельзя повторно оформлять" in result["actions"][0]


def test_looks_like_weekly_includes_daily():
    assert weekly_fmt.looks_like_weekly_itog(
        r"C:\Users\v.dubovik\Desktop\Ежедневные отчёты\на_правку\Ежедневный_отчёт_2026-08-17.docx"
    )


def test_catch_up_target_is_yesterday():
    now = dt.datetime(2026, 8, 18, 8, 30, 0)
    assert wr.catch_up_target_day(now) == dt.date(2026, 8, 17)


def test_catch_up_reopens_when_unedited(tmp_path: Path):
    edit = tmp_path / "на_правку"
    accepted = tmp_path / "принятые"
    yesterday = dt.date(2026, 8, 17)
    path = edit / wr.daily_filename(yesterday)
    _write_daily_docx(path, "создали в 16:30, не сохраняли в Word")
    created = dt.datetime(2026, 8, 17, 16, 30, 0)
    wr.write_generated_stamp(path, created)
    os.utime(path, (created.timestamp(), created.timestamp() + 5))
    action, found, source = wr.plan_daily_run(
        yesterday, edit_dir=edit, accepted_dir=accepted, catch_up=True
    )
    assert action == "reopen"
    assert source == "unedited"
    assert found == path


def test_catch_up_skips_when_accepted_exists(tmp_path: Path):
    edit = tmp_path / "на_правку"
    accepted = tmp_path / "принятые"
    yesterday = dt.date(2026, 8, 17)
    _write_daily_docx(accepted / wr.daily_filename(yesterday), "уже правил")
    action, path, source = wr.plan_daily_run(
        yesterday, edit_dir=edit, accepted_dir=accepted, catch_up=True
    )
    assert action == "skip"
    assert source == "accepted"


def test_catch_up_skips_when_saved_later(tmp_path: Path):
    edit = tmp_path / "на_правку"
    accepted = tmp_path / "принятые"
    yesterday = dt.date(2026, 8, 17)
    path = edit / wr.daily_filename(yesterday)
    _write_daily_docx(path, "правили и сохранили")
    created = dt.datetime(2026, 8, 17, 16, 30, 0)
    wr.write_generated_stamp(path, created)
    later = (created + dt.timedelta(minutes=20)).timestamp()
    os.utime(path, (later, later))
    action, found, source = wr.plan_daily_run(
        yesterday, edit_dir=edit, accepted_dir=accepted, catch_up=True
    )
    assert action == "skip"
    assert source == "edited"
    assert found == path


def test_catch_up_creates_when_yesterday_missing(tmp_path: Path):
    edit = tmp_path / "на_правку"
    accepted = tmp_path / "принятые"
    edit.mkdir()
    accepted.mkdir()
    yesterday = dt.date(2026, 8, 17)
    action, path, source = wr.plan_daily_run(
        yesterday, edit_dir=edit, accepted_dir=accepted, catch_up=True
    )
    assert action == "create"
    assert path is None
    assert source == "missing"


def test_daily_1630_regenerates_unedited(tmp_path: Path):
    edit = tmp_path / "на_правку"
    accepted = tmp_path / "принятые"
    day = dt.date(2026, 8, 17)
    path = edit / wr.daily_filename(day)
    _write_daily_docx(path, "черновик до 16:30")
    created = dt.datetime(2026, 8, 17, 10, 0, 0)
    wr.write_generated_stamp(path, created)
    os.utime(path, (created.timestamp(), created.timestamp() + 5))
    action, found, source = wr.plan_daily_run(
        day, edit_dir=edit, accepted_dir=accepted, catch_up=False
    )
    assert action == "create"
    assert source == "unedited"
    assert found == path


def test_not_edited_if_mtime_close_to_generation(tmp_path: Path):
    folder = tmp_path / "на_правку"
    path = folder / wr.daily_filename(dt.date(2026, 8, 17))
    _write_daily_docx(path, "только создали")
    now = dt.datetime.now()
    wr.write_generated_stamp(path, now)
    os.utime(path, (now.timestamp(), now.timestamp() + 10))
    assert wr.was_daily_edited_by_user(path) is False


def test_assemble_does_not_reformat_daily_files():
    import inspect

    src = inspect.getsource(wr.assemble_weekly_from_daily)
    assert "format_weekly_docx" not in src
    assert "copy_docx_body_as_is" in src
    assert "workdays_in_period" in src


def test_scheduler_and_weekend_in_main():
    import inspect

    tasks = inspect.getsource(wr.ensure_windows_daily_tasks)
    assert "16:30" in tasks
    assert "08:30" in tasks
    assert "--daily --catch-up" in tasks or "--catch-up" in tasks
    main_src = inspect.getsource(wr.main)
    assert "is_weekend(today)" in main_src
    assert "Выходной" in main_src
    assert "catch_up_target_day" in main_src
